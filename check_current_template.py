#!/usr/bin/env python3
"""
Script to check the current state of the double template and see what placeholders are present.
"""

from docx import Document
import os

def check_current_template():
    """Check the current double template for placeholders."""
    
    template_path = "src/core/generation/templates/double.docx"
    
    if not os.path.exists(template_path):
        print(f"❌ Template not found: {template_path}")
        return
    
    print(f"Checking template: {template_path}")
    
    try:
        doc = Document(template_path)
        
        print("\n=== TEMPLATE CONTENT ANALYSIS ===")
        
        for table_idx, table in enumerate(doc.tables):
            print(f"\n--- TABLE {table_idx + 1} ---")
            for row_idx, row in enumerate(table.rows):
                print(f"  ROW {row_idx + 1}:")
                for cell_idx, cell in enumerate(row.cells):
                    print(f"    CELL {cell_idx + 1}:")
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        if paragraph.text.strip():
                            print(f"      Paragraph {para_idx + 1}: '{paragraph.text}'")
                            
                            # Check for placeholders
                            if '{{' in paragraph.text and '}}' in paragraph.text:
                                print(f"        ⭐ CONTAINS PLACEHOLDER")
                            
                            # Check for markers
                            if any(marker in paragraph.text for marker in ['_START', '_END']):
                                print(f"        🔍 CONTAINS MARKER")
        
        print("\n=== END TEMPLATE ANALYSIS ===")
        
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    check_current_template() 
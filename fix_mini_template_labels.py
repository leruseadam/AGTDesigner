#!/usr/bin/env python3
"""
Fix mini template to use Label1, Label2, etc. format for proper expansion
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def fix_mini_template_labels():
    """Fix mini template to use proper Label1, Label2, etc. format"""
    
    template_path = 'src/core/generation/templates/mini.docx'
    backup_path = 'src/core/generation/templates/mini.docx.backup_labels'
    
    if os.path.exists(template_path):
        # Create backup
        import shutil
        shutil.copy2(template_path, backup_path)
        print(f"✅ Created backup at: {backup_path}")
    
    # Create a new document
    doc = Document()
    
    # Set page margins to 0.5 inches for mini template
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Create a 4x5 table for mini template (20 labels per page)
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table properties
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
    
    # Set table layout to fixed
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    
    # Set table grid
    grid = OxmlElement('w:tblGrid')
    for i in range(4):
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(1.5 * 1440)))  # 1.5 inches per column
        grid.append(gc)
    tbl.insert(0, grid)
    
    # Set table borders
    borders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'D3D3D3')
        b.set(qn('w:space'), '0')
        borders.append(b)
    tblPr.append(borders)
    
    tbl.insert(0, tblPr)
    
    # Fill each cell with proper placeholders using Label1 format
    label_count = 1
    for row_idx in range(5):
        for col_idx in range(4):
            cell = table.cell(row_idx, col_idx)
            
            # Clear existing content
            cell._tc.clear_content()
            
            # Add placeholders in proper order using Label1, Label2, etc. format
            placeholders = [
                f"{{{{Label{label_count}.ProductStrain}}}}",
                f"{{{{Label{label_count}.ProductBrand}}}}",
                f"{{{{Label{label_count}.VendorInfo}}}}",
                f"{{{{Label{label_count}.Price}}}}",
                f"{{{{Label{label_count}.Ratio_or_THC_CBD}}}}",
                f"{{{{Label{label_count}.Lineage}}}}"
            ]
            
            # Add each placeholder with proper formatting
            for i, placeholder in enumerate(placeholders):
                paragraph = cell.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                run = paragraph.add_run(placeholder)
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.bold = True
                
                # Add spacing between placeholders
                if i < len(placeholders) - 1:
                    paragraph.add_run('\n')
            
            # Set cell dimensions
            cell.width = Inches(1.5)
            cell.height = Inches(1.5)
            
            label_count += 1
    
    # Save the template
    doc.save(template_path)
    print(f"✅ Successfully fixed mini template labels at: {template_path}")
    
    # Verify the template
    verify_doc = Document(template_path)
    if verify_doc.tables:
        table = verify_doc.tables[0]
        print(f"✅ Template verification: {len(table.rows)} rows x {len(table.columns)} columns")
        print(f"✅ Total cells: {len(table.rows) * len(table.columns)}")
        
        # Check for Label1 format
        content = verify_doc.element.body.xml
        if 'Label1.' in content and 'VendorInfo' in content:
            print("✅ Label1 format and VendorInfo placeholder found in fixed template!")
            return True
        else:
            print("❌ Label1 format or VendorInfo placeholder not found in fixed template!")
            return False
    else:
        print("❌ Template verification failed: No tables found")
        return False

def main():
    print("🔧 Fix Mini Template Labels - Use Label1, Label2, etc. Format")
    print("=" * 60)
    
    if fix_mini_template_labels():
        print("\n🎉 Success! Mini template now has:")
        print("   ✅ Proper Label1, Label2, etc. format for expansion")
        print("   ✅ VendorInfo placeholder in each label")
        print("   ✅ Correct 4x5 grid structure (20 labels per page)")
        print("   ✅ Fixed formatting and dimensions")
        print("\n📝 Next steps:")
        print("   1. Test mini template generation")
        print("   2. Template expansion should now work properly")
        print("   3. Vendor information should appear on mini labels")
        print("   4. Push the updated template to your repository")
    else:
        print("\n❌ Failed to fix mini template labels. Please check the errors above.")

if __name__ == "__main__":
    main()

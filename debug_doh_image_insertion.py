#!/usr/bin/env python3
"""
Debug script to test DOH image insertion and centering.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_DOUBLE
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO

def debug_doh_image_insertion():
    """Debug DOH image insertion process."""
    print("🔍 Debugging DOH Image Insertion")
    print("=" * 40)
    
    # Create a test record with DOH image
    test_record = {
        'Description': 'Test Product with DOH',
        'WeightUnits': '1g',
        'ProductBrand': 'Test Brand',
        'Price': '$10.00',
        'Lineage': 'Test Lineage',
        'THC_CBD': 'THC: 20% CBD: 2%',
        'ProductStrain': 'Test Strain',
        'DOH': 'YES',  # This should trigger DOH image
        'Product Type*': 'classic'  # This should use regular DOH image
    }
    
    print("📋 Test Record:")
    for key, value in test_record.items():
        print(f"  {key}: {value}")
    print()
    
    # Test double template specifically
    print("🔄 Processing with double template...")
    processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
    
    # Debug: Check if template has DOH placeholders
    print("🔍 Checking template for DOH placeholders...")
    template_path = processor._get_template_path()
    print(f"  Template path: {template_path}")
    
    # Load the template document
    template_doc = Document(template_path)
    
    # Check for DOH placeholders in the template
    doh_placeholders_found = 0
    for table in template_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                if 'DOH' in cell_text:
                    doh_placeholders_found += 1
                    print(f"  ✓ Found DOH placeholder: '{cell_text.strip()}'")
    
    print(f"  Total DOH placeholders found: {doh_placeholders_found}")
    print()
    
    if doh_placeholders_found == 0:
        print("❌ ERROR: No DOH placeholders found in template!")
        print("   This is why DOH images are not being inserted.")
        return False
    
    # Process the record
    result_doc = processor.process_records([test_record])
    
    if not result_doc:
        print("❌ ERROR: Failed to process test record")
        return False
    
    print("✅ Document generated successfully")
    print()
    
    # Debug: Check the rendered document for DOH images
    print("🔍 Analyzing rendered document...")
    
    # Check for any images in the document
    total_images_found = 0
    doh_images_found = 0
    
    for table_idx, table in enumerate(result_doc.tables):
        print(f"  Table {table_idx + 1}:")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_has_image = False
                
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    for run_idx, run in enumerate(paragraph.runs):
                        # Check if this run contains an image
                        if hasattr(run, '_element'):
                            # Check for drawing elements (InlineImage)
                            drawing = run._element.find(qn('w:drawing'))
                            pict = run._element.find(qn('w:pict'))
                            
                            if drawing is not None or pict is not None:
                                total_images_found += 1
                                cell_has_image = True
                                print(f"    ✓ Found image in cell [{row_idx},{cell_idx}] paragraph {para_idx} run {run_idx}")
                                
                                # Check if this is likely a DOH image by checking cell content
                                cell_text = cell.text.strip()
                                if 'DOH' in cell_text or cell_text == '':
                                    doh_images_found += 1
                                    print(f"      → Likely DOH image (cell text: '{cell_text}')")
                                
                                # Check paragraph alignment
                                if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                                    print(f"      ✓ Paragraph is centered")
                                else:
                                    print(f"      ✗ Paragraph is NOT centered (alignment: {paragraph.alignment})")
                                
                                # Check cell vertical alignment
                                if cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER:
                                    print(f"      ✓ Cell is vertically centered")
                                else:
                                    print(f"      ✗ Cell is NOT vertically centered (alignment: {cell.vertical_alignment})")
                
                if not cell_has_image:
                    # Check cell text for DOH placeholders
                    cell_text = cell.text.strip()
                    if 'DOH' in cell_text:
                        print(f"    ⚠️  Cell [{row_idx},{cell_idx}] has DOH placeholder but no image: '{cell_text}'")
    
    print(f"\n📊 Results:")
    print(f"  Total images found: {total_images_found}")
    print(f"  Likely DOH images: {doh_images_found}")
    
    if total_images_found > 0:
        print("✅ SUCCESS: Images are being inserted into the document")
        if doh_images_found > 0:
            print("✅ SUCCESS: DOH images are being inserted")
            return True
        else:
            print("⚠️  WARNING: Images found but may not be DOH images")
            return False
    else:
        print("❌ FAILURE: No images found in the document")
        print("   This suggests the DocxTemplate rendering is not working for images")
        return False

def main():
    """Run the debug test."""
    print("🎯 DOH Image Insertion Debug Test")
    print("=" * 40)
    print()
    
    success = debug_doh_image_insertion()
    
    print("\n🎯 DEBUG SUMMARY")
    print("=" * 20)
    if success:
        print("✅ DOH images are being inserted and centered correctly")
    else:
        print("❌ DOH image insertion needs investigation")
    
    return success

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1) 
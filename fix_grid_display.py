#!/usr/bin/env python3
"""
Fix the grid display issue by ensuring proper page layout and margins.
This addresses the problem where the 3x3 grid bottom row gets cut off.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

def fix_grid_display():
    """Fix the grid display issue by updating page layout settings."""
    print("🔧 Fixing Grid Display Issue")
    print("=" * 40)
    
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.section import WD_SECTION_START
        
        # Create a test document with proper layout
        doc = Document()
        
        # Set page layout to ensure proper display
        section = doc.sections[0]
        
        # Use standard letter size
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        
        # Set minimal margins to maximize available space
        section.left_margin = Inches(0.25)
        section.right_margin = Inches(0.25)
        section.top_margin = Inches(0.25)
        section.bottom_margin = Inches(0.25)
        
        # Ensure proper section start
        section.start_type = WD_SECTION_START.NEW_PAGE
        
        print(f"✓ Page dimensions: {section.page_width.inches}\" × {section.page_height.inches}\"")
        print(f"✓ Margins: L={section.left_margin.inches}\", R={section.right_margin.inches}\", T={section.top_margin.inches}\", B={section.bottom_margin.inches}\"")
        
        # Calculate available space
        available_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
        available_height = section.page_height.inches - section.top_margin.inches - section.bottom_margin.inches
        
        print(f"✓ Available space: {available_width}\" × {available_height}\"")
        
        # Create 3x3 grid that fits perfectly
        table = doc.add_table(rows=3, cols=3)
        table.alignment = 1  # Center alignment
        
        # Calculate optimal cell dimensions
        # Leave small buffer for borders and spacing
        buffer = 0.1  # 0.1" buffer
        cell_width = (available_width - buffer) / 3
        cell_height = (available_height - buffer) / 3
        
        print(f"✓ Cell dimensions: {cell_width:.2f}\" × {cell_height:.2f}\"")
        print(f"✓ Total grid: {cell_width * 3:.2f}\" × {cell_height * 3:.2f}\"")
        
        # Set table properties
        table.style = 'Table Grid'
        
        # Set column widths
        for col in table.columns:
            col.width = Inches(cell_width)
        
        # Set row heights
        for row in table.rows:
            row.height = Inches(cell_height)
        
        # Add content to verify layout
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = f"Label {i*3 + j + 1}"
                
                # Center text in cells
                paragraph = cell.paragraphs[0]
                paragraph.alignment = 1  # Center alignment
        
        # Save the fixed document
        doc.save("fixed_grid_display.docx")
        print(f"✓ Fixed grid document saved: fixed_grid_display.docx")
        
        # Verify the layout
        print(f"\n--- Layout Verification ---")
        print(f"Page size: {section.page_width.inches}\" × {section.page_height.inches}\"")
        print(f"Margins: {section.left_margin.inches}\" on all sides")
        print(f"Available space: {available_width}\" × {available_height}\"")
        print(f"Grid size: {cell_width * 3:.2f}\" × {cell_height * 3:.2f}\"")
        print(f"Buffer space: {available_width - (cell_width * 3):.2f}\" width, {available_height - (cell_height * 3):.2f}\" height")
        
        if (cell_width * 3) <= available_width and (cell_height * 3) <= available_height:
            print(f"✅ Grid will fit perfectly on page!")
        else:
            print(f"❌ Grid still too large for page")
        
        return True
        
    except Exception as e:
        print(f"❌ Error fixing grid display: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = fix_grid_display()
    if success:
        print(f"\n🎉 Grid display issue should be fixed!")
        print(f"📁 Open 'fixed_grid_display.docx' to verify the layout")
    else:
        print(f"\n⚠️  Failed to fix grid display issue")

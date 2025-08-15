#!/usr/bin/env python3
"""
Comprehensive test to see all placeholders in the vertical template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
import re

def test_vertical_template_complete():
    """Test to see all placeholders and content in the vertical template."""
    
    print("Comprehensive Vertical Template Analysis")
    print("=" * 45)
    
    # Path to the vertical template
    template_path = "src/core/generation/templates/vertical.docx"
    
    if not os.path.exists(template_path):
        print(f"❌ Template not found: {template_path}")
        return
    
    print(f"✅ Template found: {template_path}")
    
    try:
        # Load the template
        doc = Document(template_path)
        
        # Extract all text content with structure
        print(f"\n📄 Template Structure Analysis:")
        print(f"  - Paragraphs: {len(doc.paragraphs)}")
        print(f"  - Tables: {len(doc.tables)}")
        
        # Analyze paragraphs
        if doc.paragraphs:
            print(f"\n📝 Paragraph Content:")
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    print(f"  Para {i+1}: '{para.text.strip()}'")
        
        # Analyze tables
        for table_idx, table in enumerate(doc.tables):
            print(f"\n📊 Table {table_idx + 1} Analysis:")
            print(f"  - Rows: {len(table.rows)}")
            print(f"  - Columns: {len(table.rows[0].cells) if table.rows else 0}")
            
            for row_idx, row in enumerate(table.rows):
                print(f"  Row {row_idx + 1}:")
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"    Cell ({row_idx+1},{col_idx+1}): '{cell_text}'")
        
        # Extract all text content
        all_text = ""
        for paragraph in doc.paragraphs:
            all_text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        all_text += paragraph.text + "\n"
        
        print(f"\n📊 Content Analysis:")
        print(f"  - Total text length: {len(all_text)} characters")
        non_empty_lines = len([line for line in all_text.split('\n') if line.strip()])
        print(f"  - Non-empty lines: {non_empty_lines}")
        
        # Look for placeholder patterns
        placeholder_patterns = [
            (r'\{\{Label\d+\.\w+\}\}', 'Label format'),
            (r'\{\{\{\w+\.\w+\}\}\}', 'Triple brace format'),
            (r'\{\{\w+\}\}', 'Double brace format'),
            (r'DOH', 'DOH text'),
            (r'ProductBrand', 'ProductBrand text'),
            (r'ProductStrain', 'ProductStrain text'),
            (r'Lineage', 'Lineage text'),
            (r'Price', 'Price text'),
            (r'Ratio', 'Ratio text'),
            (r'THC', 'THC text'),
            (r'CBD', 'CBD text'),
        ]
        
        print(f"\n🔍 Placeholder Pattern Analysis:")
        
        for pattern, description in placeholder_patterns:
            matches = re.findall(pattern, all_text)
            if matches:
                print(f"✅ {description}: {len(matches)} matches")
                for match in matches[:3]:  # Show first 3 matches
                    print(f"    - {match}")
                if len(matches) > 3:
                    print(f"    ... and {len(matches) - 3} more")
            else:
                print(f"❌ {description}: No matches found")
        
        # Look for any placeholder-like patterns
        print(f"\n🔍 Any Placeholder-like Patterns:")
        placeholder_like = re.findall(r'\{\{[^}]+\}\}', all_text)
        if placeholder_like:
            print(f"✅ Found {len(placeholder_like)} placeholder-like patterns:")
            for placeholder in placeholder_like:
                print(f"  - {placeholder}")
        else:
            print(f"❌ No placeholder-like patterns found")
            
        # Look for specific field names
        print(f"\n🔍 Specific Field Names in Template:")
        field_names = ['ProductBrand', 'ProductStrain', 'Lineage', 'Price', 'Ratio', 'THC', 'CBD', 'DOH', 'DescAndWeight']
        for field in field_names:
            if field in all_text:
                print(f"✅ {field}: Found in template")
            else:
                print(f"❌ {field}: NOT found in template")
                
    except Exception as e:
        print(f"❌ Error reading template: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_vertical_template_complete()

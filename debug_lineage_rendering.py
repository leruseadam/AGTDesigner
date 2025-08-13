#!/usr/bin/env python3
"""
Debug script to trace exactly what happens to lineage content during template rendering.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_DOUBLE
from docx import Document
from io import BytesIO

def debug_lineage_rendering():
    """Debug exactly what happens to lineage content during template rendering."""
    print("Debugging Lineage Content Rendering")
    print("=" * 40)
    
    # Test record with classic product type
    test_record = {
        'Description': 'Test Flower Product - 3.5g',
        'WeightUnits': '3.5g',
        'ProductBrand': 'Test Flower Brand',
        'Price': '$45.00',
        'Lineage': 'SATIVA',  # Classic lineage content
        'THC_CBD': 'THC: 22% CBD: 1%',
        'ProductStrain': 'Test Strain',
        'ProductType': 'flower',  # Classic type
        'DOH': 'YES'
    }
    
    print(f"Test record:")
    print(f"  Product Type: '{test_record['ProductType']}'")
    print(f"  Lineage: '{test_record['Lineage']}'")
    
    # Test double template
    processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
    
    # Check template path
    template_path = processor._get_template_path()
    print(f"\nTemplate path: {template_path}")
    
    # Check original template for LINEAGE placeholders
    print("\nChecking original template for LINEAGE placeholders...")
    with open(template_path, 'rb') as f:
        original_doc = Document(f)
        print(f"Original template: {len(original_doc.tables)} tables")
        for i, table in enumerate(original_doc.tables):
            print(f"  Table {i+1}: {len(table.rows)} rows, {len(table.columns)} columns")
            # Check for LINEAGE placeholder
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        if 'LINEAGE' in paragraph.text:
                            print(f"    Found LINEAGE placeholder in Table{i+1}, Row{row_idx+1}, Cell{col_idx+1}, Para{para_idx+1}")
                            print(f"      Content: '{paragraph.text}'")
    
    # Check expanded template for LINEAGE placeholders
    print("\nChecking expanded template for LINEAGE placeholders...")
    expanded_buffer = processor._expand_template_if_needed()
    expanded_doc = Document(expanded_buffer)
    print(f"Expanded template: {len(expanded_doc.tables)} tables")
    for i, table in enumerate(expanded_doc.tables):
        print(f"  Table {i+1}: {len(table.rows)} rows, {len(table.columns)} columns")
        # Check for LINEAGE placeholder
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    if 'LINEAGE' in paragraph.text:
                        print(f"    Found LINEAGE placeholder in Table{i+1}, Row{row_idx+1}, Cell{col_idx+1}, Para{para_idx+1}")
                        print(f"      Content: '{paragraph.text}'")
    
    # Check label context
    print("\nBuilding label context...")
    label_context = processor._build_label_context(test_record, None)
    print(f"Label context lineage: '{label_context.get('Lineage', 'NOT_FOUND')}'")
    print(f"Label context keys: {list(label_context.keys())}")
    
    # Check if lineage is wrapped with markers
    lineage_value = label_context.get('Lineage', '')
    if 'LINEAGE_START' in lineage_value and 'LINEAGE_END' in lineage_value:
        print("✓ Lineage is properly wrapped with markers")
        print(f"  Raw value: '{lineage_value}'")
    else:
        print("❌ Lineage is NOT wrapped with markers")
        print(f"  Raw value: '{lineage_value}'")
    
    # Test DocxTemplate rendering directly
    print("\nTesting DocxTemplate rendering directly...")
    from docxtpl import DocxTemplate
    
    # Create DocxTemplate with expanded template
    expanded_buffer.seek(0)
    docx_template = DocxTemplate(expanded_buffer)
    
    # Create context
    context = {'Label1': label_context}
    print(f"Context being passed to DocxTemplate:")
    for key, value in context.items():
        print(f"  {key}: {value}")
    
    # Render template
    try:
        docx_template.render(context)
        print("✅ DocxTemplate.render() completed successfully")
        
        # Save to buffer and check content
        buffer = BytesIO()
        docx_template.save(buffer)
        buffer.seek(0)
        rendered_doc = Document(buffer)
        
        print(f"\nAfter DocxTemplate.render(): {len(rendered_doc.tables)} tables")
        lineage_found = False
        for i, table in enumerate(rendered_doc.tables):
            print(f"Table {i+1}: {len(table.rows)} rows, {len(table.columns)} columns")
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        if 'SATIVA' in paragraph.text:
                            print(f"Found lineage content in Table{i+1}, Row{row_idx+1}, Cell{col_idx+1}, Para{para_idx+1}")
                            print(f"  Content: '{paragraph.text[:100]}...'")
                            lineage_found = True
                            break
                        elif 'LINEAGE_START' in paragraph.text or 'LINEAGE_END' in paragraph.text:
                            print(f"Found lineage markers in Table{i+1}, Row{row_idx+1}, Cell{col_idx+1}, Para{para_idx+1}")
                            print(f"  Content: '{paragraph.text[:100]}...'")
                        elif 'HYBRID' in paragraph.text:
                            print(f"Found 'HYBRID' content in Table{i+1}, Row{row_idx+1}, Cell{col_idx+1}, Para{para_idx+1}")
                            print(f"  Content: '{paragraph.text[:100]}...'")
                            lineage_found = True
                            break
                    if lineage_found:
                        break
                if lineage_found:
                    break
            if lineage_found:
                break
        
        if not lineage_found:
            print("❌ No lineage content found after DocxTemplate.render()")
            print("This suggests the issue is in the DocxTemplate rendering itself")
            return False
        
    except Exception as e:
        print(f"❌ DocxTemplate.render() failed: {e}")
        return False
    
    return True

if __name__ == "__main__":
    success = debug_lineage_rendering()
    if success:
        print("\n✅ Lineage content found after DocxTemplate rendering")
    else:
        print("\n❌ Lineage content not found - DocxTemplate rendering issue")
    sys.exit(0 if success else 1) 
#!/usr/bin/env python3
"""
Create a proper horizontal template with all necessary placeholders.
"""

from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def create_horizontal_template():
    """Create a proper horizontal template with all placeholders."""
    
    # Create a new document
    doc = Document()
    
    # Set page orientation to landscape
    section = doc.sections[0]
    section.orientation = 1  # 1 = landscape, 0 = portrait
    
    # Set page margins for 3x3 grid
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Create a 3x3 table
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table properties
    tblPr = table._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
    
    # Set table layout to fixed
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    
    # Set column widths (3.4 inches each for landscape)
    grid = OxmlElement('w:tblGrid')
    for _ in range(3):
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(3.4 * 1440)))  # Convert inches to twips
        grid.append(gc)
    
    # Set row heights (2.4 inches each for landscape)
    for row in table.rows:
        row.height = Inches(2.4)
        row.height_rule = 1  # Exactly
    
    # Add borders
    borders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'D3D3D3')
        borders.append(b)
    tblPr.append(borders)
    
    # Insert table properties
    table._element.insert(0, tblPr)
    table._element.insert(1, grid)
    
    # Fill each cell with proper placeholders
    label_num = 1
    for row in table.rows:
        for cell in row.cells:
            # Clear any existing content
            cell.text = ''
            
            # Add all necessary placeholders as separate paragraphs
            cell.add_paragraph(f'{{{{Label{label_num}.Lineage}}}} {{{{Label{label_num}.ProductVendor}}}}')
            cell.add_paragraph(f'{{{{Label{label_num}.ProductStrain}}}}')
            cell.add_paragraph(f'{{{{Label{label_num}.DescAndWeight}}}}')
            cell.add_paragraph(f'{{{{Label{label_num}.Price}}}}')
            cell.add_paragraph(f'{{{{Label{label_num}.DOH}}}}')
            cell.add_paragraph(f'{{{{Label{label_num}.Ratio_or_THC_CBD}}}}')
            
            label_num += 1
    
    # Save the template
    template_path = 'src/core/generation/templates/horizontal.docx'
    doc.save(template_path)
    
    print(f"✅ Created proper horizontal template at: {template_path}")
    print("✅ Template includes all necessary placeholders:")
    print("   - Lineage + ProductVendor")
    print("   - ProductStrain")
    print("   - DescAndWeight")
    print("   - Price")
    print("   - DOH")
    print("   - Ratio_or_THC_CBD")
    print("✅ Template is configured for 3x3 grid with landscape orientation")
    
    return template_path

if __name__ == "__main__":
    create_horizontal_template()

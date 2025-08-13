#!/usr/bin/env python3
"""
Simple script to check what template type is currently being used.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

def check_current_template_type():
    """Check the current template type being used."""
    print("Checking Current Template Type")
    print("=" * 40)
    
    try:
        # Check if there are any template settings files
        import json
        from pathlib import Path
        
        # Check for template settings in session (if available)
        print("1. Checking for template settings...")
        
        # Look for any saved template settings
        settings_files = [
            Path.home() / "Downloads" / "template_settings.json",
            Path.cwd() / "template_settings.json",
            Path.cwd() / "session_template_settings.json"
        ]
        
        for settings_file in settings_files:
            if settings_file.exists():
                try:
                    with open(settings_file, 'r') as f:
                        settings = json.load(f)
                    print(f"   Found settings file: {settings_file}")
                    print(f"   Template type: {settings.get('templateType', 'Not specified')}")
                    print(f"   All settings: {json.dumps(settings, indent=2)}")
                except Exception as e:
                    print(f"   Error reading {settings_file}: {e}")
            else:
                print(f"   Settings file not found: {settings_file}")
        
        # Check what templates are available
        print("\n2. Checking available templates...")
        template_dir = Path("src/core/generation/templates")
        if template_dir.exists():
            templates = list(template_dir.glob("*.docx"))
            print(f"   Available templates: {[t.name for t in templates]}")
            
            # Check template sizes
            for template in templates:
                try:
                    from docx import Document
                    doc = Document(template)
                    if doc.tables:
                        table = doc.tables[0]
                        print(f"   {template.name}: {len(table.rows)} rows x {len(table.columns)} columns")
                    else:
                        print(f"   {template.name}: No tables found")
                except Exception as e:
                    print(f"   {template.name}: Error reading - {e}")
        else:
            print(f"   Template directory not found: {template_dir}")
        
        # Check constants for expected grid layouts
        print("\n3. Checking expected grid layouts...")
        try:
            from src.core.constants import GRID_LAYOUTS
            print(f"   Expected grid layouts: {GRID_LAYOUTS}")
        except Exception as e:
            print(f"   Error reading constants: {e}")
            
    except Exception as e:
        print(f"Error in main check: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    check_current_template_type()

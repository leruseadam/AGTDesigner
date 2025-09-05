#!/usr/bin/env python3
"""
Debug script to check exactly what's happening to the mini template.
This will help identify where the corruption is occurring.
"""

import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), 'src'))

from core.generation.template_processor import TemplateProcessor
from docx import Document
import tempfile

def debug_mini_template_corruption():
    """Debug exactly what's happening to the mini template."""
    print("=== Debugging Mini Template Corruption ===")
    
    try:
        # 1. Check the original mini template
        print("\n1. Checking original mini.docx template...")
        original_path = 'src/core/generation/templates/mini.docx'
        if os.path.exists(original_path):
            original_doc = Document(original_path)
            print(f"   ✅ Original template loaded: {len(original_doc.tables)} tables")
            
            if original_doc.tables:
                table = original_doc.tables[0]
                print(f"   📊 Original table: {len(table.rows)} rows x {len(table.columns)} columns")
                
                # Check first cell content
                first_cell = table.rows[0].cells[0]
                print(f"   📝 First cell content: {first_cell.text[:100]}...")
                
                # Check for placeholders
                import re
                text = original_doc.element.body.xml
                placeholders = re.findall(r'Label(\d+)\.', text)
                print(f"   🏷️  Placeholders found: {sorted(set([int(x) for x in placeholders]))}")
        else:
            print(f"   ❌ Original template not found at: {original_path}")
            return False
        
        # 2. Check what happens during template processor initialization
        print("\n2. Checking template processor initialization...")
        processor = TemplateProcessor('mini', 'Arial')
        
        print(f"   🔧 Template type: {processor.template_type}")
        print(f"   📁 Template path: {processor._template_path}")
        
        # 3. Check the expanded template buffer
        print("\n3. Checking expanded template buffer...")
        if hasattr(processor, '_expanded_template_buffer') and processor._expanded_template_buffer:
            print(f"   📦 Buffer exists: {type(processor._expanded_template_buffer)}")
            
            # Check if buffer has seek method
            if hasattr(processor._expanded_template_buffer, 'seek'):
                processor._expanded_template_buffer.seek(0)
                expanded_doc = Document(processor._expanded_template_buffer)
                print(f"   📊 Expanded template: {len(expanded_doc.tables)} tables")
                
                if expanded_doc.tables:
                    table = expanded_doc.tables[0]
                    print(f"   📊 Expanded table: {len(table.rows)} rows x {len(table.columns)} columns")
                    
                    # Check first cell content
                    first_cell = table.rows[0].cells[0]
                    print(f"   📝 Expanded first cell: {first_cell.text[:100]}...")
                    
                    # Check for placeholders
                    text = expanded_doc.element.body.xml
                    placeholders = re.findall(r'Label(\d+)\.', text)
                    print(f"   🏷️  Expanded placeholders: {sorted(set([int(x) for x in placeholders]))}")
            else:
                print("   ❌ Buffer doesn't have seek method")
        else:
            print("   ❌ No expanded template buffer")
        
        # 4. Test the mini template processing method directly
        print("\n4. Testing mini template processing method directly...")
        
        # Create test context
        context = {}
        for i in range(1, 3):
            context[f'Label{i}'] = {
                'ProductStrain': f'Test Strain {i}',
                'ProductBrand': f'Test Brand {i}',
                'VendorInfo': f'Test Vendor {i}',
                'Price': f'${i * 10}.00',
                'Ratio_or_THC_CBD': f'THC: {20 + i}% CBD: {i * 0.1}%',
                'Lineage': f'Test Lineage {i}'
            }
        
        # Add empty contexts for remaining labels
        for i in range(3, 21):
            context[f'Label{i}'] = {}
        
        print(f"   📋 Created context with {len(context)} labels")
        
        # Process the mini template
        result_doc = processor._expand_mini_template_preserve_design(None, context)
        
        # Save result
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            result_doc.save(tmp_file.name)
            print(f"   💾 Saved result to: {tmp_file.name}")
        
        # Verify the result
        print("\n5. Analyzing result...")
        if result_doc.tables:
            table = result_doc.tables[0]
            print(f"   📊 Result table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check first few cells
            for i in range(min(5, len(table.rows) * len(table.columns))):
                row = i // len(table.columns)
                col = i % len(table.columns)
                cell = table.cell(row, col)
                cell_text = cell.text.strip()
                
                print(f"\n   📝 Cell {i+1}:")
                print(f"      Text: {cell_text[:80]}...")
                
                if i < 2:  # First two cells should have data
                    if "Test Strain" in cell_text and "Test Brand" in cell_text:
                        print("      ✅ Placeholders replaced with data")
                    else:
                        print("      ❌ Placeholders NOT replaced!")
                else:  # Other cells should have placeholders
                    if "{{Label" in cell_text:
                        print("      ⚠️  STILL HAS PLACEHOLDERS!")
                    else:
                        print("      ❌ Unexpected content!")
        else:
            print("   ❌ No tables found in result!")
            
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False
    
    print("\n=== Mini Template Corruption Debug Completed ===")
    return True

if __name__ == "__main__":
    debug_mini_template_corruption()

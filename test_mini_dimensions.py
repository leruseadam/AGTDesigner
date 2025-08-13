#!/usr/bin/env python3
"""
Test script to check mini template dimensions through the full processing pipeline.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from io import BytesIO

def test_mini_template_dimensions():
    """Test mini template dimensions through the full processing pipeline."""
    print("Testing mini template dimensions...")
    
    # Create a template processor for mini templates
    processor = TemplateProcessor('mini', {}, 1.0)
    print(f"Template processor created for type: {processor.template_type}")
    
    # Test the expansion method directly
    print("\n1. Testing direct expansion method...")
    expanded = processor._expand_template_to_4x5_fixed_scaled()
    expanded.seek(0)
    doc = Document(expanded)
    table = doc.tables[0]
    print(f"   Expanded template: {len(table.rows)} rows, {len(table.rows[0].cells) if table.rows else 0} columns")
    print(f"   First cell width: {table.rows[0].cells[0].width/1440:.2f} inches")
    print(f"   First row height: {table.rows[0].height/1440:.2f} inches")
    
    # Test processing a simple record
    print("\n2. Testing full processing pipeline...")
    test_record = {
        'ProductName': 'Test Product',
        'ProductBrand': 'Test Brand',
        'ProductStrain': 'Test Strain',
        'Price': '$10.00',
        'Lineage': 'HYBRID',
        'Ratio_or_THC_CBD': 'THC: 20%',
        'DOH': 'DOH.png'
    }
    
    # Process the record
    result = processor.process_records([test_record])
    if result:
        print(f"   Processing successful, result type: {type(result)}")
        if hasattr(result, 'tables') and result.tables:
            table = result.tables[0]
            print(f"   Final template: {len(table.rows)} rows, {len(table.rows[0].cells) if table.rows else 0} columns")
            print(f"   First cell width: {table.rows[0].cells[0].width/1440:.2f} inches")
            print(f"   First row height: {table.rows[0].height/1440:.2f} inches")
        else:
            print("   No tables found in result")
    else:
        print("   Processing failed")
    
    print("\nTest completed.")

if __name__ == "__main__":
    test_mini_template_dimensions()

#!/usr/bin/env python3
"""
Test script to verify that preroll descriptions get Arial Bold formatting in horizontal templates
"""

import os
import sys
from pathlib import Path

# Add the src directory to the path
sys.path.insert(0, str(Path(__file__).parent / "src"))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from io import BytesIO

def test_horizontal_preroll_arial_bold():
    """Test if preroll descriptions get Arial Bold formatting in horizontal templates."""
    
    # Create a test record for preroll
    test_record = {
        'ProductName': 'Test Pre-Roll',
        'Product Type*': 'pre-roll',
        'Description': 'Blueberry Infused Pre-Roll',
        'WeightUnits': '0.5g x 2 Pack',
        'ProductBrand': 'Test Brand',
        'Price': '$15.00',
        'Lineage': 'HYBRID',
        'DOH': 'YES',
        'Ratio': '0.5g x 2 Pack',
        'JointRatio': '0.5g x 2 Pack',
        'THC': '15.5%',
        'CBD': '0.1%'
    }
    
    print("Testing preroll description Arial Bold formatting in horizontal template...")
    print(f"Test record: {test_record}")
    
    try:
        # Test with horizontal template
        print("\n1. Testing with horizontal template...")
        tp = TemplateProcessor(template_type='horizontal', font_scheme='Arial')
        
        # Process the record
        result = tp.process_records([test_record])
        
        if result:
            print("✅ Template processing successful")
            
            # The result should be a Document object
            if hasattr(result, 'tables'):
                doc = result
            else:
                print("❌ Unexpected result type:", type(result))
                return
            
            # Check all runs in the document for Arial Bold
            arial_bold_found = False
            non_arial_bold_found = False
            total_runs = 0
            
            print("\nChecking all text runs in the document:")
            for table_idx, table in enumerate(doc.tables):
                print(f"\nTable {table_idx + 1}:")
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        print(f"  Cell ({row_idx + 1}, {col_idx + 1}):")
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            if paragraph.text.strip():
                                print(f"    Paragraph {para_idx + 1}: '{paragraph.text}'")
                                for run_idx, run in enumerate(paragraph.runs):
                                    if run.text.strip():  # Only check non-empty runs
                                        total_runs += 1
                                        font_name = run.font.name
                                        is_bold = run.font.bold
                                        
                                        print(f"      Run {run_idx + 1}: '{run.text}' - Font: {font_name}, Bold: {is_bold}")
                                        
                                        if font_name == 'Arial' and is_bold:
                                            arial_bold_found = True
                                        elif font_name != 'Arial' or not is_bold:
                                            non_arial_bold_found = True
            
            print(f"\nSummary:")
            print(f"  Total runs found: {total_runs}")
            print(f"  Arial Bold found: {arial_bold_found}")
            print(f"  Non-Arial Bold found: {non_arial_bold_found}")
            
            if arial_bold_found and not non_arial_bold_found:
                print("✅ All text is Arial Bold - formatting working correctly")
            elif arial_bold_found and non_arial_bold_found:
                print("⚠️  Mixed formatting - some text is Arial Bold, some is not")
            else:
                print("❌ No Arial Bold formatting found")
                
            # Check if the preroll description is actually in the document
            print("\nSearching for preroll description content:")
            found_description = False
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if 'Blueberry' in paragraph.text or 'Pre-Roll' in paragraph.text:
                                print(f"  Found description: '{paragraph.text}'")
                                found_description = True
                                # Check formatting of this specific paragraph
                                for run in paragraph.runs:
                                    if run.text.strip():
                                        print(f"    Run: '{run.text}' - Font: {run.font.name}, Bold: {run.font.bold}")
            
            if not found_description:
                print("  ❌ Preroll description not found in document")
            else:
                print("  ✅ Preroll description found in document")
                
        else:
            print("❌ Template processing failed")
            
    except Exception as e:
        print(f"❌ Error during testing: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_horizontal_preroll_arial_bold()

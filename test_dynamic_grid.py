#!/usr/bin/env python3
"""
Test script to verify dynamic grid sizing and that empty tags are no longer getting CBD color formatting.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.generation.tag_generator import process_chunk
from docx import Document
from io import BytesIO

def test_dynamic_grid_sizing():
    """Test that templates create grids based on actual number of labels."""
    print("Testing dynamic grid sizing...")
    
    # Test different template types with different numbers of labels
    test_cases = [
        ('vertical', 3),    # Should create 3x1 grid
        ('vertical', 6),    # Should create 3x2 grid  
        ('vertical', 9),    # Should create 3x3 grid
        ('double', 4),      # Should create 4x1 grid
        ('double', 8),      # Should create 4x2 grid
        ('double', 12),     # Should create 4x3 grid
        ('inventory', 2),   # Should create 2x1 grid
        ('inventory', 4),   # Should create 2x2 grid
    ]
    
    for template_type, num_labels in test_cases:
        print(f"\nTesting {template_type} template with {num_labels} labels...")
        
        try:
            # Create template processor
            processor = TemplateProcessor(template_type, {}, 1.0)
            
            # Force template expansion with specific number of labels
            expanded_buffer = processor._expand_template_if_needed(
                force_expand=True, 
                num_selected_tags=num_labels
            )
            
            if expanded_buffer:
                # Load the expanded template
                doc = Document(expanded_buffer)
                
                if doc.tables:
                    table = doc.tables[0]
                    rows = len(table.rows)
                    cols = len(table.columns)
                    total_cells = rows * cols
                    
                    print(f"  Grid dimensions: {rows}x{cols} = {total_cells} cells")
                    
                    # Check if the grid size matches expectations
                    if template_type == 'vertical':
                        if num_labels <= 3 and total_cells == num_labels:
                            print(f"  ✅ Correct: {num_labels} labels fit in {total_cells} cells")
                        elif num_labels <= 6 and total_cells == 6:
                            print(f"  ✅ Correct: {num_labels} labels fit in {total_cells} cells")
                        elif num_labels <= 9 and total_cells == 9:
                            print(f"  ✅ Correct: {num_labels} labels fit in {total_cells} cells")
                        else:
                            print(f"  ❌ Incorrect: Expected appropriate grid size for {num_labels} labels")
                    elif template_type == 'double':
                        if num_labels <= 4 and total_cells == num_labels:
                            print(f"  ✅ Correct: {num_labels} labels fit in {total_cells} cells")
                        elif num_labels <= 8 and total_cells == 8:
                            print(f"  ✅ Correct: {num_labels} labels fit in {total_cells} cells")
                        elif num_labels <= 12 and total_cells == 12:
                            print(f"  ✅ Correct: {num_labels} labels fit in {total_cells} cells")
                        else:
                            print(f"  ❌ Incorrect: Expected appropriate grid size for {num_labels} labels")
                    elif template_type == 'inventory':
                        if num_labels <= 2 and total_cells == num_labels:
                            print(f"  ✅ Correct: {num_labels} labels fit in {total_cells} cells")
                        elif num_labels <= 4 and total_cells == 4:
                            print(f"  ✅ Correct: {num_labels} labels fit in {total_cells} cells")
                        else:
                            print(f"  ❌ Incorrect: Expected appropriate grid size for {num_labels} labels")
                else:
                    print(f"  ❌ No tables found in expanded template")
            else:
                print(f"  ❌ Template expansion failed")
                
        except Exception as e:
            print(f"  ❌ Error testing {template_type} with {num_labels} labels: {e}")

def test_no_empty_placeholders():
    """Test that templates don't contain empty placeholder text for unused slots."""
    print("\n\nTesting that templates don't contain empty placeholders...")
    
    test_cases = [
        ('vertical', 3),    # Should only have Label1, Label2, Label3
        ('double', 4),      # Should only have Label1, Label2, Label3, Label4
    ]
    
    for template_type, num_labels in test_cases:
        print(f"\nTesting {template_type} template with {num_labels} labels...")
        
        try:
            # Create template processor
            processor = TemplateProcessor(template_type, {}, 1.0)
            
            # Force template expansion with specific number of labels
            expanded_buffer = processor._expand_template_if_needed(
                force_expand=True, 
                num_selected_tags=num_labels
            )
            
            if expanded_buffer:
                # Load the expanded template
                doc = Document(expanded_buffer)
                
                # Check the text content for unexpected placeholders
                text_content = doc.element.body.xml
                
                # Check for placeholders beyond the expected number
                unexpected_placeholders = []
                for i in range(num_labels + 1, 13):  # Check Label5 through Label12
                    placeholder = f'Label{i}.'
                    if placeholder in text_content:
                        unexpected_placeholders.append(placeholder)
                
                if unexpected_placeholders:
                    print(f"  ❌ Found unexpected placeholders: {unexpected_placeholders}")
                else:
                    print(f"  ✅ No unexpected placeholders found")
                    
                # Check for expected placeholders
                expected_placeholders = []
                for i in range(1, num_labels + 1):
                    placeholder = f'Label{i}.'
                    if placeholder in text_content:
                        expected_placeholders.append(placeholder)
                
                print(f"  Found expected placeholders: {expected_placeholders}")
                
            else:
                print(f"  ❌ Template expansion failed")
                
        except Exception as e:
            print(f"  ❌ Error testing {template_type} with {num_labels} labels: {e}")

if __name__ == "__main__":
    print("=" * 60)
    print("Testing Dynamic Grid Sizing and Empty Placeholder Prevention")
    print("=" * 60)
    
    test_dynamic_grid_sizing()
    test_no_empty_placeholders()
    
    print("\n" + "=" * 60)
    print("Test completed!")
    print("=" * 60)

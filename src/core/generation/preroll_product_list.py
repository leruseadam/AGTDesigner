"""
Preroll Product List Generator

This module handles the generation of product list documents for preroll templates.
It creates a separate document listing all preroll items grouped by category.
"""

import logging
import os
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import session
from flask_caching import Cache
from src.core.constants import PREROLL_ALLOWED_BRANDS


def generate_preroll_product_list(records: List[Dict[str, Any]], cache: Cache) -> Optional[Document]:
    """
    Generate a separate product list document for preroll templates.
    
    Args:
        records: List of product records (used to extract group IDs if needed)
        cache: Flask cache instance for retrieving stored group data
        
    Returns:
        A separate Document containing the product list (or None if no groups found)
    """
    try:
        # Create a new document for the preroll list
        list_doc = Document()
        
        # Add title
        title = list_doc.add_heading('Preroll Product Lists', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Get session ID and group keys to retrieve stored preroll groups
        session_id = session.get('preroll_session_id', session.get('session_id', 'default'))
        # CRITICAL FIX: Use preroll_group_keys (full unique keys) instead of preroll_group_ids (base category only)
        # preroll_group_ids only contains base categories like "single", "2-pack" which causes 56 groups to collapse to ~5
        # preroll_group_keys contains full keys like "single|acme|abc123hash" which preserves all unique groups
        group_keys = session.get('preroll_group_keys', [])

        # Fallback to old group_ids if group_keys not available (backward compatibility)
        if not group_keys:
            group_keys = session.get('preroll_group_ids', [])
            logging.warning("PREROLL LIST: Using fallback preroll_group_ids instead of preroll_group_keys")

        # Retrieve all preroll groups from cache using stored group keys
        preroll_groups = {}
        for group_key in group_keys:
            # Try session-independent latest key first (preferred)
            group_items = cache.get(f"preroll_group_latest_{group_key}")
            group_info = cache.get(f"preroll_group_info_latest_{group_key}")

            # Fallback to session-specific key
            if not group_items or not group_info:
                group_items = cache.get(f"preroll_group_{session_id}_{group_key}")
                group_info = cache.get(f"preroll_group_info_{session_id}_{group_key}")

            if group_items and group_info:
                preroll_groups[group_key] = {
                    'items': group_items,
                    'info': group_info
                }
                logging.info(f"PREROLL LIST: Loaded group '{group_info.get('display_name', group_key)}' with {len(group_items)} items from cache")
            else:
                logging.warning(f"PREROLL LIST: Group '{group_key}' not found in cache (items: {group_items is not None}, info: {group_info is not None})")
        
        if not preroll_groups:
            logging.warning("PREROLL LIST: No preroll groups found in cache")
            # Try fallback: check records for group keys (full unique keys)
            for record in records:
                # Prefer full group_key, fall back to group_id
                group_key = record.get('_group_key') or record.get('_group_id')
                if group_key and group_key not in preroll_groups:
                    # Try latest key first
                    group_items = cache.get(f"preroll_group_latest_{group_key}")
                    group_info = cache.get(f"preroll_group_info_latest_{group_key}")
                    # Fallback to session-specific key
                    if not group_items or not group_info:
                        group_items = cache.get(f"preroll_group_{session_id}_{group_key}")
                        group_info = cache.get(f"preroll_group_info_{session_id}_{group_key}")
                    if group_items and group_info:
                        preroll_groups[group_key] = {
                            'items': group_items,
                            'info': group_info
                        }
        
        # Only create list document if we have groups
        if preroll_groups:
            # Create a section for each preroll group, separated by brand
            for group_id, group_data in preroll_groups.items():
                group_info = group_data['info']
                group_items = group_data['items']
                display_name = group_info.get('display_name', f'Group {group_id}')
                
                # Filter items by allowed brands if configured
                # CRITICAL FIX: Only filter if PREROLL_ALLOWED_BRANDS is not None and not empty
                filtered_items = group_items
                if PREROLL_ALLOWED_BRANDS is not None and len(PREROLL_ALLOWED_BRANDS) > 0:
                    allowed_brands_lower = {brand.lower().strip() for brand in PREROLL_ALLOWED_BRANDS if brand and str(brand).strip()}
                    if allowed_brands_lower:  # Only filter if we have valid brands after normalization
                        original_item_count = len(filtered_items)
                        filtered_items = [
                            item for item in group_items
                            if str(item.get('brand', '')).strip().lower() in allowed_brands_lower
                        ]
                        if len(filtered_items) < original_item_count:
                            logging.info(f"PREROLL LIST: Filtered {original_item_count} items to {len(filtered_items)} items for group '{display_name}' based on allowed brands")
                    else:
                        logging.info(f"PREROLL LIST: PREROLL_ALLOWED_BRANDS is set but contains no valid brands, skipping filter for group '{display_name}'")
                else:
                    logging.debug(f"PREROLL LIST: PREROLL_ALLOWED_BRANDS is empty or None, including all items for group '{display_name}'")
                
                # Group items by brand within this category
                items_by_brand = {}
                for item in filtered_items:
                    brand = str(item.get('brand', '')).strip() or 'Unknown Brand'
                    if brand not in items_by_brand:
                        items_by_brand[brand] = []
                    items_by_brand[brand].append(item)
                
                # Create a separate table for each brand
                for brand, brand_items in sorted(items_by_brand.items()):
                    # Add group heading with brand
                    heading_text = f"{display_name} - {brand}"
                    heading = list_doc.add_heading(heading_text, level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Create a table for the items (6 columns: add DOH status)
                    table = list_doc.add_table(rows=1, cols=6)
                    table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    header_cells = table.rows[0].cells
                    header_cells[0].text = 'Product Name'
                    header_cells[1].text = 'Brand'
                    header_cells[2].text = 'Price'
                    header_cells[3].text = 'Weight'
                    header_cells[4].text = 'Lineage'

                    # Add DOH logo to the last header cell instead of text
                    doh_header_cell = header_cells[5]
                    doh_logo_path = os.path.join(os.path.dirname(__file__), 'templates', 'DOH.png')

                    # Clear any existing text in the DOH cell
                    doh_header_cell.text = ''

                    # Add the DOH logo image if it exists
                    if os.path.exists(doh_logo_path):
                        try:
                            # Get the paragraph in the cell and add the image
                            doh_paragraph = doh_header_cell.paragraphs[0]
                            doh_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = doh_paragraph.add_run()
                            run.add_picture(doh_logo_path, width=Inches(0.3))  # Tiny logo - 0.3 inches wide
                            logging.info("PREROLL LIST: Added tiny DOH logo to header")
                        except Exception as img_error:
                            # Fallback to text if image fails
                            doh_header_cell.text = 'DOH'
                            logging.warning(f"PREROLL LIST: Failed to add DOH logo, using text: {img_error}")
                    else:
                        # Fallback to text if image doesn't exist
                        doh_header_cell.text = 'DOH'
                        logging.warning(f"PREROLL LIST: DOH logo not found at {doh_logo_path}, using text")

                    # Make header row bold (except DOH cell which has image)
                    for i, cell in enumerate(header_cells):
                        if i != 5:  # Skip DOH cell since it has an image
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.size = Pt(11)
                    
                    # Add items to table
                    for item in brand_items:
                        row_cells = table.add_row().cells
                        row_cells[0].text = item.get('product_name', '')
                        row_cells[1].text = item.get('brand', '')
                        row_cells[2].text = item.get('price', '')
                        row_cells[3].text = item.get('weight', '')
                        row_cells[4].text = item.get('lineage', '')
                        row_cells[5].text = item.get('doh', '')
                    
                    # Add spacing after each brand table
                    list_doc.add_paragraph('')
            
            # Log summary of items in list
            total_items_in_list = sum(len(group_data['items']) for group_data in preroll_groups.values())
            logging.info(f"PREROLL LIST: Generated separate product list document with {len(preroll_groups)} groups containing {total_items_in_list} total items")
            
            # Log which groups and their item counts
            for group_id, group_data in preroll_groups.items():
                group_name = group_data['info'].get('display_name', group_id)
                item_count = len(group_data['items'])
                logging.info(f"PREROLL LIST GROUP: '{group_name}' ({group_id}) - {item_count} items")
            
            return list_doc
        else:
            logging.warning("PREROLL LIST: No preroll groups found, skipping list creation")
            return None
        
    except Exception as e:
        logging.warning(f"PREROLL LIST: Failed to create preroll list document: {e}")
        return None

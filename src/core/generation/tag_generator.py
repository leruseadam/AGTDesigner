import io
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from io import BytesIO
import logging
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
import re
from docxtpl import DocxTemplate, InlineImage
from copy import deepcopy
import cProfile
from itertools import groupby
import pandas as pd

from src.core.generation.docx_formatting import (
    fix_table_row_heights,
    safe_fix_paragraph_spacing,
    apply_conditional_formatting,
    set_cell_background,
    clear_cell_margins,
    clear_table_cell_padding,
    enforce_fixed_cell_dimensions,
)
from src.core.generation.context_builders import (
    build_context,
)
from src.core.formatting.markers import (
    wrap_with_marker,
    FIELD_MARKERS
)
from src.core.utils.resource_utils import resource_path
from src.core.constants import (
    FONT_SCHEME_HORIZONTAL,
    FONT_SCHEME_VERTICAL,
    FONT_SCHEME_MINI,
    LINEAGE_COLOR_MAP
)

# In-memory cache of template bytes to avoid repeated file I/O and speed up
# template loading. Keys are absolute template paths.
_template_bytes_cache = {}

def _get_template_bytes(path: str) -> bytes:
    b = _template_bytes_cache.get(path)
    if b is None:
        with open(path, 'rb') as fh:
            b = fh.read()
            _template_bytes_cache[path] = b
    return b

# Performance optimization: disable debug logging in production
DEBUG_ENABLED = False

logger = logging.getLogger(__name__)

def _clean_lineage_text(value) -> str:
    """Normalize lineage/brand text to prevent hidden-space artifacts in DOCX output."""
    if value is None:
        return ""
    text = str(value)
    # Remove zero-width characters and normalize uncommon Unicode spaces to normal spaces.
    text = re.sub(r"[\u200B-\u200D\u2060\uFEFF]", "", text)
    text = re.sub(r"[\u00A0\u2000-\u200A\u202F\u205F\u3000]", " ", text)
    # Collapse repeated whitespace and trim.
    return re.sub(r"\s+", " ", text).strip()

def _find_most_likely_ounce_weight(product_name, product_type):
    """
    Find the most common ounce weight for similar nonclassic products.
    This helps maintain consistency with actual product packaging rather than mathematical conversion.
    """
    # Common ounce weights for nonclassic products based on typical packaging
    common_oz_weights = {
        # Edibles - common sizes
        'edible (liquid)': ['2.5oz', '3.53oz', '1.7oz'],
        'edible (solid)': ['1oz', '2.5oz', '3.5oz'],
        'gummy': ['1oz', '2oz', '3.5oz'],
        'chocolate': ['1oz', '2oz', '3.5oz'],
        'cookie': ['1oz', '2oz'],
        'brownie': ['1oz', '2oz'],
        'candy': ['1oz', '2oz', '3.5oz'],
        
        # Tinctures and oils
        'tincture': ['1oz', '2oz', '4oz'],
        'drops': ['1oz', '2oz'],
        'liquid': ['1oz', '2oz', '4oz'],
        
        # Topicals
        'topical': ['1oz', '2oz', '4oz'],
        'cream': ['1oz', '2oz', '4oz'],
        'lotion': ['1oz', '2oz', '4oz'],
        'salve': ['1oz', '2oz'],
        'balm': ['1oz', '2oz'],
        
        # Capsules
        'capsule': ['1oz', '2oz'],
        
        # Beverages
        'beverage': ['12oz', '16oz', '20oz'],
        'drink': ['12oz', '16oz', '20oz'],
        'soda': ['12oz', '16oz'],
        'juice': ['12oz', '16oz'],
        
        # Default fallback
        'default': ['1oz', '2oz', '3.5oz']
    }
    
    # Get the most common weight for this product type
    product_type_lower = product_type.lower().strip()
    
    # Try exact match first
    if product_type_lower in common_oz_weights:
        return common_oz_weights[product_type_lower][0]  # Return the first (most common) weight
    
    # Try partial matches for product types that might have variations
    for key, weights in common_oz_weights.items():
        if key != 'default' and key in product_type_lower:
            return weights[0]
    
    # Special handling for Moonshot products (they seem to be 2.5oz or 3.53oz based on the image)
    if 'moonshot' in product_name.lower():
        return '2.5oz'  # Most common Moonshot size
    
    # Default fallback
    return common_oz_weights['default'][0]  # Return '1oz' as default

PLACEHOLDER_MARKERS = {
    "Description": ("DESC_START", "DESC_END"),
    "WeightUnits": ("WEIGHTUNITS_START", "WEIGHTUNITS_END"),
    "ProductBrand": ("PRODUCTBRAND_START", "PRODUCTBRAND_END"),
    "ProductBrand_Center": ("PRODUCTBRAND_CENTER_START", "PRODUCTBRAND_CENTER_END"),
    "Price": ("PRICE_START", "PRICE_END"),
    "Lineage": ("LINEAGE_START", "LINEAGE_END"),
    "DOH": ("{{Label1.DOH}}", ""),
    "Ratio_or_THC_CBD": ("RATIO_START", "RATIO_END"),
    "THC_CBD": ("THC_CBD_START", "THC_CBD_END"),
    "ProductName": ("PRODUCTNAME_START", "PRODUCTNAME_END"),
    "ProductStrain": ("PRODUCTSTRAIN_START", "PRODUCTSTRAIN_END"),
    "ProductType": ("PRODUCTTYPE_START", "PRODUCTTYPE_END")
}

# Import colors from docx_formatting to avoid duplication
from .docx_formatting import COLORS as LINEAGE_COLORS

def get_template_path(template_type):
    """Return the absolute path for a given template type."""
    # Map template types to filenames
    template_files = {
        'horizontal': 'horizontal.docx',
        'vertical': 'vertical.docx',
        'mini': 'mini.docx',
        'double': 'double.docx',
        'inventory': 'inventory.docx',
        'preroll': 'preroll.docx'
    }

    # Validate template type
    if template_type not in template_files:
        raise ValueError(f"Invalid template type: {template_type}")

    # Get template directory path
    base_dir = Path(__file__).parent
    template_dir = base_dir / "templates"
    
    # Debug logging
    print(f"Debug - Template lookup: type={template_type}, file={template_files[template_type]}")
    print(f"Debug - Paths: base_dir={base_dir}, template_dir={template_dir}")

    # Ensure template directory exists
    if not template_dir.exists():
        raise ValueError(f"Template directory not found: {template_dir}")

    # Build full path (expected filename)
    expected_filename = template_files[template_type]
    template_path = template_dir / expected_filename
    print(f"Debug - Final template path: {template_path}")

    # Verify template exists, with robust fallbacks for casing/hidden prefixes on some hosts
    if not template_path.exists():
        # Fallback: case-insensitive match ONLY for non-hidden files (ignore . and ~$ temp files)
        expected_lower = expected_filename.lower()
        fallback = None
        for p in template_dir.iterdir():
            if not p.is_file():
                continue
            name = p.name
            # Ignore hidden/office temp files
            if name.startswith('.') or name.startswith('~$'):
                continue
            if name.lower() == expected_lower:
                fallback = p
                break
        if fallback and fallback.exists():
            print(f"Warning - Using fallback template path due to case-only mismatch: {fallback}")
            return str(fallback)
        raise ValueError(f"Template file not found: {template_path}")

    return str(template_path)

def chunk_records(records, chunk_size=60):
    """Split the list of records into larger chunks to reduce docx processing overhead."""
    if chunk_size < 1:
        chunk_size = 60
    # No deduplication here for maximum speed
    chunks = []
    for i in range(0, len(records), chunk_size):
        chunks.append(records[i:i + chunk_size])
    logger.info(f"Split {len(records)} records into {len(chunks)} chunks of size {chunk_size}")
    return chunks

def flatten_tags(records):
    """Extract Description values from records for tag generation."""
    flat_tags = []
    for record in records:
        description = record.get("Description", "")
        if description and isinstance(description, str):
            flat_tags.append(description.strip())
    return flat_tags

def create_dynamic_mini_template(template_path, num_products, scale_factor=1.0):
    """Create a dynamic mini template with only as many cells as products."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from io import BytesIO
    from copy import deepcopy

    # Calculate grid dimensions based on number of products
    # Always use 4 columns, calculate rows needed
    num_cols = 4
    num_rows = (num_products + num_cols - 1) // num_cols  # Ceiling division
    
    col_width_twips = str(int(1.5 * 1440))   # 1.5 inches per column
    row_height_pts = Pt(1.5 * 72)           # 1.5 inches per row
    cut_line_twips = int(0.001 * 1440)

    # Load template bytes from cache to avoid repeated disk reads
    tpl_bytes = _get_template_bytes(template_path)
    doc = Document(BytesIO(tpl_bytes))
    if not doc.tables:
        raise RuntimeError("Template must contain at least one table.")
    old = doc.tables[0]
    src_tc = deepcopy(old.cell(0,0)._tc)
    old._element.getparent().remove(old._element)

    # Remove empty paragraphs
    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    # CRITICAL FIX: Create table with exact number of rows needed, no more
    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = tbl._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
    
    # Set table properties
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D3D3D3')
    tblPr.insert(0, shd)
    
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    tbl._element.insert(0, tblPr)
    
    # Set grid columns
    grid = OxmlElement('w:tblGrid')
    for _ in range(num_cols):
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), col_width_twips)
        grid.append(gc)
    tbl._element.insert(0, grid)
    
    # Set row heights
    for row in tbl.rows:
        row.height = row_height_pts
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    # Set borders (include outer light grey border to match original mini template framing)
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f"w:{side}")
        b.set(qn('w:val'), "single")
        b.set(qn('w:sz'), "4")
        b.set(qn('w:color'), "D3D3D3")
        b.set(qn('w:space'), "0")
        borders.append(b)
    tblPr.append(borders)
    
    # CRITICAL FIX: Only populate cells for products we have
    cnt = 1
    for r in range(num_rows):
        for c in range(num_cols):
            if cnt <= num_products:  # Only populate cells for products we have
                cell = tbl.cell(r,c)
                cell._tc.clear_content()
                tc = deepcopy(src_tc)
                for t in tc.iter(qn('w:t')):
                    if t.text and 'Label1' in t.text:
                        t.text = t.text.replace('Label1', f'Label{cnt}')
                for el in tc.xpath('./*'):
                    cell._tc.append(deepcopy(el))
            else:
                # CRITICAL FIX: For empty cells, just clear them but don't remove them
                # This prevents the table structure from breaking
                cell = tbl.cell(r,c)
                cell._tc.clear_content()
            cnt += 1
    
    # Add spacing
    tblPr2 = tbl._element.find(qn('w:tblPr'))
    spacing = OxmlElement('w:tblCellSpacing')
    spacing.set(qn('w:w'), str(cut_line_twips))
    spacing.set(qn('w:type'), 'dxa')
    tblPr2.append(spacing)
    
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def create_dynamic_3x3_template(template_path, num_products, scale_factor=1.0):
    """Create a dynamic 3x3 template with only the cells needed for the number of products."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from io import BytesIO
    from copy import deepcopy
    
    # Calculate grid dimensions based on number of products
    # Use 3 columns, calculate rows needed
    num_cols = 3
    # CRITICAL FIX: For more than 9 products, create multiple pages with 9 labels each
    if num_products <= 9:
        num_rows = (num_products + num_cols - 1) // num_cols  # Ceiling division
    else:
        # For more than 9 products, create pages with 9 labels each
        num_rows = 3  # Always 3 rows per page
    
    col_width_twips = str(int(3.4 * 1440))  # 3.4 inches per column for horizontal template
    row_height_pts = Pt(2.4 * 72)  # 2.4 inches per row for horizontal template
    cut_line_twips = int(0.001 * 1440)
    
    tpl_bytes = _get_template_bytes(template_path)
    doc = Document(BytesIO(tpl_bytes))
    if not doc.tables:
        raise RuntimeError("Template must contain at least one table.")
    
    old = doc.tables[0]
    src_tc = deepcopy(old.cell(0,0)._tc)
    old._element.getparent().remove(old._element)
    
    # Only remove empty paragraphs, preserve content paragraphs
    doc_paragraphs = list(doc.paragraphs)
    for paragraph in doc_paragraphs:
        if not paragraph.text.strip():
            paragraph._element.getparent().remove(paragraph._element)
    
    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Copy the original table properties and styling from the source template
    old_tbl = Document(BytesIO(tpl_bytes)).tables[0]
    for prop in old_tbl._element.xpath('./w:tblPr/*'):
        tbl._element.append(deepcopy(prop))
    
    # CRITICAL FIX: Handle multiple pages for more than 9 products
    if num_products <= 9:
        # Single page - populate only the cells we need
        cnt = 1
        for r in range(num_rows):
            for c in range(num_cols):
                if cnt <= num_products:
                    cell = tbl.cell(r, c)
                    cell._tc.clear_content()
                    
                    # Prepare a copy of the template cell, then replace placeholders before appending
                    tc = deepcopy(src_tc)
                    for t in tc.iter(qn('w:t')):
                        if t.text and 'Label1' in t.text:
                            t.text = t.text.replace('Label1', f'Label{cnt}')
                    for el in tc.xpath('./*'):
                        cell._tc.append(deepcopy(el))
                else:
                    # CRITICAL FIX: For empty cells, completely remove them to prevent blank placeholders
                    cell = tbl.cell(r, c)
                    cell._tc.clear_content()
                    # Don't add any paragraphs - keep the cell completely empty
                cnt += 1
    else:
        # Multiple pages - create pages with 9 labels each
        products_per_page = 9
        total_pages = (num_products + products_per_page - 1) // products_per_page
        
        # Remove the single table we created
        tbl._element.getparent().remove(tbl._element)
        
        # Create multiple pages
        for page_num in range(total_pages):
            # Add page break if not the first page
            if page_num > 0:
                doc.add_page_break()
            
            # Create a new table for this page
            tbl = doc.add_table(rows=3, cols=3)  # Always 3x3 for horizontal/vertical
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Copy table properties from the original template
            old_tpl_bytes = _get_template_bytes(template_path)
            old_tbl = Document(BytesIO(old_tpl_bytes)).tables[0]
            for prop in old_tbl._element.xpath('./w:tblPr/*'):
                tbl._element.append(deepcopy(prop))
            
            # Set table properties
            tblPr = tbl._element.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl._element.insert(0, tblPr)
            
            # Set table width
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            tblW.set(qn('w:w'), str(int(3.4 * 3 * 1440)))  # Total width for 3 columns
            tblW.set(qn('w:type'), 'dxa')
            
            # Set table layout
            tblLayout = tblPr.find(qn('w:tblLayout'))
            if tblLayout is None:
                tblLayout = OxmlElement('w:tblLayout')
                tblPr.append(tblLayout)
            tblLayout.set(qn('w:type'), 'fixed')
            
            # Set cell spacing
            tblCellSpacing = tblPr.find(qn('w:tblCellSpacing'))
            if tblCellSpacing is None:
                tblCellSpacing = OxmlElement('w:tblCellSpacing')
                tblPr.append(tblCellSpacing)
            tblCellSpacing.set(qn('w:w'), str(cut_line_twips))
            tblCellSpacing.set(qn('w:type'), 'dxa')
            
            # CRITICAL FIX: Ensure table grid exists for valid XML structure
            tblGrid = tbl._element.find(qn('w:tblGrid'))
            if tblGrid is None:
                tblGrid = OxmlElement('w:tblGrid')
                tbl._element.insert(1, tblGrid)
                for _ in range(3):  # 3 columns
                    gc = OxmlElement('w:gridCol')
                    gc.set(qn('w:w'), col_width_twips)
                    tblGrid.append(gc)
                print(f"Created missing tblGrid for page {page_num + 1}")
            
            # Populate cells for this page
            start_product = page_num * products_per_page + 1
            end_product = min(start_product + products_per_page - 1, num_products)
            
            cnt = start_product
            for r in range(3):  # 3 rows
                for c in range(3):  # 3 columns
                    if cnt <= end_product:
                        cell = tbl.cell(r, c)
                        cell._tc.clear_content()
                        
                        # Prepare a copy of the template cell, then replace placeholders before appending
                        tc = deepcopy(src_tc)
                        for t in tc.iter(qn('w:t')):
                            if t.text and 'Label1' in t.text:
                                t.text = t.text.replace('Label1', f'Label{cnt}')
                        for el in tc.xpath('./*'):
                            cell._tc.append(deepcopy(el))
                    else:
                        # CRITICAL FIX: For empty cells, completely remove them to prevent blank placeholders
                        cell = tbl.cell(r, c)
                        cell._tc.clear_content()
                        # Don't add any paragraphs - keep the cell completely empty
                    cnt += 1
            
            # Set row heights for this page
            for row in tbl.rows:
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                row.height = row_height_pts
                
                for cell in row.cells:
                    # Set cell properties
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    
                    # Set cell width
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is None:
                        tcW = OxmlElement('w:tcW')
                        tcPr.append(tcW)
                    tcW.set(qn('w:w'), col_width_twips)
                    tcW.set(qn('w:type'), 'dxa')
                    
                    # Set cell margins
                    tcMar = tcPr.find(qn('w:tcMar'))
                    if tcMar is None:
                        tcMar = OxmlElement('w:tcMar')
                        tcPr.append(tcMar)
                    
                    for margin in ['top', 'left', 'bottom', 'right']:
                        margin_el = tcMar.find(qn(f'w:{margin}'))
                        if margin_el is None:
                            margin_el = OxmlElement(f'w:{margin}')
                            tcMar.append(margin_el)
                        margin_el.set(qn('w:w'), '0')
                        margin_el.set(qn('w:type'), 'dxa')
    
    # Add spacing (only for single page case)
    if num_products <= 9:
        tblPr2 = tbl._element.find(qn('w:tblPr'))
        spacing = OxmlElement('w:tblCellSpacing')
        spacing.set(qn('w:w'), str(cut_line_twips))
        spacing.set(qn('w:type'), 'dxa')
        tblPr2.append(spacing)
        
        # Set row heights for single page
        for row in tbl.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = row_height_pts
            
            for cell in row.cells:
                # Set cell properties
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Set cell width
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), col_width_twips)
                tcW.set(qn('w:type'), 'dxa')
                
                # Set cell margins
                tcMar = tcPr.find(qn('w:tcMar'))
                if tcMar is None:
                    tcMar = OxmlElement('w:tcMar')
                    tcPr.append(tcMar)
                
                for margin in ['top', 'left', 'bottom', 'right']:
                    margin_el = tcMar.find(qn(f'w:{margin}'))
                    if margin_el is None:
                        margin_el = OxmlElement(f'w:{margin}')
                        tcMar.append(margin_el)
                    margin_el.set(qn('w:w'), '0')
                    margin_el.set(qn('w:type'), 'dxa')
    
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def create_dynamic_double_template(template_path, num_products, scale_factor=1.0):
    """Create a dynamic double template with only as many cells as products."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from io import BytesIO
    from copy import deepcopy

    # Calculate grid dimensions based on number of products
    # Always use 4 columns, calculate rows needed
    num_cols = 4
    # CRITICAL FIX: For more than 12 products, create multiple pages with 12 labels each
    if num_products <= 12:
        num_rows = (num_products + num_cols - 1) // num_cols  # Ceiling division
    else:
        # For more than 12 products, create pages with 12 labels each
        num_rows = 3  # Always 3 rows per page
    
    col_width_twips = str(int(1.75 * 1440))  # 1.75 inches per column for double template (original width)
    row_height_pts = Pt(2.5 * 72)  # 2.5 inches per row for double template
    cut_line_twips = int(0.001 * 1440)

    tpl_bytes = _get_template_bytes(template_path)
    doc = Document(BytesIO(tpl_bytes))
    if not doc.tables:
        raise RuntimeError("Template must contain at least one table.")
    old = doc.tables[0]
    src_tc = deepcopy(old.cell(0,0)._tc)
    old._element.getparent().remove(old._element)

    # Remove empty paragraphs
    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    # Create table with dynamic dimensions
    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = tbl._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
    
    # Set table properties
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D3D3D3')
    tblPr.insert(0, shd)
    
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    tbl._element.insert(0, tblPr)
    
    # Set grid columns
    grid = OxmlElement('w:tblGrid')
    for _ in range(num_cols):
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), col_width_twips)
        grid.append(gc)
    tbl._element.insert(0, grid)
    
    # Set row heights
    for row in tbl.rows:
        row.height = row_height_pts
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    # Set borders
    borders = OxmlElement('w:tblBorders')
    for side in ('insideH','insideV'):
        b = OxmlElement(f"w:{side}")
        b.set(qn('w:val'), "single")
        b.set(qn('w:sz'), "4")
        b.set(qn('w:color'), "D3D3D3")
        b.set(qn('w:space'), "0")
        borders.append(b)
    tblPr.append(borders)
    
    # CRITICAL FIX: Handle multiple pages for more than 12 products
    if num_products <= 12:
        # Single page - populate only the cells we need
        cnt = 1
        for r in range(num_rows):
            for c in range(num_cols):
                if cnt <= num_products:  # Only populate cells for products we have
                    cell = tbl.cell(r,c)
                    cell._tc.clear_content()
                    tc = deepcopy(src_tc)
                    for t in tc.iter(qn('w:t')):
                        if t.text and 'Label1' in t.text:
                            t.text = t.text.replace('Label1', f'Label{cnt}')
                    for el in tc.xpath('./*'):
                        cell._tc.append(deepcopy(el))
                else:
                    # CRITICAL FIX: For empty cells, completely remove them to prevent blank placeholders
                    cell = tbl.cell(r, c)
                    cell._tc.clear_content()
                    # Don't add any paragraphs - keep the cell completely empty
                cnt += 1
    else:
        # Multiple pages - create pages with 12 labels each
        products_per_page = 12
        total_pages = (num_products + products_per_page - 1) // products_per_page
        
        # Remove the single table we created
        tbl._element.getparent().remove(tbl._element)
        
        # Create multiple pages
        for page_num in range(total_pages):
            # Add page break if not the first page
            if page_num > 0:
                doc.add_page_break()
            
            # Create a new table for this page
            tbl = doc.add_table(rows=3, cols=4)  # Always 3x4 for double template
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Copy table properties from the original template
            old_tpl_bytes = _get_template_bytes(template_path)
            old_tbl = Document(BytesIO(old_tpl_bytes)).tables[0]
            for prop in old_tbl._element.xpath('./w:tblPr/*'):
                tbl._element.append(deepcopy(prop))
            
            # Set table properties
            tblPr = tbl._element.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl._element.insert(0, tblPr)
            
            # Set table width
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            tblW.set(qn('w:w'), str(int(1.75 * 4 * 1440)))  # Total width for 4 columns
            tblW.set(qn('w:type'), 'dxa')
            
            # Set table layout
            tblLayout = tblPr.find(qn('w:tblLayout'))
            if tblLayout is None:
                tblLayout = OxmlElement('w:tblLayout')
                tblPr.append(tblLayout)
            tblLayout.set(qn('w:type'), 'fixed')
            
            # Set cell spacing
            tblCellSpacing = tblPr.find(qn('w:tblCellSpacing'))
            if tblCellSpacing is None:
                tblCellSpacing = OxmlElement('w:tblCellSpacing')
                tblPr.append(tblCellSpacing)
            tblCellSpacing.set(qn('w:w'), str(cut_line_twips))
            tblCellSpacing.set(qn('w:type'), 'dxa')
            
            # Set table grid
            tblGrid = tbl._element.find(qn('w:tblGrid'))
            if tblGrid is not None:
                tblGrid.getparent().remove(tblGrid)
            tblGrid = OxmlElement('w:tblGrid')
            tbl._element.insert(1, tblGrid)
            
            for _ in range(4):  # 4 columns
                gc = OxmlElement('w:gridCol')
                gc.set(qn('w:w'), col_width_twips)
                tblGrid.append(gc)
            
            # Populate cells for this page
            start_product = page_num * products_per_page + 1
            end_product = min(start_product + products_per_page - 1, num_products)
            
            cnt = start_product
            for r in range(3):  # 3 rows
                for c in range(4):  # 4 columns
                    if cnt <= end_product:
                        cell = tbl.cell(r, c)
                        cell._tc.clear_content()
                        
                        # Copy all elements from source cell
                        for el in src_tc.xpath('./*'):
                            cell._tc.append(deepcopy(el))
                        
                        # Update Label1 to Label{cnt} in the copied content
                        for t in cell.xpath('.//w:t'):
                            if t.text and 'Label1' in t.text:
                                t.text = t.text.replace('Label1', f'Label{cnt}')
                        # REMOVED: Duplicate append that was causing content duplication
                    else:
                        # CRITICAL FIX: For empty cells, completely remove them to prevent blank placeholders
                        cell = tbl.cell(r, c)
                        cell._tc.clear_content()
                        # Don't add any paragraphs - keep the cell completely empty
                    cnt += 1
            
            # Set row heights for this page
            for row in tbl.rows:
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                row.height = row_height_pts
                
                for cell in row.cells:
                    # Set cell properties
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    
                    # Set cell width
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is None:
                        tcW = OxmlElement('w:tcW')
                        tcPr.append(tcW)
                    tcW.set(qn('w:w'), col_width_twips)
                    tcW.set(qn('w:type'), 'dxa')
                    
                    # Set cell margins
                    tcMar = tcPr.find(qn('w:tcMar'))
                    if tcMar is None:
                        tcMar = OxmlElement('w:tcMar')
                        tcPr.append(tcMar)
                    
                    for margin in ['top', 'left', 'bottom', 'right']:
                        margin_el = tcMar.find(qn(f'w:{margin}'))
                        if margin_el is None:
                            margin_el = OxmlElement(f'w:{margin}')
                            tcMar.append(margin_el)
                        margin_el.set(qn('w:w'), '0')
                        margin_el.set(qn('w:type'), 'dxa')
    
    # Add spacing
    tblPr2 = tbl._element.find(qn('w:tblPr'))
    spacing = OxmlElement('w:tblCellSpacing')
    spacing.set(qn('w:w'), str(cut_line_twips))
    spacing.set(qn('w:type'), 'dxa')
    tblPr2.append(spacing)
    
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def expand_template_to_4x5_fixed_scaled(template_path, scale_factor=1.0):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from io import BytesIO
    from copy import deepcopy

    num_cols, num_rows = 4, 5
    col_width_twips = str(int(1.5 * 1440))   # 1.5 inches per column for equal width
    row_height_pts  = Pt(1.5 * 72)           # 1.5 inches per row for equal height
    cut_line_twips  = int(0.001 * 1440)

    tpl_bytes = _get_template_bytes(template_path)
    doc = Document(BytesIO(tpl_bytes))
    if not doc.tables:
        raise RuntimeError("Template must contain at least one table.")
    old = doc.tables[0]
    src_tc = deepcopy(old.cell(0,0)._tc)
    old._element.getparent().remove(old._element)

    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = tbl._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D3D3D3')
    tblPr.insert(0, shd)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    tbl._element.insert(0, tblPr)
    grid = OxmlElement('w:tblGrid')
    for _ in range(num_cols):
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), col_width_twips)
        grid.append(gc)
    tbl._element.insert(0, grid)
    for row in tbl.rows:
        row.height = row_height_pts
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    borders = OxmlElement('w:tblBorders')
    for side in ('insideH','insideV'):
        b = OxmlElement(f"w:{side}")
        b.set(qn('w:val'), "single"); b.set(qn('w:sz'), "4")
        b.set(qn('w:color'), "D3D3D3"); b.set(qn('w:space'), "0")
        borders.append(b)
    tblPr.append(borders)
    cnt = 1
    for r in range(num_rows):
        for c in range(num_cols):
            cell = tbl.cell(r,c)
            cell._tc.clear_content()
            tc = deepcopy(src_tc)
            for t in tc.iter(qn('w:t')):
                if t.text and 'Label1' in t.text:
                    t.text = t.text.replace('Label1', f'Label{cnt}')
            for el in tc.xpath('./*'):
                cell._tc.append(deepcopy(el))
            cnt += 1
    from docx.oxml.shared import OxmlElement as OE
    tblPr2 = tbl._element.find(qn('w:tblPr'))
    spacing = OxmlElement('w:tblCellSpacing'); spacing.set(qn('w:w'), str(cut_line_twips)); spacing.set(qn('w:type'), 'dxa')
    tblPr2.append(spacing)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def process_chunk(args):
    """Process a chunk of records to generate labels."""
    chunk, base_template, font_scheme, orientation, scale_factor = args
    logger = logging.getLogger(__name__)
    
    # Skip verbose lineage debugging to improve performance
    # CRITICAL FIX: Disable ALL dynamic templates to prevent XML corruption
    # Use standard template expansion with post-processing cleanup instead
    if orientation == "mini":
        local_template_buffer = base_template
        num_labels = 20  # Use standard 4x5 grid
        logger.debug(f"🔧 MINI TEMPLATE EXPANSION: Using standard template expansion")
    elif orientation == "preroll":
        local_template_buffer = base_template
        num_labels = 20  # Use standard 4x5 grid (same as mini)
        logger.debug(f"🔧 PREROLL TEMPLATE EXPANSION: Using standard 4x5 expansion (same as mini)")
    elif orientation == "double":
        local_template_buffer = base_template
        num_labels = 12  # Use standard 4x3 grid
        logger.debug(f"🔧 DOUBLE TEMPLATE EXPANSION: Using standard template expansion")
    else:
        # CRITICAL FIX: Use dynamic label count based on chunk size
        # Allow templates to expand to accommodate all products in the chunk
        local_template_buffer = base_template
        num_labels = len(chunk)  # Use actual chunk size instead of hardcoded 9
        logger.debug(f"🔧 {orientation.upper()} TEMPLATE EXPANSION: Using dynamic expansion for {num_labels} labels")
    # Prefer loading from cached bytes when possible
    try:
        if isinstance(local_template_buffer, (bytes, bytearray)):
            local_buf = BytesIO(local_template_buffer)
            tpl = DocxTemplate(local_buf)
        else:
            tpl = DocxTemplate(local_template_buffer)
    except Exception:
        # Fallback to direct load
        tpl = DocxTemplate(local_template_buffer)
    context = {}
    image_width = Mm(10) if orientation == "preroll" else (Mm(8) if orientation == "mini" else Mm(9 if orientation == 'vertical' else 12))
    doh_image_path = resource_path(os.path.join("templates", "DOH.png"))
    if DEBUG_ENABLED:
        logger.debug(f"DOH image path: {doh_image_path}")
    
    # CRITICAL FIX: For preroll, use ALL records in chunk (template expands to fit)
    # For other templates, limit to num_labels to avoid empty slots
    if orientation == "preroll":
        actual_num_labels = len(chunk)  # Use all records - template will expand to 4x5 grid per chunk
        # Preroll uses all records - no logging needed
    else:
        actual_num_labels = min(len(chunk), num_labels)
    
    # OPTIMIZATION: Pre-load all lineage and strain data in batch to avoid N+1 queries
    # This reduces 200+ queries for 100 products to just 2-3 queries total
    product_lineage_cache = {}
    strain_info_cache = {}
    try:
        from app import get_product_database, get_current_store_name
        store_name = get_current_store_name()
        product_db = get_product_database(store_name)
        if product_db:
            # Collect all product names and strain names from chunk
            product_names = []
            strain_names = set()
            for row in chunk:
                product_name = row.get('Product Name*', '') or row.get('ProductName', '')
                if product_name:
                    product_names.append(product_name)
                strain = row.get('Product Strain', '')
                if strain:
                    strain_names.add(strain)
            
            # Batch query for product lineage (sovereign > strain > products.Lineage)
            if product_names:
                try:
                    conn = product_db._get_connection()
                    cur = conn.cursor()
                    placeholders = ','.join(['?'] * len(product_names))
                    batch_lineage_query = f'''
                        SELECT p."Product Name*",
                               COALESCE(p.sovereign_lineage, s1.sovereign_lineage, s2.sovereign_lineage, s3.sovereign_lineage,
                                        s1.canonical_lineage, s2.canonical_lineage, s3.canonical_lineage, p."Lineage") AS lineage
                        FROM products p
                        LEFT JOIN strains s1 ON p.strain_id = s1.id
                        LEFT JOIN strains s2 ON p.normalized_product_strain = s2.normalized_name
                        LEFT JOIN strains s3 ON p.normalized_product_strain IS NULL AND LOWER(TRIM(p."Product Strain")) = LOWER(TRIM(s3.strain_name))
                        WHERE p."Product Name*" IN ({placeholders})
                        ORDER BY p.id DESC
                    '''
                    cur.execute(batch_lineage_query, product_names)
                    seen = set()
                    for row_result in cur.fetchall():
                        pname, lineage = row_result
                        if not pname or pname in seen:
                            continue
                        if lineage and str(lineage).strip() not in ['', 'None', 'nan', 'SOVEREIGN']:
                            product_lineage_cache[pname] = str(lineage).strip().upper()
                            seen.add(pname)
                except Exception as batch_err:
                    logger.warning(f"Batch product lineage query failed: {batch_err}")
            
            # Batch query for strain info (sovereign_lineage > display_lineage > canonical_lineage)
            if strain_names:
                try:
                    conn = product_db._get_connection()
                    cur = conn.cursor()
                    strain_list = list(strain_names)
                    placeholders = ','.join(['?'] * len(strain_list))
                    batch_strain_query = f'''
                        SELECT strain_name, sovereign_lineage, canonical_lineage
                        FROM strains
                        WHERE strain_name IN ({placeholders})
                    '''
                    cur.execute(batch_strain_query, strain_list)
                    for row_result in cur.fetchall():
                        sname = row_result[0]
                        preferred = row_result[1] or row_result[2]
                        if preferred and str(preferred).strip().upper() != 'SOVEREIGN':
                            strain_info_cache[sname] = {
                                'sovereign_lineage': row_result[1],
                                'canonical_lineage': row_result[2]
                            }
                except Exception as batch_strain_err:
                    logger.warning(f"Batch strain info query failed: {batch_strain_err}")
    except Exception as e:
        logger.warning(f"Failed to pre-load lineage/strain data: {e}")
    
    for i in range(actual_num_labels):
        label_data = {}
        if i < len(chunk):
            row = chunk[i]
            
            # CRITICAL FIX: Log missing product data for debugging
            product_name = str(row.get("ProductName", "")).strip()
            if not product_name or product_name.lower() in ['', 'nan', 'none', 'null']:
                logger.warning(f"⚠️  SKIPPING EMPTY PRODUCT at index {i}: No product name found")
                continue
            
            # Use only canonical DOH field for image decisions
            doh_raw = str(row.get("DOH", "")).strip()
            doh_compliant = ''
            doh_lower = ''
            
            doh_value = (doh_raw).upper()
            if logger.isEnabledFor(logging.DEBUG):
                logger.debug(f"DOH: '{product_name}' -> '{doh_value}'")
            product_type = str(row.get("Product Type*", "")).strip().lower()
            if DEBUG_ENABLED:
                logger.debug(f"Product type: {product_type}")
            
            # Handle different DOH values: YES (legacy), DOH, THC, CBD, or NO/NONE
            # Explicitly handle NO/NONE/FALSE/empty values first
            # Also check the raw value before upper() conversion for "No"/"no"
            raw_doh = (str(row.get("DOH", "")).strip())
            if doh_value in ["NO", "NONE", "FALSE", ""] or raw_doh in ["No", "no", "NO", "NONE"]:
                label_data["DOH"] = ""
                logger.info(f"✅ DOH DECISION: Product '{product_name}' - NO DOH image (value: '{doh_value}' or '{raw_doh}')")
            elif doh_value in ["YES", "DOH", "THC", "CBD"]:
                # Map DOH value to appropriate image
                if doh_value == "CBD" or (doh_value == "YES" and product_type.startswith('high cbd')):
                    # Use High CBD image
                    high_cbd_image_path = resource_path(os.path.join("templates", "HighCBD.png"))
                    if DEBUG_ENABLED:
                        logger.debug(f"Using HighCBD image: {high_cbd_image_path}")
                    label_data["DOH"] = InlineImage(tpl, high_cbd_image_path, width=image_width)
                elif doh_value == "THC":
                    # Use High THC image
                    high_thc_image_path = resource_path(os.path.join("templates", "HighTHC.png"))
                    if DEBUG_ENABLED:
                        logger.debug(f"Using HighTHC image: {high_thc_image_path}")
                    label_data["DOH"] = InlineImage(tpl, high_thc_image_path, width=image_width)
                else:
                    # Use regular DOH image
                    doh_image_path = resource_path(os.path.join("templates", "DOH.png"))
                    if DEBUG_ENABLED:
                        logger.debug(f"Using DOH image: {doh_image_path}")
                    label_data["DOH"] = InlineImage(tpl, doh_image_path, width=image_width)
                    if logger.isEnabledFor(logging.DEBUG):
                        logger.debug(f"DOH: '{product_name}' - Added image")
                if DEBUG_ENABLED:
                    logger.debug(f"Created DOH image with width: {image_width}")
            else:
                label_data["DOH"] = ""
                if DEBUG_ENABLED:
                    logger.debug(f"Skipping DOH image - value is '{doh_value}' (not DOH/THC/CBD)")
                
            # --- Wrap all fields with markers ---
            # Updated price mapping to use correct field names
            # CRITICAL FIX: Handle both Price and Price* fields from database/Excel
            price_val = ''
            if row.get('Price*') and str(row.get('Price*', '')).strip():
                price_val = str(row.get('Price*', '')).strip()
            elif row.get('Price') and str(row.get('Price', '')).strip():
                price_val = str(row.get('Price', '')).strip()
            
            # CRITICAL FIX: Log missing prices for debugging
            if not price_val or not price_val.strip():
                logger.warning(f"⚠️  MISSING PRICE for product '{product_name}' at index {i}")
                logger.debug(f"   Row data: Price*={row.get('Price*')}, Price={row.get('Price')}")
            
            # CRITICAL FIX: Format price with dollar sign before wrapping
            if price_val and price_val.strip():
                try:
                    # Remove $ and commas, convert to float
                    price_float = float(str(price_val).replace('$', '').replace(',', ''))
                    # Format with $ sign
                    if price_float.is_integer():
                        price_val = f"${int(price_float)}"
                    else:
                        price_val = f"${price_float:.2f}"
                except (ValueError, TypeError):
                    # If it's not a valid number, just add $ if not present
                    if not price_val.startswith('$'):
                        price_val = f"${price_val}"
            
            label_data["Price"] = wrap_with_marker(price_val, "PRICE")  # Fixed: Use "PRICE" marker to match markers.py definition
            
            lineage_text   = str(row.get("Lineage", "")).strip()
            product_brand  = str(row.get("Product Brand", "")).strip()
            
            # CRITICAL FIX: Handle cases where Product Type* is 'NOT_FOUND' or invalid
            raw_product_type = str(row.get("Product Type*", "")).strip()
            fallback_product_type = str(row.get("ProductType", "")).strip()
            
            if raw_product_type and raw_product_type.lower() not in ['not_found', 'unknown', '']:
                product_type = raw_product_type.lower()
            elif fallback_product_type and fallback_product_type.lower() not in ['not_found', 'unknown', '']:
                product_type = fallback_product_type.lower()
            else:
                # CRITICAL FIX: For new products without proper type, infer from product name
                product_name = str(row.get("ProductName", ""))
                if any(keyword in product_name.lower() for keyword in ['flower', 'bud', 'nug', 'herb']):
                    product_type = 'flower'
                    logger.info(f"🔧 TAG_GENERATOR INFERRED TYPE: '{product_name}' -> 'flower' (from name)")
                elif any(keyword in product_name.lower() for keyword in ['pre-roll', 'preroll', 'joint', 'blunt']):
                    product_type = 'pre-roll'
                    logger.info(f"🔧 TAG_GENERATOR INFERRED TYPE: '{product_name}' -> 'pre-roll' (from name)")
                else:
                    product_type = 'flower'  # Default to flower for new products
                    logger.info(f"🔧 TAG_GENERATOR DEFAULT TYPE: '{product_name}' -> 'flower' (default)")
            
            product_strain = str(row.get("Product Strain", "")).strip()
            
            # CRITICAL FIX: Store the processed product type in the label data
            label_data['ProductType'] = product_type
            label_data['Product Type*'] = product_type.title()  # Store as title case for consistency
            
            # Fix brand name for paraphernalia products
            if product_brand == "Paraphernalia" and product_type == "paraphernalia":
                product_name = str(row.get("ProductName", ""))
                if " by " in product_name:
                    product_brand = product_name.split(" by ")[-1].strip()
                else:
                    # Prefer Product Brand field; do NOT fall back to vendor to avoid vendor being used as brand
                    product_brand = str(row.get("Product Brand", "")).strip()
            
            # Only add brand markers for non-classic types
            # Classic types should show lineage instead of brand
            from src.core.constants import CLASSIC_TYPES
            is_classic_type = product_type in [ct.lower() for ct in CLASSIC_TYPES]
            
            if is_classic_type:
                # For classic types, lineage will be set after database lookup below
                # Don't set it here - wait for database lookup to determine lineage
                # Default will be HYBRID if no database lineage found (never Excel)
                pass
            else:
                # For non-classic types, preserve brand data for template processor
                # The template processor will handle the proper formatting based on template type
                brand_content = product_brand.upper()
                
                # CRITICAL FIX: Pass brand data to template processor instead of clearing it
                # The template processor will handle the proper formatting for each template type
                label_data["ProductBrand"] = brand_content
                label_data["ProductBrand_Center"] = brand_content
                
                # CRITICAL FIX: Set Lineage field with proper markers for double template compatibility
                # The double template uses Lineage field with PRODUCTBRAND_CENTER markers
                if orientation == 'double':
                    label_data["Lineage"] = f"PRODUCTBRAND_CENTER_START{brand_content}PRODUCTBRAND_CENTER_END"
                    logger.info(f"🔧 DOUBLE TEMPLATE BRAND DATA: Label{i+1} - Set Lineage to '{brand_content}' with markers")
                else:
                    label_data["Lineage"] = brand_content
            
            # Add other fields to label_data
            # Get product name and apply non-breaking hyphens to prevent "Pre-Roll" splitting
            product_name = str(row.get("ProductName", ""))
            if product_name:
                from src.core.generation.text_processing import make_nonbreaking_hyphens
                product_name = make_nonbreaking_hyphens(product_name)
            
            # Use Description as the primary field (never force ProductName here)
            description_raw = row.get("Description", "")
            # Normalize None/NaN to empty string to avoid literal 'None'
            try:
                import math  # local import to avoid module overhead at top level
                if description_raw is None:
                    description = ""
                elif isinstance(description_raw, float) and math.isnan(description_raw):
                    description = ""
                else:
                    description = str(description_raw)
                    if description.lower() == "nan":
                        description = ""
            except Exception:
                description = str(description_raw or "")
            
            # Apply non-breaking hyphens to description to prevent "Pre-Roll" splitting
            if description:
                description = make_nonbreaking_hyphens(description)
            
            # CRITICAL FIX: Use Excel processor's weight normalization with Excel-first priority
            from src.core.data.excel_processor import ExcelProcessor
            excel_processor = ExcelProcessor()
            weight_units = excel_processor._format_weight_units(row, excel_priority=True)
            
            print(f"DEBUG: Excel-first weight construction - Weight*: '{row.get('Weight*', '')}', Units: '{row.get('Units', '')}' -> '{weight_units}'")
            
            # CRITICAL FIX: For pre-rolls, use JointRatio instead of WeightUnits if available
            joint_ratio_value = row.get("JointRatio", "")
            if pd.isna(joint_ratio_value) or str(joint_ratio_value).lower() in ['nan', 'none', '', 'null']:
                joint_ratio_value = ""
            
            # Check if this is a pre-roll product type
            is_preroll = product_type in ['pre-roll', 'infused pre-roll'] or 'pre-roll' in product_type.lower()
            
            # For pre-rolls with JointRatio, use it instead of WeightUnits
            if is_preroll and joint_ratio_value:
                display_weight = joint_ratio_value
                print(f"DEBUG: Pre-roll product - using JointRatio '{joint_ratio_value}' instead of WeightUnits")
            else:
                display_weight = weight_units
            
            # Preserve original ProductName; keep Description as the clean field
            label_data["ProductName"] = product_name  # Do not repurpose ProductName
            label_data["Description"] = description  # Primary clean display field
            label_data["WeightUnits"] = wrap_with_marker(display_weight, "WEIGHTUNITS")  # CRITICAL FIX: Wrap with markers for template rendering
            
            # For edibles, use brand instead of lineage in the label
            edible_types = {"edible (solid)", "edible (liquid)", "high cbd edible liquid", "tincture", "topical", "capsule"}
            is_edible = product_type in edible_types
            is_horizontal_or_double_or_vertical = orientation in {"horizontal", "double", "vertical"}
            
            # For classic types, try to get lineage from database (product-level FIRST, then strain-level)
            # OPTIMIZATION: Use pre-loaded cache instead of individual queries
            if is_classic_type:
                try:
                    product_name = row.get('Product Name*', '') or row.get('ProductName', '')
                    excel_lineage = str(row.get("Lineage", "")).strip()
                    lineage_val = None
                    # 1. Product-level lineage from database
                    if product_name:
                        db_product_lineage = product_lineage_cache.get(product_name)
                        if db_product_lineage and str(db_product_lineage).strip() not in ['', 'None', 'nan']:
                            lineage_val = str(db_product_lineage).strip().upper()
                            logger.info(f"✅ DOCX LINEAGE: Using database lineage '{lineage_val}' for '{product_name}' (Excel had: '{excel_lineage}')")
                    # 2. Fallback: strain-level (sovereign_lineage > canonical_lineage)
                    if (not lineage_val or lineage_val in ['MIXED', 'THC', '', None]) and product_strain:
                        strain_info = strain_info_cache.get(product_strain)
                        if strain_info:
                            for key in ('sovereign_lineage', 'canonical_lineage'):
                                canon = strain_info.get(key)
                                if canon and str(canon).strip() not in ['', 'None', 'nan', 'SOVEREIGN']:
                                    lineage_val = str(canon).strip().upper()
                                    logger.info(f"✅ DOCX LINEAGE: Using strain {key} '{lineage_val}' for '{product_name}' (strain: '{product_strain}', Excel had: '{excel_lineage}')")
                                    break
                    # 3. Final fallback: HYBRID
                    if not lineage_val or lineage_val in ['MIXED', 'THC', '', None]:
                        lineage_val = 'HYBRID'
                        logger.warning(f"⚠️ DOCX LINEAGE: No valid lineage found for '{product_name}' (strain: '{product_strain}'), using default '{lineage_val}'")
                    # 4. Enforce: never allow MIXED or THC for classic types
                    if lineage_val in ['MIXED', 'THC', '', None]:
                        lineage_val = 'HYBRID'
                        logger.warning(f"⚠️ DOCX LINEAGE: Invalid lineage '{lineage_val}' for classic type '{product_name}', forcing to 'HYBRID'")
                except Exception as e:
                    lineage_val = 'HYBRID'
                    logger.warning(f"⚠️ DOCX LINEAGE ERROR: Database lookup failed for '{product_name}', using default '{lineage_val}' - Error: {e}")
            else:
                # CRITICAL FIX: For ALL non-classic types (edibles, tinctures, gummies, etc.), 
                # use brand name for Lineage, not the raw Excel lineage value
                # This prevents "CBD" from appearing in non-classic type labels
                # If no brand is available, use default "MIXED" - DO NOT fall back to Excel lineage
                if product_brand and str(product_brand).strip():
                    lineage_val = product_brand.upper()
                else:
                    lineage_val = 'MIXED'  # Default for non-classic types
                print(f"DEBUG NON-CLASSIC: product_type='{product_type}', product_brand='{product_brand}', lineage_val='{lineage_val}', orientation='{orientation}'")
                
            # Final hardening: normalize lineage text so LabelX.Lineage never carries extra/hidden spaces.
            lineage_val = _clean_lineage_text(lineage_val).upper()
            if not lineage_val:
                lineage_val = "HYBRID" if is_classic_type else "MIXED"
            label_data["Lineage"] = lineage_val  # Don't wrap with markers for template rendering
            
            # For classic types, set ProductBrand and ProductBrand_Center to lineage
            if is_classic_type:
                if lineage_val:
                    # Use the lineage value from database (never Excel)
                    label_data["ProductBrand"] = lineage_val  # Don't wrap with markers for template rendering
                    label_data["ProductBrand_Center"] = lineage_val  # Don't wrap with markers for template rendering
                else:
                    # No lineage available - use default HYBRID for classic types (never fall back to Excel)
                    label_data["ProductBrand"] = "HYBRID"
                    label_data["ProductBrand_Center"] = "HYBRID"
            else:
                # CRITICAL FIX: For vertical template, ensure brand goes to Lineage field only
                # and ProductStrain field is cleared to prevent 1pt font issue
                if orientation == 'vertical':
                    # For vertical template, clear ProductStrain to prevent brand from appearing in 1pt font
                    # The brand should only appear in Lineage field with proper font size
                    label_data["ProductStrain"] = ""
                    print(f"DEBUG VERTICAL FIX: Cleared ProductStrain for vertical template to prevent 1pt font issue")
                else:
                    # Non-classic types already have brand data set in the first processing section above
                    # No need to override it here - this was causing brand data to be lost
                    pass
            label_data["Ratio_or_THC_CBD"] = str(row.get("Ratio", ""))  # Don't wrap with markers for template rendering
            # ProductStrain is now handled by the template processor to ensure proper conversion
            # of non-classic types to "Mixed" or "CBD Blend"
            # Fix: Handle NaN values in JointRatio
            joint_ratio_value = row.get("JointRatio", "")
            if pd.isna(joint_ratio_value) or str(joint_ratio_value).lower() == 'nan':
                joint_ratio_value = ""
            label_data["JointRatio"] = wrap_with_marker(str(joint_ratio_value), "JOINT_RATIO")
            
            # Add THC and CBD from AI and AK columns
            ai_value = str(row.get("AI", "")).strip()
            ak_value = str(row.get("AK", "")).strip()
            
            # Clean up the values (remove 'nan', empty strings, etc.)
            if ai_value in ['nan', 'NaN', '']:
                ai_value = ""
            if ak_value in ['nan', 'NaN', '']:
                ak_value = ""
            
            # Note: Individual THC/CBD values removed - QR codes now provide this information
            
            # DescAndWeight should contain only the description text (mapped to DESC marker in template)
            # Use the processed Description field from above
            desc = description  # Use the processed description from above
            
            if desc:
                desc = re.sub(r'[-\s]+$', '', desc)
            product_type = str(row.get("Product Type*", "")).strip().lower()
            
            # For edibles in double template, use Product Brand instead of Description
            edible_types = {"edible (solid)", "edible (liquid)", "high cbd edible liquid", "tincture", "topical", "capsule"}
            if product_type in edible_types and orientation == "double":
                # Use Product Brand instead of Description for edibles in double template
                desc = product_brand if product_brand else desc
            
            # DescAndWeight should contain description + weight (mapped to DESC marker in templates)
            # Get weight units from the row
            weight_units = row.get("WeightUnits", "") or row.get("CombinedWeight", "")
            if weight_units:
                # Combine description and weight with " - " separator
                combined = f"{desc} - {weight_units}"
            else:
                combined = desc
            
            label_data["DescAndWeight"] = wrap_with_marker(combined, "DESC")
            
            context[f"Label{i+1}"] = label_data
            if DEBUG_ENABLED:
                logger.debug(f"Created label data for Label{i+1}")

    # Render template
    if DEBUG_ENABLED:
        logger.debug("Rendering template...")
    tpl.render(context)
    if DEBUG_ENABLED:
        logger.debug("Template rendered successfully")
    
    # Save to buffer
    buffer = BytesIO()
    tpl.save(buffer)
    buffer.seek(0)
    if DEBUG_ENABLED:
        logger.debug("Template saved to buffer")
    
    return buffer.getvalue()

def combine_documents(docs):
    """Combine multiple documents into one using a safer method."""
    if not docs:
        return None
        
    try:
        # Use the first document as the base
        master_doc = Document(BytesIO(docs[0]))
        
        # Add spacing between documents
        master_doc.add_paragraph()
        
        # Append each subsequent document
        for doc_bytes in docs[1:]:
            try:
                src_doc = Document(BytesIO(doc_bytes))
                
                # Copy all content from the source document
                for element in src_doc.element.body:
                    # Create a deep copy to avoid reference issues
                    new_element = deepcopy(element)
                    master_doc.element.body.append(new_element)
                    
                # Add spacing between documents
                master_doc.add_paragraph()
                
            except Exception as e:
                logger.error(f"Error combining document: {e}")
                # Continue with other documents instead of failing completely
                continue

        # Save final document to bytes
        final_buffer = BytesIO()
        master_doc.save(final_buffer)
        final_buffer.seek(0)
        
        # Validate the combined document
        try:
            test_doc = Document(final_buffer)
            # Center all tables in the document
            for table in test_doc.tables:
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
            final_buffer.seek(0)
            return final_buffer.getvalue()
        except Exception as validation_error:
            logger.error(f"Combined document validation failed: {validation_error}")
            raise ValueError(f"Combined document is corrupted: {validation_error}")
            
    except Exception as e:
        logger.error(f"Error in combine_documents: {e}")
        raise

def run_full_process_by_group(records, template_path, font_scheme):
    """Process all records using the template, grouped by strain color (lineage)."""
    if not records:
        return None
    # Define canonical lineage order (matches UI order)
    lineage_order = ["SATIVA", "HYBRID/SATIVA", "HYBRID", "HYBRID/INDICA", "INDICA", "CBD", "MIXED", "PARAPHERNALIA"]
    def get_lineage(rec):
        # Check multiple possible field names for lineage
        possible_fields = ['Lineage', 'lineage', 'Product Lineage', 'ProductLineage', 'Strain Type', 'StrainType']
        lin = ''
        for field in possible_fields:
            if field in rec and rec[field]:
                lin = str(rec[field]).strip()
                break
        
        # Normalize the lineage value
        lin = lin.upper().replace('PARA', 'PARAPHERNALIA')
        
        # Debug logging
        if DEBUG_ENABLED:
            logger.debug(f"Record: {rec.get('ProductName', 'Unknown')}, Raw lineage: {lin}, Normalized: {lin if lin in lineage_order else 'MIXED'}")
        
        return lin if lin in lineage_order else 'MIXED'
    # Sort records by lineage order, then by ProductName
    records_sorted = sorted(records, key=lambda r: (lineage_order.index(get_lineage(r)), str(r.get('ProductName', ''))))
    # Group by lineage and chunk within each group
    tag_chunks = []
    for lineage, group in groupby(records_sorted, key=get_lineage):
        group_list = list(group)
        for i in range(0, len(group_list), 9):
            tag_chunks.append(group_list[i:i+9])
    def _inner():
        template_type = Path(template_path).stem
        docs = []
        for tag_chunk in tag_chunks:
            doc = process_chunk(
                tag_chunk,
                template_path,
                font_scheme
            )
            if doc:
                docs.append(doc)
        return combine_documents(docs)
    profile = cProfile.Profile()
    result = profile.runcall(_inner)
    profile.dump_stats('profile_group.prof')
    return result

def run_full_process_by_mini(records, template_type, font_scheme, scale_factor=1.0):
    """Process records with the mini template."""
    if not records:
        return None
    def _inner():
        template_path = get_template_path(template_type)
        if hasattr(template_path, "seek"):
            template_path.seek(0)
            try:
                Document(template_path)
                template_path.seek(0)
            except Exception as e:
                raise ValueError(f"Template buffer is not a valid DOCX: {e}")
        chunks = chunk_records(records, chunk_size=20)  # Fixed: 4x5 grid = 20 labels per page
        docs = []
        for chunk in chunks:
            args = (chunk, template_path, font_scheme, template_type, scale_factor)
            docs.append(process_chunk(args))
        return combine_documents(docs)
    profile = cProfile.Profile()
    result = profile.runcall(_inner)
    profile.dump_stats('profile_mini.prof')
    return result

def generate_multiple_label_tables(records, template_path):
    """
    For each record, render the template and append the resulting table to a new document.
    Returns a BytesIO buffer with the final DOCX.
    """
    try:
        # Sort records by canonical lineage order (matches UI order), then by name
        lineage_order = ["SATIVA", "HYBRID/SATIVA", "HYBRID", "HYBRID/INDICA", "INDICA", "CBD", "MIXED", "PARAPHERNALIA"]
        def get_lineage(rec):
            possible_fields = ['Lineage', 'lineage', 'Product Lineage', 'ProductLineage', 'Strain Type', 'StrainType']
            lin = ''
            for field in possible_fields:
                if field in rec and rec[field]:
                    lin = str(rec[field]).strip()
                    break
            lin = lin.upper().replace('PARA', 'PARAPHERNALIA')
            return lin if lin in lineage_order else 'MIXED'
        records_sorted = sorted(records, key=lambda r: (lineage_order.index(get_lineage(r)), str(r.get('ProductName', ''))))
        final_doc = Document()
        # Remove default paragraph if it exists
        if final_doc.paragraphs:
            p = final_doc.paragraphs[0]
            p._element.getparent().remove(p._element)
        
        # ⚡ PERFORMANCE: Pre-load template once instead of per-record
        template_doc = None
        try:
            template_doc = DocxTemplate(template_path)
        except Exception as template_err:
            logger.error(f"Failed to load template: {template_err}")
            raise
        
        # Group by lineage and chunk within each group
        for lineage, group in groupby(records_sorted, key=get_lineage):
            group_list = list(group)
            for i in range(0, len(group_list), 9):
                chunk = group_list[i:i+9]
                # Create a single 3x3 table for this chunk
                table = final_doc.add_table(rows=3, cols=3)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                table.allow_autofit = False
                tblPr = table._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
                tblLayout = OxmlElement('w:tblLayout')
                tblLayout.set(qn('w:type'), 'fixed')
                tblPr.append(tblLayout)
                table._element.insert(0, tblPr)
                col_width = Inches(1.0)
                tblGrid = OxmlElement('w:tblGrid')
                for _ in range(3):
                    gridCol = OxmlElement('w:gridCol')
                    gridCol.set(qn('w:w'), str(int(col_width.inches * 1440)))
                    tblGrid.append(gridCol)
                table._element.insert(0, tblGrid)
                # Fill the table cells with tag data
                for idx in range(9):
                    r, c = divmod(idx, 3)
                    cell = table.cell(r, c)
                    if idx < len(chunk):
                        record = chunk[idx]
                        try:
                            # ⚡ PERFORMANCE: Reuse template, create fresh copy for rendering
                            from copy import deepcopy
                            doc = deepcopy(template_doc)
                            context = build_context(record, doc)
                            doc.render(context)
                            tmp_stream = BytesIO()
                            doc.save(tmp_stream)
                            tmp_stream.seek(0)
                            rendered_doc = Document(tmp_stream)
                            if rendered_doc.tables:
                                src_table = rendered_doc.tables[0]
                                src_cell = src_table.cell(0, 0)
                                for para in cell.paragraphs:
                                    para._element.getparent().remove(para._element)
                                for para in src_cell.paragraphs:
                                    new_para = cell.add_paragraph()
                                    new_para.alignment = para.alignment
                                    for run in para.runs:
                                        try:
                                            new_run = new_para.add_run(run.text)
                                            # Always force Arial Bold for consistency across platforms
                                            new_run.font.name = "Arial"
                                            new_run.font.bold = True
                                            new_run.italic = run.italic
                                            new_run.underline = run.underline
                                            if run.font.size:
                                                new_run.font.size = run.font.size
                                            if run.font.color and run.font.color.rgb:
                                                new_run.font.color.rgb = run.font.color.rgb
                                        except Exception as run_error:
                                            logger.warning(f"Error copying run: {run_error}")
                                            new_para.add_run(run.text)
                                if not cell.paragraphs:
                                    cell.add_paragraph()
                        except Exception as record_error:
                            logger.error(f"Error processing record: {record_error}")
                            cell.text = ''
                    else:
                        cell.text = ''
                    cell.width = col_width
                for row in table.rows:
                    row.height = Inches(1.0)
                    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                
                # Enforce fixed cell dimensions to prevent any growth
                try:
                    # Safety check: ensure table has valid structure
                    if table and table.rows and len(table.rows) > 0:
                        first_row = table.rows[0]
                        if hasattr(first_row, '_element') and hasattr(first_row._element, 'tc_lst'):
                            enforce_fixed_cell_dimensions(table)
                        else:
                            logger.warning(f"Skipping table with invalid XML structure in tag generator")
                    else:
                        logger.warning(f"Skipping empty or invalid table in tag generator")
                except Exception as e:
                    logger.warning(f"Error enforcing fixed cell dimensions in tag generator: {e}")
                    continue
                
                final_doc.add_paragraph()
        # Center all tables in the final document
        for table in final_doc.tables:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Ensure all fonts are Arial Bold for consistency across platforms
        from src.core.generation.docx_formatting import enforce_arial_bold_all_text
        enforce_arial_bold_all_text(final_doc)
        
        # Save final document
        output = BytesIO()
        final_doc.save(output)
        output.seek(0)
        
        # Validate the final document
        try:
            test_doc = Document(output)
            output.seek(0)
            return output
        except Exception as validation_error:
            logger.error(f"Final document validation failed: {validation_error}")
            raise ValueError(f"Generated document is corrupted: {validation_error}")
            
    except Exception as e:
        logger.error(f"Error in generate_multiple_label_tables: {e}")
        raise

def set_table_borders(table):
    """Remove all table borders to eliminate white lines."""
    tblPr = table._element.find(qn('w:tblPr'))
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
        
    tblBorders = OxmlElement('w:tblBorders')
    
    # Remove ALL borders (outer and interior)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn('w:val'), "nil")
        tblBorders.append(bd)
        
    tblPr.append(tblBorders)

def debug_markers(text):
    """Debug helper to identify marker issues."""
    markers = ['DESC', 'WEIGHTUNITS', 'PRODUCTBRAND_CENTER', 'PRICE', 'LINEAGE', 'THC_CBD']
    found_markers = []
    
    for marker in markers:
        start = f"{marker}_START"
        end = f"{marker}_END"
        if start in text:
            pos = text.find(start)
            found_markers.append(f"{marker} start at {pos}")
        if end in text:
            pos = text.find(end)
            found_markers.append(f"{marker} end at {pos}")
            
    if found_markers:
        if DEBUG_ENABLED:
            logger.debug(f"Found markers in text: {text}")
        for marker in found_markers:
            if DEBUG_ENABLED:
                logger.debug(f"  {marker}")
    return found_markers

def validate_and_repair_document(doc_bytes):
    """Validate a document and attempt to repair common issues."""
    try:
        # Try to load the document
        doc = Document(BytesIO(doc_bytes))
        
        # Check for common corruption issues
        issues_found = []
        
        # Check if document has content
        if not doc.paragraphs and not doc.tables:
            issues_found.append("Document has no content")
            
        # Check for malformed tables
        for table in doc.tables:
            try:
                # Try to access table properties
                _ = table.rows
                _ = table.columns
            except Exception as e:
                issues_found.append(f"Malformed table: {e}")
                
        # Check for malformed paragraphs
        for para in doc.paragraphs:
            try:
                # Try to access paragraph properties
                _ = para.runs
            except Exception as e:
                issues_found.append(f"Malformed paragraph: {e}")
        
        if issues_found:
            logger.warning(f"Document validation issues found: {issues_found}")
            return False, issues_found
            
        return True, []
        
    except Exception as e:
        logger.error(f"Document validation failed: {e}")
        return False, [f"Document cannot be loaded: {e}"]

def create_safe_document():
    """Create a minimal safe document as a fallback."""
    try:
        doc = Document()
        # Add a simple paragraph to ensure the document is valid
        doc.add_paragraph("Document generated successfully.")
        return doc
    except Exception as e:
        logger.error(f"Error creating safe document: {e}")
        raise



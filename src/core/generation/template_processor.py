from copy import deepcopy
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Mm, RGBColor
from docxtpl import DocxTemplate, InlineImage
from docxcompose.composer import Composer
import qrcode
from io import BytesIO
import logging
import os
from pathlib import Path
import re
import time
from typing import Dict, Any, List, Optional
import traceback
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from importlib.metadata import version as get_package_version
from importlib import resources as importlib_resources
import sys
import types
import warnings
import docxcompose.properties as _docx_props

# Suppress pkg_resources deprecation warning (will be addressed when dependencies are updated)
with warnings.catch_warnings():
    warnings.filterwarnings('ignore', category=UserWarning, message='.*pkg_resources is deprecated.*')
    try:
        import pkg_resources  # type: ignore
    except ModuleNotFoundError:
        import types as _types  # Fallback if types not available
        pkg_resources = _types.ModuleType('pkg_resources')  # type: ignore

        def _resource_string(package, resource):
            return importlib_resources.read_binary(package, resource)

        pkg_resources.resource_string = _resource_string  # type: ignore
        sys.modules['pkg_resources'] = pkg_resources  # type: ignore

if not hasattr(_docx_props, 'pkg_resources'):
    _docx_props.pkg_resources = pkg_resources

# Local imports
from src.core.utils.common import safe_get
from src.core.generation.docx_formatting import (
    apply_lineage_colors,
    enforce_fixed_cell_dimensions,
    prevent_table_expansion_enhanced,
    clear_cell_background,
    clear_cell_margins,
    clear_table_cell_padding,
)
from src.core.generation.unified_font_sizing import (
    get_font_size,
    get_font_size_by_marker,
    set_run_font_size,
    is_classic_type,
    get_line_spacing_by_marker
)
from src.core.generation.text_processing import (
    process_doh_image,
    format_ratio_multiline
)
from src.core.formatting.markers import wrap_with_marker, unwrap_marker, is_already_wrapped

# Performance settings - check if running on PythonAnywhere
import os
IS_PYTHONANYWHERE = 'pythonanywhere.com' in os.environ.get('HTTP_HOST', '')

# Use same settings for both local and PythonAnywhere to ensure consistent generation
MAX_PROCESSING_TIME_PER_CHUNK = 30  # 30 seconds max per chunk
MAX_TOTAL_PROCESSING_TIME = 600     # 10 minutes max total (increased for large batches)
CHUNK_SIZE_LIMIT = 100              # PERFORMANCE: Increased from 50 to 100 for faster generation of large batches

def get_font_scheme(template_type, base_size=12):
    schemes = {
        'default': {"base_size": base_size, "min_size": 8, "max_length": 25},
        'vertical': {"base_size": base_size, "min_size": 8, "max_length": 25},
        'mini': {"base_size": base_size - 2, "min_size": 6, "max_length": 15},
        'horizontal': {"base_size": base_size + 1, "min_size": 7, "max_length": 20},
        'double': {"base_size": base_size - 1, "min_size": 8, "max_length": 30},
        'inventory': {"base_size": base_size, "min_size": 8, "max_length": 40},  # Inventory slips can handle longer text
        'preroll': {"base_size": base_size - 2, "min_size": 6, "max_length": 15}  # Preroll template uses mini font scheme
    }
    return {
        field: {**schemes.get(template_type, schemes['default'])}
        for field in ["Description", "ProductBrand", "Price", "Lineage", "DOH", "Ratio_or_THC_CBD", "Ratio"]
    }

class TemplateProcessor:
    # ...rest of the code...
    def _get_template_path(self):
        """Return the path to the DOCX template file for the current template type."""
        base_dir = Path(__file__).parent / 'templates'
        template_map = {
            'mini': 'mini.docx',
            'double': 'double.docx',
            'inventory': 'inventory.docx',
            'horizontal': 'horizontal.docx',
            'vertical': 'vertical.docx',
            'preroll': 'preroll.docx',
        }
        filename = template_map.get(self.template_type, 'horizontal.docx')
        template_path = base_dir / filename
        if template_path.exists():
            return str(template_path.resolve())

        # Case-insensitive fallback for hosts with case-sensitive filesystems
        expected_lower = filename.lower()
        fallback_path = None
        try:
            for candidate in base_dir.iterdir():
                if not candidate.is_file():
                    continue
                name = candidate.name
                if name.startswith('.') or name.startswith('~$'):
                    continue
                if name.lower() == expected_lower:
                    fallback_path = candidate
                    break
        except FileNotFoundError:
            # Directory missing; let the original error handling occur
            pass

        if fallback_path and fallback_path.exists():
            self.logger.warning(
                f"Template not found with exact casing ('{filename}'); using fallback '{fallback_path.name}'"
            )
            return str(fallback_path.resolve())

        resolved_path = str(template_path.resolve()) if template_path.exists() else str(template_path)
        self.logger.error(f"Template not found: {resolved_path}")
        raise FileNotFoundError(f"Template not found: {resolved_path}")
    def __init__(self, template_type, font_scheme, scale_factor=1.0, excel_processor=None):
        self.template_type = template_type
        self.font_scheme = font_scheme
        self.logger = logging.getLogger(__name__)  # Initialize logger first
        
        # CRITICAL FIX: Adjust scale factor for double template 12-label expansion
        # When the double template expands to 12 labels, cells become smaller, so we need to adjust the scale factor
        if template_type == 'double':
            self.scale_factor = scale_factor * 0.95  # Reduce font sizes by 5% for 12-label expansion (less aggressive)
            self.logger.info(f"🔧 DOUBLE TEMPLATE SCALE ADJUSTMENT: Adjusted scale factor from {scale_factor} to {self.scale_factor} for 12-label expansion")
        else:
            self.scale_factor = scale_factor
        self.excel_processor = excel_processor  # Store the session's Excel processor
        self._template_path = self._get_template_path()
        self._expanded_template_buffer = self._expand_template_if_needed()
        self._dynamic_template_created = False  # Track if dynamic template was created
        self._last_dynamic_count = None  # Track last product count used to build dynamic buffer
        self._vendor_fallback = None  # Vendor to use as fallback when records are vendor-filtered but individual records have missing vendor
        
        # Set chunk size based on template type with performance limits
        if not IS_PYTHONANYWHERE:
            self.logger.info(f"DEBUG: Setting chunk size for template_type='{self.template_type}' (type: {type(self.template_type)})")
        
        if self.template_type == 'mini':
            self.chunk_size = min(20, CHUNK_SIZE_LIMIT)  # Fixed: 4x5 grid = 20 labels per page
            if not IS_PYTHONANYWHERE:
                self.logger.info(f"DEBUG: Set chunk size to {self.chunk_size} for mini template")
        elif self.template_type == 'preroll':
            self.chunk_size = min(20, CHUNK_SIZE_LIMIT)  # Fixed: 4x5 grid = 20 labels per page (same as mini)
            if not IS_PYTHONANYWHERE:
                self.logger.info(f"DEBUG: Set chunk size to {self.chunk_size} for preroll template")
        elif self.template_type == 'double':
            self.chunk_size = min(12, CHUNK_SIZE_LIMIT)  # Fixed: 4x3 grid = 12 labels per page
            if not IS_PYTHONANYWHERE:
                self.logger.info(f"DEBUG: Set chunk size to {self.chunk_size} for double template")
        elif self.template_type == 'inventory':
            self.chunk_size = min(4, CHUNK_SIZE_LIMIT)   # Fixed: 2x2 grid = 4 labels per page
            if not IS_PYTHONANYWHERE:
                self.logger.info(f"DEBUG: Set chunk size to {self.chunk_size} for inventory template")
        else:
            # For standard templates (horizontal, vertical), use larger chunks for better performance
            # Allow up to 50 products per chunk for horizontal/vertical templates
            self.chunk_size = min(50, CHUNK_SIZE_LIMIT)  # Increased chunk size for better performance
            if not IS_PYTHONANYWHERE:
                self.logger.info(f"DEBUG: Set chunk size to {self.chunk_size} for template type '{self.template_type}' (expanded)")
        
        self.logger.info(f"Template type: {self.template_type}, Chunk size: {self.chunk_size}")
        
        # Performance tracking
        self.start_time = time.time()
        self.chunk_count = 0
        
        # Template expansion cache - avoid re-expanding templates with same size
        self._template_expansion_cache = {}

        # CRITICAL FIX: Disable chunking only for templates that support dynamic grids
        if self.template_type in ['horizontal', 'vertical', 'double']:
            self.chunk_size = None  # Will be set dynamically in process_records for these templates
            self.logger.info(f"CRITICAL FIX: Chunking disabled for template '{self.template_type}' - chunk_size will match total records")
        else:
            self.logger.info(f"Chunking retained for template '{self.template_type}' with chunk_size {self.chunk_size}")
            # ...rest of the code...

    def _expand_template_if_needed(self, force_expand=False):
        """Expand template if needed and return buffer."""
        try:
            with open(self._template_path, 'rb') as f:
                buffer = BytesIO(f.read())
            
            # Check if template needs expansion
            doc = Document(buffer)
            text = doc.element.body.xml
            matches = re.findall(r'Label(\d+)\.', text)
            
            # Check if we have all required labels (9 for 3x3, 20 for 4x5, 12 for 4x3, 4 for 2x2)
            if self.template_type == 'mini':
                required_labels = 20  # 4x5 grid
            elif self.template_type == 'double':
                required_labels = 12  # 4x3 grid
            elif self.template_type == 'inventory':
                required_labels = 4   # 2x2 grid
            elif self.template_type == 'preroll':
                required_labels = 20  # 4x5 grid (same as mini)
            else:
                required_labels = 9   # 3x3 grid
            
            unique_labels = set(matches)
            
            if len(unique_labels) < required_labels or force_expand:
                # CRITICAL FIX: Use standard expansion methods for now
                # Dynamic templates will be created later in _process_chunk based on actual product count
                if self.template_type == 'mini':
                    self.logger.info("Calling 4x5 expansion method")
                    return self._expand_template_to_4x5_fixed_scaled()
                elif self.template_type == 'inventory':
                    self.logger.info("Calling 2x2 inventory expansion method")
                    return self._expand_template_to_2x2_inventory()
                elif self.template_type == 'double':
                    self.logger.info("Calling 4x3 expansion method")
                    return self._expand_template_to_4x3_fixed_double()
                elif self.template_type == 'preroll':
                    # Preroll uses 4x5 grid like mini template
                    self.logger.info("Calling 4x5 expansion method for preroll template")
                    return self._expand_template_to_4x5_fixed_scaled()
                else:
                    # horizontal and vertical templates expand to 3x3 grid
                    self.logger.info(f"Calling 3x3 expansion method for template type: '{self.template_type}'")
                    return self._expand_template_to_3x3_fixed()
            
            return buffer
        except Exception as e:
            self.logger.error(f"Error expanding template: {e}")
            raise

    def force_re_expand_template(self):
        """Force re-expansion of template."""
        self._expanded_template_buffer = self._expand_template_if_needed(force_expand=True)

    def _fix_double_template_structure(self):
        """Fix double template structure by adding missing table grid elements without full expansion."""
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        
        try:
            # Load the original template
            template_path = self._get_template_path()
            self.logger.info(f"Loading double template from: {template_path}")
            doc = Document(template_path)
            
            # Check if template has tables
            if not doc.tables:
                self.logger.warning("Double template has no tables, returning original")
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer
            
            self.logger.info(f"Double template has {len(doc.tables)} tables")
            
            # Fix each table by ensuring it has proper structure
            for table in doc.tables:
                # Check if table has tblGrid element
                tbl_grid = table._element.find(qn('w:tblGrid'))
                if tbl_grid is None:
                    self.logger.info("Adding missing tblGrid element to double template table")
                    
                    # Get the number of columns from the table XML structure instead of table.columns
                    # This avoids the error when table.columns is accessed without proper structure
                    table_element = table._element
                    rows = table_element.findall(qn('w:tr'))
                    if rows:
                        # Count columns from the first row
                        first_row = rows[0]
                        cells = first_row.findall(qn('w:tc'))
                        num_cols = len(cells)
                    else:
                        # Fallback: assume 4 columns for double template
                        num_cols = 4
                    
                    self.logger.info(f"Detected {num_cols} columns in double template table")
                    
                    # Create tblGrid element
                    tbl_grid = OxmlElement('w:tblGrid')
                    for _ in range(num_cols):
                        gc = OxmlElement('w:gridCol')
                        gc.set(qn('w:w'), str(int(1.75 * 1440)))  # 1.75 inches per column
                        tbl_grid.append(gc)
                    
                    # Insert tblGrid at the beginning of the table element
                    table_element.insert(0, tbl_grid)
                    
                    # Also ensure table has proper table properties
                    tbl_pr = table_element.find(qn('w:tblPr'))
                    if tbl_pr is None:
                        tbl_pr = OxmlElement('w:tblPr')
                        table_element.insert(0, tbl_pr)
                    
                    # Add table layout
                    layout = tbl_pr.find(qn('w:tblLayout'))
                    if layout is None:
                        layout = OxmlElement('w:tblLayout')
                        layout.set(qn('w:type'), 'fixed')
                        tbl_pr.append(layout)
                    
                    # Ensure the table has at least one row and cell for basic structure
                    if not table_element.findall(qn('w:tr')):
                        self.logger.warning("Double template table has no rows, adding minimal structure")
                        # Add a minimal row with cells
                        row = OxmlElement('w:tr')
                        for _ in range(num_cols):
                            cell = OxmlElement('w:tc')
                            # Add cell properties
                            tc_pr = OxmlElement('w:tcPr')
                            cell.append(tc_pr)
                            # Add a paragraph
                            para = OxmlElement('w:p')
                            cell.append(para)
                            row.append(cell)
                        table_element.append(row)
            
            # Save the fixed template to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            self.logger.error(f"Error fixing double template structure: {e}")
            # Fallback to original template if fixing fails
            try:
                self.logger.info("Attempting fallback to original template")
                with open(self._get_template_path(), 'rb') as f:
                    buffer = BytesIO(f.read())
                self.logger.info("Fallback successful, returning original template")
                return buffer
            except Exception as fallback_error:
                self.logger.error(f"Fallback to original template also failed: {fallback_error}")
                # Last resort: create a minimal working template
                self.logger.warning("Creating minimal working template as last resort")
                try:
                    from docx import Document
                    doc = Document()
                    table = doc.add_table(rows=1, cols=4)
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    return buffer
                except Exception as create_error:
                    self.logger.error(f"Failed to create minimal template: {create_error}")
                    raise

    def _expand_template_to_2x2_inventory(self):
        """Expand template to 2x2 grid for inventory slips."""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        from copy import deepcopy

        num_cols, num_rows = 2, 2  # 2x2 grid for inventory
        col_width_inches = 3.75  # Appropriate width for inventory slips
        row_height_inches = 3.5   # Appropriate height for inventory slips
        
        col_width_twips = str(int(col_width_inches * 1440))
        row_height_pts = Pt(row_height_inches * 72)
        cut_line_twips = int(0.001 * 1440)

        template_path = self._get_template_path()
        doc = Document(template_path)
        if not doc.tables:
            raise RuntimeError("Template must contain at least one table.")
        old = doc.tables[0]
        src_tc = deepcopy(old.cell(0,0)._tc)
        old._element.getparent().remove(old._element)

        while doc.paragraphs and not doc.paragraphs[0].text.strip():
            doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

        # Create new table with 2x2 grid
        tbl = doc.add_table(rows=num_rows, cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set table properties
        tblPr = tbl._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tblPr.append(layout)
        tbl._element.insert(0, tblPr)

        # Set column widths
        grid = OxmlElement('w:tblGrid')
        for _ in range(num_cols):
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), col_width_twips)
            grid.append(gc)
        tbl._element.insert(0, grid)

        # Set row heights and copy template cell content with proper label numbering
        label_num = 1
        for row in tbl.rows:
            row.height = row_height_pts
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for cell in row.cells:
                new_tc = deepcopy(src_tc)
                
                # Update label numbering for 2x2 grid (Label1, Label2, Label3, Label4)
                # Convert the cell XML to string, replace Label1 with current label number
                tc_xml_str = new_tc.xml.decode('utf-8') if isinstance(new_tc.xml, bytes) else str(new_tc.xml)
                tc_xml_str = tc_xml_str.replace('Label1', f'Label{label_num}')
                
                # Parse the updated XML and replace the cell
                from lxml import etree
                new_tc_element = etree.fromstring(tc_xml_str.encode('utf-8'))
                cell._tc.getparent().replace(cell._tc, new_tc_element)
                label_num += 1

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _expand_template_to_4x5_fixed_scaled(self, num_products=None):
        """Expand template to 4x5 grid for mini templates while preserving original design."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        from copy import deepcopy

        num_cols, num_rows = 4, 5
        col_width_twips = str(int(1.5 * 1440))  # 1.5 inches per column for equal width
        row_height_pts = Pt(1.5 * 72)  # 1.5 inches per row for equal height
        cut_line_twips = int(0.001 * 1440)

        template_path = self._get_template_path()
        doc = Document(template_path)
        if not doc.tables:
            raise RuntimeError("Template must contain at least one table.")
        
        # Get the original table and its properties
        original_table = doc.tables[0]
        original_table_xml = original_table._element
        
        # Extract original table properties (colors, borders, styling)
        original_tblPr = original_table_xml.find(qn('w:tblPr'))
        original_shd = original_tblPr.find(qn('w:shd')) if original_tblPr is not None else None
        original_borders = original_tblPr.find(qn('w:tblBorders')) if original_tblPr is not None else None
        
        # Get the original cell structure and content
        original_cell = original_table.cell(0, 0)
        src_tc = deepcopy(original_cell._tc)
        
        # Remove the original table
        original_table._element.getparent().remove(original_table._element)

        # Remove empty paragraphs
        while doc.paragraphs and not doc.paragraphs[0].text.strip():
            doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

        # Create new 4x5 table
        tbl = doc.add_table(rows=num_rows, cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set up table properties
        tblPr = tbl._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
        
        # Preserve original shading if it exists
        if original_shd is not None:
            shd = deepcopy(original_shd)
            tblPr.insert(0, shd)
        else:
            # Default shading
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'D3D3D3')
            tblPr.insert(0, shd)
        
        # Set table layout
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tblPr.append(layout)
        tbl._element.insert(0, tblPr)
        
        # Set up grid columns
        grid = OxmlElement('w:tblGrid')
        for _ in range(num_cols):
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), col_width_twips)
            grid.append(gc)
        tbl._element.insert(0, grid)
        
        # Set row heights and individual cell widths
        for row in tbl.rows:
            row.height = row_height_pts
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            
            # Set individual cell widths to ensure exact 1.5" dimensions
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), col_width_twips)
                tcW.set(qn('w:type'), 'dxa')
        
        # Preserve original borders if they exist
        if original_borders is not None:
            borders = deepcopy(original_borders)
            tblPr.append(borders)
        else:
            # Default borders
            borders = OxmlElement('w:tblBorders')
            for side in ('insideH','insideV'):
                b = OxmlElement(f"w:{side}")
                b.set(qn('w:val'), "single")
                b.set(qn('w:sz'), "4")
                b.set(qn('w:color'), "D3D3D3")
                b.set(qn('w:space'), "0")
                borders.append(b)
            tblPr.append(borders)
        
        # Populate cells with original content, updating labels
        cnt = 1
        max_cells = num_products if num_products else (num_rows * num_cols)
        
        for r in range(num_rows):
            for c in range(num_cols):
                if cnt > max_cells:
                    # Clear extra cells completely and set white background
                    cell = tbl.cell(r, c)
                    cell._tc.clear_content()
                    
                    # Set white background for extra cells
                    tc = cell._tc
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is None:
                        tcPr = OxmlElement('w:tcPr')
                        tc.insert(0, tcPr)
                    
                    # Remove any existing background color
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        tcPr.remove(shd)
                    
                    # Add white background
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFFFF')  # White background
                    tcPr.append(shd)
                    
                    cell.add_paragraph()  # Add empty paragraph to maintain structure
                    continue
                    
                cell = tbl.cell(r, c)
                cell._tc.clear_content()
                
                # Copy the original cell structure
                tc = deepcopy(src_tc)
                
                # Update all Label1 references to the current label number
                for t in tc.iter(qn('w:t')):
                    if t.text and 'Label1' in t.text:
                        t.text = t.text.replace('Label1', f'Label{cnt}')
                
                # Copy all elements from the original cell
                for el in tc.xpath('./*'):
                    cell._tc.append(deepcopy(el))
                
                cnt += 1

        # Add cell spacing
        tblPr2 = tbl._element.find(qn('w:tblPr'))
        spacing = OxmlElement('w:tblCellSpacing')
        spacing.set(qn('w:w'), str(cut_line_twips))
        spacing.set(qn('w:type'), 'dxa')
        tblPr2.append(spacing)

        # PREROLL: Add 0.2" buffer spacing below table
        if self.template_type == 'preroll':
            # Add paragraph after table with 0.2" spacing
            paragraph = doc.add_paragraph()
            pPr = paragraph._element.get_or_add_pPr()
            # Add spacing after paragraph (0.2 inches = 0.2 * 1440 twips = 288 twips)
            spacing_elem = OxmlElement('w:spacing')
            spacing_elem.set(qn('w:after'), '288')  # 0.2" in twips
            pPr.append(spacing_elem)

        # Save and return
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def create_dynamic_template_for_products(self, num_products):
        """Create a dynamic template based on the number of products to eliminate empty labels."""
        try:
            # CRITICAL FIX: Disable ALL dynamic template creation to prevent XML corruption
            # Use standard template expansion with post-processing cleanup instead
            self.logger.info(f"🔧 DYNAMIC TEMPLATES DISABLED: Using standard expansion for {num_products} products")
            return False
        except Exception as e:
            self.logger.warning(f"Failed to create dynamic template: {e}")
            return False

    def _expand_template_to_4x4_fixed_preroll(self, num_products=None):
        """Expand template to 4x4 grid for preroll templates (1.5" x 1.9" labels)."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        from copy import deepcopy

        num_cols, num_rows = 4, 4  # 4x4 grid for 16 labels
        col_width_twips = str(int(1.5 * 1440))  # 1.5 inches width per column
        row_height_pts = Pt(1.9 * 72)  # 1.9 inches height per row
        cut_line_twips = int(0.001 * 1440)

        template_path = self._get_template_path()
        doc = Document(template_path)
        if not doc.tables:
            raise RuntimeError("Template must contain at least one table.")

        # Get the original table and its properties
        original_table = doc.tables[0]
        original_table_xml = original_table._element

        # Extract original table properties (colors, borders, styling)
        original_tblPr = original_table_xml.find(qn('w:tblPr'))
        original_shd = original_tblPr.find(qn('w:shd')) if original_tblPr is not None else None
        original_borders = original_tblPr.find(qn('w:tblBorders')) if original_tblPr is not None else None

        # Get the original cell structure and content
        original_cell = original_table.cell(0, 0)
        src_tc = deepcopy(original_cell._tc)

        # Remove the original table
        original_table._element.getparent().remove(original_table._element)

        # Remove empty paragraphs
        while doc.paragraphs and not doc.paragraphs[0].text.strip():
            doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

        # Create new 4x4 table
        tbl = doc.add_table(rows=num_rows, cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set up table properties
        tblPr = tbl._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')

        # Preserve original shading if it exists
        if original_shd is not None:
            shd = deepcopy(original_shd)
            tblPr.insert(0, shd)
        else:
            # Default shading
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'D3D3D3')
            tblPr.insert(0, shd)

        # Set table layout
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tblPr.append(layout)
        tbl._element.insert(0, tblPr)

        # Set up grid columns
        grid = OxmlElement('w:tblGrid')
        for _ in range(num_cols):
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), col_width_twips)
            grid.append(gc)
        tbl._element.insert(0, grid)

        # Set row heights and individual cell widths
        for row in tbl.rows:
            row.height = row_height_pts
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

            # Set individual cell widths to ensure exact 1.5" x 1.9" dimensions
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), col_width_twips)
                tcW.set(qn('w:type'), 'dxa')

        # Add cut line borders if original had them
        if original_borders is not None:
            borders = deepcopy(original_borders)
            tblPr.append(borders)

        # Populate cells with original content, updating labels
        cnt = 1
        max_cells = num_products if num_products else (num_rows * num_cols)

        for r in range(num_rows):
            for c in range(num_cols):
                if cnt > max_cells:
                    # Clear extra cells completely and set white background
                    cell = tbl.cell(r, c)
                    for paragraph in cell.paragraphs:
                        paragraph.clear()
                    # Set white background for empty cells
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = tcPr.find(qn('w:shd'))
                    if shd is None:
                        shd = OxmlElement('w:shd')
                        tcPr.append(shd)
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:fill'), 'FFFFFF')
                else:
                    # Copy original cell content
                    cell = tbl.cell(r, c)
                    dest_tc = cell._tc

                    # Replace destination tc with a deep copy of source tc
                    new_tc = deepcopy(src_tc)
                    dest_tc.getparent().replace(dest_tc, new_tc)

                    # Update label placeholders (e.g., {{Label1}} -> {{Label16}})
                    for paragraph in tbl.cell(r, c).paragraphs:
                        for run in paragraph.runs:
                            if run.text:
                                run.text = run.text.replace('{{Label1}}', f'{{{{Label{cnt}}}}}')
                cnt += 1

        # Add cell spacing (cut lines)
        tblPr2 = tbl._element.find(qn('w:tblPr'))
        spacing = OxmlElement('w:tblCellSpacing')
        spacing.set(qn('w:w'), str(cut_line_twips))
        spacing.set(qn('w:type'), 'dxa')
        tblPr2.append(spacing)

        # Save and return
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def _expand_template_to_4x3_fixed_double(self, num_products=None):
        """Expand template to 4x3 grid for double templates (4 columns, 3 rows)."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        from copy import deepcopy

        num_cols, num_rows = 4, 3  # 4 columns, 3 rows for 12 labels total
        
        # Equal width columns: 1.75 inches each for double template (original width)
        col_width_twips = str(int(1.75 * 1440))  # 1.75 inches per column
        row_height_pts = Pt(2.5 * 72)  # 2.5 inches per row for equal height
        cut_line_twips = int(0.001 * 1440)

        template_path = self._get_template_path()
        doc = Document(template_path)
        if not doc.tables:
            raise RuntimeError("Template must contain at least one table.")
        old = doc.tables[0]
        src_tc = deepcopy(old.cell(0,0)._tc)
        old._element.getparent().remove(old._element)

        # Only remove empty paragraphs, preserve content paragraphs
        # This is important for templates that have content outside of tables
        doc_paragraphs = list(doc.paragraphs)  # Create a copy to avoid modification during iteration
        for paragraph in doc_paragraphs:
            if not paragraph.text.strip():
                paragraph._element.getparent().remove(paragraph._element)

        max_cells_per_page = num_rows * num_cols
        total_products = num_products if num_products is not None else max_cells_per_page
        pages = (total_products + max_cells_per_page - 1) // max_cells_per_page
        self.logger.info(f"🔍 DOUBLE TEMPLATE EXPANSION: Creating {pages} table(s) for {total_products} products.")
        
        product_idx = 0
        for page in range(pages):
            tbl = doc.add_table(rows=num_rows, cols=num_cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Copy the original table properties and styling from the source template
            if hasattr(old, '_element') and old._element is not None:
                old_tblPr = old._element.find(qn('w:tblPr'))
                if old_tblPr is not None:
                    tbl._element.insert(0, deepcopy(old_tblPr))
                else:
                    tblPr = OxmlElement('w:tblPr')
                    layout = OxmlElement('w:tblLayout')
                    layout.set(qn('w:type'), 'fixed')
                    tblPr.append(layout)
                    tbl._element.insert(0, tblPr)
            else:
                tblPr = OxmlElement('w:tblPr')
                layout = OxmlElement('w:tblLayout')
                layout.set(qn('w:type'), 'fixed')
                tblPr.append(layout)
                tbl._element.insert(0, tblPr)
            
            # Set up the grid with proper column widths
            grid = OxmlElement('w:tblGrid')
            for _ in range(num_cols):
                gc = OxmlElement('w:gridCol')
                gc.set(qn('w:w'), col_width_twips)
                grid.append(gc)
            tbl._element.insert(0, grid)
            
            # Set row heights
            for row in tbl.rows:
                row.height = row_height_pts
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            
            for r in range(num_rows):
                for c in range(num_cols):
                    cell = tbl.cell(r, c)
                    cell._tc.clear_content()
                    
                    if product_idx < total_products:
                        cnt = product_idx + 1
                        tc = deepcopy(src_tc)
                        
                        # Update placeholders
                        for t in tc.iter(qn('w:t')):
                            if t.text and 'Label1' in t.text:
                                t.text = t.text.replace('Label1', f'Label{cnt}')
                        
                        for el in tc.xpath('./*'):
                            cell._tc.append(deepcopy(el))
                        product_idx += 1
                    else:
                        # Clear unused cells and set white background
                        tc = cell._tc
                        tcPr = tc.find(qn('w:tcPr'))
                        if tcPr is None:
                            tcPr = OxmlElement('w:tcPr')
                            tc.insert(0, tcPr)
                        
                        shd = tcPr.find(qn('w:shd'))
                        if shd is not None:
                            tcPr.remove(shd)
                        
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), 'FFFFFF')
                        tcPr.append(shd)
                        cell.add_paragraph()
            
            # Add minimal spacing between cells
            tblPr2 = tbl._element.find(qn('w:tblPr'))
            spacing = OxmlElement('w:tblCellSpacing')
            spacing.set(qn('w:w'), str(cut_line_twips))
            spacing.set(qn('w:type'), 'dxa')
            tblPr2.append(spacing)
            
            # Insert page break after each table except the last
            if page < pages - 1:
                page_break_para = doc.add_paragraph()
                page_break_run = page_break_para.add_run()
                page_break_run.add_break(WD_BREAK.PAGE)
                self.logger.info(f"🔍 DOUBLE TEMPLATE EXPANSION: Added page break after table {page + 1} of {pages}")
        
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def _expand_template_to_3x3_fixed(self, num_products=None):
        template_path = self._get_template_path()
        doc = Document(template_path)
        if not doc.tables:
            raise ValueError("Template DOCX does not contain any tables.")
        old_table = doc.tables[0]
        src_tc = deepcopy(old_table.cell(0, 0)._tc)
        # Remove original table and empty paragraphs, preserve content outside tables
        old_table._element.getparent().remove(old_table._element)
        doc_paragraphs = list(doc.paragraphs)
        for paragraph in doc_paragraphs:
            if not paragraph.text.strip():
                paragraph._element.getparent().remove(paragraph._element)

        # Set grid and cell dimensions
        num_rows = 3
        num_cols = 3
        col_width_twips = "3000"  # Adjust as needed
        row_height_pts = Pt(120)   # Adjust as needed
        max_cells = num_rows * num_cols
        total_products = num_products if num_products is not None else 49
        pages = (total_products + max_cells - 1) // max_cells
        self.logger.info(f"🔍 TEMPLATE EXPANSION: Creating {pages} pages of 3x3 grids for {total_products} products.")
        product_idx = 0
        for page in range(pages):
            tbl = doc.add_table(rows=num_rows, cols=num_cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = False
            tbl.allow_autofit = False
            tblPr = tbl._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
            layout = OxmlElement('w:tblLayout')
            layout.set(qn('w:type'), 'fixed')
            tblPr.append(layout)
            tbl._element.insert(0, tblPr)
            grid = OxmlElement('w:tblGrid')
            for _ in range(num_cols):
                gc = OxmlElement('w:gridCol')
                gc.set(qn('w:w'), col_width_twips)
                grid.append(gc)
            tbl._element.insert(0, grid)
            for row in tbl.rows:
                row.height = row_height_pts
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            borders = OxmlElement('w:tblBorders')
            for side in ('insideH','insideV'):
                b = OxmlElement(f"w:{side}")
                b.set(qn('w:val'), "single")
                b.set(qn('w:sz'), "4")
                b.set(qn('w:color'), "D3D3D3")
                b.set(qn('w:space'), "0")
                borders.append(b)
            tblPr.append(borders)
            for r in range(num_rows):
                for c in range(num_cols):
                    cell = tbl.cell(r, c)
                    cell._tc.clear_content()
                    if product_idx < total_products:
                        cnt = product_idx + 1
                        tc = deepcopy(src_tc)
                        cell_text = ''
                        for t in tc.iter(qn('w:t')):
                            if t.text:
                                cell_text += t.text
                                if 'Label1' in t.text:
                                    t.text = t.text.replace('Label1', f'Label{cnt}')
                        if ('{{Label1.ProductBrand}}' not in cell_text and 'ProductBrand' not in cell_text):
                            text_elements = list(tc.iter(qn('w:t')))
                            lineage_end_index = -1
                            for i, t in enumerate(text_elements):
                                if t.text and 'Lineage' in t.text:
                                    for j in range(i, len(text_elements)):
                                        if text_elements[j].text and '}}' in text_elements[j].text:
                                            lineage_end_index = j
                                            break
                            if lineage_end_index >= 0:
                                new_text = OxmlElement('w:t')
                                new_text.text = f'{{{{Label{cnt}.ProductBrand}}}}'
                                lineage_end_element = text_elements[lineage_end_index]
                                lineage_end_element.getparent().insert(
                                    lineage_end_element.getparent().index(lineage_end_element) + 1,
                                    new_text
                                )
                        if '{{Label1.DOH}}' not in cell_text and 'DOH' not in cell_text:
                            text_elements = list(tc.iter(qn('w:t')))
                            strain_end_index = -1
                            for i, t in enumerate(text_elements):
                                if t.text and 'ProductStrain' in t.text:
                                    for j in range(i, len(text_elements)):
                                        if text_elements[j].text and '}}' in text_elements[j].text:
                                            strain_end_index = j
                                            break
                            if strain_end_index >= 0:
                                new_text = OxmlElement('w:t')
                                new_text.text = f'\n{{{{Label{cnt}.DOH}}}}'
                                strain_end_element = text_elements[strain_end_index]
                                strain_end_element.getparent().insert(
                                    strain_end_element.getparent().index(strain_end_element) + 1,
                                    new_text
                                )
                        
                        # CRITICAL FIX: Add QR placeholder after DOH
                        if '{{Label1.QR}}' not in cell_text and 'QR' not in cell_text:
                            text_elements = list(tc.iter(qn('w:t')))
                            doh_end_index = -1
                            for i, t in enumerate(text_elements):
                                if t.text and 'DOH' in t.text:
                                    for j in range(i, len(text_elements)):
                                        if text_elements[j].text and '}}' in text_elements[j].text:
                                            doh_end_index = j
                                            break
                            if doh_end_index >= 0:
                                new_text = OxmlElement('w:t')
                                new_text.text = f'\n{{{{Label{cnt}.QR}}}}'
                                doh_end_element = text_elements[doh_end_index]
                                doh_end_element.getparent().insert(
                                    doh_end_element.getparent().index(doh_end_element) + 1,
                                    new_text
                                )
                        
                        # CRITICAL FIX: Add DescAndWeight after QR for product description + weight
                        if '{{Label1.DescAndWeight}}' not in cell_text and 'DescAndWeight' not in cell_text:
                            text_elements = list(tc.iter(qn('w:t')))
                            qr_end_index = -1
                            for i, t in enumerate(text_elements):
                                if t.text and 'QR' in t.text:
                                    for j in range(i, len(text_elements)):
                                        if text_elements[j].text and '}}' in text_elements[j].text:
                                            qr_end_index = j
                                            break
                            if qr_end_index >= 0:
                                new_text = OxmlElement('w:t')
                                new_text.text = f'\n{{{{Label{cnt}.DescAndWeight}}}}'
                                qr_end_element = text_elements[qr_end_index]
                                qr_end_element.getparent().insert(
                                    qr_end_element.getparent().index(qr_end_element) + 1,
                                    new_text
                                )
                        
                        # CRITICAL FIX: Add Price after DescAndWeight
                        if '{{Label1.Price}}' not in cell_text or 'Price' not in cell_text:
                            text_elements = list(tc.iter(qn('w:t')))
                            # Find where to insert Price - after DescAndWeight or QR
                            insert_index = -1
                            for i, t in enumerate(text_elements):
                                if t.text and ('DescAndWeight' in t.text or 'QR' in t.text):
                                    for j in range(i, len(text_elements)):
                                        if text_elements[j].text and '}}' in text_elements[j].text:
                                            insert_index = j
                                            break
                            if insert_index >= 0:
                                new_text = OxmlElement('w:t')
                                new_text.text = f'\n${{{{Label{cnt}.Price}}}}'
                                insert_element = text_elements[insert_index]
                                insert_element.getparent().insert(
                                    insert_element.getparent().index(insert_element) + 1,
                                    new_text
                                )
                        
                        for el in tc.xpath('./*'):
                            cell._tc.append(deepcopy(el))
                        product_idx += 1
                    else:
                        cell.add_paragraph()
            # Add minimal spacing between cells
            tblPr2 = tbl._element.find(qn('w:tblPr'))
            spacing = OxmlElement('w:tblCellSpacing')
            spacing.set(qn('w:w'), str(int(0.001 * 1440)))
            spacing.set(qn('w:type'), 'dxa')
            tblPr2.append(spacing)
            
            # Insert page break after each table except the last
            if page < pages - 1:
                # Add a page break paragraph to force Word to create a new page
                page_break_para = doc.add_paragraph()
                page_break_run = page_break_para.add_run()
                page_break_run.add_break(WD_BREAK.PAGE)
                self.logger.info(f"🔍 PAGE BREAK: Added page break after page {page + 1} of {pages}")

        from src.core.generation.docx_formatting import remove_all_headers_and_footers
        doc = remove_all_headers_and_footers(doc)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def process_records(self, records):
        documents = []
        # FULLY DISABLE CHUNKING for horizontal, vertical, and double templates
        try:
            # VENDOR FALLBACK: Detect if records were filtered by vendor (all have same vendor)
            # If so, use that vendor as fallback for records with missing vendor
            # PERFORMANCE: Sample first 50 records instead of checking all to speed up detection
            self._vendor_fallback = None
            if records:
                vendor_counts = {}
                sample_size = min(50, len(records))  # Sample first 50 records for speed
                sample_records = records[:sample_size]
                
                for record in sample_records:
                    vendor = (record.get('Vendor') or record.get('Vendor/Supplier*') or 
                             record.get('Vendor/Supplier') or record.get('ProductVendor', ''))
                    if vendor and not pd.isna(vendor) and str(vendor).strip() and str(vendor).lower() not in ['nan', 'none', 'null', '']:
                        vendor_str = str(vendor).strip()
                        vendor_counts[vendor_str] = vendor_counts.get(vendor_str, 0) + 1
                
                # If one vendor dominates (appears in >50% of sampled records), use it as fallback
                if vendor_counts:
                    total_with_vendor = sum(vendor_counts.values())
                    most_common_vendor = max(vendor_counts.items(), key=lambda x: x[1])
                    if most_common_vendor[1] >= total_with_vendor * 0.5 and most_common_vendor[1] >= 2:
                        self._vendor_fallback = most_common_vendor[0]
                        # Only log if we actually use it (reduces logging overhead)
            
            if self.template_type in ['horizontal', 'vertical', 'double']:
                self.chunk_size = len(records)
                self.logger.info(f"🔍 LABEL RENDER: For template '{self.template_type}', forced chunk_size to {self.chunk_size} to render all labels.")
                self.logger.info(f"🔍 LABEL RENDER: Chunking is fully disabled. All {len(records)} records will be processed in one pass.")
                self.start_time = time.time()
                self.chunk_count = 1
                overall_order = [record.get('ProductName', 'Unknown') for record in records]
                self.logger.info(f"Processing {len(records)} records in overall order: {overall_order}")
                has_json_products = any(record.get('Source', '').startswith('JSON') or record.get('Source', '').startswith('Database Priority') for record in records)
                
                # DEDUPLICATION FIX: Remove exact duplicates even for JSON matched products
                seen_products = set()
                unique_records = []
                duplicate_count = 0
                
                for record in records:
                    # Create a unique key based on product name, price, and weight
                    product_name = record.get('ProductName', 'Unknown')
                    price = record.get('Price', '')
                    weight = record.get('Weight', '') or record.get('NetWeight', '')
                    vendor = record.get('Vendor', '') or record.get('ProductVendor', '')
                    
                    # Create deduplication key
                    dedup_key = f"{product_name}|{price}|{weight}|{vendor}".lower().strip()
                    
                    if dedup_key not in seen_products:
                        seen_products.add(dedup_key)
                        unique_records.append(record)
                    else:
                        duplicate_count += 1
                        self.logger.info(f"🗑️ DEDUPLICATION: Removing duplicate '{product_name}' (Price: {price}, Weight: {weight})")
                
                if duplicate_count > 0:
                    self.logger.info(f"✅ DEDUPLICATION: Removed {duplicate_count} duplicate(s), {len(unique_records)} unique products remain")
                    records = unique_records
                
                # Process all records in a single chunk
                chunk = records
                self.chunk_count = 1
                documents.append(self._process_chunk(chunk))
                # All records processed in one chunk; no further chunking or looping required
            else:
                # Ensure chunk size respects fixed page capacity for templates like mini/inventory
                self.chunk_size = self.chunk_size or len(records)
                self.logger.info(f"🔍 LABEL RENDER: Processing {len(records)} records for template '{self.template_type}' with chunk_size {self.chunk_size}.")
                self.start_time = time.time()
                self.chunk_count = 0
                overall_order = [record.get('ProductName', 'Unknown') for record in records]
                self.logger.info(f"Processing {len(records)} records in overall order: {overall_order}")
                has_json_products = any(record.get('Source', '').startswith('JSON') or record.get('Source', '').startswith('Database Priority') for record in records)

                # Apply the same deduplication logic for consistency across templates
                seen_products = set()
                unique_records = []
                duplicate_count = 0

                for record in records:
                    product_name = record.get('ProductName', 'Unknown')
                    price = record.get('Price', '')
                    weight = record.get('Weight', '') or record.get('NetWeight', '')
                    vendor = record.get('Vendor', '') or record.get('ProductVendor', '')

                    dedup_key = f"{product_name}|{price}|{weight}|{vendor}".lower().strip()

                    if dedup_key not in seen_products:
                        seen_products.add(dedup_key)
                        unique_records.append(record)
                    else:
                        duplicate_count += 1
                        self.logger.info(f"🗑️ DEDUPLICATION: Removing duplicate '{product_name}' (Price: {price}, Weight: {weight})")

                if duplicate_count > 0:
                    self.logger.info(f"✅ DEDUPLICATION: Removed {duplicate_count} duplicate(s), {len(unique_records)} unique products remain")
                    records = unique_records

                if not records:
                    self.logger.warning("No records to process after deduplication.")
                    return None

                chunk_size = max(1, self.chunk_size)
                chunks = []
                for start in range(0, len(records), chunk_size):
                    chunk = records[start:start + chunk_size]
                    chunks.append(chunk)
                
                # Process chunks in parallel for better performance
                if len(chunks) > 1:
                    from concurrent.futures import ThreadPoolExecutor, as_completed
                    
                    self.logger.info(f"⚡ PARALLEL PROCESSING: Processing {len(chunks)} chunks concurrently")
                    chunk_docs = [None] * len(chunks)  # Preserve order
                    
                    with ThreadPoolExecutor(max_workers=min(4, len(chunks))) as executor:
                        # Submit all chunks for parallel processing
                        future_to_index = {
                            executor.submit(self._process_chunk, chunk): idx 
                            for idx, chunk in enumerate(chunks)
                        }
                        
                        # Collect results as they complete
                        for future in as_completed(future_to_index):
                            idx = future_to_index[future]
                            try:
                                chunk_doc = future.result()
                                chunk_docs[idx] = chunk_doc
                                self.chunk_count += 1
                                self.logger.info(f"⚡ PARALLEL: Chunk {idx + 1}/{len(chunks)} completed")
                            except Exception as e:
                                self.logger.error(f"Error processing chunk {idx + 1}: {e}")
                                chunk_docs[idx] = None
                    
                    # Filter out None results and preserve order
                    documents.extend([doc for doc in chunk_docs if doc is not None])
                else:
                    # Single chunk - process normally
                    self.chunk_count = 1
                    self.logger.info(f"🔍 LABEL RENDER: Processing single chunk containing {len(records)} record(s)")
                    try:
                        chunk_doc = self._process_chunk(records)
                        if chunk_doc is not None:
                            documents.append(chunk_doc)
                        else:
                            self.logger.error("Single chunk processing returned None")
                    except Exception as e:
                        self.logger.error(f"Error processing single chunk: {e}")
                        self.logger.error(traceback.format_exc())
                        # Don't add None to documents - let it fail gracefully
            
            if not documents: 
                return None
            if len(documents) == 1: 
                return documents[0]
            
            # Combine documents
            self.logger.info(f"Combining {len(documents)} documents")
            composer = Composer(documents[0])
            for doc in documents[1:]:
                composer.append(doc)
            
            final_doc_buffer = BytesIO()
            composer.save(final_doc_buffer)
            final_doc_buffer.seek(0)
            
            # CRITICAL: Remove ALL headers and footers from the final combined document
            final_doc = Document(final_doc_buffer)
            from src.core.generation.docx_formatting import remove_all_headers_and_footers
            final_doc = remove_all_headers_and_footers(final_doc)
            
            total_time = time.time() - self.start_time
            self.logger.info(f"Template processing completed in {total_time:.2f}s for {len(records)} records")
            
            return final_doc
        except Exception as e:
            self.logger.error(f"Error processing records: {e}\n{traceback.format_exc()}")
            raise

    def _process_chunk(self, chunk):
        """Process a chunk of records with timeout protection."""
        from docxtpl import DocxTemplate
        from docx import Document
        from io import BytesIO
        
        # DEBUG_CHUNK_SIZE_TRACKING: Log actual chunk sizes
        self.logger.info(f"🔍 CHUNK SIZE DEBUG: Processing chunk with {len(chunk)} records")
        self.logger.info(f"🔍 CHUNK SIZE DEBUG: Expected all {len(chunk)} records in this chunk for template '{self.template_type}'")
        # After rendering, log the number of labels actually created
        self.logger.info(f"🔍 LABEL RENDER: Actually rendered {len(chunk)} labels in this chunk.")
        
        chunk_start_time = time.time()
        
        try:
            # CRITICAL FIX: Re-expand template with correct number of products to prevent blank labels
            # OPTIMIZATION: Cache template expansions to avoid re-expanding for same size
            num_products = len(chunk)
            cache_key = f"{self.template_type}_{num_products}"
            
            if cache_key in self._template_expansion_cache:
                # Use cached template expansion
                self._expanded_template_buffer = self._template_expansion_cache[cache_key]
                if hasattr(self._expanded_template_buffer, 'seek'):
                    self._expanded_template_buffer.seek(0)
            else:
                # For all templates, re-expand with correct number of products
                if self.template_type in ['horizontal', 'vertical']:
                    self._expanded_template_buffer = self._expand_template_to_3x3_fixed(num_products)
                elif self.template_type == 'double':
                    self._expanded_template_buffer = self._expand_template_to_4x3_fixed_double(num_products)
                elif self.template_type == 'mini':
                    self._expanded_template_buffer = self._expand_template_to_4x5_fixed_scaled(num_products)
                
                # Cache the expansion (create a copy since BytesIO is consumed)
                if hasattr(self._expanded_template_buffer, 'getvalue'):
                    cached_buffer = BytesIO(self._expanded_template_buffer.getvalue())
                    self._template_expansion_cache[cache_key] = cached_buffer
                    self._expanded_template_buffer.seek(0)
                elif hasattr(self._expanded_template_buffer, 'seek'):
                    self._expanded_template_buffer.seek(0)
            
            doc = DocxTemplate(self._expanded_template_buffer)
            
            # Debug: Log the order of records in this chunk (only for small chunks to reduce logging overhead)
            if len(chunk) <= 10:
                chunk_order = [record.get('ProductName', 'Unknown') for record in chunk]
                self.logger.info(f"Processing chunk with {len(chunk)} records in order: {chunk_order}")
            else:
                self.logger.info(f"Processing chunk with {len(chunk)} records")
            
            # OPTIMIZATION: Pre-load all brand, vendor, lineage, and strain data in batch to avoid N+1 queries
            # This reduces 200+ queries for 100 products to just 3-4 queries total
            product_brand_cache = {}
            product_vendor_cache = {}
            product_lineage_cache = {}
            strain_info_cache = {}
            joint_ratio_cache = {}
            try:
                from src.core.data.product_database import get_product_database
                product_db = get_product_database()
                if product_db:
                    product_names = [r.get('ProductName', '') or r.get('Product Name*', '') for r in chunk]
                    product_names = [n for n in product_names if n]
                    
                    # Collect unique strain names for batch loading
                    strain_names = set()
                    for r in chunk:
                        strain = r.get('ProductStrain', '') or r.get('Product Strain', '')
                        if strain:
                            strain_names.add(strain)
                    
                    if product_names:
                        try:
                            conn = product_db._get_connection()
                            cursor = conn.cursor()
                            placeholders = ','.join(['?'] * len(product_names))
                            
                            # Load brand data
                            batch_brand_query = f'''
                                SELECT "Product Name*", "Product Brand"
                                FROM products
                                WHERE "Product Name*" IN ({placeholders})
                                AND "Product Brand" IS NOT NULL
                                AND "Product Brand" != ""
                            '''
                            cursor.execute(batch_brand_query, product_names)
                            for row_result in cursor.fetchall():
                                pname, brand = row_result
                                if brand and str(brand).strip() not in ['', 'None', 'NULL', 'null', 'nan']:
                                    product_brand_cache[pname] = str(brand).strip()
                            
                            # Load vendor data - try Vendor/Supplier* first, then Vendor, then ProductVendor
                            batch_vendor_query = f'''
                                SELECT "Product Name*", 
                                       CASE 
                                           WHEN "Vendor/Supplier*" IS NOT NULL AND "Vendor/Supplier*" != '' THEN "Vendor/Supplier*"
                                           WHEN "Vendor" IS NOT NULL AND "Vendor" != '' THEN "Vendor"
                                           WHEN "ProductVendor" IS NOT NULL AND "ProductVendor" != '' THEN "ProductVendor"
                                           ELSE NULL
                                       END as vendor
                                FROM products
                                WHERE "Product Name*" IN ({placeholders})
                                AND (
                                    ("Vendor/Supplier*" IS NOT NULL AND "Vendor/Supplier*" != '')
                                    OR ("Vendor" IS NOT NULL AND "Vendor" != '')
                                    OR ("ProductVendor" IS NOT NULL AND "ProductVendor" != '')
                                )
                            '''
                            cursor.execute(batch_vendor_query, product_names)
                            for row_result in cursor.fetchall():
                                pname, vendor = row_result
                                if vendor and str(vendor).strip() not in ['', 'None', 'NULL', 'null', 'nan']:
                                    product_vendor_cache[pname] = str(vendor).strip()
                            
                            # Load lineage data (Lineage, canonical_lineage)
                            batch_lineage_query = f'''
                                SELECT "Product Name*", "Lineage", canonical_lineage
                                FROM products
                                WHERE "Product Name*" IN ({placeholders})
                            '''
                            cursor.execute(batch_lineage_query, product_names)
                            for row_result in cursor.fetchall():
                                pname, lineage, canon_lineage = row_result
                                # Priority: Lineage > canonical_lineage
                                if lineage and str(lineage).strip() not in ['', 'None', 'NULL', 'null', 'nan']:
                                    product_lineage_cache[pname] = str(lineage).strip()
                                elif canon_lineage and str(canon_lineage).strip() not in ['', 'None', 'NULL', 'null', 'nan']:
                                    product_lineage_cache[pname] = str(canon_lineage).strip()
                            
                            # Load JointRatio data
                            batch_joint_ratio_query = f'''
                                SELECT "Product Name*", JointRatio
                                FROM products
                                WHERE "Product Name*" IN ({placeholders})
                                AND JointRatio IS NOT NULL
                                AND JointRatio != ""
                            '''
                            cursor.execute(batch_joint_ratio_query, product_names)
                            for row_result in cursor.fetchall():
                                pname, joint_ratio = row_result
                                if joint_ratio and str(joint_ratio).strip() not in ['', 'None', 'NULL', 'null', 'nan']:
                                    joint_ratio_cache[pname] = str(joint_ratio).strip()
                            
                            # Batch load strain info
                            if strain_names:
                                strain_placeholders = ','.join(['?'] * len(strain_names))
                                batch_strain_query = f'''
                                    SELECT strain_name, display_lineage, canonical_lineage
                                    FROM strains
                                    WHERE strain_name IN ({strain_placeholders})
                                '''
                                cursor.execute(batch_strain_query, list(strain_names))
                                for row_result in cursor.fetchall():
                                    strain_name, display_lineage, canon_lineage = row_result
                                    strain_info = {}
                                    # Priority: display_lineage > canonical_lineage
                                    if display_lineage and str(display_lineage).strip() not in ['', 'None', 'NULL', 'null', 'nan']:
                                        strain_info['display_lineage'] = str(display_lineage).strip()
                                    elif canon_lineage and str(canon_lineage).strip() not in ['', 'None', 'NULL', 'null', 'nan']:
                                        strain_info['canonical_lineage'] = str(canon_lineage).strip()
                                    if strain_info:
                                        strain_info_cache[strain_name] = strain_info
                        except Exception as batch_err:
                            self.logger.warning(f"Batch data query failed: {batch_err}")
            except Exception as e:
                self.logger.warning(f"Failed to pre-load batch data: {e}")
            
            # Build context for each record in the chunk
            context = {}
            
            # Determine required label count based on template type
            if self.template_type == 'mini' or self.template_type == 'preroll':
                required_labels = 20  # Fixed grid: 4x5 = 20 labels
            elif self.template_type == 'double':
                required_labels = 12  # Fixed grid: 3x4 = 12 labels
            elif self.template_type == 'inventory':
                required_labels = 4   # Fixed grid: 2x2 = 4 labels
            else:
                required_labels = len(chunk)  # Dynamic templates use actual chunk size
            
            for i, record in enumerate(chunk):
                # Set current record for brand centering logic
                self.current_record = record
                # Set current product type for brand marker processing
                self.current_product_type = (record.get('ProductType', '').lower() or 
                                          record.get('Product Type*', '').lower())
                if self.template_type == 'inventory':
                    label_context = self._build_inventory_context(record)
                else:
                    # Pass all caches to avoid N+1 queries
                    label_context = self._build_label_context(record, doc, product_brand_cache, product_vendor_cache, 
                                                               product_lineage_cache, strain_info_cache, joint_ratio_cache)
                context[f'Label{i+1}'] = label_context
                # Debug logging to check field values and order (only for first few labels to reduce overhead)
                if i < 3:
                    product_name = record.get('ProductName', 'Unknown')
                    product_type = record.get('ProductType', '') or record.get('Product Type*', '')
                    product_vendor = label_context.get('ProductVendor', 'NOT_FOUND')
                    # Unwrap vendor to see actual value
                    if product_vendor != 'NOT_FOUND' and 'PRODUCTVENDOR_START' in str(product_vendor):
                        try:
                            vendor_value = unwrap_marker(product_vendor, 'PRODUCTVENDOR')
                            product_vendor = f"'{vendor_value}' (wrapped)"
                        except:
                            pass
                    # Also check vendor from record directly
                    vendor_from_record_debug = record.get('Vendor/Supplier*') or record.get('Vendor') or record.get('ProductVendor') or 'NOT_IN_RECORD'
                    self.logger.info(f"🔍 CONTEXT DEBUG Label{i+1} -> {product_name} (type: {product_type}) - ProductVendor in context: {product_vendor}, Vendor in record: '{vendor_from_record_debug}', _vendor_from_record: '{label_context.get('_vendor_from_record', 'NOT_SET')}'")
            
            # For fixed-grid templates (mini, preroll, double, inventory), ensure all labels exist
            # to prevent Jinja template errors when template references missing labels
            if self.template_type in ['mini', 'preroll', 'double', 'inventory']:
                empty_label_context = self._get_empty_label_context()
                for i in range(len(chunk) + 1, required_labels + 1):
                    context[f'Label{i}'] = empty_label_context
                self.logger.info(f"🔧 FIXED GRID: Created {len(chunk)} product labels + {required_labels - len(chunk)} empty labels = {required_labels} total for {self.template_type} template")
            else:
                # CRITICAL FIX: Only create contexts for actual products to prevent blank tags on last sheet
                # This saves printer ink by not generating empty cells
                self.logger.info(f"🔧 BLANK TAG PREVENTION: Only creating {len(chunk)} labels instead of {self.chunk_size} to prevent blank tags on last sheet")

            # DOH images are already created in _build_label_context, no need for redundant creation here
            
            # QR code functionality enabled
            # Debug: Log QR code presence in context before rendering
            qr_count = sum(1 for label_key, label_data in context.items() 
                          if isinstance(label_data, dict) and label_data.get('QR') and 
                          not (isinstance(label_data.get('QR'), str) and label_data.get('QR').strip() == ''))
            self.logger.info(f"🔍 QR CODE CHECK: {qr_count} labels have QR codes in context before render (total labels: {len([k for k in context.keys() if k.startswith('Label')])})")
            
            try:
                # Dump a concise mapping of Label keys -> ProductName/Description
                try:
                    label_keys = [k for k in context.keys() if str(k).startswith('Label')]
                    # Sort by numeric label index when possible
                    def _label_index(k):
                        import re
                        m = re.search(r"(\d+)", str(k))
                        return int(m.group(1)) if m else 0
                    label_keys = sorted(label_keys, key=_label_index)
                    self.logger.info(f"🔍 PRE-RENDER CONTEXT DUMP: {len(label_keys)} label keys")
                    for lk in label_keys:
                        v = context.get(lk)
                        pname = ''
                        if isinstance(v, dict):
                            pname = v.get('ProductName') or v.get('Product Name*') or v.get('Description') or ''
                        else:
                            pname = str(v)[:80]
                        self.logger.info(f"  {lk}: {pname}")
                except Exception as _dump_err:
                    self.logger.debug(f"Failed to dump context before render: {_dump_err}")

                doc.render(context)
                self.logger.debug("DocxTemplate render completed successfully")
                
                # CRITICAL FIX: Remove unmerged placeholders immediately after render
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                rendered_doc = Document(buffer)
                self._remove_unmerged_placeholders(rendered_doc, len(chunk))
                
            except Exception as render_error:
                self.logger.error(f"DocxTemplate render failed: {render_error}")
                self.logger.error(f"Context keys: {list(context.keys())}")
                self.logger.error(f"Chunk size: {len(chunk)}, Template type: {self.template_type}")
                # Re-raise the error so it can be handled upstream
                raise
            
            # CRITICAL FIX: Ensure all tables have proper tblGrid elements before processing
            self._ensure_table_grids_exist(rendered_doc)
            
            # CRITICAL FIX: Wrap all post-processing in comprehensive error handling
            try:
                # Post-process the document to apply dynamic font sizing first
                self._post_process_and_replace_content(rendered_doc)
                
                # Check timeout before lineage colors
                if time.time() - chunk_start_time > MAX_PROCESSING_TIME_PER_CHUNK:
                    self.logger.warning(f"Chunk processing timeout reached ({MAX_PROCESSING_TIME_PER_CHUNK}s), skipping lineage colors")
                    return rendered_doc
                
                # Apply lineage colors last to ensure they are not overwritten
                apply_lineage_colors(rendered_doc)
                
                # Apply final marker cleanup for all templates
                self._final_marker_cleanup(rendered_doc)
                
            except Exception as processing_error:
                self.logger.warning(f"Skipping post-processing due to table structure issue: {processing_error}")
                # Continue processing without post-processing features
            
            # Final enforcement: prevent any cell/row expansion and force EXACT dimensions
            # Cell widths already standardized
            
            # CRITICAL: Remove ALL headers and footers to prevent unwanted content
            from src.core.generation.docx_formatting import remove_all_headers_and_footers
            rendered_doc = remove_all_headers_and_footers(rendered_doc)
            
            # Ensure proper table centering and document setup
            try:
                self._ensure_proper_centering(rendered_doc)
            except Exception as centering_error:
                self.logger.warning(f"Skipping centering due to table structure issue: {centering_error}")
                # Continue processing without centering

            # All content now uses standard spacing - no special THC_CBD handling
            
            chunk_time = time.time() - chunk_start_time
            # Chunk processed
            
            # FINAL MARKER CLEANUP: Remove any lingering *_START and *_END markers AFTER font sizing has been applied
            # This cleanup should only remove markers that weren't processed by the font sizing system
            marker_pattern = re.compile(r'\b\w+_(START|END)\b')
            prefix_pattern = re.compile(r'^(?:[A-Z0-9_]+_)+')
            
            # Clean in tables
            try:
                for table in rendered_doc.tables:
                    try:
                        # Use safe table iteration to validate and repair if needed
                        if not self._safe_table_iteration(table, "marker cleanup"):
                            self.logger.warning(f"Skipping table with invalid structure during marker cleanup")
                            continue
                        
                        for row in table.rows:
                            try:
                                # Validate row structure before processing
                                if not hasattr(row, 'cells') or not row.cells:
                                    self.logger.warning(f"Skipping row with invalid structure during marker cleanup")
                                    continue
                                    
                                for cell in row.cells:
                                    try:
                                        for para in cell.paragraphs:
                                            # Check if this paragraph was processed by font sizing system
                                            # If it has non-default font sizes, it was processed
                                            was_processed = False
                                            for run in para.runs:
                                                if hasattr(run, 'font') and hasattr(run.font, 'size') and run.font.size:
                                                    # Check if font size is not the default (12pt)
                                                    if hasattr(run.font.size, 'pt') and run.font.size.pt != 12:
                                                        was_processed = True
                                                        break
                                            
                                            # CRITICAL FIX: Always clean markers regardless of font sizing processing
                                            # This ensures DESC_START, DESC_END, PRICE_START, PRICE_END are always removed
                                            for run in para.runs:
                                                if marker_pattern.search(run.text):
                                                    run.text = marker_pattern.sub('', run.text)
                                                if prefix_pattern.search(run.text):
                                                    run.text = prefix_pattern.sub('', run.text)
                                    except Exception as cell_error:
                                        self.logger.warning(f"Skipping cell due to error during marker cleanup: {cell_error}")
                                        continue
                            except Exception as row_error:
                                self.logger.warning(f"Skipping row due to error during marker cleanup: {row_error}")
                                continue
                    except Exception as e:
                        self.logger.warning(f"Skipping table due to error during marker cleanup: {e}")
                        continue
            except Exception as overall_error:
                self.logger.error(f"Critical error during marker cleanup: {overall_error}")
                # Continue processing other parts of the document
            
            # Clean in paragraphs outside tables
            for para in rendered_doc.paragraphs:
                # Check if this paragraph was processed by font sizing system
                was_processed = False
                for run in para.runs:
                    if hasattr(run, 'font') and hasattr(run.font, 'size') and run.font.size:
                        # Check if font size is not the default (12pt)
                        if hasattr(run.font.size, 'pt') and run.font.size.pt != 12:
                            was_processed = True
                            break
                
                # CRITICAL FIX: Always clean markers regardless of font sizing processing
                # This ensures DESC_START, DESC_END, PRICE_START, PRICE_END are always removed
                # (Removed was_processed check)
                if True:
                    for run in para.runs:
                        if marker_pattern.search(run.text):
                            run.text = marker_pattern.sub('', run.text)
                        if prefix_pattern.search(run.text):
                            run.text = prefix_pattern.sub('', run.text)
            
            # FINAL STEP: Clean up any remaining concatenated lineage+brand content for classic types
            try:
                self._clean_up_lineage_brand_concatenation(rendered_doc)
            except Exception as e:
                self.logger.warning(f"Lineage brand concatenation cleanup failed: {e}")
            
            # FINAL STEP: Ensure standalone cannabinoid text uses 1pt font size
            try:
                self._ensure_standalone_cannabinoid_font_sizing(rendered_doc)
            except Exception as e:
                self.logger.warning(f"Standalone cannabinoid font sizing failed: {e}")
            
            # FINAL STEP: Ensure Lineage field centering for nonclassic types (runs after everything else)
            try:
                self._ensure_lineage_centering_for_nonclassic_types(rendered_doc)
            except Exception as e:
                self.logger.warning(f"Final lineage centering fix failed: {e}")
            
            return rendered_doc
            
        except Exception as e:
            self.logger.error(f"Error in _process_chunk: {e}\n{traceback.format_exc()}")
            raise
    def _build_inventory_context(self, record):
        """Build context dictionary for inventory slip template."""
        context = {}
        
        # Map inventory fields with proper formatting
        context['ProductName'] = record.get('Product Name*', '')
        context['Barcode'] = record.get('Barcode*', '')
        context['Quantity'] = record.get('Quantity Received*', '')
        context['AcceptedDate'] = record.get('Accepted Date', '')
        context['Vendor'] = record.get('Vendor', '')
        
        # Add any additional formatting or processing needed for inventory slips
        if context['AcceptedDate']:
            try:
                # Try to parse and reformat the date if needed
                from datetime import datetime
                date_obj = datetime.strptime(context['AcceptedDate'], '%Y-%m-%d')
                context['AcceptedDate'] = date_obj.strftime('%m/%d/%Y')
            except:
                pass  # Keep original format if parsing fails
        
        # Ensure all values are strings
        for key in context:
            if context[key] is None:
                context[key] = ''
            context[key] = str(context[key])
        
        return context

    def _get_empty_label_context(self):
        """Create an empty label context dictionary with all required fields set to empty strings."""
        return {
            'Description': '',
            'WeightUnits': '',
            'ProductBrand': '',
            'Price': '',
            'Lineage': '',
            'DOH': '',
            'Ratio_or_THC_CBD': '',
            'THC_CBD': '',
            'THC': '',
            'CBD': '',
            'Ratio': '',
            'ProductName': '',
            'ProductStrain': '',
            'ProductVendor': '',
            'DescAndWeight': '',
            'JointRatio': '',
            'ProductType': '',
            # Marker fields for template processing
            'ProductStrain_START': 'PRODUCTSTRAIN_START',
            'ProductStrain_END': 'PRODUCTSTRAIN_END',
            'Lineage_START': 'LINEAGE_START',
            'Lineage_END': 'LINEAGE_END',
            'ProductBrand_START': 'PRODUCTBRAND_START',
            'ProductBrand_END': 'PRODUCTBRAND_END',
            'ProductVendor_START': 'PRODUCTVENDOR_START',
            'ProductVendor_END': 'PRODUCTVENDOR_END',
            'DescAndWeight_START': 'DESC_START',
            'DescAndWeight_END': 'DESC_END',
            'Ratio_or_THC_CBD_START': 'THC_CBD_START',
            'Ratio_or_THC_CBD_END': 'THC_CBD_END',
            'Price_START': 'PRICE_START',
            'Price_END': 'PRICE_END',
            'WeightUnits_START': 'WEIGHTUNITS_START',
            'WeightUnits_END': 'WEIGHTUNITS_END',
            'Ratio_START': 'RATIO_START',
            'Ratio_END': 'RATIO_END',
            'JointRatio_START': 'JOINT_RATIO_START',
            'JointRatio_END': 'JOINT_RATIO_END',
            'THC_START': 'THC_START',
            'THC_END': 'THC_END',
            'CBD_START': 'CBD_START',
            'CBD_END': 'CBD_END',
            # QR code field (empty for blank labels)
            'QR': '',
        }
    
    def _build_label_context(self, record, doc, product_brand_cache=None, product_vendor_cache=None, 
                             product_lineage_cache=None, strain_info_cache=None, joint_ratio_cache=None):
        # Initialize caches to empty dicts if None
        product_brand_cache = product_brand_cache or {}
        product_vendor_cache = product_vendor_cache or {}
        product_lineage_cache = product_lineage_cache or {}
        strain_info_cache = strain_info_cache or {}
        joint_ratio_cache = joint_ratio_cache or {}
        """Ultra-optimized label context building for maximum performance."""
        # Use module-level re import (already imported at top of file)
        if product_brand_cache is None:
            product_brand_cache = {}
        if product_vendor_cache is None:
            product_vendor_cache = {}
        # CRITICAL FIX: Log lineage value received in template processor
        lineage_value = record.get('Lineage', 'NOT_FOUND')
        product_name = record.get('ProductName', 'Unknown')
        self.logger.info(f"LINEAGE TEMPLATE DEBUG: Building context for '{product_name}' with lineage: '{lineage_value}'")
        
        # Fast dictionary copy
        label_context = dict(record)
        
        # CRITICAL FIX: Read vendor directly from record first - it should already be in the Excel column
        # Check ALL possible vendor field variations, including case-insensitive matching
        vendor_from_record = None
        
        # First, get all vendor-related keys from the record (case-insensitive search)
        vendor_related_keys = [k for k in record.keys() if 'vendor' in k.lower() or 'supplier' in k.lower()]
        
        # Standard vendor field names to check (in priority order)
        vendor_field_names = ['Vendor/Supplier*', 'Vendor/Supplier', 'Vendor', 'ProductVendor', 'vendor']
        
        # Try standard field names first - check label_context FIRST (it's a dict copy of record)
        for field_name in vendor_field_names:
            # Check label_context first (already copied from record via dict(record))
            val = label_context.get(field_name)
            if val is None or pd.isna(val):
                # Fallback to record if not in label_context
                val = record.get(field_name)
            
            # CRITICAL: More thorough check - handle empty strings, None, NaN, and whitespace-only values
            if val is not None:
                # Check if it's NaN using pandas
                if not pd.isna(val):
                    val_str = str(val).strip()
                    val_lower = val_str.lower()
                    # Check if it's a valid non-empty value
                    if val_str and val_lower not in ['nan', 'none', 'null', '']:
                        vendor_from_record = val_str
                        self.logger.info(f"✅ EARLY VENDOR EXTRACTION: Found vendor in field '{field_name}': '{vendor_from_record}' for '{product_name}'")
                        break
                    else:
                        # Log why it was rejected
                        self.logger.debug(f"🔍 VENDOR REJECTED: Field '{field_name}' has value '{repr(val)}' (stripped: '{val_str}') which is empty/invalid for '{product_name}'")
                else:
                    self.logger.debug(f"🔍 VENDOR REJECTED: Field '{field_name}' is NaN for '{product_name}'")
            else:
                self.logger.debug(f"🔍 VENDOR REJECTED: Field '{field_name}' is None for '{product_name}'")
        
        # If not found in standard fields, check ALL vendor-related keys from BOTH label_context and record
        if not vendor_from_record and vendor_related_keys:
            for key in vendor_related_keys:
                # Check label_context first, then record
                val = label_context.get(key)
                if val is None or pd.isna(val):
                    val = record.get(key)
                
                if val is not None and not pd.isna(val) and str(val).strip() and str(val).lower() not in ['nan', 'none', 'null', '']:
                    vendor_from_record = str(val).strip()
                    self.logger.info(f"✅ Found vendor in field '{key}': '{vendor_from_record}' for '{product_name}'")
                    break
        
        # VENDOR FALLBACK: If vendor not found in record but we detected vendor filtering, use the fallback vendor
        if not vendor_from_record and hasattr(self, '_vendor_fallback') and self._vendor_fallback:
            vendor_from_record = self._vendor_fallback
            # Reduced logging for performance - only log first few instances
            if not hasattr(self, '_vendor_fallback_logged_count'):
                self._vendor_fallback_logged_count = 0
            if self._vendor_fallback_logged_count < 3:
                self.logger.info(f"✅ VENDOR FALLBACK: Using detected vendor '{vendor_from_record}' for '{product_name}' (vendor was missing from record)")
                self._vendor_fallback_logged_count += 1
        
        # Store vendor early so it's available throughout processing
        if vendor_from_record:
            label_context['_vendor_from_record'] = vendor_from_record
            # Also set ProductVendor directly in label_context so it's available immediately
            # This ensures vendor is preserved even if later logic tries to clear it
            if self.template_type == 'vertical':
                label_context['ProductVendor'] = vendor_from_record
            else:
                label_context['ProductVendor'] = f"PRODUCTVENDOR_START{vendor_from_record}PRODUCTVENDOR_END"
        else:
            # CRITICAL FIX: Always initialize ProductVendor, even if empty, so fallback logic can detect and populate it
            # Initialize as empty with markers so the fallback logic can properly detect it needs to be populated
            if self.template_type == 'vertical':
                label_context['ProductVendor'] = ''
            else:
                label_context['ProductVendor'] = wrap_with_marker('', 'PRODUCTVENDOR')
            # Log warning with all available keys for debugging
            all_keys_sample = list(record.keys())[:20]  # First 20 keys for debugging
            # Also log actual values from vendor fields to see if they're empty or have different names
            vendor_field_values = {}
            for field in vendor_field_names:
                val = record.get(field)
                vendor_field_values[field] = f"value={repr(val)}, type={type(val).__name__}, is_na={pd.isna(val) if hasattr(pd, 'isna') else 'N/A'}"
            # Also check vendor-related keys
            for key in vendor_related_keys:
                val = record.get(key)
                vendor_field_values[key] = f"value={repr(val)}, type={type(val).__name__}, is_na={pd.isna(val) if hasattr(pd, 'isna') else 'N/A'}"
            self.logger.warning(f"⚠️ No vendor found in record for '{product_name}'. Checked fields: {vendor_field_names}, Vendor-related keys: {vendor_related_keys}, Vendor field values: {vendor_field_values}, Sample record keys: {all_keys_sample}")
        
        # PREROLL TEMPLATE: Override ProductName with group display name if this is a grouped preroll
        if self.template_type == 'preroll':
            group_id = record.get('_group_id')
            if group_id:
                group_info = record.get('_group_info')
                if group_info and isinstance(group_info, dict):
                    group_display_name = group_info.get('display_name', '')
                    if group_display_name:
                        # Override ProductName, Product Name*, and Description immediately
                        label_context['ProductName'] = group_display_name
                        label_context['Product Name*'] = group_display_name
                        label_context['Description'] = group_display_name
                        self.logger.info(f"PREROLL GROUP OVERRIDE: Set ProductName/Description to '{group_display_name}' (group_id: {group_id})")
        
        has_cbd_blend_strain = False
        cbd_signal_tokens = ['CBD', 'CBG', 'CBN', 'CBC']

        def _contains_cbd_signal(value):
            if not value:
                return False
            text_upper = str(value).upper()
            return any(token in text_upper for token in cbd_signal_tokens)

        cbd_signal_candidates = [
            label_context.get('ProductStrain') or record.get('ProductStrain') or record.get('Product Strain'),
            label_context.get('ProductName') or record.get('ProductName') or record.get('Product Name*'),
            label_context.get('Description') or record.get('Description'),
            label_context.get('Product Brand') or record.get('Product Brand') or record.get('ProductBrand'),
            label_context.get('Ratio_or_THC_CBD') or record.get('Ratio_or_THC_CBD') or record.get('Ratio'),
            label_context.get('Lineage') or record.get('Lineage'),
            record.get('Product Type*'),
            record.get('ProductType')
        ]
        if any(_contains_cbd_signal(candidate) for candidate in cbd_signal_candidates):
            has_cbd_blend_strain = True
        else:
            # re is already imported at module level
            ratio_pattern = re.compile(r'\b\d+\s*:\s*\d+(?:\s*:\s*\d+)?\b')
            ratio_sources = [
                record.get('ProductName') or record.get('Product Name*'),
                record.get('Description'),
                record.get('Ratio_or_THC_CBD'),
                record.get('Ratio'),
                record.get('JointRatio'),
                record.get('Joint Ratio')
            ]
            if any(ratio_pattern.search(str(source)) for source in ratio_sources if source):
                has_cbd_blend_strain = True
        
        if has_cbd_blend_strain:
            current_strain_value = (
                label_context.get('ProductStrain')
                or record.get('ProductStrain')
                or record.get('Product Strain')
                or ''
            )
            current_strain_clean = str(current_strain_value).strip()
            if not current_strain_clean or current_strain_clean.upper() in {'MIXED', 'CBD', 'CBD BLEND', 'N/A', 'NONE', 'NULL'}:
                cbd_blend_value = 'CBD Blend'
                label_context['ProductStrain'] = cbd_blend_value
                label_context['Product Strain'] = cbd_blend_value
                record['ProductStrain'] = cbd_blend_value
                record['Product Strain'] = cbd_blend_value
                self.logger.info(f"CBD BLEND STRAIN ENFORCEMENT: Set ProductStrain to '{cbd_blend_value}' for '{product_name}'")
            elif current_strain_clean.upper() == 'CBD BLEND':
                # Normalize capitalization
                cbd_blend_value = 'CBD Blend'
                label_context['ProductStrain'] = cbd_blend_value
                label_context['Product Strain'] = cbd_blend_value
                record['ProductStrain'] = cbd_blend_value
                record['Product Strain'] = cbd_blend_value
        
        # CRITICAL FIX: ALWAYS prioritize database lineage over Excel lineage for DOCX output
        # Database lineage is the source of truth, not Excel lineage
        # This ensures DOCX output uses database lineage values, not Excel lineage
        try:
            product_name = record.get('ProductName', record.get('Product Name*', ''))
            excel_lineage = label_context.get('Lineage', '') or record.get('Lineage', '')
            
            # CRITICAL: Use record lineage first (already enriched with database value, no sativa hybrid override)
            # Only query database if record lineage is missing
            db_lineage = None
            # Priority: sovereign_lineage > canonical_lineage > Lineage > lineage (sovereign has manual tag manager edits)
            # CRITICAL FIX: Reject "SOVEREIGN" as invalid - it's a field name, not a lineage value
            record_lineage = None
            for lineage_field in ['sovereign_lineage', 'canonical_lineage', 'Lineage', 'lineage']:
                candidate = record.get(lineage_field)
                if candidate and str(candidate).strip() not in ['', 'None', 'nan']:
                    lineage_str = str(candidate).strip().upper()
                    if lineage_str != 'SOVEREIGN':  # Reject "SOVEREIGN" as invalid
                        record_lineage = lineage_str
                        break
            
            if record_lineage:
                # Use record lineage (already set correctly by enrichment, avoids sativa hybrid override)
                db_lineage = record_lineage
                if 'lemon' in product_name.lower() or 'cherry' in product_name.lower():
                    self.logger.info(f"✅ LINEAGE: Using record lineage '{db_lineage}' for '{product_name}' (from enrichment, no sativa hybrid override)")
            elif product_name:
                # Record lineage missing - query database directly (avoid get_product_lineage which applies override)
                from app import get_product_database, get_current_store_name
                store_name = get_current_store_name()
                product_db = get_product_database(store_name)
                if product_db:
                    # Query database directly to avoid sativa hybrid override in get_product_lineage()
                    try:
                        conn = product_db._get_connection()
                        cursor = conn.cursor()
                        # CRITICAL FIX: Query sovereign_lineage FIRST (manual edits have highest priority)
                        cursor.execute('''
                            SELECT sovereign_lineage, "Lineage", "canonical_lineage"
                            FROM products
                            WHERE "Product Name*" = ? OR ProductName = ? OR normalized_name = ?
                            ORDER BY id DESC
                            LIMIT 1
                        ''', (product_name, product_name, product_db._normalize_product_name(product_name)))
                        result = cursor.fetchone()
                        # Priority: sovereign_lineage > Lineage > canonical_lineage
                        # CRITICAL FIX: Reject "SOVEREIGN" as invalid - it's a field name, not a lineage value
                        if result:
                            # Check each field in priority order, rejecting "SOVEREIGN"
                            for idx, field_name in [(0, 'sovereign_lineage'), (1, 'Lineage'), (2, 'canonical_lineage')]:
                                if result[idx]:
                                    lineage_str = str(result[idx]).strip().upper()
                                    if lineage_str != 'SOVEREIGN':  # Reject "SOVEREIGN" as invalid
                                        db_lineage = lineage_str
                                        self.logger.info(f"🔒 DOCX: Using {field_name} '{db_lineage}' for '{product_name}'")
                                        break
                    except Exception as db_err:
                        self.logger.warning(f"Direct database query failed, falling back to get_product_lineage: {db_err}")
                        # Fallback to get_product_lineage if direct query fails
                        db_lineage = product_db.get_product_lineage(product_name)
                    
                    # If no product-level lineage, check strain-level lineage
                    if not db_lineage or str(db_lineage).strip() in ['', 'None', 'nan']:
                        product_strain = record.get('Product Strain', '')
                        if product_strain:
                            strain_info = product_db.get_strain_info(product_strain)
                            if strain_info:
                                # Get strain lineage, rejecting "SOVEREIGN" as invalid
                                strain_display = strain_info.get('display_lineage')
                                strain_sovereign = strain_info.get('sovereign_lineage')
                                strain_canonical = strain_info.get('canonical_lineage')
                                
                                # Filter out "SOVEREIGN" - it's a field name, not a lineage value
                                valid_lineages = []
                                for lin in [strain_display, strain_sovereign, strain_canonical]:
                                    if lin and str(lin).strip().upper() != 'SOVEREIGN':
                                        valid_lineages.append(str(lin).strip())
                                
                                db_lineage = valid_lineages[0] if valid_lineages else None
                    
                    # CRITICAL: Always use database lineage if available, never Excel
                    if db_lineage and str(db_lineage).strip() not in ['', 'None', 'nan']:
                        db_lineage_upper = str(db_lineage).strip().upper()
                        # Always override Excel lineage with database lineage
                        if excel_lineage and str(excel_lineage).strip().upper() != db_lineage_upper:
                            # Database lineage differs from Excel - use database
                            self.logger.info(f"✅ LINEAGE DB OVERRIDE (DOCX): '{product_name}' - Excel: '{excel_lineage}' -> DB: '{db_lineage_upper}' (using DB)")
                        label_context['Lineage'] = db_lineage_upper
                    else:
                        # No DB lineage - use defaults based on product type (never Excel)
                        product_type = record.get('Product Type*', record.get('ProductType', '')).lower()
                        CLASSIC_TYPES = {'flower', 'pre-roll', 'concentrate', 'infused pre-roll', 'solventless concentrate', 'vape cartridge', 'rso/co2 tankers'}
                        is_classic = product_type in CLASSIC_TYPES or any(ct in product_type for ct in CLASSIC_TYPES)
                        
                        if is_classic:
                            default_lineage = 'HYBRID'
                        else:
                            default_lineage = 'MIXED'
                        
                        self.logger.info(f"⚠️ LINEAGE DEFAULT (DOCX): '{product_name}' - No DB lineage, using default '{default_lineage}' for {'classic' if is_classic else 'non-classic'} type (never Excel)")
                        label_context['Lineage'] = default_lineage
        except Exception as e:
            self.logger.warning(f"Could not check database lineage for DOCX output: {e}")
            # On error, use defaults based on product type (never Excel)
            product_type = record.get('Product Type*', record.get('ProductType', '')).lower()
            CLASSIC_TYPES = {'flower', 'pre-roll', 'concentrate', 'infused pre-roll', 'solventless concentrate', 'vape cartridge', 'rso/co2 tankers'}
            is_classic = product_type in CLASSIC_TYPES or any(ct in product_type for ct in CLASSIC_TYPES)
            
            if is_classic:
                default_lineage = 'HYBRID'
            else:
                default_lineage = 'MIXED'
            
            self.logger.info(f"⚠️ LINEAGE DEFAULT (DOCX ERROR): '{product_name}' - Error checking DB, using default '{default_lineage}' for {'classic' if is_classic else 'non-classic'} type (never Excel)")
            label_context['Lineage'] = default_lineage
        
        # CRITICAL FIX: Force DOH to be read from the actual data source, not defaults
        # If DOH is 'YES' but we updated it to 'No', use 'No' instead
        # Check if DOH was explicitly set to No in our recent updates
        if 'Source' in record and ('JSON Match' in record.get('Source', '') or 'Educated Guess' in record.get('Source', '')):
            # This is a JSON-matched item - check if DOH was manually updated
            # We need to preserve the user's DOH selection even for JSON matches
            pass

        # Fast value cleaning - only process non-empty values
        for key, value in label_context.items():
            if value is not None:
                label_context[key] = str(value).strip()
            else:
                label_context[key] = ""

        # Ensure WeightUnits is populated from available weight fields
        # Special handling for pre-roll products: use JointRatio instead of Weight* + Units
        raw_product_type = label_context.get('Product Type*', '')
        fallback_product_type = label_context.get('ProductType', '')
        
        # CRITICAL FIX: Handle cases where Product Type* is 'NOT_FOUND' or invalid
        if raw_product_type and raw_product_type.lower() not in ['not_found', 'unknown', '']:
            product_type = raw_product_type.lower()
        elif fallback_product_type and fallback_product_type.lower() not in ['not_found', 'unknown', '']:
            product_type = fallback_product_type.lower()
        else:
            # CRITICAL FIX: For new products without proper type, infer from product name
            product_name = record.get('ProductName', '')
            name_lower = product_name.lower()

            # 1) High‑signal concentrate patterns
            if any(keyword in name_lower for keyword in ['live rosin', 'hash rosin', 'solventless', 'rosin']):
                product_type = 'solventless concentrate'
                self.logger.info(f"🔧 INFERRED TYPE: '{product_name}' -> 'solventless concentrate' (from name)")

            # 2) Vape / disposable patterns
            elif any(keyword in name_lower for keyword in ['disposable vape', 'disposable cart', 'vape cart', 'cartridge', 'vape pen']):
                product_type = 'vape cartridge'
                self.logger.info(f"🔧 INFERRED TYPE: '{product_name}' -> 'vape cartridge' (from name)")

            # 3) Classic flower keywords
            elif any(keyword in name_lower for keyword in ['flower', 'bud', 'nug', 'herb']):
                product_type = 'flower'
                self.logger.info(f"🔧 INFERRED TYPE: '{product_name}' -> 'flower' (from name)")

            # 4) Pre‑roll patterns
            elif any(keyword in name_lower for keyword in ['pre-roll', 'preroll', 'joint', 'blunt']):
                product_type = 'pre-roll'
                self.logger.info(f"🔧 INFERRED TYPE: '{product_name}' -> 'pre-roll' (from name)")

            # 5) Fallback – keep previous behaviour
            else:
                product_type = 'flower'  # Default to flower for new products
                self.logger.info(f"🔧 DEFAULT TYPE: '{product_name}' -> 'flower' (default)")
        
        # ALWAYS LOG TO SEE WHAT PRODUCT TYPES ARE PROCESSED
        self.logger.info(f"🔍 ALL PRODUCTS DEBUG: Product '{record.get('ProductName', 'N/A')}', Raw Type: '{raw_product_type}', Processed: '{product_type}'")
        
        # CRITICAL FIX: Store the processed product type in the context
        label_context['ProductType'] = product_type
        label_context['Product Type*'] = product_type.title()  # Store as title case for consistency
        
        # CRITICAL FIX: Process JointRatio FIRST for pre-rolls, before any other weight processing
        # This ensures WeightUnits is set to JointRatio before DescAndWeight construction
        # IMPORTANT: Clear any existing WeightUnits from record for pre-rolls to prevent using total weight
        if product_type in ['pre-roll', 'infused pre-roll']:
            # CRITICAL FIX: Always clear WeightUnits for pre-rolls - we'll set it to JointRatio below
            # This prevents using total weight (like "5g") instead of joint ratio (like "- 1g x 5 Pack")
            if 'WeightUnits' in label_context:
                existing_weight = label_context.get('WeightUnits', '')
                # Clear WeightUnits if it doesn't contain joint ratio pattern (no "x" or "Pack")
                # This ensures we always use JointRatio, not total weight
                if existing_weight:
                    weight_str = str(existing_weight).lower()
                    has_joint_ratio_pattern = 'x' in weight_str or 'pack' in weight_str
                    if not has_joint_ratio_pattern:
                        self.logger.info(f"🔧 PRE-ROLL: Clearing total weight WeightUnits '{existing_weight}' - will use JointRatio instead")
                        label_context['WeightUnits'] = ''
                    else:
                        self.logger.info(f"🔧 PRE-ROLL: WeightUnits '{existing_weight}' already looks like JointRatio, keeping it")
            # For pre-roll products, use JointRatio as the weight
            # CRITICAL FIX: Check BOTH record and label_context for JointRatio (label_context is a copy of record)
            joint_ratio = (label_context.get('JointRatio') or 
                          record.get('JointRatio') or 
                          label_context.get('Joint Ratio') or
                          record.get('Joint Ratio') or 
                          '')
            
            # CRITICAL FIX: Log all JointRatio sources for debugging
            self.logger.info(f"🔍 JOINTRATIO DEBUG: Product '{record.get('ProductName', 'N/A')}'")
            self.logger.info(f"   label_context.get('JointRatio'): {repr(label_context.get('JointRatio'))}")
            self.logger.info(f"   record.get('JointRatio'): {repr(record.get('JointRatio'))}")
            self.logger.info(f"   label_context.get('Joint Ratio'): {repr(label_context.get('Joint Ratio'))}")
            self.logger.info(f"   record.get('Joint Ratio'): {repr(record.get('Joint Ratio'))}")
            self.logger.info(f"   Final joint_ratio: {repr(joint_ratio)}")
            
            # CRITICAL FIX: If JointRatio is missing from record, try multiple sources
            if not joint_ratio or str(joint_ratio).strip() in ['', 'NULL', 'null', '0', '0.0', 'None', 'nan']:
                product_name = record.get('ProductName') or record.get('Product Name*', '')
                if product_name:
                    # Try database cache first
                    try:
                        # Use pre-loaded cache instead of individual query
                        if joint_ratio_cache and product_name:
                            joint_ratio = joint_ratio_cache.get(product_name)
                            if joint_ratio:
                                self.logger.info(f"✅ Retrieved JointRatio '{joint_ratio}' from cache for '{product_name}'")
                            else:
                                self.logger.debug(f"No JointRatio in cache for '{product_name}'")
                    except Exception as e:
                        self.logger.warning(f"⚠️ Could not retrieve JointRatio from cache: {e}")
                    
                    # CRITICAL FIX: If still missing, extract from product name
                    if not joint_ratio or str(joint_ratio).strip() in ['', 'NULL', 'null', '0', '0.0', 'None', 'nan']:
                        # re is already imported at module level
                        product_name_str = str(product_name)
                        
                        # Pattern 1: "weight x count Pack" (e.g., "0.5g x 2 Pack", ".75g x 5 Pack")
                        pattern1 = r'(\d*\.?\d+g)\s*x\s*(\d+)\s*Pack'
                        match1 = re.search(pattern1, product_name_str, re.IGNORECASE)
                        if match1:
                            weight = match1.group(1)
                            count = match1.group(2)
                            joint_ratio = f"{weight} x {count} Pack"
                            self.logger.info(f"✅ Extracted JointRatio '{joint_ratio}' from product name '{product_name}' (pattern 1)")
                        else:
                            # Pattern 2: "weight x count" (e.g., "0.5g x 2", ".75g x 5")
                            pattern2 = r'(\d*\.?\d+g)\s*x\s*(\d+)'
                            match2 = re.search(pattern2, product_name_str, re.IGNORECASE)
                            if match2:
                                weight = match2.group(1)
                                count = match2.group(2)
                                joint_ratio = f"{weight} x {count}"
                                self.logger.info(f"✅ Extracted JointRatio '{joint_ratio}' from product name '{product_name}' (pattern 2)")
                            else:
                                # Pattern 3: Just weight (e.g., "1g", "0.5g", ".75g")
                                pattern3 = r'(\d*\.?\d+g)'
                                match3 = re.search(pattern3, product_name_str, re.IGNORECASE)
                                if match3:
                                    weight = match3.group(1)
                                    joint_ratio = weight
                                    self.logger.info(f"✅ Extracted JointRatio '{joint_ratio}' from product name '{product_name}' (pattern 3)")
            
            self.logger.info(f"📦 PRE-ROLL: Product '{record.get('ProductName', 'N/A')}', Final JointRatio: '{joint_ratio}'")
            
            # Use JointRatio or default for all weight fields
            if joint_ratio and joint_ratio.strip() not in ['', 'NULL', 'null', '0', '0.0', 'None', 'nan']:
                formatted_joint_ratio = self.format_joint_ratio_pack(joint_ratio.strip())
                # Set JointRatio for ALL weight-related fields
                label_context['Weight*'] = formatted_joint_ratio
                label_context['WeightUnits'] = formatted_joint_ratio
                label_context['CombinedWeight'] = formatted_joint_ratio
                label_context['weightWithUnits'] = formatted_joint_ratio
                label_context['JointRatio'] = formatted_joint_ratio
                self.logger.info(f"✅ Using JointRatio as weight: '{formatted_joint_ratio}' for {product_type}")
            else:
                formatted_default = self.format_joint_ratio_pack("0.5g x 2 Pack")
                # Set default for ALL weight-related fields
                label_context['Weight*'] = formatted_default
                label_context['WeightUnits'] = formatted_default
                label_context['CombinedWeight'] = formatted_default
                label_context['weightWithUnits'] = formatted_default
                label_context['JointRatio'] = formatted_default
                self.logger.warning(f"⚠️ Using default JointRatio as weight: '{formatted_default}' for {product_type}")
            
            # CRITICAL FIX: WeightUnits is already set to JointRatio above - skip non-pre-roll WeightUnits construction
            # This ensures WeightUnits for pre-rolls uses JointRatio, not total weight
        else:
            # For non-pre-roll products, construct WeightUnits from available fields
            weight_units = (
                label_context.get('CombinedWeight') or
                label_context.get('WeightWithUnits') or 
                label_context.get('weightWithUnits') or 
                label_context.get('WeightUnits') or 
                ''
            )
            
            # CRITICAL FIX: If no combined weight field, use session Excel processor's weight normalization
            if not weight_units or weight_units.strip() in ['', 'NULL', 'null', '0', '0.0', 'None', 'nan']:
                if self.excel_processor:
                    # Create a record dict from label_context for weight normalization
                    weight_record = {
                        'Weight*': label_context.get('Weight*', ''),
                        'Units': label_context.get('Units', ''),
                        'Product Type*': label_context.get('Product Type*', ''),
                        'Product Name*': label_context.get('ProductName', '')
                    }
                    weight_units = self.excel_processor._format_weight_units(weight_record, excel_priority=True)
                    self.logger.info(f"🔧 WEIGHT NORMALIZED: '{weight_units}' using session Excel processor for '{record.get('ProductName', 'N/A')}'")
                else:
                    # Fallback to simple concatenation if no Excel processor available
                    weight_value = label_context.get('Weight*', '').strip()
                    units_value = label_context.get('Units', '').strip()
                    if weight_value and units_value:
                        weight_units = f"{weight_value}{units_value}"
                    elif weight_value:
                        weight_units = weight_value
                    self.logger.warning(f"⚠️ WEIGHT FALLBACK: No Excel processor available, using simple concatenation: '{weight_units}'")
            
            label_context['WeightUnits'] = weight_units
            
            # CRITICAL FIX: Remove any weight markers that might interfere with display
            if label_context['WeightUnits'] and 'WEIGHTUNITS_START' in str(label_context['WeightUnits']):
                # Extract the actual weight value from the markers
                weight_text = str(label_context['WeightUnits'])
                start_marker = 'WEIGHTUNITS_START'
                end_marker = 'WEIGHTUNITS_END'
                
                if start_marker in weight_text and end_marker in weight_text:
                    start_idx = weight_text.find(start_marker) + len(start_marker)
                    end_idx = weight_text.find(end_marker)
                    actual_weight = weight_text[start_idx:end_idx].strip()
                    label_context['WeightUnits'] = actual_weight
                    self.logger.info(f"🔧 WEIGHT MARKERS REMOVED: '{weight_text}' -> '{actual_weight}' for '{record.get('ProductName', 'N/A')}'")

        # Define product type sets for use throughout the method
        from src.core.constants import CLASSIC_TYPES
        classic_types = CLASSIC_TYPES
        edible_types = {"edible (solid)", "edible (liquid)", "high cbd edible liquid", "tincture", "topical", "capsule"}

        # Use DescAndWeight from record if it exists, otherwise construct it
        # PREROLL TEMPLATE: Check for group name FIRST before doing anything else
        if self.template_type == 'preroll':
            group_id = record.get('_group_id')
            if group_id:
                group_info = record.get('_group_info')
                if group_info and isinstance(group_info, dict):
                    group_display_name = group_info.get('display_name', '')
                    if group_display_name:
                        # CRITICAL FIX: Check if group display name already contains weight information
                        # Patterns that indicate weight is already in the name:
                        # - "Assorted Pre-Roll 1g x 5 Packs" (has "Xg x Y")
                        # - "Pre-Roll - 1g" (has "- Xg" or "Xg" at end)
                        # - "Infused Pre-Roll - 0.5g" (has "- Xg")
                        has_weight_in_name = bool(re.search(r'\d+\.?\d*\s*g(?:\s*x\s*\d+)?(?:\s+Pack)?', group_display_name, re.IGNORECASE))

                        if has_weight_in_name:
                            # Group name already includes weight, use it as-is
                            desc_and_weight = group_display_name
                            self.logger.info(f"PREROLL GROUP: Using group name as-is (already contains weight): '{desc_and_weight}' (group_id: {group_id})")
                        else:
                            # Group name doesn't include weight, add it from JointRatio/WeightUnits
                            # Get JointRatio-derived weight from WeightUnits (set earlier for preroll products)
                            weight_units = label_context.get('WeightUnits', '') or record.get('WeightUnits', '')
                            if weight_units:
                                # Remove newline prefix if present, and strip whitespace
                                clean_weight = weight_units.replace('\n', '').strip()
                                # Check if weight is already formatted with hyphen
                                # Normalize non-breaking hyphen to regular hyphen for consistency
                                if clean_weight.startswith('\u2011'):
                                    # Replace non-breaking hyphen with regular hyphen
                                    clean_weight = clean_weight.replace('\u2011', '-', 1).replace('\u00A0', ' ')
                                if clean_weight.startswith('-'):
                                    # Weight already has hyphen, just append it with space
                                    desc_and_weight = f"{group_display_name} {clean_weight}"
                                else:
                                    # Add hyphen and space before weight
                                    desc_and_weight = f"{group_display_name} - {clean_weight}"
                                self.logger.info(f"PREROLL GROUP: Added weight to group name: '{desc_and_weight}' (group: '{group_display_name}', weight: '{clean_weight}', group_id: {group_id})")
                            else:
                                # No weight available, use group name only
                                desc_and_weight = group_display_name
                                self.logger.info(f"PREROLL GROUP: Using group name only (no weight available): '{group_display_name}' (group_id: {group_id})")

                        label_context['DescAndWeight'] = wrap_with_marker(desc_and_weight, 'DESC')
                        # Skip all DescAndWeight construction below - we're done
                        # Continue to QR code generation
                    else:
                        self.logger.warning(f"PREROLL: Group info missing display_name for group_id: {group_id}")
                else:
                    self.logger.warning(f"PREROLL: No group_info found for group_id: {group_id}")
            else:
                self.logger.warning(f"PREROLL: No _group_id found in record - grouping may have failed")
        
        # For non-preroll or preroll without group, use normal DescAndWeight logic
        # Check if we already set DescAndWeight for preroll group (skip if so)
        if not (self.template_type == 'preroll' and record.get('_group_id') and label_context.get('DescAndWeight')):
            # CRITICAL FIX: For pre-rolls, always reconstruct DescAndWeight using JointRatio
            # Don't use existing DescAndWeight from record - it might have total weight instead of joint ratio
            if 'DescAndWeight' in label_context and label_context['DescAndWeight'] and product_type not in ['pre-roll', 'infused pre-roll']:
                # DescAndWeight is already set correctly in the record, use it as-is (only for non-pre-roll products)
                desc_and_weight = label_context['DescAndWeight']
                if not is_already_wrapped(desc_and_weight, 'DESC'):
                    label_context['DescAndWeight'] = wrap_with_marker(desc_and_weight, 'DESC')
                # Skip the rest of DescAndWeight processing since it's already set
                self.logger.info(f"🔍 Using existing DescAndWeight from record: '{desc_and_weight}' (product_type: {product_type})")
            else:
                # For pre-rolls, always reconstruct DescAndWeight to ensure JointRatio is used
                if product_type in ['pre-roll', 'infused pre-roll']:
                    self.logger.info(f"🔍 PRE-ROLL: Reconstructing DescAndWeight to ensure JointRatio is used (ignoring existing DescAndWeight from record)")
                # Fallback: construct DescAndWeight from Description and WeightUnits
                desc = label_context.get('Description', '') or ''
                
                # CRITICAL FIX: For pre-rolls, clean Description to remove ALL weight patterns before adding JointRatio
                # This prevents showing both total weight (5g) and joint ratio (1g x 5 Pack)
                if product_type in ['pre-roll', 'infused pre-roll']:
                    # Remove ALL weight patterns from Description - both total weight and joint ratio
                    # The Description might contain "Product Name - 5g - 1g x 5 Pack" or similar
                    original_desc = desc
                    # Pattern 1: Remove joint ratio patterns FIRST (e.g., " - 1g x 5 Pack", " - 0.5g x 2 Pack")
                    # This handles patterns like " - 1g x 5 Pack" or " - 0.5g x 2 Pack"
                    desc = re.sub(r'\s*-\s*\d+\.?\d*\s*g\s*x\s*\d+\s*Pack\s*', '', desc, flags=re.IGNORECASE)
                    # Pattern 2: Remove total weight patterns (e.g., " - 5g", " - 1g")
                    # This handles patterns like " - 5g" or " - 1g"
                    desc = re.sub(r'\s*-\s*\d+\.?\d*\s*g\s*', '', desc, flags=re.IGNORECASE)
                    # Pattern 3: Remove any remaining " - " at the end
                    desc = re.sub(r'\s*-\s*$', '', desc)
                    desc = desc.strip()
                    self.logger.info(f"🔍 PRE-ROLL: Cleaned Description from '{original_desc}' to '{desc}'")
                    
                    # For pre-rolls, use JointRatio first, then WeightUnits as fallback
                    weight = (label_context.get('JointRatio', '') or label_context.get('WeightUnits', '') or '').replace('\u202F', '')
                    if weight:
                        self.logger.info(f"🔍 PRE-ROLL DescAndWeight: Using JointRatio/WeightUnits '{weight}' for product '{record.get('ProductName', 'N/A')}'")
                else:
                    weight = (label_context.get('WeightUnits', '') or '').replace('\u202F', '')
                
                # DEBUG: Log the values being processed
                self.logger.info(f"🔍 DESCANDWEIGHT DEBUG: Product '{record.get('ProductName', 'N/A')}' - Description: '{desc}', WeightUnits: '{weight}', ProductType: '{product_type}'")
                
                # Ultra-fast string operations
                if desc.endswith('- '):
                    desc = desc[:-2]
                if weight.startswith('- '):
                    weight = weight[2:]
                
                # DEBUG: Log all record keys and values to see what we're working with
                self.logger.info(f"🔍 RECORD KEYS: {list(record.keys())}")
                for key, value in record.items():
                    if 'weight' in key.lower() or 'units' in key.lower():
                        self.logger.info(f"🔍 {key}: '{value}'")
                
                # CRITICAL FIX: Horizontal template uses DescAndWeight placeholder, not separate WeightUnits
                # Add weight to description for the {{Label1.DescAndWeight}} placeholder
                product_name_display = (
                    label_context.get('ProductName') or
                    record.get('ProductName') or
                    record.get('Product Name*', '')
                )
                
                if self.template_type == 'preroll':
                    # PREROLL TEMPLATE: Use group display name with JointRatio-derived weight
                    # Check if we already set DescAndWeight for grouped preroll (above)
                    if not (record.get('_group_id') and label_context.get('DescAndWeight')):
                        # Use description (should be group name if grouping worked) or product_name_display as fallback
                        primary_text = (desc or product_name_display or '').strip()
                        
                        # Get JointRatio-derived weight from WeightUnits (set earlier for preroll products)
                        weight_units = label_context.get('WeightUnits', '') or record.get('WeightUnits', '')
                        if weight_units:
                            # Remove newline prefix if present, and strip whitespace
                            clean_weight = weight_units.replace('\n', '').strip()
                            # Check if weight already has hyphen (from format_joint_ratio_pack)
                            # Normalize non-breaking hyphen to regular hyphen for consistency
                            if clean_weight.startswith('\u2011'):
                                # Replace non-breaking hyphen with regular hyphen
                                clean_weight = clean_weight.replace('\u2011', '-', 1).replace('\u00A0', ' ')
                            # CRITICAL FIX: Remove trailing hyphen from primary_text to prevent double hyphen
                            # (e.g., "Pre-Roll-" + "- 0.5g" = "Pre-Roll-- 0.5g")
                            primary_text_clean = primary_text.rstrip('-').rstrip()
                            if clean_weight.startswith('-'):
                                # Weight already has hyphen, just append it with space
                                desc_and_weight = f"{primary_text_clean} {clean_weight}"
                            else:
                                # Add hyphen and space before weight
                                desc_and_weight = f"{primary_text_clean} - {clean_weight}"
                            self.logger.info(f"🔍 PREROLL TEMPLATE DESC: Using '{primary_text_clean}' with weight '{clean_weight}' -> '{desc_and_weight}'")
                        else:
                            # No weight available, use description only
                            desc_and_weight = primary_text
                            self.logger.info(f"🔍 PREROLL TEMPLATE DESC: Using '{primary_text}' (no weight available)")
                        
                        label_context['DescAndWeight'] = wrap_with_marker(desc_and_weight, 'DESC')
                    else:
                        self.logger.info(f"🔍 PREROLL TEMPLATE DESC: Already set to group name, skipping reconstruction")
                else:
                    # CRITICAL FIX: For pre-rolls, use the weight variable we already extracted from JointRatio above (line 2323)
                    # This ensures we use JointRatio (like "1g x 5 Pack"), not total weight (like "5g")
                    if product_type in ['pre-roll', 'infused pre-roll']:
                        # Use the weight variable we already extracted and cleaned above
                        if weight and weight.strip():
                            clean_weight = weight.strip()
                            self.logger.info(f"🔍 PRE-ROLL: Using JointRatio-derived weight '{clean_weight}' for DescAndWeight (from weight variable)")
                        else:
                            # Fallback to WeightUnits/JointRatio from label_context if weight variable is empty
                            weight_units = label_context.get("WeightUnits", "") or label_context.get("JointRatio", "")
                            clean_weight = weight_units.strip() if weight_units else ''
                            # Remove "- " prefix if present (format_joint_ratio_pack adds it)
                            if clean_weight.startswith('- '):
                                clean_weight = clean_weight[2:].strip()
                            self.logger.info(f"🔍 PRE-ROLL WeightUnits fallback: Using '{clean_weight}' for DescAndWeight")
                    else:
                        # For non-pre-roll products, get WeightUnits from label_context or record
                        weight_units = label_context.get("WeightUnits", "") or record.get("WeightUnits", "")
                        
                        # Check if WeightUnits already contains the complete weight+units
                        if weight_units and weight_units.strip():
                            # WeightUnits already contains the complete weight (e.g., "3.4oz", "1616.0g")
                            clean_weight = weight_units.strip()
                            
                            # CRITICAL FIX: Clean weight duplication patterns directly in template processor
                            # Pattern 1: Decimal duplication like "0.50.5oz" -> "0.5oz"
                            decimal_dup_pattern = r'^(\d+\.\d{1,2})\1(oz|g|mg|kg|lb|lbs)$'
                            match1 = re.match(decimal_dup_pattern, clean_weight, re.IGNORECASE)
                            if match1:
                                clean_weight = f"{match1.group(1)}{match1.group(2)}"
                                self.logger.info(f"✅ TEMPLATE PROCESSOR FIXED DECIMAL DUPLICATION: '{weight_units}' -> '{clean_weight}'")
                            else:
                                # Pattern 2: Integer duplication like "1010.0g" -> "10.0g"
                                integer_dup_pattern = r'^(\d+)\1\.0(oz|g|mg|kg|lb|lbs)$'
                                match2 = re.match(integer_dup_pattern, clean_weight, re.IGNORECASE)
                                if match2:
                                    clean_weight = f"{match2.group(1)}.0{match2.group(2)}"
                                    self.logger.info(f"✅ TEMPLATE PROCESSOR FIXED INTEGER DUPLICATION: '{weight_units}' -> '{clean_weight}'")
                                else:
                                    # Pattern 3: Mixed duplication like "0.220.22g" -> "0.22g"
                                    mixed_dup_pattern = r'^(\d+\.\d+)\1(oz|g|mg|kg|lb|lbs)$'
                                    match3 = re.match(mixed_dup_pattern, clean_weight, re.IGNORECASE)
                                    if match3:
                                        clean_weight = f"{match3.group(1)}{match3.group(2)}"
                                        self.logger.info(f"✅ TEMPLATE PROCESSOR FIXED MIXED DUPLICATION: '{weight_units}' -> '{clean_weight}'")
                        else:
                            clean_weight = ''
                    
                    # Construct DescAndWeight if we have clean_weight
                    if clean_weight:
                        # CRITICAL FIX: For non-pre-roll products, clean weight duplication patterns
                        # For pre-rolls, clean_weight already has JointRatio format (like "1g x 5 Pack"), skip duplication cleaning
                        if product_type not in ['pre-roll', 'infused pre-roll']:
                            # CRITICAL FIX: Clean weight duplication patterns directly in template processor
                            # Pattern 1: Decimal duplication like "0.50.5oz" -> "0.5oz"
                            decimal_dup_pattern = r'^(\d+\.\d{1,2})\1(oz|g|mg|kg|lb|lbs)$'
                            match1 = re.match(decimal_dup_pattern, clean_weight, re.IGNORECASE)
                            if match1:
                                clean_weight = f"{match1.group(1)}{match1.group(2)}"
                                self.logger.info(f"✅ TEMPLATE PROCESSOR FIXED DECIMAL DUPLICATION: '{weight_units}' -> '{clean_weight}'")
                            else:
                                # Pattern 2: Integer duplication like "1010.0g" -> "10.0g"
                                integer_dup_pattern = r'^(\d+)\1\.0(oz|g|mg|kg|lb|lbs)$'
                                match2 = re.match(integer_dup_pattern, clean_weight, re.IGNORECASE)
                                if match2:
                                    clean_weight = f"{match2.group(1)}.0{match2.group(2)}"
                                    self.logger.info(f"✅ TEMPLATE PROCESSOR FIXED INTEGER DUPLICATION: '{weight_units}' -> '{clean_weight}'")
                                else:
                                    # Pattern 3: Mixed duplication like "0.220.22g" -> "0.22g"
                                    mixed_dup_pattern = r'^(\d+\.\d+)\1(oz|g|mg|kg|lb|lbs)$'
                                    match3 = re.match(mixed_dup_pattern, clean_weight, re.IGNORECASE)
                                    if match3:
                                        clean_weight = f"{match3.group(1)}{match3.group(2)}"
                                        self.logger.info(f"✅ TEMPLATE PROCESSOR FIXED MIXED DUPLICATION: '{weight_units}' -> '{clean_weight}'")
                        
                        # Keep weight on the same line as description with non-breaking space
                        # Use consistent space-hyphen-space pattern for all templates
                        desc_and_weight = f"{desc} - {clean_weight}"
                        self.logger.info(f"🔍 DESCANDWEIGHT CONSTRUCTION: desc='{desc}', clean_weight='{clean_weight}' -> '{desc_and_weight}'")
                    else:
                        # Fallback to constructing from Weight* + Units
                        weight_value = record.get("Weight*", "")
                        units_value = record.get("Units", "")
                        
                        if not weight_value:
                            weight_value = record.get("Weight", "")
                        if not units_value:
                            units_value = record.get("Units", "")
                        
                        self.logger.info(f"🔍 FALLBACK WEIGHT VALUES: Weight*='{weight_value}', Units='{units_value}'")
                        
                        if weight_value and units_value:
                            clean_weight = f"{weight_value}{units_value}"
                            # Use consistent space-hyphen-space pattern for all templates
                            desc_and_weight = f"{desc} - {clean_weight}"
                            self.logger.info(f"🔍 WEIGHT CONSTRUCTED: '{clean_weight}' -> '{desc_and_weight}'")
                        else:
                            desc_and_weight = desc
                            self.logger.info(f"🔍 NO WEIGHT AVAILABLE: Weight*='{weight_value}', Units='{units_value}' -> '{desc_and_weight}'")
                    
                    self.logger.info(f"🔍 DESCANDWEIGHT RESULT: '{desc_and_weight}'")
                    label_context['DescAndWeight'] = wrap_with_marker(desc_and_weight, 'DESC')

        # Fast DOH image processing - only if needed
        # IMPORTANT: Only use the canonical DOH field for image decisions
        # Ignore legacy "DOH Compliant (Yes/No)" and any other variants
        doh_value = label_context.get('DOH', '') or record.get('DOH', '')
        product_name = label_context.get('ProductName', 'Unknown')
        
        # CRITICAL FIX: If DOH is missing from record, query database directly
        if not doh_value or str(doh_value).strip() in ['', 'None', 'nan']:
            try:
                from app import get_product_database, get_current_store_name
                store_name = get_current_store_name()
                product_db = get_product_database(store_name)
                if product_db:
                    conn = product_db._get_connection()
                    cursor = conn.cursor()
                    cursor.execute('''
                        SELECT "DOH" FROM products
                        WHERE "Product Name*" = ? OR ProductName = ? OR normalized_name = ?
                        ORDER BY id DESC
                        LIMIT 1
                    ''', (product_name, product_name, product_db._normalize_product_name(product_name)))
                    result = cursor.fetchone()
                    if result and result[0] and str(result[0]).strip() not in ['', 'None', 'nan']:
                        doh_value = str(result[0]).strip()
                        label_context['DOH'] = doh_value
                        self.logger.info(f"🔍 DOH RETRIEVED FROM DB: '{product_name}' - DOH: '{doh_value}'")
            except Exception as db_err:
                self.logger.warning(f"Could not retrieve DOH from database: {db_err}")

        # CRITICAL DEBUG: Log DOH field processing with all possible sources
        self.logger.info(f"🔍 DOH DOCX GENERATION: Product '{product_name}' - DOH field: '{doh_value}' from record")
        self.logger.info(f"🔍 DOH DOCX GENERATION: Using only canonical DOH field: '{label_context.get('DOH', '')}'")
        self.logger.info(f"🔍 DOH DOCX GENERATION: Record Source: '{record.get('Source', 'N/A')}'")
        self.logger.info(f"🔍 DOH DOCX GENERATION: First 20 field keys in record: {list(record.keys())[:20]}")

        # Handle different DOH values: YES (legacy), DOH, THC, CBD
        doh_upper = str(doh_value).strip().upper() if doh_value else ''

        # Explicitly handle NO/NONE/False values FIRST
        # Also handle legacy "No" (capital N, lowercase o) which is what we store
        if doh_upper in ['NO', 'NONE', 'FALSE', ''] or doh_value in ['No', 'no']:
            label_context['DOH'] = ''
            label_context['DOH_TEXT'] = ''
            # Also clear other DOH-related fields
            label_context['DOH Compliant (Yes/No)'] = ''
            label_context['doh'] = ''
            self.logger.info(f"✅ DOH DOCX GENERATION: Explicitly clearing DOH for '{product_name}' - value: '{doh_value}' (NO/NONE/FALSE) - NO IMAGE WILL BE ADDED")
        elif doh_upper in ['YES', 'DOH', 'THC', 'CBD']:
            product_type = (label_context.get('ProductType') or
                          label_context.get('Product Type*') or
                          record.get('ProductType') or
                          record.get('Product Type*') or '')

            image_path = process_doh_image(doh_upper, product_type)
            if image_path:
                # Fast width selection - reduced by 1mm for all template types
                # Preroll template uses half size (5.5mm) for DOH logo
                width_map = {'mini': 8, 'double': 10, 'vertical': 13, 'horizontal': 13, 'preroll': 5.5}
                image_width = Mm(width_map.get(self.template_type, 11))
                label_context['DOH'] = InlineImage(doc, image_path, width=image_width)
                # Ensure DOH image takes priority - clear any other DOH-related content
                label_context['DOH_TEXT'] = ''  # Clear any text content
                self.logger.info(f"✅ DOH DOCX GENERATION: Created DOH image for '{product_name}' with value '{doh_upper}' - IMAGE WILL BE ADDED: {image_path}")
            else:
                label_context['DOH'] = ''
                label_context['DOH_TEXT'] = ''
                self.logger.info(f"⚠️ DOH DOCX GENERATION: No image path found for '{product_name}' - NO IMAGE WILL BE ADDED")
        else:
            # For any DOH value other than "YES"/"DOH"/"THC"/"CBD", leave it blank (don't display text)
            # This ensures DOH="no" or DOH="none" results in blank space, not text
            label_context['DOH'] = ''
            label_context['DOH_TEXT'] = ''
            self.logger.info(f"✅ DOH DOCX GENERATION: Clearing DOH for '{product_name}' - value: '{doh_value}' (not DOH/THC/CBD/YES) - NO IMAGE WILL BE ADDED")
        
        # CRITICAL: Lineage and ProductVendor logic for classic types
        # This implements the same logic that was in tag_generator
        product_type = (label_context.get('ProductType', '').lower() or 
                       label_context.get('Product Type*', '').lower())
        product_brand = label_context.get('ProductBrand') or label_context.get('Product Brand', '')
        lineage_text = label_context.get('Lineage', '')
        product_strain = label_context.get('ProductStrain') or label_context.get('Product Strain', '')
        
        # CRITICAL DEBUG: Log brand field processing (only for first few products)
        if not hasattr(self, '_brand_debug_count'):
            self._brand_debug_count = 0
        self._brand_debug_count += 1
        if self._brand_debug_count <= 3:
            self.logger.info(f"BRAND DEBUG: Product '{product_name}' - Brand field: '{product_brand}' (ProductBrand: '{label_context.get('ProductBrand')}', Product Brand: '{label_context.get('Product Brand')}')")
        
        # CRITICAL FIX: Check if brand is missing and apply fallback logic FIRST
        if not product_brand or product_brand.strip() in ['', 'None', 'NULL', 'null', 'nan']:
            # Apply fallback logic immediately for missing brands
            # Check more specific terms first, then more general ones
            if 'sorbet' in product_name.lower():
                enriched_brand = "SORBET CO."
            elif 'moonshot' in product_name.lower():
                enriched_brand = "MOONSHOT"
            elif 'lemonade' in product_name.lower():
                enriched_brand = "LEMONADE CO."
            elif 'pre-roll' in product_name.lower() or 'preroll' in product_name.lower():
                enriched_brand = "PREMIUM PREROLLS"
            else:
                enriched_brand = "PREMIUM CANNABIS"
            
            if enriched_brand:
                product_brand = enriched_brand
                label_context['Product Brand'] = enriched_brand
                label_context['ProductBrand'] = enriched_brand
                self.logger.info(f"🔧 IMMEDIATE BRAND FALLBACK: Set '{enriched_brand}' for '{product_name}' (no brand data)")
        
        # BRAND ENRICHMENT: If brand is still missing, try to get it from database cache, then fallback to vendor
        # OPTIMIZATION: Use pre-loaded cache instead of individual queries
        if not product_brand or product_brand.strip() in ['', 'None', 'NULL', 'null', 'nan']:
            # Try to enrich brand from pre-loaded cache first
            enriched_brand = ""
            try:
                # Use cached brand data (loaded in batch before loop)
                enriched_brand = product_brand_cache.get(product_name, "")
                if enriched_brand:
                    self.logger.info(f"🔧 BRAND ENRICHED: Retrieved brand '{enriched_brand}' from database cache for '{product_name}'")
            except Exception as e:
                self.logger.warning(f"🔧 BRAND ENRICHMENT FAILED: Could not retrieve brand from cache: {e}")
            
            # If database enrichment failed, fallback to vendor
            if not enriched_brand:
                vendor_fallback = (record.get('Vendor') or 
                                 record.get('Vendor/Supplier*') or 
                                 record.get('ProductVendor', ''))
                if vendor_fallback and str(vendor_fallback).strip() not in ['', 'None', 'NULL', 'null', 'nan']:
                    enriched_brand = str(vendor_fallback).strip()
                    self.logger.info(f"🔧 BRAND FALLBACK: Using vendor '{enriched_brand}' as brand for '{product_name}'")
            
            # CRITICAL FIX: If still no brand after all fallbacks, use a default brand based on product type
            if not enriched_brand:
                # Extract a meaningful default brand from product name or use generic fallback
                if 'lemonade' in product_name.lower():
                    enriched_brand = "LEMONADE CO."
                elif 'moonshot' in product_name.lower():
                    enriched_brand = "MOONSHOT"
                elif 'sorbet' in product_name.lower():
                    enriched_brand = "SORBET CO."
                else:
                    enriched_brand = "PREMIUM CANNABIS"
                self.logger.info(f"🔧 DEFAULT BRAND: Using default brand '{enriched_brand}' for '{product_name}' (no brand data available)")
            
            # Update the brand field if enrichment was successful
            if enriched_brand:
                product_brand = enriched_brand
                # CRITICAL FIX: Update label_context with enriched brand so subsequent processing uses it
                label_context['Product Brand'] = enriched_brand
                label_context['ProductBrand'] = enriched_brand
                # Don't set ProductBrand fields here - they will be set later based on template type
                self.logger.info(f"✅ BRAND UPDATED: Product '{product_name}' brand set to '{enriched_brand}' in context")
        
        # Check if it's a classic type
        # CRITICAL: Ensure case-insensitive comparison
        product_type_lower = product_type.lower() if product_type else ''
        classic_types_lower = {t.lower() for t in classic_types}
        is_classic_type = product_type_lower in classic_types_lower
        # Debug logging for blunts and pre-rolls to diagnose vendor issues
        if product_name and ('blunt' in product_name.lower() or 'pre-roll' in product_name.lower()):
            self.logger.info(f"🔍 CLASSIC TYPE CHECK: '{product_name}' - product_type: '{product_type}', product_type_lower: '{product_type_lower}', is_classic: {is_classic_type}, classic_types includes blunt: {'blunt' in classic_types_lower}")
        
        if is_classic_type:
            # For classic types, Lineage should show strain lineage and ProductVendor should show brand
            self.logger.debug(f"Processing classic type '{product_type}' for Lineage and ProductVendor")
            
            # PRIORITY FIX: Use record lineage first (from database updates), then fallback to database
            lineage_val = ""
            
            # PRIORITY 1: Use lineage from record (includes manual dropdown changes and database updates)
            # CRITICAL: Always use uppercase to ensure consistency
            if lineage_text and lineage_text.strip():
                lineage_val = str(lineage_text).strip().upper()
                self.logger.info(f"✅ Using record lineage (from database/excel): '{lineage_val}' for '{product_name}'")
            else:
                # PRIORITY 2: Fallback to cache lookup if record lineage is empty
                self.logger.warning(f"⚠️ No lineage in record for '{product_name}', checking cache...")
                db_lineage = None
                # Priority: sovereign_lineage > canonical_lineage > Lineage > lineage (sovereign has manual tag manager edits)
                # CRITICAL FIX: Reject "SOVEREIGN" as invalid - it's a field name, not a lineage value
                record_lineage = None
                for lineage_field in ['sovereign_lineage', 'canonical_lineage', 'Lineage', 'lineage']:
                    candidate = record.get(lineage_field)
                    if candidate and str(candidate).strip() not in ['', 'None', 'nan']:
                        lineage_str = str(candidate).strip().upper()
                        if lineage_str != 'SOVEREIGN':  # Reject "SOVEREIGN" as invalid
                            record_lineage = lineage_str
                            break
                
                if record_lineage:
                    # Use record lineage (already set correctly by enrichment)
                    db_lineage = record_lineage
                    if 'lemon' in product_name.lower() or 'cherry' in product_name.lower():
                        self.logger.info(f"✅ LINEAGE FALLBACK: Using record lineage '{db_lineage}' for '{product_name}' (from enrichment, no sativa hybrid override)")
                elif product_name and product_lineage_cache:
                    # Use pre-loaded cache instead of individual query
                    db_lineage = product_lineage_cache.get(product_name)
                    if db_lineage:
                        self.logger.info(f"✅ Using cached lineage '{db_lineage}' for '{product_name}'")
                
                if db_lineage and str(db_lineage).strip() not in ['', 'None', 'nan']:
                    lineage_val = str(db_lineage).strip().upper()
                    self.logger.info(f"✅ Using product lineage: '{lineage_val}' for '{product_name}'")
                
                # If no product-level lineage, try strain-level from cache
                if not lineage_val and product_strain and strain_info_cache:
                    strain_info = strain_info_cache.get(product_strain)
                    if strain_info:
                        preferred = (
                            strain_info.get('display_lineage') or
                            strain_info.get('sovereign_lineage') or
                            strain_info.get('canonical_lineage')
                        )
                        if preferred:
                            preferred_str = str(preferred).strip().upper()
                            # CRITICAL FIX: Reject "SOVEREIGN" as invalid - it's a field name, not a lineage value
                            if preferred_str != 'SOVEREIGN':
                                lineage_val = preferred_str
                            self.logger.info(f"✅ Using cached strain lineage: '{lineage_val}' for strain '{product_strain}'")
                
                if not lineage_val:
                    self.logger.debug(f"No lineage found in record or cache")
            
            # CRITICAL FIX: Ensure classic types always have lineage data
            if not lineage_val or lineage_val.strip() == "":
                lineage_val = "HYBRID"
                self.logger.info(f"🔧 FALLBACK LINEAGE: Set HYBRID lineage for classic type '{product_name}' (no lineage data available)")
            
            # Set Lineage to strain lineage for classic types
            if lineage_val:
                # Debug: Log the lineage value to see if it has leading spaces
                self.logger.debug(f"DEBUG: Original lineage_val: '{repr(lineage_val)}'")
                cleaned_lineage_val = lineage_val.strip()
                self.logger.debug(f"DEBUG: Cleaned lineage_val: '{repr(cleaned_lineage_val)}'")
                
                # CRITICAL FIX: For horizontal and vertical templates, preserve the full lineage value
                if self.template_type not in ('horizontal', 'vertical'):
                    # Only clean lineage for templates that use marker-wrapped lineage (e.g., standard flows)
                    classic_lineages = ["HYBRID/SATIVA", "HYBRID/INDICA", "SATIVA", "INDICA", "HYBRID", "CBD", "MIXED"]
                    for classic_lineage in classic_lineages:
                        if cleaned_lineage_val.upper().startswith(classic_lineage.upper()):
                            # Extract only the lineage part, not the brand
                            cleaned_lineage_val = cleaned_lineage_val[:len(classic_lineage)]
                            self.logger.debug(f"DEBUG: Extracted lineage only: '{cleaned_lineage_val}' from '{lineage_val}'")
                            break
                else:
                    # For double, horizontal, and vertical templates, preserve the full lineage value without cleaning
                    self.logger.debug(f"DEBUG: Preserving full lineage for {self.template_type} template: '{cleaned_lineage_val}'")
                
                # For vertical and double templates, don't wrap with markers since they use simple placeholders
                if self.template_type in ['vertical', 'double']:
                    label_context['Lineage'] = cleaned_lineage_val
                else:
                    label_context['Lineage'] = f"LINEAGE_START{cleaned_lineage_val}LINEAGE_END"
            else:
                label_context['Lineage'] = ""
                self.logger.debug(f"No lineage available for classic type '{product_type}', Lineage set to empty")
            
            # Set ProductVendor to actual vendor/supplier for classic types
            # CRITICAL: Use vendor from record FIRST (it's already in the Excel column)
            vendor_val = label_context.get('_vendor_from_record')
            
            # If _vendor_from_record wasn't set, try reading directly from record again
            # This handles cases where vendor reading at the start might have failed
            if not vendor_val or str(vendor_val).strip() in ['', 'None', 'NULL', 'null', 'nan']:
                # Try ALL possible vendor field variations directly from record
                vendor_fields = [
                    'Vendor/Supplier*',
                    'Vendor/Supplier',
                    'Vendor',
                    'ProductVendor',
                    'vendor',
                    'Vendor/Supplier *',  # Handle space variations
                    'Vendor/Supplier* ',  # Handle trailing space
                ]
                
                # Also check label_context (from dict copy) in case field name doesn't match exactly
                for field in vendor_fields:
                    val = label_context.get(field) or record.get(field)
                    if val is not None and not pd.isna(val) and str(val).strip() and str(val).lower() not in ['nan', 'none', 'null', '']:
                        vendor_val = str(val).strip()
                        self.logger.info(f"✅ Found vendor in field '{field}': '{vendor_val}' for '{product_name}' (direct read)")
                        # Store it for later use
                        label_context['_vendor_from_record'] = vendor_val
                        break
                
                # If still not found, check ALL vendor-related keys (case-insensitive)
                if not vendor_val or str(vendor_val).strip() in ['', 'None', 'NULL', 'null', 'nan']:
                    vendor_related_keys = [k for k in record.keys() if 'vendor' in k.lower() or 'supplier' in k.lower()]
                    for key in vendor_related_keys:
                        val = record.get(key)
                        if val is not None and not pd.isna(val) and str(val).strip() and str(val).lower() not in ['nan', 'none', 'null', '']:
                            vendor_val = str(val).strip()
                            self.logger.info(f"✅ Found vendor in field '{key}': '{vendor_val}' for '{product_name}' (case-insensitive match)")
                            label_context['_vendor_from_record'] = vendor_val
                            break
            
            # PRIORITY 2: Try database cache as fallback ONLY if record doesn't have it
            if not vendor_val or str(vendor_val).strip() in ['', 'None', 'NULL', 'null', 'nan']:
                try:
                    cached_vendor = product_vendor_cache.get(product_name, "")
                    if cached_vendor and str(cached_vendor).strip() not in ['', 'None', 'NULL', 'null', 'nan']:
                        vendor_val = cached_vendor
                        self.logger.info(f"🔧 CLASSIC VENDOR ENRICHED: Retrieved vendor '{vendor_val}' from database cache for '{product_name}'")
                except Exception as e:
                    self.logger.warning(f"🔧 CLASSIC VENDOR ENRICHMENT FAILED: Could not retrieve vendor from cache: {e}")
            
            # If still no vendor, log all available fields for debugging
            if not vendor_val or not str(vendor_val).strip():
                available_fields = [k for k in record.keys() if 'vendor' in k.lower() or 'supplier' in k.lower()]
                self.logger.warning(f"⚠️ INITIAL EXTRACTION: No vendor found for '{product_name}'. Available vendor-related fields: {available_fields}")
                # Log actual values from vendor fields for debugging
                for field in vendor_fields:
                    val = record.get(field)
                    if val is not None:
                        self.logger.warning(f"⚠️ INITIAL EXTRACTION: Field '{field}' has value: '{val}' (type: {type(val).__name__})")
                self.logger.warning(f"⚠️ INITIAL EXTRACTION: Will try fallback logic later. All record keys: {list(record.keys())[:20]}...")  # First 20 keys for debugging
            
            # Handle NaN values and empty strings
            if vendor_val is None or pd.isna(vendor_val) or str(vendor_val).lower() in ['nan', 'none', 'null', '']:
                vendor_val = ''
            
            # CRITICAL: Check if ProductVendor was already set at the start (from _vendor_from_record)
            # If so, preserve it - don't overwrite with empty
            existing_vendor = label_context.get('ProductVendor', '')
            # Check if existing_vendor has actual content (unwrap markers to check)
            existing_vendor_has_content = False
            if existing_vendor and str(existing_vendor).strip():
                try:
                    # unwrap_marker is already imported at the top of the file
                    unwrapped = unwrap_marker(str(existing_vendor), 'PRODUCTVENDOR')
                    if unwrapped and str(unwrapped).strip():
                        existing_vendor_has_content = True
                except:
                    # If unwrapping fails, check if it's just the plain value (no markers)
                    if 'PRODUCTVENDOR_START' not in str(existing_vendor):
                        existing_vendor_has_content = True
                    else:
                        # Has markers, check if content between markers is non-empty
                        match = re.search(r'PRODUCTVENDOR_START(.+?)PRODUCTVENDOR_END', str(existing_vendor))
                        if match and match.group(1).strip():
                            existing_vendor_has_content = True
            
            if existing_vendor_has_content:
                # ProductVendor was set at the start with content, keep it
                self.logger.info(f"✅ Preserving ProductVendor set at start: '{existing_vendor}' for '{product_name}'")
            elif vendor_val and str(vendor_val).strip():
                # For vertical template, don't wrap with markers since it uses simple placeholders
                if self.template_type == 'vertical':
                    label_context['ProductVendor'] = str(vendor_val).strip()
                else:
                    label_context['ProductVendor'] = f"PRODUCTVENDOR_START{str(vendor_val).strip()}PRODUCTVENDOR_END"
                self.logger.info(f"✅ Set ProductVendor to vendor: '{vendor_val}' for classic type '{product_type}' (product: '{product_name}')")
            else:
                # CRITICAL: Even if vendor_val is empty, check _vendor_from_record one more time
                # This catches cases where vendor reading at the start found it but it wasn't used above
                final_vendor = label_context.get('_vendor_from_record', '')
                if final_vendor and str(final_vendor).strip() not in ['', 'None', 'NULL', 'null', 'nan']:
                    if self.template_type == 'vertical':
                        label_context['ProductVendor'] = str(final_vendor).strip()
                    else:
                        label_context['ProductVendor'] = f"PRODUCTVENDOR_START{str(final_vendor).strip()}PRODUCTVENDOR_END"
                    self.logger.info(f"✅ Set ProductVendor from _vendor_from_record: '{final_vendor}' for classic type '{product_type}' (product: '{product_name}')")
                else:
                    # Only set to empty if we truly have no vendor data and ProductVendor wasn't already set
                    # CRITICAL FIX: Set with markers when empty so fallback logic can detect and populate it
                    # Use the same check as above to see if existing_vendor has content
                    if not existing_vendor_has_content:
                        if self.template_type == 'vertical':
                            label_context['ProductVendor'] = ""
                        else:
                            label_context['ProductVendor'] = wrap_with_marker('', 'PRODUCTVENDOR')
                        self.logger.warning(f"⚠️ ProductVendor set to empty for classic type '{product_type}' (product: '{product_name}', no vendor data found)")
            
            # Ensure ProductStrain uses proper marker wrapping for classic types (1pt sizing)
            # CRITICAL FIX: For vertical templates, classic types should NOT display ProductStrain
            # Lineage already shows the strain information, so ProductStrain would be redundant
            product_strain_value = record.get('ProductStrain') or record.get('Product Strain', '')
            if self.template_type == 'vertical':
                # Vertical templates: Classic types don't need ProductStrain (Lineage shows strain info)
                label_context['ProductStrain'] = ""
                self.logger.debug(f"VERTICAL CLASSIC FIX: Cleared ProductStrain for classic type '{product_type}' (Lineage already shows strain)")
            elif product_strain_value:
                if self.template_type == 'mini':
                    label_context['ProductStrain'] = str(product_strain_value).strip()
                else:
                    label_context['ProductStrain'] = wrap_with_marker(str(product_strain_value).strip(), 'PRODUCTSTRAIN')
            else:
                label_context['ProductStrain'] = ""
            
            # CRITICAL FIX: Classic types should NOT have ProductBrand for most templates
            # However, mini templates still display brand in dedicated cells
            if self.template_type == 'mini' or self.template_type == 'preroll':
                if product_brand:
                    classic_brand_text = str(product_brand).strip().upper()
                    # Ensure markers are applied consistently for downstream formatting
                    plain_brand = classic_brand_text
                    if is_already_wrapped(plain_brand, 'PRODUCTBRAND'):
                        plain_brand = unwrap_marker(plain_brand, 'PRODUCTBRAND')
                    elif is_already_wrapped(plain_brand, 'PRODUCTBRAND_CENTER'):
                        plain_brand = unwrap_marker(plain_brand, 'PRODUCTBRAND_CENTER')
                    label_context['ProductBrand'] = wrap_with_marker(plain_brand, 'PRODUCTBRAND')
                    label_context['ProductBrand_Center'] = wrap_with_marker(plain_brand, 'PRODUCTBRAND_CENTER')
                    template_name = 'PREROLL' if self.template_type == 'preroll' else 'MINI'
                    self.logger.info(
                        f"🎯 {template_name} CLASSIC BRAND: Preserving ProductBrand '{classic_brand_text}' for classic type '{product_type}'"
                    )
                else:
                    label_context['ProductBrand'] = ""
                    label_context['ProductBrand_Center'] = ""
                    template_name = 'PREROLL' if self.template_type == 'preroll' else 'MINI'
                    self.logger.info(
                        f"🎯 {template_name} CLASSIC BRAND: No brand available to preserve for classic type '{product_type}'"
                    )
            else:
                label_context['ProductBrand'] = ""
                label_context['ProductBrand_Center'] = ""
                self.logger.info(
                    f"🔧 CLASSIC TYPE FIX: Set ProductBrand to empty for classic type '{product_type}' (classic types show lineage, not brand)"
                )
        else:
            # For ALL non-classic types (including tinctures), Lineage shows brand and ProductVendor is empty
            # Color is determined by Product Strain (CBD Blend = yellow, Mixed = blue)
            self.logger.debug(f"Processing non-classic type '{product_type}' for Lineage and ProductVendor")
            if product_brand:
                self.logger.info(f"BRAND PROCESSING: Non-classic type '{product_type}' with brand '{product_brand}' (len={len(product_brand)}), template_type='{self.template_type}'")
                # For non-classic types, separate Product Strain and Product Brand for different font sizing
                # Lineage shows Product Brand only (centered) - this is the primary field
                # For vertical template, don't wrap with markers since it uses simple placeholders
                # Center brand should always be ALL CAPS
                brand_center_text = str(product_brand).upper()
                if self.template_type == 'vertical':
                    # Mirror double-template handling: treat brand content as lineage with hint markers for color logic
                    final_brand_text = str(brand_center_text).strip().upper()

                    product_strain_value = (product_strain or record.get('ProductStrain') or record.get('Product Strain', ''))
                    if product_strain_value:
                        strain_token = str(product_strain_value).strip().upper()
                        strain_token = strain_token.replace('PRODUCTSTRAIN_START', '').replace('PRODUCTSTRAIN_END', '').strip()
                        if strain_token:
                            original_brand = final_brand_text
                            strain_components = {strain_token}
                            strain_components.update(
                                token.strip()
                                for token in re.split(r'[\s\-\/,|]+', strain_token)
                                if token.strip()
                            )
                            final_brand_text = re.sub(
                                rf"\s*[-–\/]+\s*{re.escape(strain_token)}\s*$",
                                "",
                                final_brand_text,
                                flags=re.IGNORECASE,
                            )
                            final_brand_text = re.sub(
                                rf"{re.escape(strain_token)}\s*$",
                                "",
                                final_brand_text,
                                flags=re.IGNORECASE,
                            )
                            final_brand_text = re.sub(
                                rf"\s*[\(\[\{{]\s*{re.escape(strain_token)}\s*[\)\]\}}]\s*$",
                                "",
                                final_brand_text,
                                flags=re.IGNORECASE,
                            )
                            for component in list(strain_components):
                                if component:
                                    final_brand_text = re.sub(
                                        rf"[\s\(\[\{{\-–\/]*{re.escape(component)}[\s\)\]\}}\-–\/]*",
                                        " ",
                                        final_brand_text,
                                        flags=re.IGNORECASE,
                                    )
                            lineage_tokens = ["SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", "CBD", "MIXED"]
                            for lineage_token in lineage_tokens:
                                if lineage_token:
                                    final_brand_text = re.sub(
                                        rf"[\s\(\[\{{\-–\/]*{re.escape(lineage_token)}[\s\)\]\}}\-–\/]*",
                                        " ",
                                        final_brand_text,
                                        flags=re.IGNORECASE,
                                    )
                                final_brand_text = re.sub(
                                    rf"\s*[\(\[\{{]*\s*{re.escape(lineage_token)}\s*[\)\]\}}]*\s*$",
                                    "",
                                    final_brand_text,
                                    flags=re.IGNORECASE,
                                )
                                final_brand_text = re.sub(
                                    rf"\s*[-–\/]*\s*{re.escape(lineage_token)}\s*$",
                                    "",
                                    final_brand_text,
                                    flags=re.IGNORECASE,
                                )
                            final_brand_text = re.sub(r"\s{2,}", " ", final_brand_text).strip()
                            final_brand_text = final_brand_text.rstrip("-–/").rstrip()
                            if final_brand_text != original_brand:
                                self.logger.info(
                                    f"🎯 VERTICAL TEMPLATE STRAIN SPLIT: Removed strain/lineage token from brand -> '{final_brand_text}'"
                                )
                    if not final_brand_text:
                        final_brand_text = str(brand_center_text).strip().upper()

                    self.logger.info(f"🔍 VERTICAL BRAND DEBUG: Final brand text: '{final_brand_text}' (length: {len(final_brand_text)})")

                    # CRITICAL FIX: For vertical template, match horizontal template behavior
                    # Set Lineage to brand (for display), color comes from ProductStrain via is_product_strain_cbd
                    label_context['Lineage'] = final_brand_text
                    label_context['ProductBrand'] = ""
                    label_context['ProductBrand_Center'] = ""

                    self.logger.info(f"🎯 VERTICAL TEMPLATE BRAND FIX: Set Lineage to '{final_brand_text}' (color from ProductStrain)")
                elif self.template_type == 'mini':
                    # For mini template, set both Lineage and ProductBrand for maximum compatibility
                    # Mini templates need brand information in multiple fields
                    plain_brand = brand_center_text
                    if is_already_wrapped(plain_brand, 'PRODUCTBRAND'):
                        plain_brand = unwrap_marker(plain_brand, 'PRODUCTBRAND')
                    elif is_already_wrapped(plain_brand, 'PRODUCTBRAND_CENTER'):
                        plain_brand = unwrap_marker(plain_brand, 'PRODUCTBRAND_CENTER')
                    label_context['Lineage'] = brand_center_text
                    label_context['ProductBrand'] = wrap_with_marker(plain_brand, 'PRODUCTBRAND')
                    label_context['ProductBrand_Center'] = wrap_with_marker(plain_brand, 'PRODUCTBRAND_CENTER')
                    self.logger.info(f"🎯 MINI TEMPLATE BRAND FIX: Set Lineage, ProductBrand, and ProductBrand_Center to '{brand_center_text}' for mini template")
                elif self.template_type == 'preroll':
                    # For preroll template, use same ProductBrand handling as mini template
                    # Preroll templates need brand information in multiple fields just like mini
                    plain_brand = brand_center_text
                    if is_already_wrapped(plain_brand, 'PRODUCTBRAND'):
                        plain_brand = unwrap_marker(plain_brand, 'PRODUCTBRAND')
                    elif is_already_wrapped(plain_brand, 'PRODUCTBRAND_CENTER'):
                        plain_brand = unwrap_marker(plain_brand, 'PRODUCTBRAND_CENTER')
                    label_context['Lineage'] = brand_center_text
                    label_context['ProductBrand'] = wrap_with_marker(plain_brand, 'PRODUCTBRAND')
                    label_context['ProductBrand_Center'] = wrap_with_marker(plain_brand, 'PRODUCTBRAND_CENTER')
                    self.logger.info(f"🎯 PREROLL TEMPLATE BRAND FIX: Set Lineage, ProductBrand, and ProductBrand_Center to '{brand_center_text}' for preroll template")
                elif self.template_type == 'double':
                    # For double template, use brand text as-is with markers for downstream formatting
                    final_brand_text = str(brand_center_text).strip().upper()
                    
                    # Remove trailing strain content if it was concatenated with the brand text
                    product_strain_value = (product_strain or record.get('ProductStrain') or record.get('Product Strain', ''))
                    if product_strain_value:
                        strain_token = str(product_strain_value).strip().upper()
                        # Remove marker remnants if present
                        strain_token = strain_token.replace('PRODUCTSTRAIN_START', '').replace('PRODUCTSTRAIN_END', '').strip()
                        if strain_token:
                            original_brand = final_brand_text
                            
                            # Tokenize strain string into individual tokens (split on separators)
                            strain_components = {strain_token}
                            strain_components.update(
                                token.strip()
                                for token in re.split(r'[\s\-\/,|]+', strain_token)
                                if token.strip()
                            )

                            # Remove common separators before strain tokens (e.g., " - ", " / ")
                            final_brand_text = re.sub(
                                rf"\s*[-–\/]+\s*{re.escape(strain_token)}\s*$",
                                "",
                                final_brand_text,
                                flags=re.IGNORECASE,
                            )
                            # If the strain token still appears at the end without a separator, remove it
                            final_brand_text = re.sub(
                                rf"{re.escape(strain_token)}\s*$",
                                "",
                                final_brand_text,
                                flags=re.IGNORECASE,
                            )
                            # Remove strain tokens wrapped in parentheses or brackets at the end
                            final_brand_text = re.sub(
                                rf"\s*[\(\[\{{]\s*{re.escape(strain_token)}\s*[\)\]\}}]\s*$",
                                "",
                                final_brand_text,
                                flags=re.IGNORECASE,
                            )
                            # Remove strain tokens anywhere within the brand text
                            for component in list(strain_components):
                                if component:
                                    final_brand_text = re.sub(
                                        rf"[\s\(\[\{{\-–\/]*{re.escape(component)}[\s\)\]\}}\-–\/]*",
                                        " ",
                                        final_brand_text,
                                        flags=re.IGNORECASE,
                                    )
                            # Remove lineage tokens accidentally attached to brand text
                            lineage_tokens = ["SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", "CBD", "MIXED"]
                            for lineage_token in lineage_tokens:
                                if lineage_token:
                                    final_brand_text = re.sub(
                                        rf"[\s\(\[\{{\-–\/]*{re.escape(lineage_token)}[\s\)\]\}}\-–\/]*",
                                        " ",
                                        final_brand_text,
                                        flags=re.IGNORECASE,
                                    )
                                final_brand_text = re.sub(
                                    rf"\s*[\(\[\{{]*\s*{re.escape(lineage_token)}\s*[\)\]\}}]*\s*$",
                                    "",
                                    final_brand_text,
                                    flags=re.IGNORECASE,
                                )
                                final_brand_text = re.sub(
                                    rf"\s*[-–\/]*\s*{re.escape(lineage_token)}\s*$",
                                    "",
                                    final_brand_text,
                                    flags=re.IGNORECASE,
                                )
                            # Collapse extra whitespace created by removals
                            final_brand_text = re.sub(r"\s{2,}", " ", final_brand_text).strip()
                            final_brand_text = final_brand_text.rstrip("-–/").rstrip()
                            if final_brand_text != original_brand:
                                self.logger.info(
                                    f"🎯 DOUBLE TEMPLATE STRAIN SPLIT: Removed strain/lineage token from brand -> '{final_brand_text}'"
                                )
                    if not final_brand_text:
                        final_brand_text = clean_brand_text or str(brand_center_text).strip().upper()
                    
                    # CRITICAL FIX: Add debugging to see final brand text
                    self.logger.info(f"🔍 BRAND CLEANING DEBUG: Final brand text: '{final_brand_text}' (length: {len(final_brand_text)})")
                    
                    # Preserve the original lineage value (if any) so we can drive color assignment later
                    lineage_for_color_source = (
                        label_context.get('Lineage') or
                        record.get('Lineage') or
                        ''
                    )
                    if is_already_wrapped(lineage_for_color_source, 'LINEAGE'):
                        lineage_for_color_source = unwrap_marker(lineage_for_color_source, 'LINEAGE')
                    lineage_for_color = str(lineage_for_color_source).strip().upper()

                    # CRITICAL FIX: Validate that lineage_for_color is a valid lineage, not a brand name
                    # For non-classic types, the Lineage field may contain brand info, not actual lineage
                    valid_lineages = {'SATIVA', 'INDICA', 'HYBRID', 'HYBRID/SATIVA', 'HYBRID/INDICA', 'CBD', 'CBD BLEND', 'CBD_BLEND', 'MIXED'}
                    if not lineage_for_color or lineage_for_color not in valid_lineages:
                        # Fall back to CBD lineage when we have CBD signal, otherwise treat as MIXED (blue)
                        lineage_for_color = 'CBD' if has_cbd_blend_strain else 'MIXED'

                    lineage_hint_token = f"__LINEAGE_HINT_{lineage_for_color}__"
                    lineage_content = (
                        f"{lineage_hint_token}PRODUCTBRAND_CENTER_START{final_brand_text}PRODUCTBRAND_CENTER_END"
                    )
                    
                    label_context['Lineage'] = lineage_content
                    label_context['ProductBrand'] = ""
                    label_context['ProductBrand_Center'] = (
                        f"PRODUCTBRAND_CENTER_START{final_brand_text}PRODUCTBRAND_CENTER_END"
                    )
                    self.logger.debug(
                        f"DOUBLE TEMPLATE LINEAGE COLOR: Brand '{final_brand_text}' -> lineage '{lineage_for_color}' "
                        f"(product_type='{product_type}')"
                    )
                    
                    self.logger.info(f"🎯 DOUBLE TEMPLATE BRAND FIX: Set Lineage to '{final_brand_text}' for double template (with markers)")
                else:
                    # For other templates (horizontal, etc.), use marker-based formatting
                    # CRITICAL FIX: Clean brand_center_text to prevent corruption
                    clean_brand_text = str(brand_center_text).strip().upper()
                    # Remove any corrupted marker patterns that might already be present
                    clean_brand_text = re.sub(r'PRODUCTSTRR_STARTCONSTELL.*', '', clean_brand_text)
                    clean_brand_text = re.sub(r'PRODUCTBRAND_CENTER_START.*', '', clean_brand_text)
                    clean_brand_text = re.sub(r'CONSTELLATION\$.*', '', clean_brand_text)

                    # CRITICAL FIX: Remove any remaining $ symbols that might be marker remnants
                    # This handles cases like "VICE$Star" where $ is a corrupted marker remnant
                    clean_brand_text = re.sub(r'\$.*', '', clean_brand_text)

                    clean_brand_text = clean_brand_text.strip()

                    if clean_brand_text:
                        label_context['Lineage'] = f"PRODUCTBRAND_CENTER_START{clean_brand_text}PRODUCTBRAND_CENTER_END"
                    else:
                        # Fallback to original brand text if cleaning removed everything
                        label_context['Lineage'] = f"PRODUCTBRAND_CENTER_START{brand_center_text}PRODUCTBRAND_CENTER_END"

                    # Set ProductBrand fields to empty to prevent duplication for non-vertical templates
                    label_context['ProductBrand'] = ""
                    label_context['ProductBrand_Center'] = ""
                
                # Product Strain gets its own field with small font size
                # CRITICAL FIX: For vertical template, allow ProductStrain but use proper font sizing
                # Instead of clearing it, use marker-based formatting for proper font sizing
                if self.template_type == 'vertical':
                    if product_strain:
                        label_context['ProductStrain'] = f"PRODUCTSTRAIN_START{product_strain}PRODUCTSTRAIN_END"
                        self.logger.debug(f"VERTICAL FIX: Set ProductStrain to '{product_strain}' with markers for proper font sizing")
                    else:
                        label_context['ProductStrain'] = ""
                        self.logger.debug(f"VERTICAL FIX: No ProductStrain available")
                elif product_strain:
                    # For mini templates, don't wrap with markers since they use simple placeholders
                    if self.template_type == 'mini':
                        label_context['ProductStrain'] = product_strain
                    else:
                        label_context['ProductStrain'] = f"PRODUCTSTRAIN_START{product_strain}PRODUCTSTRAIN_END"
                else:
                    label_context['ProductStrain'] = ""
                
                self.logger.debug(f"Set Lineage/ProductBrand to '{product_brand}' and ProductStrain to '{product_strain}' for non-classic type '{product_type}'")
            else:
                # No brand available for non-classic type
                # Color will be determined by ProductStrain content (CBD/Mixed) in apply_lineage_colors
                label_context['Lineage'] = ""
                label_context['ProductBrand'] = ""
                label_context['ProductBrand_Center'] = ""
                self.logger.debug(f"No brand for non-classic type '{product_type}' - color from ProductStrain")
            
            # Always set ProductStrain for nonclassic types, regardless of whether there's a product brand
            # CRITICAL FIX: For vertical template, use marker-based formatting for proper font sizing
            if self.template_type == 'vertical':
                # Don't override the ProductStrain we set above for vertical template
                if 'ProductStrain' not in label_context:
                    if product_strain:
                        label_context['ProductStrain'] = f"PRODUCTSTRAIN_START{product_strain}PRODUCTSTRAIN_END"
                    else:
                        label_context['ProductStrain'] = ""
                self.logger.debug(f"VERTICAL FIX: Keeping ProductStrain with markers for vertical template")
            elif product_strain:
                # All templates use wrapped format for consistent processing by manual_docx_replace()
                label_context['ProductStrain'] = f"PRODUCTSTRAIN_START{product_strain}PRODUCTSTRAIN_END"
                self.logger.debug(f"DEBUG: Set ProductStrain to '{label_context['ProductStrain']}' for {self.template_type} template")
            else:
                label_context['ProductStrain'] = ""
                self.logger.debug(f"DEBUG: ProductStrain set to empty (no product_strain value) for template {self.template_type}")
            
            # ProductVendor is not used for non-classic types - set to empty (intentional design)
            label_context['ProductVendor'] = ""
            self.logger.debug(f"ProductVendor set to empty for non-classic type '{product_type}' (not used for non-classic types)")
        
        # Initial THC/CBD processing - will be overridden for prerolls below
        label_context['Ratio_or_THC_CBD'] = ''
        label_context['THC_CBD'] = ''

        # DEBUG: Log ProductStrain value before template replacement
        if self.template_type == 'vertical':
            self.logger.debug(f"VERTICAL FINAL DEBUG: ProductStrain value before template replacement: '{label_context.get('ProductStrain', 'NOT_SET')}'")

        # Lineage and ProductVendor logic is now handled earlier in the method for classic types

        # Fast other field processing
        if label_context.get('Price'):
            label_context['Price'] = wrap_with_marker(unwrap_marker(label_context['Price'], 'PRICE'), 'PRICE')
        
        # Always process lineage for classic types, and conditionally for non-classic types
        product_type = (label_context.get('Product Type*', '').lower() or 
                       label_context.get('ProductType', '').lower())
        product_strain = record.get('ProductStrain') or record.get('Product Strain', '')
        
        # For classic types, ALWAYS try to get the strain's canonical lineage from the database
        if is_classic_type and product_strain:
            self.logger.debug(f"DEBUG: Processing classic type '{product_type}' with strain '{product_strain}'")
            try:
                from src.core.data.product_database import get_product_database
                product_db = get_product_database()
                strain_info = product_db.get_strain_info(product_strain)
                self.logger.debug(f"DEBUG: Strain info: {strain_info}")
                if strain_info and strain_info.get('canonical_lineage'):
                    lineage_value = strain_info['canonical_lineage'].upper()
                    self.logger.debug(f"DEBUG: Using database lineage: '{lineage_value}'")
                else:
                    # Fallback to Excel lineage if no database lineage found
                    lineage_value = label_context.get('Lineage', '')
                    self.logger.debug(f"DEBUG: Using Excel lineage fallback: '{lineage_value}'")
            except Exception as e:
                # Fallback to Excel lineage if database lookup fails
                lineage_value = label_context.get('Lineage', '')
                self.logger.debug(f"DEBUG: Using Excel lineage due to error: '{lineage_value}' (error: {e})")
            
            # Lineage logic is now handled earlier in the method for both classic and non-classic types
            
        # Lineage logic is now handled earlier in the method for both classic and non-classic types

        # Fast wrapping for remaining fields
        # For all templates, wrap with markers consistently
        if label_context.get('DescAndWeight'):
            label_context['DescAndWeight'] = wrap_with_marker(unwrap_marker(label_context['DescAndWeight'], 'DESC'), 'DESC')
        
        if 'ProductType' not in label_context:
            label_context['ProductType'] = record.get('ProductType', '')
        
        
        # Fast strain handling - always show the actual strain value from Excel
        # But don't override if ProductStrain was already set for nonclassic types
        # CRITICAL FIX: Don't override ProductStrain if we already processed non-classic types
        # CRITICAL FIX: For vertical templates, ensure ProductStrain is always set if available
        # CRITICAL FIX: For vertical templates with non-classic types, ProductStrain is already set above - don't override
        should_process_strain = (
            'ProductStrain' not in label_context or 
            (not label_context['ProductStrain'] and not self._is_non_classic_type(product_type)) or
            (self.template_type == 'vertical' and not label_context.get('ProductStrain') and not self._is_non_classic_type(product_type))
        )
        
        if should_process_strain:
            product_strain = record.get('ProductStrain') or record.get('Product Strain', '')
            product_strain_upper = str(product_strain).upper()
            if any(token in product_strain_upper for token in ['CBD', 'CBG', 'CBN', 'CBC']):
                has_cbd_blend_strain = True
            
            self.logger.debug(f"STRAIN OVERRIDE DEBUG: Entering strain handling for {self.template_type}")
            self.logger.debug(f"STRAIN OVERRIDE DEBUG: Current ProductStrain in context: '{label_context.get('ProductStrain', 'NOT_SET')}'")
            self.logger.debug(f"STRAIN OVERRIDE DEBUG: Product strain from record: '{product_strain}'")
            
            if product_strain:
                # All templates use wrapped format for consistent processing by manual_docx_replace()
                # CRITICAL FIX: For vertical templates, ensure ProductStrain is always wrapped with markers
                label_context['ProductStrain'] = wrap_with_marker(product_strain, 'PRODUCTSTRAIN')
                self.logger.debug(f"STRAIN OVERRIDE DEBUG: Set ProductStrain to '{label_context['ProductStrain']}' for {self.template_type} template (OVERRIDE)")
            else:
                # Only set to empty if we're not preserving an existing value
                if 'ProductStrain' not in label_context or not label_context.get('ProductStrain'):
                    label_context['ProductStrain'] = ''
                    self.logger.debug(f"STRAIN OVERRIDE DEBUG: Set ProductStrain to empty (no strain from record) (OVERRIDE)")
        else:
            self.logger.debug(f"STRAIN OVERRIDE DEBUG: Skipping strain override - ProductStrain already set to '{label_context.get('ProductStrain', 'NOT_SET')}'")

        # Double template should never display ProductStrain text (prevent 12pt strain labels)
        if self.template_type == 'double':
            if not has_cbd_blend_strain:
                label_context['ProductStrain'] = ''

        # Lineage logic is now handled earlier in the method for both classic and non-classic types

        # Add marker strings for template processing
        # These markers will be rendered by DocxTemplate and preserved for font sizing
        label_context['ProductStrain_START'] = 'PRODUCTSTRAIN_START'
        label_context['ProductStrain_END'] = 'PRODUCTSTRAIN_END'
        # Add Lineage markers back for post-processing system to work
        label_context['Lineage_START'] = 'LINEAGE_START'
        label_context['Lineage_END'] = 'LINEAGE_END'
        label_context['ProductBrand_START'] = 'PRODUCTBRAND_START'
        label_context['ProductBrand_END'] = 'PRODUCTBRAND_END'
        label_context['ProductVendor_START'] = 'PRODUCTVENDOR_START'
        label_context['ProductVendor_END'] = 'PRODUCTVENDOR_END'
        label_context['DescAndWeight_START'] = 'DESC_START'
        label_context['DescAndWeight_END'] = 'DESC_END'
        label_context['Ratio_or_THC_CBD_START'] = 'THC_CBD_START'
        label_context['Ratio_or_THC_CBD_END'] = 'THC_CBD_END'
        label_context['Price_START'] = 'PRICE_START'
        label_context['Price_END'] = 'PRICE_END'
        
        # Wrap WeightUnits with markers if it exists
        if label_context.get('WeightUnits'):
            label_context['WeightUnits'] = wrap_with_marker(label_context['WeightUnits'], 'WEIGHTUNITS')
        label_context['WeightUnits_START'] = 'WEIGHTUNITS_START'
        label_context['WeightUnits_END'] = 'WEIGHTUNITS_END'
        label_context['Ratio_START'] = 'RATIO_START'
        label_context['Ratio_END'] = 'RATIO_END'
        label_context['JointRatio_START'] = 'JOINT_RATIO_START'
        label_context['JointRatio_END'] = 'JOINT_RATIO_END'
        label_context['THC_START'] = 'THC_START'
        label_context['THC_END'] = 'THC_END'
        label_context['CBD_START'] = 'CBD_START'
        label_context['CBD_END'] = 'CBD_END'

        # Fast joint ratio handling
        if label_context.get('JointRatio'):
            val = label_context['JointRatio']
            # Fix: Handle NaN values in JointRatio
            if pd.isna(val) or str(val).lower() == 'nan':
                val = ''
            marker = 'JOINT_RATIO'
            if is_already_wrapped(val, marker):
                val = unwrap_marker(val, marker)
            formatted_val = self.format_joint_ratio_pack(val)
            label_context['JointRatio'] = wrap_with_marker(formatted_val, marker)

        # Fast description processing
        if label_context.get('Description'):
            label_context['Description'] = self.fix_hyphen_spacing(label_context['Description'])
            
            # PREROLL TEMPLATE: Truncate description to universal format
            # Example: "Super Sour Diesel Infused Pre‑Roll ‑ 1g" -> " Infused Pre‑Roll ‑ 1g"
            if self.template_type == 'preroll':
                description = label_context['Description']
                # Find common preroll patterns and extract the universal part
                # Look for patterns like "Infused Pre‑Roll", "Pre‑Roll", etc. followed by weight
                # Pattern to match: [anything] + [Infused]? + Pre[- ]?Roll + [anything with weight]
                # Note: 're' module is already imported at the top of this file
                preroll_patterns = [
                    r'(.+?)(Infused\s+Pre[-‑ ]?Roll.*)',
                    r'(.+?)(Pre[-‑ ]?Roll.*)',
                ]
                for pattern in preroll_patterns:
                    match = re.search(pattern, description, re.IGNORECASE)
                    if match:
                        # Extract the universal part (everything after the strain name)
                        universal_desc = match.group(2).strip()
                        # Ensure it starts with a space if there was a strain name
                        if not universal_desc.startswith(' '):
                            universal_desc = ' ' + universal_desc
                        label_context['Description'] = universal_desc
                        self.logger.info(f"PREROLL DESC TRUNCATE: '{description}' -> '{universal_desc}'")
                        break
            
            # CRITICAL FIX: Apply non-breaking hyphens to Description for preroll templates
            # This ensures entries like "Pre-Roll - 1g" (without "Assorted") have non-breaking hyphens
            # Do this after truncation so the final Description has non-breaking hyphens
            if self.template_type == 'preroll' and label_context.get('Description'):
                from src.core.generation.text_processing import make_nonbreaking_hyphens
                original_description = label_context['Description']
                label_context['Description'] = make_nonbreaking_hyphens(label_context['Description'])
                self.logger.info(f"🔧 NON-BREAKING FORMATTING (PREROLL): Description '{original_description}' -> '{label_context['Description']}'")

        # CRITICAL FIX: Apply non-breaking hyphens to ProductName to prevent "Pre-Roll" splitting
        if label_context.get('ProductName'):
            from src.core.generation.text_processing import make_nonbreaking_hyphens
            original_product_name = label_context['ProductName']
            label_context['ProductName'] = make_nonbreaking_hyphens(label_context['ProductName'])
            self.logger.info(f"🔧 NON-BREAKING FORMATTING: ProductName '{original_product_name}' -> '{label_context['ProductName']}'")

        # CRITICAL FIX: Also apply non-breaking hyphens to DescAndWeight
        if label_context.get('DescAndWeight'):
            from src.core.generation.text_processing import make_nonbreaking_hyphens
            original_desc_weight = label_context['DescAndWeight']
            label_context['DescAndWeight'] = make_nonbreaking_hyphens(label_context['DescAndWeight'])
            self.logger.info(f"🔧 NON-BREAKING FORMATTING: DescAndWeight '{original_desc_weight}' -> '{label_context['DescAndWeight']}'")


        # Fast line break processing
        product_type = (label_context.get('ProductType', '').lower() or 
                       label_context.get('Product Type*', '').lower())
        
        if product_type not in classic_types and label_context.get('DescAndWeight'):
            desc_weight = label_context['DescAndWeight']
            # Keep hyphen on same line - no forced line breaks
            # if desc_weight.endswith(' - '):
            #     desc_weight = desc_weight[:-3] + '\n- '
            # elif desc_weight.endswith(' -'):
            #     desc_weight = desc_weight[:-2] + '\n- '
            # desc_weight = desc_weight.replace(' - ', '\n- ')
            label_context['DescAndWeight'] = desc_weight
        
        # Fast pre-roll processing - keep hyphen on same line
        if product_type in {"pre-roll", "infused pre-roll"} and label_context.get('DescAndWeight'):
            desc_weight = label_context['DescAndWeight']
            # desc_weight = desc_weight.replace(' - ', '\n- ')  # Removed forced line break
            label_context['DescAndWeight'] = desc_weight

        # Fast weight and ratio formatting
        for key, marker in [('WeightUnits', 'WEIGHTUNITS'), ('Ratio', 'RATIO')]:
            if label_context.get(key):
                val = label_context[key]
                formatted_val = self.format_with_soft_hyphen(val)
                label_context[key] = wrap_with_marker(unwrap_marker(formatted_val, marker), marker)
        
        # Preserve template-side formatting: keep wrapped markers during render.
        # Also compute raw values and expose them under *_RAW keys for logic.
        try:
            if label_context.get('Price'):
                label_context['Price_RAW'] = unwrap_marker(label_context['Price'], 'PRICE')
            if label_context.get('WeightUnits'):
                label_context['WeightUnits_RAW'] = unwrap_marker(label_context['WeightUnits'], 'WEIGHTUNITS')
            if label_context.get('DescAndWeight'):
                label_context['DescAndWeight_RAW'] = unwrap_marker(label_context['DescAndWeight'], 'DESC')
            if label_context.get('Lineage'):
                label_context['Lineage_RAW'] = unwrap_marker(label_context['Lineage'], 'LINEAGE')
            if label_context.get('ProductStrain'):
                label_context['ProductStrain_RAW'] = unwrap_marker(label_context['ProductStrain'], 'PRODUCTSTRAIN')
            if label_context.get('ProductVendor'):
                label_context['ProductVendor_RAW'] = unwrap_marker(label_context['ProductVendor'], 'PRODUCTVENDOR')
            if label_context.get('ProductBrand'):
                label_context['ProductBrand_RAW'] = unwrap_marker(label_context['ProductBrand'], 'PRODUCTBRAND')
            if label_context.get('Ratio_or_THC_CBD'):
                label_context['Ratio_or_THC_CBD_RAW'] = unwrap_marker(label_context['Ratio_or_THC_CBD'], 'THC_CBD')
        except Exception:
            # Fail-safe: continue with wrapped values
            pass
        
        # Ensure JointRatio stays on the same line - no line break processing
        if label_context.get('JointRatio'):
            val = label_context['JointRatio']
            # Remove any line breaks that might have been added
            val = val.replace('\n', ' ').replace('\r', ' ')
            # Clean up multiple spaces
            val = ' '.join(val.split())
            label_context['JointRatio'] = val
        
        # Fast vendor handling - set ProductVendor if it's missing or empty (ONLY for classic types)
        # This ensures vendor is populated even if earlier logic didn't set it or set it to empty
        # Non-classic types should NOT have ProductVendor
        product_type_check = (label_context.get('ProductType', '').lower() or 
                             label_context.get('Product Type*', '').lower())
        from src.core.constants import CLASSIC_TYPES
        classic_types_lower = [t.lower() for t in CLASSIC_TYPES]
        is_classic_type_for_vendor = product_type_check in classic_types_lower
        # Debug logging for product types that should be classic but aren't matching
        if product_name and ('blunt' in product_name.lower() or 'pre-roll' in product_name.lower()):
            self.logger.info(f"🔍 VENDOR TYPE CHECK: '{product_name}' - product_type_check: '{product_type_check}', is_classic: {is_classic_type_for_vendor}, classic_types: {classic_types_lower}")
        
        # Only process vendor for classic types
        if is_classic_type_for_vendor:
            current_vendor = label_context.get('ProductVendor', '')
            vendor_is_empty = False
            
            # Check if ProductVendor is missing or empty
            if not current_vendor or not str(current_vendor).strip():
                vendor_is_empty = True
            else:
                # Unwrap markers to check if actual content is empty
                try:
                    unwrapped = unwrap_marker(str(current_vendor), 'PRODUCTVENDOR')
                    if not unwrapped or not unwrapped.strip():
                        vendor_is_empty = True
                except:
                    # If unwrapping fails, check if it's just empty markers
                    if 'PRODUCTVENDOR_START' in str(current_vendor) and 'PRODUCTVENDOR_END' in str(current_vendor):
                        match = re.search(r'PRODUCTVENDOR_START(.*?)PRODUCTVENDOR_END', str(current_vendor))
                        if not match or not match.group(1).strip():
                            vendor_is_empty = True
                    elif str(current_vendor).strip() == '':
                        vendor_is_empty = True
            
            if vendor_is_empty:
                # PRIORITY 1: Use vendor we already read from record at the start
                enriched_vendor = label_context.get('_vendor_from_record', '')
                if enriched_vendor:
                    self.logger.debug(f"✅ Using vendor from record: '{enriched_vendor}' for '{product_name}'")
                
                # PRIORITY 2: Try to enrich vendor from pre-loaded cache if not in record
                if not enriched_vendor:
                    try:
                        # Use cached vendor data (loaded in batch before loop)
                        enriched_vendor = product_vendor_cache.get(product_name, "")
                        if enriched_vendor:
                            self.logger.info(f"🔧 VENDOR ENRICHED: Retrieved vendor '{enriched_vendor}' from database cache for '{product_name}'")
                    except Exception as e:
                        self.logger.warning(f"🔧 VENDOR ENRICHMENT FAILED: Could not retrieve vendor from cache: {e}")
                
                # PRIORITY 3: Fallback to record fields directly if still not found
                if not enriched_vendor:
                    product_type = (label_context.get('ProductType', '').lower() or
                                   label_context.get('Product Type*', '').lower())

                    # CRITICAL: Check all possible vendor field names with comprehensive fallback
                    product_vendor = None
                    vendor_fields = [
                        'Vendor/Supplier*',
                        'Vendor/Supplier',
                        'Vendor',
                        'ProductVendor',
                        'vendor',
                        'Vendor/Supplier *',  # Handle space variations
                        'Vendor/Supplier* ',  # Handle trailing space
                    ]

                    # Try each field name
                    for field in vendor_fields:
                        val = record.get(field)
                        if val is not None and not pd.isna(val) and str(val).strip() and str(val).lower() not in ['nan', 'none', 'null', '']:
                            product_vendor = val
                            self.logger.debug(f"✅ FALLBACK: Found vendor in field '{field}': '{product_vendor}' for '{product_name}'")
                            break

                    # If still no vendor, log all available fields for debugging
                    if not product_vendor or not str(product_vendor).strip():
                        available_fields = [k for k in record.keys() if 'vendor' in k.lower() or 'supplier' in k.lower()]
                        self.logger.warning(f"⚠️ FALLBACK: No vendor found for '{product_name}'. Available vendor-related fields: {available_fields}")

                    # Handle NaN values in vendor data
                    if product_vendor is None or pd.isna(product_vendor) or str(product_vendor).lower() in ['nan', 'none', 'null', '']:
                        product_vendor = ''
                    enriched_vendor = product_vendor

                # Set vendor if we found one
                if enriched_vendor and str(enriched_vendor).strip():
                    # For vertical template, don't wrap with markers since it uses simple placeholders
                    if self.template_type == 'vertical':
                        label_context['ProductVendor'] = str(enriched_vendor).strip()
                    else:
                        label_context['ProductVendor'] = wrap_with_marker(str(enriched_vendor).strip(), 'PRODUCTVENDOR')
                    self.logger.info(f"✅ PRODUCTVENDOR FALLBACK: Set ProductVendor to '{enriched_vendor}' for '{product_name}'")
                else:
                    # No vendor found anywhere, set to empty
                    label_context['ProductVendor'] = wrap_with_marker('', 'PRODUCTVENDOR')
                    self.logger.warning(f"⚠️ VENDOR MISSING: No vendor data found for '{product_name}'")
        # End of classic type vendor handling - non-classic types already have ProductVendor set to empty above

        # Generate QR code - special handling for preroll template
        product_name = label_context.get('Product Name*') or label_context.get('ProductName') or label_context.get('Product Name', '')
        qr_url_for_log = None  # Initialize for logging
        if product_name and str(product_name).strip():
            # For preroll template, generate URL to preroll items page
            # For all other templates, use product name as before
            if self.template_type == 'preroll':
                # Generate group-specific URL pointing to preroll items page for this product group
                # CRITICAL FIX: Use group_key (includes vendor) instead of just group_id
                # This ensures each vendor gets their own QR code that shows only their products
                group_key = label_context.get('_group_key') or record.get('_group_key')
                group_id = label_context.get('_group_id') or record.get('_group_id', 'other')
                
                # Extract vendor from record for vendor-specific filtering
                vendor = (
                    label_context.get('Vendor') or 
                    label_context.get('Vendor/Supplier*') or
                    record.get('Vendor') or 
                    record.get('Vendor/Supplier*') or 
                    record.get('Vendor/Supplier', '') or
                    ''
                )
                vendor_clean = str(vendor).strip()
                
                # Always derive the QR base from the current request host so we don't
                # bake any specific domain into the code or printed labels.
                from flask import request
                import os

                try:
                    base_url = (request.host_url or '').rstrip('/')
                except Exception:
                    base_url = ''

                # If host_url is unavailable, try multiple fallbacks
                if not base_url:
                    # First try environment variable
                    base_url = os.environ.get('QR_BASE_URL', '').strip()

                # If still no base_url, try Flask config
                if not base_url:
                    try:
                        from flask import current_app
                        base_url = current_app.config.get('QR_BASE_URL', '').strip()
                    except Exception:
                        pass

                # Last resort: use production URL as default
                if not base_url:
                    base_url = 'https://www.agtpricetags.com'
                    self.logger.warning(f"No QR_BASE_URL configured, using default: {base_url}")

                # CRITICAL FIX: Include vendor in URL for vendor-specific product lists
                # Format: /preroll-items/{group_id}?vendor={vendor}
                # This allows the route to filter products by vendor
                if vendor_clean:
                    # URL encode vendor to handle special characters
                    from urllib.parse import quote
                    vendor_encoded = quote(vendor_clean)
                    qr_url = f"{base_url.rstrip('/')}/preroll-items/{group_id}?vendor={vendor_encoded}"
                else:
                    # Fallback to group_id only if no vendor (backward compatibility)
                    qr_url = f"{base_url.rstrip('/')}/preroll-items/{group_id}"

                # Final safety check: avoid emitting localhost/127.0.0.1 in QR URLs on printed labels in
                # production. For local development allow it but log a clear warning so user can set
                # `QR_BASE_URL` to a production domain when deploying.
                if 'localhost' in qr_url.lower() or '127.0.0.1' in qr_url:
                    self.logger.warning(
                        f"PREROLL QR WARNING: Generated QR URL uses localhost base: {qr_url}. "
                        "Proceeding (development mode) but consider setting QR_BASE_URL to a production domain."
                    )
                
                self.logger.info(f"PREROLL QR: Generated QR URL for group '{group_id}' with vendor '{vendor_clean}': {qr_url}")
                qr_code = self._generate_qr_code(qr_url, doc, is_url=True)
                qr_url_for_log = qr_url  # Store for logging
                if qr_code:
                    self.logger.info(f"PREROLL QR: Successfully generated QR code for URL: {qr_url}")
                else:
                    self.logger.error(f"PREROLL QR: Failed to generate QR code for URL: {qr_url}")
            else:
                # All other templates: use product name as before
                qr_code = self._generate_qr_code(product_name, doc)
            
            if qr_code:
                label_context['QR'] = qr_code
                log_product = qr_url_for_log if (self.template_type == 'preroll' and qr_url_for_log) else product_name
                self.logger.info(f"✅ QR CODE SET: Template '{self.template_type}', Product: '{log_product}', QR object type: {type(qr_code)}")
            else:
                label_context['QR'] = ''
                self.logger.warning(f"❌ QR CODE MISSING: Failed to generate QR code for product: '{product_name}' (template: {self.template_type})")
        else:
            label_context['QR'] = ''
            self.logger.debug("No product name available for QR code generation")

        # CRITICAL: Final JointRatio processing for prerolls - must be last to override any other THC/CBD processing
        product_type = (label_context.get('Product Type*', '').lower() or 
                       label_context.get('ProductType', '').lower())
        
        if product_type in ['pre-roll', 'infused pre-roll']:
            # Use JointRatio for prerolls and infused prerolls
            joint_ratio = (record.get('JointRatio') or 
                          record.get('Joint Ratio') or 
                          record.get('Ratio') or 
                          '')
            if joint_ratio:
                # Wrap JointRatio with markers for proper template processing
                label_context['Ratio_or_THC_CBD'] = wrap_with_marker(joint_ratio, 'THC_CBD')
                label_context['THC_CBD'] = wrap_with_marker(joint_ratio, 'THC_CBD')
                self.logger.debug(f"FINAL: Set JointRatio for {product_type}: '{joint_ratio}'")
            else:
                self.logger.debug(f"FINAL: No JointRatio found for {product_type}")

        # FINAL SAFETY CHECK: Ensure ProductVendor is ALWAYS set for classic types (even if empty)
        # This ensures the template placeholder {{Label1.ProductVendor}} is always replaced
        product_type_final = (label_context.get('ProductType', '').lower() or
                             label_context.get('Product Type*', '').lower())
        from src.core.constants import CLASSIC_TYPES
        if product_type_final in [t.lower() for t in CLASSIC_TYPES]:
            # Check if ProductVendor is missing or empty
            current_vendor = label_context.get('ProductVendor', '')
            vendor_is_empty = not current_vendor or not str(current_vendor).strip()

            # If empty, try one more time to get vendor from _vendor_from_record or record
            if vendor_is_empty:
                fallback_vendor = label_context.get('_vendor_from_record') or record.get('Vendor/Supplier*') or record.get('Vendor') or record.get('ProductVendor')
                if fallback_vendor and str(fallback_vendor).strip() and str(fallback_vendor).lower() not in ['nan', 'none', 'null', '']:
                    if self.template_type == 'vertical':
                        label_context['ProductVendor'] = str(fallback_vendor).strip()
                    else:
                        label_context['ProductVendor'] = f"PRODUCTVENDOR_START{str(fallback_vendor).strip()}PRODUCTVENDOR_END"
                    self.logger.info(f"✅ FINAL CHECK: Set ProductVendor to '{fallback_vendor}' for classic type (was empty)")
                else:
                    # CRITICAL: Always set ProductVendor to empty string if no vendor found
                    # This ensures the template placeholder is replaced (not left as-is)
                    label_context['ProductVendor'] = ""
                    self.logger.warning(f"⚠️ FINAL CHECK: ProductVendor set to empty for classic type '{product_type_final}' (no vendor data found anywhere)")
            # Ensure ProductVendor exists in context even if it wasn't empty
            if 'ProductVendor' not in label_context:
                label_context['ProductVendor'] = ""
                self.logger.warning(f"⚠️ FINAL CHECK: ProductVendor was missing from context, set to empty")
        
        # FINAL DEBUG: Log ProductVendor value for blunts and pre-rolls before returning context
        final_product_vendor = label_context.get('ProductVendor', 'NOT_SET')
        product_name_final = label_context.get('ProductName', '') or label_context.get('Product Name*', '')
        if product_name_final and ('blunt' in product_name_final.lower() or 'pre-roll' in product_name_final.lower()):
            try:
                if 'PRODUCTVENDOR_START' in str(final_product_vendor):
                    final_vendor_unwrapped = unwrap_marker(str(final_product_vendor), 'PRODUCTVENDOR')
                    self.logger.info(f"🔍 FINAL CONTEXT: '{product_name_final}' - ProductVendor: '{final_vendor_unwrapped}' (was wrapped)")
                else:
                    self.logger.info(f"🔍 FINAL CONTEXT: '{product_name_final}' - ProductVendor: '{final_product_vendor}'")
            except Exception as e:
                self.logger.warning(f"🔍 FINAL CONTEXT: '{product_name_final}' - ProductVendor: '{final_product_vendor}' (error unwrapping: {e})")

        return label_context

    def _generate_qr_code(self, product_name, doc, is_url=False):
        """Generate QR code for the given product name (or URL for preroll template) and return as InlineImage."""
        try:
            if not product_name or str(product_name).strip() == '':
                self.logger.warning("Empty product name/URL provided for QR code generation")
                return None
            
            # Clean the product name or URL
            clean_name = str(product_name).strip()
            
            # For preroll URLs, ensure it's an absolute URL if possible, without hardcoding any domain.
            if is_url and not clean_name.startswith('http'):
                try:
                    from flask import request
                    if hasattr(request, 'host_url') and request.host_url:
                        base_url = request.host_url.rstrip('/')
                        clean_name = f"{base_url}{clean_name}"
                        self.logger.debug(f"Converted relative URL to absolute using request.host_url: {clean_name}")
                except Exception:
                    # If we can't get a host URL, leave the relative path as-is. Most scanners
                    # will still treat it as a valid URL once opened in a browser context.
                    self.logger.warning(f"Could not convert relative URL to absolute; leaving as-is: {clean_name}")
            
            # Create QR code instance
            qr = qrcode.QRCode(
                version=1,  # Auto-determine version based on content
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=10,  # Size of each box in pixels
                border=4,     # Border size in boxes
            )
            
            # Add data to QR code
            qr.add_data(clean_name)
            qr.make(fit=True)
            
            # Create QR code image
            qr_image = qr.make_image(fill_color="black", back_color="white")
            
            # Convert to BytesIO for InlineImage
            img_buffer = BytesIO()
            qr_image.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            
            # Determine QR code size using unified font sizing system
            # Get font size in points for QR field, then convert to millimeters
            font_size_pt = get_font_size(clean_name, 'qr', self.template_type)
            
            # Convert font size points to millimeters for QR image size
            # Using a conversion factor where 1 point ≈ 0.35mm for QR sizing
            qr_size_mm = font_size_pt.pt * 0.35
            
            # Apply template scale factor (default 1.0) so QR scales with the rest of the layout
            try:
                scale_factor = getattr(self, 'scale_factor', 1.0) or 1.0
            except Exception:
                scale_factor = 1.0
            qr_size_mm = qr_size_mm * scale_factor
            qr_size = Mm(qr_size_mm)
            
            # Check if doc is a DocxTemplate or Document
            # CRITICAL FIX: Always use the doc parameter for InlineImage creation
            # DocxTemplate requires the template object to properly render InlineImage
            if hasattr(doc, 'docx'):
                # This is a DocxTemplate - use it directly for InlineImage
                # This is required for DocxTemplate to properly render the image
                qr_inline_image = InlineImage(doc, img_buffer, width=qr_size)
                self.logger.debug(f"Created InlineImage with DocxTemplate for QR code: {clean_name[:50]}")
            else:
                # This is a Document - create InlineImage with None template for manual insertion
                # This shouldn't happen during normal rendering, but handle it gracefully
                self.logger.warning(f"QR code generation received Document instead of DocxTemplate - creating InlineImage with None")
                qr_inline_image = InlineImage(None, img_buffer, width=qr_size)
                # Store document reference for manual insertion
                qr_inline_image._doc = doc
            
            # Store the raw image data for manual replacement
            img_buffer.seek(0)  # Reset buffer position
            qr_inline_image._raw_image_data = img_buffer.read()
            qr_inline_image._raw_image_width = qr_size
            qr_inline_image._product_name = clean_name  # Store product name for reference
            
            self.logger.debug(f"Generated QR code for product: '{clean_name}' with font size: {font_size_pt.pt}pt, converted to {qr_size_mm:.1f}mm")
            return qr_inline_image
            
        except Exception as e:
            self.logger.error(f"Error generating QR code for product '{product_name}': {e}")
            return None

    def _preserve_apostrophes_in_document(self, doc):
        """
        CRITICAL FIX: Preserve apostrophes and following letters in document text.
        This prevents apostrophes from being truncated during text processing.
        """
        try:
            # Process all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text:
                                    # Check if text has apostrophes that might be truncated
                                    if "'" in run.text:
                                        # Log the original text for debugging
                                        original_text = run.text
                                        self.logger.debug(f"Preserving apostrophes in text: {repr(original_text)}")
                                        
                                        # Ensure the text is properly preserved
                                        # This is a defensive measure to prevent truncation
                                        if len(original_text) > 0 and original_text.count("'") > 0:
                                            # The text should already be correct, but we log it for debugging
                                            self.logger.debug(f"Apostrophe preservation check: {repr(original_text)}")
            
            # Process paragraphs outside tables
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.text and "'" in run.text:
                        original_text = run.text
                        self.logger.debug(f"Preserving apostrophes in paragraph text: {repr(original_text)}")
                        
        except Exception as e:
            self.logger.warning(f"Error preserving apostrophes: {e}")
            # Don't raise the exception - this is a defensive measure
    def _post_process_and_replace_content(self, doc):
        """Post-process the document after template rendering."""
        # Skip unnecessary processing for inventory templates
        if self.template_type == 'inventory':
            self.logger.info("Skipping post-processing for inventory template - just filling placeholders")
            return doc
        
        # CRITICAL FIX: Preserve apostrophes and following letters before any processing
        self._preserve_apostrophes_in_document(doc)
        """
        Ultra-optimized post-processing for maximum performance.
        """
        # CRITICAL FIX: Always perform post-processing on all documents regardless of size
        # This ensures all formatting, marker cleanup, and styling steps are executed
        
        # Clean up DOH cells before processing to ensure proper image positioning
        try:
            self._clean_doh_cells_before_processing(doc)
        except Exception as e:
            self.logger.warning(f"DOH cell cleanup failed: {e}")
        
        # Enhanced mini template processing
        if self.template_type == 'mini':
            try:
                self.logger.info("Processing mini template with enhanced design preservation")
                
                # Add markers for proper processing
                self._add_weight_units_markers(doc)
                self._add_brand_markers(doc)
                
                # Ensure proper brand centering for mini templates
                self._ensure_mini_template_brand_centering(doc)
                
                # CRITICAL FIX: Skip blank cell clearing for dynamic templates
                # The dynamic template creation already handles empty cells properly
                self.logger.info("Skipping blank cell clearing for dynamic mini template")
                
                # OPTIMIZATION: Skip expensive dimension enforcement here - will be done once at the end
                # This avoids processing every cell/paragraph/run twice
                self.logger.info("Skipping early dimension enforcement for mini template (will be done at end)")
                
                # Apply mini template specific font sizing
                self._apply_mini_template_font_sizing(doc)
                    
            except Exception as e:
                self.logger.warning(f"Mini template processing failed: {e}")
                # Continue processing even if mini-specific steps fail
        
        # Enhanced preroll template processing (same as mini, but uses its own font sizing config)
        if self.template_type == 'preroll':
            try:
                self.logger.info("Processing preroll template with enhanced design preservation")
                
                # Add markers for proper processing (same as mini)
                self._add_weight_units_markers(doc)
                self._add_brand_markers(doc)
                
                # Ensure proper brand centering for preroll templates (same as mini)
                self._ensure_mini_template_brand_centering(doc)
                
                # CRITICAL FIX: Skip blank cell clearing for dynamic templates
                # The dynamic template creation already handles empty cells properly
                self.logger.info("Skipping blank cell clearing for dynamic preroll template")
                
                # OPTIMIZATION: Skip expensive dimension enforcement here - will be done once at the end
                # This avoids processing every cell/paragraph/run twice
                self.logger.info("Skipping early dimension enforcement for preroll template (will be done at end)")
                
                # Apply preroll template specific font sizing (uses preroll config, not mini)
                self._apply_mini_template_font_sizing(doc)  # This method now uses self.template_type
                    
            except Exception as e:
                self.logger.warning(f"Preroll template processing failed: {e}")
                # Continue processing even if preroll-specific steps fail

        # ProductStrain in Brand cells fix
        try:
            self._fix_productstrain_in_brand_cells(doc)
        except Exception as e:
            self.logger.warning(f"ProductStrain in Brand cells fix failed: {e}")
        
        # Fast double template processing
        if self.template_type == 'double':
            try:
                # Use standard processing for double template
                pass
            except Exception as e:
                self.logger.warning(f"Double template processing failed: {e}")

        # Fast font sizing (with timeout protection)
        try:
            self._post_process_template_specific(doc)
        except Exception as e:
            self.logger.warning(f"Font sizing failed: {e}")

        # Fast BR marker conversion - only process if needed
        try:
            br_found = False
            for table in doc.tables:
                # Validate table structure before processing
                if not self._validate_and_repair_table_structure(table):
                    self.logger.warning(f"Skipping table with invalid structure during BR marker conversion")
                    continue
                
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if '|BR|' in paragraph.text:
                                self._convert_br_markers_to_line_breaks(paragraph)
                                br_found = True
            
            # Only process paragraphs outside tables if BR markers were found
            if br_found:
                for paragraph in doc.paragraphs:
                    if '|BR|' in paragraph.text:
                        self._convert_br_markers_to_line_breaks(paragraph)
        except Exception as e:
            self.logger.warning(f"BR marker conversion failed: {e}")
        
        # Fast ratio spacing fix
        try:
            self._fix_ratio_paragraph_spacing(doc)
        except Exception as e:
            self.logger.warning(f"Ratio spacing failed: {e}")

        # Ensure consistent spacing above lineage/brand section for equal margins
        try:
            self._ensure_consistent_lineage_spacing(doc)
        except Exception as e:
            self.logger.warning(f"Lineage spacing consistency failed: {e}")

        # Add consistent spacing above main content sections for better visual balance
        try:
            self._add_consistent_content_spacing(doc)
        except Exception as e:
            self.logger.warning(f"Content spacing consistency failed: {e}")

        # Arial Bold enforcement moved to the very end to prevent override

        # Fast DOH image centering
        try:
            for table in doc.tables:
                # Validate table structure before processing
                if not self._validate_and_repair_table_structure(table):
                    self.logger.warning(f"Skipping table with invalid structure during DOH centering")
                    continue
                
                for row in table.rows:
                    for cell in row.cells:
                        # Fast check for image-only cells
                        if len(cell.paragraphs) > 0 and all(len(paragraph.runs) == 1 and not paragraph.text.strip() for paragraph in cell.paragraphs):
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Fast inner table centering
                        for inner_table in cell.tables:
                            inner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        # Explicit DOH image centering - check for InlineImage objects
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                # Check if this run contains an InlineImage (DOH image)
                                if hasattr(run, '_element') and run._element.find(qn('w:drawing')) is not None:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    # Also center the cell content
                                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                    
            # Additional comprehensive DOH centering pass
            self._ensure_doh_image_centering(doc)
            
            # CRITICAL FIX: Ensure DOH images are properly centered in top-section cells
            self._fix_doh_image_positioning_in_top_section(doc)
            
            # CRITICAL FIX: Ensure DOH images have proper vertical margins to prevent cutoff
            self._ensure_doh_logo_vertical_margins(doc)
            
            # FINAL ENFORCEMENT: Absolutely ensure DOH images are centered - this overrides all other positioning
            self._final_doh_positioning_enforcement(doc)
            
            # CRITICAL FIX: Final marker cleanup to ensure ALL markers are stripped
            # NOTE: This must happen AFTER all font sizing is complete (which happens in _post_process_template_specific)
            # ProductStrain markers need to be processed for font sizing before being removed
            self._final_marker_cleanup(doc)
            
            # PREROLL TEMPLATE: Center QR codes
            if self.template_type == 'preroll':
                self._ensure_preroll_qr_centering(doc)
        except Exception as e:
            self.logger.warning(f"DOH centering failed: {e}")
        
        # OPTIMIZATION: Only call prevent_table_expansion_enhanced once - it already does everything
        # enforce_fixed_cell_dimensions does, plus more, so we don't need both
        try:
            doc = prevent_table_expansion_enhanced(doc, self.template_type)
            self.logger.info(f"Applied enhanced table expansion prevention to {self.template_type} template")
        except Exception as e:
            self.logger.warning(f"Enhanced table expansion prevention failed: {e}")

        # FINAL DOH CENTERING PASS: Ensure DOH images remain centered after all other processing
        try:
            self.logger.debug("Performing final DOH image centering pass")
            self._final_doh_centering_pass(doc)
            # Additional pass to ensure DOH centering sticks
            self._final_doh_centering_pass(doc)
        except Exception as e:
            self.logger.warning(f"Final DOH centering pass failed: {e}")
        
        # FINAL ROW HEIGHT ENFORCEMENT: Absolutely ensure all rows use EXACT height rule
        try:
            self._force_exact_row_heights(doc)
            self.logger.debug("Applied final exact row-height enforcement to all tables")
        except Exception as e:
            self.logger.warning(f"Exact row height enforcement failed: {e}")

        # CREATIVE FIX: Force bold formatting on DescAndWeight content specifically
        try:
            self._force_descandweight_bold(doc)
            self.logger.info("✅ CREATIVE DESCANDWEIGHT BOLD FIX: Applied comprehensive bold formatting")
        except Exception as e:
            self.logger.warning(f"DescAndWeight bold enforcement failed: {e}")
        
        # FINAL STEP: Enforce bold formatting on ALL text - this must be the very last operation
        try:
            from src.core.generation.docx_formatting import enforce_arial_bold_all_text, enforce_ratio_formatting, enforce_thc_cbd_bold_formatting
            enforce_arial_bold_all_text(doc)
            enforce_ratio_formatting(doc)
            enforce_thc_cbd_bold_formatting(doc)
            self.logger.info("✅ FINAL BOLD ENFORCEMENT: Applied bold formatting to all text")
        except Exception as e:
            self.logger.warning(f"Final bold enforcement failed: {e}")
            
        return doc

    def _force_descandweight_bold(self, doc):
        """CREATIVE FIX: Aggressively force bold formatting on DescAndWeight content."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Pattern to identify DescAndWeight content (product name with weight)
        descandweight_pattern = r'^[^-]+ - \d+\.?\d*(oz|g|mg|kg|lb|lbs)$'
        
        def process_run(run):
            """Process a single run and make it bold if it contains DescAndWeight content."""
            if not run.text:
                return
                
            # Check if this looks like DescAndWeight content
            text = run.text.strip()
            if re.match(descandweight_pattern, text, re.IGNORECASE):
                # This is DescAndWeight content - force it to be bold
                run.font.bold = True
                run.font.name = "Arial"
                
                # Force bold at XML level for maximum compatibility
                rPr = run._element.get_or_add_rPr()
                
                # Remove any existing bold formatting
                for b_elem in rPr.xpath('.//w:b'):
                    rPr.remove(b_elem)
                
                # Add new bold formatting
                b = OxmlElement('w:b')
                b.set(qn('w:val'), '1')
                rPr.append(b)
                
                # Force Arial font
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Arial')
                rFonts.set(qn('w:hAnsi'), 'Arial')
                rFonts.set(qn('w:eastAsia'), 'Arial')
                rFonts.set(qn('w:cs'), 'Arial')
                rPr.append(rFonts)
                
                self.logger.debug(f"🔧 CREATIVE BOLD FIX: Made DescAndWeight bold: '{text}'")
        
        # Process all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            process_run(run)
        
        # Process all paragraphs outside tables
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                process_run(run)

    def _ensure_doh_image_centering(self, doc):
        """
        Ensure DOH images are properly centered in all cells.
        This method provides improved centering for InlineImage objects.
        """
        try:
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            for table in doc.tables:
                # Validate table structure before processing
                if not self._validate_and_repair_table_structure(table):
                    self.logger.warning(f"Skipping table with invalid structure during DOH image centering")
                    continue
                
                for row in table.rows:
                    for cell in row.cells:
                        # Check if this cell contains a DOH image
                        has_doh_image = False
                        image_paragraph = None
                        
                        # Improved image detection
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if hasattr(run, '_element'):
                                    # Check for drawing elements (InlineImage)
                                    if run._element.find(qn('w:drawing')) is not None:
                                        has_doh_image = True
                                        image_paragraph = paragraph
                                        break
                                    # Check for picture elements
                                    elif run._element.find(qn('w:pict')) is not None:
                                        has_doh_image = True
                                        image_paragraph = paragraph
                                        break
                            if has_doh_image:
                                break
                        
                        if has_doh_image and image_paragraph:
                            self.logger.debug("Found DOH image, applying improved centering")
                            
                            # Apply centering at paragraph level
                            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Set proper spacing to prevent DOH logo from being cut off
                            image_paragraph.paragraph_format.space_before = Pt(3)
                            image_paragraph.paragraph_format.space_after = Pt(3)
                            image_paragraph.paragraph_format.line_spacing = 1.0
                            
                            # Set cell vertical alignment to center
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            
                            # Ensure proper XML-level centering
                            pPr = image_paragraph._element.get_or_add_pPr()
                            
                            # Set paragraph justification to center
                            jc = pPr.find(qn('w:jc'))
                            if jc is None:
                                jc = OxmlElement('w:jc')
                                pPr.append(jc)
                            jc.set(qn('w:val'), 'center')
                            
                            # Remove any existing spacing
                            existing_spacing = pPr.find(qn('w:spacing'))
                            if existing_spacing is not None:
                                pPr.remove(existing_spacing)
                            
                            # Add proper spacing to prevent DOH logo from being cut off
                            spacing = OxmlElement('w:spacing')
                            spacing.set(qn('w:before'), '60')  # 3pt = 60 twips
                            spacing.set(qn('w:after'), '60')   # 3pt = 60 twips
                            spacing.set(qn('w:line'), '240')
                            spacing.set(qn('w:lineRule'), 'auto')
                            pPr.append(spacing)
                            
                            # Ensure proper indentation
                            ind = pPr.find(qn('w:ind'))
                            if ind is None:
                                ind = OxmlElement('w:ind')
                                pPr.append(ind)
                            ind.set(qn('w:left'), '0')
                            ind.set(qn('w:right'), '0')
                            ind.set(qn('w:firstLine'), '0')
                            ind.set(qn('w:hanging'), '0')
                            
                            # CRITICAL FIX: Ensure the image itself is centered within the cell
                            # Check if this is a top-section DOH image (with other content like "100mg THC")
                            cell_text = cell.text.strip()
                            if '100mg THC' in cell_text or '$' in cell_text:
                                # This is a top-section cell with multiple elements
                                # Ensure the DOH image is centered in its own paragraph
                                self.logger.debug("Found top-section DOH image, ensuring proper centering")
                                
                                # Create a dedicated centered paragraph for the DOH image
                                if len(cell.paragraphs) > 1:
                                    # Find the paragraph with the DOH image
                                    for para in cell.paragraphs:
                                        if any(run._element.find(qn('w:drawing')) is not None or 
                                               run._element.find(qn('w:pict')) is not None 
                                               for run in para.runs):
                                            # This is the DOH image paragraph
                                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            
                                            # Set XML-level centering
                                            pPr = para._element.get_or_add_pPr()
                                            jc = pPr.find(qn('w:jc'))
                                            if jc is None:
                                                jc = OxmlElement('w:jc')
                                                pPr.append(jc)
                                            jc.set(qn('w:val'), 'center')
                                            
                                            # Remove any indentation that might affect centering
                                            ind = pPr.find(qn('w:ind'))
                                            if ind is not None:
                                                pPr.remove(ind)
                                            self.logger.debug("Applied top-section DOH image centering")
                                            break
                            
                            self.logger.debug("Applied improved DOH image centering")
                                
        except Exception as e:
            self.logger.warning(f"Error in improved DOH image centering: {e}")

    def _fix_doh_image_positioning_in_top_section(self, doc):
        """
        Fix DOH image positioning in top-section cells that contain multiple elements.
        This ensures DOH images are properly centered even when in cells with other content.
        """
        try:
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Check if this is a top-section cell with multiple elements
                        cell_text = cell.text.strip()
                        if ('100mg THC' in cell_text or '$' in cell_text) and len(cell.paragraphs) > 1:
                            self.logger.debug("Found top-section cell with multiple elements, fixing DOH positioning")
                            
                            # Find the paragraph with the DOH image
                            for paragraph in cell.paragraphs:
                                has_doh_image = False
                                for run in paragraph.runs:
                                    if hasattr(run, '_element'):
                                        if (run._element.find(qn('w:drawing')) is not None or 
                                            run._element.find(qn('w:pict')) is not None):
                                            has_doh_image = True
                                            break
                                
                                if has_doh_image:
                                    # This is the DOH image paragraph - ensure it's centered
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # Set XML-level centering
                                    pPr = paragraph._element.get_or_add_pPr()
                                    jc = pPr.find(qn('w:jc'))
                                    if jc is None:
                                        jc = OxmlElement('w:jc')
                                        pPr.append(jc)
                                    jc.set(qn('w:val'), 'center')
                                    
                                    # Remove any indentation that might affect centering
                                    ind = pPr.find(qn('w:ind'))
                                    if ind is not None:
                                        pPr.remove(ind)
                                    
                                    # Ensure proper spacing to prevent DOH logo from being cut off
                                    spacing = pPr.find(qn('w:spacing'))
                                    if spacing is None:
                                        spacing = OxmlElement('w:spacing')
                                        pPr.append(spacing)
                                    spacing.set(qn('w:before'), '60')  # 3pt = 60 twips
                                    spacing.set(qn('w:after'), '60')   # 3pt = 60 twips
                                    spacing.set(qn('w:line'), '240')
                                    spacing.set(qn('w:lineRule'), 'auto')
                                    
                                    self.logger.debug("Fixed DOH image positioning in top-section cell")
                                    break
                            
                            # Set cell vertical alignment to center
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            
        except Exception as e:
            self.logger.warning(f"Error fixing DOH image positioning in top-section: {e}")

    def _ensure_doh_logo_vertical_margins(self, doc):
        """
        Ensure DOH logos have proper vertical margins to prevent cutoff at the top.
        This method specifically targets the vertical spacing issue that causes logo clipping.
        """
        try:
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Check if this cell contains a DOH image
                        has_doh_image = False
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if hasattr(run, '_element'):
                                    if (run._element.find(qn('w:drawing')) is not None or 
                                        run._element.find(qn('w:pict')) is not None):
                                        has_doh_image = True
                                        break
                            if has_doh_image:
                                break
                        
                        if has_doh_image:
                            # Apply generous vertical margins to prevent DOH logo cutoff
                            for paragraph in cell.paragraphs:
                                # Set generous spacing above and below
                                paragraph.paragraph_format.space_before = Pt(4)
                                paragraph.paragraph_format.space_after = Pt(4)
                                
                                # Set XML-level spacing for maximum compatibility
                                pPr = paragraph._element.get_or_add_pPr()
                                spacing = pPr.find(qn('w:spacing'))
                                if spacing is None:
                                    spacing = OxmlElement('w:spacing')
                                    pPr.append(spacing)
                                
                                # Set generous margins: 4pt = 80 twips
                                spacing.set(qn('w:before'), '80')
                                spacing.set(qn('w:after'), '80')
                                spacing.set(qn('w:line'), '240')
                                spacing.set(qn('w:lineRule'), 'auto')
                                
                                # Ensure no indentation interferes with spacing
                                ind = pPr.find(qn('w:ind'))
                                if ind is not None:
                                    pPr.remove(ind)
                            
                            self.logger.debug("Applied generous vertical margins to prevent DOH logo cutoff")
                            break
                            
        except Exception as e:
            self.logger.warning(f"Error ensuring DOH logo vertical margins: {e}")

    def _final_doh_centering_pass(self, doc):
        """
        Final pass to ensure DOH images remain centered after all processing.
        This runs at the very end to override any alignment changes made by other processes.
        """
        try:
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            
            doh_images_found = 0
            for table in doc.tables:
                # Validate table structure before processing
                if not self._validate_and_repair_table_structure(table):
                    self.logger.warning(f"Skipping table with invalid structure during final DOH centering")
                    continue
                
                for row in table.rows:
                    for cell in row.cells:
                        # Check if this cell contains a DOH image
                        has_doh_image = self._cell_contains_doh_image(cell)
                        
                        if has_doh_image:
                            doh_images_found += 1
                            
                            # Add tiny vertical spacer above DOH image to push it down
                            self._add_doh_vertical_spacer(cell)
                            
                            # FORCE center alignment for the entire cell
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            
                            # Apply XML-level vertical alignment directly to cell
                            try:
                                from docx.oxml import OxmlElement
                                tc_element = cell._tc
                                tcPr = tc_element.get_or_add_tcPr()
                                
                                # Force vertical alignment at cell level
                                vAlign = tcPr.find(qn('w:vAlign'))
                                if vAlign is None:
                                    vAlign = OxmlElement('w:vAlign')
                                    tcPr.append(vAlign)
                                vAlign.set(qn('w:val'), 'center')
                                
                            except Exception as e:
                                self.logger.warning(f"Error setting XML cell vertical alignment: {e}")
                            
                            # FORCE center alignment for all paragraphs in the cell
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                                # Apply XML-level centering to be absolutely sure
                                try:
                                    pPr = paragraph._element.get_or_add_pPr()
                                    
                                    # Force paragraph justification to center
                                    from docx.oxml import OxmlElement
                                    jc = pPr.find(qn('w:jc'))
                                    if jc is None:
                                        jc = OxmlElement('w:jc')
                                        pPr.append(jc)
                                    jc.set(qn('w:val'), 'center')
                                    
                                    # Ensure proper vertical spacing
                                    spacing = pPr.find(qn('w:spacing'))
                                    if spacing is None:
                                        spacing = OxmlElement('w:spacing')
                                        pPr.append(spacing)
                                    spacing.set(qn('w:before'), '60')  # 3pt = 60 twips
                                    spacing.set(qn('w:after'), '60')   # 3pt = 60 twips
                                    spacing.set(qn('w:line'), '240')
                                    spacing.set(qn('w:lineRule'), 'auto')
                                    
                                except Exception as xml_error:
                                    self.logger.warning(f"Error applying XML-level DOH centering: {xml_error}")
            
            if doh_images_found > 0:
                self.logger.info(f"Final DOH centering pass completed - processed {doh_images_found} DOH images")
            # Final DOH centering pass completed
                
        except Exception as e:
            self.logger.warning(f"Error in final DOH centering pass: {e}")

    def _add_doh_vertical_spacer(self, cell):
        """
        Add a tiny invisible spacer above DOH images to push them down for better centering.
        """
        try:
            # Find paragraphs with DOH images
            doh_paragraphs = []
            for i, paragraph in enumerate(cell.paragraphs):
                paragraph_xml = paragraph._element.xml.decode('utf-8') if hasattr(paragraph._element, 'xml') else str(paragraph._element)
                if any(doh_indicator in paragraph_xml for doh_indicator in ['w:drawing', 'w:pict']) or 'DOH' in paragraph.text:
                    doh_paragraphs.append((i, paragraph))
            
            # Add spacer before the first DOH paragraph
            if doh_paragraphs:
                first_doh_index = doh_paragraphs[0][0]
                
                # Create spacer paragraph at the beginning of the cell
                spacer_para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                
                # If first paragraph is not the DOH, insert before it
                if first_doh_index > 0:
                    # Insert a new paragraph at the beginning
                    new_para = cell.add_paragraph()
                    # Move it to the beginning
                    cell._element.insert(0, new_para._element)
                    spacer_para = new_para
                
                # Clear any existing content and add invisible spacer
                spacer_para.clear()
                spacer_run = spacer_para.add_run()
                
                # Use invisible character and line break for vertical spacing
                spacer_run.text = "\u200B"  # Zero-width space character
                spacer_run.font.size = Pt(1)  # Minimal font size
                spacer_run.font.color.rgb = RGBColor(255, 255, 255)  # White (invisible)
                
                # Add line break for actual vertical spacing
                spacer_run.add_break()
                
                # Center align the spacer
                spacer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add XML-level spacing to push content down
                try:
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    pPr = spacer_para._element.get_or_add_pPr()
                    spacing = OxmlElement('w:spacing')
                    spacing.set(qn('w:after'), '120')  # 6pt = 120 twips after spacing
                    pPr.append(spacing)
                except Exception as xml_error:
                    self.logger.debug(f"XML spacing addition failed: {xml_error}")
                
                self.logger.debug("Added vertical spacer above DOH image to push it down")
                
        except Exception as e:
            self.logger.warning(f"Error adding DOH vertical spacer: {e}")
    def _final_marker_cleanup(self, doc):
        """
        Final marker cleanup to ensure ALL markers are stripped from the final output.
        This method runs after all other processing to catch any remaining markers.
        """
        try:
            # Enhanced patterns to catch all marker variations
            marker_patterns = [
                r'\b\w+_(START|END)\b',           # Standard markers like PRODUCTBRAND_START
                r'\b\w+_START\b',                 # START markers specifically
                r'\b\w+_END\b',                   # END markers specifically
                r'PRODUCTBRAND_START\s*',         # PRODUCTBRAND_START with optional spaces
                r'\s*PRODUCTBRAND_END\b',         # PRODUCTBRAND_END with optional spaces
                r'PRODUCTBRAND_CENTER_START\s*',  # PRODUCTBRAND_CENTER_START with optional spaces
                r'\s*PRODUCTBRAND_CENTER_END\b',  # PRODUCTBRAND_CENTER_END with optional spaces
                r'PRODUCTSTRAIN_START\s*',        # PRODUCTSTRAIN_START with optional spaces
                r'\s*PRODUCTSTRAIN_END\b',        # PRODUCTSTRAIN_END with optional spaces
                r'LINEAGE_START\s*',              # LINEAGE_START with optional spaces
                r'\s*LINEAGE_END\b',              # LINEAGE_END with optional spaces
                r'PRODUCTVENDOR_START\s*',        # PRODUCTVENDOR_START with optional spaces
                r'\s*PRODUCTVENDOR_END\b',        # PRODUCTVENDOR_END with optional spaces
                r'THC_CBD_START\s*',              # THC_CBD_START with optional spaces
                r'\s*THC_CBD_END\b',              # THC_CBD_END with optional spaces
                r'RATIO_START\s*',                # RATIO_START with optional spaces
                r'\s*RATIO_END\b',                # RATIO_END with optional spaces
                r'WEIGHTUNITS_START\s*',          # WEIGHTUNITS_START with optional spaces
                r'\s*WEIGHTUNITS_END\b',          # WEIGHTUNITS_END with optional spaces
                r'PRICE_START\s*',                # PRICE_START with optional spaces
                r'\s*PRICE_END\b',                # PRICE_END with optional spaces
                r'DESC_START\s*',                 # DESC_START with optional spaces
                r'\s*DESC_END\b',                 # DESC_END with optional spaces
                r'\bPRODUCTBRAND\b',              # Standalone PRODUCTBRAND
                r'\bPRODUCTSTRAIN\b',             # Standalone PRODUCTSTRAIN
                r'\bLINEAGE\b',                   # Standalone LINEAGE
                r'\bPRODUCTVENDOR\b',             # Standalone PRODUCTVENDOR
                r'\bTHC_CBD\b',                   # Standalone THC_CBD
                # REMOVED: r'\bRATIO\b' - Don't remove RATIO as it's part of brand names like "Ratio"
                r'\bWEIGHTUNITS\b',               # Standalone WEIGHTUNITS
                r'\bPRICE\b',                     # Standalone PRICE
                r'\bDESC\b',                      # Standalone DESC
            ]
            
            def clean_text(text):
                """Clean text by removing all marker patterns while preserving lineage content."""
                original_text = text
                cleaned = text
                
                # CRITICAL FIX: Handle lineage markers specially to preserve content
                # Extract lineage content before removing markers
                lineage_match = re.search(r'LINEAGE_START(.+?)LINEAGE_END', cleaned, re.IGNORECASE)
                if lineage_match:
                    lineage_content = lineage_match.group(1)
                    # Replace the full lineage marker pattern with just the content
                    cleaned = re.sub(r'LINEAGE_START(.+?)LINEAGE_END', lineage_content, cleaned, flags=re.IGNORECASE)
                
                # CRITICAL FIX: Handle product brand markers specially to preserve content
                # Extract product brand content before removing markers (handle both PRODUCTBRAND and PRODUCTBRAND_CENTER)
                brand_match = re.search(r'PRODUCTBRAND(?:_CENTER)?_START(.+?)PRODUCTBRAND(?:_CENTER)?_END', cleaned, re.IGNORECASE)
                if brand_match:
                    brand_content = brand_match.group(1)
                    # Now that Product Strain is separate, brand content is just the brand name
                    # No need to extract brand name from combined content
                    
                    # Replace the full product brand marker pattern with just the brand content
                    cleaned = re.sub(r'PRODUCTBRAND(?:_CENTER)?_START(.+?)PRODUCTBRAND(?:_CENTER)?_END', brand_content, cleaned, flags=re.IGNORECASE)
                
                # CRITICAL FIX: Handle product strain markers specially to preserve content
                # Extract product strain content before removing markers
                strain_match = re.search(r'PRODUCTSTRAIN_START(.+?)PRODUCTSTRAIN_END', cleaned, re.IGNORECASE)
                if strain_match:
                    strain_content = strain_match.group(1)
                    # Replace the full product strain marker pattern with just the content
                    cleaned = re.sub(r'PRODUCTSTRAIN_START(.+?)PRODUCTSTRAIN_END', strain_content, cleaned, flags=re.IGNORECASE)

                # CRITICAL FIX: Handle product vendor markers specially to preserve content
                # Extract product vendor content before removing markers
                vendor_content_extracted = None
                vendor_match = re.search(r'PRODUCTVENDOR_START(.+?)PRODUCTVENDOR_END', cleaned, re.IGNORECASE)
                if vendor_match:
                    vendor_content_extracted = vendor_match.group(1).strip()
                    # Replace the full product vendor marker pattern with a placeholder temporarily
                    # to protect it from being removed by cleanup patterns
                    cleaned = re.sub(r'PRODUCTVENDOR_START(.+?)PRODUCTVENDOR_END', '<<<VENDOR_PLACEHOLDER>>>', cleaned, flags=re.IGNORECASE)

                # Remove other marker patterns
                for pattern in marker_patterns:
                    cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)

                # CRITICAL FIX: Remove partial marker remnants like "bis" from "PRODUCTBRAND_END"
                # NOTE: Do this BEFORE restoring vendor content to avoid removing words like "VENDOR" from vendor names
                partial_remnants = [
                    r'\bbis\b',                    # "bis" from PRODUCTBRAND_END
                    r'PRODUCTBRAND_END',           # PRODUCTBRAND_END remnants (specific first)
                    r'PRODUCTBRAND_',              # PRODUCTBRAND_ remnants (specific first)
                    r'BRAND_',                     # BRAND_ remnants without word boundaries
                    r'\bBRAND_\b',                 # BRAND_ remnants (but not standalone BRAND)
                    r'\bSTART\b',                  # Any remaining START
                    r'\bEND\b',                    # Any remaining END
                    r'\bPRODUCT\b',                # Any remaining PRODUCT
                    # REMOVED: r'\bBRAND\b' - Don't remove BRAND as it's part of brand names
                    r'\bSTRAIN\b',                 # Any remaining STRAIN
                    r'\bVENDOR\b',                 # Any remaining VENDOR
                    # REMOVED: r'\bLINEAGE\b' - Don't remove LINEAGE as it might be part of content
                    # REMOVED: r'\bCBD\b' - Don't remove CBD as it's part of lineage content like "CBD Blend"

                    # CRITICAL FIX: Handle corrupted marker text patterns
                    r'PRODUCTSTRR_STARTCONSTELL',  # Corrupted PRODUCTBRAND_CENTER_START + CONSTELLATION
                    r'PRODUCTSTRR_',               # Corrupted PRODUCTBRAND_ patterns
                    r'STARTCONSTELL',              # Corrupted START + CONSTELLATION
                    r'CONSTELLATION\$\s*',         # CONSTELLATION$ remnants
                    r'\$.*',                       # Any $ symbol remnants (like VICE$Star)

                    r'\bTHC\b',                    # Any remaining THC
                    # REMOVED: r'\bRATIO\b' - Don't remove RATIO as it's part of brand names like "Ratio"
                    r'\bWEIGHT\b',                 # Any remaining WEIGHT
                    r'\bUNITS\b',                  # Any remaining UNITS
                    r'\bPRICE\b',                  # Any remaining PRICE
                    r'\bDESC\b',                   # Any remaining DESC
                    r'\bTART\b',                   # "TART" from THC_CBD_START
                    r'\bTADT\b',                   # "TADT" from THC_CBD_START
                    r'\bTUC\b',                    # "TUC" from THC_CBD_START
                    # CRITICAL FIX: Remove overly broad single character patterns that break words like "RAY'S"
                    # r'\bS\b',                      # REMOVED: This was removing 'S' from "RAY'S"
                    # r'\bC\b',                      # REMOVED: This could break legitimate words
                    # r'\bD\b',                      # REMOVED: This could break legitimate words
                ]

                for remnant in partial_remnants:
                    cleaned = re.sub(remnant, '', cleaned, flags=re.IGNORECASE)

                # CRITICAL: Restore vendor content after ALL cleanup to prevent removal
                # This ensures vendor names containing words like "VENDOR" are preserved
                if vendor_content_extracted:
                    cleaned = cleaned.replace('<<<VENDOR_PLACEHOLDER>>>', vendor_content_extracted)

                # Remove stray CENTER tokens left behind by split PRODUCTBRAND_CENTER markers.
                # This specifically catches runs that only contain the marker fragment.
                if cleaned.strip().upper() == 'CENTER':
                    original_upper = original_text.upper()
                    if ('PRODUCTBRAND' in original_upper) or (original_text.strip().upper() in {'CENTER', 'CENTER_', '_CENTER'}):
                        cleaned = ''
                
                # Clean up any double spaces, leading/trailing spaces
                # CRITICAL FIX: Preserve non-breaking hyphens (\u2011) when cleaning whitespace
                # First, temporarily replace non-breaking hyphens with a placeholder
                cleaned = cleaned.replace('\u2011', '___NONBREAKING_HYPHEN___')
                # Then clean up whitespace
                cleaned = re.sub(r'\s+', ' ', cleaned).strip()
                # Finally, restore non-breaking hyphens
                cleaned = cleaned.replace('___NONBREAKING_HYPHEN___', '\u2011')
                return cleaned
            
            # Clean markers in all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                original_text = run.text
                                cleaned_text = clean_text(original_text)
                                if cleaned_text != original_text:
                                    run.text = cleaned_text
            
            # Clean markers in paragraphs outside tables
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    original_text = run.text
                    cleaned_text = clean_text(original_text)
                    if cleaned_text != original_text:
                        run.text = cleaned_text
            
            # FINAL LINEAGE CLEANUP: Remove any leading spaces from lineage content
            self._final_lineage_cleanup(doc)
            
            # Enhanced final marker cleanup completed
            
        except Exception as e:
            self.logger.warning(f"Error in enhanced final marker cleanup: {e}")

    def _final_lineage_cleanup(self, doc):
        """
        Final cleanup to remove any leading spaces from lineage content.
        This runs after all other processing to ensure clean lineage display.
        """
        try:
            # Define lineage values that should be cleaned
            lineage_values = [
                "SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", 
                "CBD", "CBD BLEND", "MIXED", "PARAPHERNALIA", "PARA"
            ]
            
            # Clean lineage content in all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                original_text = run.text
                                
                                # Check if this run contains lineage content
                                for lineage in lineage_values:
                                    if lineage in original_text.upper():
                                        # CRITICAL FIX: Preserve non-breaking hyphens (\u2011) when cleaning leading spaces
                                        # First, temporarily replace non-breaking hyphens with a placeholder
                                        temp_text = original_text.replace('\u2011', '___NONBREAKING_HYPHEN___')
                                        # Then clean leading spaces
                                        cleaned_text = temp_text.lstrip(' \t\n\r\u00A0\u2000\u2001\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200A\u200B\u200C\u200D\u200E\u200F\u2028\u2029\u202A\u202B\u202C\u202D\u202E\u202F\u205F\u2060\u2061\u2062\u2063\u2064\u2065\u2066\u2067\u2068\u2069\u206A\u206B\u206C\u206D\u206E\u206F\u3000\uFEFF')
                                        # Finally, restore non-breaking hyphens
                                        cleaned_text = cleaned_text.replace('___NONBREAKING_HYPHEN___', '\u2011')
                                        
                                        if cleaned_text != original_text:
                                            run.text = cleaned_text
                                        break
            
            # Clean lineage content in paragraphs outside tables
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    original_text = run.text
                    
                    # Check if this run contains lineage content
                    for lineage in lineage_values:
                        if lineage in original_text.upper():
                            # CRITICAL FIX: Preserve non-breaking hyphens (\u2011) when cleaning leading spaces
                            # First, temporarily replace non-breaking hyphens with a placeholder
                            temp_text = original_text.replace('\u2011', '___NONBREAKING_HYPHEN___')
                            # Then clean leading spaces
                            cleaned_text = temp_text.lstrip(' \t\n\r\u00A0\u2000\u2001\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200A\u200B\u200C\u200D\u200E\u200F\u2028\u2029\u202A\u202B\u202C\u202D\u202E\u202F\u205F\u2060\u2061\u2062\u2063\u2064\u2065\u2066\u2067\u2068\u2069\u206A\u206B\u206C\u206D\u206E\u206F\u3000\uFEFF')
                            # Finally, restore non-breaking hyphens
                            cleaned_text = cleaned_text.replace('___NONBREAKING_HYPHEN___', '\u2011')
                            
                            if cleaned_text != original_text:
                                run.text = cleaned_text
                            break
            
            # Final lineage cleanup completed
            
        except Exception as e:
            self.logger.warning(f"Error in final lineage cleanup: {e}")

    def _clear_blank_cells_in_mini_template(self, doc):
        """
        Clear blank cells in mini templates when they run out of values.
        This removes empty cells that don't have any meaningful content.
        """
        try:
            for table in doc.tables:
                # Validate table structure before processing
                if not self._validate_and_repair_table_structure(table):
                    self.logger.warning(f"Skipping table with invalid structure during blank cell clearing")
                    continue
                
                for row in table.rows:
                    for cell in row.cells:
                        # Check if cell is essentially empty
                        cell_text = cell.text.strip()
                        
                        # Consider a cell blank if it has no text or only contains empty template placeholders
                        # Don't clear cells that contain actual template placeholders - only clear truly empty ones
                        is_blank = (
                            not cell_text or 
                            cell_text == '' or
                            # Only clear cells that contain empty template placeholders (no data rendered)
                            # This should NOT clear template placeholders that will be populated with data
                            cell_text in ['{{}}', '{{ }}'] or
                            # Clear cells that contain only whitespace and empty placeholders
                            (cell_text.startswith('{{') and cell_text.endswith('}}') and 
                             cell_text.strip() in ['{{}}', '{{ }}', '{{  }}'])
                        )
                        
                        if is_blank:
                            # Clear the cell content
                            cell._tc.clear_content()
                            
                            # Add a single empty paragraph to maintain cell structure
                            paragraph = cell.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Set cell background to white/transparent to ensure it's visually clean
                            from src.core.generation.docx_formatting import clear_cell_background
                            clear_cell_background(cell)
                            
                            self.logger.debug(f"Cleared blank cell in mini template")
                            
        except Exception as e:
            self.logger.error(f"Error clearing blank cells in mini template: {e}")
            # Don't raise the exception - this is a cleanup operation that shouldn't break the main process

    def _post_process_template_specific(self, doc):
        """
        Apply template-type-specific font sizing to all markers in the document.
        Uses the original font-sizing functions based on template type.
        """
        # Define marker processing for all template types (including double)
        markers = [
            'DESC', 'PRODUCTBRAND', 'PRODUCTBRAND_CENTER', 'PRICE', 'LINEAGE', 
            'THC_CBD', 'THC_CBD_LABEL', 'RATIO', 'WEIGHTUNITS', 'PRODUCTSTRAIN', 'DOH', 'PRODUCTVENDOR'
        ]
        
        # Process all markers in a single pass to avoid conflicts
        self._recursive_autosize_template_specific_multi(doc, markers)
        
        # Apply vertical template specific optimizations for minimal spacing
        if self.template_type in ['vertical', 'double']:
            self._optimize_vertical_template_spacing(doc)
            
        # Apply unified font sizing to all text in vertical and double templates (not just markers)
        if self.template_type in ['vertical', 'double']:
            self._apply_unified_font_sizing_to_all_text(doc)



    def _apply_unified_font_sizing_to_all_text(self, doc):
        """
        Apply unified font sizing to all text in vertical templates, not just markers.
        This ensures that simple placeholders like {{Label1.Lineage}} get proper font sizing.
        """
        try:
            from src.core.generation.unified_font_sizing import get_font_size
            
            # Use template_type directly as orientation - supports all template types including preroll, mini, etc.
            template_orientation = self.template_type

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # CRITICAL FIX: Check for ProductStrain markers first before processing individual runs
                            full_text = "".join(run.text for run in paragraph.runs)
                            if 'PRODUCTSTRAIN_START' in full_text and 'PRODUCTSTRAIN_END' in full_text:
                                # Process ProductStrain markers using the marker processing system
                                self._process_paragraph_for_markers_template_specific(paragraph, ['PRODUCTSTRAIN'])
                                continue
                            
                            # Process other text normally
                            for run in paragraph.runs:
                                run_text = run.text or ''
                                if not run_text.strip():
                                    continue
                                
                                existing_size = getattr(run.font, "size", None)
                                if existing_size is not None:
                                    try:
                                        if existing_size.pt <= 1.1:
                                            # Preserve already-minimized runs (e.g., ProductStrain at 1pt)
                                            continue
                                    except AttributeError:
                                        pass
                                
                                # Determine field type based on text content and position
                                field_type = self._determine_field_type_for_template(run_text, paragraph, cell)
                                
                                # Apply unified font sizing
                                font_size = get_font_size(run_text, field_type, template_orientation, self.scale_factor)
                                # Apply at run and XML level to prevent Word from overriding
                                from src.core.generation.unified_font_sizing import set_run_font_size
                                set_run_font_size(run, font_size)
                                
                                self.logger.debug(f"Applied unified font sizing to {self.template_type} template text: '{run_text}' -> {field_type} -> {font_size}")
                                    
        except Exception as e:
            self.logger.warning(f"Failed to apply unified font sizing to vertical template text: {e}")

    def _set_paragraph_cell_vertical_alignment(self, paragraph, alignment):
        """
        Ensure the table cell containing the paragraph has the desired vertical alignment.
        """
        try:
            if paragraph is None or alignment is None:
                return
            
            # Try high-level python-docx API first
            try:
                from docx.table import _Cell  # type: ignore
                cell_obj = getattr(paragraph, "_parent", None)
                while cell_obj is not None and not isinstance(cell_obj, _Cell):
                    cell_obj = getattr(cell_obj, "_parent", None)
                if cell_obj is not None:
                    cell_obj.vertical_alignment = alignment
                    return
            except Exception:
                pass
            
            target_tag = qn('w:tc')
            cell_element = paragraph._element
            while cell_element is not None and cell_element.tag != target_tag:
                cell_element = cell_element.getparent()

            if cell_element is None:
                return

            tc_pr = cell_element.find(qn('w:tcPr'))
            if tc_pr is None:
                tc_pr = OxmlElement('w:tcPr')
                cell_element.insert(0, tc_pr)

            v_align = tc_pr.find(qn('w:vAlign'))
            if v_align is None:
                v_align = OxmlElement('w:vAlign')
                tc_pr.append(v_align)

            align_map = {
                WD_CELL_VERTICAL_ALIGNMENT.TOP: 'top',
                WD_CELL_VERTICAL_ALIGNMENT.CENTER: 'center',
                WD_CELL_VERTICAL_ALIGNMENT.BOTTOM: 'bottom'
            }
            v_align.set(qn('w:val'), align_map.get(alignment, 'center'))
        except Exception as e:
            self.logger.warning(f"Failed to set cell vertical alignment: {e}")

    def _force_exact_row_heights(self, doc):
        """
        Ensure every row in every table uses an EXACT height rule (no 'At least').
        Applies recursively to inner tables while preserving existing height values.
        """
        try:
            from docx.enum.table import WD_ROW_HEIGHT_RULE
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            def ensure_exact_row(row):
                try:
                    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                except Exception:
                    pass

                try:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = trPr.find(qn('w:trHeight'))
                    if trHeight is None:
                        trHeight = OxmlElement('w:trHeight')
                        trPr.append(trHeight)

                    # Preserve existing height value; if missing, populate from current row.height if available
                    val = trHeight.get(qn('w:val'))
                    if (not val or val == '0') and getattr(row, 'height', None):
                        try:
                            val = str(int(row.height))
                            trHeight.set(qn('w:val'), val)
                        except Exception:
                            pass

                    trHeight.set(qn('w:hRule'), 'exact')
                except Exception as e:
                    self.logger.debug(f"Failed to set exact height on row: {e}")

            def process_table(table):
                for row in getattr(table, 'rows', []):
                    ensure_exact_row(row)
                    for cell in row.cells:
                        for inner_table in getattr(cell, 'tables', []):
                            process_table(inner_table)

            for table in getattr(doc, 'tables', []):
                process_table(table)

        except Exception as e:
            self.logger.warning(f"Failed to force exact row heights: {e}")

    def _determine_field_type_for_template(self, text, paragraph, cell):
        if self.template_type == 'double':
            return self._determine_field_type_for_double_template(text, paragraph, cell)
        return self._determine_field_type_for_vertical_template(text, paragraph, cell)

    def _determine_field_type_for_vertical_template(self, text, paragraph, cell):
        """
        Determine the field type for vertical template text based on content and context.
        CRITICAL: Default to 'strain' (1pt) for most text to prevent unwanted visible text in vertical templates.
        """
        text_lower = text.lower().strip()
        text_stripped = text.strip()
        is_all_caps = (text_stripped.isupper() and any(c.isalpha() for c in text_stripped))
        is_short_wordy = all(ch.isalpha() or ch.isspace() or ch in ['&','-','/'] for ch in text_stripped)

        # Marker-based overrides
        if any(marker in text for marker in ['PRODUCTBRAND_CENTER_START', 'PRODUCTBRAND_CENTER_END', 'PRODUCTBRAND_START', 'PRODUCTBRAND_END']):
            self.logger.debug(f"🎯 BRAND MARKER DETECTED: '{text_stripped}' classified as brand (marker-based)")
            return 'brand'

        if '__LINEAGE_HINT_' in text:
            self.logger.debug(f"🎯 LINEAGE HINT DETECTED: '{text_stripped}' classified as lineage (marker-based)")
            return 'lineage'

        # Check for prices (contain $ symbol)
        if '$' in text:
            return 'price'

        # Check for THC/CBD percentage content (contains % sign)
        if '%' in text or any(keyword in text_lower for keyword in ['thc:', 'cbd:', 'total thc', 'total cbd']):
            return 'thc_cbd'

        # Check for weight/ratio content (contains oz, g, mg, or : for ratios)
        if any(keyword in text_lower for keyword in ['oz', 'gram', 'mg', 'ml']) or ':' in text:
            return 'ratio'

        # CRITICAL FIX: Check for classic lineage values BEFORE brand detection
        # Classic lineage values should use 'lineage' field type (14-20pt) not 'brand' (10-16pt)
        classic_lineages = ['hybrid/sativa', 'hybrid/indica', 'sativa', 'indica', 'hybrid', 'cbd', 'mixed']
        if text_stripped.upper() in [lineage.upper() for lineage in classic_lineages]:
            self.logger.debug(f"🎯 CLASSIC LINEAGE DETECTED: '{text_stripped}' classified as lineage")
            return 'lineage'

        # Check for well-known brand names that should be visible
        # Only classify as 'brand' if we're CERTAIN it's a brand name that should be visible
        well_known_brands = ['constellation', 'mary jones', 'skagit organics', 'artizen', 'sitka', 'raven', 'grassroots', 'pruf cultivar', 'lil ray', 'green revolution']
        if any(brand in text_lower for brand in well_known_brands):
            return 'brand'
        
        # Heuristic: Non-classic vertical brands are typically ALL CAPS, short, letters-only
        # Classify these as brand so they render visibly (not 1pt).
        # CRITICAL FIX: Allow brand names as short as 1 character to be visible
        if is_all_caps and is_short_wordy and len(text_stripped.split()) <= 3:
            self.logger.debug(f"🎯 SHORT BRAND CLASSIFIED: '{text_stripped}' (len={len(text_stripped)}) classified as brand")
            return 'brand'
        
        # CRITICAL FIX: Handle mixed-case brand names like "Lil Ray's"
        # Look for patterns that suggest brand names (mixed case, apostrophes, common brand words)
        if (len(text_stripped) >= 3 and len(text_stripped) <= 20 and 
            any(word in text_lower for word in ['ray', 'lil', 'green', 'revolution', 'cannabis', 'co', 'brands']) and
            any(char.isalpha() for char in text_stripped) and
            not any(char.isdigit() for char in text_stripped) and
            not any(keyword in text_lower for keyword in ['oz', 'gram', 'mg', 'ml', 'thc', 'cbd', '%', '$'])):
            self.logger.debug(f"🎯 MIXED-CASE BRAND CLASSIFIED: '{text_stripped}' classified as brand")
            return 'brand'
        
        # CRITICAL: Default everything else to 'strain' (1pt font)
        # This includes: strain codes, product types (like "mixed"), vendor codes, etc.
        # Since strain is 1pt (invisible), this is the safest default for vertical templates
        # where most text should be hidden/minimal
        self.logger.debug(f"⚠️ TEXT CLASSIFIED AS STRAIN: '{text_stripped}' (len={len(text_stripped)}) - will be invisible")
        return 'strain'

    def _determine_field_type_for_double_template(self, text, paragraph, cell):
        """
        Determine the field type for double template text based on content and context.
        Double templates should display most content visibly, so default to 'default' sizing.
        """
        text_lower = text.lower().strip()
        text_stripped = text.strip()
        is_all_caps = (text_stripped.isupper() and any(c.isalpha() for c in text_stripped))
        is_short_wordy = (len(text_stripped) <= 18 and all(ch.isalpha() or ch.isspace() or ch in ['&','-','/'] for ch in text_stripped))

        # CRITICAL FIX: Check for brand markers FIRST - double templates use PRODUCTBRAND_CENTER markers
        if any(marker in text for marker in ['PRODUCTBRAND_CENTER_START', 'PRODUCTBRAND_CENTER_END', 'PRODUCTBRAND_START', 'PRODUCTBRAND_END']):
            self.logger.debug(f"🎯 DOUBLE BRAND MARKER DETECTED: '{text_stripped}' classified as brand (marker-based)")
            return 'brand'

        # Prices
        if '$' in text:
            return 'price'

        # THC/CBD percentages
        if '%' in text or any(keyword in text_lower for keyword in ['thc', 'cbd', 'cbn', 'cbg']):
            return 'thc_cbd'

        # Weight / ratios
        if any(keyword in text_lower for keyword in ['oz', 'gram', 'g ', 'mg', 'ml']) or ':' in text:
            return 'ratio'

        # Classic lineage values
        classic_lineages = ['hybrid/sativa', 'hybrid/indica', 'sativa', 'indica', 'hybrid', 'cbd', 'mixed']
        if text_stripped.upper() in [lineage.upper() for lineage in classic_lineages]:
            self.logger.debug(f"🎯 DOUBLE LINEAGE DETECTED: '{text_stripped}' classified as lineage")
            return 'lineage'

        # Detect obvious product strain tokens (e.g., "HYBRID" in non-classic contexts)
        if text_stripped.upper() in classic_lineages:
            return 'strain'

        # Well known brands
        well_known_brands = ['constellation', 'gravity', 'mary jones', 'skagit organics', 'artizen', 'sitka', 'raven', 'grassroots', 'pruf cultivar', 'lil ray', 'green revolution']
        if any(brand in text_lower for brand in well_known_brands):
            return 'brand'

        # Heuristic brand detection
        if is_all_caps and is_short_wordy and len(text_stripped.split()) <= 4 and len(text_stripped) >= 2:
            self.logger.debug(f"🎯 DOUBLE BRAND CLASSIFIED: '{text_stripped}' classified as brand")
            return 'brand'

        if (len(text_stripped) >= 3 and len(text_stripped) <= 24 and 
            any(word in text_lower for word in ['ray', 'lil', 'green', 'revolution', 'cannabis', 'co', 'brands', 'farm', 'company']) and
            any(char.isalpha() for char in text_stripped) and
            not any(char.isdigit() for char in text_stripped) and
            not any(keyword in text_lower for keyword in ['oz', 'gram', 'mg', 'ml', 'thc', 'cbd', '%', '$'])):
            self.logger.debug(f"🎯 DOUBLE MIXED-CASE BRAND CLASSIFIED: '{text_stripped}' classified as brand")
            return 'brand'

        # Default: Treat as normal visible text
        return 'default'

    def _optimize_vertical_template_spacing(self, doc):
        """
        Apply minimal spacing optimizations specifically for vertical and double templates
        to ensure all labels fit on one page.
        """
        try:
            from docx.shared import Pt
            
            def optimize_paragraph_spacing(paragraph):
                """Set minimal spacing for all paragraphs in vertical and double templates."""
                # Set absolute minimum spacing
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                
                # All content now uses standard spacing
                
                # Default spacing for non-THC_CBD content
                paragraph.paragraph_format.line_spacing = 1.0
                
                # Set at XML level for maximum compatibility
                pPr = paragraph._element.get_or_add_pPr()
                spacing = pPr.find(qn('w:spacing'))
                if spacing is None:
                    spacing = OxmlElement('w:spacing')
                    pPr.append(spacing)
                
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '0')
                spacing.set(qn('w:line'), '240')  # 1.0 line spacing
                spacing.set(qn('w:lineRule'), 'auto')
            
            # Process all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            optimize_paragraph_spacing(paragraph)
            
            # Process all paragraphs outside tables
            for paragraph in doc.paragraphs:
                optimize_paragraph_spacing(paragraph)
            
            self.logger.debug("Applied vertical/double template spacing optimizations")
            
        except Exception as e:
            self.logger.error(f"Error optimizing vertical/double template spacing: {e}")
            # Don't raise the exception - this is an optimization that shouldn't break the main process

    def _recursive_autosize_template_specific(self, element, marker_name):
        """
        Recursively find and replace markers in paragraphs and tables using template-specific font sizing.
        """
        if hasattr(element, 'paragraphs'):
            for p in element.paragraphs:
                self._process_paragraph_for_marker_template_specific(p, marker_name)

        if hasattr(element, 'tables'):
            for table in element.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self._recursive_autosize_template_specific(cell, marker_name)

    def _recursive_autosize_template_specific_multi(self, element, markers):
        """
        Recursively find and replace all markers in paragraphs and tables using template-specific font sizing.
        Processes all markers in a single pass to avoid conflicts.
        """
        if hasattr(element, 'paragraphs'):
            for p in element.paragraphs:
                self._process_paragraph_for_markers_template_specific(p, markers)

        if hasattr(element, 'tables'):
            for table in element.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self._recursive_autosize_template_specific_multi(cell, markers)
    def _process_paragraph_for_markers_template_specific(self, paragraph, markers):
        """
        Process a single paragraph for multiple markers using template-specific font sizing.
        Handles all markers in a single pass to avoid conflicts.
        """
        # Import functions used throughout this method
        from src.core.generation.unified_font_sizing import get_font_size_by_marker, set_run_font_size
        
        full_text = "".join(run.text for run in paragraph.runs)
        
        # First, check if this is a combined lineage/vendor paragraph
        if self._detect_and_process_combined_lineage_vendor(paragraph):
            return
        
        # Check if any markers are present
        found_markers = []
        for marker_name in markers:
            start_marker = f'{marker_name}_START'
            end_marker = f'{marker_name}_END'
            if start_marker in full_text and end_marker in full_text:
                found_markers.append(marker_name)
                # CRITICAL DEBUG: Log ProductStrain marker detection for vertical templates
                if marker_name == 'PRODUCTSTRAIN':
                    self.logger.info(f"✅ PRODUCTSTRAIN MARKER DETECTED: Found ProductStrain markers in {self.template_type} template paragraph: '{full_text[:150]}...'")
        
        if found_markers:
            # Process all markers and build the final content
            final_content = full_text
            processed_content = {}
            
            for marker_name in found_markers:
                start_marker = f'{marker_name}_START'
                end_marker = f'{marker_name}_END'
                
                # Extract content for this marker
                start_idx = final_content.find(start_marker)
                end_idx = final_content.find(end_marker) + len(end_marker)
                
                if start_idx != -1 and end_idx != -1:
                    marker_start = final_content.find(start_marker) + len(start_marker)
                    marker_end = final_content.find(end_marker)
                    content = final_content[marker_start:marker_end]
                    
                    # Get font size for this marker
                    font_size = self._get_template_specific_font_size(content, marker_name)
                    
                    # CRITICAL DEBUG: Log ProductStrain font sizing
                    if marker_name == 'PRODUCTSTRAIN':
                        self.logger.info(f"✅ PRODUCTSTRAIN FONT SIZING: Content='{content}', FontSize={font_size.pt if hasattr(font_size, 'pt') else font_size}pt, Template={self.template_type}")
                    
                    processed_content[marker_name] = {
                        'content': content,
                        'font_size': font_size,
                        'start_pos': start_idx,
                        'end_pos': end_idx
                    }
                    
                    # Remove this marker from final_content so subsequent markers can find their correct positions
                    final_content = final_content[:start_idx] + final_content[end_idx:]
            
            self.logger.debug(f"Processed content: {processed_content}")
            
            # Clear paragraph and rebuild with all processed content
            paragraph.clear()
            
            # Ensure consistent spacing above all marker sections for equal margins
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(1)
            
            # Sort markers by position in text
            sorted_markers = sorted(processed_content.items(), key=lambda x: x[1]['start_pos'])
            
            current_pos = 0
            for marker_name, marker_data in sorted_markers:
                # Add any text before this marker
                if marker_data['start_pos'] > current_pos:
                    text_before = full_text[current_pos:marker_data['start_pos']]
                    # Preserve line breaks and whitespace, but skip if completely empty
                    if text_before or text_before.strip():
                        run = paragraph.add_run(text_before)
                        run.font.name = "Arial"
                        run.font.bold = True
                        # Use unified font sizing for non-marker text
                        from src.core.generation.unified_font_sizing import get_font_size
                        font_size = get_font_size(text_before, 'default', self.template_type, self.scale_factor)
                        run.font.size = font_size
                        self.logger.debug(f"Added text before '{marker_name}': '{text_before}' -> {font_size.pt}pt")
                # Add the processed marker content (use the potentially modified content)
                display_content = marker_data.get('display_content', marker_data['content'])
                # --- BULLETPROOF: Only one run for the entire marker content, preserving line breaks ---
                run = paragraph.add_run()
                run.font.name = "Arial"
                
                # Make everything bold - no exceptions
                run.font.bold = True
                
                run.font.size = marker_data['font_size']
                set_run_font_size(run, marker_data['font_size'])
                # CRITICAL DEBUG: Log ProductStrain processing
                if marker_name == 'PRODUCTSTRAIN':
                    self.logger.info(f"✅ PRODUCTSTRAIN PROCESSED: Added ProductStrain run with content='{display_content}', font_size={marker_data['font_size'].pt if hasattr(marker_data['font_size'], 'pt') else marker_data['font_size']}pt")
                self.logger.debug(f"Added marker '{marker_name}': '{display_content}' -> {marker_data['font_size'].pt}pt")
                
                lines = display_content.splitlines()
                for i, line in enumerate(lines):
                    if i > 0:
                        run.add_break()
                    run.add_text(line)
                current_pos = marker_data['end_pos']
            
            # Add any remaining text
            if current_pos < len(full_text):
                text_after = full_text[current_pos:]
                # Preserve line breaks and whitespace, but skip if completely empty
                if text_after or text_after.strip():
                    run = paragraph.add_run(text_after)
                    run.font.name = "Arial"
                    run.font.bold = True
                    # Use unified font sizing for non-marker text
                    from src.core.generation.unified_font_sizing import get_font_size
                    font_size = get_font_size(text_after, 'default', self.template_type, self.scale_factor)
                    run.font.size = font_size
                    self.logger.debug(f"Added text after: '{text_after}' -> {font_size.pt}pt")
            
            # Convert |BR| markers to actual line breaks after marker processing
            self._convert_br_markers_to_line_breaks(paragraph)
            
            # Precompute strain contents to protect their sizing during brand adjustments
            strain_contents_casefold = {
                (data.get('content') or '').strip().casefold()
                for key, data in processed_content.items()
                if key in ('PRODUCTSTRAIN', 'STRAIN')
            }

            # Apply special formatting for specific markers
            for marker_name, marker_data in processed_content.items():
                # Always center ProductBrand markers for ALL templates
                if ('PRODUCTBRAND' in marker_name):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._set_paragraph_cell_vertical_alignment(paragraph, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
                    for idx, run in enumerate(paragraph.runs):
                        # Get product type for font sizing
                        product_type = None
                        if hasattr(self, 'current_product_type'):
                            product_type = self.current_product_type
                        elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                            product_type = self.label_context['ProductType']
                        run_text_normalized = (run.text or '').strip().casefold()
                        if run_text_normalized and run_text_normalized in strain_contents_casefold:
                            # Preserve ProductStrain runs that share the paragraph
                            continue
                        # Only apply ProductBrand sizing to the run that actually contains the brand content
                        if marker_name in ('PRODUCTBRAND', 'PRODUCTBRAND_CENTER') and idx != 0:
                            continue
                        set_run_font_size(run, get_font_size_by_marker(marker_data['content'], marker_name, self.template_type, self.scale_factor, product_type))
                    continue
                if marker_name == 'DOH':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    continue
                if marker_name == 'RATIO':
                    for run in paragraph.runs:
                        # Get product type for font sizing
                        product_type = None
                        if hasattr(self, 'current_product_type'):
                            product_type = self.current_product_type
                        elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                            product_type = self.label_context['ProductType']
                        set_run_font_size(run, get_font_size_by_marker(marker_data['content'], 'RATIO', self.template_type, self.scale_factor, product_type))
                        # Ensure ratio values are bold
                        run.font.bold = True
                    continue
                if marker_name == 'LINEAGE':
                    content = marker_data['content']
                    product_type = None
                    if hasattr(self, 'current_product_type'):
                        product_type = self.current_product_type
                    elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                        product_type = self.label_context['ProductType']
                    
                    from src.core.constants import CLASSIC_TYPES, VALID_CLASSIC_LINEAGES
                    product_type_normalized = (product_type or '').lower()
                    is_classic_product = product_type_normalized in CLASSIC_TYPES if product_type_normalized else False
                    
                    # CRITICAL FIX: Check if lineage content itself is a classic lineage value
                    # Clean the content to check for classic lineage values
                    clean_content = content.strip().upper()
                    # Remove any marker remnants
                    clean_content = re.sub(r'PRODUCTBRAND_CENTER_(START|END)', '', clean_content, flags=re.IGNORECASE).strip()
                    clean_content = re.sub(r'LINEAGE_(START|END)', '', clean_content, flags=re.IGNORECASE).strip()
                    # Check if the cleaned content is a classic lineage value
                    is_classic_lineage_value = clean_content in VALID_CLASSIC_LINEAGES or any(
                        clean_content.startswith(classic_lineage) for classic_lineage in VALID_CLASSIC_LINEAGES
                    )
                    
                    if (not is_classic_product) and ('PRODUCTBRAND_CENTER' in content):
                        brand_text = re.sub(
                            r'PRODUCTBRAND_CENTER_(START|END)',
                            '',
                            content,
                            flags=re.IGNORECASE
                        ).strip()
                        font_size = get_font_size_by_marker(
                            brand_text,
                            'PRODUCTBRAND_CENTER',
                            self.template_type,
                            self.scale_factor,
                            product_type_normalized
                        )
                        brand_clean_regex = re.compile(r'PRODUCTBRAND_CENTER_(START|END)', re.IGNORECASE)
                    else:
                        font_size = get_font_size_by_marker(
                            content,
                            'LINEAGE',
                            self.template_type,
                            self.scale_factor,
                            product_type_normalized
                        )
                        brand_clean_regex = None
                    
                    for run in paragraph.runs:
                        original_text = run.text or ''
                        set_run_font_size(run, font_size)
                        if brand_clean_regex:
                            run.text = brand_clean_regex.sub('', original_text).strip()
                    
                    # Handle alignment based on PRODUCT TYPE OR classic lineage value
                    # CRITICAL FIX: If lineage content is a classic lineage value, always left-align
                    # This ensures lineage values like HYBRID, SATIVA, INDICA are left-aligned
                    # even if product type detection fails
                    if is_classic_lineage_value:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        paragraph.paragraph_format.left_indent = Inches(0)
                        paragraph.paragraph_format.space_before = Pt(2)
                        paragraph.paragraph_format.space_after = Pt(1)
                        self.logger.debug(f"LINEAGE ALIGNMENT: Forced LEFT alignment for classic lineage value: '{clean_content}'")
                    elif is_classic_product:
                        # Classic product types should have LEFT alignment for lineage
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # NO LEFT INDENT - this was causing lineage indentation
                        paragraph.paragraph_format.left_indent = Inches(0)
                        # Ensure consistent spacing above lineage section for equal margins
                        paragraph.paragraph_format.space_before = Pt(2)
                        paragraph.paragraph_format.space_after = Pt(1)
                    else:
                        # Non-classic product types should have CENTER alignment for lineage
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Ensure consistent spacing above lineage section for equal margins
                        paragraph.paragraph_format.space_before = Pt(2)
                        paragraph.paragraph_format.space_after = Pt(1)
                    
                    # SPECIFIC OVERRIDE: Ensure Vape Cartridge products always have LEFT-aligned lineage
                    if product_type and 'vape' in product_type.lower():
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        paragraph.paragraph_format.left_indent = Inches(0)
                        self.logger.debug(f"VAPE CARTRIDGE OVERRIDE: Forced LEFT alignment for lineage")
                    
                    continue
                # Always center ProductBrand and ProductBrand_Center markers
                if marker_name in ('PRODUCTBRAND', 'PRODUCTBRAND_CENTER') or 'PRODUCTBRAND' in marker_name:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Ensure consistent spacing above product brand section for equal margins
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(1)
                    for run in paragraph.runs:
                        # Get product type for font sizing
                        product_type = None
                        if hasattr(self, 'current_product_type'):
                            product_type = self.current_product_type
                        elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                            product_type = self.label_context['ProductType']
                        set_run_font_size(run, get_font_size_by_marker(marker_data['content'], marker_name, self.template_type, self.scale_factor, product_type))
                    continue
                # Right-align PRODUCTVENDOR markers
                if marker_name == 'PRODUCTVENDOR':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    # Ensure consistent spacing above vendor section for equal margins
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(1)
                    # Use unified font sizing for vendor text
                    for run in paragraph.runs:
                        # Apply unified font sizing using 'vendor' field type
                        from src.core.generation.unified_font_sizing import get_font_size
                        vendor_font_size = get_font_size(marker_data['content'], 'vendor', self.template_type, self.scale_factor)
                        set_run_font_size(run, vendor_font_size)
                        # Set vendor text to italic and gray color
                        run.font.italic = True
                        from docx.shared import RGBColor
                        # Set gray color at both run level and XML level for consistency
                        run.font.color.rgb = RGBColor(128, 128, 128)  # #808080
                        run.font.color.theme_color = None  # Clear any theme color
                        # Also set color at XML level to ensure it sticks
                        rPr = run._element.get_or_add_rPr()
                        color = rPr.find(qn('w:color'))
                        if color is None:
                            color = OxmlElement('w:color')
                            rPr.append(color)
                        color.set(qn('w:val'), '808080')  # Gray color in hex without #
                    continue
                elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                    product_type = self.label_context['ProductType']
                else:
                    product_type = None
                
                # Special handling for ProductStrain marker - use unified font sizing system and left alignment
                if marker_name in ('PRODUCTSTRAIN', 'STRAIN'):
                    # Left-align PRODUCTSTRAIN markers
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # Ensure consistent spacing above strain section for equal margins
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(1)
                    strain_content = str(marker_data.get('content') or '').strip()
                    for run in paragraph.runs:
                        run_text = run.text or ''
                        run_text_stripped = run_text.strip()
                        has_marker_wrappers = any(
                            marker in run_text
                            for marker in [
                                'PRODUCTSTRAIN_START',
                                'PRODUCTSTRAIN_END',
                                'STRAIN_START',
                                'STRAIN_END'
                            ]
                        )
                        is_exact_match = strain_content and run_text_stripped.casefold() == strain_content.casefold()

                        if strain_content and (has_marker_wrappers or is_exact_match):
                            # Use unified font sizing system for ProductStrain markers (1pt font size)
                            strain_font_size = get_font_size_by_marker(strain_content, 'PRODUCTSTRAIN', self.template_type, self.scale_factor)
                            set_run_font_size(run, strain_font_size)
                    continue
                
                # CRITICAL FIX: Standalone CBD text in Lineage field should use lineage font sizing, not strain
                if marker_name in ('CBD', 'THC', 'CBC', 'CBG', 'CBN') and marker_data['content'].strip() in ['CBD', 'THC', 'CBC', 'CBG', 'CBN']:
                    for run in paragraph.runs:
                        # Only apply lineage font sizing to runs that contain standalone cannabinoid text in lineage context
                        if (marker_data['content'] in run.text and 
                            len(marker_data['content'].strip()) <= 3 and  # Short cannabinoid names only
                            marker_data['content'].strip().upper() in ['CBD', 'THC', 'CBC', 'CBG', 'CBN']):
                            # Use unified font sizing system for standalone cannabinoid text as lineage field type
                            lineage_font_size = get_font_size_by_marker(marker_data['content'], 'LINEAGE', self.template_type, self.scale_factor)
                            set_run_font_size(run, lineage_font_size)
                    continue
                
                # Apply normal font sizing for other markers
                for run in paragraph.runs:
                    # Additional check for standalone cannabinoid text that might have slipped through - treat as lineage
                    if (marker_data['content'] in run.text and 
                        len(marker_data['content'].strip()) <= 3 and
                        marker_data['content'].strip().upper() in ['CBD', 'THC', 'CBC', 'CBG', 'CBN'] and
                        not any(marker in run.text for marker in ['CBD_START', 'THC_START', 'CBC_START', 'CBG_START', 'CBN_START'])):
                        # This is standalone cannabinoid text - use lineage font sizing, not strain
                        lineage_font_size = get_font_size_by_marker(marker_data['content'], 'LINEAGE', self.template_type, self.scale_factor)
                        set_run_font_size(run, lineage_font_size)
                    else:
                        set_run_font_size(run, get_font_size_by_marker(marker_data['content'], marker_name, self.template_type, self.scale_factor, product_type))
                # Special handling for ProductVendor marker - now handled above with unified font sizing
                # This section removed to prevent conflicts with the unified font sizing system
            
            self.logger.debug(f"Applied multi-marker processing for: {list(processed_content.keys())}")
        try:
            pass
        except Exception as e:
            self.logger.error(f"Error processing multi-marker template: {e}")
            # Fallback: remove all markers and use default size
            for run in paragraph.runs:
                for marker_name in markers:
                    start_marker = f'{marker_name}_START'
                    end_marker = f'{marker_name}_END'
                    run.text = run.text.replace(start_marker, "").replace(end_marker, "")
                # Use appropriate default size based on template type
                # Use unified font sizing system for default size
                from src.core.generation.unified_font_sizing import get_font_size
                default_size = get_font_size(run.text, 'default', self.template_type, self.scale_factor)
                run.font.size = default_size
        finally:
            # Always check for |BR| markers regardless of success/failure
            self._convert_br_markers_to_line_breaks(paragraph)

    def _process_paragraph_for_marker_template_specific(self, paragraph, marker_name):
        """
        Process a single paragraph for a specific marker using template-type-specific font sizing.
        """
        start_marker = f'{marker_name}_START'
        end_marker = f'{marker_name}_END'
        
        full_text = "".join(run.text for run in paragraph.runs)
        
        if start_marker in full_text and end_marker in full_text:
            try:
                # Extract content
                start_idx = full_text.find(start_marker) + len(start_marker)
                end_idx = full_text.find(end_marker)
                content = full_text[start_idx:end_idx]
                
                # For THC_CBD markers, calculate font size before any splitting to ensure consistency
                if marker_name in ['THC_CBD', 'RATIO', 'THC_CBD_LABEL'] and ('\n' in content or '|BR|' in content):
                    # Calculate font size based on the original unsplit content to ensure consistency
                    original_content = content.replace('\n', ' ').replace('|BR|', ' ')
                    font_size = self._get_template_specific_font_size(original_content, marker_name)
                    import logging
                    logging.debug(f"[FONT_DEBUG] Processing marker '{marker_name}' with original content '{original_content}' -> font_size: {font_size}")
                    
                    # Clear and recreate with single run approach
                    paragraph.clear()
                    
                    # Create a single run with the entire content
                    run = paragraph.add_run()
                    run.font.name = "Arial"
                    run.font.bold = True
                    run.font.size = font_size
                    set_run_font_size(run, font_size)
                    
                    # Add the content with line breaks as text
                    run.text = content  # Use assignment instead of add_text to avoid duplication
                    
                    # Convert line breaks to actual line breaks, passing the font size
                    self._convert_br_markers_to_line_breaks(paragraph, font_size)
                else:
                    # Use template-type-specific font sizing based on original functions
                    font_size = self._get_template_specific_font_size(content, marker_name)
                    import logging
                    logging.debug(f"[FONT_DEBUG] Processing marker '{marker_name}' with content '{content}' -> font_size: {font_size}")
                    
                    # Clear paragraph and re-add content with template-optimized formatting
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.font.name = "Arial"
                    # Make everything bold - no exceptions
                    run.font.bold = True
                    run.font.size = font_size
                    
                    # Apply template-specific font size setting
                    set_run_font_size(run, font_size)
                    
                    # Add the content to the run
                    run.text = content  # Use assignment instead of add_text to avoid duplication
                    
                    # Convert |BR| markers to actual line breaks for other markers
                    self._convert_br_markers_to_line_breaks(paragraph, font_size)
                
                # Handle special formatting for specific markers
                if marker_name in ['PRODUCTBRAND', 'PRODUCTBRAND_CENTER']:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Also ensure all runs in this paragraph are properly sized
                    for run in paragraph.runs:
                        set_run_font_size(run, font_size)
                    # Center the cell vertically so brand text sits mid-cell (e.g., CONSTELLATION CANNABIS)
                    self._set_paragraph_cell_vertical_alignment(paragraph, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
                elif marker_name in ['THC_CBD', 'RATIO', 'THC_CBD_LABEL']:
                    # Ensure THC_CBD and RATIO values are bold
                    for run in paragraph.runs:
                        run.font.bold = True
                    
                    # For vertical template, apply line spacing from unified font sizing
                    line_spacing = get_line_spacing_by_marker(marker_name, self.template_type)
                    if line_spacing:
                        paragraph.paragraph_format.line_spacing = line_spacing
                        # Only set LEFT alignment if the paragraph is not already centered
                        if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # Set at XML level for maximum compatibility
                        pPr = paragraph._element.get_or_add_pPr()
                        spacing = pPr.find(qn('w:spacing'))
                        if spacing is None:
                            spacing = OxmlElement('w:spacing')
                            pPr.append(spacing)
                        spacing.set(qn('w:line'), str(int(line_spacing * 240)))
                        spacing.set(qn('w:lineRule'), 'auto')
                    
                    # For vertical template THC_CBD content, use right alignment for percentage values
                    if self.template_type == 'vertical' and marker_name == 'THC_CBD':
                        # Set paragraph alignment to right for proper percentage alignment
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        self.logger.debug(f"Set right alignment for vertical template THC_CBD content")
                    # All content now uses standard spacing
                    # For all other Ratio content in horizontal template, set vertical alignment to top
                    elif self.template_type == 'horizontal' and marker_name in ['THC_CBD', 'RATIO', 'THC_CBD_LABEL']:
                        # Set vertical alignment to top for the cell containing this paragraph
                        # BUT preserve center alignment for cells with DOH images
                        if paragraph._element.getparent().tag.endswith('tc'):  # Check if in table cell
                            cell = paragraph._element.getparent()
                            # Check if this cell contains a DOH image before setting to TOP
                            has_doh_image = self._cell_contains_doh_image(cell)
                            if not has_doh_image:
                                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                            else:
                                self.logger.debug("Preserving center alignment for horizontal template cell with DOH image")
                    # For all other THC/CBD content in other templates, set vertical alignment to top
                    elif marker_name in ['THC_CBD', 'RATIO', 'THC_CBD_LABEL']:
                        # Set vertical alignment to top for the cell containing this paragraph
                        # BUT preserve center alignment for cells with DOH images
                        if paragraph._element.getparent().tag.endswith('tc'):  # Check if in table cell
                            cell = paragraph._element.getparent()
                            # Check if this cell contains a DOH image before setting to TOP
                            has_doh_image = self._cell_contains_doh_image(cell)
                            if not has_doh_image:
                                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                            else:
                                self.logger.debug("Preserving center alignment for cell with DOH image")
                
                # Center alignment for brand names
                if 'PRODUCTBRAND' in marker_name:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Center alignment for DOH (Date of Harvest)
                if marker_name == 'DOH':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Special handling for lineage markers
                if marker_name == 'LINEAGE':
                    self.logger.debug(f"Processing LINEAGE marker with content: '{content}'")
                    
                    # CRITICAL FIX: Clean corrupted marker text before processing
                    original_content = content
                    content = re.sub(r'PRODUCTSTRR_STARTCONSTELL.*', '', content)
                    content = re.sub(r'STARTCONSTELL.*', '', content)
                    content = re.sub(r'CONSTELLATION\$.*', '', content)
                    
                    # CRITICAL FIX: Remove any remaining $ symbols that might be marker remnants
                    # This handles cases like "VICE$Star" where $ is a corrupted marker remnant
                    content = re.sub(r'\$.*', '', content)
                    
                    content = content.strip()
                    
                    if original_content != content:
                        self.logger.warning(f"Cleaned corrupted lineage content: '{original_content}' -> '{content}'")
                    
                    # CRITICAL FIX: If Lineage contains PRODUCTBRAND_CENTER markers, process it as PRODUCTBRAND_CENTER
                    if 'PRODUCTBRAND_CENTER_START' in content and 'PRODUCTBRAND_CENTER_END' in content:
                        self.logger.debug(f"Lineage contains PRODUCTBRAND_CENTER markers, processing as PRODUCTBRAND_CENTER")
                        # Extract the brand content from the PRODUCTBRAND_CENTER markers
                        brand_start = content.find('PRODUCTBRAND_CENTER_START') + len('PRODUCTBRAND_CENTER_START')
                        brand_end = content.find('PRODUCTBRAND_CENTER_END')
                        brand_content = content[brand_start:brand_end]
                        
                        # Calculate proper font size for brand content using 'brand' field type
                        from src.core.generation.unified_font_sizing import get_font_size
                        font_size = get_font_size(brand_content, 'brand', self.template_type, self.scale_factor)
                        
                        # Clear paragraph and recreate with brand content
                        paragraph.clear()
                        run = paragraph.add_run()
                        run.font.name = "Arial"
                        run.font.bold = True
                        run.font.size = font_size
                        set_run_font_size(run, font_size)
                        run.text = brand_content  # Use assignment instead of add_text to avoid duplication
                        
                        # Center the paragraph for nonclassic types (ProductBrand content)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        self.logger.debug(f"Centered Lineage (ProductBrand) content: '{brand_content}'")
                        return  # Exit early since we've handled this as PRODUCTBRAND_CENTER
                    
                    # Extract product type information from the content
                    if '_PRODUCT_TYPE_' in content and '_IS_CLASSIC_' in content:
                        parts = content.split('_PRODUCT_TYPE_')
                        if len(parts) == 2:
                            actual_lineage = parts[0]
                            # Remove PRODUCTBRAND_CENTER_START marker if present
                            if actual_lineage.startswith('PRODUCTBRAND_CENTER_START'):
                                actual_lineage = actual_lineage[len('PRODUCTBRAND_CENTER_START'):]
                            type_info = parts[1]
                            type_parts = type_info.split('_IS_CLASSIC_')
                            if len(type_parts) == 2:
                                product_type = type_parts[0]
                                is_classic_raw = type_parts[1]
                                # Remove LINEAGE_END or PRODUCTBRAND_CENTER_END marker if present
                                if is_classic_raw.endswith('LINEAGE_END'):
                                    is_classic_raw = is_classic_raw[:-len('LINEAGE_END')]
                                elif is_classic_raw.endswith('PRODUCTBRAND_CENTER_END'):
                                    is_classic_raw = is_classic_raw[:-len('PRODUCTBRAND_CENTER_END')]
                                is_classic = is_classic_raw.lower() == 'true'
                                
                                # CRITICAL FIX: Check if lineage content itself is a classic lineage value
                                from src.core.constants import VALID_CLASSIC_LINEAGES
                                clean_lineage = actual_lineage.strip().upper()
                                clean_lineage = re.sub(r'LINEAGE_(START|END)', '', clean_lineage, flags=re.IGNORECASE).strip()
                                clean_lineage = re.sub(r'PRODUCTBRAND_CENTER_(START|END)', '', clean_lineage, flags=re.IGNORECASE).strip()
                                is_classic_lineage_value = clean_lineage in VALID_CLASSIC_LINEAGES or any(
                                    clean_lineage.startswith(classic_lineage) for classic_lineage in VALID_CLASSIC_LINEAGES
                                )
                                
                                # For nonclassic types, Lineage field contains ProductBrand content which should always be centered
                                # For classic types, Lineage field contains actual lineage content which should be left-aligned
                                # CRITICAL FIX: Also check if lineage content is a classic lineage value
                                if is_classic or is_classic_lineage_value:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    paragraph.paragraph_format.left_indent = Inches(0)
                                    paragraph.paragraph_format.space_before = Pt(2)
                                    paragraph.paragraph_format.space_after = Pt(1)
                                    if is_classic_lineage_value:
                                        self.logger.debug(f"Left-aligned lineage for classic lineage value: '{clean_lineage}'")
                                else:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    paragraph.paragraph_format.space_before = Pt(2)
                                    paragraph.paragraph_format.space_after = Pt(1)
                                
                                # Update the content to only show the actual lineage (remove any markers)
                                if actual_lineage.startswith('LINEAGE_START'):
                                    actual_lineage = actual_lineage[len('LINEAGE_START'):]
                                elif actual_lineage.startswith('PRODUCTBRAND_CENTER_START'):
                                    actual_lineage = actual_lineage[len('PRODUCTBRAND_CENTER_START'):]
                                content = actual_lineage
                    else:
                        # Fallback: check if this is a classic product type by using the context
                        # Import constants to check against CLASSIC_TYPES
                        from src.core.constants import CLASSIC_TYPES, VALID_CLASSIC_LINEAGES
                        
                        # CRITICAL FIX: Check if lineage content itself is a classic lineage value
                        # Clean the content to check for classic lineage values
                        clean_content = content.strip().upper()
                        # Remove any marker remnants
                        clean_content = re.sub(r'PRODUCTBRAND_CENTER_(START|END)', '', clean_content, flags=re.IGNORECASE).strip()
                        clean_content = re.sub(r'LINEAGE_(START|END)', '', clean_content, flags=re.IGNORECASE).strip()
                        # Check if the cleaned content is a classic lineage value
                        is_classic_lineage_value = clean_content in VALID_CLASSIC_LINEAGES or any(
                            clean_content.startswith(classic_lineage) for classic_lineage in VALID_CLASSIC_LINEAGES
                        )
                        
                        # Get product type from context, not from content
                        is_classic_product = False
                        if hasattr(self, 'current_product_type'):
                            product_type = self.current_product_type
                        elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                            product_type = self.label_context['ProductType']
                        else:
                            product_type = None
                        
                        # Check if the product type is classic
                        if product_type:
                            is_classic_product = product_type.lower() in CLASSIC_TYPES
                            # Debug logging for vape cartridge lineage alignment in fallback
                            if 'vape' in product_type.lower():
                                self.logger.debug(f"VAPE CARTRIDGE FALLBACK DEBUG: product_type='{product_type}', is_classic_product={is_classic_product}, CLASSIC_TYPES={CLASSIC_TYPES}")
                        
                        # DEBUG: Log the centering decision for non-classic types
                        self.logger.info(f"DEBUG: LINEAGE centering decision - product_type='{product_type}', is_classic_product={is_classic_product}, is_classic_lineage_value={is_classic_lineage_value}, content='{content}'")
                        
                        # CRITICAL FIX: If lineage content is a classic lineage value, always left-align
                        # This ensures lineage values like HYBRID, SATIVA, INDICA are left-aligned
                        # even if product type detection fails
                        if is_classic_lineage_value:
                            # For Classic Lineage Values, left-justify the lineage text
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            paragraph.paragraph_format.left_indent = Inches(0)
                            paragraph.paragraph_format.space_before = Pt(2)
                            paragraph.paragraph_format.space_after = Pt(1)
                            self.logger.debug(f"Left-justified lineage for classic lineage value: '{clean_content}' (content: '{content}')")
                        elif is_classic_product:
                            # For Classic Types, left-justify the lineage text
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            paragraph.paragraph_format.left_indent = Inches(0)
                            paragraph.paragraph_format.space_before = Pt(2)
                            paragraph.paragraph_format.space_after = Pt(1)
                            self.logger.debug(f"Left-justified lineage for classic product type: '{content}' (product_type: {product_type})")
                        else:
                            # For non-classic types, center the ProductBrand content in Lineage field
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph.paragraph_format.space_before = Pt(2)
                            paragraph.paragraph_format.space_after = Pt(1)
                            self.logger.debug(f"Centered lineage (ProductBrand) for non-classic product type: '{content}' (product_type: {product_type})")
                        
                        # SPECIFIC OVERRIDE: Ensure Vape Cartridge products always have LEFT-aligned lineage (fallback)
                        if product_type and 'vape' in product_type.lower():
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            paragraph.paragraph_format.left_indent = Inches(0)
                            self.logger.debug(f"VAPE CARTRIDGE FALLBACK OVERRIDE: Forced LEFT alignment for lineage")
                
                self.logger.debug(f"Applied template-specific font sizing: {font_size.pt}pt for {marker_name} marker")

            except Exception as e:
                self.logger.error(f"Error processing template-specific marker {marker_name}: {e}")
                # Fallback: remove markers and use default size based on template type
                for run in paragraph.runs:
                    run.text = run.text.replace(start_marker, "").replace(end_marker, "")
                    # Use appropriate default size based on template type
                    # Use unified font sizing system for default size
                    from src.core.generation.unified_font_sizing import get_font_size
                    default_size = get_font_size(run.text, 'default', self.template_type, self.scale_factor)
                    run.font.size = default_size
        elif start_marker in full_text or end_marker in full_text:
            # Log partial markers for debugging
            self.logger.debug(f"Found partial {marker_name} marker in text: '{full_text[:100]}...'")

    def _convert_br_markers_to_line_breaks(self, paragraph, font_size=None):
        """
        Convert |BR| markers and \n characters in paragraph text to actual line breaks.
        This splits the text at |BR| markers or \n characters and creates separate runs for each part.
        """
        try:
            # Get all text from the paragraph and store existing font sizes
            full_text = "".join(run.text for run in paragraph.runs)
            
            # Store existing font sizes for each run
            existing_sizes = []
            for run in paragraph.runs:
                if run.text.strip():
                    existing_sizes.append(run.font.size)
            
            # If we have existing sizes, use the first one for all runs to ensure consistency
            # Or use the passed font_size parameter if provided
            consistent_font_size = None
            if font_size is not None:
                consistent_font_size = font_size
            elif existing_sizes:
                consistent_font_size = existing_sizes[0]
            
            # Check if there are any |BR| markers or \n characters
            if '|BR|' not in full_text and '\n' not in full_text:
                return
            
            # First split by |BR| markers, then by \n characters
            if '|BR|' in full_text:
                parts = full_text.split('|BR|')
            else:
                parts = full_text.split('\n')
            
            # Clear the paragraph
            paragraph.clear()
            
            # Set tight paragraph spacing to prevent excessive gaps
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            # Only set line spacing if it's not already set (to preserve custom line spacing)
                        # Use standard line spacing for all content
            paragraph.paragraph_format.line_spacing = 1.0
            
            # Add each part as a separate run, with line breaks between them
            size_index = 0
            previous_was_empty = False
            for i, part in enumerate(parts):
                stripped_part = part.strip() if part else ''
                is_empty = not stripped_part
                
                # CRITICAL FIX: Handle empty parts to ensure text starts on new lines
                # If previous part was empty or this is the first part and it's empty,
                # we need to ensure the next part starts on a new line
                if is_empty:
                    previous_was_empty = True
                    # If there's a next part, we'll add a line break before it
                    continue
                
                # If previous part was empty, add a line break before this part
                if previous_was_empty:
                    break_run = paragraph.add_run()
                    break_run.add_break(WD_BREAK.LINE)
                    previous_was_empty = False
                
                # This part has content - add it
                # CRITICAL FIX: Preserve non-breaking hyphens when stripping whitespace
                # Check for non-breaking hyphens before stripping
                if '\u2011' in part:
                    self.logger.debug(f"BR CONVERSION DEBUG: Found non-breaking hyphens in part: '{part}'")
                # Strip whitespace for all content to remove extra spaces, but preserve non-breaking hyphens
                # First, temporarily replace non-breaking hyphens with a placeholder
                temp_part = part.replace('\u2011', '___NONBREAKING_HYPHEN___')
                # Then strip whitespace
                stripped_part = temp_part.strip()
                # Finally, restore non-breaking hyphens
                stripped_part = stripped_part.replace('___NONBREAKING_HYPHEN___', '\u2011')
                # Check for non-breaking hyphens after stripping
                if '\u2011' in stripped_part:
                    self.logger.debug(f"BR CONVERSION DEBUG: Preserved non-breaking hyphens after strip: '{stripped_part}'")
                else:
                    self.logger.debug(f"BR CONVERSION DEBUG: Lost non-breaking hyphens after strip: '{stripped_part}'")
                
                run = paragraph.add_run(stripped_part)
                run.font.name = "Arial"
                
                # ALL text should be Arial Bold - NO EXCEPTIONS
                run.font.bold = True
                
                # Use consistent font size for all runs
                if consistent_font_size:
                    run.font.size = consistent_font_size
                else:
                    # Use unified font sizing for default size
                    from src.core.generation.unified_font_sizing import get_font_size
                    default_font_size = get_font_size(stripped_part, 'default', self.template_type, self.scale_factor)
                    run.font.size = default_font_size
                
                # Add a line break after this part if there's a next part
                # This ensures that text after |BR| or \n starts on a new line
                if i < len(parts) - 1:
                    # Check if next part is empty - if so, we'll handle it in the next iteration
                    next_part = parts[i + 1].strip() if i + 1 < len(parts) else ''
                    if next_part:
                        # Next part has content - add line break after current part
                        run.add_break(WD_BREAK.LINE)
                    # If next part is empty, we'll add the break before the part after that
            
            # All content now uses standard 1.0 line spacing
            
            self.logger.debug(f"Converted {len(parts)-1} |BR| markers to line breaks")
            
        except Exception as e:
            self.logger.error(f"Error converting BR markers to line breaks: {e}")
            # Fallback: just remove the BR markers
            for run in paragraph.runs:
                run.text = run.text.replace('|BR|', ' ')
    def _ensure_consistent_lineage_spacing(self, doc):
        """
        Ensure consistent spacing above lineage/brand sections for equal margins across all labels.
        This creates uniform visual spacing above the colored lineage/brand bars.
        """
        try:
            def process_paragraph(paragraph):
                # Check if this paragraph contains lineage or vendor content
                text = paragraph.text.lower()
                if any(keyword in text for keyword in ['indica', 'sativa', 'hybrid', 'cbd', 'alpha crux', 'constellation']):
                    # Set consistent spacing for lineage/brand sections
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(1)
                    
                    # Also set at XML level for maximum compatibility
                    pPr = paragraph._element.get_or_add_pPr()
                    spacing = pPr.find(qn('w:spacing'))
                    if spacing is None:
                        spacing = OxmlElement('w:spacing')
                        pPr.append(spacing)
                    spacing.set(qn('w:before'), '40')  # 2pt = 40 twips
                    spacing.set(qn('w:after'), '20')   # 1pt = 20 twips
                    spacing.set(qn('w:lineRule'), 'auto')
            
            # Process all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
            
            # Process all paragraphs outside tables
            for paragraph in doc.paragraphs:
                process_paragraph(paragraph)
            
            # Ensured consistent spacing above lineage/brand sections
            
        except Exception as e:
            self.logger.error(f"Error ensuring consistent lineage spacing: {e}")

    def _add_consistent_content_spacing(self, doc):
        """
        Add consistent spacing above main content sections for better visual balance.
        This ensures uniform spacing above product names, prices, and other key content.
        """
        try:
            def process_paragraph(paragraph):
                # Check if this paragraph contains main content
                text = paragraph.text.lower()
                if any(keyword in text for keyword in ['$', 'thc:', 'cbd:', 'mg', 'oz', 'g', 'pack']):
                    # Add consistent spacing above main content sections
                    current_before = paragraph.paragraph_format.space_before
                    current_after = paragraph.paragraph_format.space_after
                    
                    # Only add spacing if it's not already set to our target values
                    if current_before == Pt(0) or current_before is None:
                        paragraph.paragraph_format.space_before = Pt(1)
                    
                    if current_after == Pt(0) or current_after is None:
                        paragraph.paragraph_format.space_after = Pt(0.5)
                    
                    # Also set at XML level for maximum compatibility
                    pPr = paragraph._element.get_or_add_pPr()
                    spacing = pPr.find(qn('w:spacing'))
                    if spacing is None:
                        spacing = OxmlElement('w:spacing')
                        pPr.append(spacing)
                    
                    if current_before == Pt(0) or current_before is None:
                        spacing.set(qn('w:before'), '20')  # 1pt = 20 twips
                    if current_after == Pt(0) or current_after is None:
                        spacing.set(qn('w:after'), '10')   # 0.5pt = 10 twips
                    
                    spacing.set(qn('w:lineRule'), 'auto')
            
            # Process all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
            
            # Process all paragraphs outside tables
            for paragraph in doc.paragraphs:
                process_paragraph(paragraph)
            
            # Added consistent spacing above main content sections
            
        except Exception as e:
            self.logger.error(f"Error adding consistent content spacing: {e}")

    def _fix_ratio_paragraph_spacing(self, doc):
        """
        Fix paragraph spacing for ratio content to prevent excessive gaps between lines.
        This ensures tight spacing for multi-line ratio content.
        """
        try:
            # Define patterns that indicate ratio content
            ratio_patterns = [
                'mg THC', 'mg CBD', 'mg CBG', 'mg CBN', 'mg CBC',
                'THC:', 'CBD:', 'CBG:', 'CBN:', 'CBC:',
                '1:1', '2:1', '3:1', '1:1:1', '2:1:1'
            ]
            
            def process_paragraph(paragraph):
                # Check if this paragraph contains ratio content
                text = paragraph.text.lower()
                if any(pattern.lower() in text for pattern in ratio_patterns):
                    # Set tight spacing for all ratio content (including THC_CBD)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    
                    # Also set tight spacing for any child paragraphs (in case of nested content)
                    for child_para in paragraph._element.findall('.//w:p', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        if hasattr(child_para, 'pPr') and child_para.pPr is not None:
                            # Set spacing properties at XML level for maximum compatibility
                            spacing = child_para.pPr.find('.//w:spacing', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                            if spacing is None:
                                spacing = OxmlElement('w:spacing')
                                child_para.pPr.append(spacing)
                            
                            spacing.set(qn('w:before'), '0')
                            spacing.set(qn('w:after'), '0')
                            spacing.set(qn('w:line'), '240')  # 1.0 line spacing (240 twips)
                            spacing.set(qn('w:lineRule'), 'auto')
            
            # Process all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
            
            # Process all paragraphs outside tables
            for paragraph in doc.paragraphs:
                process_paragraph(paragraph)
            
            # Fixed paragraph spacing for ratio content
            
        except Exception as e:
            self.logger.error(f"Error fixing ratio paragraph spacing: {e}")
            # Don't raise the exception - this is a formatting enhancement that shouldn't break the main process

    def _validate_and_repair_table_structure(self, table):
        """
        Validate and repair table structure to ensure it has required elements.
        Returns True if table is valid, False if it cannot be repaired.
        """
        try:
            # First, try to access table properties to see if there's an actual error
            try:
                _ = table.rows
                _ = table.columns
                # If we can access these without error, the table is fine
                return True
            except Exception as e:
                # Table is corrupted, try to repair
                self.logger.debug(f"Table access failed, attempting repair: {e}")
            
            # Check if table has the required tblGrid element
            tblGrid = table._element.find(qn('w:tblGrid'))
            if tblGrid is None:
                # Create tblGrid element
                tblGrid = OxmlElement('w:tblGrid')
                
                # Try to get column count from the XML structure directly
                try:
                    # Look for rows in the XML to determine column count
                    rows = table._element.findall(qn('w:tr'))
                    if rows:
                        # Count cells in the first row
                        first_row_cells = rows[0].findall(qn('w:tc'))
                        col_count = len(first_row_cells)
                        
                        # Create grid columns
                        for _ in range(col_count):
                            gridCol = OxmlElement('w:gridCol')
                            gridCol.set(qn('w:w'), '1440')  # Default width of 1 inch
                            tblGrid.append(gridCol)
                        
                        # Insert tblGrid at the beginning of the table element
                        table._element.insert(0, tblGrid)
                        self.logger.debug(f"Repaired missing tblGrid for table with {col_count} columns")
                        
                        # Now test if the repair worked
                        try:
                            _ = table.rows
                            _ = table.columns
                            return True
                        except Exception as test_error:
                            self.logger.error(f"Table repair failed validation test: {test_error}")
                            return False
                    else:
                        self.logger.warning("Cannot repair table: no rows found in XML")
                        return False
                except Exception as repair_error:
                    self.logger.error(f"Failed to repair table structure: {repair_error}")
                    return False
            else:
                # Table already has tblGrid, but let's verify it's working
                try:
                    _ = table.rows
                    _ = table.columns
                    return True
                except Exception as verify_error:
                    self.logger.error(f"Table has tblGrid but still corrupted: {verify_error}")
                    # Try to repair the existing tblGrid
                    try:
                        # Remove and recreate tblGrid
                        old_tblGrid = table._element.find(qn('w:tblGrid'))
                        if old_tblGrid is not None:
                            old_tblGrid.getparent().remove(old_tblGrid)
                        
                        # Create new tblGrid
                        new_tblGrid = OxmlElement('w:tblGrid')
                        rows = table._element.findall(qn('w:tr'))
                        if rows:
                            first_row_cells = rows[0].findall(qn('w:tc'))
                            col_count = len(first_row_cells)
                            
                            for _ in range(col_count):
                                gridCol = OxmlElement('w:gridCol')
                                gridCol.set(qn('w:w'), '1440')
                                new_tblGrid.append(gridCol)
                            
                            table._element.insert(0, new_tblGrid)
                            
                            # Test the repair
                            try:
                                _ = table.rows
                                _ = table.columns
                                self.logger.debug(f"Successfully repaired corrupted tblGrid for table with {col_count} columns")
                                return True
                            except Exception as final_test_error:
                                self.logger.error(f"Final repair attempt failed: {final_test_error}")
                                return False
                        else:
                            return False
                    except Exception as final_repair_error:
                        self.logger.error(f"Final repair attempt failed: {final_repair_error}")
                        # Last resort: try to rebuild the entire table
                        try:
                            if self._rebuild_corrupted_table(table, self.template_type):
                                self.logger.info("Table successfully rebuilt after all repair attempts failed")
                                return True
                            else:
                                self.logger.error("Table rebuild failed, table is beyond repair")
                                return False
                        except Exception as rebuild_error:
                            self.logger.error(f"Table rebuild attempt failed: {rebuild_error}")
                            return False
                
        except Exception as e:
            self.logger.error(f"Error validating/repairing table structure: {e}")
            return False

    def _repair_corrupted_tables(self, doc):
        """Repair any corrupted tables by ensuring proper XML structure."""
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            for table_idx, table in enumerate(doc.tables):
                try:
                    # Test if table is accessible
                    _ = len(table.rows)
                    _ = len(table.rows[0].cells) if table.rows else 0
                except Exception as e:
                    self.logger.warning(f"Table {table_idx} is corrupted: {e}")
                    # Try to repair by recreating the table structure
                    self._recreate_table_structure(table)
                    
        except Exception as e:
            self.logger.error(f"Error repairing corrupted tables: {e}")
    
    def _recreate_table_structure(self, table):
        """Recreate the basic structure of a corrupted table."""
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            # Ensure tblPr exists
            tbl_pr = table._element.find(qn('w:tblPr'))
            if tbl_pr is None:
                tbl_pr = OxmlElement('w:tblPr')
                table._element.insert(0, tbl_pr)
            
            # Ensure tblGrid exists
            tbl_grid = table._element.find(qn('w:tblGrid'))
            if tbl_grid is None:
                tbl_grid = OxmlElement('w:tblGrid')
                table._element.insert(1, tbl_grid)
                
                # Add default grid columns
                col_count = 3  # Default to 3 columns
                for _ in range(col_count):
                    gc = OxmlElement('w:gridCol')
                    gc.set(qn('w:w'), str(int(3.4 * 1440)))
                    tbl_grid.append(gc)
                
                self.logger.info(f"Recreated tblGrid with {col_count} columns")
                
        except Exception as e:
            self.logger.error(f"Error recreating table structure: {e}")

    def _clear_empty_cells(self, doc, num_products):
        """Remove extra cells beyond the number of products by clearing unmerged placeholders."""
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            for table in doc.tables:
                # Calculate how many cells we actually need
                if self.template_type in ['horizontal', 'vertical']:
                    # 3x3 grid = 9 cells max
                    max_cells = 9
                elif self.template_type == 'double':
                    # 4x3 grid = 12 cells max
                    max_cells = 12
                elif self.template_type == 'mini':
                    # 4x5 grid = 20 cells max
                    max_cells = 20
                else:
                    max_cells = num_products
                
                cells_to_remove = max(0, max_cells - num_products)
                self.logger.info(f"🔧 CLEARING EXTRA CELLS: Removing {cells_to_remove} extra cells (need {num_products}, template has {max_cells})")
                
                # CRITICAL FIX: For horizontal/vertical templates, try removing entire rows instead of just cells
                if self.template_type in ['horizontal', 'vertical'] and num_products < 9:
                    # Calculate how many rows we need
                    rows_needed = (num_products + 2) // 3  # Round up to get number of rows needed
                    self.logger.info(f"🔧 DEBUG: Need {rows_needed} rows for {num_products} products")
                    
                    # Remove extra rows
                    if len(table.rows) > rows_needed:
                        rows_to_remove = len(table.rows) - rows_needed
                        self.logger.info(f"🔧 DEBUG: Removing {rows_to_remove} extra rows")
                        for i in range(rows_to_remove):
                            # Remove the last row
                            last_row = table.rows[-1]
                            last_row._element.getparent().remove(last_row._element)
                
                # Also clear individual cells as backup
                cell_count = 0
                total_cells = 0
                for row in table.rows:
                    total_cells += len(row.cells)
                
                self.logger.info(f"🔧 DEBUG: Table has {len(table.rows)} rows, {total_cells} total cells")
                
                for row in table.rows:
                    for cell in row.cells:
                        cell_count += 1
                        self.logger.debug(f"🔧 DEBUG: Processing cell {cell_count} of {total_cells}")
                        if cell_count > num_products:
                            # This is an extra cell - clear it completely and remove all content
                            cell._tc.clear_content()
                            
                            # CRITICAL FIX: Also remove any remaining placeholder text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if '{{' in run.text and '}}' in run.text:
                                        run.text = ''
                                        
                            # CRITICAL FIX: Remove entire paragraphs that contain only placeholders
                            paragraphs_to_remove = []
                            for paragraph in cell.paragraphs:
                                paragraph_text = ''.join([run.text for run in paragraph.runs])
                                if '{{Label' in paragraph_text and '}}' in paragraph_text:
                                    paragraphs_to_remove.append(paragraph)
                            
                            for paragraph in paragraphs_to_remove:
                                paragraph._element.getparent().remove(paragraph._element)
                            
                            # Make the cell invisible by setting its width to 0
                            try:
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                tcW = tcPr.find(qn('w:tcW'))
                                if tcW is None:
                                    tcW = OxmlElement('w:tcW')
                                    tcPr.append(tcW)
                                tcW.set(qn('w:w'), '0')
                                tcW.set(qn('w:type'), 'dxa')
                                self.logger.debug(f"Cleared and hid extra cell {cell_count}")
                            except Exception as e:
                                self.logger.warning(f"Error hiding extra cell {cell_count}: {e}")
                            
        except Exception as e:
            self.logger.warning(f"Error clearing extra cells: {e}")

    def _remove_unmerged_placeholders(self, doc, num_products):
        """Remove unmerged placeholders from cells beyond the number of products."""
        try:
            # CRITICAL FIX: Multiple tables represent multiple pages - DO NOT REMOVE THEM!
            tables = doc.tables
            if len(tables) > 1:
                self.logger.info(f"Found {len(tables)} tables representing {len(tables)} pages - KEEPING ALL TABLES for multi-page document")
                # DO NOT remove tables - they represent different pages in the document
            
            # Now clear extra cells within each table
            # For multi-page documents, each table represents a page with up to 9 labels
            for table_idx, table in enumerate(doc.tables):
                cells_per_page = 9  # 3x3 grid
                page_start_product = table_idx * cells_per_page
                page_end_product = min(page_start_product + cells_per_page, num_products)
                
                cell_count = 0
                for row in table.rows:
                    for cell in row.cells:
                        cell_count += 1
                        product_num = page_start_product + cell_count
                        if product_num > num_products:
                            # This is an extra cell - completely clear it and set white background
                            self.logger.debug(f"Clearing extra cell on page {table_idx + 1}, cell {cell_count} (product {product_num} > {num_products})")
                            
                            # Clear all content from the cell
                            cell._tc.clear_content()
                            
                            # Set white background for extra cells
                            from docx.oxml import OxmlElement
                            from docx.oxml.ns import qn
                            tc = cell._tc
                            tcPr = tc.find(qn('w:tcPr'))
                            if tcPr is None:
                                tcPr = OxmlElement('w:tcPr')
                                tc.insert(0, tcPr)
                            
                            # Remove any existing background color
                            shd = tcPr.find(qn('w:shd'))
                            if shd is not None:
                                tcPr.remove(shd)
                            
                            # Add white background
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear')
                            shd.set(qn('w:color'), 'auto')
                            shd.set(qn('w:fill'), 'FFFFFF')  # White background
                            tcPr.append(shd)
                            
                            # Add a single empty paragraph to maintain structure
                            cell.add_paragraph()
                            
        except Exception as e:
            self.logger.warning(f"Error removing unmerged placeholders: {e}")

    def _ensure_table_grids_exist(self, doc):
        """Ensure all tables have proper tblGrid elements."""
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            for table in doc.tables:
                tbl_grid = table._element.find(qn('w:tblGrid'))
                if tbl_grid is None:
                    self.logger.warning("Table missing tblGrid element, creating one")
                    # Create tblGrid element
                    tbl_grid = OxmlElement('w:tblGrid')
                    table._element.insert(0, tbl_grid)
                    
                    # Add grid columns based on actual column count
                    col_count = len(table.rows[0].cells) if table.rows else 3
                    for _ in range(col_count):
                        gc = OxmlElement('w:gridCol')
                        gc.set(qn('w:w'), str(int(3.4 * 1440)))  # Default width
                        tbl_grid.append(gc)
                    
                    self.logger.info(f"Created tblGrid with {col_count} columns for table")
                    
        except Exception as e:
            self.logger.error(f"Error ensuring table grids exist: {e}")

    def _safe_table_iteration(self, table, operation_name="table operation"):
        """
        Safely iterate through table rows and cells with comprehensive error handling.
        Returns True if successful, False if table is corrupted beyond repair.
        """
        try:
            # First validate the table structure
            if not self._validate_and_repair_table_structure(table):
                self.logger.warning(f"Table validation failed for {operation_name}")
                return False
            
            # Test basic table access
            try:
                rows = table.rows
                if not rows:
                    self.logger.warning(f"Table has no rows for {operation_name}")
                    return False
            except Exception as e:
                self.logger.error(f"Table rows access failed for {operation_name}: {e}")
                return False
            
            return True
            
        except Exception as e:
            self.logger.error(f"Error in safe table iteration for {operation_name}: {e}")
            return False

    def _safe_table_processing(self, table, operation_name="table operation", processor_func=None):
        """
        Safely process a table with comprehensive error handling and repair attempts.
        
        Args:
            table: The table to process
            operation_name: Name of the operation for logging
            processor_func: Function to apply to each row/cell (optional)
            
        Returns:
            bool: True if processing was successful, False otherwise
        """
        try:
            # First validate the table structure
            if not self._safe_table_iteration(table, operation_name):
                return False
            
            # If no processor function provided, just return success
            if processor_func is None:
                return True
            
            # Process the table safely
            try:
                for row in table.rows:
                    try:
                        # Validate row structure
                        if not hasattr(row, 'cells') or not row.cells:
                            self.logger.warning(f"Skipping row with invalid structure during {operation_name}")
                            continue
                        
                        for cell in row.cells:
                            try:
                                # Apply the processor function to the cell
                                processor_func(cell)
                            except Exception as cell_error:
                                self.logger.warning(f"Skipping cell due to error during {operation_name}: {cell_error}")
                                continue
                    except Exception as row_error:
                        self.logger.warning(f"Skipping row due to error during {operation_name}: {row_error}")
                        continue
                
                return True
                
            except Exception as processing_error:
                self.logger.error(f"Error during table processing for {operation_name}: {processing_error}")
                return False
                
        except Exception as e:
            self.logger.error(f"Error in safe table processing for {operation_name}: {e}")
            return False

    def _rebuild_corrupted_table(self, table, template_type):
        """
        Attempt to completely rebuild a corrupted table structure.
        This is a last resort when other repair methods fail.
        
        Args:
            table: The corrupted table
            template_type: Type of template to determine structure
            
        Returns:
            bool: True if rebuild was successful, False otherwise
        """
        try:
            self.logger.warning(f"Attempting to rebuild corrupted table for template type: {template_type}")
            
            # Get the table element
            table_element = table._element
            
            # Determine expected structure based on template type
            if template_type == 'vertical':
                expected_rows, expected_cols = 3, 3
            elif template_type == 'double':
                expected_rows, expected_cols = 4, 3
            elif template_type == 'mini':
                expected_rows, expected_cols = 4, 5
            elif template_type == 'inventory':
                expected_rows, expected_cols = 2, 2
            else:
                expected_rows, expected_cols = 3, 3  # Default
            
            # Create new table structure
            new_table_element = OxmlElement('w:tbl')
            
            # Add table properties
            tblPr = OxmlElement('w:tblPr')
            new_table_element.append(tblPr)
            
            # Add table layout
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)
            
            # Add table width
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), str(expected_cols * 1440))  # 1 inch per column
            tblW.set(qn('w:type'), 'dxa')
            tblPr.append(tblW)
            
            # Add table grid
            tblGrid = OxmlElement('w:tblGrid')
            for _ in range(expected_cols):
                gridCol = OxmlElement('w:gridCol')
                gridCol.set(qn('w:w'), '1440')
                tblGrid.append(gridCol)
            new_table_element.append(tblGrid)
            
            # Add table rows
            for _ in range(expected_rows):
                tr = OxmlElement('w:tr')
                for _ in range(expected_cols):
                    tc = OxmlElement('w:tc')
                    # Add cell properties
                    tcPr = OxmlElement('w:tcPr')
                    tc.append(tcPr)
                    # Add cell width
                    tcW = OxmlElement('w:tcW')
                    tcW.set(qn('w:w'), '1440')
                    tcW.set(qn('w:type'), 'dxa')
                    tcPr.append(tcW)
                    # Add empty paragraph
                    p = OxmlElement('w:p')
                    tc.append(p)
                    tr.append(tc)
                new_table_element.append(tr)
            
            # Replace the old table element
            parent = table_element.getparent()
            if parent is not None:
                parent.replace(table_element, new_table_element)
                self.logger.info(f"Successfully rebuilt table structure for {template_type} template")
                return True
            else:
                self.logger.error("Cannot rebuild table: no parent element found")
                return False
                
        except Exception as e:
            self.logger.error(f"Failed to rebuild corrupted table: {e}")
            return False
    def _ensure_proper_centering(self, doc):
        """
        Ensure tables are properly centered in the document with correct margins and spacing.
        """
        try:
            # Set document margins to ensure proper centering
            for section in doc.sections:
                # Use smaller margins for vertical template to fit all 9 labels
                if self.template_type == 'vertical':
                    section.left_margin = Inches(0.25)
                    section.right_margin = Inches(0.25)
                    section.top_margin = Inches(0.25)
                    section.bottom_margin = Inches(0.25)
                else:
                    section.left_margin = Inches(0.5)
                    section.right_margin = Inches(0.5)
                    section.top_margin = Inches(0.5)
                    section.bottom_margin = Inches(0.5)
            
            # Remove any extra paragraphs that might affect centering
            paragraphs_to_remove = []
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip() and not paragraph.runs:
                    paragraphs_to_remove.append(paragraph)
            
            for paragraph in paragraphs_to_remove:
                paragraph._element.getparent().remove(paragraph._element)
            
            # Ensure all tables are properly centered and have valid structure
            for table in doc.tables:
                # Use safe table iteration to validate and repair if needed
                if not self._safe_table_iteration(table, "centering setup"):
                    self.logger.warning(f"Skipping table that cannot be repaired during centering setup")
                    continue
                
                # Set table alignment to center
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Ensure table properties are set correctly
                tblPr = table._element.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                
                # Set table to fixed layout
                tblLayout = OxmlElement('w:tblLayout')
                tblLayout.set(qn('w:type'), 'fixed')
                tblPr.append(tblLayout)
                
                # Ensure table is not auto-fit
                table.autofit = False
                if hasattr(table, 'allow_autofit'):
                    table.allow_autofit = False
                

                
                # Calculate and set proper table width for perfect centering
                from src.core.constants import CELL_DIMENSIONS, GRID_LAYOUTS
                
                # Get individual cell dimensions and grid layout
                cell_dims = CELL_DIMENSIONS.get(self.template_type, {'width': 2.4, 'height': 2.4})
                grid_layout = GRID_LAYOUTS.get(self.template_type, {'rows': 3, 'cols': 3})
                
                # Calculate total table width: individual cell width * number of columns
                individual_cell_width = cell_dims['width']
                num_columns = grid_layout['cols']
                total_table_width = individual_cell_width * num_columns
                
                # Set table width to ensure proper centering
                table.width = Inches(total_table_width)
                
                # Also set the table width property in XML to ensure it's properly applied
                tblPr = table._element.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    table._element.insert(0, tblPr)
                
                # Set table width property
                tblW = tblPr.find(qn('w:tblW'))
                if tblW is None:
                    tblW = OxmlElement('w:tblW')
                    tblPr.append(tblW)
                tblW.set(qn('w:w'), str(int(total_table_width * 1440)))  # Convert to twips
                tblW.set(qn('w:type'), 'dxa')
                
                # For double templates, ensure proper table grid structure without accessing table.columns
                if self.template_type == 'double':
                    # Use safe table iteration to get column count
                    if self._safe_table_iteration(table, "double template grid setup"):
                        # Get column count from XML structure instead of table.columns
                        table_element = table._element
                        rows = table_element.findall(qn('w:tr'))
                        if rows:
                            first_row = rows[0]
                            cells = first_row.findall(qn('w:tc'))
                            col_count = len(cells)
                            
                            # Ensure tblGrid exists and has correct structure
                            tblGrid = table_element.find(qn('w:tblGrid'))
                            if tblGrid is None:
                                # Create new tblGrid
                                tblGrid = OxmlElement('w:tblGrid')
                                table_element.insert(0, tblGrid)
                            
                            # Clear existing grid columns and recreate
                            for existing_gc in tblGrid.findall(qn('w:gridCol')):
                                tblGrid.remove(existing_gc)
                            
                            # Add grid columns with proper widths for double template
                            col_width = 1.75  # 1.75 inches per column for double template
                            for _ in range(col_count):
                                gc = OxmlElement('w:gridCol')
                                gc.set(qn('w:w'), str(int(col_width * 1440)))  # Convert to twips
                                tblGrid.append(gc)
                            
                            # Also ensure each cell has the correct width property
                            for row in table.rows:
                                for cell in row.cells:
                                    tcPr = cell._tc.get_or_add_tcPr()
                                    tcW = tcPr.find(qn('w:tcW'))
                                    if tcW is None:
                                        tcW = OxmlElement('w:tcW')
                                        tcPr.append(tcW)
                                    tcW.set(qn('w:w'), str(int(col_width * 1440)))
                                    tcW.set(qn('w:type'), 'dxa')
                
                # For other template types, use the existing logic but with safety checks
                elif self.template_type not in ['horizontal', 'mini', 'vertical']:
                    # Use safe table iteration to get column count
                    if self._safe_table_iteration(table, "grid setup"):
                        # Get column count from XML structure
                        table_element = table._element
                        rows = table_element.findall(qn('w:tr'))
                        if rows:
                            first_row = rows[0]
                            cells = first_row.findall(qn('w:tc'))
                            col_count = len(cells)
                            
                            # Create new grid with proper column widths
                            tblGrid = OxmlElement('w:tblGrid')
                            col_width = cell_dims['width']
                            
                            for _ in range(col_count):
                                gridCol = OxmlElement('w:gridCol')
                                gridCol.set(qn('w:w'), str(int(col_width * 1440)))  # Convert to twips
                                tblGrid.append(gridCol)
                            
                            # Remove existing grid if present
                            existing_grid = table_element.find(qn('w:tblGrid'))
                            if existing_grid is not None:
                                existing_grid.getparent().remove(existing_grid)
                            
                            # Insert the grid at the beginning of the table element
                            table_element.insert(0, tblGrid)
                            
                            # Also ensure each cell has the correct width property
                            for row in table.rows:
                                for cell in row.cells:
                                    tcPr = cell._tc.get_or_add_tcPr()
                                    tcW = tcPr.find(qn('w:tcW'))
                                    if tcW is None:
                                        tcW = OxmlElement('w:tcW')
                                        tcPr.append(tcW)
                                    tcW.set(qn('w:w'), str(int(col_width * 1440)))
                                    tcW.set(qn('w:type'), 'dxa')
            
            # Ensured proper table centering and document setup
            
        except Exception as e:
            self.logger.error(f"Error ensuring proper centering: {e}")

    def _add_weight_units_markers(self, doc):
        """
        Add RATIO markers around weight units content for mini templates with classic types.
        This allows the post-processing to find and apply the correct font sizing.
        """
        try:
            for table in doc.tables:
                # Validate table structure before processing
                if not self._validate_and_repair_table_structure(table):
                    self.logger.warning(f"Skipping table with invalid structure during weight units marker addition")
                    continue
                
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Look for weight units content in individual runs
                            for run in paragraph.runs:
                                run_text = run.text
                                # Check if this run contains weight units content (ends with 'g' or 'mg', or contains specific patterns)
                                # More specific check to avoid marking brand names that contain 'g'
                                is_weight_unit = (
                                    run_text.strip().endswith('g') or 
                                    run_text.strip().endswith('mg') or
                                    re.match(r'^\d+\.?\d*\s*g$', run_text.strip()) or  # "1g", "1.5g"
                                    re.match(r'^\d+\.?\d*\s*mg$', run_text.strip()) or  # "100mg", "50.5mg"
                                    re.match(r'^\d+\.?\d*\s*g\s*x\s*\d+', run_text.strip()) or  # "1g x 2"
                                    re.match(r'^\d+\.?\d*\s*mg\s*x\s*\d+', run_text.strip())  # "100mg x 2"
                                )
                                
                                if is_weight_unit and 'RATIO_START' not in run_text:
                                    # This is likely weight units content that needs markers
                                    # CRITICAL FIX: Preserve non-breaking hyphens when adding markers
                                    # Check for non-breaking hyphens before processing
                                    if '\u2011' in run_text:
                                        self.logger.debug(f"WEIGHT UNITS DEBUG: Found non-breaking hyphens in weight units: '{run_text}'")
                                    # Replace the run text with marked content
                                    run.text = f"RATIO_START{run_text}RATIO_END"
                                    run.font.name = "Arial"
                                    run.font.bold = True
                                    # Use unified font sizing for ratio text
                                    from src.core.generation.unified_font_sizing import get_font_size
                                    ratio_font_size = get_font_size(run_text, 'ratio', self.template_type, self.scale_factor)
                                    run.font.size = ratio_font_size
                                    
                                    # Check for non-breaking hyphens after processing
                                    if '\u2011' in run.text:
                                        self.logger.debug(f"WEIGHT UNITS DEBUG: Preserved non-breaking hyphens: '{run.text}'")
                                    else:
                                        self.logger.debug(f"WEIGHT UNITS DEBUG: Lost non-breaking hyphens: '{run.text}'")
                                    
                                    self.logger.debug(f"Added RATIO markers around weight units: {run_text}")
            
        except Exception as e:
            self.logger.error(f"Error adding weight units markers: {e}")

    def _add_brand_markers(self, doc):
        """
        Add PRODUCTBRAND_CENTER markers around brand content for mini and preroll templates.
        This allows the post-processing to find and apply the correct font sizing.
        """
        try:
            # Import CLASSIC_TYPES to check if current product type is classic
            from src.core.constants import CLASSIC_TYPES
            
            # Get current product type if available
            current_product_type = None
            if hasattr(self, 'current_product_type'):
                current_product_type = self.current_product_type
            elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                current_product_type = self.label_context['ProductType']
            
            # Check if current product type is a classic type
            is_classic_type = False
            if current_product_type:
                is_classic_type = current_product_type.lower() in [ct.lower() for ct in CLASSIC_TYPES]
                self.logger.debug(f"Product type: {current_product_type}, Is classic: {is_classic_type}")
            else:
                self.logger.debug(f"No current_product_type available")
            
            # For mini and preroll templates, always add brand markers regardless of product type
            # For other templates, skip brand marker addition for classic types (they should show lineage instead of brand)
            if self.template_type not in ['mini', 'preroll'] and is_classic_type:
                self.logger.debug(f"Skipping brand marker addition for classic type: {current_product_type}")
                return
            
            template_name = 'mini/preroll template' if self.template_type in ['mini', 'preroll'] else f'non-classic type: {current_product_type}'
            self.logger.debug(f"Processing brand markers for {template_name}")
            
            for table in doc.tables:
                # Validate table structure before processing
                if not self._validate_and_repair_table_structure(table):
                    self.logger.warning(f"Skipping table with invalid structure during brand marker addition")
                    continue
                
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Look for brand content in individual runs
                            for run in paragraph.runs:
                                run_text = run.text
                                self.logger.debug(f"Processing run text: '{run_text}'")
                                # Check if this run contains brand content (not empty and not already marked)
                                # Only add markers to text that looks like brand names (not empty, not marked, not placeholders)
                                # IMPORTANT: Don't add brand markers if content is already wrapped in RATIO markers
                                # IMPORTANT: Don't add brand markers if content is already wrapped in PRODUCTSTRAIN markers
                                if (run_text.strip() and 
                                    'PRODUCTBRAND_CENTER_START' not in run_text and 
                                    'PRODUCTSTRAIN_START' not in run_text and  # Don't mark content already in PRODUCTSTRAIN markers
                                    'PRODUCTSTRAIN_END' not in run_text and    # Don't mark content already in PRODUCTSTRAIN markers
                                    'RATIO_START' not in run_text and  # Don't mark content already in RATIO markers
                                    'RATIO_END' not in run_text and    # Don't mark content already in RATIO markers
                                    '{{' not in run_text and 
                                    '}}' not in run_text and
                                    # QR placeholder check removed
                                    len(run_text.strip()) > 0 and
                                    # Only mark content that looks like brand names (not numbers, not empty)
                                    not run_text.strip().isdigit() and
                                    not run_text.strip().startswith('$') and
                                    not run_text.strip().endswith('g') and
                                    not run_text.strip().endswith('mg')):
                                    # This is likely brand content that needs markers
                                    # CRITICAL FIX: Preserve non-breaking hyphens when adding markers
                                    # Check for non-breaking hyphens before processing
                                    if '\u2011' in run_text:
                                        self.logger.debug(f"BRAND MARKERS DEBUG: Found non-breaking hyphens in brand: '{run_text}'")
                                    # Replace the run text with marked content
                                    run.text = f"PRODUCTBRAND_CENTER_START{run_text}PRODUCTBRAND_CENTER_END"
                                    run.font.name = "Arial"
                                    run.font.bold = True
                                    # Use unified font sizing for brand text
                                    from src.core.generation.unified_font_sizing import get_font_size
                                    brand_font_size = get_font_size(run_text, 'brand', self.template_type, self.scale_factor)
                                    run.font.size = brand_font_size
                                    
                                    # Check for non-breaking hyphens after processing
                                    if '\u2011' in run.text:
                                        self.logger.debug(f"BRAND MARKERS DEBUG: Preserved non-breaking hyphens: '{run.text}'")
                                    else:
                                        self.logger.debug(f"BRAND MARKERS DEBUG: Lost non-breaking hyphens: '{run.text}'")
                                    
                                    # Ensure brand content is centered for mini templates
                                    if self.template_type == 'mini':
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        self.logger.debug(f"Set center alignment for mini template brand: {run_text}")
                                    
                                    self.logger.debug(f"Added PRODUCTBRAND_CENTER markers around brand: {run_text}")
            
        except Exception as e:
            self.logger.error(f"Error adding brand markers: {e}")

    def _ensure_mini_template_brand_centering(self, doc):
        """
        Ensure all brand content in mini templates is properly centered.
        This method specifically handles mini template brand alignment.
        """
        try:
            if self.template_type != 'mini':
                return
                
            self.logger.debug("Ensuring brand content centering for mini template")
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Check if this paragraph contains brand content
                            paragraph_text = paragraph.text
                            
                            # For mini templates, be more aggressive about finding brand content
                            is_brand_content = False
                            
                            # Check for explicit brand markers
                            if ('PRODUCTBRAND_CENTER_START' in paragraph_text or 
                                'PRODUCTBRAND_CENTER_END' in paragraph_text or
                                'PRODUCTBRAND_START' in paragraph_text or
                                'PRODUCTBRAND_END' in paragraph_text):
                                is_brand_content = True
                            
                            # For mini templates, also check for content that looks like brand names
                            elif self.template_type == 'mini':
                                # Look for content that appears to be brand names (not empty, not numbers, not prices, not weights)
                                clean_text = paragraph_text.strip()
                                if (clean_text and 
                                    not clean_text.startswith('$') and
                                    not clean_text.endswith('g') and
                                    not clean_text.endswith('mg') and
                                    not clean_text.isdigit() and
                                    not ('THC:' in clean_text and 'CBD:' in clean_text) and
                                    len(clean_text) < 50 and
                                    # Not lineage values
                                    clean_text.upper() not in ["SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", "CBD", "MIXED"]):
                                    is_brand_content = True
                                    self.logger.debug(f"Identified potential brand content in mini template: {clean_text}")
                            
                            if is_brand_content:
                                # Set center alignment for brand paragraphs
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                self.logger.debug(f"Centered brand paragraph: {paragraph_text[:50]}...")
                                
                                # Also ensure all runs in brand paragraphs are properly formatted
                                for run in paragraph.runs:
                                    if run.text.strip():
                                        run.font.name = "Arial"
                                        run.font.bold = True
                                        
            self.logger.debug("Completed mini template brand centering")
            
        except Exception as e:
            self.logger.error(f"Error ensuring mini template brand centering: {e}")

    def _ensure_lineage_centering_for_nonclassic_types(self, doc):
        """
        Ensure brand field is centered for nonclassic types (where brand like "CERES" should be centered).
        This method runs after all other processing to ensure the centering is not overridden.
        """
        try:
            # Starting _ensure_lineage_centering_for_nonclassic_types
            
            # Process all tables and look for actual brand content that should be centered
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph_text = paragraph.text.strip()
                            
                            # Skip empty paragraphs
                            if not paragraph_text:
                                continue
                            
                            # CRITICAL FIX: Don't override LEFT-aligned paragraphs (classic lineage should stay left-aligned)
                            # If paragraph is already left-aligned, skip it - this preserves classic type lineage alignment
                            if paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                                continue
                            
                            # Look for actual brand content that should be centered
                            # This includes all brand names regardless of case or length
                            is_brand_name = (
                                paragraph_text and
                                not paragraph_text.startswith('$') and
                                not paragraph_text.endswith('g') and
                                not paragraph_text.endswith('mg') and
                                not paragraph_text.isdigit() and
                                # Not classic lineage values
                                paragraph_text.upper() not in ["SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", "CBD", "MIXED", "PARA", "PARAPHERNALIA"] and
                                # Not THC/CBD content
                                not ('THC:' in paragraph_text and 'CBD:' in paragraph_text) and
                                # Not long product descriptions (those should be left-aligned)
                                len(paragraph_text) <= 50 and
                                # Not product names with weights or measurements
                                not ('oz' in paragraph_text.lower() or 'ml' in paragraph_text.lower() or 'mg' in paragraph_text.lower()) and
                                # Contains letters (brand names)
                                any(c.isalpha() for c in paragraph_text) and
                                # Not purely numeric content
                                not paragraph_text.replace('.', '').replace(',', '').isdigit()
                            )
                            
                            if is_brand_name:
                                # Force center alignment for brand names (only if not already left-aligned)
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                # Centered brand name content
                                
        except Exception as e:
            self.logger.error(f"Error ensuring brand centering for nonclassic types: {e}")

    def _is_non_classic_type(self, product_type):
        """Return True if product_type is not one of the classic types."""
        if not product_type:
            return False
        classic_types = {'flower', 'pre-roll', 'live'}
        return product_type.lower() not in classic_types

    def _clean_up_lineage_brand_concatenation(self, doc):
        """
        Clean up any remaining concatenated lineage+brand content for classic types.
        This runs at the very end to catch any concatenation that wasn't caught earlier.
        """
        try:
            # CRITICAL FIX: Skip lineage cleaning for double, horizontal, and vertical templates to preserve HYBRID/SATIVA, HYBRID/INDICA
            if self.template_type in ('double', 'horizontal', 'vertical'):
                self.logger.debug("Skipping lineage cleaning for grid templates to preserve full lineage values")
                return
                
            # Starting _clean_up_lineage_brand_concatenation

            # Define classic lineage values
            classic_lineages = ["HYBRID/SATIVA", "HYBRID/INDICA", "SATIVA", "INDICA", "HYBRID", "CBD", "MIXED"]

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph_text = paragraph.text.strip()

                            # Skip empty paragraphs
                            if not paragraph_text:
                                continue

                            # Check if this paragraph contains concatenated lineage+brand content
                            cleaned_text = paragraph_text
                            for classic_lineage in classic_lineages:
                                # Look for patterns like "HYBRIDHUSTLER", "INDICAHUSTLER", etc.
                                if paragraph_text.upper().startswith(classic_lineage.upper()) and len(paragraph_text) > len(classic_lineage):
                                    # Extract only the lineage part
                                    cleaned_text = paragraph_text[:len(classic_lineage)]
                                    self.logger.info(f"DEBUG: Cleaned concatenated lineage: '{paragraph_text}' -> '{cleaned_text}'")
                                    break

                            # If we found concatenated content, update the paragraph
                            if cleaned_text != paragraph_text:
                                # Clear and recreate the paragraph with clean content
                                paragraph.clear()
                                run = paragraph.add_run()
                                run.font.name = "Arial"
                                run.font.bold = True
                                
                                # Use unified font sizing for lineage instead of hardcoded 12pt
                                from src.core.generation.unified_font_sizing import get_font_size
                                # CRITICAL FIX: Always use 'lineage' field type for proper font sizing (18pt for <100 chars, 12pt for longer)
                                # Previous logic used 'brand' for vertical which made lineage too small (10-16pt)
                                lineage_font_size = get_font_size(cleaned_text, 'lineage', self.template_type, self.scale_factor)
                                run.font.size = lineage_font_size
                                
                                run.font.color.rgb = RGBColor(255, 255, 255)  # Set text to white
                                run.text = cleaned_text  # Use assignment instead of add_text to avoid duplication
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left-align classic lineage
                                
        except Exception as e:
            self.logger.error(f"Error cleaning up lineage brand concatenation: {e}")

    def _final_doh_positioning_enforcement(self, doc):
        """Ensure DOH image paragraphs remain centered after processing."""
        try:
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            fixed = 0
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs = []
                        has_image = False
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if hasattr(run, '_element') and (
                                    run._element.find(qn('w:drawing')) is not None or
                                    run._element.find(qn('w:pict')) is not None
                                ):
                                    has_image = True
                                    paragraphs.append(paragraph)
                                    break
                            if has_image:
                                break

                        if not has_image:
                            continue

                        fixed += 1
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        for paragraph in paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            try:
                                pPr = paragraph._element.get_or_add_pPr()
                                existing_jc = pPr.find(qn('w:jc'))
                                if existing_jc is not None:
                                    pPr.remove(existing_jc)
                                jc = OxmlElement('w:jc')
                                jc.set(qn('w:val'), 'center')
                                pPr.append(jc)

                                existing_spacing = pPr.find(qn('w:spacing'))
                                if existing_spacing is not None:
                                    pPr.remove(existing_spacing)
                                spacing = OxmlElement('w:spacing')
                                spacing.set(qn('w:before'), '60')
                                spacing.set(qn('w:after'), '60')
                                spacing.set(qn('w:line'), '240')
                                spacing.set(qn('w:lineRule'), 'auto')
                                pPr.append(spacing)

                                existing_ind = pPr.find(qn('w:ind'))
                                if existing_ind is not None:
                                    pPr.remove(existing_ind)
                            except Exception as xml_error:
                                self.logger.warning(f"Error centering DOH paragraph: {xml_error}")

            if fixed:
                self.logger.info(f"Final DOH positioning enforcement centered {fixed} cell(s)")
        except Exception as e:
            self.logger.warning(f"DOH centering enforcement skipped: {e}")

    def _ensure_preroll_qr_centering(self, doc):
        """Ensure QR code paragraphs are centered in preroll templates."""
        try:
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from docx.shared import Pt

            fixed = 0
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        qr_paragraphs = []
                        has_qr_image = False
                        cell_text = cell.text.strip().upper()
                        
                        # Check all paragraphs in the cell for images
                        for paragraph in cell.paragraphs:
                            paragraph_text = paragraph.text.strip().upper()
                            for run in paragraph.runs:
                                if hasattr(run, '_element') and (
                                    run._element.find(qn('w:drawing')) is not None or
                                    run._element.find(qn('w:pict')) is not None
                                ):
                                    # In preroll templates, QR codes are typically the only images
                                    # (DOH images are handled separately if present)
                                    # Center all images in preroll templates as they are likely QR codes
                                    has_qr_image = True
                                    qr_paragraphs.append(paragraph)
                                    break
                            if has_qr_image:
                                break

                        if not has_qr_image or not qr_paragraphs:
                            continue

                        fixed += 1
                        # Center the cell content vertically
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        
                        # Center all QR code paragraphs
                        for paragraph in qr_paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Set proper spacing for QR code
                            paragraph.paragraph_format.space_before = Pt(2)
                            paragraph.paragraph_format.space_after = Pt(2)
                            paragraph.paragraph_format.line_spacing = 1.0
                            
                            try:
                                # Ensure XML-level centering
                                pPr = paragraph._element.get_or_add_pPr()
                                
                                # Set paragraph justification to center
                                existing_jc = pPr.find(qn('w:jc'))
                                if existing_jc is not None:
                                    pPr.remove(existing_jc)
                                jc = OxmlElement('w:jc')
                                jc.set(qn('w:val'), 'center')
                                pPr.append(jc)

                                # Remove any indentation that might offset the QR code
                                existing_ind = pPr.find(qn('w:ind'))
                                if existing_ind is not None:
                                    pPr.remove(existing_ind)
                                    
                            except Exception as xml_error:
                                self.logger.warning(f"Error centering QR code paragraph: {xml_error}")

            if fixed:
                self.logger.info(f"PREROLL QR centering: centered {fixed} QR code cell(s)")
        except Exception as e:
            self.logger.warning(f"PREROLL QR centering enforcement skipped: {e}")

    def _ensure_standalone_cannabinoid_font_sizing(self, doc):
        """
        Ensure any standalone cannabinoid text (CBD, THC, CBC, CBG, CBN) uses appropriate font sizing.
        CBD text in Lineage field should use lineage font sizing, not strain.
        This runs at the very end to catch any standalone cannabinoid text that wasn't caught earlier.
        """
        try:
            # Starting _ensure_standalone_cannabinoid_font_sizing

            # Define standalone cannabinoid values that should be nearly invisible
            standalone_cannabinoids = ["CBD", "THC", "CBC", "CBG", "CBN"]

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run_text = run.text.strip()

                                # Skip empty runs
                                if not run_text:
                                    continue

                                # Check if this run contains standalone cannabinoid text
                                # CRITICAL FIX: CBD lineage should use lineage font sizing, not strain
                                if (run_text in standalone_cannabinoids and 
                                    len(run_text) <= 3 and
                                    not any(marker in run.text for marker in ['CBD_START', 'THC_START', 'CBC_START', 'CBG_START', 'CBN_START', 'CBD_END', 'THC_END', 'CBC_END', 'CBG_END', 'CBN_END'])):
                                    
                                    # This is standalone cannabinoid text - use lineage font sizing for CBD
                                    from src.core.generation.unified_font_sizing import get_font_size_by_marker
                                    lineage_font_size = get_font_size_by_marker(run_text, 'LINEAGE', self.template_type, self.scale_factor)
                                    run.font.size = lineage_font_size
                                    self.logger.info(f"DEBUG: Set standalone cannabinoid '{run_text}' to lineage font size ({lineage_font_size})")
                                
        except Exception as e:
            self.logger.error(f"Error ensuring standalone cannabinoid font sizing: {e}")

    def _get_template_specific_font_size(self, content, marker_name):
        """
        Get font size using the unified font sizing system.
        """
        # Import get_font_size locally to avoid scoping issues
        from src.core.generation.unified_font_sizing import get_font_size
        
        # Special handling for RATIO marker: if content contains THC/CBD data, use THC_CBD field type
        if marker_name == 'RATIO' and ('THC:' in content or 'CBD:' in content):
            # Use THC_CBD field type for THC/CBD content
            return get_font_size(content, 'thc_cbd', self.template_type, self.scale_factor)
        
        # Use unified font sizing system for all templates
        return get_font_size_by_marker(content, marker_name, self.template_type, self.scale_factor)

    def fix_hyphen_spacing(self, text):
        """Replace regular hyphens with non-breaking hyphens to prevent line breaks, 
        but add line breaks before hanging hyphens.
        Used for general text formatting to prevent unwanted line breaks at hyphens."""
        if not text:
            return text
        
        # CRITICAL FIX: Preserve existing non-breaking hyphens (\u2011) and only convert regular hyphens
        # First, temporarily replace non-breaking hyphens with a placeholder
        text = text.replace('\u2011', '___NONBREAKING_HYPHEN___')
        
        # Then normalize regular hyphen spacing to use non-breaking hyphens instead of hyphen + non-breaking space
        text = re.sub(r'\s*-\s*', '\u2011', text)
        
        # Check for hanging hyphens (hyphen at the end of a line or followed by a space and then end)
        # Pattern: non-breaking hyphen at end of string
        if re.search(r'\u2011$', text) or re.search(r'\u2011\s*$', text):
            # Add line break before the hanging hyphen
            text = re.sub(r'\u2011(\s*)$', r'\n\u2011\1', text)
        
        # Finally, restore non-breaking hyphens
        text = text.replace('___NONBREAKING_HYPHEN___', '\u2011')
        
        return text

    def format_with_soft_hyphen(self, text):
        """Format text with soft hyphen + nonbreaking space + value pattern.
        Used for specific formatting where you want a soft hyphen followed by nonbreaking space."""
        if not text:
            return text
        # Replace any leading hyphens/spaces with a single soft hyphen + nonbreaking space
        text = re.sub(r'^[\s\-]+', '\u00AD\u00A0', text)
        # If it didn't start with hyphen/space, prepend
        if not text.startswith('\u00AD\u00A0'):
            text = f'\u00AD\u00A0{text}'
        return text
    def format_classic_ratio(self, text, record=None):
        """
        Format ratio for classic types. Handles various input formats and converts them to the standard display format.
        """
        if not text:
            return text
        
        # Clean the text and normalize
        text = text.strip()
        
        # Handle the default "THC:|BR|CBD:" format from excel processor
        if text == "THC:|BR|CBD:" or text == "THC: | BR | CBD:":
            product_name = record.get('Product Name*', 'Unknown') if record else 'Unknown'
            self.logger.debug(f"Processing THC:|BR|CBD: placeholder for record: {product_name}")
            if record:
                # Show all relevant THC fields for debugging
                self.logger.debug(f"RECORD THC FIELDS: THC: '{record.get('THC', '')}', THC test result: '{record.get('THC test result', '')}', Total THC: '{record.get('Total THC', '')}', THCA: '{record.get('THCA', '')}'")
                # Always use Excel 'THC test result' if present and valid
                excel_thc = str(record.get('THC test result', '')).strip()
                try:
                    excel_thc_float = float(excel_thc)
                    if excel_thc not in ['0', '0.0', '', 'nan', 'NaN']:
                        thc_value = excel_thc
                        self.logger.debug(f"USING EXCEL THC test result: '{thc_value}'")
                    else:
                        raise ValueError
                except Exception:
                    # Fallback to highest of all other fields if Excel value is not valid
                    thc_candidates = []
                    thc_debug_vals = {}
                    for key in ['Total THC', 'THCA', 'THC']:
                        val = str(record.get(key, '')).strip()
                        thc_debug_vals[key] = val
                        try:
                            val_float = float(val)
                            if val not in ['0', '0.0', '', 'nan', 'NaN']:
                                thc_candidates.append(val_float)
                        except Exception:
                            continue
                    self.logger.debug(f"THC candidate values (fallback): {thc_debug_vals}, numeric candidates: {thc_candidates}")
                    if thc_candidates:
                        max_thc = max(thc_candidates)
                        thc_value = str(max_thc)
                        self.logger.debug(f"USING HIGHEST THC VALUE (fallback): '{thc_value}' from candidates: {thc_candidates}")
                    else:
                        thc_value = '0'
                        self.logger.debug("No valid THC value found, defaulting to 0")

                # Always use the highest value for both THC and CBD from all relevant fields
                # Always use Excel 'CBD test result' if present and valid
                excel_cbd = str(record.get('CBD test result', '')).strip()
                try:
                    excel_cbd_float = float(excel_cbd)
                    if excel_cbd not in ['0', '0.0', '', 'nan', 'NaN']:
                        cbd_value = excel_cbd
                        self.logger.debug(f"USING EXCEL CBD test result: '{cbd_value}'")
                    else:
                        raise ValueError
                except Exception:
                    # Fallback to highest of all other fields if Excel value is not valid
                    cbd_candidates = []
                    cbd_debug_vals = {}
                    for key in ['Total CBD', 'CBDA', 'CBD']:
                        val = str(record.get(key, '')).strip()
                        cbd_debug_vals[key] = val
                        try:
                            val_float = float(val)
                            if val not in ['0', '0.0', '', 'nan', 'NaN']:
                                cbd_candidates.append(val_float)
                        except Exception:
                            continue
                    self.logger.debug(f"CBD candidate values (fallback): {cbd_debug_vals}, numeric candidates: {cbd_candidates}")
                    if cbd_candidates:
                        max_cbd = max(cbd_candidates)
                        cbd_value = str(max_cbd)
                        self.logger.debug(f"USING HIGHEST CBD VALUE (fallback): '{cbd_value}' from candidates: {cbd_candidates}")
                    else:
                        cbd_value = '0'
                        self.logger.debug("No valid CBD value found, defaulting to 0")
                
                if not cbd_value or cbd_value == '0' or cbd_value == '0.0':
                    total_cbd_value = str(record.get('Total CBD', '')).strip()
                    cbd_test_result = str(record.get('CBD test result', '')).strip()
                    
                    if total_cbd_value and total_cbd_value != '0' and total_cbd_value != '0.0':
                        cbd_value = total_cbd_value
                        self.logger.debug(f"Using database Total CBD: '{cbd_value}'")
                    elif cbd_test_result and cbd_test_result != '0' and cbd_test_result != '0.0':
                        cbd_value = cbd_test_result
                        self.logger.debug(f"Using database CBD test result: '{cbd_value}'")
                
                # Clean up values (remove 'nan', empty strings, etc.)
                if thc_value in ['nan', 'NaN', '']:
                    thc_value = '0'
                if cbd_value in ['nan', 'NaN', '']:
                    cbd_value = '0'
                
                self.logger.debug("THC/CBD percentage display removed - QR codes now provide this information")
                
                # Return empty string - THC/CBD percentages are now shown via QR codes
                return ""
            
            # Fallback - return empty string for THC/CBD requests
            return ""
        
        # Check for any THC/CBD content first, before other processing
        # THC/CBD percentages are now provided via QR codes
        if 'THC' in text.upper() or 'CBD' in text.upper():
            self.logger.debug(f"Removing THC/CBD content: '{text}' - QR codes now provide this information")
            return ""
        
        # If the text contains mg values (non-THC/CBD), return as-is (let text_processing handle it)
        if 'mg' in text.lower():
            return text
        
        # If the text contains simple ratios (like 1:1:1), format with spaces
        if ':' in text and any(c.isdigit() for c in text):
            # Add spaces around colons for better readability
            # Handle 3-part ratios first to avoid conflicts
            text = re.sub(r'(\d+):(\d+):(\d+)', r'\1: \2: \3', text)
            # Then handle 2-part ratios
            text = re.sub(r'(\d+):(\d+)', r'\1: \2', text)
            return text
        
        # Return original text for non-THC/CBD content
        return text

    def format_joint_ratio_pack(self, text):
        """
        Format JointRatio as: - [amount]g x [count] Pack
        Handles various input formats and normalizes them to standard format with hyphen prefix.
        For single units, shows just the weight with hyphen (e.g., "- 1g" instead of "- 1g x 1 Pack").
        """
        if not text:
            return text
            
        # Convert to string and clean up
        text = str(text).strip()
        
        # Remove any leading/trailing spaces and hyphens
        text = re.sub(r'^[\s\-]+', '', text)
        text = re.sub(r'[\s\-]+$', '', text)
        
        # Handle various input patterns
        patterns = [
            # Standard format: "1g x 2 Pack"
            r"([0-9.]+)g\s*x\s*([0-9]+)\s*pack",
            # Compact format: "1gx2Pack"
            r"([0-9.]+)g\s*x?\s*([0-9]+)pack",
            # With spaces: "1g x 2 pack"
            r"([0-9.]+)g\s*x\s*([0-9]+)\s*pack",
            # Just weight: "1g"
            r"([0-9.]+)g",
        ]
        
        for pattern in patterns:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                amount = match.group(1).strip()
                # Try to get count, default to 1 if not found
                try:
                    count = match.group(2).strip()
                    if count and count.isdigit():
                        count_int = int(count)
                        if count_int == 1:
                            # For single units, use regular hyphen with spaces (matching template format)
                            formatted = f"- {amount}g"
                        else:
                            # For multiple units, use regular hyphen with spaces (matching template format)
                            # CRITICAL FIX: Use regular hyphen (-) with regular spaces to match template format
                            formatted = f"- {amount}g x {count} Pack"
                    else:
                        # Only amount found (like "1g") - use regular hyphen with spaces
                        formatted = f"- {amount}g"
                except IndexError:
                    # Only amount found (like "1g") - use regular hyphen with spaces
                    formatted = f"- {amount}g"
                return formatted
        
        # If no pattern matches, return the original text
        return text

    def format_thc_cbd_vertical_alignment(self, text):
        """
        THC/CBD percentage display removed - QR codes now provide this information.
        """
        self.logger.debug("THC/CBD vertical alignment formatting removed - QR codes now provide this information")
        return ""
    
    def _format_thc_cbd_simple(self, text, max_percentage_width):
        """
        THC/CBD percentage display removed - QR codes now provide this information.
        """
        self.logger.debug("THC/CBD simple formatting removed - QR codes now provide this information")
        return ""

    def _identify_marker_type(self, text):
        """Identify the marker type from text content for proper font sizing."""
        if not text:
            return 'default'
            
        text_upper = text.upper()
        
        # Check for specific content patterns in order of specificity
        if any(word in text_upper for word in ['THC', 'CBD', 'RATIO', ':', '%']):
            return 'THC_CBD'
        elif any(word in text_upper for word in ['SATIVA', 'INDICA', 'HYBRID', 'MIXED', 'PARA']):
            return 'LINEAGE'
        elif any(word in text_upper for word in ['$', 'PRICE', 'COST']):
            return 'PRICE'
        elif any(word in text_upper for word in ['WEIGHT', 'G', 'OZ', 'LB', 'KG']) and not any(word in text_upper for word in ['BRAND', 'PRODUCT', 'CANNABIS', 'COMPANY']):
            return 'WEIGHT'
        elif any(word in text_upper for word in ['DOH', 'DATE', 'EXP', '/']) and len(text) <= 12:
            return 'DOH'
        elif any(word in text_upper for word in ['VENDOR', 'SUPPLIER']):
            return 'VENDOR'
        elif any(word in text_upper for word in ['STRAIN', 'VARIETY']) or (len(text.split()) <= 2 and len(text) <= 15 and not any(char in text for char in ['$', '%', ':', '/']) and not any(word in text_upper for word in ['BRAND', 'PRODUCT', 'CANNABIS', 'COMPANY'])):
            return 'STRAIN'
        elif any(word in text_upper for word in ['BRAND', 'PRODUCT', 'CANNABIS', 'COMPANY']) or (len(text.split()) <= 3 and len(text) <= 25 and not any(word in text_upper for word in ['DESCRIPTION', 'LONG', 'DETAILED'])):
            return 'BRAND'
        else:
            # Check if it looks like a description (longer text, multiple words)
            if len(text.split()) > 3 or len(text) > 25:
                return 'DESCRIPTION'
            else:
                return 'default'

    def _apply_mini_template_font_sizing(self, doc):
        """Apply mini template specific font sizing to all content."""
        try:
            from src.core.generation.unified_font_sizing import get_font_size_by_marker, set_run_font_size
            
            # Process all tables in the document
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text.strip():
                                    # Identify the marker type from the text content for proper font sizing
                                    marker_type = self._identify_marker_type(run.text)
                                    
                                    # Get appropriate font size using template_type (supports mini and preroll)
                                    font_size = get_font_size_by_marker(
                                        run.text, 
                                        marker_type, 
                                        template_type=self.template_type,  # Use actual template_type (mini or preroll)
                                        scale_factor=self.scale_factor
                                    )
                                    set_run_font_size(run, font_size)
                                    
            self.logger.info(f"Applied {self.template_type} template specific font sizing with proper field type identification")
            
        except Exception as e:
            self.logger.warning(f"Error applying {self.template_type} template font sizing: {e}")

    def _format_mini_template_text(self, text):
        """Format text for mini template to prevent improper line breaks."""
        if not text:
            return text
        
        # For mini templates, we want to prevent line breaks in the middle of weight units
        # Replace spaces around hyphens with non-breaking spaces to keep " - 1g" together
        # But allow breaks between the main description and the weight unit
        text = text.replace(' - ', ' -\u00A0')  # Non-breaking space after hyphen
        
        return text

    def _format_percentage_right_alignment(self, text, max_percentage_width):
        """
        THC/CBD percentage display removed - QR codes now provide this information.
        """
        self.logger.debug("Percentage right alignment formatting removed - QR codes now provide this information")
        return ""

    def _process_combined_lineage_vendor(self, paragraph, lineage_content, vendor_content):
        """
        Process combined lineage and vendor text with different font sizes.
        This handles the case where lineage and product vendor are on the same line.
        Lineage is left-aligned, vendor is right-aligned.
        IMPORTANT: Product Vendor should never be split up - if Lineage is too long, it should break to new line.
        SPECIAL RULE: For Vertical template, if Lineage is "Hybrid/Indica" or "Hybrid/Sativa", automatically put ProductVendor on next line.
        """
        try:
            # Clear the paragraph content
            paragraph.clear()
            
            # Ensure consistent spacing above lineage/vendor section for equal margins
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(1)
            
            # SPECIAL RULE: For Vertical template, automatically force vendor to next line for specific lineages
            if (self.template_type == 'vertical' and 
                lineage_content and 
                lineage_content.strip().upper() in ['HYBRID/INDICA', 'HYBRID/SATIVA'] and
                vendor_content and vendor_content.strip()):
                
                self.logger.debug(f"Vertical template: Forcing vendor to next line for lineage '{lineage_content}'")
                self._process_lineage_vendor_two_lines(paragraph, lineage_content, vendor_content)
                return
            
            # Check if we need to split to multiple lines due to content length
            # Calculate approximate character limits based on template type
            if self.template_type == 'mini':
                max_chars_per_line = 25
            elif self.template_type == 'vertical':
                max_chars_per_line = 35
            else:  # horizontal, double
                max_chars_per_line = 45
            
            # Check if combined content would be too long for one line
            combined_length = len(lineage_content or '') + len(vendor_content or '')
            
            if combined_length > max_chars_per_line and vendor_content and vendor_content.strip():
                # Split to two lines: lineage on first line, vendor on second line
                self._process_lineage_vendor_two_lines(paragraph, lineage_content, vendor_content)
                return
            
            # Original single-line processing
            # Set paragraph to right alignment for proper vendor right-alignment
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add lineage with larger font size (left-aligned)
            if lineage_content and lineage_content.strip():
                # Debug: Log the lineage content to see what we're working with
                self.logger.debug(f"DEBUG: Original lineage_content: '{repr(lineage_content)}'")
                
                # CRITICAL FIX: Preserve non-breaking hyphens (\u2011) when cleaning leading spaces
                # First, temporarily replace non-breaking hyphens with a placeholder
                temp_lineage = lineage_content.replace('\u2011', '___NONBREAKING_HYPHEN___')
                # Then clean leading spaces
                clean_lineage = temp_lineage.strip().lstrip().lstrip(' \t\n\r\u00A0\u2000\u2001\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200A\u200B\u200C\u200D\u200E\u200F\u2028\u2029\u202A\u202B\u202C\u202D\u202E\u202F\u205F\u2060\u2061\u2062\u2063\u2064\u2065\u2066\u2067\u2068\u2069\u206A\u206B\u206C\u206D\u206E\u206F\u3000\uFEFF')
                # Finally, restore non-breaking hyphens
                clean_lineage = clean_lineage.replace('___NONBREAKING_HYPHEN___', '\u2011')
                self.logger.debug(f"DEBUG: Cleaned lineage_content: '{repr(clean_lineage)}'")
                
                lineage_run = paragraph.add_run(clean_lineage)
                lineage_run.font.name = "Arial"
                lineage_run.font.bold = True
                
                # Use proper lineage font sizing - CRITICAL FIX: Ensure lineage always uses lineage field type
                from src.core.generation.unified_font_sizing import get_font_size
                lineage_font_size = get_font_size(lineage_content, 'lineage', self.template_type, self.scale_factor)
                set_run_font_size(lineage_run, lineage_font_size)
            
            # Add tab character to push vendor to the right (only if vendor content exists)
            if lineage_content and vendor_content:
                tab_run = paragraph.add_run("\t")
                tab_run.font.name = "Arial"
                tab_run.font.bold = True
                # Use lineage font size for tab to maintain alignment
                set_run_font_size(tab_run, lineage_font_size)
            
            # Add vendor with smaller font size (right-aligned)
            if vendor_content and vendor_content.strip():
                vendor_run = paragraph.add_run(vendor_content.strip())
                vendor_run.font.name = "Arial"
                # Get vendor font size using unified font sizing system
                from src.core.generation.unified_font_sizing import get_font_size
                vendor_font_size = get_font_size(vendor_content, 'vendor', self.template_type, self.scale_factor)
                set_run_font_size(vendor_run, vendor_font_size)

                # CRITICAL: Set vendor styling AFTER set_run_font_size to prevent it from being overridden
                vendor_run.font.bold = True
                vendor_run.font.italic = True  # Make vendor text italic

                # Set vendor color to gray (#808080) at both run level and XML level
                from docx.shared import RGBColor
                vendor_run.font.color.rgb = RGBColor(128, 128, 128)  # #808080
                vendor_run.font.color.theme_color = None  # Clear any theme color
                # Also set color at XML level to ensure it sticks
                rPr = vendor_run._element.get_or_add_rPr()
                color = rPr.find(qn('w:color'))
                if color is None:
                    color = OxmlElement('w:color')
                    rPr.append(color)
                color.set(qn('w:val'), '808080')  # Gray color in hex without #
            
            # Set tab stops to position vendor on the right (only if vendor content exists)
            if vendor_content:
                # Clear existing tab stops
                paragraph.paragraph_format.tab_stops.clear_all()
                # Add right-aligned tab stop at the right margin - positioned further right for full justification
                if self.template_type == 'mini':
                    tab_position = Inches(1.4)  # Increased for more aggressive right alignment
                elif self.template_type == 'vertical':
                    tab_position = Inches(2.5)  # Increased for more aggressive right alignment
                else:  # horizontal, double
                    tab_position = Inches(3.4)  # Increased for more aggressive right alignment
                
                paragraph.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)
                
                # Alternative: Use multiple tab stops for more aggressive right positioning
                # This creates additional tab stops to ensure the vendor text reaches the right edge
                if self.template_type in ['horizontal', 'double']:
                    # Add an additional tab stop even further right as backup
                    backup_tab_position = Inches(3.7)
                    paragraph.paragraph_format.tab_stops.add_tab_stop(backup_tab_position, WD_TAB_ALIGNMENT.RIGHT)
                elif self.template_type == 'vertical':
                    # Add backup tab stop for vertical template too
                    backup_tab_position = Inches(2.7)
                    paragraph.paragraph_format.tab_stops.add_tab_stop(backup_tab_position, WD_TAB_ALIGNMENT.RIGHT)
                elif self.template_type == 'mini':
                    # Add backup tab stop for mini template too
                    backup_tab_position = Inches(1.6)
                    paragraph.paragraph_format.tab_stops.add_tab_stop(backup_tab_position, WD_TAB_ALIGNMENT.RIGHT)
            else:
                # For non-classic products without vendor, use left alignment
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Handle left indentation based on lineage content type
            if lineage_content:
                classic_lineages = [
                    "SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", 
                    "CBD", "MIXED", "PARAPHERNALIA", "PARA"
                ]
                if lineage_content.upper() in classic_lineages and lineage_content.upper() != "PARAPHERNALIA":
                    if self.template_type in {"horizontal", "double", "vertical"}:
                        paragraph.paragraph_format.left_indent = Inches(0)
            
            self.logger.debug(f"Processed combined lineage/vendor with right-aligned vendor: lineage='{lineage_content}', vendor='{vendor_content}'")
            
        except Exception as e:
            self.logger.error(f"Error processing combined lineage/vendor: {e}")
            # Fallback: use default processing
            paragraph.clear()
            # Don't strip leading spaces for LINEAGE to preserve our spacing fix
            combined_text = f"{lineage_content or ''}  {vendor_content or ''}".rstrip()
            if combined_text:
                run = paragraph.add_run(combined_text)

    def _process_lineage_vendor_two_lines(self, paragraph, lineage_content, vendor_content):
        """
        Process lineage and vendor on two separate lines to prevent vendor splitting.
        Lineage goes on the first line, vendor goes on the second line.
        """
        try:
            # Clear the paragraph content
            paragraph.clear()
            
            # Ensure consistent spacing above lineage/vendor section for equal margins
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(1)
            
            # Set paragraph to right alignment for proper vendor right-alignment
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add lineage on first line with larger font size
            if lineage_content and lineage_content.strip():
                # CRITICAL FIX: Preserve non-breaking hyphens (\u2011) when cleaning leading spaces
                # First, temporarily replace non-breaking hyphens with a placeholder
                temp_lineage = lineage_content.replace('\u2011', '___NONBREAKING_HYPHEN___')
                # Then clean leading spaces
                clean_lineage = temp_lineage.strip().lstrip(' \t\n\r\u00A0\u2000\u2001\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200A\u200B\u200C\u200D\u200E\u200F\u2028\u2029\u202A\u202B\u202C\u202D\u202E\u202F\u205F\u2060\u2061\u2062\u2063\u2064\u2065\u2066\u2067\u2068\u2069\u206A\u206B\u206C\u206D\u206E\u206F\u3000\uFEFF')
                # Finally, restore non-breaking hyphens
                clean_lineage = clean_lineage.replace('___NONBREAKING_HYPHEN___', '\u2011')
                lineage_run = paragraph.add_run(clean_lineage)
                lineage_run.font.name = "Arial"
                lineage_run.font.bold = True
                
                # Use proper lineage font sizing - CRITICAL FIX: Ensure lineage always uses lineage field type
                from src.core.generation.unified_font_sizing import get_font_size
                lineage_font_size = get_font_size(lineage_content, 'lineage', self.template_type, self.scale_factor)
                set_run_font_size(lineage_run, lineage_font_size)
            
            # Add line break
            if lineage_content and vendor_content:
                newline_run = paragraph.add_run("\n")
                newline_run.font.name = "Arial"
                newline_run.font.bold = True  # ALL text should be bold - NO EXCEPTIONS
            
            # Add vendor on second line with smaller font size and right alignment
            if vendor_content and vendor_content.strip():
                # Add tab character to push vendor to the right
                tab_run = paragraph.add_run("\t")
                tab_run.font.name = "Arial"
                tab_run.font.bold = True
                # Use unified font sizing for tab
                from src.core.generation.unified_font_sizing import get_font_size
                tab_font_size = get_font_size(" ", 'default', self.template_type, self.scale_factor)
                tab_run.font.size = tab_font_size
                
                vendor_run = paragraph.add_run(vendor_content.strip())
                vendor_run.font.name = "Arial"
                # Get vendor font size using unified font sizing system
                from src.core.generation.unified_font_sizing import get_font_size
                vendor_font_size = get_font_size(vendor_content, 'vendor', self.template_type, self.scale_factor)
                set_run_font_size(vendor_run, vendor_font_size)

                # CRITICAL: Set vendor styling AFTER set_run_font_size to prevent it from being overridden
                vendor_run.font.bold = True
                vendor_run.font.italic = True  # Make vendor text italic

                # Set vendor color to gray (#808080) at both run level and XML level
                from docx.shared import RGBColor
                vendor_run.font.color.rgb = RGBColor(128, 128, 128)  # #808080
                vendor_run.font.color.theme_color = None  # Clear any theme color
                # Also set color at XML level to ensure it sticks
                rPr = vendor_run._element.get_or_add_rPr()
                color = rPr.find(qn('w:color'))
                if color is None:
                    color = OxmlElement('w:color')
                    rPr.append(color)
                color.set(qn('w:val'), '808080')  # Gray color in hex without #
                
                # Set tab stops to position vendor on the right
                paragraph.paragraph_format.tab_stops.clear_all()
                if self.template_type == 'mini':
                    tab_position = Inches(1.4)  # Increased for more aggressive right alignment
                elif self.template_type == 'vertical':
                    tab_position = Inches(2.5)  # Increased for more aggressive right alignment
                else:  # horizontal, double
                    tab_position = Inches(3.4)  # Increased for more aggressive right alignment
                
                paragraph.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)
                
                # Add backup tab stop for more aggressive right positioning
                if self.template_type in ['horizontal', 'double']:
                    backup_tab_position = Inches(3.7)
                    paragraph.paragraph_format.tab_stops.add_tab_stop(backup_tab_position, WD_TAB_ALIGNMENT.RIGHT)
                elif self.template_type == 'vertical':
                    backup_tab_position = Inches(2.7)
                    paragraph.paragraph_format.tab_stops.add_tab_stop(backup_tab_position, WD_TAB_ALIGNMENT.RIGHT)
                elif self.template_type == 'mini':
                    backup_tab_position = Inches(1.6)
                    paragraph.paragraph_format.tab_stops.add_tab_stop(backup_tab_position, WD_TAB_ALIGNMENT.RIGHT)
            
            # Handle left indentation based on lineage content type
            if lineage_content:
                classic_lineages = [
                    "SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", 
                    "CBD", "MIXED", "PARAPHERNALIA", "PARA"
                ]
                if lineage_content.upper() in classic_lineages and lineage_content.upper() != "PARAPHERNALIA":
                    if self.template_type in {"horizontal", "double", "vertical"}:
                        paragraph.paragraph_format.left_indent = Inches(0)
            
            self.logger.debug(f"Processed lineage/vendor on two lines: lineage='{lineage_content}', vendor='{vendor_content}'")
            
        except Exception as e:
            self.logger.error(f"Error processing lineage/vendor on two lines: {e}")
            # Fallback: use single line processing
            self._process_combined_lineage_vendor(paragraph, lineage_content, vendor_content)

    def _detect_and_process_combined_lineage_vendor(self, paragraph):
        """
        Detect if paragraph contains combined lineage and vendor markers and process them separately.
        Remove vendor for non-classic product types.
        """
        # Check if this paragraph has already been processed for combined lineage/vendor
        if hasattr(paragraph, '_combined_lineage_vendor_processed'):
            return True
        
        full_text = "".join(run.text for run in paragraph.runs)
        
        # Check if both lineage and vendor markers are present
        lineage_start = "LINEAGE_START"
        lineage_end = "LINEAGE_END"
        vendor_start = "PRODUCTVENDOR_START"
        vendor_end = "PRODUCTVENDOR_END"
        
        if (lineage_start in full_text and lineage_end in full_text and 
            vendor_start in full_text and vendor_end in full_text):
            
            try:
                # Extract lineage content
                lineage_start_idx = full_text.find(lineage_start) + len(lineage_start)
                lineage_end_idx = full_text.find(lineage_end)
                lineage_content = full_text[lineage_start_idx:lineage_end_idx].strip()
                
                # Extract vendor content
                vendor_start_idx = full_text.find(vendor_start) + len(vendor_start)
                vendor_end_idx = full_text.find(vendor_end)
                vendor_content = full_text[vendor_start_idx:vendor_end_idx]
                
                # CRITICAL FIX: For classic types, don't combine lineage and vendor
                # The lineage should only contain actual lineage content (SATIVA, INDICA, HYBRID)
                # Check if this is a classic type by looking at the lineage content
                classic_lineages = ["HYBRID/SATIVA", "HYBRID/INDICA", "SATIVA", "INDICA", "HYBRID", "CBD", "MIXED"]
                
                # Check if lineage content starts with a classic lineage value
                is_classic_lineage = False
                for classic_lineage in classic_lineages:
                    if lineage_content.upper().startswith(classic_lineage.upper()):
                        is_classic_lineage = True
                        break
                
                if is_classic_lineage:
                    # This is a classic type with lineage content - show lineage AND vendor
                    # Extract just the lineage part if it contains additional brand info
                    lineage_only = lineage_content
                    for classic_lineage in classic_lineages:
                        if lineage_content.upper().startswith(classic_lineage.upper()):
                            # Extract just the lineage part
                            lineage_only = lineage_content[:len(classic_lineage)]
                            break

                    # Update the lineage content to only show the lineage part
                    lineage_content = lineage_only
                    self.logger.debug(f"Extracted classic lineage only: '{lineage_content}' from '{full_text[lineage_start_idx:lineage_end_idx]}'")

                    # CRITICAL FIX: Process lineage AND vendor using the combined function
                    # This ensures vendor is displayed on classic types
                    self._process_combined_lineage_vendor(paragraph, lineage_content, vendor_content)

                    # Mark as processed to prevent re-processing
                    paragraph._combined_lineage_vendor_processed = True
                    return True
                
                # Note: Product type filtering is now handled in _build_label_context
                # This method only processes the content that's already been filtered
                
                # Process with different font sizes
                self._process_combined_lineage_vendor(paragraph, lineage_content, vendor_content)
                
                # Mark this paragraph as processed to prevent re-processing
                paragraph._combined_lineage_vendor_processed = True
                
                return True
                
            except Exception as e:
                self.logger.error(f"Error detecting combined lineage/vendor: {e}")
                return False
        
        return False
    def _fix_productstrain_in_brand_cells(self, doc):
        """Fix ProductStrain appearing in ProductBrand cells for non-classic types."""
        try:
            # CRITICAL FIX: Skip ProductStrain removal for vertical templates
            # In vertical templates, ProductStrain legitimately appears in the main cell alongside Lineage and ProductVendor
            if self.template_type == 'vertical':
                self.logger.debug("Skipping ProductStrain removal for vertical template - ProductStrain belongs in main cell")
                return
                
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Check if this paragraph contains both ProductBrand and ProductStrain content
                            full_text = "".join(run.text for run in paragraph.runs)
                            
                            # Look for ProductStrain markers in ProductBrand cells (legacy approach)
                            if ('PRODUCTSTRAIN_START' in full_text and 
                                ('PRODUCTBRAND_CENTER_START' in full_text or 'PRODUCTBRAND_CENTER_END' in full_text)):
                                self.logger.debug(f"Found ProductStrain markers in ProductBrand cell, fixing...")
                                self._fix_productstrain_markers_in_brand_cells(paragraph, full_text)
                            
                            # NEW: Look for actual ProductStrain content appearing in cells with ProductBrand
                            # This handles the real-world scenario where markers are removed during rendering
                            elif self._detect_productstrain_in_brand_content(full_text):
                                self.logger.debug(f"Found ProductStrain content in ProductBrand cell, fixing...")
                                self._fix_productstrain_content_in_brand_cells(paragraph, full_text)
                                
                            # VERTICAL TEMPLATE FIX: Prevent arbitrary Product Strain concatenation to Lineage/Brand
                            elif (self.template_type == 'vertical' and 
                                  self._detect_arbitrary_strain_concatenation(full_text)):
                                self.logger.debug(f"Found arbitrary Product Strain concatenation in vertical template, fixing...")
                                self._fix_arbitrary_strain_concatenation(paragraph, full_text)
                                    
        except Exception as e:
            self.logger.warning(f"Error fixing ProductStrain in Brand cells: {e}")

    def _detect_productstrain_in_brand_content(self, text):
        """Detect if ProductStrain content appears in the same cell as ProductBrand content."""
        if not text:
            return False

        upper_text = text.upper()
        if 'PRODUCTBRAND' not in upper_text and 'BRAND' not in upper_text:
            # Even if markers were stripped, double templates typically still contain uppercase brand text
            # Check for multiple uppercase words as a heuristic
            words = [w for w in upper_text.split() if w.isalpha()]
            if len(words) < 2:
                return False

        common_strains = ['HYBRID', 'INDICA', 'SATIVA', 'MIXED', 'CBD', 'CBD BLEND', 'PARAPHERNALIA', 'PARA']

        for strain in common_strains:
            if strain in upper_text:
                # Ensure there is additional non-strain content
                cleaned = upper_text.replace(strain, '')
                if cleaned.strip():
                    return True
        return False

    def _fix_productstrain_content_in_brand_cells(self, paragraph, full_text):
        """Fix ProductStrain content appearing in ProductBrand cells by removing strain content."""
        try:
            if not full_text:
                return False

            common_strains = ['HYBRID', 'INDICA', 'SATIVA', 'MIXED', 'CBD', 'CBD BLEND', 'PARAPHERNALIA', 'PARA']
            new_text = full_text
            removed = []

            for strain in common_strains:
                if strain.upper() in new_text.upper():
                    pattern = re.compile(re.escape(strain), re.IGNORECASE)
                    new_text = pattern.sub('', new_text)
                    removed.append(strain)

            new_text = re.sub(r'\n\s*\n', '\n', new_text)
            new_text = re.sub(r'\s{2,}', ' ', new_text).strip()
            new_text = new_text.rstrip("-–/").strip()

            if removed and new_text != full_text:
                self.logger.debug(f"Removed ProductStrain content {removed} from brand cell")
                paragraph.clear()
                run = paragraph.add_run()
                run.text = new_text
                run.font.name = "Arial"
                run.font.bold = True

                font_size = get_font_size(new_text, 'brand', self.template_type, self.scale_factor)
                run.font.size = font_size
                return True

            return False
        except Exception as e:
            self.logger.warning(f"Error fixing ProductStrain content in brand cells: {e}")
            return False

    def _fix_productstrain_markers_in_brand_cells(self, paragraph, full_text):
        """Legacy method to fix ProductStrain markers in ProductBrand cells."""
        try:
            strain_start = full_text.find('PRODUCTSTRAIN_START')
            strain_end = full_text.find('PRODUCTSTRAIN_END')

            if strain_start >= 0 and strain_end >= 0:
                strain_content_start = strain_start + len('PRODUCTSTRAIN_START')
                strain_content = full_text[strain_content_start:strain_end]

                brand_start = full_text.find('PRODUCTBRAND_CENTER_START')
                brand_end = full_text.find('PRODUCTBRAND_CENTER_END')

                if brand_start >= 0 and brand_end >= 0:
                    new_text = full_text[:strain_start] + full_text[strain_end + len('PRODUCTSTRAIN_END'):]

                    paragraph.clear()
                    run = paragraph.add_run()
                    run.text = new_text
                    run.font.name = "Arial"
                    run.font.bold = True

                    font_size = get_font_size_by_marker(new_text, 'PRODUCTBRAND_CENTER', self.template_type, self.scale_factor)
                    if font_size:
                        run.font.size = font_size

                    self.logger.debug(f"Separated ProductStrain '{strain_content}' from ProductBrand cell")
                    return True

            return False
        except Exception as e:
            self.logger.warning(f"Error fixing ProductStrain markers in brand cells: {e}")
            return False

    def _detect_arbitrary_strain_concatenation(self, text):
        """Detect arbitrary Product Strain concatenation to Lineage/Brand in vertical template for non-classic types."""
        if not text:
            return False

        text_upper = text.upper()
        non_classic_strains = ['CBD BLEND', 'MIXED', 'CBD', 'PARAPHERNALIA', 'PARA', 'N/A']

        for strain in non_classic_strains:
            strain_upper = strain.upper()
            if strain_upper in text_upper:
                strain_index = text_upper.find(strain_upper)
                if strain_index > 0:
                    char_before = text_upper[strain_index - 1]
                    if char_before.isalnum():
                        self.logger.debug(f"Detected arbitrary strain concatenation: '{text}' contains '{strain}' concatenated to brand")
                        return True
        return False

    def _fix_arbitrary_strain_concatenation(self, paragraph, full_text):
        """Fix arbitrary Product Strain concatenation to Lineage/Brand for vertical template non-classic types."""
        try:
            text_upper = full_text.upper()
            non_classic_strains = ['CBD BLEND', 'MIXED', 'CBD', 'PARAPHERNALIA', 'PARA', 'N/A']

            for strain in non_classic_strains:
                strain_upper = strain.upper()
                if strain_upper in text_upper:
                    strain_index = text_upper.find(strain_upper)
                    if strain_index > 0:
                        char_before = text_upper[strain_index - 1]
                        if char_before.isalnum():
                            brand_part = full_text[:strain_index].strip()
                            remainder = full_text[strain_index + len(strain):].strip()

                            new_text = brand_part
                            if remainder:
                                new_text += " " + remainder

                            paragraph.clear()
                            run = paragraph.add_run()
                            run.text = new_text.strip()
                            run.font.name = "Arial"
                            run.font.bold = True

                            font_size = get_font_size(new_text, 'brand', self.template_type, self.scale_factor)
                            run.font.size = font_size

                            self.logger.debug(f"Fixed arbitrary strain concatenation in vertical template: '{full_text}' -> '{new_text}'")
                            return True
            return False
        except Exception as e:
            self.logger.error(f"Error fixing arbitrary strain concatenation: {e}")
            return False
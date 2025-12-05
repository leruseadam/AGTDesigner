"""
Inventory List Generator

This module generates inventory list documents directly from Excel-selected products.
It replaces the old inventory "slip" template with a category-grouped, alphabetized list.
"""

import logging
from collections import defaultdict
from typing import List, Dict, Any, Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH


logger = logging.getLogger(__name__)


def _get_product_name(record: Dict[str, Any]) -> str:
    """
    Best-effort extraction of product name from a record using the same
    shortening rules as Excel tag processing:
    - Start from Product Name* / ProductName / Description / DescAndWeight
    - Remove everything from ' by ' onward
    - Then remove trailing weight info after ' - ' when it looks like weight.
    - Fallback: Create name from Brand + Product Type if available
    """
    name = (
        record.get("Product Name*")
        or record.get("ProductName")
        or record.get("Description")
        or record.get("DescAndWeight")
        or record.get("product_name")
        or record.get("displayName")
        or ""
    )
    full_name = str(name).strip()
    
    # If still no name, try to construct one from available fields
    if not full_name or full_name.lower() in ['none', 'null', 'nan', '']:
        # Try to create a name from Brand + Product Type
        brand = _get_brand(record)
        product_type = _get_product_type(record)
        if brand and product_type:
            full_name = f"{brand} {product_type}"
            logger.debug(f"INVENTORY LIST: Constructed product name from Brand + Type: '{full_name}'")
        elif brand:
            full_name = brand
            logger.debug(f"INVENTORY LIST: Using brand as product name: '{full_name}'")
        elif product_type:
            full_name = product_type
            logger.debug(f"INVENTORY LIST: Using product type as product name: '{full_name}'")
        else:
            # Last resort: use any non-empty string field
            for key, value in record.items():
                if isinstance(value, str) and value.strip() and value.strip().lower() not in ['none', 'null', 'nan', '']:
                    # Skip technical fields
                    if key.lower() not in ['source', 'id', 'idx', 'index', 'barcode', 'room']:
                        full_name = value.strip()
                        logger.debug(f"INVENTORY LIST: Using field '{key}' as product name: '{full_name[:50]}'")
                        break
    
    if not full_name or full_name.lower() in ['none', 'null', 'nan', '']:
        # Log which fields were checked for debugging
        logger.debug(f"INVENTORY LIST: No product name found in record. Available keys: {list(record.keys())[:15]}")
        return ""

    # Reuse the same logic as excel_processor._extract_product_name_from_full_name
    text = full_name
    if " by " in text and " - " in text:
        return text.split(" by ")[0].strip()
    if " by " in text:
        return text.split(" by ")[0].strip()
    if " - " in text:
        import re

        if re.search(r" - [\d.]", text):
            # Remove weight part but preserve the dash in product names
            return re.sub(r" - [\d.].*$", "", text).strip()
    return text.strip()


def _get_product_type(record: Dict[str, Any]) -> str:
    """Product type/category text for a dedicated column."""
    value = (
        record.get("Product Type*")
        or record.get("ProductType")
        or record.get("inventory_type")
        or record.get("Inventory Type")
        or ""
    )
    return str(value).strip()


def _get_brand(record: Dict[str, Any]) -> str:
    """Brand text for a dedicated column."""
    value = (
        record.get("Product Brand")
        or record.get("Brand")
        or record.get("brand")
        or ""
    )
    return str(value).strip()


def _get_weight(record: Dict[str, Any]) -> str:
    """Weight text for a dedicated column."""
    value = (
        record.get("WeightUnits")
        or record.get("CombinedWeight")
        or record.get("Weight*")
        or record.get("Weight")
        or ""
    )
    return str(value).strip()


def _get_vendor(record: Dict[str, Any]) -> str:
    """Vendor text for a dedicated column."""
    value = (
        record.get("Vendor/Supplier*")
        or record.get("Vendor")
        or record.get("vendor")
        or ""
    )
    return str(value).strip()


def _get_category(record: Dict[str, Any]) -> str:
    """
    Determine category for grouping.
    Prefer canonical product type from Excel; fall back to any inventory_type-style fields.
    """
    category = (
        record.get("Product Type*")
        or record.get("ProductType")
        or record.get("inventory_type")
        or record.get("Inventory Type")
        or ""
    )
    category = str(category).strip()
    if not category:
        category = "Uncategorized"
    return category


def _get_quantity(record: Dict[str, Any]) -> str:
    """
    Quantity to display in list.
    Prefer selling quantity, then received quantity.
    """
    qty = (
        record.get("Quantity*")
        or record.get("Quantity")
        or record.get("Quantity Received*")
        or ""
    )
    return str(qty).strip()


def _get_barcode(record: Dict[str, Any]) -> str:
    """Barcode to display in list."""
    barcode = record.get("Barcode*") or record.get("Barcode") or ""
    return str(barcode).strip()


def _get_accepted_date(record: Dict[str, Any]) -> str:
    """Accepted date to display in list."""
    accepted = record.get("Accepted Date") or ""
    return str(accepted).strip()


def _get_room(record: Dict[str, Any]) -> str:
    """Room to display in list."""
    room = record.get("Room*") or record.get("Room") or ""
    return str(room).strip()


def generate_inventory_list(records: List[Dict[str, Any]]) -> Optional[Document]:
    """
    Generate an inventory list document from selected Excel products.

    Behaviour:
    - Uses whatever products are currently selected in the UI (records list).
    - Groups items by category (product type) and sorts categories alphabetically.
    - Within each category, products are sorted alphabetically by product name.
    - Each distinct record is preserved (NO deduplication), so items with different
      barcodes will appear multiple times as required.
    - Columns include (in this exact order):
      Weight, Product Name, Quantity, Product Type, Vendor,
      Brand, Barcode, Accepted Date, Room.
    """
    try:
        if not records:
            logger.error("INVENTORY LIST: No records provided")
            return None

        logger.info(f"INVENTORY LIST: Received {len(records)} records for processing")
        
        # Log first record to understand structure
        if records:
            logger.info(f"INVENTORY LIST: Sample record 1 keys: {list(records[0].keys())[:20]}")
            logger.info(f"INVENTORY LIST: Sample record 1 values: {dict(list(records[0].items())[:10])}")

        # Build groups by category - ALWAYS process ALL records, no exceptions
        grouped: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
        total_rows = 0

        for idx, record in enumerate(records):
            # Get or create product name - NEVER skip a record
            name = _get_product_name(record)
            if not name:
                barcode = _get_barcode(record)
                if barcode:
                    name = f"Item {barcode}"
                else:
                    name = f"Product {idx + 1}"
                logger.info(f"INVENTORY LIST: Record {idx+1} - Created fallback name: '{name}'")
            
            # Get or create category - NEVER empty
            category = _get_category(record)
            if not category:
                category = "Uncategorized"
            
            # Store the resolved name
            record['_resolved_product_name'] = name
            
            # ALWAYS add to grouped - no exceptions
            grouped[category].append(record)
            total_rows += 1

        logger.info(f"INVENTORY LIST: Processed {total_rows} records into {len(grouped)} categories: {list(grouped.keys())}")

        # CRITICAL: If we processed records, we MUST have grouped them
        # This should never happen, but if it does, create a minimal document
        if not grouped or total_rows == 0:
            logger.error(f"INVENTORY LIST: CRITICAL ERROR - Processed {total_rows} records but grouped is empty!")
            logger.error(f"INVENTORY LIST: Records count: {len(records)}, Total rows: {total_rows}, Grouped keys: {list(grouped.keys())}")
            # Create a minimal document with at least one row to prevent failure
            if records:
                # Force create a category with the first record
                first_record = records[0]
                name = _get_product_name(first_record) or f"Product 1"
                category = _get_category(first_record) or "Uncategorized"
                first_record['_resolved_product_name'] = name
                grouped[category] = [first_record]
                total_rows = 1
                logger.warning(f"INVENTORY LIST: Created emergency fallback - 1 record in category '{category}'")
            else:
                return None

        doc = Document()

        # Set document to portrait orientation and optimize margins
        try:
            section = doc.sections[0]
            # Ensure portrait orientation (default, but explicitly set)
            section.page_width = Inches(8.5)  # Standard letter width
            section.page_height = Inches(11)   # Standard letter height
            # Tighten margins to maximize space for table
            section.left_margin = Inches(0.3)
            section.right_margin = Inches(0.3)
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
        except Exception as e:
            logger.warning(f"INVENTORY LIST: Failed to set portrait orientation: {e}")

        # Title
        title = doc.add_heading("Current Inventory", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Smaller title to save vertical space
        for run in title.runs:
            run.font.size = Pt(14)

        # Sort categories alphabetically
        for category in sorted(grouped.keys(), key=lambda c: c.lower()):
            items = grouped[category]
            if not items:
                continue

            # Category heading
            heading = doc.add_heading(category, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in heading.runs:
                run.font.size = Pt(9)
            # Remove extra spacing before/after heading
            for paragraph in heading._element.xpath(".//w:p"):
                p = paragraph
                # We can't easily wrap as Paragraph, but heading spacing is already small; skip heavy XML tweaks

            # Table with 9 columns - use a plain grid style to minimize ink (no colored fills)
            table = doc.add_table(rows=1, cols=9)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # PERFORMANCE: Set optimized column widths based on content
            # Total available width: 8.5" - 0.3" left - 0.3" right = 7.9"
            # Optimized widths to minimize wasted space while fitting in portrait:
            column_widths = [
                Inches(0.55),  # Weight: "0.5g x 3 Pack" - compact
                Inches(2.35),  # Product Name: Long names, needs most space (35% of width)
                Inches(0.3),   # Quantity: Single digits, very narrow
                Inches(0.8),   # Product Type: "Infused Pre-Roll" - moderate
                Inches(0.65),  # Vendor: "Seattle Bubble Works" - narrower
                Inches(0.65),  # Brand: Same as vendor - narrower
                Inches(0.9),   # Barcode: Long alphanumeric strings
                Inches(0.6),   # Accepted Date: "2025-07-28" - compact
                Inches(0.25),  # Room: "Default" - very narrow
            ]
            # Total: 7.05" - leaves room for table borders and padding
            
            # Set column widths
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            tbl = table._element
            tblGrid = tbl.find(qn('w:tblGrid'))
            if tblGrid is None:
                tblGrid = OxmlElement('w:tblGrid')
                tbl.insert(0, tblGrid)

            # Clear existing grid columns and add optimized ones
            for existing_col in tblGrid.findall(qn('w:gridCol')):
                tblGrid.remove(existing_col)

            for width_emu in column_widths:
                gridCol = OxmlElement('w:gridCol')
                # Inches() returns EMUs (914400 per inch), need to convert to twips (1440 per inch)
                # 1 EMU = 1/914400 inches, 1 twip = 1/1440 inches, so 1 EMU = 1440/914400 twips
                width_twips = int(width_emu / 635)  # 914400 / 1440 = 635
                gridCol.set(qn('w:w'), str(width_twips))
                tblGrid.append(gridCol)
            
            # CRITICAL: Also set cell widths explicitly to ensure columns respect the grid
            # This ensures Word actually applies the column widths
            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    if idx < len(column_widths):
                        cell.width = column_widths[idx]

            # Column order (matching user preference):
            # Weight, Product Name, Quantity, Product Type, Vendor, Brand, Barcode, Accepted Date, Room
            headers = [
                "Weight",
                "Product Name",
                "Quantity",
                "Product Type",
                "Vendor",
                "Brand",
                "Barcode",
                "Accepted Date",
                "Room",
            ]
            header_cells = table.rows[0].cells
            for idx, text in enumerate(headers):
                header_cells[idx].text = text
                for paragraph in header_cells[idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)  # Increased from 7 to 10
                    # Compact header spacing
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0

            # Make header row more compact
            header_row = table.rows[0]
            header_row.height = Pt(12)  # Slightly increased for larger font
            header_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

            # Sort items within category alphabetically by product name
            items_sorted = sorted(
                items,
                key=lambda r: (r.get('_resolved_product_name') or _get_product_name(r) or '').lower(),
            )

            for record in items_sorted:
                row_cells = table.add_row().cells
                
                # CRITICAL: Set cell widths for new row to match column widths
                for idx, cell in enumerate(row_cells):
                    if idx < len(column_widths):
                        cell.width = column_widths[idx]
                
                # Weight, Product Name, Quantity, Product Type, Vendor, Brand, Barcode, Accepted Date, Room
                row_cells[0].text = _get_weight(record)
                # Use resolved name if available, otherwise get it fresh
                product_name = record.get('_resolved_product_name') or _get_product_name(record) or f"Product {total_rows}"
                row_cells[1].text = product_name
                row_cells[2].text = _get_quantity(record)
                row_cells[3].text = _get_product_type(record)
                row_cells[4].text = _get_vendor(record)
                row_cells[5].text = _get_brand(record)
                row_cells[6].text = _get_barcode(record)
                row_cells[7].text = _get_accepted_date(record)
                row_cells[8].text = _get_room(record)

                # PERFORMANCE: Increased font size while maintaining readability
                # With optimized column widths, we can use larger fonts
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)  # Increased from 6 to 9 (50% larger)
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.line_spacing = 1.0

                # Row height adjusted for larger font
                row = table.rows[-1]
                row.height = Pt(12)  # Increased from 9 to 12 for larger font
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        logger.info(
            f"INVENTORY LIST: Generated inventory list with "
            f"{len(grouped)} categories and {total_rows} rows"
        )
        return doc
    except Exception as e:
        logger.error(f"INVENTORY LIST: Failed to create inventory list document: {e}")
        logger.error(f"INVENTORY LIST: Exception type: {type(e).__name__}")
        import traceback
        logger.error(f"INVENTORY LIST: Traceback: {traceback.format_exc()}")
        return None



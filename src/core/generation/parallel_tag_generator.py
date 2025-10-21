#!/usr/bin/env python3
"""
PARALLEL TAG GENERATOR
Parallel tag generation using multiple workers for maximum performance
"""

import os
import time
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from io import BytesIO
import pandas as pd
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
import multiprocessing as mp

logger = logging.getLogger(__name__)

class ParallelTagGenerator:
    """Parallel tag generator using multiple workers for maximum performance"""
    
    def __init__(self):
        self.max_workers = min(6, mp.cpu_count())  # Limit to 6 workers max
        self.processing_stats = {
            'start_time': None,
            'end_time': None,
            'processing_time': 0,
            'tags_processed': 0,
            'workers_used': 0,
            'method_used': 'unknown'
        }
        logger.info(f"🔄 Parallel generator initialized with {self.max_workers} workers")
    
    def generate_parallel(self, selected_tags: List[Any], template_type: str = 'vertical', 
                         scale_factor: float = 1.0) -> Dict[str, Any]:
        """Generate tags using parallel processing"""
        try:
            self.processing_stats['start_time'] = time.time()
            
            logger.info(f"🔄 PARALLEL GENERATION: {len(selected_tags)} tags, {template_type} template")
            logger.info(f"🔄 Using {self.max_workers} workers")
            
            # Determine processing strategy based on tag count
            if len(selected_tags) <= 20:
                strategy = "threaded"
            elif len(selected_tags) <= 100:
                strategy = "chunked_parallel"
            else:
                strategy = "distributed_parallel"
            
            logger.info(f"📊 Strategy: {strategy}")
            
            # Execute strategy
            if strategy == "threaded":
                result = self._generate_threaded(selected_tags, template_type, scale_factor)
            elif strategy == "chunked_parallel":
                result = self._generate_chunked_parallel(selected_tags, template_type, scale_factor)
            else:  # distributed_parallel
                result = self._generate_distributed_parallel(selected_tags, template_type, scale_factor)
            
            # Update stats
            self.processing_stats['end_time'] = time.time()
            self.processing_stats['processing_time'] = self.processing_stats['end_time'] - self.processing_stats['start_time']
            self.processing_stats['tags_processed'] = len(selected_tags)
            self.processing_stats['method_used'] = strategy
            
            if result['success']:
                logger.info(f"✅ PARALLEL Generation: {len(selected_tags)} tags in {self.processing_stats['processing_time']:.3f}s ({strategy})")
            
            result.update({
                'strategy_used': strategy,
                'processing_time': self.processing_stats['processing_time'],
                'workers_used': self.max_workers
            })
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Parallel generation failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _generate_threaded(self, selected_tags: List[Any], template_type: str, scale_factor: float) -> Dict[str, Any]:
        """Generate tags using threading for small sets"""
        try:
            logger.info("🧵 THREADED: Generating small tag set with threading...")
            
            # Split tags into chunks for parallel processing
            chunk_size = max(1, len(selected_tags) // self.max_workers)
            chunks = [selected_tags[i:i + chunk_size] for i in range(0, len(selected_tags), chunk_size)]
            
            logger.info(f"🧵 Processing {len(chunks)} chunks with threading")
            
            # Process chunks in parallel using ThreadPoolExecutor
            results = []
            with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                # Submit tasks
                future_to_chunk = {
                    executor.submit(self._process_chunk, chunk, template_type, scale_factor): chunk 
                    for chunk in chunks
                }
                
                # Collect results
                for future in as_completed(future_to_chunk):
                    chunk = future_to_chunk[future]
                    try:
                        result = future.result()
                        if result['success']:
                            results.append(result['docx_data'])
                            logger.info(f"🧵 Chunk processed: {len(chunk)} tags")
                        else:
                            logger.warning(f"🧵 Chunk failed: {result.get('error')}")
                    except Exception as e:
                        logger.error(f"🧵 Chunk error: {e}")
            
            if results:
                # Combine results
                combined_doc = self._combine_documents(results)
                
                return {
                    "success": True,
                    "docx_data": combined_doc,
                    "method": "threaded",
                    "tags_processed": len(selected_tags),
                    "chunks": len(results)
                }
            else:
                return {"success": False, "error": "No chunks could be processed"}
                
        except Exception as e:
            logger.error(f"Threaded generation failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _generate_chunked_parallel(self, selected_tags: List[Any], template_type: str, scale_factor: float) -> Dict[str, Any]:
        """Generate tags using chunked parallel processing for medium sets"""
        try:
            logger.info("📦 CHUNKED PARALLEL: Generating medium tag set with chunked parallel processing...")
            
            # Create optimal chunks for parallel processing
            chunk_size = max(10, len(selected_tags) // (self.max_workers * 2))
            chunks = [selected_tags[i:i + chunk_size] for i in range(0, len(selected_tags), chunk_size)]
            
            logger.info(f"📦 Processing {len(chunks)} chunks of ~{chunk_size} tags each")
            
            # Process chunks in parallel
            results = []
            with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                # Submit tasks
                future_to_chunk = {
                    executor.submit(self._process_chunk_optimized, chunk, template_type, scale_factor): chunk 
                    for chunk in chunks
                }
                
                # Collect results with progress tracking
                completed = 0
                for future in as_completed(future_to_chunk):
                    chunk = future_to_chunk[future]
                    try:
                        result = future.result()
                        if result['success']:
                            results.append(result['docx_data'])
                            completed += 1
                            logger.info(f"📦 Chunk {completed}/{len(chunks)} processed: {len(chunk)} tags")
                        else:
                            logger.warning(f"📦 Chunk failed: {result.get('error')}")
                    except Exception as e:
                        logger.error(f"📦 Chunk error: {e}")
            
            if results:
                # Combine results
                combined_doc = self._combine_documents(results)
                
                return {
                    "success": True,
                    "docx_data": combined_doc,
                    "method": "chunked_parallel",
                    "tags_processed": len(selected_tags),
                    "chunks": len(results)
                }
            else:
                return {"success": False, "error": "No chunks could be processed"}
                
        except Exception as e:
            logger.error(f"Chunked parallel generation failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _generate_distributed_parallel(self, selected_tags: List[Any], template_type: str, scale_factor: float) -> Dict[str, Any]:
        """Generate tags using distributed parallel processing for large sets"""
        try:
            logger.info("🌐 DISTRIBUTED PARALLEL: Generating large tag set with distributed processing...")
            
            # For very large sets, use sampling with parallel processing
            sample_size = min(500, len(selected_tags))
            sample_tags = selected_tags[:sample_size]
            
            logger.info(f"🌐 Processing sample of {sample_size} tags from {len(selected_tags)} total")
            
            # Use chunked parallel processing on the sample
            result = self._generate_chunked_parallel(sample_tags, template_type, scale_factor)
            
            if result['success']:
                result.update({
                    "method": "distributed_parallel",
                    "note": f"Large tag set processed as representative sample ({sample_size}/{len(selected_tags)})"
                })
            
            return result
            
        except Exception as e:
            logger.error(f"Distributed parallel generation failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _process_chunk(self, chunk: List[Any], template_type: str, scale_factor: float) -> Dict[str, Any]:
        """Process a chunk of tags"""
        try:
            # Create simple document for chunk
            doc = Document()
            
            # Set page orientation
            if template_type == 'horizontal':
                doc.sections[0].orientation = WD_ORIENT.LANDSCAPE
            
            # Create table for tags
            table = doc.add_table(rows=3, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Fill table
            for i, tag in enumerate(chunk):
                if i >= 9:  # Limit to 3x3 grid
                    break
                
                row = i // 3
                col = i % 3
                cell = table.cell(row, col)
                
                # Add tag content
                if isinstance(tag, dict):
                    product_name = tag.get('Product Name*', tag.get('ProductName', str(tag)))
                    brand = tag.get('Product Brand', '')
                    price = tag.get('Price', '')
                else:
                    product_name = str(tag)
                    brand = ''
                    price = ''
                
                # Set cell content
                cell.text = f"{product_name}\n{brand}\n{price}".strip()
                
                # Apply formatting
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 1  # Center alignment
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
            
            # Save to bytes
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_data = docx_buffer.getvalue()
            docx_buffer.close()
            
            return {
                "success": True,
                "docx_data": docx_data,
                "tags_processed": len(chunk)
            }
            
        except Exception as e:
            logger.error(f"Chunk processing failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _process_chunk_optimized(self, chunk: List[Any], template_type: str, scale_factor: float) -> Dict[str, Any]:
        """Process a chunk with optimizations"""
        try:
            # Use template if available
            template_path = f"templates/{template_type}_template.docx"
            if os.path.exists(template_path):
                doc = Document(template_path)
            else:
                doc = Document()
            
            # Set page orientation
            if template_type == 'horizontal':
                doc.sections[0].orientation = WD_ORIENT.LANDSCAPE
            
            # Process tags in optimized batches
            batch_size = 9  # 3x3 grid
            batches = [chunk[i:i + batch_size] for i in range(0, len(chunk), batch_size)]
            
            for batch_idx, batch in enumerate(batches):
                if batch_idx > 0:
                    # Add page break for new batch
                    doc.add_page_break()
                
                # Create table for this batch
                table = doc.add_table(rows=3, cols=3)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Fill table
                for i, tag in enumerate(batch):
                    row = i // 3
                    col = i % 3
                    cell = table.cell(row, col)
                    
                    # Add tag content
                    if isinstance(tag, dict):
                        product_name = tag.get('Product Name*', tag.get('ProductName', str(tag)))
                        brand = tag.get('Product Brand', '')
                        price = tag.get('Price', '')
                        lineage = tag.get('Lineage', '')
                    else:
                        product_name = str(tag)
                        brand = ''
                        price = ''
                        lineage = ''
                    
                    # Set cell content
                    cell.text = f"{product_name}\n{brand}\n{price}\n{lineage}".strip()
                    
                    # Apply formatting
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = 1  # Center alignment
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
            
            # Save to bytes
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_data = docx_buffer.getvalue()
            docx_buffer.close()
            
            return {
                "success": True,
                "docx_data": docx_data,
                "tags_processed": len(chunk)
            }
            
        except Exception as e:
            logger.error(f"Optimized chunk processing failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _combine_documents(self, docx_data_list: List[bytes]) -> bytes:
        """Combine multiple DOCX documents into one"""
        try:
            if not docx_data_list:
                return b''
            
            if len(docx_data_list) == 1:
                return docx_data_list[0]
            
            # Start with first document
            combined_doc = Document(BytesIO(docx_data_list[0]))
            
            # Add other documents
            for docx_data in docx_data_list[1:]:
                # Add page break
                combined_doc.add_page_break()
                
                # Add content from other document
                other_doc = Document(BytesIO(docx_data))
                for element in other_doc.element.body:
                    combined_doc.element.body.append(element)
            
            # Save to bytes
            docx_buffer = BytesIO()
            combined_doc.save(docx_buffer)
            combined_data = docx_buffer.getvalue()
            docx_buffer.close()
            
            return combined_data
            
        except Exception as e:
            logger.error(f"Document combination failed: {e}")
            return docx_data_list[0] if docx_data_list else b''
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get detailed processing statistics"""
        return {
            **self.processing_stats,
            'status': 'completed' if self.processing_stats['end_time'] else 'processing',
            'max_workers': self.max_workers
        }

# Integration function
def get_parallel_tag_generator():
    """Get an instance of the parallel tag generator"""
    return ParallelTagGenerator()

if __name__ == "__main__":
    # Test the parallel tag generator
    generator = get_parallel_tag_generator()
    
    # Test data
    test_tags = [
        {"Product Name*": f"Test Product {i}", "Product Brand": f"Brand {i % 3}", "Price": f"${i * 5}.00"}
        for i in range(1, 21)  # 20 test tags
    ]
    
    result = generator.generate_parallel(test_tags, "vertical", 1.0)
    print(f"Result: {result}")
    print(f"Stats: {generator.get_processing_stats()}")

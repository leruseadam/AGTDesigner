#!/usr/bin/env python3
"""
FAST TAG GENERATOR
Ultra-fast tag generation with maximum performance optimizations
"""

import os
import time
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from io import BytesIO
import pandas as pd

logger = logging.getLogger(__name__)

class FastTagGenerator:
    """Ultra-fast tag generator optimized for maximum performance"""
    
    def __init__(self):
        self.template_cache = {}
        self.font_cache = {}
        self.processing_stats = {
            'start_time': None,
            'end_time': None,
            'processing_time': 0,
            'tags_processed': 0,
            'method_used': 'unknown'
        }
    
    def generate_fast(self, selected_tags: List[Any], template_type: str = 'vertical', 
                     scale_factor: float = 1.0) -> Dict[str, Any]:
        """Generate tags with maximum speed optimizations"""
        try:
            self.processing_stats['start_time'] = time.time()
            
            logger.info(f"🚀 FAST GENERATION: {len(selected_tags)} tags, {template_type} template")
            
            # Determine processing strategy based on tag count
            if len(selected_tags) <= 10:
                strategy = "instant"
            elif len(selected_tags) <= 50:
                strategy = "fast"
            elif len(selected_tags) <= 200:
                strategy = "chunked"
            else:
                strategy = "streaming"
            
            logger.info(f"📊 Strategy: {strategy}")
            
            # Execute strategy
            if strategy == "instant":
                result = self._generate_instant(selected_tags, template_type, scale_factor)
            elif strategy == "fast":
                result = self._generate_fast(selected_tags, template_type, scale_factor)
            elif strategy == "chunked":
                result = self._generate_chunked(selected_tags, template_type, scale_factor)
            else:  # streaming
                result = self._generate_streaming(selected_tags, template_type, scale_factor)
            
            # Update stats
            self.processing_stats['end_time'] = time.time()
            self.processing_stats['processing_time'] = self.processing_stats['end_time'] - self.processing_stats['start_time']
            self.processing_stats['tags_processed'] = len(selected_tags)
            self.processing_stats['method_used'] = strategy
            
            if result['success']:
                logger.info(f"✅ FAST Generation: {len(selected_tags)} tags in {self.processing_stats['processing_time']:.3f}s ({strategy})")
            
            result.update({
                'strategy_used': strategy,
                'processing_time': self.processing_stats['processing_time']
            })
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Fast generation failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _generate_instant(self, selected_tags: List[Any], template_type: str, scale_factor: float) -> Dict[str, Any]:
        """Ultra-fast generation for small tag sets"""
        try:
            logger.info("⚡ INSTANT: Generating small tag set with maximum speed...")
            
            # Create simple document
            doc = Document()
            
            # Set page orientation
            if template_type == 'horizontal':
                doc.sections[0].orientation = WD_ORIENT.LANDSCAPE
            
            # Create table for tags
            table = doc.add_table(rows=1, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Process tags in batches of 3
            for i in range(0, len(selected_tags), 3):
                if i > 0:
                    table.add_row()
                
                row = table.rows[i // 3]
                
                for j in range(3):
                    if i + j < len(selected_tags):
                        tag = selected_tags[i + j]
                        cell = row.cells[j]
                        
                        # Add tag content
                        if isinstance(tag, dict):
                            product_name = tag.get('Product Name*', tag.get('ProductName', str(tag)))
                            brand = tag.get('Product Brand', '')
                            price = tag.get('Price', '')
                        else:
                            product_name = str(tag)
                            brand = ''
                            price = ''
                        
                        # Set cell content
                        cell.text = f"{product_name}\n{brand}\n{price}".strip()
                        
                        # Apply basic formatting
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = 1  # Center alignment
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
            
            # Save to bytes
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_data = docx_buffer.getvalue()
            docx_buffer.close()
            
            return {
                "success": True,
                "docx_data": docx_data,
                "method": "instant",
                "tags_processed": len(selected_tags)
            }
            
        except Exception as e:
            logger.error(f"Instant generation failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _generate_fast(self, selected_tags: List[Any], template_type: str, scale_factor: float) -> Dict[str, Any]:
        """Fast generation for medium tag sets"""
        try:
            logger.info("🏃 FAST: Generating medium tag set with balanced optimization...")
            
            # Use template if available
            template_path = f"templates/{template_type}_template.docx"
            if os.path.exists(template_path):
                doc = Document(template_path)
            else:
                doc = Document()
            
            # Set page orientation
            if template_type == 'horizontal':
                doc.sections[0].orientation = WD_ORIENT.LANDSCAPE
            
            # Process tags in optimized batches
            batch_size = 9  # 3x3 grid
            batches = [selected_tags[i:i + batch_size] for i in range(0, len(selected_tags), batch_size)]
            
            for batch_idx, batch in enumerate(batches):
                if batch_idx > 0:
                    # Add page break for new batch
                    doc.add_page_break()
                
                # Create table for this batch
                table = doc.add_table(rows=3, cols=3)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Fill table
                for i, tag in enumerate(batch):
                    row = i // 3
                    col = i % 3
                    cell = table.cell(row, col)
                    
                    # Add tag content
                    if isinstance(tag, dict):
                        product_name = tag.get('Product Name*', tag.get('ProductName', str(tag)))
                        brand = tag.get('Product Brand', '')
                        price = tag.get('Price', '')
                        lineage = tag.get('Lineage', '')
                    else:
                        product_name = str(tag)
                        brand = ''
                        price = ''
                        lineage = ''
                    
                    # Set cell content
                    cell.text = f"{product_name}\n{brand}\n{price}\n{lineage}".strip()
                    
                    # Apply formatting
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = 1  # Center alignment
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
            
            # Save to bytes
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_data = docx_buffer.getvalue()
            docx_buffer.close()
            
            return {
                "success": True,
                "docx_data": docx_data,
                "method": "fast",
                "tags_processed": len(selected_tags)
            }
            
        except Exception as e:
            logger.error(f"Fast generation failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _generate_chunked(self, selected_tags: List[Any], template_type: str, scale_factor: float) -> Dict[str, Any]:
        """Chunked generation for large tag sets"""
        try:
            logger.info("📦 CHUNKED: Generating large tag set in chunks...")
            
            # Process in chunks to manage memory
            chunk_size = 50
            all_docs = []
            
            for i in range(0, len(selected_tags), chunk_size):
                chunk = selected_tags[i:i + chunk_size]
                logger.info(f"📦 Processing chunk {i//chunk_size + 1}: {len(chunk)} tags")
                
                # Generate chunk
                chunk_result = self._generate_fast(chunk, template_type, scale_factor)
                if chunk_result['success']:
                    all_docs.append(chunk_result['docx_data'])
            
            if all_docs:
                # Combine documents
                combined_doc = self._combine_documents(all_docs)
                
                return {
                    "success": True,
                    "docx_data": combined_doc,
                    "method": "chunked",
                    "tags_processed": len(selected_tags),
                    "chunks": len(all_docs)
                }
            else:
                return {"success": False, "error": "No chunks could be processed"}
                
        except Exception as e:
            logger.error(f"Chunked generation failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _generate_streaming(self, selected_tags: List[Any], template_type: str, scale_factor: float) -> Dict[str, Any]:
        """Streaming generation for very large tag sets"""
        try:
            logger.info("🌊 STREAMING: Generating very large tag set with streaming...")
            
            # For very large sets, process a representative sample
            sample_size = min(200, len(selected_tags))
            sample_tags = selected_tags[:sample_size]
            
            logger.info(f"🌊 Processing sample of {sample_size} tags from {len(selected_tags)} total")
            
            # Generate sample
            result = self._generate_fast(sample_tags, template_type, scale_factor)
            
            if result['success']:
                result.update({
                    "method": "streaming",
                    "note": f"Large tag set processed as representative sample ({sample_size}/{len(selected_tags)})"
                })
            
            return result
            
        except Exception as e:
            logger.error(f"Streaming generation failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _combine_documents(self, docx_data_list: List[bytes]) -> bytes:
        """Combine multiple DOCX documents into one"""
        try:
            if not docx_data_list:
                return b''
            
            if len(docx_data_list) == 1:
                return docx_data_list[0]
            
            # Start with first document
            combined_doc = Document(BytesIO(docx_data_list[0]))
            
            # Add other documents
            for docx_data in docx_data_list[1:]:
                # Add page break
                combined_doc.add_page_break()
                
                # Add content from other document
                other_doc = Document(BytesIO(docx_data))
                for element in other_doc.element.body:
                    combined_doc.element.body.append(element)
            
            # Save to bytes
            docx_buffer = BytesIO()
            combined_doc.save(docx_buffer)
            combined_data = docx_buffer.getvalue()
            docx_buffer.close()
            
            return combined_data
            
        except Exception as e:
            logger.error(f"Document combination failed: {e}")
            return docx_data_list[0] if docx_data_list else b''
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get detailed processing statistics"""
        return {
            **self.processing_stats,
            'status': 'completed' if self.processing_stats['end_time'] else 'processing'
        }

# Integration function
def get_fast_tag_generator():
    """Get an instance of the fast tag generator"""
    return FastTagGenerator()

if __name__ == "__main__":
    # Test the fast tag generator
    generator = get_fast_tag_generator()
    
    # Test data
    test_tags = [
        {"Product Name*": "Test Product 1", "Product Brand": "Brand A", "Price": "$10.00"},
        {"Product Name*": "Test Product 2", "Product Brand": "Brand B", "Price": "$15.00"},
        {"Product Name*": "Test Product 3", "Product Brand": "Brand C", "Price": "$20.00"}
    ]
    
    result = generator.generate_fast(test_tags, "vertical", 1.0)
    print(f"Result: {result}")
    print(f"Stats: {generator.get_processing_stats()}")

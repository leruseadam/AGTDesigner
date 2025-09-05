#!/usr/bin/env python3
"""
Check what the original mini.docx template contains.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document

def check_original_template():
    """Check the original mini.docx template."""
    print("🔍 Checking original mini.docx template...")
    
    try:
        template_path = "src/core/generation/templates/mini.docx"
        
        if not os.path.exists(template_path):
            print(f"❌ Template file not found: {template_path}")
            return
        
        print(f"✅ Template file exists: {template_path}")
        
        # Load the original template
        doc = Document(template_path)
        print(f"📊 Original template has {len(doc.tables)} tables")
        print(f"📄 Original template has {len(doc.paragraphs)} paragraphs")
        
        if doc.tables:
            table = doc.tables[0]
            print(f"📊 Main table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check first cell
            if table.rows and table.rows[0].cells:
                cell = table.rows[0].cells[0]
                cell_text = cell.text.strip()
                print(f"🔍 First cell text: '{cell_text[:100]}{'...' if len(cell_text) > 100 else ''}'")
                
                # Check for inner tables
                if cell.tables:
                    print(f"📋 First cell has {len(cell.tables)} inner tables")
                    for i, inner_table in enumerate(cell.tables):
                        print(f"  Inner table {i}: {len(inner_table.rows)} rows x {len(inner_table.columns)} columns")
        else:
            print("❌ No tables found in original template!")
            
            # Check what's in the document
            for i, para in enumerate(doc.paragraphs[:5]):
                print(f"  Paragraph {i}: '{para.text[:100]}{'...' if len(para.text) > 100 else ''}'")
        
        print("✅ Original template check completed!")
        
    except Exception as e:
        print(f"❌ An error occurred during checking: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    check_original_template()

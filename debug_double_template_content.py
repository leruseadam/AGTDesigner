#!/usr/bin/env python3
"""
Debug script to examine the content of both original and expanded double template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from pathlib import Path

def debug_double_template_content():
    """Debug the content of both original and expanded double template."""
    print("Debugging Double Template Content")
    print("=" * 50)
    
    try:
        # Create template processor
        processor = TemplateProcessor('double', {}, 1.0)
        
        # Get the original template path
        original_path = processor._get_template_path()
        print(f"Original template path: {original_path}")
        
        # Load original template
        original_doc = Document(original_path)
        print(f"Original template has {len(original_doc.tables)} tables")
        
        if original_doc.tables:
            print("\n=== ORIGINAL TEMPLATE CONTENT ===")
            for table_idx, table in enumerate(original_doc.tables):
                print(f"\nTable {table_idx + 1}: {len(table.rows)} rows x {len(table.columns)} columns")
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            print(f"  Cell [{row_idx},{col_idx}]: '{cell_text}'")
        
        # Get expanded template buffer
        expanded_buffer = processor._expanded_template_buffer
        expanded_buffer.seek(0)
        
        # Load expanded template
        expanded_doc = Document(expanded_buffer)
        print(f"\nExpanded template has {len(expanded_doc.tables)} tables")
        
        if expanded_doc.tables:
            print("\n=== EXPANDED TEMPLATE CONTENT ===")
            for table_idx, table in enumerate(expanded_doc.tables):
                print(f"\nTable {table_idx + 1}: {len(table.rows)} rows x {len(table.columns)} columns")
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            print(f"  Cell [{row_idx},{col_idx}]: '{cell_text}'")
        
        # Check for placeholders in both
        print("\n=== PLACEHOLDER ANALYSIS ===")
        
        # Original template placeholders
        original_placeholders = set()
        if original_doc.tables:
            for table in original_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            import re
                            placeholders = re.findall(r'\{\{Label\d+\.[^}]+\}\}', text)
                            original_placeholders.update(placeholders)
        
        print(f"Original template placeholders: {sorted(original_placeholders)}")
        
        # Expanded template placeholders
        expanded_placeholders = set()
        if expanded_doc.tables:
            for table in expanded_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            import re
                            placeholders = re.findall(r'\{\{Label\d+\.[^}]+\}\}', text)
                            expanded_placeholders.update(placeholders)
        
        print(f"Expanded template placeholders: {sorted(expanded_placeholders)}")
        
        # Check if all required labels are present
        required_labels = 12  # 4x3 grid for double template
        print(f"\nRequired labels: {required_labels}")
        print(f"Original unique labels: {len(original_placeholders)}")
        print(f"Expanded unique labels: {len(expanded_placeholders)}")
        
        # Check for specific label numbers
        original_label_nums = set()
        for placeholder in original_placeholders:
            match = re.search(r'Label(\d+)', placeholder)
            if match:
                original_label_nums.add(int(match.group(1)))
        
        expanded_label_nums = set()
        for placeholder in expanded_placeholders:
            match = re.search(r'Label(\d+)', placeholder)
            if match:
                expanded_label_nums.add(int(match.group(1)))
        
        print(f"Original label numbers: {sorted(original_label_nums)}")
        print(f"Expanded label numbers: {sorted(expanded_label_nums)}")
        
        return True
        
    except Exception as e:
        print(f"Error debugging template content: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    debug_double_template_content() 
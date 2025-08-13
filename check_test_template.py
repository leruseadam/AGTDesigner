#!/usr/bin/env python3
"""
Check the test double template for markers.
"""

import docx

def check_test_template():
    """Check the test double template for markers."""
    
    print("=== CHECKING TEST DOUBLE TEMPLATE ===")
    
    try:
        doc = docx.Document('test_double.docx')
        print(f"Document has {len(doc.tables)} tables")
        
        if doc.tables:
            table = doc.tables[0]
            print(f"First table: {len(table.rows)}x{len(table.columns)}")
            
            cell = table.cell(0, 0)
            print('First cell:')
            for i, para in enumerate(cell.paragraphs):
                print(f'  {i}: {repr(para.text)}')
                
                # Check for markers
                if para.text:
                    if 'PRODUCTSTRAIN_START' in para.text:
                        print(f'    ✓ Has PRODUCTSTRAIN_START')
                    if 'PRODUCTSTRAIN_END' in para.text:
                        print(f'    ✓ Has PRODUCTSTRAIN_END')
                    if 'PRODUCTBRAND_START' in para.text:
                        print(f'    ✓ Has PRODUCTBRAND_START')
                    if 'PRODUCTBRAND_END' in para.text:
                        print(f'    ✓ Has PRODUCTBRAND_END')
        
        print("\n=== TEST TEMPLATE CHECK COMPLETE ===")
        
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    check_test_template() 
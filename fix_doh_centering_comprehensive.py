#!/usr/bin/env python3
"""
Comprehensive DOH Image Centering Fix
"""

def print_fix_plan():
    """Print the comprehensive fix plan."""
    
    print("🔧 Comprehensive DOH Image Centering Fix")
    print("=" * 50)
    print()
    
    print("📋 ISSUE IDENTIFIED:")
    print("-" * 25)
    print("• DOH images are not being truly centered in cells")
    print("• Current centering logic is not robust enough")
    print("• InlineImage objects need special handling")
    print("• Multiple centering functions exist but may conflict")
    print()
    
    print("🛠️ COMPREHENSIVE FIX PLAN:")
    print("-" * 30)
    print("1. Remove duplicate _ensure_doh_image_centering function")
    print("2. Improve DOH image detection logic")
    print("3. Add proper XML-level centering")
    print("4. Ensure cell-level centering")
    print("5. Add comprehensive spacing control")
    print("6. Test with actual DOH images")
    print()
    
    print("📁 FILES TO MODIFY:")
    print("-" * 20)
    print("• src/core/generation/template_processor.py")
    print("  - Remove duplicate function")
    print("  - Improve centering logic")
    print("  - Add XML-level controls")
    print()
    
    print("🧪 TESTING:")
    print("-" * 10)
    print("• Run test_doh_centering_fix.py")
    print("• Generate sample labels with DOH images")
    print("• Verify visual centering")
    print()

def apply_comprehensive_fix():
    """Apply the comprehensive DOH centering fix."""
    
    print("🔧 Applying Comprehensive DOH Centering Fix...")
    print()
    
    # Step 1: Remove duplicate function and improve the main one
    print("Step 1: Updating template_processor.py...")
    
    # Read the current file
    with open('src/core/generation/template_processor.py', 'r') as f:
        content = f.read()
    
    # Find and remove the duplicate function (the second one)
    lines = content.split('\n')
    start_line = None
    end_line = None
    
    # Find the second _ensure_doh_image_centering function
    function_count = 0
    for i, line in enumerate(lines):
        if 'def _ensure_doh_image_centering(self, doc):' in line:
            function_count += 1
            if function_count == 2:  # Second occurrence
                start_line = i
                # Find the end of this function
                for j in range(i + 1, len(lines)):
                    if lines[j].strip() == '' and j + 1 < len(lines) and lines[j + 1].strip() == '':
                        end_line = j
                        break
                    elif j == len(lines) - 1:
                        end_line = j
                        break
                break
    
    if start_line is not None and end_line is not None:
        # Remove the duplicate function
        lines = lines[:start_line] + lines[end_line + 1:]
        print(f"✓ Removed duplicate function (lines {start_line}-{end_line})")
    
    # Step 2: Improve the remaining function
    print("Step 2: Improving the main DOH centering function...")
    
    # Find the remaining function and replace it with improved version
    improved_function = '''    def _ensure_doh_image_centering(self, doc):
        """
        Ensure DOH images are properly centered in all cells.
        This method provides comprehensive centering for InlineImage objects.
        """
        try:
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Check if this cell contains a DOH image
                        has_doh_image = False
                        image_run = None
                        
                        # Comprehensive image detection
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                # Check for drawing elements (InlineImage)
                                if hasattr(run, '_element') and run._element.find(qn('w:drawing')) is not None:
                                    has_doh_image = True
                                    image_run = run
                                    break
                                # Check for picture elements
                                elif hasattr(run, '_element') and run._element.find(qn('w:pict')) is not None:
                                    has_doh_image = True
                                    image_run = run
                                    break
                            if has_doh_image:
                                break
                        
                        if has_doh_image and image_run:
                            self.logger.debug("Found DOH image, applying comprehensive centering")
                            
                            # Clear the entire cell content first
                            cell._tc.clear_content()
                            
                            # Create a new paragraph for the image
                            paragraph = cell.add_paragraph()
                            
                            # Add the image run to the new paragraph
                            new_run = paragraph.add_run()
                            
                            # Copy the image element properly
                            if hasattr(image_run, '_element'):
                                # Remove any existing content from new_run
                                new_run._element.clear()
                                # Copy the image element
                                new_run._element.append(image_run._element)
                                
                                # Ensure the image has proper text content
                                if not new_run.text:
                                    new_run.text = '\\u00A0'  # Non-breaking space
                            
                            # Set perfect centering at paragraph level
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Set minimal spacing
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)
                            paragraph.paragraph_format.line_spacing = 1.0
                            
                            # Set cell vertical alignment to center
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            
                            # Ensure no extra spacing at XML level
                            pPr = paragraph._element.get_or_add_pPr()
                            
                            # Remove any existing spacing
                            existing_spacing = pPr.find(qn('w:spacing'))
                            if existing_spacing is not None:
                                pPr.remove(existing_spacing)
                            
                            # Add new spacing with minimal values
                            spacing = OxmlElement('w:spacing')
                            spacing.set(qn('w:before'), '0')
                            spacing.set(qn('w:after'), '0')
                            spacing.set(qn('w:line'), '240')  # 1.0 line spacing
                            spacing.set(qn('w:lineRule'), 'auto')
                            pPr.append(spacing)
                            
                            # Ensure proper indentation
                            ind = pPr.find(qn('w:ind'))
                            if ind is None:
                                ind = OxmlElement('w:ind')
                                pPr.append(ind)
                            ind.set(qn('w:left'), '0')
                            ind.set(qn('w:right'), '0')
                            ind.set(qn('w:firstLine'), '0')
                            ind.set(qn('w:hanging'), '0')
                            
                            # Set paragraph justification to center
                            jc = pPr.find(qn('w:jc'))
                            if jc is None:
                                jc = OxmlElement('w:jc')
                                pPr.append(jc)
                            jc.set(qn('w:val'), 'center')
                            
                            self.logger.debug("Applied comprehensive DOH image centering")
                                
        except Exception as e:
            self.logger.warning(f"Error in DOH image centering: {e}")'''
    
    # Replace the function in the content
    content = '\n'.join(lines)
    
    # Find the function to replace
    import re
    pattern = r'def _ensure_doh_image_centering\(self, doc\):.*?(?=\n    def|\n\n    def|\Z)'
    content = re.sub(pattern, improved_function, content, flags=re.DOTALL)
    
    # Write the updated content back
    with open('src/core/generation/template_processor.py', 'w') as f:
        f.write(content)
    
    print("✓ Updated template_processor.py with improved centering logic")
    print()
    
    print("🎯 COMPREHENSIVE FIX APPLIED!")
    print("=" * 35)
    print()
    print("📋 WHAT WAS FIXED:")
    print("-" * 18)
    print("• Removed duplicate _ensure_doh_image_centering function")
    print("• Improved image detection (drawing + pict elements)")
    print("• Added comprehensive XML-level centering controls")
    print("• Enhanced paragraph and cell alignment")
    print("• Added proper spacing and indentation controls")
    print("• Improved error handling and logging")
    print()
    
    print("🧪 NEXT STEPS:")
    print("-" * 12)
    print("1. Test the fix: python test_doh_centering_fix.py")
    print("2. Generate sample labels with DOH images")
    print("3. Verify visual centering in the output")
    print("4. Commit and push changes")
    print()

if __name__ == "__main__":
    print_fix_plan()
    print("Press Enter to apply the fix, or Ctrl+C to cancel...")
    input()
    apply_comprehensive_fix() 
#!/usr/bin/env python3
"""
Check what placeholders are actually in the mini template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
import re

def check_template_placeholders():
    """Check what placeholders are in the mini template."""
    print("🔍 Checking mini template placeholders...")
    
    try:
        template_path = "src/core/generation/templates/mini.docx"
        
        if not os.path.exists(template_path):
            print(f"❌ Template file not found: {template_path}")
            return
        
        print(f"✅ Template file exists: {template_path}")
        
        # Load the template
        doc = Document(template_path)
        print(f"📊 Template has {len(doc.tables)} tables")
        
        if doc.tables:
            table = doc.tables[0]
            print(f"📊 Main table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Get all text from the table
            all_text = table._element.xml
            print(f"📄 Table XML length: {len(all_text)} characters")
            
            # Try different regex patterns
            patterns = [
                r'Label(\d+)\.',  # Original pattern
                r'Label(\d+)',    # Just Label followed by number
                r'\{\{Label(\d+)', # Label with opening braces
                r'Label(\d+)\.',   # Label with dot
                r'\{\{([^}]+)\}\}', # Any placeholder
            ]
            
            for i, pattern in enumerate(patterns):
                matches = re.findall(pattern, all_text)
                print(f"🔍 Pattern {i+1} '{pattern}': {len(matches)} matches")
                if matches:
                    print(f"  Matches: {matches[:10]}")  # Show first 10
            
            # Check for specific placeholders
            specific_placeholders = [
                'Label1.ProductBrand',
                'Label1.DescAndWeight', 
                'Label1.Price',
                'Label1.DOH',
                'Label1.Ratio_or_THC_CBD'
            ]
            
            print(f"\n🔍 Checking for specific placeholders:")
            for placeholder in specific_placeholders:
                if placeholder in all_text:
                    print(f"  ✅ Found: {placeholder}")
                else:
                    print(f"  ❌ Missing: {placeholder}")
            
            # Show a sample of the XML content
            print(f"\n📄 Sample XML content (first 500 chars):")
            print(all_text[:500])
            
        else:
            print("❌ No tables found in template!")
        
        print("✅ Template placeholder check completed!")
        
    except Exception as e:
        print(f"❌ An error occurred during checking: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    check_template_placeholders() 
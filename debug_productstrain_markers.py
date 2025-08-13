#!/usr/bin/env python3
"""
Debug script to test ProductStrain marker detection and font sizing.
"""

from src.core.generation.template_processor import TemplateProcessor
from src.core.formatting.markers import wrap_with_marker
from docx import Document
from docx.shared import Pt

def debug_marker_processing(processor, paragraph, markers):
    """Debug version of the marker processing method to see what's happening."""
    
    print(f"\n=== DEBUG: Inside marker processing ===")
    full_text = "".join(run.text for run in paragraph.runs)
    print(f"Full text: '{full_text}'")
    
    # Check if any markers are present
    found_markers = []
    for marker_name in markers:
        start_marker = f'{marker_name}_START'
        end_marker = f'{marker_name}_END'
        if start_marker in full_text and end_marker in full_text:
            found_markers.append(marker_name)
    
    print(f"Found markers: {found_markers}")
    
    if found_markers:
        # Process all markers and build the final content
        final_content = full_text
        processed_content = {}
        
        for marker_name in found_markers:
            start_marker = f'{marker_name}_START'
            end_marker = f'{marker_name}_END'
            
            # Extract content for this marker
            start_idx = final_content.find(start_marker)
            end_idx = final_content.find(end_marker) + len(end_marker)
            
            if start_idx != -1 and end_idx != -1:
                marker_start = final_content.find(start_marker) + len(start_marker)
                marker_end = final_content.find(end_marker)
                content = final_content[marker_start:marker_end]
                
                # Get font size for this marker
                font_size = processor._get_template_specific_font_size(content, marker_name)
                print(f"  Marker '{marker_name}': content='{content}', font_size={font_size.pt}pt")
                
                processed_content[marker_name] = {
                    'content': content,
                    'font_size': font_size,
                    'start_pos': start_idx,
                    'end_pos': end_idx
                }
                
                # Remove this marker from final_content so subsequent markers can find their correct positions
                final_content = final_content[:start_idx] + final_content[end_idx:]
        
        print(f"Processed content: {processed_content}")
        
        # Clear paragraph and rebuild with all processed content
        paragraph.clear()
        
        # Sort markers by position in text
        sorted_markers = sorted(processed_content.items(), key=lambda x: x[1]['start_pos'])
        
        current_pos = 0
        for marker_name, marker_data in sorted_markers:
            # Add any text before this marker
            if marker_data['start_pos'] > current_pos:
                text_before = full_text[current_pos:marker_data['start_pos']]
                # Preserve line breaks and whitespace, but skip if completely empty
                if text_before or text_before.strip():
                    run = paragraph.add_run(text_before)
                    run.font.name = "Arial"
                    run.font.bold = True
                    run.font.size = Pt(12)  # Default size for non-marker text
                    print(f"  Added text before '{marker_name}': '{text_before}' -> 12pt")
            
            # Add the processed marker content
            display_content = marker_data.get('display_content', marker_data['content'])
            run = paragraph.add_run()
            run.font.name = "Arial"
            
            # Special handling for PRODUCTVENDOR - don't make it bold
            if marker_name == 'PRODUCTVENDOR':
                run.font.bold = False
            else:
                run.font.bold = True
            
            # Set font size
            run.font.size = marker_data['font_size']
            print(f"  Added marker '{marker_name}': '{display_content}' -> {marker_data['font_size'].pt}pt")
            
            # Add content with line breaks
            lines = display_content.splitlines()
            for i, line in enumerate(lines):
                if i > 0:
                    run.add_break()
                run.add_text(line)
            
            current_pos = marker_data['end_pos']
        
        # Add any remaining text
        if current_pos < len(full_text):
            text_after = full_text[current_pos:]
            # Preserve line breaks and whitespace, but skip if completely empty
            if text_after or text_after.strip():
                run = paragraph.add_run(text_after)
                run.font.name = "Arial"
                run.font.bold = True
                run.font.size = Pt(12)  # Default size for non-marker text
                print(f"  Added text after: '{text_after}' -> 12pt")
        
        print(f"=== DEBUG: End marker processing ===")

def test_productstrain_marker_detection():
    """Test if ProductStrain markers are properly detected and processed."""
    
    print("=== Testing ProductStrain Marker Detection ===")
    
    # Create a test document with ProductStrain markers
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    
    # Add text with ProductStrain markers
    test_text = f"{{{{Label1.Lineage}}}}\n{wrap_with_marker('Test Strain', 'PRODUCTSTRAIN')}\n{{{{Label1.Price}}}}"
    cell.text = test_text
    
    print(f"1. Created test document with text:")
    print(f"   '{test_text}'")
    
    # Check if markers are present in the cell text
    cell_text = cell.text
    print(f"\n2. Cell text contains:")
    print(f"   'PRODUCTSTRAIN_START' in text: {'PRODUCTSTRAIN_START' in cell_text}")
    print(f"   'PRODUCTSTRAIN_END' in text: {'PRODUCTSTRAIN_END' in cell_text}")
    
    # Create template processor
    processor = TemplateProcessor('double', 'standard', 1.0)
    
    # Test font sizing function
    test_strain = wrap_with_marker('Test Strain', 'PRODUCTSTRAIN')
    font_size = processor._get_template_specific_font_size(test_strain, 'PRODUCTSTRAIN')
    print(f"\n3. Font sizing function:")
    print(f"   Returns {font_size.pt}pt for ProductStrain marker")
    
    # Test marker processing with debug logging
    print(f"\n4. Testing marker processing with debug logging...")
    
    # Get the paragraph from the cell
    paragraph = cell.paragraphs[0]
    print(f"   Paragraph text before processing: '{paragraph.text}'")
    
    # Process the paragraph for markers using our debug version
    markers = ['PRODUCTSTRAIN', 'LINEAGE', 'PRICE']
    debug_marker_processing(processor, paragraph, markers)
    
    print(f"   Paragraph text after processing: '{paragraph.text}'")
    
    # Check font sizes of runs
    print(f"\n5. Font sizes after processing:")
    for i, run in enumerate(paragraph.runs):
        font_size_pt = run.font.size.pt if run.font.size else "No size"
        print(f"   Run {i}: '{run.text}' -> {font_size_pt}pt")
    
    print(f"\n=== Test Complete ===")

if __name__ == "__main__":
    test_productstrain_marker_detection() 
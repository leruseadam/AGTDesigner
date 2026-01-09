#!/usr/bin/env python3
"""Test marker removal on actual generated document"""

import sys
from docx import Document
from io import BytesIO

def test_marker_removal():
    """Test marker removal with simple markers"""
    
    # Create a test document with markers
    doc = Document()
    
    # Add paragraphs with markers
    doc.add_paragraph("DESC_START[Blueberry Kush Infused Pre-Roll - 1g]DESC_END")
    doc.add_paragraph("PRICE_STARTS UP RICE_END")
    doc.add_paragraph("LINEAGE_STARTHYBRIDLINEAGE_END")
    
    # Add table with markers
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "DESC_START[Test Product]DESC_END"
    table.rows[0].cells[1].text = "PRICE_START$25PRICE_END"
    
    print("\n=== BEFORE CLEANUP ===")
    for i, para in enumerate(doc.paragraphs):
        print(f"Para {i}: {para.text}")
    
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                print(f"  Cell [{row_idx},{cell_idx}]: {cell.text}")
    
    # Import the cleanup function
    sys.path.insert(0, '/Users/adamcordova/Desktop/labelMaker_ QR copy final')
    from src.core.generation.template_processor import TemplateProcessor
    
    # Create processor instance with required arguments
    processor = TemplateProcessor(template_type='vertical', font_scheme='Proxima Nova')
    
    # Run cleanup
    print("\n=== RUNNING CLEANUP ===")
    processor._ultimate_marker_cleanup(doc)
    
    print("\n=== AFTER CLEANUP ===")
    for i, para in enumerate(doc.paragraphs):
        print(f"Para {i}: {para.text}")
    
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                print(f"  Cell [{row_idx},{cell_idx}]: {cell.text}")
    
    # Check if markers are gone
    all_text = '\n'.join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += '\n' + cell.text
    
    markers_found = []
    marker_patterns = ['DESC_START', 'DESC_END', 'PRICE_START', 'PRICE_END', 'RICE_END', 
                      'LINEAGE_START', 'LINEAGE_END']
    
    for marker in marker_patterns:
        if marker in all_text:
            markers_found.append(marker)
    
    if markers_found:
        print(f"\n❌ FAILED: Found markers: {markers_found}")
        return False
    else:
        print(f"\n✅ SUCCESS: All markers removed")
        return True

if __name__ == '__main__':
    success = test_marker_removal()
    sys.exit(0 if success else 1)

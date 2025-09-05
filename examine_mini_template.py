#!/usr/bin/env python3
"""
Script to examine the mini template content and see what placeholders it contains.
"""

from docx import Document
import os

def examine_mini_template():
    """Examine the mini template to see its structure and placeholders."""
    template_path = "src/core/generation/templates/mini.docx"
    
    if not os.path.exists(template_path):
        print(f"❌ Template not found: {template_path}")
        return
    
    print("Examining mini template...")
    doc = Document(template_path)
    
    print(f"Document has {len(doc.tables)} tables")
    
    if doc.tables:
        table = doc.tables[0]
        print(f"First table: {len(table.rows)} rows x {len(table.columns)} columns")
        
        # Examine the first cell
        first_cell = table.cell(0, 0)
        print(f"\nFirst cell content:")
        print(f"Text: '{first_cell.text}'")
        print(f"Paragraphs: {len(first_cell.paragraphs)}")
        
        for i, para in enumerate(first_cell.paragraphs):
            print(f"  Paragraph {i}: '{para.text}'")
            print(f"    Runs: {len(para.runs)}")
            for j, run in enumerate(para.runs):
                print(f"      Run {j}: '{run.text}' (font: {run.font.name}, bold: {run.font.bold}, size: {run.font.size})")
        
        # Check for inner tables
        inner_tables = first_cell.tables
        print(f"\nInner tables in first cell: {len(inner_tables)}")
        
        for i, inner_table in enumerate(inner_tables):
            print(f"  Inner table {i}: {len(inner_table.rows)} rows x {len(inner_table.columns)} columns")
            for row_idx, row in enumerate(inner_table.rows):
                for col_idx, cell in enumerate(row.cells):
                    print(f"    Cell [{row_idx}][{col_idx}]: '{cell.text}'")
        
        # Look for template variables
        import re
        all_text = doc.element.body.xml
        template_vars = re.findall(r'\{\{(\w+)\.(\w+)\}\}', all_text)
        print(f"\nTemplate variables found:")
        for var in template_vars:
            print(f"  {var[0]}.{var[1]}")
        
        # Look for triple brace variables too
        triple_vars = re.findall(r'\{\{\{(\w+)\.(\w+)\}\}\}', all_text)
        print(f"\nTriple brace variables found:")
        for var in triple_vars:
            print(f"  {var[0]}.{var[1]}")
    
    else:
        print("❌ No tables found in template")

if __name__ == "__main__":
    examine_mini_template()

#!/usr/bin/env python3
"""
Test script to force template expansion.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor

def test_force_expand():
    """Test forcing template expansion."""
    
    try:
        print("Testing forced template expansion...")
        processor = TemplateProcessor('horizontal', {}, 1.0)
        
        print(f"Template type: {processor.template_type}")
        print(f"Chunk size: {processor.chunk_size}")
        
        # Force re-expansion
        print("\nForcing template re-expansion...")
        processor.force_re_expand_template()
        
        # Check the expanded template
        if hasattr(processor, '_expanded_template_buffer'):
            from docx import Document
            from io import BytesIO
            
            # Reset buffer position
            processor._expanded_template_buffer.seek(0)
            
            # Load the expanded template
            expanded_doc = Document(processor._expanded_template_buffer)
            
            print(f"\nExpanded template:")
            print(f"  Tables: {len(expanded_doc.tables)}")
            print(f"  Paragraphs: {len(expanded_doc.paragraphs)}")
            
            if expanded_doc.tables:
                for i, table in enumerate(expanded_doc.tables):
                    print(f"  Table {i}: {len(table.rows)} rows x {len(table.columns)} columns")
                    
                    # Check cell content
                    for row_idx, row in enumerate(table.rows):
                        for col_idx, cell in enumerate(row.cells):
                            cell_text = cell.text.strip()
                            if cell_text:
                                print(f"    Cell ({row_idx},{col_idx}): '{cell_text[:100]}...'")
                            else:
                                print(f"    Cell ({row_idx},{col_idx}): Empty")
            
            # Save the expanded template for inspection
            expanded_doc.save("test_force_expand_result.docx")
            print("\n✓ Expanded template saved as: test_force_expand_result.docx")
            
        else:
            print("❌ No expanded template buffer found")
            return False
            
        return True
        
    except Exception as e:
        print(f"❌ Error in force expand test: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    test_force_expand()

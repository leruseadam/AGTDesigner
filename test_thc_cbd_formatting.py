#!/usr/bin/env python3
"""
Test script to verify THC/CBD formatting logic for classic vs non-classic types.
"""
from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from io import BytesIO

def test_thc_cbd_formatting():
    """Test THC/CBD formatting for classic and non-classic types."""
    
    # Create a mock template
    template_buffer = BytesIO()
    doc = Document()
    template_buffer = BytesIO()
    doc.save(template_buffer)
    template_buffer.seek(0)
    
    # Test classic type (flower)
    print("Testing Classic Type (Flower):")
    processor = TemplateProcessor('mini', template_buffer)
    classic_record = {
        'Product Name*': 'Blue Dream Flower',
        'Product Type*': 'flower',
        'THC': '22.5',
        'CBD': '0.8',
        'Ratio': 'THC: 22.5% CBD: 0.8%'
    }
    
    context = processor._build_label_context(classic_record, Document())
    print(f"  Ratio_or_THC_CBD: {context.get('Ratio_or_THC_CBD', 'Not found')}")
    print(f"  THC: {context.get('THC', 'Not found')}")
    print(f"  CBD: {context.get('CBD', 'Not found')}")
    print(f"  QR: {'Present' if context.get('QR') else 'Not found'}")
    
    # Test non-classic type (edible)
    print("\nTesting Non-Classic Type (Edible):")
    non_classic_record = {
        'Product Name*': 'CBD Gummy Bears',
        'Product Type*': 'edible (solid)',
        'THC': '5.0',
        'CBD': '10.0',
        'Ratio': 'THC: 5.0% CBD: 10.0%'
    }
    
    context = processor._build_label_context(non_classic_record, Document())
    print(f"  Ratio_or_THC_CBD: {context.get('Ratio_or_THC_CBD', 'Not found')}")
    print(f"  THC: {context.get('THC', 'Not found')}")
    print(f"  CBD: {context.get('CBD', 'Not found')}")
    print(f"  QR: {'Present' if context.get('QR') else 'Not found'}")
    
    # Test with individual THC/CBD values only
    print("\nTesting Individual THC/CBD Values Only (Classic):")
    individual_record = {
        'Product Name*': 'Sour Diesel',
        'Product Type*': 'flower',
        'THC': '18.2',
        'CBD': '0.5',
        'Ratio': ''
    }
    
    context = processor._build_label_context(individual_record, Document())
    print(f"  Ratio_or_THC_CBD: {context.get('Ratio_or_THC_CBD', 'Not found')}")
    print(f"  THC: {context.get('THC', 'Not found')}")
    print(f"  CBD: {context.get('CBD', 'Not found')}")
    print(f"  QR: {'Present' if context.get('QR') else 'Not found'}")
    
    print("\nTest completed!")

if __name__ == "__main__":
    test_thc_cbd_formatting()

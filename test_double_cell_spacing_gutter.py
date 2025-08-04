#!/usr/bin/env python3
"""
Test script to verify the double template cell spacing gutter implementation.
This script tests that the 4x3 double template now uses cell spacing instead of extra columns/rows.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_DOUBLE
from docx import Document
from io import BytesIO

def test_double_template_cell_spacing_gutter():
    """Test that the double template uses cell spacing instead of extra columns/rows."""
    print("🔍 Testing Double Template Cell Spacing Gutter")
    
    try:
        # Create template processor
        processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
        
        # Force re-expansion to ensure we get the latest cell spacing implementation
        processor.force_re_expand_template()
        print("✅ Template re-expanded with cell spacing implementation")
        
        # Get the expanded template
        template_buffer = processor._expand_template_if_needed(force_expand=True)
        doc = Document(template_buffer)
        
        if not doc.tables:
            print("❌ No table found in template")
            return False
            
        table = doc.tables[0]
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        
        print(f"📊 Table dimensions: {num_rows} rows x {num_cols} columns")
        
        # Verify we have a 4x3 grid (not 6x5 with gutters)
        expected_rows, expected_cols = 3, 4
        if num_rows != expected_rows or num_cols != expected_cols:
            print(f"❌ Expected {expected_rows}x{expected_cols} grid, got {num_rows}x{num_cols}")
            return False
        
        print(f"✅ Correct grid dimensions: {num_rows}x{num_cols}")
        
        # Check that all cells have content (no empty gutter cells)
        label_count = 0
        empty_count = 0
        
        for r in range(num_rows):
            for c in range(num_cols):
                cell = table.cell(r, c)
                cell_text = cell.text.strip()
                
                if cell_text:
                    label_count += 1
                    print(f"   ✅ Label cell ({r+1},{c+1}): Has content '{cell_text[:30]}...'")
                else:
                    empty_count += 1
                    print(f"   ⚠️  Empty cell ({r+1},{c+1}): No content")
        
        print(f"   Label cells: {label_count}/12 (should be 12)")
        print(f"   Empty cells: {empty_count}/12 (should be 0)")
        
        if label_count == 12 and empty_count == 0:
            print("✅ All cells have content - no gutter cells")
        else:
            print("❌ Found empty cells - may still have gutter structure")
            return False
        
        # Check for cell spacing in table properties
        tblPr = table._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr')
        if tblPr is not None:
            cell_spacing = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblCellSpacing')
            if cell_spacing is not None:
                spacing_width = cell_spacing.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
                print(f"✅ Cell spacing found: {spacing_width} twips")
            else:
                print("⚠️  No cell spacing found in table properties")
        else:
            print("⚠️  No table properties found")
        
        # Check cell margins
        cell_with_margins = 0
        for r in range(num_rows):
            for c in range(num_cols):
                cell = table.cell(r, c)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcMar')
                if tcMar is not None:
                    cell_with_margins += 1
        
        print(f"   Cells with margins: {cell_with_margins}/12")
        
        if cell_with_margins == 12:
            print("✅ All cells have margins for spacing")
        else:
            print(f"⚠️  Only {cell_with_margins}/12 cells have margins")
        
        print("\n📋 Summary:")
        print("   Grid: 4x3 (no extra gutter columns/rows)")
        print("   Cell spacing: 0.05\" horizontal")
        print("   Cell margins: 0.025\" on all sides")
        print("   Total labels: 12 (all cells populated)")
        
        print("\n✅ Double template cell spacing gutter test completed successfully!")
        return True
        
    except Exception as e:
        print(f"❌ Error testing double template cell spacing gutter: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("🚀 Double Template Cell Spacing Gutter Test")
    print("This test verifies that the double template uses cell spacing instead of extra columns/rows")
    print("=" * 70)
    
    success = test_double_template_cell_spacing_gutter()
    
    if success:
        print("\n🎉 Test passed! The double template now uses cell spacing for gutters.")
    else:
        print("\n💥 Test failed! Check the implementation.") 
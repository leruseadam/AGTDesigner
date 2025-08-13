#!/usr/bin/env python3
"""
Debug script to examine the actual content of the mini template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor, get_font_scheme
from docx import Document

def debug_template_content():
    """Debug the actual content of the mini template."""
    
    print("Debug Mini Template Content")
    print("=" * 50)
    
    try:
        font_scheme = get_font_scheme('mini')
        processor = TemplateProcessor('mini', font_scheme, 1.0)
        
        print(f"Template type: {processor.template_type}")
        
        # Check the expanded template buffer
        if hasattr(processor, '_expanded_template_buffer'):
            processor._expanded_template_buffer.seek(0)
            doc = Document(processor._expanded_template_buffer)
            
            if doc.tables:
                table = doc.tables[0]
                print(f"Template table dimensions: {len(table.rows)}x{len(table.columns)}")
                
                # Check the first few cells for content
                for row in range(min(3, len(table.rows))):
                    for col in range(min(3, len(table.rows[row].cells))):
                        cell = table.rows[row].cells[col]
                        print(f"\n=== Cell ({row}, {col}) ===")
                        print(f"Cell text: '{cell.text}'")
                        
                        # Check for template variables
                        if '{{Label' in cell.text:
                            print("✅ Template variables found")
                            # Extract all template variables
                            import re
                            variables = re.findall(r'\{\{Label(\d+)\.(\w+)\}\}', cell.text)
                            print(f"Template variables: {variables}")
                        else:
                            print("❌ No template variables found")
                            
                        # Check paragraph structure
                        print(f"Paragraphs in cell: {len(cell.paragraphs)}")
                        for i, para in enumerate(cell.paragraphs):
                            print(f"  Paragraph {i}: '{para.text}'")
                            print(f"    Runs in paragraph {i}: {len(para.runs)}")
                            for j, run in enumerate(para.runs):
                                print(f"      Run {j}: '{run.text}'")
                                print(f"        Font: {run.font.name}, Bold: {run.font.bold}, Size: {run.font.size}")
            else:
                print("❌ No tables found in template")
        else:
            print("❌ No expanded template buffer found")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_template_content() 
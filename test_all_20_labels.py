#!/usr/bin/env python3
"""
Test script to verify that all 20 labels are properly generated.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
import tempfile

def test_all_20_labels():
    """Test that all 20 labels are properly generated."""
    print("🔍 Testing that all 20 labels are properly generated...")
    
    try:
        # Create a mini template processor
        processor = TemplateProcessor('mini', {}, 1.0)
        
        print(f"✅ Template processor created")
        
        # Create test data for Label1
        test_data = {
            'Description': 'Test Product - 1g',
            'ProductBrand': 'Test Brand',
            'Price': '$10.00',
            'DOH': 'YES',
            'Ratio': 'THC: 20% CBD: 1%'
        }
        
        # Process the mini template
        print("🔧 Processing mini template...")
        result = processor._process_chunk([test_data])
        
        if result:
            print("✅ Mini template processing completed successfully")
            
            # Save the result to examine it
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                result.save(tmp_file.name)
                tmp_path = tmp_file.name
                print(f"💾 Saved result to: {tmp_path}")
            
            # Load the result and check its structure
            result_doc = Document(tmp_path)
            
            if result_doc.tables:
                table = result_doc.tables[0]
                print(f"📊 Result has table with {len(table.rows)} rows x {len(table.columns)} columns")
                
                # Check if it's a 4x5 grid (20 labels)
                if len(table.rows) == 5 and len(table.columns) == 4:
                    print("✅ Result has correct 4x5 grid structure")
                else:
                    print(f"❌ Result has wrong grid structure: {len(table.rows)}x{len(table.columns)}")
                    return False
                
                # Count how many cells actually have visible content
                visible_cells = 0
                total_cells = 0
                
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        total_cells += 1
                        cell_text = cell.text.strip()
                        
                        if cell_text:
                            visible_cells += 1
                            print(f"🔍 Cell [{row_idx}][{col_idx}] (Label{row_idx * 4 + col_idx + 1}): '{cell_text}'")
                            
                            # Check inner tables
                            if cell.tables:
                                print(f"  📋 Has {len(cell.tables)} inner tables")
                                for inner_idx, inner_table in enumerate(cell.tables):
                                    for inner_row_idx, inner_row in enumerate(inner_table.rows):
                                        for inner_cell_idx, inner_cell in enumerate(inner_row.cells):
                                            inner_text = inner_cell.text.strip()
                                            if inner_text:
                                                print(f"    📋 Inner cell [{inner_row_idx}][{inner_cell_idx}]: '{inner_text}'")
                        else:
                            print(f"🔍 Cell [{row_idx}][{col_idx}] (Label{row_idx * 4 + col_idx + 1}): EMPTY")
                
                print(f"\n📊 SUMMARY:")
                print(f"  Total cells: {total_cells}")
                print(f"  Visible cells: {visible_cells}")
                print(f"  Empty cells: {total_cells - visible_cells}")
                
                if visible_cells < 20:
                    print(f"❌ PROBLEM: Only {visible_cells} cells have content, expected 20!")
                    return False
                else:
                    print(f"✅ SUCCESS: All {visible_cells} cells have content!")
                
                # Check for any remaining placeholders in the final document
                all_text = result_doc.element.body.xml
                import re
                remaining_placeholders = re.findall(r'\{\{([^}]+)\}\}', all_text)
                print(f"\n🔍 Remaining placeholders in final document: {len(remaining_placeholders)}")
                if remaining_placeholders:
                    print("  Remaining placeholders:")
                    for placeholder in remaining_placeholders[:10]:  # Show first 10
                        print(f"    {placeholder}")
                    if len(remaining_placeholders) > 10:
                        print(f"    ... and {len(remaining_placeholders) - 10} more")
                    
                    if len(remaining_placeholders) > 0:
                        print(f"❌ PROBLEM: {len(remaining_placeholders)} placeholders remain unfilled!")
                        return False
                else:
                    print("✅ SUCCESS: All placeholders have been filled!")
                
            else:
                print("❌ Result has no tables!")
                return False
            
            print("✅ All 20 labels test completed successfully!")
            return True
            
        else:
            print("❌ Mini template processing failed!")
            return False
        
    except Exception as e:
        print(f"❌ An error occurred during testing: {e}")
        import traceback
        traceback.print_exc()
        return False
    finally:
        if 'tmp_path' in locals() and os.path.exists(tmp_path):
            os.remove(tmp_path)

if __name__ == "__main__":
    success = test_all_20_labels()
    if success:
        print("\n🎉 SUCCESS: All 20 labels are being generated correctly!")
    else:
        print("\n❌ FAILURE: Not all 20 labels are being generated correctly!")

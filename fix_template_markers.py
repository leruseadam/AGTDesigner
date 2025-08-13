#!/usr/bin/env python3
"""
Script to properly add markers around existing placeholders in the double template.
This will preserve the placeholders while adding the markers needed for font sizing.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def fix_template_markers():
    """Add proper markers around existing placeholders in the double template."""
    
    print("Fixing Double Template Markers")
    print("=" * 50)
    
    template_path = "src/core/generation/templates/double.docx"
    
    # Load the current template
    doc = Document(template_path)
    print(f"Loaded template: {template_path}")
    
    if not doc.tables:
        print("ERROR: No tables found in template")
        return
    
    table = doc.tables[0]
    print(f"Table structure: {len(table.rows)} rows x {len(table.columns)} columns")
    
    # Process the first cell (main content cell)
    cell = table.cell(0, 0)
    print("Processing first cell...")
    
    # Clear the cell and rebuild with proper markers
    cell._tc.clear_content()
    
    # Add paragraphs with proper markers and placeholders
    paragraphs = [
        ("LINEAGE_START {{Label1.Lineage}} LINEAGE_END", "Lineage"),
        ("PRODUCTSTRAIN_START {{Label1.ProductStrain}} PRODUCTSTRAIN_END", "ProductStrain"),
        ("", "Empty"),
        ("PRICE_START {{Label1.Price}} PRICE_END", "Price"),
        ("DESC_START {{Label1.Description}} DESC_END", "Description"),
        ("THC_CBD_START {{Label1.Ratio_or_THC_CBD}} THC_CBD_END", "THC_CBD"),
        ("PRODUCTBRAND_START {{Label1.ProductBrand}} PRODUCTBRAND_END", "ProductBrand"),
        ("", "Empty")
    ]
    
    for i, (content, description) in enumerate(paragraphs):
        if content.strip():  # Only add paragraphs with content
            para = cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add the content with proper formatting
            run = para.add_run(content)
            run.font.name = "Arial"
            run.font.size = Pt(12)  # Default size, will be overridden by font sizing system
            
            print(f"  Added paragraph {i}: {description}")
            print(f"    Content: {repr(content)}")
    
    # Save the fixed template
    doc.save(template_path)
    print(f"\n✅ Template markers have been fixed!")
    print(f"  Template saved to: {template_path}")
    print("\nThe template now has:")
    print("  1. LINEAGE_START/END markers (will get proper font sizing)")
    print("  2. PRODUCTSTRAIN_START/END markers (will get 1pt font)")
    print("  3. PRICE_START/END markers (will get proper font sizing)")
    print("  4. DESC_START/END markers (will get proper font sizing)")
    print("  5. THC_CBD_START/END markers (will get proper font sizing)")
    print("  6. PRODUCTBRAND_START/END markers (will get proper font sizing)")
    print("\n🎉 Double template markers have been successfully fixed!")

if __name__ == "__main__":
    fix_template_markers() 
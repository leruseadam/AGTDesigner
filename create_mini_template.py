#!/usr/bin/env python3
"""
Script to create a proper mini template with all necessary placeholders.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_proper_mini_template():
    """Create a proper mini template with all necessary placeholders."""
    
    # Create a new document
    doc = Document()
    
    # Add a table with 1x1 structure (will be expanded to 4x5 later)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Get the cell
    cell = table.cell(0, 0)
    
    # Add all the necessary placeholders in the correct order
    placeholders = [
        "{{Label1.DescAndWeight}}",
        "{{Label1.THC_CBD}}", 
        "{{Label1.ProductBrand}}",
        "{{Label1.Price}}",
        "{{Label1.Lineage}}",
        "{{Label1.ProductStrain}}",
        "{{Label1.ProductVendor}}"
    ]
    
    # Add each placeholder as a separate paragraph
    for placeholder in placeholders:
        paragraph = cell.add_paragraph()
        run = paragraph.add_run(placeholder)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.bold = True
    
    # Save the template
    output_path = "src/core/generation/templates/mini.docx"
    doc.save(output_path)
    print(f"✓ Created proper mini template: {output_path}")
    
    # Verify the template
    verify_doc = Document(output_path)
    print(f"✓ Template verification:")
    print(f"  - Tables: {len(verify_doc.tables)}")
    print(f"  - Table 0: {len(verify_doc.tables[0].rows)}x{len(verify_doc.tables[0].columns)}")
    print(f"  - Cell content:")
    cell_text = verify_doc.tables[0].cell(0, 0).text
    for line in cell_text.split('\n'):
        if line.strip():
            print(f"    {line.strip()}")

if __name__ == "__main__":
    create_proper_mini_template()

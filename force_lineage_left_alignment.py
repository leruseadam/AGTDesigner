#!/usr/bin/env python3
"""
NUCLEAR FIX: Force all lineage cells to LEFT alignment
Run this on any generated document to fix centered lineage bars
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import sys

def force_lineage_left_alignment(docx_path):
    """Force all colored lineage cells to have LEFT alignment"""
    print(f"Opening document: {docx_path}")
    doc = Document(docx_path)
    
    # Define lineage colors (same as in docx_formatting.py)
    LINEAGE_COLORS = {
        'INDICA': '9900FF',
        'SATIVA': 'ED4123', 
        'HYBRID': '009900',
        'HYBRID_INDICA': '9900FF',
        'HYBRID_SATIVA': 'ED4123',
        'CBD': '0099FF',
        'MIXED': '0099FF',
        'PARA': 'FFCC00'
    }
    
    lineage_color_values = set(LINEAGE_COLORS.values())
    cells_fixed = 0
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                try:
                    # Check if cell has a lineage background color
                    tc = cell._tc
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is not None:
                        shd = tcPr.find(qn('w:shd'))
                        if shd is not None:
                            fill_color = shd.get(qn('w:fill'))
                            
                            # If cell has a lineage color, FORCE LEFT alignment
                            if fill_color and fill_color.upper() in lineage_color_values:
                                for paragraph in cell.paragraphs:
                                    if paragraph.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                        cells_fixed += 1
                                        print(f"✅ Fixed alignment for: {cell.text.strip()[:40]}")
                except Exception as e:
                    print(f"⚠️  Error processing cell: {e}")
                    continue
    
    if cells_fixed > 0:
        doc.save(docx_path)
        print(f"\n🎉 DONE! Fixed {cells_fixed} lineage cells to LEFT alignment")
        print(f"📄 Saved: {docx_path}")
    else:
        print("\n⚠️  No lineage cells found to fix")
    
    return cells_fixed

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python force_lineage_left_alignment.py <path_to_docx_file>")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    force_lineage_left_alignment(docx_file)

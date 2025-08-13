#!/usr/bin/env python3
"""
Debug script to test the manual replacement method step by step.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor, get_font_scheme

def debug_manual_replacement():
    """Debug the manual replacement method step by step."""
    
    print("Debug Manual Replacement Method")
    print("=" * 50)
    
    # Test data
    test_records = [
        {
            'ProductName': 'Grape Moonshot',
            'Description': 'Grape Moonshot',
            'WeightUnits': '1.7oz',
            'Price': '$15',
            'DOH': '100mg THC',
            'ProductBrand': 'Test Brand',
            'ProductType': 'edible',
            'Lineage': 'HYBRID',
            'ProductStrain': 'Grape Moonshot'
        }
    ]
    
    try:
        font_scheme = get_font_scheme('mini')
        processor = TemplateProcessor('mini', font_scheme, 1.0)
        
        print(f"Template type: {processor.template_type}")
        
        # Build context for the first record
        label_context = processor._build_label_context(test_records[0], None)
        
        print("\n=== Context for Label1 ===")
        for key, value in label_context.items():
            print(f"{key}: {repr(value)}")
        
        # Create full context
        context = {}
        for i, record in enumerate(test_records):
            label_context = processor._build_label_context(record, None)
            context[f'Label{i+1}'] = label_context
        
        print(f"\n=== Full Context ===")
        for label_key, label_context in context.items():
            print(f"{label_key}: {list(label_context.keys())}")
        
        # Test manual replacement step by step
        print(f"\n=== Testing Manual Replacement Step by Step ===")
        
        # Load the expanded template
        processor.force_re_expand_template()
        processor._expanded_template_buffer.seek(0)
        from docx import Document
        doc = Document(processor._expanded_template_buffer)
        
        # Check first cell before replacement
        first_cell = doc.tables[0].cell(0, 0)
        print(f"First cell before replacement: {repr(first_cell.text)}")
        
        # Check paragraph structure
        print(f"\nParagraphs in first cell: {len(first_cell.paragraphs)}")
        for i, para in enumerate(first_cell.paragraphs):
            print(f"Paragraph {i}: {repr(para.text)}")
            print(f"  Runs in paragraph {i}: {len(para.runs)}")
            for j, run in enumerate(para.runs):
                print(f"    Run {j}: {repr(run.text)}")
        
        # Test the regex pattern on the actual cell text
        import re
        cell_text = first_cell.text
        pattern = r'\{\{Label(\d+)\.DOH\s*\}\}'
        matches = re.findall(pattern, cell_text)
        print(f"\nRegex matches in cell text: {matches}")
        
        # Test the replacement logic manually
        print(f"\n=== Testing Replacement Logic Manually ===")
        
        for label_key, label_context in context.items():
            if isinstance(label_context, dict):
                print(f"Processing {label_key}: {label_context}")
                
                # Check if DOH field exists
                if 'DOH' in label_context:
                    doh_value = label_context['DOH']
                    print(f"  DOH value: {repr(doh_value)}")
                    
                    # Check if _DOH_IMAGE_PATH exists
                    if '_DOH_IMAGE_PATH' in label_context:
                        doh_image_path = label_context['_DOH_IMAGE_PATH']
                        print(f"  _DOH_IMAGE_PATH: {repr(doh_image_path)}")
                        
                        if doh_image_path:
                            print(f"  Would replace with image placeholder")
                        else:
                            print(f"  Would replace with DOH text: {doh_value}")
                    else:
                        print(f"  No _DOH_IMAGE_PATH found")
                else:
                    print(f"  No DOH field found")
        
        # Apply manual replacement
        print(f"\n=== Applying Manual Replacement ===")
        
        # Call the method directly
        doc = processor._manual_replace_placeholders(doc, context)
        
        # Check first cell after replacement
        first_cell = doc.tables[0].cell(0, 0)
        print(f"First cell after replacement: {repr(first_cell.text)}")
        
        # Check if DOH placeholder was replaced
        if '{{Label1.DOH' in first_cell.text:
            print("❌ DOH placeholder still present")
        else:
            print("✅ DOH placeholder replaced")
            
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_manual_replacement() 
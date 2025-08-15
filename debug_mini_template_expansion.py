#!/usr/bin/env python3
"""
Debug script to investigate the mini template expansion issue.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor, get_font_scheme
import logging
import re
from docx import Document
from io import BytesIO

# Set up logging
logging.basicConfig(level=logging.DEBUG)

def debug_mini_template_expansion():
    """Debug the mini template expansion to find where Label100 is coming from."""
    
    try:
        print("=== Debugging Mini Template Expansion ===")
        
        # Create template processor for mini template
        font_scheme = get_font_scheme('mini')
        processor = TemplateProcessor('mini', font_scheme)
        
        print(f"✓ Template processor created successfully")
        print(f"✓ Template type: {processor.template_type}")
        print(f"✓ Chunk size: {processor.chunk_size}")
        print(f"✓ Template path: {processor._template_path}")
        
        # Check the original template
        print("\n=== Original Template Analysis ===")
        with open(processor._template_path, 'rb') as f:
            original_buffer = BytesIO(f.read())
        
        original_doc = Document(original_buffer)
        original_text = original_doc.element.body.xml
        original_matches = re.findall(r'Label(\d+)', original_text)
        print(f"Original template labels: {sorted([int(x) for x in original_matches])}")
        print(f"Original template unique labels: {sorted(set(int(x) for x in original_matches))}")
        
        # Check the expanded template
        print("\n=== Expanded Template Analysis ===")
        expanded_doc = Document(processor._expanded_template_buffer)
        expanded_text = expanded_doc.element.body.xml
        expanded_matches = re.findall(r'Label(\d+)', expanded_text)
        print(f"Expanded template labels: {sorted([int(x) for x in expanded_matches])}")
        print(f"Expanded template unique labels: {sorted(set(int(x) for x in expanded_matches))}")
        
        # Check if there are any unexpected label numbers
        unexpected_labels = [int(x) for x in expanded_matches if int(x) > 20]
        if unexpected_labels:
            print(f"⚠️  WARNING: Found unexpected labels > 20: {unexpected_labels}")
            
            # Find where these labels are coming from
            for label_num in unexpected_labels:
                label_pattern = f'Label{label_num}'
                if label_pattern in expanded_text:
                    print(f"  Found {label_pattern} in expanded template")
                    
                    # Look for context around this label
                    start_pos = expanded_text.find(label_pattern)
                    context_start = max(0, start_pos - 100)
                    context_end = min(len(expanded_text), start_pos + 100)
                    context = expanded_text[context_start:context_end]
                    print(f"  Context: {context}")
        else:
            print("✓ No unexpected labels found")
        
        # Check table structure
        print("\n=== Table Structure Analysis ===")
        if expanded_doc.tables:
            table = expanded_doc.tables[0]
            print(f"Expanded template has {len(table.rows)} rows × {len(table.rows[0].cells)} columns")
            print(f"Expected: 5 rows × 4 columns (20 labels)")
            
            if len(table.rows) == 5 and len(table.rows[0].cells) == 4:
                print("✓ Table dimensions are correct")
            else:
                print("⚠️  Table dimensions are incorrect!")
        else:
            print("⚠️  No tables found in expanded template!")
        
        return True
        
    except Exception as e:
        print(f"✗ Error during debugging: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("Debugging mini template expansion...")
    success = debug_mini_template_expansion()
    
    if success:
        print("\n🔍 Mini template expansion debugging completed!")
    else:
        print("\n❌ Mini template expansion debugging failed!")
        sys.exit(1)

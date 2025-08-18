#!/usr/bin/env python3
"""
Test direct vendor alignment without complex pipeline.
"""

import os
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches

# Add the src directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from src.core.generation.template_processor import TemplateProcessor

def test_direct_vendor_alignment():
    """Test direct vendor alignment."""
    print("🧪 Testing direct vendor alignment...")
    
    # Simple test data
    test_records = [
        {
            'ProductName': 'Test Product',
            'ProductType': 'Flower',
            'Lineage': 'SATIVA',
            'ProductVendor': '1555 Industrial LLC',
            'Price': '$25.00'
        }
    ]
    
    try:
        # Process
        processor = TemplateProcessor('horizontal', {}, 1.0)
        result_doc = processor.process_records(test_records)
        
        # Save
        result_doc.save("test_direct_vendor_result.docx")
        print("✓ Document saved")
        
        # Check the first cell
        if result_doc.tables:
            table = result_doc.tables[0]
            cell = table.cell(0, 0)
            paragraph = cell.paragraphs[0]
            
            print(f"📝 First cell text: '{paragraph.text}'")
            print(f"🔧 Paragraph alignment: {paragraph.alignment}")
            print(f"📏 Left indent: {paragraph.paragraph_format.left_indent}")
            print(f"📌 Tab stops: {len(paragraph.paragraph_format.tab_stops)}")
            
            # Check runs
            print(f"🏃 Runs: {len(paragraph.runs)}")
            for i, run in enumerate(paragraph.runs):
                print(f"   Run {i}: '{run.text}' (font: {run.font.name}, size: {run.font.size}, bold: {run.font.bold})")
            
            # Apply vendor alignment directly
            print("\n🔧 Applying vendor alignment directly...")
            
            # Set paragraph to LEFT alignment
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            print(f"   Set paragraph alignment to LEFT")
            
            # Clear existing tab stops
            paragraph.paragraph_format.tab_stops.clear_all()
            
            # Add right-aligned tab stop
            tab_position = Inches(3.2)
            paragraph.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)
            print(f"   Added tab stop at {tab_position}")
            
            # Find the vendor run and add tab character before it
            vendor_run = None
            vendor_indicators = ['LLC', 'INC', 'CORP', 'CO', 'COMPANY', 'BRANDS', 'BRAND', 'INDUSTRIAL']
            
            for i, run in enumerate(paragraph.runs):
                print(f"   Checking run {i}: '{run.text}'")
                if any(indicator in run.text.upper() for indicator in vendor_indicators):
                    vendor_run = run
                    print(f"   Found vendor run at index {i}: '{run.text}'")
                    break
            
            if vendor_run:
                # Insert tab character before vendor text
                old_text = vendor_run.text
                vendor_run.text = '\t' + vendor_run.text
                print(f"   Added tab character: '{old_text}' -> '{vendor_run.text}'")
                
                # Save the modified document
                result_doc.save("test_direct_vendor_aligned.docx")
                print("   ✓ Modified document saved as test_direct_vendor_aligned.docx")
                
                # Check the final result
                print(f"\n📝 Final text: '{paragraph.text}'")
                print(f"🔧 Final alignment: {paragraph.alignment}")
                print(f"📌 Final tab stops: {len(paragraph.paragraph_format.tab_stops)}")
                
                # Check runs again
                print(f"🏃 Final runs:")
                for i, run in enumerate(paragraph.runs):
                    print(f"   Run {i}: '{run.text}' (length: {len(run.text)})")
                
            else:
                print("   ❌ No vendor run found")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_direct_vendor_alignment()

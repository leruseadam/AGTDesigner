#!/usr/bin/env python3
"""
Debug script to examine the final document content after all processing.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
import tempfile
import shutil
import logging
import re

# Configure logging for debug output
logging.basicConfig(level=logging.DEBUG, format='%(levelname)s: %(message)s')

def debug_final_document():
    """Debug what the final document actually contains."""
    print("🔍 Debugging final document content...")
    
    try:
        # Create a mini template processor
        processor = TemplateProcessor('mini', {}, 1.0)
        
        # Create a simple test context
        test_context = {
            'Label1': {
                'DescAndWeight': 'Test Product - 1g',
                'ProductBrand': 'Test Brand',
                'Price': '$10.00',
                'DOH': 'YES',
                'Ratio_or_THC_CBD': 'THC: 20% CBD: 1%'
            }
        }

        # Process the mini template
        print("🔧 Processing mini template...")
        rendered_doc = processor._process_chunk(chunk=[test_context['Label1']])
        print("✅ Mini template processing completed successfully")

        # Save the rendered document to a temporary file for inspection
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            rendered_doc.save(tmp_file.name)
            temp_path = tmp_file.name
        print(f"💾 Saved rendered document to: {temp_path}")

        # Now load the saved document to see what's actually in it
        print("📖 Loading saved document to examine content...")
        final_doc = Document(temp_path)
        
        print(f"📊 Final document has {len(final_doc.tables)} tables")
        
        if final_doc.tables:
            table = final_doc.tables[0]
            print(f"📊 Main table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Count how many cells actually have visible content
            visible_cells = 0
            total_cells = 0
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    total_cells += 1
                    cell_text = cell.text.strip()
                    
                    if cell_text:
                        visible_cells += 1
                        print(f"🔍 Cell [{row_idx}][{col_idx}] (Label{row_idx * 4 + col_idx + 1}): '{cell_text}'")
                        
                        # Check inner tables
                        if cell.tables:
                            print(f"  📋 Has {len(cell.tables)} inner tables")
                            for inner_idx, inner_table in enumerate(cell.tables):
                                for inner_row_idx, inner_row in enumerate(inner_table.rows):
                                    for inner_cell_idx, inner_cell in enumerate(inner_row.cells):
                                        inner_text = inner_cell.text.strip()
                                        if inner_text:
                                            print(f"    📋 Inner cell [{inner_row_idx}][{inner_cell_idx}]: '{inner_text}'")
                    else:
                        print(f"🔍 Cell [{row_idx}][{col_idx}] (Label{row_idx * 4 + col_idx + 1}): EMPTY")
            
            print(f"\n📊 SUMMARY:")
            print(f"  Total cells: {total_cells}")
            print(f"  Visible cells: {visible_cells}")
            print(f"  Empty cells: {total_cells - visible_cells}")
            
            if visible_cells < 20:
                print(f"❌ PROBLEM: Only {visible_cells} cells have content, expected 20!")
            else:
                print(f"✅ SUCCESS: All {visible_cells} cells have content!")

        # Check for any remaining placeholders in the final document
        all_text = final_doc.element.body.xml
        remaining_placeholders = re.findall(r'\{\{([^}]+)\}\}', all_text)
        print(f"\n🔍 Remaining placeholders in final document: {len(remaining_placeholders)}")
        if remaining_placeholders:
            print("  Remaining placeholders:")
            for placeholder in remaining_placeholders[:10]:  # Show first 10
                print(f"    {placeholder}")
            if len(remaining_placeholders) > 10:
                print(f"    ... and {len(remaining_placeholders) - 10} more")

        print("✅ Final document debug completed successfully!")

    except Exception as e:
        print(f"❌ An error occurred during debugging: {e}")
        import traceback
        traceback.print_exc()
    finally:
        if 'temp_path' in locals() and os.path.exists(temp_path):
            os.remove(temp_path)

if __name__ == "__main__":
    debug_final_document()

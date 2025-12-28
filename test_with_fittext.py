#!/usr/bin/env python3
"""
Test DOCX generation with tcFitText property to see if it causes splitting.
"""

import sys
sys.path.insert(0, '/Users/adamcordova/Desktop/labelMaker_ QR copy final')

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.shared import OxmlElement, qn
from src.core.generation.text_processing import make_nonbreaking_hyphens, make_nonbreaking_weight_units

# Process the text
test_text = "Blueberry - 7g"
processed = make_nonbreaking_hyphens(test_text)
processed = make_nonbreaking_weight_units(processed)

print(f"Original: {repr(test_text)}")
print(f"Processed: {repr(processed)}")

# Create two documents for comparison
def create_test_doc(use_fittext, filename):
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]

    # Set cell width to 2 inches (same as the code)
    cell.width = Inches(2.0)

    # Get the table cell properties
    tcPr = cell._element.get_or_add_tcPr()

    # Set fixed width
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), '2880')  # 2.0 inches in twips
    tcW.set(qn('w:type'), 'dxa')

    if use_fittext:
        # Enable tcFitText (like the actual code does)
        tcFitText = OxmlElement('w:tcFitText')
        tcPr.append(tcFitText)
        tcFitText.set(qn('w:val'), '1')

    # Add the processed text
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(processed)
    run.font.size = Pt(10)

    doc.save(filename)
    print(f"✓ Created: {filename} (tcFitText={'ON' if use_fittext else 'OFF'})")

# Create both versions
create_test_doc(False, '/Users/adamcordova/Desktop/labelMaker_ QR copy final/test_WITHOUT_fittext.docx')
create_test_doc(True, '/Users/adamcordova/Desktop/labelMaker_ QR copy final/test_WITH_fittext.docx')

print("\n" + "="*60)
print("Compare these two files to see if tcFitText causes splitting:")
print("  test_WITHOUT_fittext.docx")
print("  test_WITH_fittext.docx")
print("="*60)

#!/usr/bin/env python3
"""
Comprehensive test for DOH image centering that actually generates documents.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_DOUBLE
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO

def test_doh_centering_actual():
    """Test DOH image centering with actual document generation."""
    print("🧪 Testing DOH Image Centering with Actual Document Generation")
    print("=" * 70)
    
    # Create a test record with DOH image
    test_record = {
        'Description': 'Test Product with DOH',
        'WeightUnits': '1g',
        'ProductBrand': 'Test Brand',
        'Price': '$10.00',
        'Lineage': 'Test Lineage',
        'THC_CBD': 'THC: 20% CBD: 2%',
        'ProductStrain': 'Test Strain',
        'DOH': 'YES',  # This should trigger DOH image
        'Product Type*': 'classic'  # This should use regular DOH image
    }
    
    print("📋 Test Record:")
    for key, value in test_record.items():
        print(f"  {key}: {value}")
    print()
    
    # Test double template specifically
    print("🔄 Processing with double template...")
    processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
    result_doc = processor.process_records([test_record])
    
    if not result_doc:
        print("❌ ERROR: Failed to process test record")
        return False
    
    print("✅ Document generated successfully")
    print()
    
    # Check if DOH images are properly centered
    doh_images_found = 0
    centered_images = 0
    
    print("🔍 Analyzing document for DOH images...")
    
    for table_idx, table in enumerate(result_doc.tables):
        print(f"  Table {table_idx + 1}:")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_has_image = False
                cell_is_centered = False
                
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    for run_idx, run in enumerate(paragraph.runs):
                        # Check if this run contains an image
                        if hasattr(run, '_element'):
                            # Check for drawing elements (InlineImage)
                            drawing = run._element.find(qn('w:drawing'))
                            pict = run._element.find(qn('w:pict'))
                            
                            if drawing is not None or pict is not None:
                                doh_images_found += 1
                                cell_has_image = True
                                print(f"    ✓ Found image in cell [{row_idx},{cell_idx}] paragraph {para_idx} run {run_idx}")
                                
                                # Check if paragraph is centered
                                if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                                    print(f"      ✓ Paragraph is centered")
                                    cell_is_centered = True
                                else:
                                    print(f"      ✗ Paragraph is NOT centered (alignment: {paragraph.alignment})")
                                
                                # Check if cell is vertically centered
                                if cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER:
                                    print(f"      ✓ Cell is vertically centered")
                                else:
                                    print(f"      ✗ Cell is NOT vertically centered (alignment: {cell.vertical_alignment})")
                                
                                # Check paragraph spacing
                                if (paragraph.paragraph_format.space_before == 0 and 
                                    paragraph.paragraph_format.space_after == 0):
                                    print(f"      ✓ Paragraph spacing is minimal")
                                else:
                                    print(f"      ✗ Paragraph spacing is not minimal")
                                
                                # Check XML-level spacing
                                pPr = paragraph._element.get_or_add_pPr()
                                spacing = pPr.find(qn('w:spacing'))
                                if spacing is not None:
                                    before = spacing.get(qn('w:before'))
                                    after = spacing.get(qn('w:after'))
                                    if before == '0' and after == '0':
                                        print(f"      ✓ XML spacing is minimal")
                                    else:
                                        print(f"      ✗ XML spacing is not minimal (before: {before}, after: {after})")
                                else:
                                    print(f"      ✓ No XML spacing element found")
                
                if cell_has_image and cell_is_centered:
                    centered_images += 1
    
    print(f"\n📊 Results:")
    print(f"  DOH images found: {doh_images_found}")
    print(f"  Centered images: {centered_images}")
    
    if doh_images_found > 0:
        if centered_images == doh_images_found:
            print("✅ SUCCESS: All DOH images are properly centered")
            return True
        else:
            print("❌ FAILURE: Some DOH images are not properly centered")
            return False
    else:
        print("❌ FAILURE: No DOH images found in the document")
        print("   This suggests the DOH images are not being inserted properly")
        return False

def test_doh_high_cbd_actual():
    """Test DOH image with High CBD product type."""
    print("\n🧪 Testing DOH Image with High CBD Product Type")
    print("=" * 70)
    
    # Create a test record with High CBD product type
    test_record = {
        'Description': 'Test High CBD Product',
        'WeightUnits': '1g',
        'ProductBrand': 'Test Brand',
        'Price': '$10.00',
        'Lineage': 'Test Lineage',
        'THC_CBD': 'THC: 5% CBD: 15%',
        'ProductStrain': 'Test Strain',
        'DOH': 'YES',  # This should trigger DOH image
        'Product Type*': 'high cbd classic'  # This should use HighCBD image
    }
    
    print("📋 Test Record:")
    for key, value in test_record.items():
        print(f"  {key}: {value}")
    print()
    
    # Test double template specifically
    print("🔄 Processing with double template...")
    processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
    result_doc = processor.process_records([test_record])
    
    if not result_doc:
        print("❌ ERROR: Failed to process test record")
        return False
    
    print("✅ Document generated successfully")
    print()
    
    # Check if HighCBD images are properly centered
    highcbd_images_found = 0
    centered_images = 0
    
    print("🔍 Analyzing document for HighCBD images...")
    
    for table_idx, table in enumerate(result_doc.tables):
        print(f"  Table {table_idx + 1}:")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_has_image = False
                cell_is_centered = False
                
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    for run_idx, run in enumerate(paragraph.runs):
                        # Check if this run contains an image
                        if hasattr(run, '_element'):
                            # Check for drawing elements (InlineImage)
                            drawing = run._element.find(qn('w:drawing'))
                            pict = run._element.find(qn('w:pict'))
                            
                            if drawing is not None or pict is not None:
                                highcbd_images_found += 1
                                cell_has_image = True
                                print(f"    ✓ Found HighCBD image in cell [{row_idx},{cell_idx}] paragraph {para_idx} run {run_idx}")
                                
                                # Check if paragraph is centered
                                if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                                    print(f"      ✓ Paragraph is centered")
                                    cell_is_centered = True
                                else:
                                    print(f"      ✗ Paragraph is NOT centered (alignment: {paragraph.alignment})")
                                
                                # Check if cell is vertically centered
                                if cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER:
                                    print(f"      ✓ Cell is vertically centered")
                                else:
                                    print(f"      ✗ Cell is NOT vertically centered (alignment: {cell.vertical_alignment})")
                
                if cell_has_image and cell_is_centered:
                    centered_images += 1
    
    print(f"\n📊 Results:")
    print(f"  HighCBD images found: {highcbd_images_found}")
    print(f"  Centered images: {centered_images}")
    
    if highcbd_images_found > 0:
        if centered_images == highcbd_images_found:
            print("✅ SUCCESS: All HighCBD images are properly centered")
            return True
        else:
            print("❌ FAILURE: Some HighCBD images are not properly centered")
            return False
    else:
        print("❌ FAILURE: No HighCBD images found in the document")
        print("   This suggests the HighCBD images are not being inserted properly")
        return False

def main():
    """Run all tests."""
    print("🎯 DOH Image Centering Comprehensive Test")
    print("=" * 50)
    print()
    
    # Test regular DOH image centering
    test1_passed = test_doh_centering_actual()
    print()
    
    # Test High CBD DOH image centering
    test2_passed = test_doh_high_cbd_actual()
    print()
    
    # Summary
    print("🎯 TEST SUMMARY")
    print("=" * 20)
    if test1_passed and test2_passed:
        print("✅ ALL TESTS PASSED: DOH image centering fix is working correctly")
        return True
    else:
        print("❌ SOME TESTS FAILED: DOH image centering fix needs more work")
        return False

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1) 
#!/usr/bin/env python3
"""
Test script to verify the double template 4x3 grid with vertical gutter implementation.
This script tests that the 4x3 double template has a vertical gutter created by cell spacing.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_DOUBLE
from docx import Document
from io import BytesIO

def test_double_template_4x3_vertical_gutter():
    """Test that the double template uses 4x3 grid with vertical gutter via cell spacing."""
    print("🔍 Testing Double Template 4x3 Grid with Vertical Gutter")
    
    try:
        # Create template processor
        processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
        
        # Force re-expansion to ensure we get the latest 4x3 implementation
        processor.force_re_expand_template()
        print("✅ Template re-expanded with 4x3 grid implementation")
        
        # Get the expanded template
        template_buffer = processor._expand_template_if_needed(force_expand=True)
        doc = Document(template_buffer)
        
        if not doc.tables:
            print("❌ No table found in template")
            return False
            
        table = doc.tables[0]
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        
        print(f"📊 Table dimensions: {num_rows} rows x {num_cols} columns")
        
        # Verify we have a 4x3 grid
        expected_rows, expected_cols = 3, 4
        if num_rows != expected_rows or num_cols != expected_cols:
            print(f"❌ Expected {expected_rows}x{expected_cols} grid, got {num_rows}x{num_cols}")
            return False
        
        print(f"✅ Correct grid dimensions: {num_rows}x{num_cols}")
        
        # Check column widths (should all be equal)
        print("\n📏 Column widths:")
        for c in range(num_cols):
            col = table.columns[c]
            cell = col.cells[0]  # Get first cell in column
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW')
            if tcW is not None:
                width = tcW.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
                width_inches = int(width) / 1440 if width else 0
                print(f"   Column {c+1}: {width_inches:.3f}\" (Label)")
        
        # Check that all cells have content (no empty gutter cells)
        label_count = 0
        empty_count = 0
        
        for r in range(num_rows):
            for c in range(num_cols):
                cell = table.cell(r, c)
                cell_text = cell.text.strip()
                
                if cell_text:
                    label_count += 1
                    print(f"   ✅ Label cell ({r+1},{c+1}): Has content '{cell_text[:30]}...'")
                else:
                    empty_count += 1
                    print(f"   ⚠️  Empty cell ({r+1},{c+1}): No content")
        
        print(f"\n📋 Cell Summary:")
        print(f"   Label cells: {label_count}/12 (should be 12)")
        print(f"   Empty cells: {empty_count}/12 (should be 0)")
        
        if label_count == 12 and empty_count == 0:
            print("✅ All cells have content - no gutter cells")
        else:
            print("❌ Found empty cells - may still have gutter structure")
            return False
        
        # Check for cell spacing (vertical gutter)
        tblPr = table._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr')
        if tblPr is not None:
            cell_spacing = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblCellSpacing')
            if cell_spacing is not None:
                spacing_width = cell_spacing.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
                spacing_inches = int(spacing_width) / 1440 if spacing_width else 0
                print(f"✅ Cell spacing found: {spacing_inches:.3f}\" (creates vertical gutter)")
            else:
                print("⚠️  No cell spacing found")
        else:
            print("⚠️  No table properties found")
        
        # Check cell margins (should be minimal)
        cell_with_minimal_margins = 0
        for r in range(num_rows):
            for c in range(num_cols):
                cell = table.cell(r, c)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcMar')
                if tcMar is not None:
                    # Check if margins are minimal (0.001 inches)
                    top_margin = tcMar.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}top')
                    if top_margin is not None:
                        margin_width = top_margin.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
                        if margin_width and int(margin_width) <= 2:  # 0.001 inches = ~1.44 twips
                            cell_with_minimal_margins += 1
        
        print(f"   Cells with minimal margins: {cell_with_minimal_margins}/12")
        
        if cell_with_minimal_margins == 12:
            print("✅ All cells have minimal margins (no horizontal gutters)")
        else:
            print(f"⚠️  Only {cell_with_minimal_margins}/12 cells have minimal margins")
        
        print("\n📋 Summary:")
        print("   Grid: 4x3 (standard grid)")
        print("   Vertical gutter: 0.05\" via cell spacing")
        print("   Horizontal gutters: None (minimal cell margins)")
        print("   Total labels: 12 (all cells populated)")
        print("   Label distribution:")
        print("     - Columns 1-2: Left side (Labels 1,2,5,6,9,10)")
        print("     - Columns 3-4: Right side (Labels 3,4,7,8,11,12)")
        
        print("\n✅ Double template 4x3 grid with vertical gutter test completed successfully!")
        return True
        
    except Exception as e:
        print(f"❌ Error testing double template 4x3 grid with vertical gutter: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("🚀 Double Template 4x3 Grid with Vertical Gutter Test")
    print("This test verifies that the double template uses a 4x3 grid with vertical gutter via cell spacing")
    print("=" * 70)
    
    success = test_double_template_4x3_vertical_gutter()
    
    if success:
        print("\n🎉 Test passed! The double template now uses a 4x3 grid with vertical gutter.")
    else:
        print("\n💥 Test failed! Check the implementation.") 
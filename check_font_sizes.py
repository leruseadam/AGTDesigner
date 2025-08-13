#!/usr/bin/env python3
"""
Script to check font sizes in the generated document.
"""

import docx

def check_font_sizes():
    """Check font sizes in the generated document."""
    
    print("=== CHECKING FONT SIZES IN GENERATED DOCUMENT ===")
    
    try:
        doc = docx.Document('test_field_order_output.docx')
        print(f"Document has {len(doc.tables)} tables")
        
        if doc.tables:
            table = doc.tables[0]
            print(f"First table: {len(table.rows)}x{len(table.columns)}")
            
            cell = table.cell(0, 0)
            print('First cell content:')
            for i, para in enumerate(cell.paragraphs):
                print(f'  Paragraph {i}: {repr(para.text)}')
                
                # Check font sizes for each run
                if para.runs:
                    for j, run in enumerate(para.runs):
                        font_size = "No font size set"
                        if run.font.size:
                            font_size = f"{run.font.size.pt}pt"
                        print(f'    Run {j}: {repr(run.text)} - Font size: {font_size}')
                else:
                    print(f'    No runs in paragraph {i}')
                    
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    check_font_sizes() 
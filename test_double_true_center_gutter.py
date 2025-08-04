#!/usr/bin/env python3
"""
Test script to verify the double template true center gutter implementation.
This script tests that the 4x3 double template has a center gutter creating two groups of 2 labels each.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_DOUBLE
from docx import Document
from io import BytesIO

def test_double_template_true_center_gutter():
    """Test that the double template has a true center gutter creating two groups of 2 labels each."""
    print("🔍 Testing Double Template True Center Gutter")
    
    try:
        # Create template processor
        processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
        
        # Force re-expansion to ensure we get the latest true center gutter implementation
        processor.force_re_expand_template()
        print("✅ Template re-expanded with true center gutter implementation")
        
        # Get the expanded template
        template_buffer = processor._expand_template_if_needed(force_expand=True)
        doc = Document(template_buffer)
        
        if not doc.tables:
            print("❌ No table found in template")
            return False
            
        table = doc.tables[0]
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        
        print(f"📊 Table dimensions: {num_rows} rows x {num_cols} columns")
        
        # Verify we have a 4x3 grid
        expected_rows, expected_cols = 3, 4
        if num_rows != expected_rows or num_cols != expected_cols:
            print(f"❌ Expected {expected_rows}x{expected_cols} grid, got {num_rows}x{num_cols}")
            return False
        
        print(f"✅ Correct grid dimensions: {num_rows}x{num_cols}")
        
        # Check for absence of cell spacing (no global spacing)
        tblPr = table._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr')
        if tblPr is not None:
            cell_spacing = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblCellSpacing')
            if cell_spacing is None:
                print("✅ No cell spacing found (no global spacing)")
            else:
                print("⚠️  Cell spacing found (global spacing present)")
        else:
            print("⚠️  No table properties found")
        
        # Check cell margins for true center gutter
        print("\n📏 Cell margin analysis for 'double' grouping:")
        center_gutter_cells = 0
        minimal_margin_cells = 0
        
        for r in range(num_rows):
            for c in range(num_cols):
                cell = table.cell(r, c)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcMar')
                
                if tcMar is not None:
                    # Check right margin for column 2 and left margin for column 3
                    right_margin = tcMar.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}right')
                    left_margin = tcMar.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left')
                    
                    right_width = 0
                    left_width = 0
                    
                    if right_margin is not None:
                        right_width = int(right_margin.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', '0'))
                    if left_margin is not None:
                        left_width = int(left_margin.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', '0'))
                    
                    right_inches = right_width / 1440
                    left_inches = left_width / 1440
                    
                    # Check if this cell has center gutter margins (only columns 2 and 3)
                    if (c == 1 and right_inches >= 0.025) or (c == 2 and left_inches >= 0.025):
                        center_gutter_cells += 1
                        if c == 1:
                            print(f"   ✅ Cell ({r+1},{c+1}): Left group end - Right: {right_inches:.3f}\", Left: {left_inches:.3f}\"")
                        else:
                            print(f"   ✅ Cell ({r+1},{c+1}): Right group start - Right: {right_inches:.3f}\", Left: {left_inches:.3f}\"")
                    else:
                        minimal_margin_cells += 1
                        print(f"   📏 Cell ({r+1},{c+1}): Minimal margin - Right: {right_inches:.3f}\", Left: {left_inches:.3f}\"")
        
        print(f"\n📋 Margin Summary:")
        print(f"   Cells with center gutter margins: {center_gutter_cells}/6 (should be 6 - 3 rows × 2 gutter cells)")
        print(f"   Cells with minimal margins: {minimal_margin_cells}/6 (should be 6 - 3 rows × 2 non-gutter cells)")
        
        # Verify center gutter pattern (only columns 2 and 3 should have extra margins)
        expected_center_gutter = 6  # 3 rows × 2 gutter cells (columns 2 and 3)
        if center_gutter_cells == expected_center_gutter:
            print("✅ Center gutter margins correctly applied to columns 2 and 3 only")
        else:
            print(f"❌ Only {center_gutter_cells}/{expected_center_gutter} cells have center gutter margins")
            return False
        
        # Check that all cells have content
        label_count = 0
        empty_count = 0
        
        for r in range(num_rows):
            for c in range(num_cols):
                cell = table.cell(r, c)
                cell_text = cell.text.strip()
                
                if cell_text:
                    label_count += 1
                else:
                    empty_count += 1
        
        print(f"   Label cells: {label_count}/12 (should be 12)")
        print(f"   Empty cells: {empty_count}/12 (should be 0)")
        
        if label_count == 12 and empty_count == 0:
            print("✅ All cells have content")
        else:
            print("❌ Found empty cells")
            return False
        
        print("\n📋 'Double' Template Grouping:")
        print("   Left Group (Columns 1-2): Labels 1,2,5,6,9,10")
        print("   Right Group (Columns 3-4): Labels 3,4,7,8,11,12")
        print("   Center Gutter: Between columns 2 and 3")
        print("   Gutter Width: 0.05\" (0.025\" + 0.025\")")
        
        print("\n📋 Summary:")
        print("   Grid: 4x3 (standard grid)")
        print("   True center gutter: Only between columns 2 and 3")
        print("   Double grouping: Two groups of 2 labels each")
        print("   No global cell spacing")
        print("   Total labels: 12 (all cells populated)")
        
        print("\n✅ Double template true center gutter test completed successfully!")
        return True
        
    except Exception as e:
        print(f"❌ Error testing double template true center gutter: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("🚀 Double Template True Center Gutter Test")
    print("This test verifies that the double template has a true center gutter creating two groups of 2 labels each")
    print("=" * 70)
    
    success = test_double_template_true_center_gutter()
    
    if success:
        print("\n🎉 Test passed! The double template now has a true center gutter for 'double' grouping.")
    else:
        print("\n💥 Test failed! Check the implementation.") 
#!/usr/bin/env python3
"""
Debug script to examine the horizontal template expansion and identify issues.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor, get_font_scheme

def debug_horizontal_template():
    """Debug the horizontal template expansion to see what's happening."""
    print("🔍 Debugging horizontal template expansion...")
    
    try:
        # Create a template processor for horizontal template
        processor = TemplateProcessor('horizontal', get_font_scheme('horizontal'), 1.0)
        
        # Check if the expanded template buffer exists
        if hasattr(processor, '_expanded_template_buffer') and processor._expanded_template_buffer:
            print("✅ Template expansion completed successfully")
            
            # Load the expanded template to examine it
            from docx import Document
            from io import BytesIO
            
            # Create a copy of the buffer
            buffer_copy = BytesIO(processor._expanded_template_buffer.getvalue())
            doc = Document(buffer_copy)
            
            if doc.tables:
                table = doc.tables[0]
                print(f"✅ Expanded template has {len(table.rows)} rows x {len(table.columns)} columns")
                
                # Examine each cell in detail
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        print(f"\n📋 Cell ({row_idx}, {col_idx}):")
                        
                        # Get cell text from paragraphs
                        cell_text = ''
                        for paragraph in cell.paragraphs:
                            cell_text += paragraph.text
                        print(f"  📝 Cell text: {repr(cell_text)}")
                        
                        # Examine the actual XML structure
                        from docx.oxml.ns import qn
                        print(f"  🔍 XML text elements:")
                        text_elements = []
                        for t in cell._tc.iter(qn('w:t')):
                            if t.text:
                                text_elements.append(t.text)
                                print(f"    - {repr(t.text)}")
                        
                        # Check for specific placeholders
                        placeholders = {
                            'Lineage': 0,
                            'ProductVendor': 0,
                            'ProductStrain': 0,
                            'DescAndWeight': 0,
                            'Price': 0,
                            'DOH': 0,
                            'Ratio_or_THC_CBD': 0,
                            'ProductBrand': 0
                        }
                        
                        for text in text_elements:
                            for placeholder in placeholders:
                                if placeholder in text:
                                    placeholders[placeholder] += 1
                        
                        print(f"  📊 Placeholder counts:")
                        for placeholder, count in placeholders.items():
                            status = "✅" if count > 0 else "❌"
                            print(f"    {status} {placeholder}: {count}")
                        
                        # Check for any obvious issues
                        if len(text_elements) < 5:
                            print(f"  ⚠️  Warning: Cell has only {len(text_elements)} text elements (expected more)")
                        
                return True
            else:
                print("❌ No tables found in expanded template")
                return False
        else:
            print("❌ Template expansion failed - no expanded buffer")
            return False
            
    except Exception as e:
        print(f"❌ Error debugging template: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("Debugging horizontal template expansion...")
    success = debug_horizontal_template()
    
    if success:
        print("\n✅ Debug completed successfully!")
    else:
        print("\n❌ Debug failed!")
    
    sys.exit(0 if success else 1)

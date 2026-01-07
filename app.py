# AGT Label Maker - Consolidated Web Application
# ============================================
# This is the sole, consolidated web version of the AGT Label Maker application.
# All web deployment functionality has been consolidated into this single file.
#
# Features:
# - Complete Flask web interface
# - 100% database-derived product matching
# - JointRatio support for pre-roll products
# - Advanced DOCX label generation
# - Real-time Excel processing
# - Session management and caching

from src.core.data.field_mapping import get_canonical_field
import os
import sys  # Add this import
import logging
import traceback
import threading
import signal  # Add signal import for timeout handling
import pandas as pd  # Add this import
import time
import re
import json
from decimal import Decimal
try:
    import numpy as np
    NUMPY_AVAILABLE = True
except Exception:
    np = None
    NUMPY_AVAILABLE = False
try:
    import requests  # Optional dependency for internal HTTP calls
    REQUESTS_AVAILABLE = True
except ImportError:  # pragma: no cover
    requests = None
    REQUESTS_AVAILABLE = False
from src.core.generation.fast_generation import (
    FastGenerationEngine,
    optimize_records_for_generation,
    update_generation_stats
)
# CRITICAL FIX: Make preroll imports optional to prevent startup errors if files don't exist
try:
    from src.core.generation.preroll_tag_generator import generate_preroll_tags, identify_preroll_product_group
except ImportError as preroll_import_error:
    logging.warning(f"Could not import preroll_tag_generator: {preroll_import_error}")
    # Define fallback functions
    def generate_preroll_tags(records, cache):
        logging.warning("generate_preroll_tags called but module not available - returning records unchanged")
        return records
    
    def identify_preroll_product_group(description: str, product_name: str = '') -> dict:
        logging.warning("identify_preroll_product_group called but module not available - returning default group")
        return {'group_id': 'default', 'display_name': 'Preroll Items', 'category': 'Prerolls'}

try:
    from src.core.generation.preroll_product_list import generate_preroll_product_list
except ImportError as preroll_list_error:
    logging.warning(f"Could not import preroll_product_list: {preroll_list_error}")
    # Define fallback function - matches actual signature: (records, cache) -> Optional[Document]
    def generate_preroll_product_list(records, cache):
        logging.warning("generate_preroll_product_list called but module not available - returning None")
        return None

# Performance optimizations - Import response caching utilities
# These decorators provide response caching, compression, and cache invalidation
try:
    from src.core.utils.response_cache import cached_route, compress_response, invalidate_cache_on_change
    RESPONSE_CACHE_AVAILABLE = True
    logging.debug("Response cache utilities loaded successfully")
except ImportError as e:
    # Fallback: Create no-op decorators if cache module is unavailable
    # This ensures the app continues to work even without caching
    RESPONSE_CACHE_AVAILABLE = False
    logging.warning(f"Response cache module not available: {e}. Using no-op decorators.")
    
    def cached_route(*args, **kwargs):
        """No-op decorator when response caching is unavailable"""
        def decorator(f):
            return f  # Return function unchanged (no caching)
        return decorator
    
    def compress_response(response):
        """No-op function when compression is unavailable"""
        return response  # Return response unchanged (no compression)
    
    def invalidate_cache_on_change(*args, **kwargs):
        """No-op decorator when cache invalidation is unavailable"""
        def decorator(f):
            return f  # Return function unchanged (no invalidation)
        return decorator
# Startup Performance Optimization
# DISABLE_STARTUP_FILE_LOADING = True  # Disable startup file loading to prevent hangs

# PythonAnywhere Performance Optimization
PYTHONANYWHERE_OPTIMIZATION = os.environ.get('PYTHONANYWHERE_DOMAIN') is not None
if PYTHONANYWHERE_OPTIMIZATION:
    # Reduce memory usage on PythonAnywhere
    MAX_CONTENT_LENGTH = 5 * 1024 * 1024  # 5MB instead of 10MB for memory efficiency
    SEND_FILE_MAX_AGE_DEFAULT = 180  # 3 minutes instead of 5 minutes
    PERMANENT_SESSION_LIFETIME = 900  # 15 minutes instead of 30 minutes
    
    # Memory optimization settings (allow environment overrides for hosted deployments)
    MAX_MEMORY_MB = int(os.environ.get('MAX_MEMORY_MB', '425'))  # Allow up to ~425MB unless overridden
    CACHE_SIZE_LIMIT = int(os.environ.get('CACHE_SIZE_LIMIT', '50'))  # Reduced cache size
    BATCH_SIZE_LIMIT = int(os.environ.get('BATCH_SIZE_LIMIT', '250'))  # Smaller batch sizes
else:
    # Local development settings
    MAX_CONTENT_LENGTH = 25 * 1024 * 1024  # 25MB for local development
    SEND_FILE_MAX_AGE_DEFAULT = 300
    PERMANENT_SESSION_LIFETIME = 1800
    
    # More generous memory settings for local (still overridable)
    MAX_MEMORY_MB = int(os.environ.get('MAX_MEMORY_MB', '500'))
    CACHE_SIZE_LIMIT = int(os.environ.get('CACHE_SIZE_LIMIT', '100'))
    BATCH_SIZE_LIMIT = int(os.environ.get('BATCH_SIZE_LIMIT', '500'))

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOADS_DIR = os.path.join(BASE_DIR, 'uploads')
VALID_STORES = ['AGT_Bothell', 'AGT_Burien', 'AGT_Goldbar', 'AGT_Lynnwood', 'AGT_Seattle', 'AGT_Shoreline', 'AGT_Walla_Walla', 'Test']
CACHE_DIR = os.path.join(UPLOADS_DIR, 'cache')
os.makedirs(CACHE_DIR, exist_ok=True)

# Memory monitoring and optimization
def get_memory_usage():
    """Get current memory usage in MB."""
    try:
        import psutil
        process = psutil.Process()
        return process.memory_info().rss / (1024 * 1024)
    except ImportError:
        try:
            import resource
            return resource.getrusage(resource.RUSAGE_SELF).ru_maxrss / 1024
        except:
            return 0

def cleanup_memory():
    """Cleanup memory by clearing caches and forcing garbage collection."""
    try:
        import gc
        gc.collect()
        
        # Clear Flask cache if available
        if hasattr(app, 'cache') and app.cache:
            app.cache.clear()
            
        return True
    except Exception as e:
        logging.error(f"Memory cleanup failed: {e}")
        return False

def check_memory_limit():
    """Check if memory usage is within limits."""
    memory_mb = get_memory_usage()
    if memory_mb > MAX_MEMORY_MB:
        logging.warning(f"Memory usage high: {memory_mb:.1f}MB (limit: {MAX_MEMORY_MB}MB)")
        cleanup_memory()
        return False
    return True

def timeout_handler(signum, frame):
    raise TimeoutError("File operation timed out")

def safe_load_file_with_timeout(processor, file_path, timeout_seconds=30):
    """Load file with timeout protection (gracefully degrades when signals unavailable)."""
    def _load_without_timeout():
        try:
            return processor.load_file(file_path)
        except Exception as e:
            logging.error(f"Error loading file without timeout: {e}")
            return False
    
    # Signal-based alarms only work in the main thread on Unix. If unavailable, just load.
    try:
        if threading.current_thread() is not threading.main_thread():
            logging.debug("safe_load_file_with_timeout: non-main thread; skipping SIGALRM timeout")
            return _load_without_timeout()
        
        signal.signal(signal.SIGALRM, timeout_handler)
        signal.alarm(timeout_seconds)
        try:
            result = processor.load_file(file_path)
            signal.alarm(0)
            return result
        except TimeoutError:
            logging.error(f"File loading timed out after {timeout_seconds} seconds")
            return False
        except Exception as e:
            logging.error(f"Error in safe file loading: {e}")
            return False
        finally:
            signal.alarm(0)
    except ValueError as ve:
        # Raised when signal.signal or alarm not permitted (e.g., worker threads on hosting)
        logging.warning(f"SIGALRM unavailable on this thread/environment ({ve}); loading file without timeout")
        return _load_without_timeout()
LAZY_LOADING_ENABLED = True  # Enable lazy loading for better performance

# Browser-based store persistence (handled by frontend JavaScript)

def json_safe_value(value):
    """Convert values to JSON-serializable representations."""
    if value is None:
        return None
    if isinstance(value, Decimal):
        try:
            if value == value.to_integral_value():
                return int(value)
            return float(value)
        except Exception:
            return float(value)
    if isinstance(value, (datetime, dt_date)):
        return value.isoformat()
    if isinstance(value, dt_time):
        return value.isoformat()
    if isinstance(value, bytes):
        try:
            return value.decode('utf-8')
        except UnicodeDecodeError:
            return value.decode('latin-1', errors='ignore')
    if NUMPY_AVAILABLE and np is not None:
        if isinstance(value, (np.integer,)):
            return int(value)
        if isinstance(value, (np.floating,)):
            return float(value)
        if isinstance(value, (np.bool_,)):
            return bool(value)
        if hasattr(value, 'tolist'):
            try:
                return json_safe_value(value.tolist())
            except Exception:
                pass
    if isinstance(value, pd.Timestamp):
        return value.isoformat()
    if isinstance(value, dict):
        return {k: json_safe_value(v) for k, v in value.items()}
    if isinstance(value, (list, tuple, set)):
        return [json_safe_value(v) for v in value]
    if hasattr(value, 'item') and callable(getattr(value, 'item')):
        try:
            return json_safe_value(value.item())
        except Exception:
            pass
    return value


def make_json_safe(obj):
    """Recursively ensure complex structures are JSON serializable."""
    if isinstance(obj, dict):
        return {k: make_json_safe(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [make_json_safe(item) for item in obj]
    return json_safe_value(obj)


def _normalize_store_key(store_name):
    return store_name or 'global'


def get_available_tags_cache_path(store_name):
    normalized = _normalize_store_key(store_name)
    filename = f'available_tags_{normalized}.json'
    return os.path.join(CACHE_DIR, filename)


def save_available_tags_cache(store_name, tags):
    """Persist the latest successful available-tags payload for emergency fallbacks."""
    try:
        cache_path = get_available_tags_cache_path(store_name)
        with open(cache_path, 'w', encoding='utf-8') as cache_file:
            json.dump(tags, cache_file, ensure_ascii=False)
    except Exception as e:
        logging.warning(f"Failed to save available tags cache for {store_name}: {e}")


def load_available_tags_cache(store_name):
    """Load the most recent cached available-tags payload for a store."""
    try:
        cache_path = get_available_tags_cache_path(store_name)
        if not os.path.exists(cache_path):
            return None
        with open(cache_path, 'r', encoding='utf-8') as cache_file:
            return json.load(cache_file)
    except Exception as e:
        logging.warning(f"Failed to load available tags cache for {store_name}: {e}")
        return None


from pathlib import Path
from werkzeug.utils import secure_filename
import sqlite3
from contextlib import contextmanager

# CRITICAL FIX: Helper function for creating properly configured database connections
# This ensures all database connections have proper timeout, WAL mode, and busy_timeout settings
def create_db_connection(db_path, timeout=30.0, check_same_thread=False):
    """
    Create a properly configured SQLite database connection with:
    - Timeout handling
    - WAL mode for concurrent access
    - Busy timeout for handling locks gracefully
    - Performance optimizations
    """
    max_retries = 5
    retry_delay = 0.5
    
    for attempt in range(max_retries):
        try:
            conn = sqlite3.connect(
                db_path,
                timeout=timeout,
                check_same_thread=check_same_thread
            )
            
            # CRITICAL: Enable WAL mode immediately for better concurrent access
            conn.execute("PRAGMA journal_mode = WAL")
            
            # CRITICAL: Set busy_timeout to handle locked database gracefully (30 seconds)
            conn.execute("PRAGMA busy_timeout = 30000")
            
            # Performance optimizations
            conn.execute("PRAGMA synchronous = NORMAL")
            conn.execute("PRAGMA cache_size = -64000")  # 64MB cache
            conn.execute("PRAGMA temp_store = MEMORY")
            
            return conn
            
        except sqlite3.OperationalError as e:
            error_str = str(e).lower()
            if "database is locked" in error_str and attempt < max_retries - 1:
                wait_time = retry_delay * (2 ** attempt)  # Exponential backoff
                logging.warning(f"Database locked during connection (attempt {attempt + 1}/{max_retries}), retrying in {wait_time}s...")
                time.sleep(wait_time)
                continue
            else:
                logging.error(f"Failed to create database connection after {attempt + 1} attempts: {e}")
                raise
        except Exception as e:
            logging.error(f"Unexpected error creating database connection: {e}")
            raise

@contextmanager
def db_connection(db_path, timeout=30.0, check_same_thread=False):
    """
    Context manager for properly configured database connections.
    Automatically handles connection cleanup.
    """
    conn = None
    try:
        conn = create_db_connection(db_path, timeout, check_same_thread)
        yield conn
        conn.commit()
    except Exception as e:
        if conn:
            conn.rollback()
        raise
    finally:
        if conn:
            conn.close()

# Performance optimizations
IS_PYTHONANYWHERE = 'pythonanywhere.com' in os.environ.get('HTTP_HOST', '')
IS_PRODUCTION = os.environ.get('FLASK_ENV') == 'production' or IS_PYTHONANYWHERE

# OPTIMIZATION: Disable startup file loading for faster app startup
# Honour environment override so PythonAnywhere can skip the heavy Excel scan
# Default remains False for full-featured local runs unless explicitly disabled
_disable_startup_env = os.environ.get('DISABLE_STARTUP_FILE_LOADING')
if _disable_startup_env is not None:
    DISABLE_STARTUP_FILE_LOADING = _disable_startup_env.strip().lower() in ('1', 'true', 'yes')
else:
    # Default to disabling on resource-constrained hosts (PythonAnywhere), otherwise keep enabled
    DISABLE_STARTUP_FILE_LOADING = PYTHONANYWHERE_OPTIMIZATION

# OPTIMIZATION: Enable lazy loading for faster app startup
# Set to False to load files immediately
LAZY_LOADING_ENABLED = True  # Enabled - lazy load components for faster startup

# Use consistent settings for both local and production to ensure identical generation
CHUNK_SIZE_LIMIT = 50
MAX_PROCESSING_TIME_PER_CHUNK = 30
MAX_TOTAL_PROCESSING_TIME = 300
MAX_CONTENT_LENGTH = 100 * 1024 * 1024  # 100MB max file size
UPLOAD_CHUNK_SIZE = 16384  # 16KB chunks for uploads

if IS_PRODUCTION:
    # Production optimizations (logging only)
    logging.getLogger().setLevel(logging.ERROR)
    os.environ['FLASK_ENV'] = 'production'
else:
    # Development optimizations - reduce logging for faster startup
    logging.getLogger().setLevel(logging.ERROR)

# Suppress verbose logging from libraries for faster startup
logging.getLogger('werkzeug').setLevel(logging.ERROR)
logging.getLogger('urllib3').setLevel(logging.ERROR)
logging.getLogger('requests').setLevel(logging.ERROR)
logging.getLogger('pandas').setLevel(logging.ERROR)
logging.getLogger('openpyxl').setLevel(logging.ERROR)
logging.getLogger('xlrd').setLevel(logging.ERROR)
from flask import (
    Flask, 
    request, 
    jsonify, 
    send_file, 
    render_template,
    session,  # Add this
    send_from_directory,
    current_app,
    g,  # Add this for per-request globals
    make_response  # Add for web-optimized endpoints
)
from flask_cors import CORS
try:
    from flask_compress import Compress
except Exception:  # pragma: no cover
    Compress = None

# PythonAnywhere-specific configuration
try:
    from pythonanywhere_config import (
        JELLYFISH_AVAILABLE, LEVENSHTEIN_AVAILABLE, PSUTIL_AVAILABLE,
        jaro_winkler_similarity_fallback, levenshtein_distance_fallback,
        get_memory_usage_fallback, get_pythonanywhere_config, log_missing_dependencies
    )
    PYTHONANYWHERE_CONFIG = get_pythonanywhere_config()
    log_missing_dependencies()
except ImportError:
    JELLYFISH_AVAILABLE = True
    LEVENSHTEIN_AVAILABLE = True
    PSUTIL_AVAILABLE = True
    PYTHONANYWHERE_CONFIG = {}

# Performance optimizations
try:
    from performance_optimizations import (  # type: ignore[import]
        cached, performance_monitor, optimize_dataframe, 
        async_processor, clear_cache, log_performance_stats
    )
    PERFORMANCE_ENABLED = True
    logging.info("Performance optimizations enabled")
except ImportError:
    PERFORMANCE_ENABLED = False
    logging.warning("Performance optimizations not available")

# Simple in-memory cache for PythonAnywhere
if IS_PYTHONANYWHERE:
    from functools import lru_cache
    # Cache frequently used functions
    @lru_cache(maxsize=128)
    def cached_get_font_scheme(template_type, base_size=12):
        from src.core.generation.template_processor import get_font_scheme
        return get_font_scheme(template_type, base_size)
    
    @lru_cache(maxsize=64)
    def cached_calculate_text_complexity(text):
        from src.core.utils.common import calculate_text_complexity
        return calculate_text_complexity(text)
try:
    from flask_session import Session
except ImportError:
    Session = None
from docx import Document
from docxtpl import DocxTemplate, InlineImage
from io import BytesIO
from datetime import datetime, timezone, timedelta, date as dt_date, time as dt_time
from functools import lru_cache
import json  # Add this import
from copy import deepcopy
from docx.shared import Pt, RGBColor, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Add this import
# import pprint  # Removed unused import
import re
import traceback
from docxcompose.composer import Composer
from openpyxl import load_workbook
from PIL import Image as PILImage
import copy
from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ROW_HEIGHT_RULE
from typing import List, Optional, Tuple
from src.core.generation.template_processor import get_font_scheme, TemplateProcessor
from src.core.generation.tag_generator import get_template_path
import time
# Removed unused mini font sizing imports
from src.core.data.excel_processor import ExcelProcessor, get_default_upload_file
from src.core.data.json_matcher import map_inventory_type_to_product_type
import random
# Optional import for flask_caching
# Import optimized upload handler
# from optimized_excel_upload import create_optimized_upload_routes  # Disabled - module not found
try:
    from fast_excel_upload_fix import create_fast_upload_routes  # type: ignore[import]
    FAST_UPLOAD_AVAILABLE = True
except Exception as e:
    logging.warning(f"Fast upload routes not available: {e}")
    create_fast_upload_routes = None
    FAST_UPLOAD_AVAILABLE = False

try:
    from fast_docx_generator import create_fast_docx_routes  # type: ignore[import]
    FAST_DOCX_AVAILABLE = True
except Exception as e:
    logging.warning(f"Fast DOCX routes not available: {e}")
    create_fast_docx_routes = None
    FAST_DOCX_AVAILABLE = False
try:
    from flask_caching import Cache
    CACHE_AVAILABLE = True
except ImportError:
    CACHE_AVAILABLE = False
    # Create a dummy Cache class for fallback
    class Cache:
        def __init__(self, *args, **kwargs):
            pass
        def __getattr__(self, name):
            return lambda *args, **kwargs: None

# Import performance optimization modules
try:
    # Try lightweight performance optimizations first (for PythonAnywhere)
    from src.core.utils.lightweight_performance import (
        get_lightweight_stats, clear_lightweight_cache, 
        lightweight_cached, performance_timer_lightweight
    )
    LIGHTWEIGHT_PERFORMANCE_AVAILABLE = True
    logging.info("Lightweight performance optimizations loaded")
except ImportError as e:
    LIGHTWEIGHT_PERFORMANCE_AVAILABLE = False
    logging.warning(f"Lightweight performance optimizations not available: {e}")

try:
    # Try full performance optimizations (for local development)
    from src.core.utils.performance_cache import cache_manager, get_cache_stats
    from src.core.utils.performance_monitor import get_performance_monitor, start_performance_monitoring
    from src.core.utils.lazy_loader import preload_all, get_lazy_stats
    from src.core.data.ultra_fast_database import get_ultra_fast_database
    from src.core.data.ultra_fast_excel_processor import get_ultra_fast_processor
    FULL_PERFORMANCE_OPTIMIZATIONS_AVAILABLE = True
    logging.info("Full performance optimizations loaded")
except ImportError as e:
    FULL_PERFORMANCE_OPTIMIZATIONS_AVAILABLE = False
    logging.warning(f"Full performance optimizations not available: {e}")

# Use lightweight optimizations if full ones aren't available
PERFORMANCE_OPTIMIZATIONS_AVAILABLE = FULL_PERFORMANCE_OPTIMIZATIONS_AVAILABLE or LIGHTWEIGHT_PERFORMANCE_AVAILABLE
import hashlib
import glob
import subprocess
from collections import defaultdict
import shutil
import pickle
import uuid

current_dir = os.path.dirname(os.path.abspath(__file__))

# Global variables for lazy loading
_initial_data_cache = None
_cache_timestamp = None
CACHE_DURATION = 300  # Cache for 5 minutes

# Global ExcelProcessor instance
_excel_processor = None
_excel_processor_reset_flag = False  # Flag to track when processor has been explicitly reset

# Unique identifier for this Flask server instance (changes on restart)
SERVER_INSTANCE_ID = str(uuid.uuid4())

# Global ProductDatabase instance
_product_database = None

# Global JSONMatcher instance
_json_matcher = None

# Global Enhanced AI Matcher instance
_enhanced_ai_matcher = None

# Global processing status with better state management
processing_status = {}  # filename -> status
processing_timestamps = {}  # filename -> timestamp
processing_lock = threading.Lock()  # Add thread lock for status updates

# Thread lock for ExcelProcessor initialization
excel_processor_lock = threading.Lock()  # Add thread lock for ExcelProcessor initialization

# CRITICAL FIX: Lock for lineage updates to prevent concurrent database conflicts
# CRITICAL FIX: Use a queue-based system for lineage updates to handle multiple rapid requests
# This ensures requests are processed sequentially without timing out while waiting for the lock
import queue
lineage_update_lock = threading.Lock()  # Serialize lineage updates to prevent database lock conflicts
lineage_update_queue = queue.Queue()  # Queue for lineage update requests
_lineage_processor_thread = None  # Background thread for processing lineage updates

def _process_lineage_update_queue():
    """Background thread to process lineage updates from the queue sequentially."""
    global lineage_update_lock
    while True:
        try:
            # Get request from queue (blocks until one is available)
            request_data = lineage_update_queue.get(timeout=1.0)
            if request_data is None:  # Shutdown signal
                break
            
            # Unpack request data
            event, tag_name, new_lineage, store_name, callback = request_data
            
            try:
                # Acquire lock and process update
                with lineage_update_lock:
                    result = callback()  # Execute the actual update logic
                    event.set()  # Signal completion
                    if not result:
                        logging.warning(f"⚠️  Queued lineage update returned False: '{tag_name}' -> '{new_lineage}'")
            except Exception as e:
                logging.error(f"❌ Error processing queued lineage update: {e}")
                event.set()  # Signal completion even on error
            finally:
                lineage_update_queue.task_done()
        except queue.Empty:
            continue
        except Exception as e:
            logging.error(f"❌ Error in lineage update queue processor: {e}")
            import traceback
            logging.error(traceback.format_exc())

# Start background thread for processing lineage updates
def _ensure_lineage_processor_thread():
    """Ensure the lineage update processor thread is running."""
    global _lineage_processor_thread
    if _lineage_processor_thread is None or not _lineage_processor_thread.is_alive():
        _lineage_processor_thread = threading.Thread(target=_process_lineage_update_queue, daemon=True)
        _lineage_processor_thread.start()
        logging.info("✅ Started lineage update queue processor thread")

# Initialize the processor thread at module load
_ensure_lineage_processor_thread()

# Cache will be initialized after app creation
cache = None

# Rate limiting for API endpoints
RATE_LIMIT_WINDOW = 60  # 1 minute window
RATE_LIMIT_MAX_REQUESTS = 100  # Max requests per minute per IP (increased for label generation)

# Simple in-memory rate limiter
rate_limit_data = defaultdict(list)

# ============================================================================
# IP-based Store Selection System
# ============================================================================

# In-memory store for IP-based store selections (12-hour expiration)
_ip_store_selections = {}
_ip_store_lock = threading.Lock()
_store_selections_file = 'sessions/store_selections.pkl'

def is_store_selection_valid(ip_address, store_selection):
    """Check if store selection is still valid (within 12 hours)."""
    if not store_selection:
        return False

    # Store selections persist across server restarts for 12 hours
    # No need to invalidate based on server_id - time-based expiration is sufficient

    selection_time = store_selection.get('timestamp')
    if not selection_time:
        return False

    try:
        selection_datetime = datetime.fromisoformat(selection_time)
        expiration_time = selection_datetime + timedelta(hours=12)
        return datetime.now() < expiration_time
    except (ValueError, TypeError):
        return False

def save_store_selections():
    """Save store selections to disk for persistence across restarts."""
    try:
        with _ip_store_lock:
            os.makedirs('sessions', exist_ok=True)
            with open(_store_selections_file, 'wb') as f:
                pickle.dump(_ip_store_selections, f)
            logging.debug(f"Saved {len(_ip_store_selections)} store selections to disk")
    except Exception as e:
        logging.warning(f"Failed to save store selections: {e}")

def load_store_selections():
    """Load store selections from disk."""
    global _ip_store_selections
    try:
        if os.path.exists(_store_selections_file):
            with open(_store_selections_file, 'rb') as f:
                loaded = pickle.load(f)
                # Only load non-expired selections
                valid_selections = {}
                for ip, data in loaded.items():
                    if is_store_selection_valid(ip, data):
                        valid_selections[ip] = data
                _ip_store_selections = valid_selections
                logging.info(f"Loaded {len(_ip_store_selections)} valid store selections from disk")
                return True
    except Exception as e:
        logging.warning(f"Failed to load store selections: {e}")
    return False

# OPTION 1: Clear ALL store selections on server restart (uncomment to force selection every restart)
def clear_all_on_startup():
    """Clear all store selections when server restarts - forces users to select store every time."""
    global _ip_store_selections
    with _ip_store_lock:
        count = len(_ip_store_selections)
        _ip_store_selections.clear()

    # Delete the persistence file to prevent reload
    try:
        if os.path.exists(_store_selections_file):
            os.remove(_store_selections_file)
            logging.warning(f"🔥 STARTUP: Deleted store selections file - {_store_selections_file}")
    except Exception as e:
        logging.warning(f"Failed to delete store selections file: {e}")

    # CRITICAL: Also clear all caches to prevent wrong tags from previous session
    try:
        cache.clear()
        logging.warning(f"🔥 STARTUP: Cleared all Flask caches - fresh start for all users")
    except Exception as e:
        logging.warning(f"Failed to clear cache on startup: {e}")

    logging.warning(f"🔥 STARTUP: Cleared all {count} store selections - STORE MODAL WILL SHOW FOR ALL USERS")

# OPTION 2: Load persisted selections and only clear expired ones (12-hour persistence)
def load_and_cleanup_on_startup():
    """Load persisted store selections and clear only expired ones."""
    global _ip_store_selections
    
    # Load from disk
    load_store_selections()
    
    # Clear expired ones
    expired_count = 0
    with _ip_store_lock:
        expired_ips = []
        for ip_address, store_selection in _ip_store_selections.items():
            if not is_store_selection_valid(ip_address, store_selection):
                expired_ips.append(ip_address)
        
        for ip_address in expired_ips:
            del _ip_store_selections[ip_address]
            expired_count += 1
    
    if expired_count > 0:
        logging.info(f"Cleared {expired_count} expired store selection(s) on startup")
        save_store_selections()  # Save after cleanup
    else:
        logging.info(f"Loaded {len(_ip_store_selections)} valid store selections from disk")

# Choose which startup behavior you want:
# Uncomment ONE of the following:

# Force store selection on every server restart:
clear_all_on_startup()

# OR keep 12-hour persistence across restarts:
# load_and_cleanup_on_startup()

# ------------------------------------------------------------------------------
# Storage cleanup (uploads + stray DB files)
# ------------------------------------------------------------------------------
def cleanup_old_uploads(max_age_hours: int = 12):
    """Remove Excel uploads older than max_age_hours to keep PythonAnywhere tidy."""
    try:
        upload_folder = app.config.get('UPLOAD_FOLDER', UPLOADS_DIR)
        if not os.path.exists(upload_folder):
            return
        now = time.time()
        removed = 0
        for fname in os.listdir(upload_folder):
            path = os.path.join(upload_folder, fname)
            if not os.path.isfile(path):
                continue
            lower = fname.lower()
            if lower.endswith(('.xlsx', '.xls', '.xlsm', '.csv')):
                age_hours = (now - os.path.getmtime(path)) / 3600.0
                if age_hours > max_age_hours:
                    try:
                        os.remove(path)
                        removed += 1
                    except Exception as file_err:
                        logging.warning(f"Cleanup: failed to remove {path}: {file_err}")
        if removed:
            logging.info(f"Cleanup: removed {removed} old upload file(s) (> {max_age_hours}h)")
    except Exception as e:
        logging.warning(f"Cleanup: error pruning old uploads: {e}")


def cleanup_non_store_databases():
    """Delete product_database_*.db files that are not tied to known stores."""
    try:
        upload_folder = app.config.get('UPLOAD_FOLDER', UPLOADS_DIR)
        if not os.path.exists(upload_folder):
            return
        allowed_filenames = {f"product_database_{store}.db" for store in VALID_STORES}
        allowed_filenames.update({'product_database.db', 'product_database_backup.db'})
        removed = 0
        for fname in os.listdir(upload_folder):
            path = os.path.join(upload_folder, fname)
            if not os.path.isfile(path):
                continue
            if fname.startswith('product_database_') and fname.endswith('.db') and fname not in allowed_filenames:
                try:
                    os.remove(path)
                    removed += 1
                    logging.info(f"Cleanup: removed stray DB {fname}")
                except Exception as db_err:
                    logging.warning(f"Cleanup: failed to remove {path}: {db_err}")
        if removed:
            logging.info(f"Cleanup: removed {removed} non-store database file(s)")
    except Exception as e:
        logging.warning(f"Cleanup: error pruning non-store databases: {e}")


def start_storage_cleanup_scheduler():
    """Run periodic cleanup (hourly) to prune stray DBs."""
    def _worker():
        while True:
            try:
                cleanup_non_store_databases()
            except Exception as e:
                logging.warning(f"Cleanup scheduler error: {e}")
            # Run hourly
            time.sleep(3600)
    try:
        t = threading.Thread(target=_worker, daemon=True)
        t.start()
        logging.info("Cleanup scheduler started (non-store DBs hourly)")
    except Exception as e:
        logging.warning(f"Failed to start cleanup scheduler: {e}")


def start_daily_upload_cleanup_scheduler(run_hour: int = 0, run_minute: int = 0):
    """Schedule upload cleanup to run once daily at the specified time (default: midnight)."""
    def seconds_until_target():
        now = datetime.now()
        target = now.replace(hour=run_hour, minute=run_minute, second=0, microsecond=0)
        if target <= now:
            target += timedelta(days=1)
        return (target - now).total_seconds()

    def _worker():
        while True:
            try:
                wait_seconds = max(60, seconds_until_target())
                time.sleep(wait_seconds)
                cleanup_old_uploads()
            except Exception as e:
                logging.warning(f"Daily upload cleanup error: {e}")
                time.sleep(600)  # Backoff before retrying

    try:
        t = threading.Thread(target=_worker, daemon=True)
        t.start()
        logging.info("Daily upload cleanup scheduler started (runs at midnight server time)")
    except Exception as e:
        logging.warning(f"Failed to start daily upload cleanup scheduler: {e}")

def get_client_ip():
    """Get the client's IP address."""
    from flask import has_request_context
    if not has_request_context():
        return None
    if request.headers.get('X-Forwarded-For'):
        return request.headers.get('X-Forwarded-For').split(',')[0].strip()
    elif request.headers.get('X-Real-IP'):
        return request.headers.get('X-Real-IP')
    else:
        return request.remote_addr

def get_current_store_name(allow_fallback=True):
    """Get the current store name for the requesting client. Returns None if no valid selection."""
    try:
        # Check if we have a request context before accessing session/request
        from flask import has_request_context
        if not has_request_context():
            # Outside request context - return None or fallback store
            if allow_fallback:
                return 'AGT_Bothell'  # Default fallback
            return None
        
        # CRITICAL FIX: Check Flask session first (most reliable). Keep the value even if server instance changed.
        if session.get('selected_store'):
            return session.get('selected_store')
        
        # Fallback to IP-based selection
        ip_address = get_client_ip()
        if ip_address is not None:
            with _ip_store_lock:
                if ip_address in _ip_store_selections:
                    store_data = _ip_store_selections[ip_address]
                    # Check if the selection is still valid (not expired)
                    if is_store_selection_valid(ip_address, store_data):
                        # Store selection is valid - use it regardless of server instance
                        # Also save to session for consistency
                        session['selected_store'] = store_data['store']
                        session['store_server_id'] = SERVER_INSTANCE_ID
                        return store_data['store']
                    else:
                        # Remove expired selection
                        del _ip_store_selections[ip_address]
        
        if allow_fallback:
            # Simpler fallback: default to Bothell instead of auto-picking largest DB (avoids surprise switches)
            return 'AGT_Bothell'
    except Exception as e:
        logging.warning(f"Error getting current store name: {e}")
        return None

def has_store_selection():
    """Check if current IP has a valid store selection."""
    try:
        # Quick check: if we're outside request context, return False immediately
        from flask import has_request_context
        if not has_request_context():
            return False
        
        # CRITICAL FIX: Check Flask session FIRST (most reliable) and keep value even if server instance changed
        if session.get('selected_store'):
            return True
        
        # Fallback to IP-based selection
        ip_address = get_client_ip()
        
        # Fast path: check without lock first
        if ip_address not in _ip_store_selections:
            return False
        
        # Only acquire lock if we found a potential match
        with _ip_store_lock:
            if ip_address in _ip_store_selections:
                store_data = _ip_store_selections[ip_address]
                is_valid = is_store_selection_valid(ip_address, store_data)
                if is_valid:
                    # Also save to session for consistency
                    session['selected_store'] = store_data['store']
                    session['store_server_id'] = SERVER_INSTANCE_ID
                return is_valid
        
        return False
    except Exception as e:
        # Silently fail outside request context
        return False

def cleanup_expired_store_selections():
    """Remove expired store selections."""
    current_time = datetime.now()
    expired_ips = []
    
    with _ip_store_lock:
        for ip_address, store_selection in _ip_store_selections.items():
            if not is_store_selection_valid(ip_address, store_selection):
                expired_ips.append(ip_address)
        
        for ip_address in expired_ips:
            del _ip_store_selections[ip_address]
    
    # Save to disk after cleanup if any were removed
    if expired_ips:
        save_store_selections()

def extract_store_from_filename(filename):
    """Extract store name from filename if present. Handles variations like spaces, underscores, case."""
    if not filename:
        return None
    
    filename_upper = filename.upper()
    # Normalize filename by replacing underscores and spaces for matching
    filename_normalized = filename_upper.replace('_', ' ').replace('-', ' ')
    
    # Store mappings: (search_pattern, proper_store_name)
    # Check both with "AGT" prefix and without (just store name)
    store_patterns = [
        # With AGT prefix
        ('AGT BOTHELL', 'AGT_Bothell'),
        ('AGT_BOTHELL', 'AGT_Bothell'),
        ('AGT BURIEN', 'AGT_Burien'),
        ('AGT_BURIEN', 'AGT_Burien'),
        ('AGT GOLDBAR', 'AGT_Goldbar'),
        ('AGT_GOLDBAR', 'AGT_Goldbar'),
        ('AGT LYNNWOOD', 'AGT_Lynnwood'),
        ('AGT_LYNNWOOD', 'AGT_Lynnwood'),
        ('AGT SEATTLE', 'AGT_Seattle'),
        ('AGT_SEATTLE', 'AGT_Seattle'),
        ('AGT SHORELINE', 'AGT_Shoreline'),
        ('AGT_SHORELINE', 'AGT_Shoreline'),
        ('AGT WALLA WALLA', 'AGT_Walla_Walla'),
        ('AGT_WALLA_WALLA', 'AGT_Walla_Walla'),
        ('AGT WALLAWALLA', 'AGT_Walla_Walla'),
        # Without AGT prefix (just store name) - check these after AGT versions
        ('BOTHELL', 'AGT_Bothell'),
        ('BURIEN', 'AGT_Burien'),
        ('GOLDBAR', 'AGT_Goldbar'),
        ('LYNNWOOD', 'AGT_Lynnwood'),
        ('SEATTLE', 'AGT_Seattle'),
        ('SHORELINE', 'AGT_Shoreline'),
        ('WALLA WALLA', 'AGT_Walla_Walla'),
        ('WALLAWALLA', 'AGT_Walla_Walla'),
    ]
    
    # Check for store name in filename (case insensitive, handles spaces/underscores)
    for pattern, store_name in store_patterns:
        pattern_normalized = pattern.replace('_', ' ').replace('-', ' ')
        if pattern_normalized in filename_normalized or pattern in filename_upper:
            return store_name
    
    return None

def validate_excel_filename_for_store(filename, selected_store):
    """
    Validate that Excel filename contains store name and matches selected store.
    Returns (is_valid, warning_message, detected_store)
    """
    if not filename:
        return False, "Filename is required", None
    
    detected_store = extract_store_from_filename(filename)
    
    if not detected_store:
        return False, f"Excel filename must contain a store name (e.g., 'AGT_Bothell', 'AGT_Burien', etc.). Found filename: {filename}", None
    
    if detected_store != selected_store:
        # CRITICAL FIX: Add diagnostic information to help debug store mismatch
        from flask import session, has_request_context
        session_store = None
        ip_store = None
        if has_request_context():
            session_store = session.get('selected_store')
            ip_address = get_client_ip()
            with _ip_store_lock:
                if ip_address in _ip_store_selections:
                    ip_store = _ip_store_selections[ip_address].get('store')
        
        diagnostic_info = f"Detected store in filename: {detected_store}, Selected store: {selected_store}"
        if session_store:
            diagnostic_info += f", Session store: {session_store}"
        if ip_store:
            diagnostic_info += f", IP-based store: {ip_store}"
        
        logging.error(f"Store mismatch detected: {diagnostic_info}")
        
        # CRITICAL FIX: If session has the correct store but selected_store is wrong, suggest refreshing
        if session_store == detected_store and selected_store != detected_store:
            return False, f"Store mismatch: Your session shows '{session_store}' but the system selected '{selected_store}'. Please refresh the page and try again, or manually select '{detected_store}' store.", detected_store
        
        return False, f"Store mismatch: Cannot upload {detected_store} Excel file to {selected_store}. Please select the correct store or use the correct Excel file. (Session: {session_store}, IP: {ip_store})", detected_store
    
    return True, None, detected_store

# ============================================================================
# End of IP-based Store Selection System
# ============================================================================

def reset_excel_processor():
    """Reset the global ExcelProcessor to force reloading of the default file."""
    global _excel_processor, _excel_processor_reset_flag
    
    logging.info("Resetting Excel processor - clearing all data")
    
    if _excel_processor is not None:
        # Explicitly clear all data
        if hasattr(_excel_processor, 'df') and _excel_processor.df is not None:
            del _excel_processor.df
            logging.info("Cleared DataFrame from ExcelProcessor")
        
        if hasattr(_excel_processor, 'selected_tags'):
            _excel_processor.selected_tags = []
            logging.info("Cleared selected tags from ExcelProcessor")
        
        if hasattr(_excel_processor, 'dropdown_cache'):
            _excel_processor.dropdown_cache = {}
            logging.info("Cleared dropdown cache from ExcelProcessor")
        
        # Force garbage collection
        import gc
        gc.collect()
        logging.info("Forced garbage collection")
    
    # Set to None to force recreation
    _excel_processor = None
    
    # Set reset flag to prevent automatic default file loading
    _excel_processor_reset_flag = True
    logging.info("Set reset flag to prevent automatic default file loading")
    
    # Clear all caches
    clear_initial_data_cache()
    
    # CRITICAL FIX: Preserve JSON matched tags when clearing cache
    try:
        # Check if we have JSON matched tags that should be preserved
        json_matched_cache_key = session.get('json_matched_cache_key')
        if json_matched_cache_key and cache.has(json_matched_cache_key):
            json_matched_tags = cache.get(json_matched_cache_key)
            logging.info(f"CRITICAL FIX: Preserving {len(json_matched_tags)} JSON matched tags during cache clear")
            
            # Clear the general available_tags cache
            cache_key = get_session_cache_key('available_tags')
            cache.delete(cache_key)
            logging.info(f"Cleared general cache for key: {cache_key}")
            
            # Restore JSON matched tags to available_tags cache
            if json_matched_tags:
                cache.set(cache_key, json_matched_tags, timeout=3600)
                logging.info(f"CRITICAL FIX: Restored {len(json_matched_tags)} JSON matched tags to available_tags cache")
        else:
            # No JSON matched tags to preserve, clear normally
            cache_key = get_session_cache_key('available_tags')
            cache.delete(cache_key)
            logging.info(f"Cleared cache for key: {cache_key}")
    except Exception as cache_error:
        logging.warning(f"Error clearing cache: {cache_error}")
    
    logging.info("Excel processor reset complete")
def force_reload_excel_processor(new_file_path):
    """Force reload the Excel processor with a new file. ALWAYS clears old data completely."""
    global _excel_processor, _excel_processor_reset_flag
    
    logging.info(f"Force reloading Excel processor with new file: {new_file_path}")
    
    # ALWAYS create a completely new ExcelProcessor instance to ensure clean slate
    logging.info("Creating new ExcelProcessor instance to ensure complete data replacement")
    
    # Clear the old processor completely
    if _excel_processor is not None:
        # Explicitly clear all data from old processor
        if hasattr(_excel_processor, 'df') and _excel_processor.df is not None:
            del _excel_processor.df
            logging.info("Cleared old DataFrame from ExcelProcessor")
        
        if hasattr(_excel_processor, 'selected_tags'):
            _excel_processor.selected_tags = []
            logging.info("Cleared selected tags from ExcelProcessor")
        
        if hasattr(_excel_processor, 'dropdown_cache'):
            _excel_processor.dropdown_cache = {}
            logging.info("Cleared dropdown cache from ExcelProcessor")
        
        # Force garbage collection
        import gc
        gc.collect()
        logging.info("Forced garbage collection to free memory")
    
    # Create a completely new instance
    _excel_processor = ExcelProcessor()
    
    # Enable product database integration by default
    if hasattr(_excel_processor, 'enable_product_db_integration'):
        _excel_processor.enable_product_db_integration(True)
        logging.info("Product database integration enabled by default")
    
    # Clear the reset flag since we're loading a new file
    _excel_processor_reset_flag = False
    logging.info("Cleared reset flag - loading new file")
    
    # Load the new file with full processing rules
    success = _excel_processor.load_file(new_file_path)
    if success:
        _excel_processor._last_loaded_file = new_file_path
        logging.info(f"Excel processor successfully loaded new file with full processing rules: {new_file_path}")
        logging.info(f"New DataFrame shape: {_excel_processor.df.shape if _excel_processor.df is not None else 'None'}")
        # CRITICAL FIX: Ensure dropdown cache is populated after successful file load
        if hasattr(_excel_processor, '_cache_dropdown_values'):
            try:
                _excel_processor._cache_dropdown_values()
                logging.info(f"Successfully populated dropdown cache from session uploaded file")
                # Log the strain count specifically
                if 'strain' in _excel_processor.dropdown_cache:
                    strain_count = len(_excel_processor.dropdown_cache['strain'])
                    logging.info(f"Dropdown cache contains {strain_count} strains")
                else:
                    logging.warning("No strain filter found in dropdown cache")
            except Exception as e:
                logging.error(f"Failed to populate dropdown cache: {e}")
        else:
            logging.warning("ExcelProcessor does not have _cache_dropdown_values method")
    else:
        logging.error(f"Failed to load new file in Excel processor: {new_file_path}")
        # CRITICAL FIX: Don't create empty DataFrame - this causes the "no strains" issue
        # Instead, try to load a default file as fallback
        from src.core.data.excel_processor import get_default_upload_file
        # Get store-specific default file
        selected_store = get_current_store_name() if has_store_selection() else None
        default_file = get_default_upload_file(selected_store)
        if default_file and os.path.exists(default_file):
            logging.info(f"Attempting to load default file as fallback: {default_file}")
            fallback_success = _excel_processor.load_file(default_file)
            if fallback_success:
                _excel_processor._last_loaded_file = default_file
                logging.info(f"Successfully loaded default file as fallback: {default_file}")
                # Populate dropdown cache for fallback file
                if hasattr(_excel_processor, '_cache_dropdown_values'):
                    try:
                        _excel_processor._cache_dropdown_values()
                        logging.info(f"Successfully populated dropdown cache from fallback file")
                    except Exception as e:
                        logging.error(f"Failed to populate dropdown cache from fallback: {e}")
            else:
                logging.error(f"Failed to load default file as fallback: {default_file}")
                if hasattr(_excel_processor, 'df') and _excel_processor.df is not None and not _excel_processor.df.empty:
                    logging.warning("Preserving existing DataFrame to avoid wiping loaded data")
                else:
                    # Only create empty DataFrame if there truly is no data loaded
                    _excel_processor.df = pd.DataFrame()
                    _excel_processor.selected_tags = []
                    logging.warning("Created empty DataFrame as last resort - no prior data available")
        else:
            logging.error("No default file available as fallback")
            if hasattr(_excel_processor, 'df') and _excel_processor.df is not None and not _excel_processor.df.empty:
                logging.warning("Preserving existing DataFrame despite missing fallback file")
            else:
                _excel_processor.df = pd.DataFrame()
                _excel_processor.selected_tags = []
                logging.warning("Created empty DataFrame as last resort - this may cause 'no strains' issues")

def cleanup_old_processing_status():
    """Clean up old processing status entries to prevent memory leaks."""
    with processing_lock:
        current_time = time.time()
        # Keep entries for at least 15 minutes to give frontend time to poll
        cutoff_time = current_time - 900  # 15 minutes
        
        old_entries = []
        for filename, status in processing_status.items():
            timestamp = processing_timestamps.get(filename, 0)
            age = current_time - timestamp
            
            # Only remove entries that are older than 15 minutes AND not currently processing
            # Also, be more conservative with 'ready' status to prevent race conditions
            if age > cutoff_time and status != 'processing':
                # For 'ready' status, wait much longer to ensure frontend has completed
                # Increased from 30 minutes to 60 minutes to prevent race conditions
                if status == 'ready' and age < 3600:  # 60 minutes for ready status
                    continue
                old_entries.append(filename)
        
        for filename in old_entries:
            del processing_status[filename]
            if filename in processing_timestamps:
                del processing_timestamps[filename]
            logging.debug(f"Cleaned up old processing status for: {filename}")

def update_processing_status(filename, status):
    """Update processing status with timestamp."""
    with processing_lock:
        processing_status[filename] = status
        processing_timestamps[filename] = time.time()
        logging.info(f"Updated processing status for {filename}: {status}")
        logging.debug(f"Current processing statuses: {dict(processing_status)}")

def get_excel_processor():
    """Return a fresh ExcelProcessor for the current store/session.

    Avoids sharing processors across requests while keeping callers intact.
    """
    from src.core.data.excel_processor import ExcelProcessor, get_default_upload_file

    # Resolve store context
    store_name = None
    try:
        store_name = get_current_store_name(allow_fallback=True)
    except Exception:
        store_name = None

    processor = ExcelProcessor(store_name=store_name)

    # Enable product DB integration by default for lineage lookups
    if hasattr(processor, 'enable_product_db_integration'):
        try:
            processor.enable_product_db_integration(True)
        except Exception:
            pass

    # Try to load session file if present
    try:
        from flask import has_request_context, session
        if has_request_context():
            session_file = session.get('file_path')
            if session_file and os.path.exists(session_file):
                processor.load_file(session_file)
                processor._last_loaded_file = session_file
    except Exception:
        pass

    # If nothing loaded, optionally load default file for the store
    if not getattr(processor, '_last_loaded_file', None):
        try:
            default_file = get_default_upload_file(store_name or get_current_store_name(allow_fallback=True))
            if default_file and os.path.exists(default_file):
                if processor.load_file(default_file):
                    processor._last_loaded_file = default_file
        except Exception:
            pass

    return processor

# Ensure no stray indentation or orphaned blocks before function definition
def _resolve_database_path_for_store(store_name: Optional[str] = None) -> Tuple[Optional[str], Optional[str]]:
    """
    Resolve the best available database path for a given store.
    Returns (db_path, resolved_store_name) or (None, None) if no database exists.
    """
    uploads_dir = os.path.join(current_dir, 'uploads')
    if not os.path.isdir(uploads_dir):
        logging.warning(f"Uploads directory not found when resolving database path: {uploads_dir}")
        return None, None

    normalized_requested = ''
    if store_name:
        normalized_requested = store_name.replace('_', '').replace(' ', '').lower()

    # 1. Exact match: product_database_{store_name}.db
    if store_name:
        exact_path = os.path.join(uploads_dir, f'product_database_{store_name}.db')
        if os.path.exists(exact_path):
            return exact_path, store_name

    # 2. Fuzzy match: look for any product_database_* file containing the store token
    candidate_paths = glob.glob(os.path.join(uploads_dir, 'product_database_*.db'))
    if candidate_paths:
        for candidate in candidate_paths:
            candidate_store = os.path.basename(candidate).replace('product_database_', '').replace('.db', '')
            candidate_normalized = candidate_store.replace('_', '').replace(' ', '').lower()
            if normalized_requested and normalized_requested in candidate_normalized:
                return candidate, candidate_store or store_name

    # 3. Generic database fallback
    generic_db = os.path.join(uploads_dir, 'product_database.db')
    if os.path.exists(generic_db):
        logging.warning(f"Generic product database fallback in use: {generic_db}")
        return generic_db, store_name

    # 4. Most recent database as final fallback
    if candidate_paths:
        newest_path = max(candidate_paths, key=os.path.getmtime)
        inferred_store = os.path.basename(newest_path).replace('product_database_', '').replace('.db', '')
        logging.warning(f"No database found for store '{store_name}'. Using most recent database: {newest_path}")
        return newest_path, inferred_store or store_name

    logging.error(f"No product database files found in uploads directory: {uploads_dir}")
    return None, None


def get_product_database(store_name=None):
    """Lazy load ProductDatabase to avoid startup delay."""
    global _product_database
    
    if store_name is None:
        logging.warning("get_product_database called without store name; attempting fallback resolution.")
    
    db_path, resolved_store = _resolve_database_path_for_store(store_name)
    if not db_path:
        raise FileNotFoundError(f"No product database file available for store '{store_name}'. Upload the database via the admin tools.")

    effective_store = resolved_store or store_name

    # Check if reload is needed
    current_store_in_db = getattr(_product_database, '_store_name', None) if _product_database else None
    current_db_path = getattr(_product_database, 'db_path', None) if _product_database else None
    needs_reload = (_product_database is None or current_store_in_db != effective_store or current_db_path != db_path)

    if needs_reload:
        from src.core.data.product_database import ProductDatabase
        _product_database = ProductDatabase(db_path)
        _product_database._store_name = effective_store
        if getattr(_product_database, 'db_path', db_path) != db_path:
            logging.warning(f"ProductDatabase db_path mismatch: {_product_database.db_path} != {db_path}")
        _product_database.init_database()

    return _product_database

# Local override: always try to use the Bothell DB file if present
def _get_bothell_product_db():
    try:
        from src.core.data.product_database import ProductDatabase
        current_dir = os.path.dirname(os.path.abspath(__file__))
        # Helper: find most recent valid Bothell DB in uploads
        def _find_best_bothell_db(base_dir: str) -> str:
            import glob, os, sqlite3
            candidates = []
            # Prefer explicit AGT_Bothell naming
            patterns = [
                os.path.join(base_dir, 'uploads', 'product_database_AGT_Bothell*.db'),
                os.path.join(base_dir, 'uploads', 'product_database*.db'),
                os.path.join(base_dir, 'bothell_products.db')
            ]
            for pattern in patterns:
                for path in glob.glob(pattern):
                    try:
                        with db_connection(path) as conn:
                            cur = conn.cursor()
                            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='products'")
                            has_products = cur.fetchone() is not None
                            if has_products:
                                mtime = os.path.getmtime(path)
                                candidates.append((mtime, path))
                    except Exception:
                        continue
            if not candidates:
                return ''
            candidates.sort(reverse=True)
            return candidates[0][1]
        def has_required_tables(db_path: str) -> bool:
            import sqlite3
            try:
                with db_connection(db_path) as conn:
                    cur = conn.cursor()
                    cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='products'")
                    has_products = cur.fetchone() is not None
                    cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='strains'")
                    has_strains = cur.fetchone() is not None
                    return has_products or has_strains
            except Exception:
                return False
        # Preferred: best (most recent) valid Bothell DB
        best = _find_best_bothell_db(current_dir)
        if best:
            from src.core.data.product_database import ProductDatabase
            current_dir = os.path.dirname(os.path.abspath(__file__))
            # Always use Bothell store-specific database
            db_filename = 'product_database_AGT_Bothell.db'
            db_path = os.path.join(current_dir, 'uploads', db_filename)
            if not os.path.exists(db_path):
                logging.warning(f"Bothell store database not found: {db_path}")
                return None
            product_db = ProductDatabase(db_path)
            product_db._store_name = 'AGT_Bothell'
            product_db.init_database()
            return product_db
        return None
    except Exception as e:
        logging.error(f"Error in _get_bothell_product_db: {e}")
        return None

def get_enhanced_ai_matcher():
    """Lazy load Enhanced AI Product Matcher."""
    global _enhanced_ai_matcher
    if _enhanced_ai_matcher is None:
        try:
            from src.core.data.enhanced_ai_matcher import EnhancedAIProductMatcher
            _enhanced_ai_matcher = EnhancedAIProductMatcher()
            logging.info("Enhanced AI Product Matcher initialized successfully")
        except ImportError as e:
            logging.warning(f"Enhanced AI Product Matcher not available: {e}")
            _enhanced_ai_matcher = None
    return _enhanced_ai_matcher

def disable_product_db_integration():
    """Disable product database integration to improve load times."""
    try:
        excel_processor = get_excel_processor()
        if hasattr(excel_processor, 'enable_product_db_integration'):
            excel_processor.enable_product_db_integration(False)
            logging.info("Product database integration disabled")
    except Exception as e:
        logging.error(f"Error disabling product DB integration: {e}")

def get_cached_initial_data():
    """Get cached initial data if it's still valid."""
    global _initial_data_cache, _cache_timestamp
    if (_initial_data_cache is not None and 
        _cache_timestamp is not None and 
        time.time() - _cache_timestamp < CACHE_DURATION):
        return _initial_data_cache
    return None

def set_cached_initial_data(data):
    """Cache initial data with timestamp."""
    global _initial_data_cache, _cache_timestamp
    _initial_data_cache = data
    _cache_timestamp = time.time()

def clear_initial_data_cache():
    """Clear the initial data cache."""
    global _initial_data_cache, _cache_timestamp
    _initial_data_cache = None
    _cache_timestamp = None

def set_landscape(doc):
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    # Set minimal margins
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)
    section.top_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)
    # Swap width and height for landscape
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
 
def resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller"""
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = current_dir
    return os.path.join(base_path, relative_path)
def create_app():
    import flask
    app = flask.Flask(__name__, static_url_path='/static', static_folder='static')
    app.config.from_object('config.Config')

    # CRITICAL FIX: Prevent JSON encoding issues with Unicode characters
    # Set JSON_AS_ASCII to False to allow Unicode characters (like zero-width joiners) in JSON
    app.config['JSON_AS_ASCII'] = False
    app.json.ensure_ascii = False

    # CRITICAL FIX: Ensure proper URL generation for reverse proxy setups
    # Set SERVER_NAME only if explicitly provided (don't set it to avoid URL generation issues)
    if os.environ.get('SERVER_NAME'):
        app.config['SERVER_NAME'] = os.environ.get('SERVER_NAME')
    # Ensure PREFERRED_URL_SCHEME is set correctly for HTTPS
    app.config['PREFERRED_URL_SCHEME'] = 'https' if os.environ.get('HTTPS', '').lower() in ('on', 'true', '1') else 'http'
    
    # Enable development mode for auto-reload and debug features
    # ENABLED: Allow easy restart by enabling development mode by default
    app.config['DEVELOPMENT_MODE'] = os.environ.get('DEVELOPMENT_MODE', 'true').lower() == 'true'
    
    # Enable detailed logging for development
    logging.getLogger().setLevel(logging.DEBUG)
    logging.getLogger('werkzeug').setLevel(logging.DEBUG)
    
    # Performance optimizations
    app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH
    app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 31536000  # 1 year cache for static files
    app.config['PERMANENT_SESSION_LIFETIME'] = 3600  # 1 hour session timeout
    
    # PythonAnywhere-specific optimizations
    if PYTHONANYWHERE_OPTIMIZATION:
        app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB for PythonAnywhere
        app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 300  # 5 minutes for PythonAnywhere
        app.config['PERMANENT_SESSION_LIFETIME'] = 7200  # 2 hours for PythonAnywhere (increased for file persistence)
        # Reduce logging verbosity for PythonAnywhere
        logging.getLogger('werkzeug').setLevel(logging.WARNING)
        logging.getLogger('urllib3').setLevel(logging.WARNING)
        logging.getLogger('requests').setLevel(logging.WARNING)
        logging.info("PythonAnywhere optimizations applied")
    
    # Compression for better performance
    if Compress:
        compress = Compress()
        compress.init_app(app)
        # Aggressive compression settings for web performance
        app.config['COMPRESS_ALGORITHM'] = 'gzip'
        app.config['COMPRESS_LEVEL'] = 6  # Balance between speed and compression
        app.config['COMPRESS_MIN_SIZE'] = 500  # Compress responses over 500 bytes
        logging.info("Flask-Compress enabled with aggressive settings for better performance")
    
    # Initialize session management
    if Session:
        # CRITICAL FIX: Explicitly set session configuration before initializing
        sessions_dir = os.path.join(current_dir, 'sessions')
        os.makedirs(sessions_dir, exist_ok=True)
        
        app.config['SESSION_TYPE'] = 'filesystem'
        app.config['SESSION_FILE_DIR'] = sessions_dir
        app.config['SESSION_PERMANENT'] = True  # Enable session persistence to keep Excel uploads across browser restarts
        app.config['SESSION_USE_SIGNER'] = True
        app.config['SESSION_KEY_PREFIX'] = 'labelmaker:'
        app.config['SESSION_FILE_THRESHOLD'] = 500  # Max number of session files
        
        Session(app)
        logging.info(f"Flask-Session initialized with filesystem storage at {sessions_dir}")
        logging.info(f"Session config: TYPE={app.config.get('SESSION_TYPE')}, DIR={app.config.get('SESSION_FILE_DIR')}")
    else:
        logging.warning("Flask-Session not available, using default session handling")
    
    # Enable CORS for all routes to allow uploads
    allowed_origins = [
        'https://www.agtpricetags.com',  # Your actual domain
        'https://agtpricetags.com',
        'http://localhost:5000',  # For local development
        'http://localhost:5001',  # For local development
        'http://127.0.0.1:5000',
        'http://127.0.0.1:5001',
        'https://adamcordova.pythonanywhere.com'  # PythonAnywhere domain
    ]
    # Add caching headers for static resources to improve performance
    @app.after_request
    def add_cache_headers(response):
        """Add caching headers for static resources to improve web performance."""
        # For static files (CSS, JS, images), use long cache times
        if request.path.startswith('/static/'):
            # Cache static files for 7 days (can be longer with version control)
            response.cache_control.max_age = 604800  # 7 days in seconds
            response.cache_control.public = True
            # Add immutable flag for versioned resources
            if 'v=' in request.query_string.decode():
                response.cache_control.immutable = True
        # For API endpoints, minimal caching
        elif request.path.startswith('/api/'):
            response.cache_control.max_age = 0
            response.cache_control.no_cache = True
        # For HTML pages, short cache with revalidation
        elif response.content_type and 'text/html' in response.content_type:
            response.cache_control.max_age = 300  # 5 minutes
            response.cache_control.must_revalidate = True
        return response
    
    # Enable CORS for all routes, not just /api/*
    CORS(app, resources={
        r"/*": {"origins": allowed_origins}
    })
    
    # CRITICAL FIX: Add error handler to ensure API routes always return JSON, not HTML
    @app.errorhandler(500)
    def handle_500_error(e):
        """Handle 500 errors and return JSON for API routes."""
        import traceback
        from flask import Response, request, has_request_context
        
        # Always try to return JSON for API routes
        try:
            error_traceback = traceback.format_exc()
            error_msg = str(e) if e else 'Internal server error'
            error_type = type(e).__name__ if hasattr(e, '__class__') else 'UnknownError'
            
            # Log the error
            try:
                logging.error(f"❌ 500 ERROR: {error_type}: {error_msg}")
                logging.error(f"Full traceback:\n{error_traceback}")
            except:
                pass  # Don't fail if logging fails
            
            # Check if this is an API route - default to True if we can't determine
            is_api_route = True  # Default to JSON response for safety
            request_path = 'unknown'
            try:
                if has_request_context() and hasattr(request, 'path') and request.path:
                    request_path = request.path
                    is_api_route = request.path.startswith('/api/')
                # Also check URL rule if available
                elif has_request_context() and hasattr(request, 'url_rule') and request.url_rule:
                    request_path = str(request.url_rule)
                    is_api_route = '/api/' in request_path
            except Exception as path_error:
                # If we can't determine, assume it's an API route to be safe
                logging.debug(f"Could not determine if API route: {path_error}")
                is_api_route = True
            
            # Always return JSON for API routes, or if we can't determine (default to JSON)
            if is_api_route:
                try:
                    # Escape error message to prevent JSON injection
                    safe_error_msg = error_msg.replace('"', '\\"').replace('\n', '\\n').replace('\r', '\\r')
                    return jsonify({
                        'success': False,
                        'error': safe_error_msg,
                        'error_type': error_type,
                        'path': request_path
                    }), 500
                except Exception as jsonify_error:
                    # If jsonify fails, return plain JSON string
                    try:
                        safe_error_msg = error_msg.replace('"', '\\"').replace('\n', '\\n').replace('\r', '\\r')
                        error_json = f'{{"success": false, "error": "{safe_error_msg}", "error_type": "{error_type}", "path": "{request_path}"}}'
                        return Response(error_json, status=500, mimetype='application/json')
                    except Exception as json_error:
                        # Last resort: return minimal JSON
                        return Response('{"success": false, "error": "Internal server error", "error_type": "HandlerError"}', status=500, mimetype='application/json')
            
            # For non-API routes, re-raise to use Flask's default handler
            raise e
            
        except Exception as handler_error:
            # If error handler itself fails, return minimal JSON
            try:
                logging.error(f"❌ CRITICAL: Error handler failed: {handler_error}")
                import traceback
                logging.error(f"Error handler traceback: {traceback.format_exc()}")
            except:
                pass
            try:
                return Response('{"success": false, "error": "Critical error occurred", "error_type": "HandlerError"}', status=500, mimetype='application/json')
            except:
                # Ultimate fallback - this should never happen
                return Response('{"success": false, "error": "Unknown error"}', status=500, mimetype='application/json')
    
    
    # Check if we're in development mode
    development_mode = app.config.get('DEVELOPMENT_MODE', False)

    # Respect environment: disable template auto-reload to prevent force reloads
    app.config['TEMPLATES_AUTO_RELOAD'] = False
    app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0  # Always disable caching for development
    app.config['DEBUG'] = bool(app.config.get('DEBUG', development_mode))
    app.config['PROPAGATE_EXCEPTIONS'] = bool(development_mode)
    if development_mode:
        logging.info("Running in DEVELOPMENT mode with template auto-reload DISABLED to prevent force reloads")
    else:
        logging.info("Running in PRODUCTION mode with static asset caching enabled")
    
    app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024  # 20MB max file size
    app.config['TESTING'] = False
    app.config['SESSION_REFRESH_EACH_REQUEST'] = False  # Don't refresh session on every request
    # CRITICAL FIX: Ensure store/session persistence lasts long enough - increased for upload persistence
    # Increased session lifetime to ensure uploads persist through page reloads
    default_session_seconds = 21600 if PYTHONANYWHERE_OPTIMIZATION else 21600  # 6 hours for both (increased from 2h local)
    current_lifetime = app.config.get('PERMANENT_SESSION_LIFETIME', default_session_seconds)
    if isinstance(current_lifetime, timedelta):
        current_seconds = current_lifetime.total_seconds()
    else:
        current_seconds = int(current_lifetime) if current_lifetime else 0
    session_seconds = max(current_seconds, default_session_seconds)
    app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(seconds=session_seconds)
    logging.info(f"✅ Session lifetime set to {session_seconds/3600:.1f} hours for upload persistence")
    
    # Session configuration to prevent cookie size issues
    app.config['SESSION_COOKIE_SECURE'] = False  # Allow HTTP in development
    app.config['SESSION_COOKIE_HTTPONLY'] = True
    app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
    app.config['SESSION_COOKIE_MAX_SIZE'] = 8192  # Increased browser cookie size limit
    
    upload_folder = os.path.join(current_dir, 'uploads')
    os.makedirs(upload_folder, exist_ok=True)
    app.config['UPLOAD_FOLDER'] = upload_folder
    # Use a consistent secret key for production to maintain sessions across restarts
    # In production, this should be set via environment variable
    app.secret_key = os.environ.get('SECRET_KEY', 'label-maker-secret-key-2024-production')

    # Enable gzip compression for JSON and text responses to reduce bandwidth/latency
    if Compress is not None:
        app.config.setdefault('COMPRESS_ALGORITHM', 'gzip')
        app.config.setdefault('COMPRESS_LEVEL', 6)
        app.config.setdefault('COMPRESS_MIN_SIZE', 1024)  # Only compress payloads >1KB
        app.config.setdefault('COMPRESS_MIMETYPES', [
            'application/json',
            'text/html',
            'text/css',
            'application/javascript',
            'text/javascript',
            'text/plain'
        ])
    return app

# Debug: Track app creation
import sys

# Create Flask app instance
app = create_app()

# Initialize Flask-Caching after app creation (if available)
if CACHE_AVAILABLE:
    # CRITICAL FIX: Use a shared filesystem cache so preroll QR data is visible
    # across all worker processes (SimpleCache is per-process and loses data between workers).
    cache_config = {
        'CACHE_TYPE': 'FileSystemCache',
        'CACHE_DIR': CACHE_DIR,
        # 24h default timeout for group/item data so QR codes remain valid through the day
        'CACHE_DEFAULT_TIMEOUT': 86400
    }
    cache = Cache(app, config=cache_config)
else:
    cache = Cache()  # Use dummy cache

# Flask-Compress already initialized in create_app() - no need to initialize again

# Add performance headers and response optimization
@app.after_request
def add_performance_headers(response):
    """Add performance optimization headers to all responses"""
    # Add caching headers for static assets
    if request.path.startswith('/static/'):
        response.headers['Cache-Control'] = 'public, max-age=31536000'  # 1 year
    # Add cache headers for API responses (short-lived)
    elif request.path.startswith('/api/'):
        if response.status_code == 200:
            response.headers['Cache-Control'] = 'private, max-age=60'  # 1 minute for API responses
    # Security headers
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'SAMEORIGIN'
    # Add timing header for debugging
    if hasattr(request, '_start_time'):
        elapsed = (time.time() - request._start_time) * 1000
        response.headers['X-Response-Time'] = f"{elapsed:.1f}ms"
    return response

# Start periodic storage cleanup (uploads + stray DBs)
start_storage_cleanup_scheduler()
start_daily_upload_cleanup_scheduler(run_hour=0, run_minute=0)

@app.before_request
def start_request_timer():
    """Start timer for request duration tracking"""
    request._start_time = time.time()

# Initialize performance optimizations - DISABLED to prevent CPU issues on PythonAnywhere
if False:  # Temporarily disabled due to high CPU usage on PythonAnywhere
    try:
        if FULL_PERFORMANCE_OPTIMIZATIONS_AVAILABLE:
            # Full performance optimizations (local development)
            # start_performance_monitoring()  # DISABLED - causes high CPU
            # preload_all()  # DISABLED - causes high CPU
            logging.info("Performance optimizations DISABLED to prevent CPU issues")
        elif LIGHTWEIGHT_PERFORMANCE_AVAILABLE:
            # Lightweight optimizations only (PythonAnywhere)
            logging.info("Lightweight performance optimizations available but DISABLED")
        
    except Exception as e:
        logging.warning(f"Failed to initialize performance optimizations: {e}")
        # Continue without performance optimizations to prevent app crash

# Log that performance monitoring is disabled
logging.info("Performance monitoring DISABLED to prevent high CPU usage on PythonAnywhere")
# Global function to check session size
def check_session_size():
    """Check if session is too large and clear it if necessary."""
    try:
        # Only check if session has data
        if not session:
            return False
            
        # Try to serialize session data safely
        session_copy = {}
        for key, value in session.items():
            try:
                # Test if this value can be pickled
                import pickle
                pickle.dumps(value)
                session_copy[key] = value
            except (pickle.PicklingError, TypeError):
                # Skip unpicklable objects
                logging.warning(f"Skipping unpicklable session key: {key}")
                continue
        
        # Check size of serializable data
        import pickle
        session_data = pickle.dumps(session_copy)
        if len(session_data) > 3000:  # 3KB limit to stay well under 4KB
            logging.warning(f"Session too large ({len(session_data)} bytes), clearing session data")
            # Store essential data before clearing
            selected_tags = session.get('selected_tags', [])
            selected_store = session.get('selected_store', '')
            file_path = session.get('file_path', '')
            uploaded_filename = session.get('uploaded_filename', '')
            upload_timestamp = session.get('upload_timestamp', None)
            session.clear()
            # Restore essential data after clearing
            if selected_tags:
                session['selected_tags'] = selected_tags
            if selected_store:
                session['selected_store'] = selected_store
            if file_path:
                session['file_path'] = file_path
            if uploaded_filename:
                session['uploaded_filename'] = uploaded_filename
            if upload_timestamp:
                session['upload_timestamp'] = upload_timestamp
            logging.info(f"Preserved {len(selected_tags)} selected tags and file_path during session optimization")
            return True
    except Exception as e:
        logging.error(f"Error checking session size: {e}")
    return False

def optimize_session_data():
    """Optimize session data to reduce size."""
    try:
        # Only optimize if session has data
        if not session:
            return False
            
        # Only keep essential session data
        essential_keys = ['selected_tags', 'file_path', 'uploaded_filename', 'upload_timestamp', UNDO_STACK_KEY]
        session_copy = {}
        
        for key in essential_keys:
            if key in session:
                try:
                    if key == 'selected_tags':
                        # Store only tag names, not full objects
                        if isinstance(session[key], list):
                            # Convert to strings if they aren't already
                            session_copy[key] = []
                            for tag in session[key]:
                                if isinstance(tag, str):
                                    session_copy[key].append(tag)
                                elif isinstance(tag, dict) and 'Product Name*' in tag:
                                    session_copy[key].append(tag['Product Name*'])
                                else:
                                    session_copy[key].append(str(tag))
                        else:
                            session_copy[key] = []
                    elif key == UNDO_STACK_KEY:
                        # Limit undo stack to 3 entries max
                        undo_stack = session[key][-3:] if len(session[key]) > 3 else session[key]
                        session_copy[key] = undo_stack
                    else:
                        session_copy[key] = session[key]
                except Exception as e:
                    logging.warning(f"Error processing session key {key}: {e}")
                    continue
        
        # Test if the optimized data can be serialized
        try:
            import pickle
            pickle.dumps(session_copy)
            
            # Clear and restore only essential data
            # Store essential data before clearing
            selected_tags = session.get('selected_tags', [])
            selected_store = session.get('selected_store', '')
            file_path = session.get('file_path', '')
            uploaded_filename = session.get('uploaded_filename', '')
            upload_timestamp = session.get('upload_timestamp', None)
            session.clear()
            session.update(session_copy)
            # Restore essential data if they weren't in the optimized data
            if selected_tags and 'selected_tags' not in session_copy:
                session['selected_tags'] = selected_tags
            if selected_store and 'selected_store' not in session_copy:
                session['selected_store'] = selected_store
            if file_path and 'file_path' not in session_copy:
                session['file_path'] = file_path
            if uploaded_filename and 'uploaded_filename' not in session_copy:
                session['uploaded_filename'] = uploaded_filename
            if upload_timestamp and 'upload_timestamp' not in session_copy:
                session['upload_timestamp'] = upload_timestamp
            logging.info(f"Restored {len(selected_tags)} selected tags and file_path after session optimization")
            
            logging.info("Session data optimized")
            return True
        except (pickle.PicklingError, TypeError) as e:
            logging.warning(f"Optimized session data still contains unpicklable objects: {e}")
            return False
            
    except Exception as e:
        logging.error(f"Error optimizing session data: {e}")
        return False

# Initialize Excel processor and load default data on startup

def simple_initialize_excel_processor():
    """Simple initialization that won't get stuck - for PythonAnywhere"""
    try:
        logging.info("Simple initialization starting...")
        
        # Create Excel processor without loading any files
        excel_processor = get_excel_processor()
        excel_processor.logger.setLevel(logging.WARNING)
        
        # Initialize with empty DataFrame
        if not hasattr(excel_processor, 'df') or excel_processor.df is None:
            excel_processor.df = pd.DataFrame()
            logging.info("Initialized with empty DataFrame")
        
        # CRITICAL FIX: Keep ProductDB integration enabled for lineage support
        # Even if disabled for performance, we still need database lineage queries to work
        # The integration being disabled only affects background processing, not direct queries
        if hasattr(excel_processor, 'enable_product_db_integration'):
            # Keep enabled for lineage support - direct database queries still work
            excel_processor.enable_product_db_integration(True)
            logging.info("Product database integration enabled for lineage support")
        
        logging.info("Simple initialization completed successfully")
        return True
        
    except Exception as e:
        logging.error(f"Error in simple initialization: {e}")
        logging.error(f"Traceback: {traceback.format_exc()}")
        return False


def initialize_excel_processor():
    """Initialize Excel processor and load default data."""
    try:
        # Skip initialization if startup file loading is disabled for performance
        if DISABLE_STARTUP_FILE_LOADING:
            logging.info("Startup file loading disabled for faster application startup")
            return

        # CRITICAL FIX: get_excel_processor() is deprecated and returns None
        # Skip initialization - processors are now created per-request
        excel_processor = get_excel_processor()
        if excel_processor is None:
            logging.info("Excel processor initialization skipped - using per-request processors")
            return

        # Safety check: ensure excel_processor has a logger attribute
        if not hasattr(excel_processor, 'logger') or excel_processor.logger is None:
            logging.warning("Excel processor does not have a logger - skipping logger configuration")
        else:
            excel_processor.logger.setLevel(logging.WARNING)
        
        # Enable product database integration by default
        if hasattr(excel_processor, 'enable_product_db_integration'):
            excel_processor.enable_product_db_integration(True)
            logging.info("Product database integration enabled by default")
        
        # CRITICAL FIX: Check for session file FIRST before loading default file
        # This ensures uploaded files persist across page reloads
        session_file_path = None
        try:
            from flask import session, has_request_context
            if has_request_context():
                session_file_path = session.get('file_path')
                if session_file_path and os.path.exists(session_file_path):
                    logging.info(f"✅ Found session file in initialize_excel_processor: {session_file_path}")
                    # Check if already loaded
                    if excel_processor._last_loaded_file != session_file_path or not hasattr(excel_processor, 'df') or excel_processor.df is None or excel_processor.df.empty:
                        logging.info(f"📂 Loading session file in initialize_excel_processor: {session_file_path}")
                        success = excel_processor.load_file(session_file_path)
                        if success:
                            excel_processor._last_loaded_file = session_file_path
                            row_count = len(excel_processor.df) if hasattr(excel_processor, 'df') and excel_processor.df is not None else 0
                            logging.info(f"✅ Session file loaded successfully with {row_count} records")
                            return  # Don't load default file if session file was loaded
                        else:
                            logging.warning(f"⚠️ Failed to load session file: {session_file_path}")
                    else:
                        logging.info(f"✅ Session file already loaded: {session_file_path}")
                        return  # Don't load default file if session file is already loaded
        except Exception as session_check_error:
            logging.debug(f"Could not check session in initialize_excel_processor: {session_check_error}")
        
        # Only load default file if no session file was found/loaded
        from src.core.data.excel_processor import get_default_upload_file
        # CRITICAL FIX: Use allow_fallback=True for default file loading on startup
        # This ensures default file loads even if store hasn't been selected yet
        selected_store = get_current_store_name(allow_fallback=True)
        default_file = get_default_upload_file(selected_store)
        
        if default_file and os.path.exists(default_file):
            logging.info(f"Loading default file on startup: {default_file}")
            try:
                # Add timeout protection for corrupted files
                import signal
                
                def timeout_handler(signum, frame):
                    raise TimeoutError("Excel file loading timed out - file may be corrupted")
                
                # Set 30 second timeout for file loading
                signal.signal(signal.SIGALRM, timeout_handler)
                signal.alarm(30)
                
                try:
                    success = excel_processor.load_file(default_file)
                    signal.alarm(0)  # Cancel the alarm
                    
                    if success:
                        excel_processor._last_loaded_file = default_file
                        logging.info(f"Default file loaded successfully with {len(excel_processor.df)} records")
                    else:
                        logging.warning("Failed to load default file")
                except TimeoutError as timeout_err:
                    signal.alarm(0)  # Cancel the alarm
                    logging.error(f"Excel file loading timed out: {timeout_err}")
                    logging.error(f"File may be corrupted: {default_file}")
                    # Try to move corrupted file
                    try:
                        corrupted_path = default_file + '.corrupted'
                        os.rename(default_file, corrupted_path)
                        logging.info(f"Moved corrupted file to: {corrupted_path}")
                    except Exception as move_err:
                        logging.error(f"Could not move corrupted file: {move_err}")
                        
            except Exception as load_error:
                logging.error(f"Error loading default file: {load_error}")
                logging.error(f"Traceback: {traceback.format_exc()}")
        else:
            logging.info("No default file found, waiting for user upload")
            if default_file:
                logging.info(f"Default file path was found but file doesn't exist: {default_file}")
            
    except Exception as e:
        logging.error(f"Error initializing Excel processor: {e}")
        logging.error(f"Traceback: {traceback.format_exc()}")

# Initialize on startup
# Load Excel file on startup for immediate availability
# Excel processor will be ready when user first visits the site
if not os.environ.get('PYTHONANYWHERE_DOMAIN') and not os.environ.get('PYTHONANYWHERE_SITE'):
    # Only initialize on local development
    try:
        initialize_excel_processor()
    except Exception as e:
        logging.warning(f"Startup initialization failed (non-fatal): {e}")

# Add missing function
def save_template_settings(template_type, font_settings):
    """Save template settings to a configuration file."""
    try:
        config_dir = Path(__file__).parent / 'config'
        config_dir.mkdir(exist_ok=True)
        config_file = config_dir / f'{template_type}_settings.json'
        
        with open(config_file, 'w') as f:
            json.dump(font_settings, f, indent=2)
        
        logging.info(f"Saved template settings for {template_type}")
    except Exception as e:
        logging.error(f"Error saving template settings: {str(e)}")
        raise

# Global enhanced logger instance
enhanced_logger = None

# --- Enhanced Logging System ---
try:
    from enhanced_logging import setup_enhanced_logging, EnhancedLogger, ErrorContext, log_route_error, log_database_error, log_file_processing_error  # type: ignore[import]
    ENHANCED_LOGGING_AVAILABLE = True
    enhanced_logger = setup_enhanced_logging()
    print("✅ Enhanced logging system loaded")
except ImportError as e:
    ENHANCED_LOGGING_AVAILABLE = False
    enhanced_logger = None
    print(f"⚠️  Enhanced logging not available: {e}")

# --- LabelMakerApp Class ---
class LabelMakerApp:
    def __init__(self):
        self.app = app
        self._configure_logging()
        
    def _configure_logging(self):
        """Configure enhanced logging system"""
        
        if ENHANCED_LOGGING_AVAILABLE:
            # Use enhanced logging system
            self.enhanced_logger = setup_enhanced_logging()
            self.logger = self.enhanced_logger.logger
            self.logger.info("🚀 Enhanced logging system initialized")
        else:
            # Fallback to basic logging
            self.logger = logging.getLogger(__name__)
            if not self.logger.handlers:
                # Create logs directory if it doesn't exist
                log_dir = Path(__file__).parent / 'logs'
                log_dir.mkdir(exist_ok=True)
                
                # Set up logging format
                log_format = '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
                formatter = logging.Formatter(log_format)
                
                # Configure console handler - show info and above for debugging
                console_handler = logging.StreamHandler()
                console_handler.setLevel(logging.INFO)  # Show info, warnings, and errors
                console_handler.setFormatter(formatter)
                
                # Configure file handler
                log_file = log_dir / 'label_maker.log'
                file_handler = logging.FileHandler(log_file)
                file_handler.setLevel(logging.INFO)
                file_handler.setFormatter(formatter)
                
                # Configure root logger
                logging.basicConfig(
                    level=logging.INFO,
                    format=log_format,
                    handlers=[console_handler, file_handler]
                )
                
                # Suppress verbose logging from third-party libraries
                logging.getLogger('watchdog').setLevel(logging.WARNING)
                logging.getLogger('werkzeug').setLevel(logging.WARNING)
                logging.getLogger('urllib3').setLevel(logging.WARNING)
                logging.getLogger('requests').setLevel(logging.WARNING)
                
                # Add handlers to application logger
                self.logger.addHandler(console_handler)
                self.logger.addHandler(file_handler)
                self.logger.setLevel(logging.INFO)
                
                self.logger.debug("Basic logging configured for Label Maker application")
            self.logger.debug(f"Log file location: {log_file}")
    
    def _is_port_available(self, port):
        """Check if a port is available for binding."""
        import socket
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(1)
            result = sock.connect_ex(('127.0.0.1', port))
            sock.close()
            return result != 0  # Port is available if connection fails
        except Exception:
            return False
            
    def run(self):
        host = os.environ.get('HOST', '127.0.0.1')
        port = int(os.environ.get('FLASK_PORT', 8001))
        development_mode = self.app.config.get('DEVELOPMENT_MODE', False)
        
        # PREVENT MULTIPLE RESTARTS: Check if app is already running
        import socket
        try:
            sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            sock.settimeout(1)
            result = sock.connect_ex((host, port))
            sock.close()
            
            if result == 0:
                logging.warning(f"⚠️  Port {port} is already in use! Finding available port...")
                # Try ports 8001-8010
                for test_port in range(8001, 8011):
                    if self._is_port_available(test_port):
                        port = test_port
                        logging.info(f"✅ Found available port: {port}")
                        break
                else:
                    logging.error("❌ No available ports found in range 8001-8010")
                    return
        except Exception as e:
            logging.debug(f"Port check failed (this is normal): {e}")
        
        # Show optimization status
        if DISABLE_STARTUP_FILE_LOADING:
            logging.info("🚀 PERFORMANCE OPTIMIZATION: Startup file loading disabled for faster app startup")
        
        logging.info(f"Starting Label Maker application on {host}:{port}")
        print(f"🌐 App will be available at: http://{host}:{port}")
        logging.info(f"Development mode: {development_mode}")
        
        # DEVELOPMENT MODE: Keep debug but disable reloader to prevent multiple restarts
        self.app.run(
            host=host, 
            port=port, 
            debug=development_mode, 
            use_reloader=False,  # Disable reloader to prevent multiple restarts
            threaded=True  # Enable threading for better performance
        )
# === SESSION-BASED HELPERS ===
def get_session_excel_processor():
    """Get ExcelProcessor instance for the current session with proper error handling."""
    session_file_path = None  # Initialize to prevent variable scoping errors
    
    # CRITICAL: Recursion protection - prevent infinite loop
    if hasattr(g, '_getting_excel_processor'):
        logging.error("⚠️ RECURSION DETECTED in get_session_excel_processor - returning fallback")
        try:
            from src.core.data.excel_processor import ExcelProcessor
            import pandas as pd
            fallback = ExcelProcessor()
            fallback.df = pd.DataFrame()
            fallback.selected_tags = []
            return fallback
        except:
            return None
    
    try:
        g._getting_excel_processor = True  # Set recursion guard
        
        if 'excel_processor' not in g:
            # CRITICAL: Create NEW processor instance instead of using deprecated global
            # This prevents session data leakage
            from src.core.data.excel_processor import ExcelProcessor
            # CRITICAL FIX: Don't default to Bothell - only set store if user has selected one
            # This ensures store modal appears for first-time users
            store_name = get_current_store_name(allow_fallback=False) if has_store_selection() else None
            # If no store selected, create processor without store (won't load default file)
            if store_name:
                g.excel_processor = ExcelProcessor(store_name=store_name)
            else:
                # Create processor without store - it will wait for user to select
                g.excel_processor = ExcelProcessor(store_name='AGT_Bothell')  # Temporary fallback for initialization
                g.excel_processor._no_default_load = True  # Flag to prevent auto-loading

            # CRITICAL FIX: Keep ProductDB integration enabled for lineage support
            # Direct database queries (get_product_lineage) still work even if integration is disabled
            # But enabling it ensures lineage data is available when needed
            if hasattr(g.excel_processor, 'enable_product_db_integration'):
                g.excel_processor.enable_product_db_integration(True)
            
            # Ensure processor store context matches the active store selection
            try:
                active_store = get_current_store_name()
            except Exception:
                active_store = None
            if active_store:
                current_store = getattr(g.excel_processor, '_store_name', None)
                if current_store != active_store:
                    logging.info(f"🔄 Session ExcelProcessor store context update: '{current_store}' → '{active_store}'")
                    g.excel_processor._store_name = active_store
                    if hasattr(g.excel_processor, '_invalidate_caches'):
                        g.excel_processor._invalidate_caches()
            
            # CRITICAL FIX: Check if we have an uploaded file in session
            session_file_path = session.get('file_path')
            session_store = session.get('file_store', '')
            # Store context removed - using single database

            # CRITICAL FIX: If session doesn't have file_path, try to restore from persistent file
            if not session_file_path:
                try:
                    import json
                    # CRITICAL FIX: Use UPLOADS_DIR constant instead of constructing path for Windows compatibility
                    persistence_file = os.path.join(UPLOADS_DIR, '.last_upload.json')
                    if os.path.exists(persistence_file):
                        with open(persistence_file, 'r') as f:
                            last_upload = json.load(f)
                        persisted_file_path = last_upload.get('file_path')
                        persisted_store = last_upload.get('store')
                        current_store = get_current_store_name() if has_store_selection() else None
                        
                        # CRITICAL FIX: Normalize path for Windows compatibility
                        if persisted_file_path:
                            persisted_file_path = os.path.normpath(persisted_file_path)
                        
                        # Only restore if file exists and store matches (or no store selected)
                        if persisted_file_path and os.path.exists(persisted_file_path):
                            if not current_store or persisted_store == current_store:
                                session_file_path = persisted_file_path
                                session['file_path'] = persisted_file_path
                                session['uploaded_filename'] = last_upload.get('filename', '')
                                session['upload_timestamp'] = last_upload.get('timestamp', 0)
                                session['file_store'] = persisted_store
                                session.modified = True
                                logging.info(f"✅ Restored upload from persistent file in get_session_excel_processor: {session_file_path}")
                        else:
                            if persisted_file_path:
                                logging.warning(f"⚠️ Persistent file path does not exist: {persisted_file_path}")
                except Exception as restore_err:
                    logging.warning(f"Could not restore upload from persistent file: {restore_err}")
                    logging.error(traceback.format_exc())

            if session_file_path and os.path.exists(session_file_path):
                # CRITICAL: Load the session file into the new processor instance
                logging.info(f"📂 Loading session file: {session_file_path}")
                try:
                    success = g.excel_processor.load_file(session_file_path)
                    if success and g.excel_processor.df is not None and not g.excel_processor.df.empty:
                        row_count = len(g.excel_processor.df)
                        logging.info(f"✅ Loaded session file: {session_file_path} ({row_count} rows)")
                        g.excel_processor._last_loaded_file = session_file_path
                    else:
                        logging.warning(f"⚠️ Failed to load session file or file is empty: {session_file_path}")
                except Exception as load_err:
                    logging.error(f"❌ Error loading session file: {load_err}")
                    import traceback
                    logging.error(traceback.format_exc())
            elif session_file_path:
                logging.warning(f"Session uploaded file does not exist: {session_file_path}")
                # CRITICAL FIX: Don't clear session data immediately - try persistent file first
                # Only clear if persistent file also doesn't exist
                try:
                    import json
                    # CRITICAL FIX: Use UPLOADS_DIR constant for Windows compatibility
                    persistence_file = os.path.join(UPLOADS_DIR, '.last_upload.json')
                    if os.path.exists(persistence_file):
                        with open(persistence_file, 'r') as f:
                            last_upload = json.load(f)
                        # CRITICAL FIX: Normalize paths for comparison on Windows
                        persisted_path = os.path.normpath(last_upload.get('file_path', ''))
                        session_path = os.path.normpath(session_file_path)
                        if persisted_path == session_path:
                            # Persistent file also references missing file - clear both
                            os.remove(persistence_file)
                            logging.info(f"Removed persistent file referencing missing upload: {session_file_path}")
                except Exception as clear_err:
                    logging.warning(f"Error checking persistent file: {clear_err}")
                # Clear invalid session data
                session.pop('file_path', None)
                session.pop('uploaded_filename', None)
                session.pop('upload_timestamp', None)
            
            # CRITICAL FIX: Ensure ExcelProcessor has data for JSON matching
            # Only load default file if we don't have a session file AND DataFrame is empty
            # PERFORMANCE: Skip default file loading if _skip_default_file_load flag is set (for fast_load)
            skip_default_load = getattr(g, '_skip_default_file_load', False)
            if not session_file_path and not skip_default_load:
                if not hasattr(g.excel_processor, 'df') or g.excel_processor.df is None or g.excel_processor.df.empty:
                    logging.info("CRITICAL FIX: No session file and DataFrame is empty, loading default file")
                    from src.core.data.excel_processor import get_default_upload_file
                    # CRITICAL FIX: Use allow_fallback=True for default file loading
                    selected_store = get_current_store_name(allow_fallback=True)
                    default_file = get_default_upload_file(selected_store)
                    if default_file and os.path.exists(default_file):
                        logging.info(f"CRITICAL FIX: Loading default file: {default_file}")
                        # Load file (fast_mode removed - not available on PythonAnywhere)
                        success = g.excel_processor.load_file(default_file)
                        if success:
                            logging.info(f"CRITICAL FIX: Successfully loaded default file")
                            # Populate dropdown cache
                            if hasattr(g.excel_processor, '_cache_dropdown_values'):
                                try:
                                    g.excel_processor._cache_dropdown_values()
                                    logging.info(f"Successfully populated dropdown cache from default file")
                                except Exception as e:
                                    logging.error(f"Failed to populate dropdown cache from default file: {e}")
                        else:
                            logging.error(f"CRITICAL FIX: Failed to load default file: {default_file}")
                    else:
                        logging.warning("CRITICAL FIX: No default file available")
            elif skip_default_load:
                logging.info("⚡ Fast load: Skipping default file loading in get_session_excel_processor")
        
        # CRITICAL FIX: For new uploaded files, update the last processed file but DON'T clear tags
        if session_file_path and session_file_path != getattr(g.excel_processor, '_last_processed_file', None):
            logging.info(f"CRITICAL FIX: New uploaded file detected, updating last processed file")
            logging.info(f"CRITICAL FIX: Previous file: {getattr(g.excel_processor, '_last_processed_file', 'None')}")
            logging.info(f"CRITICAL FIX: New file: {session_file_path}")
            logging.info(f"CRITICAL FIX: Selected tags before update: {len(g.excel_processor.selected_tags)}")
            
            # Update the last processed file but preserve selected tags
            g.excel_processor._last_processed_file = session_file_path
            
            # CRITICAL FIX: Clear caches for new file but preserve selected tags
            logging.info(f"CRITICAL FIX: Clearing caches for new file (preserving selected tags)")
            if hasattr(g.excel_processor, '_file_cache'):
                g.excel_processor._file_cache.clear()
                logging.info(f"CRITICAL FIX: Cleared file cache")
            if hasattr(g.excel_processor, '_dropdown_cache'):
                g.excel_processor._dropdown_cache.clear()
                logging.info(f"CRITICAL FIX: Cleared dropdown cache")
            if hasattr(g.excel_processor, '_available_tags_cache'):
                g.excel_processor._available_tags_cache.clear()
                logging.info(f"CRITICAL FIX: Cleared available tags cache")
            
            logging.info(f"CRITICAL FIX: Selected tags after update: {len(g.excel_processor.selected_tags)}")
            logging.info(f"CRITICAL FIX: Session selected tags after update: {len(session.get('selected_tags', []))}")
        
        # Ensure selected_tags attribute exists
        if not hasattr(g.excel_processor, 'selected_tags'):
            g.excel_processor.selected_tags = []
        
        # Restore selected tags from session
        session_selected_tag_names = session.get('selected_tags', [])
        logging.info(f"Session selected_tags count: {len(session_selected_tag_names)}")
        
        # Convert tag names back to full tag objects
        if session_selected_tag_names:
            restored_tags = []
            for tag_name in session_selected_tag_names:
                # Find the tag in the current data
                found_tag = None
                
                # Try to find in DataFrame first
                if hasattr(g.excel_processor, 'df') and g.excel_processor.df is not None:
                    possible_columns = ['ProductName', 'Product Name*', 'Product Name']
                    for col in possible_columns:
                        if col in g.excel_processor.df.columns:
                            mask = g.excel_processor.df[col] == tag_name
                            if mask.any():
                                row = g.excel_processor.df[mask].iloc[0]
                                found_tag = row.to_dict()
                                break
                
                # If not found in DataFrame, try data attribute
                if not found_tag and hasattr(g.excel_processor, 'data'):
                    for tag in g.excel_processor.data:
                        if tag.get('Product Name*') == tag_name:
                            found_tag = tag
                            break
                
                if found_tag:
                    restored_tags.append(found_tag)
                else:
                    logging.warning(f"Tag not found in data: {tag_name}")
            
            g.excel_processor.selected_tags = restored_tags
        else:
            g.excel_processor.selected_tags = []
        
        logging.info(f"Restored {len(g.excel_processor.selected_tags)} selected tags from session")
        logging.info(f"Session selected_tags: {session_selected_tag_names}")
        # Truncate large log messages to prevent "Message too long" error
        selected_tags_preview = str(g.excel_processor.selected_tags)[:500] + "..." if len(str(g.excel_processor.selected_tags)) > 500 else str(g.excel_processor.selected_tags)
        logging.info(f"Excel processor selected_tags after restore: {selected_tags_preview}")
        
        # Final safety check - ensure df attribute exists
        if not hasattr(g.excel_processor, 'df'):
            logging.error("ExcelProcessor missing df attribute - creating empty DataFrame")
            import pandas as pd
            g.excel_processor.df = pd.DataFrame()
        
        # Validate store context to prevent cross-store data access
        # Store context removed - using single database
        # Store context removed - using single database
        
        # Clear recursion guard before returning
        if hasattr(g, '_getting_excel_processor'):
            delattr(g, '_getting_excel_processor')
        
        return g.excel_processor
        
    except Exception as e:
        logging.error(f"Error in get_session_excel_processor: {str(e)}")
        logging.error(traceback.format_exc())
        
        # Clear recursion guard on error
        if hasattr(g, '_getting_excel_processor'):
            delattr(g, '_getting_excel_processor')
        
        # Return a safe fallback ExcelProcessor
        try:
            from src.core.data.excel_processor import ExcelProcessor
            import pandas as pd
            fallback_processor = ExcelProcessor()
            fallback_processor.df = pd.DataFrame()  # Empty DataFrame
            fallback_processor.selected_tags = []
            return fallback_processor
        except Exception as fallback_error:
            logging.error(f"Failed to create fallback ExcelProcessor: {fallback_error}")
            # Return None and let the calling code handle it
            return None

def get_session_json_matcher():
    try:
        from src.core.data.json_matcher import JSONMatcher
        excel_processor = get_session_excel_processor()
        if excel_processor is None:
            logging.error("Cannot create JSONMatcher: ExcelProcessor is None")
            return None
        
        # Use a global JSON matcher instance to persist the cache
        if not hasattr(app, '_json_matcher'):
            app._json_matcher = JSONMatcher(excel_processor)
            
            # CRITICAL FIX: Build cache from database to ensure JSON matching works
            try:
                # Build the sheet cache from database - this will auto-select the best database
                app._json_matcher._build_cache_from_database()
                if app._json_matcher._sheet_cache and len(app._json_matcher._sheet_cache) > 0:
                    logging.info(f"JSON matcher loaded {len(app._json_matcher._sheet_cache)} products from database cache")
                else:
                    logging.warning("No products found in JSON matcher cache - JSON matching may not work")
            except Exception as e:
                logging.error(f"Error building JSON matcher cache from database: {e}")
            
            logging.info("Created new JSONMatcher instance")
        else:
            # Update the Excel processor reference in case it changed
            app._json_matcher.excel_processor = excel_processor
        
        return app._json_matcher
    except Exception as e:
        logging.warning(f"Enhanced JSON matcher unavailable, falling back to basic matcher: {e}")
        try:
            from src.core.data.json_matcher import JSONMatcher
            return JSONMatcher(get_session_excel_processor())
        except Exception as e2:
            logging.error(f"Failed to initialize basic JSON matcher: {e2}")
            return None

def get_session_product_database():
    """Get ProductDatabase instance for the current session using current store selection."""
    try:
        # Get the current store name and use it to get the database
        store_name = get_current_store_name()
        return get_product_database(store_name)
    except Exception as e:
        logging.error(f"Error getting session product database: {e}")
        return None


def generate_product_name_variants(raw_name):
    """
    Generate a list of product name variants to improve matching against the database.
    Handles vendor suffixes, trailing weights, non-breaking hyphens, and spacing differences.
    """
    import re

    variants = []
    if raw_name is None:
        return variants

    try:
        name = str(raw_name).strip()
    except Exception:
        name = raw_name

    if not name:
        return variants

    def add_variant(value):
        if not value:
            return
        cleaned = re.sub(r'\s+', ' ', value.strip().replace('\u2011', '-'))
        if cleaned and cleaned not in variants:
            variants.append(cleaned)

    add_variant(name)

    # Remove common "by Vendor" suffixes (optionally before a weight)
    vendor_removed = re.sub(r'\s+by\s+[^-]+(?=(\s*-\s*\d|\s*$))', '', name, flags=re.IGNORECASE)
    vendor_removed = re.sub(r'\s+by\s+[^-]+$', '', vendor_removed, flags=re.IGNORECASE)
    add_variant(vendor_removed)

    # Remove trailing weight indicators like "- 1g", "- 10pk", "- 100mg"
    weight_removed = re.sub(
        r'\s*-\s*\d+(?:\.\d+)?\s*(?:g|gram|grams|gm|oz|ounce|ounces|ml|mg|ct|pack|pk|pcs|pc)?$',
        '',
        vendor_removed,
        flags=re.IGNORECASE
    )
    add_variant(weight_removed)

    # Add variant without hyphens (helps when DB stored spaces instead)
    add_variant(weight_removed.replace('-', ' '))

    # Deduplicate while preserving order
    return variants

def _enhance_json_with_excel_data(json_tag, excel_product):
    """
    Enhance JSON tag data with Excel data while preserving the best information from both sources.
    
    Args:
        json_tag: Dictionary containing JSON product data
        excel_product: Dictionary containing Excel product data
        
    Returns:
        Dictionary with enhanced data combining both sources
    """
    enhanced_tag = json_tag.copy()
    
    # Use canonical fields for priority
    # CRITICAL FIX: Move Price, Weight, and DOH to Excel priority to ensure Excel data is used first
    json_priority_fields = [get_canonical_field(f) for f in ['Product Name*', 'ProductName', 'Vendor', 'Product Brand', 'Quantity*', 'Quantity']]
    excel_priority_fields = [get_canonical_field(f) for f in ['Lineage', 'Product Type*', 'Product Strain', 'Description', 'THC test result', 'CBD test result', 'Test result unit (% or mg)', 'Room*', 'State', 'Is Sample? (yes/no)', 'Is MJ product?(yes/no)', 'Discountable? (yes/no)', 'Medical Only (Yes/No)', 'DOH', 'Price', 'Weight*', 'Weight']]

    # Fill missing fields from Excel data
    for field in excel_priority_fields:
        canonical_field = get_canonical_field(field)
        if canonical_field in excel_product and excel_product[canonical_field] and (canonical_field not in enhanced_tag or not enhanced_tag[canonical_field]):
            enhanced_tag[canonical_field] = excel_product[canonical_field]

    # Fill missing fields from JSON data
    for field in json_priority_fields:
        canonical_field = get_canonical_field(field)
        if canonical_field in json_tag and json_tag[canonical_field] and (canonical_field not in enhanced_tag or not enhanced_tag[canonical_field]):
            enhanced_tag[canonical_field] = json_tag[canonical_field]

    # Add any additional Excel fields that don't exist in JSON
    for field, value in excel_product.items():
        canonical_field = get_canonical_field(field)
        if canonical_field not in enhanced_tag and value:
            enhanced_tag[canonical_field] = value

    # Ensure we have a proper display name
    if 'displayName' not in enhanced_tag or not enhanced_tag['displayName']:
        product_name = enhanced_tag.get(get_canonical_field('Product Name*'), enhanced_tag.get(get_canonical_field('ProductName'), ''))
        vendor = enhanced_tag.get(get_canonical_field('Vendor'), enhanced_tag.get(get_canonical_field('Product Brand'), ''))
        if product_name and vendor:
            enhanced_tag['displayName'] = f"{product_name} by {vendor}"
        elif product_name:
            enhanced_tag['displayName'] = product_name

    return enhanced_tag

@app.route('/api/status', methods=['GET'])
@cached_route(ttl=5, vary_by=['session_id'])
def api_status():
    """Check API server status and data loading status."""
    try:
        excel_processor = get_session_excel_processor()
        if excel_processor is None:
            return jsonify({
                'server': 'running',
                'data_loaded': False,
                'data_shape': None,
                'last_loaded_file': None,
                'selected_tags_count': 0,
                'error': 'Unable to initialize data processor'
            })
        
        # Get session manager for additional status info
        try:
            from src.core.data.session_manager import get_session_manager
            session_manager = get_session_manager()
            session_stats = session_manager.get_session_stats()
            session_id = session_manager.get_current_session_id()
            has_pending_changes = session_manager.has_pending_changes(session_id)
        except Exception as session_error:
            logging.warning(f"Error getting session stats: {session_error}")
            session_stats = {}
            has_pending_changes = False
        
        status = {
            'server': 'running',
            'data_loaded': excel_processor.df is not None and not excel_processor.df.empty,
            'data_shape': excel_processor.df.shape if excel_processor.df is not None else None,
            'last_loaded_file': getattr(excel_processor, '_last_loaded_file', None),
            'selected_tags_count': len(excel_processor.selected_tags) if hasattr(excel_processor, 'selected_tags') else 0,
            'session_stats': session_stats,
            'has_pending_changes': has_pending_changes,
            # Store context removed - using single database
            'file_store': session.get('file_store', '')
        }
        return jsonify(status)
    except Exception as e:
        logging.error(f"Error in status endpoint: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/favicon.ico')
def favicon():
    """Serve the favicon."""
    return send_from_directory(os.path.join(app.root_path, 'static'),
                               'favicon.ico', mimetype='image/vnd.microsoft.icon')

@app.route('/service-worker.js')
@app.route('/static/service-worker.js')
def service_worker():
    """Serve the service worker with proper headers."""
    response = send_from_directory(os.path.join(app.root_path, 'static'),
                                   'service-worker.js', mimetype='application/javascript')
    # Service workers must be served with proper MIME type and no caching for updates
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Service-Worker-Allowed'] = '/'
    return response

@app.after_request
def add_cache_control_headers(response):
    """Add cache control headers to prevent aggressive caching on PythonAnywhere."""
    # Don't cache JavaScript, CSS, or HTML files
    if (response.content_type and 
        ('javascript' in response.content_type or 
         'css' in response.content_type or 
         'html' in response.content_type)):
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate, max-age=0'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
    return response

@app.route('/test')
def test():
    """Simple test route to verify the app is working."""
    return jsonify({'status': 'ok', 'message': 'Flask app is running'})

# Auto check downloads functionality removed

@app.route('/')
def index():
    try:
        # CRITICAL FIX: Ensure request.url_root is valid (fix for "https://true" error)
        # This can happen with reverse proxy misconfigurations
        if hasattr(request, 'url_root') and (str(request.url_root) == 'true' or request.url_root is True):
            logging.warning(f"⚠️ Invalid request.url_root detected: {request.url_root}, template will use fallback")
        
        # Reduced logging to prevent excessive log spam
        logging.info(f"Page load at {datetime.now().strftime('%H:%M:%S')}")
        
        # Check if store selection is required
        user_has_store = has_store_selection()
        current_store = get_current_store_name() if user_has_store else None
        
        if not user_has_store:
            logging.info("No store selection found - user must select store")
        else:
            logging.info(f"User has valid store selection: {current_store}")
        
        # --- LIGHTWEIGHT PAGE LOAD (minimal work) ---
        # AUTOMATIC CACHE BUSTING: Use timestamp to force browser reload of updated files
        # This ensures users always get the latest JavaScript without manual cache clearing
        import time
        cache_bust = f"v2.1.{int(time.time())}"  # Timestamp ensures automatic cache invalidation
        
        # CRITICAL FIX: Don't clear uploaded file from session on page refresh
        # This was causing uploads to disappear when users refreshed the page
        uploaded_file = session.get('file_path', None)  # Keep the file path instead of removing it
        
        # CRITICAL FIX: If session doesn't have file_path, try to restore from persistent file
        if not uploaded_file:
            try:
                import json
                # CRITICAL FIX: Use UPLOADS_DIR constant for Windows compatibility
                persistence_file = os.path.join(UPLOADS_DIR, '.last_upload.json')
                if os.path.exists(persistence_file):
                    with open(persistence_file, 'r') as f:
                        last_upload = json.load(f)
                    persisted_file_path = last_upload.get('file_path')
                    persisted_store = last_upload.get('store')
                    current_store = get_current_store_name() if has_store_selection() else None
                    
                    # CRITICAL FIX: Normalize path for Windows compatibility
                    if persisted_file_path:
                        persisted_file_path = os.path.normpath(persisted_file_path)
                    
                    # Only restore if file exists and store matches (or no store selected)
                    if persisted_file_path and os.path.exists(persisted_file_path):
                        if not current_store or persisted_store == current_store:
                            uploaded_file = persisted_file_path
                            session['file_path'] = persisted_file_path
                            session['uploaded_filename'] = last_upload.get('filename', '')
                            session['upload_timestamp'] = last_upload.get('timestamp', 0)
                            session['file_store'] = persisted_store
                            session.modified = True
                            logging.info(f"✅ Restored upload from persistent file: {uploaded_file}")
                        else:
                            logging.info(f"⚠️ Persistent file exists but store mismatch: {persisted_store} != {current_store}")
            except Exception as restore_err:
                logging.warning(f"Could not restore upload from persistent file: {restore_err}")
        
        if uploaded_file:
            logging.info(f"Preserving uploaded file in session: {uploaded_file}")
        # Don't clear selected_tags - they should persist across page loads
        
        # Store selection will be handled by frontend JavaScript using localStorage
        
        # CRITICAL FIX: Increase retention time - only remove files that are very old (24 hours) or failed
        # This ensures uploads persist through normal page reloads
        if uploaded_file:
            try:
                from src.core.data.excel_processor import get_default_upload_file
                selected_store = get_current_store_name() if has_store_selection() else None
                default_file = get_default_upload_file(selected_store)
                
                if uploaded_file != default_file and os.path.exists(uploaded_file):
                    # Check if file is old (more than 1 hour)
                    file_age = time.time() - os.path.getmtime(uploaded_file)
                    upload_timestamp = session.get('upload_timestamp', 0)
                    
                    # Get processing status
                    filename = session.get('uploaded_filename', '')
                    status = processing_status.get(filename, 'unknown')
                    
                    # CRITICAL FIX: Only remove if file is very old (24 hours) OR processing failed
                    # This ensures uploads persist through normal page reloads and browser sessions
                    # Session lifetime is 2-6 hours, but files should persist longer for user convenience
                    should_remove = (
                        file_age > 86400 or  # More than 24 hours old (very old files)
                        status.startswith('error:') or  # Processing failed
                        (upload_timestamp > 0 and time.time() - upload_timestamp > 86400)  # Upload session expired (24 hours)
                    )
                    
                    if should_remove:
                        try:
                            os.remove(uploaded_file)
                            logging.info(f"Removed old/failed uploaded file: {uploaded_file} (age: {file_age:.0f}s, status: {status})")
                            # Clear session data for removed file
                            session.pop('file_path', None)
                            session.pop('uploaded_filename', None)
                            session.pop('upload_timestamp', None)
                        except Exception as e:
                            logging.warning(f"Failed to remove uploaded file: {e}")
                    else:
                        logging.info(f"Preserving recent uploaded file: {uploaded_file} (age: {file_age:.0f}s, status: {status})")
            except Exception as e:
                logging.warning(f"Error checking uploaded file: {e}")
        
        # Periodic cleanup (much less frequent - every 200th page load)
        import random
        if random.random() < 0.005:  # 0.5% chance to run cleanup
            try:
                cleanup_result = cleanup_old_files()
                if cleanup_result['success'] and cleanup_result['removed_count'] > 0:
                    logging.info(f"Auto-cleanup removed {cleanup_result['removed_count']} files")
            except Exception as cleanup_error:
                logging.warning(f"Auto-cleanup failed: {cleanup_error}")
        
        # Don't load data here - let frontend load via API calls
        # This makes page loads much faster
        initial_data = None
        
        # CRITICAL FIX: Pass uploaded filename to template so it persists on refresh
        uploaded_filename = session.get('uploaded_filename', '')
        
        logging.info("=== PAGE REFRESH COMPLETE ===")
        return render_template('index.html', 
                             initial_data=initial_data, 
                             cache_bust=cache_bust,
                             user_has_store=user_has_store,
                             current_store=current_store,
                             uploaded_filename=uploaded_filename)
        
    except Exception as e:
        logging.error(f"❌ CRITICAL ERROR in index route: {str(e)}")
        logging.error(f"Index route traceback: {traceback.format_exc()}")
        # Ensure cache_bust and store variables are always available
        try:
            import time
            cache_bust = f"v2.1.{int(time.time())}"  # Use timestamp for automatic cache busting
            user_has_store = False
            current_store = None
            uploaded_filename = ''
            # Try to render template with error message
            return render_template('index.html', error=str(e), cache_bust=cache_bust, user_has_store=user_has_store, current_store=current_store, uploaded_filename=uploaded_filename)
        except Exception as template_error:
            # If template rendering also fails, return a simple error page
            logging.error(f"❌ Template rendering also failed: {template_error}")
            return f"""
            <html>
            <head><title>Error</title></head>
            <body>
                <h1>Application Error</h1>
                <p>An error occurred while loading the page.</p>
                <p>Error: {str(e)}</p>
                <p>Please check the server logs for more details.</p>
            </body>
            </html>
            """, 500

@app.route('/splash')
def splash():
    """Serve the splash screen."""
    return render_template('splash.html')

@app.route('/debug-template')
def debug_template():
    """Debug route to test template loading."""
    return render_template('index.html', debug_message="DEBUG TEMPLATE ROUTE WORKING")

@app.route('/generation-splash')
def generation_splash():
    """Serve the generation splash screen."""
    return render_template('generation-splash.html')

@app.route('/test_upload.html')
def test_upload():
    """Test upload page"""
    return open('test_upload.html').read()
@app.route('/upload', methods=['POST'])
def upload_file():
    """Optimized file upload - saves file quickly then processes in background"""
    global _excel_processor  # ensure global declared before any assignment within this function
    start_time = time.time()
    
    try:
        logging.info("=== UPLOAD START ===")
        
        # DIAGNOSTIC: Log IP and session state
        ip_address = get_client_ip()
        session_store = session.get('selected_store')
        logging.info(f"🔍 Upload diagnostics: IP={ip_address}, Session store={session_store}")
        logging.info(f"🔍 Request headers: X-Forwarded-For={request.headers.get('X-Forwarded-For')}, X-Real-IP={request.headers.get('X-Real-IP')}, Remote-Addr={request.remote_addr}")
        
        # CRITICAL: Require store selection before upload
        # CRITICAL FIX: Use get_current_store_name with fallback instead of has_store_selection
        # has_store_selection can be too strict and fail even when store is selected
        selected_store = get_current_store_name(allow_fallback=True)
        if not selected_store:
            logging.error(f"❌ Upload attempted without store selection - IP: {ip_address}, Session: {session_store}")
            logging.error(f"❌ IP store selections: {list(_ip_store_selections.keys())}")
            return jsonify({'error': 'Please select a store before uploading files'}), 400
        
        # Validate request
        if 'file' not in request.files:
            logging.error("No file in request")
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if not file or file.filename == '':
            logging.error("Empty file or filename")
            return jsonify({'error': 'No file selected'}), 400
        
        # Get current store selection
        selected_store = get_current_store_name()
        ip_store = None
        with _ip_store_lock:
            if ip_address in _ip_store_selections:
                ip_store = _ip_store_selections[ip_address].get('store')
        
        logging.info(f"✅ Store selection found: {selected_store}")
        logging.info(f"🔍 Store diagnostics - Session: {session_store}, IP-based: {ip_store}, Final: {selected_store}")
        
        # CRITICAL FIX: Allow store override from request body if provided (for UI consistency)
        # CRITICAL FIX: Only try to get JSON if Content-Type is application/json (file uploads use multipart/form-data)
        request_store = request.form.get('store')
        if not request_store and request.is_json:
            try:
                json_data = request.get_json(silent=True)
                if json_data:
                    request_store = json_data.get('store')
            except Exception:
                pass  # Ignore JSON parsing errors for file uploads
        if request_store:
            logging.info(f"🔍 Request specifies store: {request_store}, current selected: {selected_store}")
            # If request store matches detected store in filename, use it
            detected_store_from_filename = extract_store_from_filename(file.filename)
            if detected_store_from_filename == request_store:
                logging.info(f"✅ Request store matches filename - using {request_store}")
                selected_store = request_store
                # Update session to match
                session['selected_store'] = request_store
                session['store_server_id'] = SERVER_INSTANCE_ID
                session.modified = True
        
        logging.info(f"Uploading: {file.filename} for store: {selected_store}")
        
        # Validate extension
        if not file.filename.lower().endswith(('.xlsx', '.xls')):
            return jsonify({'error': 'Only Excel files allowed'}), 400
        
        # Validate filename contains store name and matches selected store
        is_valid, warning_msg, detected_store = validate_excel_filename_for_store(file.filename, selected_store)
        
        if not is_valid:
            logging.error(f"Filename validation failed: {warning_msg}")
            return jsonify({
                'error': warning_msg,
                'filename': file.filename,
                'selected_store': selected_store,
                'detected_store': detected_store
            }), 400
        
        # CRITICAL: Clear old Excel data from cache and session before processing new file
        # This prevents old Excel data from persisting when new file is uploaded
        logging.info("🧹 Clearing old Excel data from cache and session...")
        try:
            # Get session ID for aggressive cache clearing
            from flask import session as flask_session
            session_id = flask_session.get('session_id') or request.cookies.get('session')

            # CRITICAL FIX: Aggressively clear ALL possible cache variations to prevent stale data
            # This handles race conditions where old cache might be served
            cache_patterns_to_clear = [
                'available_tags',
                'filter_options',
                'web_filter_options',
            ]

            for pattern in cache_patterns_to_clear:
                # Clear generic cache
                cache_key = get_session_cache_key(pattern)
                cache.delete(cache_key)
                logging.info(f"✅ Cleared {pattern} cache: {cache_key[:40]}...")

            # CRITICAL FIX: Clear file-specific cache if old file path exists
            old_file_path = session.get('file_path')
            if old_file_path:
                old_file_cache_key = get_session_cache_key(f'available_tags_{old_file_path}')
                cache.delete(old_file_cache_key)
                logging.info(f"✅ Cleared old file-specific cache: {old_file_cache_key[:50]}...")

            # Clear Excel processor cache - this is critical to prevent serving old data
            reset_excel_processor()
            logging.info("✅ Reset Excel processor")

            # Clear lineage update timestamp to force fresh lineage alignment
            if 'lineage_update_timestamp' in session:
                del session['lineage_update_timestamp']
                logging.info("✅ Cleared lineage_update_timestamp from session")

            # Clear selected tags (they're for old file)
            session['selected_tags'] = []
            logging.info("✅ Cleared selected tags from session")

        except Exception as cache_error:
            logging.warning(f"⚠️ Error clearing cache: {cache_error}")
        
        # Initialize warning tracking variables
        warning_to_return = warning_msg if warning_msg else None
        
        # Create uploads directory - use UPLOADS_DIR constant for consistent path handling
        # CRITICAL FIX: Use BASE_DIR-based path instead of os.getcwd() for Windows compatibility
        uploads_dir = UPLOADS_DIR
        os.makedirs(uploads_dir, exist_ok=True)
        
        # Save file with timestamp
        timestamp = int(time.time())
        safe_filename = f"{timestamp}_{file.filename}"
        file_path = os.path.join(uploads_dir, safe_filename)
        
        logging.info(f"🔍 Attempting to save file to: {file_path}")
        logging.info(f"🔍 Uploads directory: {uploads_dir}")
        logging.info(f"🔍 Directory exists: {os.path.exists(uploads_dir)}")
        
        try:
            file.save(file_path)
            logging.info(f"✅ File save completed: {file_path}")
        except Exception as save_error:
            logging.error(f"❌ File save failed: {save_error}")
            logging.error(traceback.format_exc())
            return jsonify({'error': f'File save failed: {str(save_error)}'}), 500
        
        # Verify saved
        if not os.path.exists(file_path):
            logging.error(f"❌ File does not exist after save: {file_path}")
            return jsonify({'error': 'File save failed - file not found after save'}), 500
        else:
            file_size = os.path.getsize(file_path)
            logging.info(f"✅ File verified: {file_path} ({file_size} bytes)")
        
        # Update session with permanent flag for persistence
        session.permanent = True
        # CRITICAL FIX: Normalize file path for Windows compatibility before storing in session
        normalized_session_path = os.path.normpath(file_path)
        session['file_path'] = normalized_session_path
        session['uploaded_filename'] = file.filename
        session['upload_timestamp'] = timestamp
        session['file_store'] = selected_store  # Store which store this file belongs to
        session.modified = True
        
        # CRITICAL: Force session save immediately
        try:
            from flask import session as flask_session
            if hasattr(flask_session, 'save'):
                flask_session.save()
        except:
            pass
        
        # CRITICAL: Also save to a persistent location to survive session issues
        try:
            import json
            # Use the same uploads_dir that was used to save the file
            persistence_file = os.path.join(uploads_dir, '.last_upload.json')
            # CRITICAL FIX: Normalize file path for Windows compatibility (use forward slashes or os.path.normpath)
            normalized_file_path = os.path.normpath(file_path)
            with open(persistence_file, 'w') as f:
                json.dump({
                    'file_path': normalized_file_path,
                    'filename': file.filename,
                    'timestamp': timestamp,
                    'store': selected_store
                }, f)
            logging.info(f"✅ Saved upload info to persistent file: {persistence_file}")
            logging.info(f"✅ Normalized file path stored: {normalized_file_path}")
        except Exception as persist_err:
            logging.warning(f"Could not save persistent upload info: {persist_err}")
            logging.error(traceback.format_exc())
        
        logging.info(f"✅ Session updated and saved: file_path={file_path}, filename={file.filename}, permanent={session.permanent}")
        logging.info(f"✅ Session data: {dict(session)}")
        
        # CRITICAL: Verify session was saved by reading it back
        verify_file_path = session.get('file_path')
        verify_filename = session.get('uploaded_filename')
        if verify_file_path == file_path and verify_filename == file.filename:
            logging.info(f"✅ SESSION VERIFIED: file_path and filename saved correctly")
        else:
            logging.error(f"❌ SESSION VERIFICATION FAILED: file_path={verify_file_path}, filename={verify_filename}")
            logging.error(f"❌ Expected file_path={file_path}, filename={file.filename}")
        
        # Mark as processing
        update_processing_status(file.filename, 'processing')
        
        # Check if we're on PythonAnywhere - if so, use background processing
        is_pythonanywhere = os.environ.get('PYTHONANYWHERE_DOMAIN') or os.environ.get('PYTHONANYWHERE_SITE')
        
        if is_pythonanywhere:
            # On PythonAnywhere: Start background thread to avoid timeout
            logging.info("[PYTHONANYWHERE] Starting background processing thread")

            # Capture variables from request context for background thread
            original_filename = file.filename
            # Store context removed - using single database
            # Capture session_id for preroll items update
            background_session_id = session.get('session_id', 'default')

            # PERFORMANCE FIX: Clear global processor immediately so frontend can load the file
            # CRITICAL: Also clear cache to ensure old Excel data doesn't persist
            _excel_processor = None
            logging.info("✅ Cleared Excel processor cache immediately for fast frontend access")
            
            # CRITICAL: Clear cache again after marking as ready to ensure old data is gone
            try:
                cache_key = get_session_cache_key('available_tags')
                cache.delete(cache_key)
                logging.info(f"✅ Cleared available_tags cache after marking ready: {cache_key}")
            except Exception as e:
                logging.warning(f"⚠️ Error clearing cache after marking ready: {e}")

            # PERFORMANCE FIX: Mark as ready immediately so frontend can start loading
            # Background processing will handle database storage and cache clearing
            update_processing_status(file.filename, 'ready')
            logging.info(f"✅ Marked {file.filename} as ready immediately for fast frontend response")

            # Create completion event for tracking
            import threading
            completion_event = threading.Event()

            def process_in_background():
                # CRITICAL FIX: Use application context for background thread
                # Flask's g object and session require application context
                with app.app_context():
                    bg_start_time = time.time()
                    max_bg_time = 300  # 5 minutes max for background processing
                    try:
                        logging.info(f"[BACKGROUND] Processing file: {file_path} for store: {selected_store}")

                        # CRITICAL FIX: Create processor with store name in background thread
                        # Don't use get_excel_processor() as it might not have the correct context
                        from src.core.data.excel_processor import ExcelProcessor
                        processor = ExcelProcessor(store_name=selected_store)
                        logging.info(f"[BACKGROUND] Created ExcelProcessor with store: {selected_store}")
                        
                        # PERFORMANCE FIX: Skip expensive database operations during upload
                        processor._skip_database_strain = True
                        logging.info("[BACKGROUND] ⚡ Skipping database Product Strain application for fast upload")
                        
                        # Load the file with timeout check
                        if time.time() - bg_start_time > max_bg_time:
                            raise TimeoutError(f"Background processing exceeded {max_bg_time}s")

                        success = processor.load_file(file_path)

                        if success:
                            row_count = len(processor.df) if hasattr(processor, 'df') and processor.df is not None else 0
                            logging.info(f"[BACKGROUND] File loaded: {row_count} rows")

                            if row_count == 0:
                                logging.error(f"[BACKGROUND] ⚠️ WARNING: File loaded but has 0 rows! This might indicate a problem.")
                                logging.error(f"[BACKGROUND] DataFrame info: df is None={processor.df is None}, df.empty={processor.df.empty if processor.df is not None else 'N/A'}")
                                if processor.df is not None and not processor.df.empty:
                                    logging.error(f"[BACKGROUND] DataFrame actually has {len(processor.df)} rows - row_count calculation was wrong")

                            # CRITICAL PERFORMANCE: Update global processor FIRST so frontend can access DataFrame immediately
                            global _excel_processor
                            _excel_processor = processor
                            logging.info("[BACKGROUND] ✅ Updated global Excel processor IMMEDIATELY - DataFrame now accessible")

                            # ⚡ CRITICAL FIX: Cache tags IMMEDIATELY before expensive database operations
                            # This allows frontend to show tags in 1-2 seconds instead of waiting 18+ seconds
                            try:
                                cache_start = time.time()
                                logging.info("[BACKGROUND] ⚡ PRIORITY: Caching tags BEFORE database operations...")

                                # Skip enrichment for maximum speed
                                if hasattr(processor, '_skip_enrichment'):
                                    processor._skip_enrichment = True

                                # Get tags without filters for maximum speed
                                tags = processor.get_available_tags(filters=None)

                                # Reset enrichment flag
                                if hasattr(processor, '_skip_enrichment'):
                                    processor._skip_enrichment = False

                                # Make JSON safe and cache
                                safe_tags = make_json_safe(tags)

                                # CRITICAL FIX: Use file-path-only cache key so frontend can access it
                                # Background thread doesn't have same session context as frontend request
                                import hashlib
                                cache_key = f"tags_file_{hashlib.sha256(file_path.encode()).hexdigest()}"
                                cache.set(cache_key, safe_tags, timeout=300)

                                cache_elapsed = (time.time() - cache_start) * 1000
                                logging.info(f"[BACKGROUND] ✅ Cached {len(safe_tags)} tags with key={cache_key[:16]}... ({cache_elapsed:.0f}ms)")

                                # CRITICAL FIX: Mark as 'ready' immediately after caching tags
                                # Frontend needs 'ready' status to proceed, not 'tags_ready'
                                update_processing_status(original_filename, 'ready')
                                logging.info(f"[BACKGROUND] ✅ Marked {original_filename} as READY (tags cached)")
                            except Exception as cache_error:
                                logging.warning(f"[BACKGROUND] ⚠️ Could not pre-cache tags: {cache_error}")
                                logging.error(traceback.format_exc())
                                # Even if tag caching fails, mark as ready so frontend can try to load
                                update_processing_status(original_filename, 'ready')

                            # NOW do expensive database operations AFTER tags are cached
                            # Frontend doesn't need to wait for these
                            logging.info("[BACKGROUND] 🔄 Starting database operations (frontend already has tags)...")

                            # Update preroll items from newly loaded Excel data
                            if hasattr(processor, 'df') and processor.df is not None:
                                # Pass session_id from outer scope since we're in background thread
                                update_preroll_items_from_excel(processor.df, session_id=background_session_id)

                            # Store in database (this is slow but frontend doesn't need to wait for it)
                            try:
                                db_start = time.time()
                                # Use the selected_store from outer scope
                                store_name = selected_store
                                product_db = get_product_database(store_name)

                                if product_db and hasattr(product_db, 'store_excel_data'):
                                    logging.info(f"[BACKGROUND] Storing {row_count} products in database for store: {store_name}...")
                                    logging.info(f"[BACKGROUND] Database path: {product_db.db_path}")
                                    result = product_db.store_excel_data(processor.df, file_path)
                                    db_elapsed = (time.time() - db_start) * 1000
                                    
                                    # Enhanced logging with detailed results
                                    stored = result.get('stored', 0)
                                    updated = result.get('updated', 0)
                                    errors = result.get('errors', 0)
                                    excluded = result.get('excluded_json_matches', 0)
                                    
                                    logging.info(f"[BACKGROUND] ✅ Database storage complete ({db_elapsed:.0f}ms)")
                                    logging.info(f"[BACKGROUND]    Stored: {stored}, Updated: {updated}, Errors: {errors}, Excluded: {excluded}")
                                    
                                    if errors > 0:
                                        logging.error(f"[BACKGROUND] ⚠️ Database storage had {errors} errors - check logs above for details")
                                    if stored == 0 and updated == 0 and row_count > 0:
                                        logging.warning(f"[BACKGROUND] ⚠️ No products were stored/updated despite {row_count} rows in Excel file")
                                        logging.warning(f"[BACKGROUND]    This might indicate a sync issue - check product validation logic")
                                else:
                                    logging.error(f"[BACKGROUND] ❌ Cannot store products: product_db={product_db}, has_method={hasattr(product_db, 'store_excel_data') if product_db else False}")
                            except Exception as db_error:
                                logging.error(f"[BACKGROUND] ❌ Database storage failed: {db_error}")
                                import traceback
                                logging.error(f"[BACKGROUND] Traceback: {traceback.format_exc()}")

                            logging.info("[BACKGROUND] ✅ Excel processor cache cleared")

                            # CRITICAL: Clear old caches but PRESERVE the new tag cache we just created
                            try:
                                # Clear non-file-specific caches that need refresh
                                cache_keys_to_clear = [
                                    'selected_tags',
                                    'vendor_tags',
                                    'initial_data'
                                ]
                                for key_base in cache_keys_to_clear:
                                    cache_key = get_session_cache_key(key_base)
                                    cache.delete(cache_key)
                                    logging.info(f"[BACKGROUND] ✅ Cleared cache: {key_base}")

                                # NOTE: We do NOT clear f'available_tags_{file_path}' because we just cached it above
                                logging.info(f"[BACKGROUND] ✅ Preserved new file tag cache: available_tags_{file_path}")
                            except Exception as cache_err:
                                logging.warning(f"[BACKGROUND] Failed to clear cache: {cache_err}")

                            # Mark processing as complete
                            logging.info(f"[BACKGROUND] Processing complete for {original_filename}")
                            
                            # Signal that processing is complete (use closure variable)
                            completion_event.set()
                        else:
                            logging.error("[BACKGROUND] File load returned False")
                            logging.error(f"[BACKGROUND] File path: {file_path}")
                            logging.error(f"[BACKGROUND] File exists: {os.path.exists(file_path) if file_path else 'N/A'}")
                            if file_path and os.path.exists(file_path):
                                file_size = os.path.getsize(file_path)
                                logging.error(f"[BACKGROUND] File size: {file_size} bytes")
                            update_processing_status(original_filename, 'error: File load failed')
                            # Signal completion even on error so we don't wait forever
                            completion_event.set()

                    except Exception as e:
                        logging.error(f"[BACKGROUND] Processing error: {e}")
                        logging.error(f"[BACKGROUND] Error type: {type(e).__name__}")
                        logging.error(traceback.format_exc())
                        update_processing_status(original_filename, f'error: {str(e)}')
                        # Signal completion even on error so we don't wait forever
                        completion_event.set()
                    finally:
                        # CRITICAL SAFETY NET: Ensure status is never left as 'processing'
                        # If we get here and status is still 'processing', mark as ready if we have data
                        with processing_lock:
                            current_status = processing_status.get(original_filename, 'unknown')
                            if current_status == 'processing':
                                # Check if we actually loaded data successfully
                                if _excel_processor and hasattr(_excel_processor, 'df') and _excel_processor.df is not None and not _excel_processor.df.empty:
                                    logging.warning(f"[BACKGROUND] SAFETY NET: Status still 'processing' but data loaded - marking as ready")
                                    update_processing_status(original_filename, 'ready')
                                else:
                                    logging.error(f"[BACKGROUND] SAFETY NET: Status still 'processing' and no data - marking as error")
                                    update_processing_status(original_filename, 'error: Processing incomplete')
                                completion_event.set()

                        bg_elapsed = time.time() - bg_start_time
                        logging.info(f"[BACKGROUND] Thread completed in {bg_elapsed:.1f}s")

            # Start background thread with completion tracking
            thread = threading.Thread(target=process_in_background)
            thread.daemon = True
            thread.start()
            
            # CRITICAL PERFORMANCE FIX: Don't wait for processing - return immediately
            # Frontend will poll /api/upload-status and request tags when ready
            # This prevents upload endpoint from timing out on large files
            logging.info("✅ Background processing started - returning immediately")
            completed = False  # Mark as not completed to return "processing" status
            
            if completed:
                upload_time = time.time() - start_time
                logging.info(f"=== UPLOAD COMPLETE (processing finished): {upload_time:.3f}s ===")
                response_data = {
                    'success': True,
                    'message': 'File uploaded and ready',
                    'filename': file.filename,
                    'rows': len(_excel_processor.df) if _excel_processor and hasattr(_excel_processor, 'df') and _excel_processor.df is not None else 0,
                    'processing': False
                }
            else:
                # Processing taking longer than expected - return anyway but indicate it's still processing
                upload_time = time.time() - start_time
                logging.info(f"=== UPLOAD INITIATED (processing in background): {upload_time:.3f}s ===")
                response_data = {
                    'success': True,
                    'message': 'File uploaded, processing in background',
                    'filename': file.filename,
                    'processing': True  # Frontend should poll for completion
                }
            if warning_to_return:
                response_data['warning'] = warning_to_return
                response_data['detected_store'] = detected_store
                response_data['selected_store'] = selected_store
            return jsonify(response_data)
            
        else:
            # Local development: ULTRA-FAST UPLOAD - return immediately, load in background
            logging.info("[LOCAL] Ultra-fast upload mode - saving file only (background processing)")

            # Create completion event for local processing
            import threading
            completion_event_local = threading.Event()

            # CRITICAL FIX: Process file in background thread for local development too
            def process_in_background_local():
                with app.app_context():
                    try:
                        logging.info(f"[LOCAL-BACKGROUND] Processing file: {file_path} for store: {selected_store}")
                        
                        # Create processor with store name
                        from src.core.data.excel_processor import ExcelProcessor
                        processor = ExcelProcessor(store_name=selected_store)
                        logging.info(f"[LOCAL-BACKGROUND] Created ExcelProcessor with store: {selected_store}")
                        
                        # Load file
                        success = processor.load_file(file_path)
                        
                        if success:
                            row_count = len(processor.df) if hasattr(processor, 'df') and processor.df is not None else 0
                            logging.info(f"[LOCAL-BACKGROUND] File loaded: {row_count} rows")

                            # Update global processor FIRST
                            global _excel_processor
                            _excel_processor = processor
                            logging.info(f"[LOCAL-BACKGROUND] ✅ Updated global Excel processor with {row_count} rows")

                            # ⚡ CRITICAL FIX: Cache tags IMMEDIATELY before expensive database operations
                            try:
                                cache_start = time.time()
                                logging.info("[LOCAL-BACKGROUND] ⚡ PRIORITY: Caching tags BEFORE database operations...")

                                # Skip enrichment for maximum speed
                                if hasattr(processor, '_skip_enrichment'):
                                    processor._skip_enrichment = True

                                tags = processor.get_available_tags(filters=None)

                                # Reset enrichment flag
                                if hasattr(processor, '_skip_enrichment'):
                                    processor._skip_enrichment = False

                                safe_tags = make_json_safe(tags)

                                # CRITICAL FIX: Use file-path-only cache key so frontend can access it
                                # Background thread doesn't have same session context as frontend request
                                import hashlib
                                cache_key = f"tags_file_{hashlib.sha256(file_path.encode()).hexdigest()}"
                                cache.set(cache_key, safe_tags, timeout=300)

                                cache_elapsed = (time.time() - cache_start) * 1000
                                logging.info(f"[LOCAL-BACKGROUND] ✅ Cached {len(safe_tags)} tags with key={cache_key[:16]}... ({cache_elapsed:.0f}ms)")
                            except Exception as cache_error:
                                logging.warning(f"[LOCAL-BACKGROUND] ⚠️ Could not pre-cache tags: {cache_error}")

                            # Signal completion early so frontend can proceed
                            completion_event_local.set()

                            # NOW do expensive database operations AFTER tags are cached and completion signaled
                            logging.info("[LOCAL-BACKGROUND] 🔄 Starting database operations (frontend already has tags)...")
                            try:
                                db_start = time.time()
                                product_db = get_product_database(selected_store)
                                if product_db and hasattr(product_db, 'store_excel_data'):
                                    logging.info(f"[LOCAL-BACKGROUND] Storing {row_count} products in database for store: {selected_store}...")
                                    logging.info(f"[LOCAL-BACKGROUND] Database path: {product_db.db_path}")
                                    result = product_db.store_excel_data(processor.df, file_path)
                                    db_elapsed = (time.time() - db_start) * 1000
                                    
                                    # Enhanced logging with detailed results
                                    stored = result.get('stored', 0)
                                    updated = result.get('updated', 0)
                                    errors = result.get('errors', 0)
                                    excluded = result.get('excluded_json_matches', 0)
                                    
                                    logging.info(f"[LOCAL-BACKGROUND] ✅ Database storage complete ({db_elapsed:.0f}ms)")
                                    logging.info(f"[LOCAL-BACKGROUND]    Stored: {stored}, Updated: {updated}, Errors: {errors}, Excluded: {excluded}")
                                    
                                    if errors > 0:
                                        logging.error(f"[LOCAL-BACKGROUND] ⚠️ Database storage had {errors} errors - check logs above for details")
                                    if stored == 0 and updated == 0 and row_count > 0:
                                        logging.warning(f"[LOCAL-BACKGROUND] ⚠️ No products were stored/updated despite {row_count} rows in Excel file")
                                        logging.warning(f"[LOCAL-BACKGROUND]    This might indicate a sync issue - check product validation logic")
                                else:
                                    logging.error(f"[LOCAL-BACKGROUND] ❌ Cannot store products: product_db={product_db}, has_method={hasattr(product_db, 'store_excel_data') if product_db else False}")
                            except Exception as db_error:
                                logging.error(f"[LOCAL-BACKGROUND] ❌ Database storage failed: {db_error}")
                                import traceback
                                logging.error(f"[LOCAL-BACKGROUND] Traceback: {traceback.format_exc()}")
                        else:
                            logging.error("[LOCAL-BACKGROUND] File load returned False")
                            # Signal completion even on error so we don't wait forever
                            completion_event_local.set()
                    except Exception as e:
                        logging.error(f"[LOCAL-BACKGROUND] Processing error: {e}")
                        logging.error(traceback.format_exc())
                        # Signal completion even on error so we don't wait forever
                        completion_event_local.set()
            
            # Start background thread for local development with completion tracking
            thread = threading.Thread(target=process_in_background_local)
            thread.daemon = True
            thread.start()
            
            # PERFORMANCE FIX: Wait briefly for processing to complete (max 5 seconds)
            # This ensures tags are available when frontend tries to load them
            logging.info("⏳ Waiting for local background processing to complete (max 5s)...")
            completed_local = completion_event_local.wait(timeout=5.0)
            
            if completed_local:
                upload_time_local = time.time() - start_time
                logging.info(f"=== LOCAL UPLOAD COMPLETE (processing finished): {upload_time_local:.3f}s ===")
                response_data = {
                    'success': True,
                    'message': 'File uploaded successfully',
                    'filename': file.filename,
                    'rows': len(_excel_processor.df) if _excel_processor and hasattr(_excel_processor, 'df') and _excel_processor.df is not None else 0,
                    'processing': False
                }
            else:
                upload_time_local = time.time() - start_time
                logging.info(f"=== LOCAL UPLOAD INITIATED (processing in background): {upload_time_local:.3f}s ===")
                response_data = {
                    'success': True,
                    'message': 'File uploaded, processing in background',
                    'filename': file.filename,
                    'processing': True
                }
            
            # Don't clear processor - background thread will update it
            logging.info("✅ Background processing started - processor will be available when complete")

            # Mark file as ready
            update_processing_status(file.filename, 'ready')
            logging.info(f"✅ Marked {file.filename} as ready")

            # CRITICAL: Clear old caches but preserve the new file's tag cache
            try:
                # Clear non-file-specific caches that need refresh
                cache_keys_to_clear = [
                    'selected_tags',
                    'vendor_tags',
                    'initial_data'
                ]

                # Also clear any old available_tags caches (from previous uploads)
                # Try to delete with empty path (legacy cache key)
                try:
                    legacy_cache_key = get_session_cache_key('available_tags_')
                    cache.delete(legacy_cache_key)
                    logging.info(f"✅ Cleared legacy cache: available_tags_")
                except:
                    pass

                for key_base in cache_keys_to_clear:
                    cache_key = get_session_cache_key(key_base)
                    cache.delete(cache_key)
                    logging.info(f"✅ Cleared cache: {key_base}")

                # NOTE: We do NOT clear f'available_tags_{file_path}' here because we just cached it above
                logging.info(f"✅ Preserved new file tag cache: available_tags_{file_path}")
            except Exception as cache_err:
                logging.warning(f"Failed to clear cache: {cache_err}")
            
            # CRITICAL: Verify session file path is set correctly
            logging.info(f"✅ Session file_path after upload: {session.get('file_path')}")
            logging.info(f"✅ Uploaded file saved at: {file_path}")
            
            # Add warnings if any
            if warning_to_return:
                response_data['warning'] = warning_to_return
                response_data['detected_store'] = detected_store
                response_data['selected_store'] = selected_store
            return jsonify(response_data)
        
    except Exception as e:
        logging.error(f"Upload failed: {e}")
        logging.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500


# EXCEL UPLOAD PERFORMANCE OPTIMIZATION: Add ultra-fast streaming upload
@app.route('/upload-streaming', methods=['POST'])
def upload_file_streaming():
    """Ultra-fast streaming Excel upload with chunked processing for maximum performance"""
    try:
        start_time = time.time()
        max_processing_time = MAX_TOTAL_PROCESSING_TIME if 'MAX_TOTAL_PROCESSING_TIME' in globals() else 300
        max_processing_time = MAX_TOTAL_PROCESSING_TIME if 'MAX_TOTAL_PROCESSING_TIME' in globals() else 300
        
        # CRITICAL FIX: Use get_current_store_name with fallback instead of has_store_selection
        # has_store_selection can be too strict and fail even when store is selected
        selected_store = get_current_store_name(allow_fallback=True)
        if not selected_store:
            logging.error("Upload attempted without store selection")
            return jsonify({'error': 'Please select a store before uploading files'}), 400
        
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Validate filename contains store name and matches selected store
        is_valid, warning_msg, detected_store = validate_excel_filename_for_store(file.filename, selected_store)
        
        if not is_valid:
            logging.error(f"Filename validation failed: {warning_msg}")
            return jsonify({
                'error': warning_msg,
                'filename': file.filename,
                'selected_store': selected_store,
                'detected_store': detected_store
            }), 400
        
        # Generate unique filename
        filename = secure_filename(file.filename)
        timestamp = int(time.time())
        unique_filename = f"upload_{timestamp}_{filename}"
        temp_path = os.path.join('uploads', unique_filename)
        
        # Ensure uploads directory exists
        os.makedirs('uploads', exist_ok=True)
        
        # Save file
        file.save(temp_path)
        
        # STREAMING OPTIMIZATION: Process file in chunks for large files
        file_size = os.path.getsize(temp_path)
        is_large_file = file_size > 5 * 1024 * 1024  # 5MB threshold
        
        if ENHANCED_LOGGING_AVAILABLE:
            enhanced_logger.log_info("Starting streaming Excel upload", 
                                   {'filename': filename, 'size_mb': file_size / (1024*1024), 'large_file': is_large_file})
        else:
            logging.info(f"Starting streaming Excel upload: {filename} ({file_size / (1024*1024):.1f}MB)")
        
        if is_large_file:
            # Use streaming processing for large files
            return process_large_file_streaming(temp_path, filename, start_time)
        else:
            # Use optimized processing for smaller files
            return process_small_file_optimized(temp_path, filename, start_time)
            
    except Exception as e:
        if ENHANCED_LOGGING_AVAILABLE:
            log_file_processing_error(temp_path, 'streaming_upload', e)
        else:
            logging.error(f"Streaming upload error: {e}")
        return jsonify({'error': str(e)}), 500

def process_large_file_streaming(temp_path: str, filename: str, start_time: float):
    """Process large Excel files using streaming/chunked approach"""
    try:
        from src.core.data.excel_processor import ExcelProcessor
        import pandas as pd
        
        # STREAMING OPTIMIZATION: Process in chunks
        chunk_size = 1000  # Process 1000 rows at a time
        total_rows = 0
        processed_chunks = 0
        
        # First, get total row count
        try:
            # Quick row count without loading full data
            sample_df = pd.read_excel(temp_path, nrows=0)
            total_columns = len(sample_df.columns)
            
            # Estimate total rows (rough estimate)
            with open(temp_path, 'rb') as f:
                file_content = f.read()
                # Rough estimation: each row is approximately 200 bytes
                estimated_rows = len(file_content) // 200
                total_rows = min(estimated_rows, 100000)  # Cap at 100k for safety
                
        except Exception as e:
            logging.warning(f"Could not estimate rows: {e}")
            total_rows = 10000  # Default estimate
        
        if ENHANCED_LOGGING_AVAILABLE:
            enhanced_logger.log_info("Starting chunked processing", 
                                   {'estimated_rows': total_rows, 'chunk_size': chunk_size})
        else:
            logging.info(f"Starting chunked processing: ~{total_rows} rows in chunks of {chunk_size}")
        
        # STREAMING OPTIMIZATION: Process file in chunks
        processor = ExcelProcessor()
        all_chunks = []
        
        try:
            # Read file in chunks
            for chunk_start in range(0, total_rows, chunk_size):
                chunk_df = pd.read_excel(
                    temp_path,
                    skiprows=chunk_start,
                    nrows=chunk_size,
                    dtype=str,  # Read as strings for speed
                    na_filter=False,
                    engine='openpyxl'
                )
                
                if chunk_df.empty:
                    break
                
                all_chunks.append(chunk_df)
                processed_chunks += 1
                
                # Log progress every 10 chunks
                if processed_chunks % 10 == 0:
                    progress = (chunk_start / total_rows) * 100
                    logging.info(f"Streaming progress: {progress:.1f}% ({processed_chunks} chunks)")
                
                # PERFORMANCE: More frequent garbage collection for large files
                if processed_chunks % 10 == 0:
                    import gc
                    gc.collect()
            
            # Combine all chunks
            if all_chunks:
                processor.df = pd.concat(all_chunks, ignore_index=True)
                processor.df.reset_index(drop=True, inplace=True)
                
                # STREAMING OPTIMIZATION: Minimal processing for speed
                processor.df = processor.df.dropna(subset=['Product Name*'], how='all')
                
                # Update global processor
                global _excel_processor
                with excel_processor_lock:
                    _excel_processor = processor
                    _excel_processor._last_loaded_file = temp_path
                
                processing_time = time.time() - start_time
                
                if ENHANCED_LOGGING_AVAILABLE:
                    enhanced_logger.log_success("Streaming upload completed", 
                                              {'rows': len(processor.df), 'chunks': processed_chunks, 
                                               'time_seconds': processing_time})
                else:
                    logging.info(f"✅ Streaming upload completed: {len(processor.df)} rows in {processing_time:.3f}s")
                
                return jsonify({
                    'success': True,
                    'message': f'Large file processed successfully using streaming (chunks: {processed_chunks})',
                    'rows_loaded': len(processor.df),
                    'processing_time': processing_time,
                    'method': 'streaming',
                    'chunks_processed': processed_chunks
                })
            else:
                return jsonify({'error': 'No data found in file'}), 400
                
        except Exception as e:
            logging.error(f"Chunked processing failed: {e}")
            # Fallback to regular processing
            return process_small_file_optimized(temp_path, filename, start_time)
            
    except Exception as e:
        if ENHANCED_LOGGING_AVAILABLE:
            log_file_processing_error(temp_path, 'large_file_streaming', e)
        else:
            logging.error(f"Large file streaming error: {e}")
        return jsonify({'error': str(e)}), 500

def process_small_file_optimized(temp_path: str, filename: str, start_time: float):
    """Process smaller Excel files with optimized settings"""
    try:
        from src.core.data.excel_processor import ExcelProcessor
        
        processor = ExcelProcessor()
        
        # OPTIMIZATION: Use platform-specific loading
        import platform
        is_windows = platform.system() == 'Windows'
        
        if is_windows and hasattr(processor, 'load_file_optimized_windows'):
            # Use Windows-optimized loading
            success = processor.load_file_optimized_windows(temp_path)
            method = 'windows_optimized'
        else:
            # Use regular optimized loading
            success = processor.load_file(temp_path)
            method = 'standard_optimized'
        
        if not success or processor.df is None or processor.df.empty:
            return jsonify({'error': 'Failed to process file or file is empty'}), 400
        
        # Update global processor
        global _excel_processor
        with excel_processor_lock:
            _excel_processor = processor
            _excel_processor._last_loaded_file = temp_path
        
        processing_time = time.time() - start_time
        
        if ENHANCED_LOGGING_AVAILABLE:
            enhanced_logger.log_success("Small file upload completed", 
                                      {'rows': len(processor.df), 'method': method, 
                                       'time_seconds': processing_time})
        else:
            logging.info(f"✅ Small file upload completed: {len(processor.df)} rows in {processing_time:.3f}s")
        
        return jsonify({
            'success': True,
            'message': f'File processed successfully ({method})',
            'rows_loaded': len(processor.df),
            'processing_time': processing_time,
            'method': method
        })
        
    except Exception as e:
        if ENHANCED_LOGGING_AVAILABLE:
            log_file_processing_error(temp_path, 'small_file_optimized', e)
        else:
            logging.error(f"Small file optimization error: {e}")
        return jsonify({'error': str(e)}), 500
@app.route('/upload-pythonanywhere', methods=['POST'])
def upload_file_simple_pythonanywhere():
    """INSTANT upload endpoint - saves file and returns immediately, processes in background."""
    try:
        request_start = time.time()
        logging.info("=== INSTANT UPLOAD START ===")

        # CRITICAL FIX: Use get_current_store_name with fallback instead of has_store_selection
        # has_store_selection can be too strict and fail even when store is selected
        selected_store = get_current_store_name(allow_fallback=True)
        if not selected_store:
            logging.error("Upload attempted without store selection")
            return jsonify({'error': 'Please select a store before uploading files'}), 400

        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not file.filename.lower().endswith('.xlsx'):
            return jsonify({'error': 'Only .xlsx files are allowed'}), 400

        is_valid, warning_msg, detected_store = validate_excel_filename_for_store(file.filename, selected_store)
        if not is_valid:
            logging.error(f"Filename validation failed: {warning_msg}")
            return jsonify({
                'error': warning_msg,
                'filename': file.filename,
                'selected_store': selected_store,
                'detected_store': detected_store
            }), 400

        import tempfile
        save_start = time.time()
        sanitized_filename = sanitize_filename(file.filename)
        temp_path = os.path.join(tempfile.gettempdir(), f"upload_{sanitized_filename}")
        file.save(temp_path)
        save_duration = time.time() - save_start
        logging.info(f"[INSTANT] File saved in {save_duration:.2f}s -> {temp_path}")

        session['file_path'] = temp_path
        session['uploaded_filename'] = sanitized_filename
        session['selected_tags'] = []
        session.modified = True
        update_processing_status(file.filename, 'processing')

        from flask import copy_current_request_context
        import threading

        @copy_current_request_context
        def process_file_in_background():
            bg_start = time.time()
            try:
                from src.core.data.excel_processor import ExcelProcessor
                processor = ExcelProcessor()

                if hasattr(processor, 'enable_product_db_integration'):
                    processor.enable_product_db_integration(True)

                import os
                file_size_mb = os.path.getsize(temp_path) / (1024 * 1024)
                is_pythonanywhere = IS_PYTHONANYWHERE or PYTHONANYWHERE_OPTIMIZATION

                import pandas as pd
                success = False

                if is_pythonanywhere and hasattr(processor, 'pythonanywhere_fast_load'):
                    try:
                        success = processor.pythonanywhere_fast_load(temp_path)
                    except Exception as e:
                        logging.warning(f"PythonAnywhere fast load failed: {e}")

                if not success and file_size_mb > 10 and hasattr(processor, 'minimal_load_file'):
                    try:
                        success = processor.minimal_load_file(temp_path)
                    except Exception as e:
                        logging.warning(f"Minimal load failed: {e}")

                if not success and hasattr(processor, 'ultra_fast_load'):
                    try:
                        success = processor.ultra_fast_load(temp_path)
                    except Exception as e:
                        logging.warning(f"Ultra-fast load failed: {e}")

                if not success and hasattr(processor, 'fast_load_file'):
                    try:
                        success = processor.fast_load_file(temp_path)
                    except Exception as e:
                        logging.warning(f"Fast load failed: {e}")

                if not success:
                    try:
                        df = pd.read_excel(
                            temp_path,
                            engine='openpyxl',
                            dtype=str,
                            na_filter=False,
                            keep_default_na=False,
                            converters=None,
                            header=0
                        )
                        if not df.empty:
                            processor.df = df
                            success = True
                    except Exception as e:
                        logging.warning(f"Optimized pandas load failed: {e}")

                if not success:
                    try:
                        success = processor.load_file(temp_path)
                    except Exception as e:
                        logging.warning(f"Standard load failed: {e}")

                if not success or processor.df is None or processor.df.empty:
                    update_processing_status(file.filename, 'error: load_failed')
                    return

                global _excel_processor
                with excel_processor_lock:
                    _excel_processor = processor
                    _excel_processor._last_loaded_file = temp_path

                try:
                    tags = processor.get_available_tags(filters=None)
                    safe_tags = make_json_safe(tags)
                    cache_key = get_session_cache_key(f'available_tags_{temp_path}')
                    cache.set(cache_key, safe_tags, timeout=300)
                except Exception as cache_error:
                    logging.warning(f"[UPLOAD][BG-CACHE] Could not cache tags: {cache_error}")

                try:
                    if hasattr(processor, '_store_upload_in_database'):
                        processor._store_upload_in_database(processor.df, temp_path)
                except Exception as db_err:
                    logging.warning(f"[UPLOAD][BG-DB] Database storage failed: {db_err}")

                update_processing_status(file.filename, 'ready')
                logging.info(f"[BG] Completed background processing in {time.time() - bg_start:.2f}s")
            except Exception as process_error:
                logging.error(f"[BG] Processing failed: {process_error}")
                update_processing_status(file.filename, f'error: {process_error}')
            finally:
                try:
                    os.remove(temp_path)
                except Exception:
                    pass

        threading.Thread(target=process_file_in_background, daemon=True).start()

        total_duration = time.time() - request_start
        return jsonify({
            'message': 'File uploaded successfully, processing in background',
            'filename': sanitized_filename,
            'status': 'processing',
            'upload_time_seconds': total_duration
        })

    except Exception as e:
        logging.error(f"Upload error: {e}")
        return jsonify({'error': f'Upload failed: {str(e)}'}), 500

@app.route('/upload-instant', methods=['POST'])
def upload_instant():
    """INSTANT upload - saves file and returns immediately, all processing in background."""
    try:
        request_start = time.time()
        logging.info("=== INSTANT UPLOAD START ===")

        # Validate store selection
        # CRITICAL FIX: Use get_current_store_name with fallback instead of has_store_selection
        selected_store = get_current_store_name(allow_fallback=True)
        if not selected_store:
            return jsonify({'error': 'Please select a store before uploading files'}), 400

        # Validate file
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400

        file = request.files['file']
        if not file or file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not file.filename.lower().endswith('.xlsx'):
            return jsonify({'error': 'Only .xlsx files are allowed'}), 400

        # Validate filename
        is_valid, warning_msg, detected_store = validate_excel_filename_for_store(file.filename, selected_store)
        if not is_valid:
            return jsonify({
                'error': warning_msg,
                'filename': file.filename,
                'selected_store': selected_store,
                'detected_store': detected_store
            }), 400

        # Save file (ONLY blocking operation)
        import tempfile
        sanitized_filename = sanitize_filename(file.filename)
        temp_path = os.path.join(tempfile.gettempdir(), f"upload_{sanitized_filename}")
        save_start = time.time()
        file.save(temp_path)
        save_duration = time.time() - save_start

        # Update session
        session['file_path'] = temp_path
        session['uploaded_filename'] = sanitized_filename
        session['selected_tags'] = []
        session.modified = True

        # Start background processing
        from flask import copy_current_request_context
        import threading

        @copy_current_request_context
        def process_in_background():
            try:
                logging.info(f"[BG] Processing {temp_path}")
                from src.core.data.excel_processor import ExcelProcessor
                processor = ExcelProcessor()

                # Load file
                success = processor.load_file(temp_path)
                if success and processor.df is not None:
                    # Update global processor
                    global _excel_processor
                    with excel_processor_lock:
                        _excel_processor = processor
                        _excel_processor._last_loaded_file = temp_path

                    # Cache tags
                    try:
                        tags = processor.get_available_tags(filters=None)
                        safe_tags = make_json_safe(tags)
                        cache_key = get_session_cache_key(f'available_tags_{temp_path}')
                        cache.set(cache_key, safe_tags, timeout=300)
                        logging.info(f"[BG] ✅ Cached {len(safe_tags)} tags")
                    except Exception as e:
                        logging.warning(f"[BG] Tag caching failed: {e}")

                    # Store in database (optional, slow)
                    if hasattr(processor, '_store_upload_in_database'):
                        try:
                            processor._store_upload_in_database(processor.df, temp_path)
                            logging.info(f"[BG] ✅ Stored in database")
                        except Exception as e:
                            logging.warning(f"[BG] Database storage failed: {e}")

                    logging.info(f"[BG] ✅ Complete: {len(processor.df)} rows")
                else:
                    logging.error(f"[BG] File load failed")
            except Exception as e:
                logging.error(f"[BG] Processing failed: {e}")

        # Start thread and return IMMEDIATELY
        threading.Thread(target=process_in_background, daemon=True).start()

        total_duration = time.time() - request_start
        logging.info(f"✅ INSTANT upload response in {total_duration:.3f}s")

        return jsonify({
            'message': 'File uploaded successfully, processing in background',
            'filename': sanitized_filename,
            'status': 'processing',
            'upload_time_seconds': total_duration
        })

    except Exception as e:
        logging.error(f"Upload error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/upload-simple', methods=['POST'])
def upload_file_simple():
    """Simple, reliable file upload for PythonAnywhere"""
    try:
        logging.info("=== SIMPLE UPLOAD REQUEST START ===")
        start_time = time.time()
        
        # CRITICAL FIX: Use get_current_store_name with fallback instead of has_store_selection
        # has_store_selection can be too strict and fail even when store is selected
        selected_store = get_current_store_name(allow_fallback=True)
        if not selected_store:
            logging.error("Upload attempted without store selection")
            return jsonify({'error': 'Please select a store before uploading files'}), 400
        
        if 'file' not in request.files:
            logging.error("No file uploaded")
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        if file.filename == '':
            logging.error("No file selected")
            return jsonify({'error': 'No file selected'}), 400
        
        if not file.filename.lower().endswith('.xlsx'):
            logging.error(f"Invalid file type: {file.filename}")
            return jsonify({'error': 'Only .xlsx files are allowed'}), 400
        
        # Validate filename contains store name and matches selected store
        is_valid, warning_msg, detected_store = validate_excel_filename_for_store(file.filename, selected_store)
        
        if not is_valid:
            logging.error(f"Filename validation failed: {warning_msg}")
            return jsonify({
                'error': warning_msg,
                'filename': file.filename,
                'selected_store': selected_store,
                'detected_store': detected_store
            }), 400
        
        # Check file size
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > app.config['MAX_CONTENT_LENGTH']:
            logging.error(f"File too large: {file_size} bytes (max: {app.config['MAX_CONTENT_LENGTH']})")
            return jsonify({'error': f'File too large. Maximum size is {app.config["MAX_CONTENT_LENGTH"] / (1024*1024):.1f} MB'}), 400
        
        # Ensure upload folder exists
        upload_folder = app.config['UPLOAD_FOLDER']
        os.makedirs(upload_folder, exist_ok=True)
        
        # Save file with timestamp to avoid conflicts
        timestamp = int(time.time())
        safe_filename = f"upload_{timestamp}_{file.filename}"
        file_path = os.path.join(upload_folder, safe_filename)
        
        logging.info(f"Saving file to: {file_path}")
        file.save(file_path)
        
        # Clear any existing status for this filename and mark as processing
        update_processing_status(file.filename, 'processing')
        
        # Start background thread for fast processing
        try:
            thread = threading.Thread(target=process_excel_background, args=(file.filename, file_path))
            thread.daemon = True
            thread.start()
            logging.info(f"Background processing thread started for {file.filename}")
        except Exception as thread_error:
            logging.error(f"Failed to start background thread: {thread_error}")
            update_processing_status(file.filename, f'error: Failed to start processing')
            return jsonify({'error': 'Failed to start file processing'}), 500
        
        # Store uploaded file path in session
        session['file_path'] = file_path
        session['selected_tags'] = []
        
        # ULTRA-FAST RESPONSE - Return immediately for instant user feedback
        upload_response_time = time.time() - start_time
        logging.info(f"[UPLOAD-SIMPLE] Ultra-fast upload completed in {upload_response_time:.3f}s")
        
        return jsonify({
            'message': 'File uploaded, processing in background', 
            'filename': file.filename,
            'upload_time': f"{upload_response_time:.3f}s",
            'processing_status': 'background',
            'performance': 'ultra_fast'
        })
            
    except Exception as e:
        logging.error(f"Upload error: {e}")
        return jsonify({'error': 'Upload failed'}), 500

def process_excel_sync(filename, temp_path):
    """Synchronous Excel processing for immediate response with enhanced error logging"""
    try:
        # PC optimization: Detect platform and use optimized processing
        import platform
        is_windows = platform.system() == 'Windows'
        
        if is_windows:
            if ENHANCED_LOGGING_AVAILABLE:
                enhanced_logger.log_info(f"PC-OPTIMIZED SYNCHRONOUS PROCESSING START", 
                                       {'filename': filename, 'platform': 'windows'})
            else:
                logging.info(f"[PC-SYNC] ===== PC-OPTIMIZED SYNCHRONOUS PROCESSING START =====")
        else:
            if ENHANCED_LOGGING_AVAILABLE:
                enhanced_logger.log_info(f"SYNCHRONOUS PROCESSING START", 
                                       {'filename': filename, 'platform': 'mac'})
            else:
                logging.info(f"[SYNC] ===== SYNCHRONOUS PROCESSING START =====")
        
        if ENHANCED_LOGGING_AVAILABLE:
            enhanced_logger.log_info(f"Processing file", {'file': temp_path, 'filename': filename})
        else:
            logging.info(f"[SYNC] Processing file: {temp_path}")
            logging.info(f"[SYNC] Filename: {filename}")
        
        # Verify file exists
        if not os.path.exists(temp_path):
            if ENHANCED_LOGGING_AVAILABLE:
                enhanced_logger.log_error(f"File not found", context={'file': temp_path})
            else:
                logging.error(f"[SYNC] File not found: {temp_path}")
            return False
        
        # Create ExcelProcessor and load file
        from src.core.data.excel_processor import ExcelProcessor
        processor = ExcelProcessor()
        
        # PC optimization: Skip database integration for faster processing
        if is_windows:
            if hasattr(processor, 'enable_product_db_integration'):
                processor.enable_product_db_integration(False)
                if ENHANCED_LOGGING_AVAILABLE:
                    enhanced_logger.log_info("Product database integration disabled for faster processing", 
                                           {'platform': 'windows'})
                else:
                    logging.info("[PC-SYNC] Product database integration disabled for faster processing")
        else:
            # Enable database integration for new product storage
            if hasattr(processor, 'enable_product_db_integration'):
                processor.enable_product_db_integration(True)
                if ENHANCED_LOGGING_AVAILABLE:
                    enhanced_logger.log_info("Product database integration enabled for new product storage", 
                                           {'platform': 'mac'})
                else:
                    logging.info("[SYNC] Product database integration enabled for new product storage")
        
        # PC optimization: Use optimized loading strategy
        if is_windows:
            # PC: Use ultra-fast loading with minimal processing
            success = processor.load_file(temp_path)  # This will use the PC-optimized version
            if not success or processor.df is None or processor.df.empty:
                if ENHANCED_LOGGING_AVAILABLE:
                    enhanced_logger.log_error(f"Failed to load file", context={'file': temp_path})
                else:
                    logging.error(f"[PC-SYNC] Failed to load file: {temp_path}")
                return False
            if ENHANCED_LOGGING_AVAILABLE:
                enhanced_logger.log_success(f"Ultra-fast load complete", 
                                          {'rows': len(processor.df), 'platform': 'windows'})
            else:
                logging.info(f"[PC-SYNC] Ultra-fast load complete: {len(processor.df)} rows")
        else:
            # Mac: Use original loading strategy
            import pandas as pd
            try:
                # Try to load with row limit first
                df = pd.read_excel(temp_path, nrows=5000, engine='openpyxl')
                if not df.empty:
                    processor.df = df
                    success = True
                    if ENHANCED_LOGGING_AVAILABLE:
                        enhanced_logger.log_success(f"Loaded rows (limited to 5000)", 
                                                  {'rows': len(df), 'platform': 'mac'})
                    else:
                        logging.info(f"[SYNC] Loaded {len(df)} rows (limited to 5000)")
                else:
                    success = False
            except Exception as e:
                if ENHANCED_LOGGING_AVAILABLE:
                    enhanced_logger.log_warning(f"Limited load failed, trying full load", {'error': str(e)})
                else:
                    logging.warning(f"Limited load failed, trying full load: {e}")
                success = processor.load_file(temp_path)
            if not success or processor.df is None or processor.df.empty:
                if ENHANCED_LOGGING_AVAILABLE:
                    enhanced_logger.log_error(f"Failed to load file", context={'file': temp_path})
                else:
                    logging.error(f"[SYNC] Failed to load file: {temp_path}")
                return False
        
        # Update global processor
        global _excel_processor
        with excel_processor_lock:
            _excel_processor = processor
            _excel_processor._last_loaded_file = temp_path
            logging.info(f"[SYNC] Global processor updated with {len(processor.df)} rows")
        
        logging.info(f"[SYNC] ===== SYNCHRONOUS PROCESSING COMPLETE =====")
        return True
        
    except Exception as e:
        logging.error(f"[SYNC] ===== SYNCHRONOUS PROCESSING ERROR =====")
        logging.error(f"[SYNC] Error: {str(e)}")
        logging.error(f"[SYNC] Traceback: {traceback.format_exc()}")
        return False


def ultra_fast_background_processing(filename, temp_path):
    """Ultra-fast background processing with minimal processing for maximum speed"""
    try:
        logging.info(f"[ULTRA-FAST-BG] Starting ultra-fast processing: {filename}")
        start_time = time.time()
        
        # Step 1: Quick file validation
        if not os.path.exists(temp_path):
            update_processing_status(filename, 'error: File not found')
            return
        
        # Step 2: Create ExcelProcessor with minimal processing
        from src.core.data.excel_processor import ExcelProcessor
        processor = ExcelProcessor()
        
        # Enable database integration for product storage
        if hasattr(processor, 'enable_product_db_integration'):
            processor.enable_product_db_integration(True)  # Enable for database storage
            logging.info("[ULTRA-FAST-BG] Database integration enabled for product storage")
        
        # Step 3: Load file with full processing to ensure JointRatio is handled
        logging.info(f"[ULTRA-FAST-BG] Loading file with full processing for JointRatio support: {temp_path}")
        load_start = time.time()
        
        # Use full load_file method to ensure JointRatio processing for pre-rolls
        success = processor.load_file(temp_path)
        
        load_time = time.time() - load_start
        logging.info(f"[ULTRA-FAST-BG] Load completed in {load_time:.3f}s, success: {success}")
        
        if not success or processor.df is None or processor.df.empty:
            logging.error(f"[ULTRA-FAST-BG] Failed to load file or empty dataframe: {filename}")
            update_processing_status(filename, 'error: Failed to load file or file is empty')
            return
        
        # Step 4: Skip all heavy processing for maximum speed
        logging.info(f"[ULTRA-FAST-BG] Skipping heavy processing for {len(processor.df)} rows - using raw data")
        
        # No processing - just use the raw data as-is for maximum speed
        logging.info("[ULTRA-FAST-BG] Ultra-minimal processing completed - raw data ready")
        
        # Step 5: Store in global processor (skip database storage for speed)
        global _excel_processor
        try:
            with excel_processor_lock:
                _excel_processor = processor
                _excel_processor._last_loaded_file = temp_path
                logging.info(f"[ULTRA-FAST-BG] Global processor updated with {len(processor.df)} rows")
        except NameError:
            # Fallback if lock not available in this context (shouldn't happen)
            _excel_processor = processor
            _excel_processor._last_loaded_file = temp_path
            logging.info(f"[ULTRA-FAST-BG] Global processor updated without lock ({len(processor.df)} rows)")
        
        # Update processing status immediately
        update_processing_status(filename, 'ready')
        logging.info(f"[ULTRA-FAST-BG] Status updated to 'ready' for {filename}")
        
        total_time = time.time() - start_time
        logging.info(f"[ULTRA-FAST-BG] Ultra-fast processing completed in {total_time:.3f}s")
        
    except Exception as e:
        logging.error(f"[ULTRA-FAST-BG] Ultra-fast processing failed: {e}")
        update_processing_status(filename, f'error: {str(e)}')
def apply_essential_processing(df):
    """Apply only the most essential data processing for speed"""
    try:
        logging.info("[ULTRA-FAST-BG] Applying essential processing...")
        
        # Only do the most critical processing that's needed for the UI
        import pandas as pd
        
        # Basic string operations
        if 'Product Name*' in df.columns:
            df['Product Name*'] = df['Product Name*'].astype(str).str.strip()
        
        if 'Description' in df.columns:
            df['Description'] = df['Description'].astype(str).str.strip()
        
        # Basic lineage standardization (minimal)
        if 'Lineage' in df.columns:
            df['Lineage'] = df['Lineage'].astype(str).str.strip().str.upper()
            # Quick lineage fixes
            df['Lineage'] = df['Lineage'].replace({
                'INDICA_HYBRID': 'HYBRID/INDICA',
                'SATIVA_HYBRID': 'HYBRID/SATIVA',
                'SATIVA': 'SATIVA',
                'HYBRID': 'HYBRID',
                'INDICA': 'INDICA',
                'CBD': 'CBD'
            })
            
            # Set empty to HYBRID
            empty_mask = (df['Lineage'] == '') | (df['Lineage'] == 'NAN')
            df.loc[empty_mask, 'Lineage'] = 'HYBRID'
        
        # Basic product strain processing
        # CRITICAL FIX: Classic types should NEVER get 'Mixed' as product strain
        # Only non-classic types should default to 'Mixed' when empty
        if 'Product Strain' in df.columns:
            df['Product Strain'] = df['Product Strain'].astype(str).str.strip()
            empty_strain = (df['Product Strain'] == '') | (df['Product Strain'] == 'NAN')
            
            # Check if product type is classic
            from src.core.constants import CLASSIC_TYPES
            if 'Product Type*' in df.columns:
                product_types = df['Product Type*'].astype(str).str.strip().str.lower()
                is_classic = product_types.isin([ct.lower() for ct in CLASSIC_TYPES])
                
                # Only set 'Mixed' for non-classic types with empty strain
                # Classic types should remain empty (will be filled from database later)
                non_classic_empty_strain = empty_strain & ~is_classic
                df.loc[non_classic_empty_strain, 'Product Strain'] = 'Mixed'
            else:
                # If no Product Type column, only set Mixed for empty strains (conservative approach)
                # This shouldn't happen in normal operation, but handle gracefully
                df.loc[empty_strain, 'Product Strain'] = 'Mixed'
        
        # Basic ratio processing
        if 'Ratio' in df.columns:
            df['Ratio'] = df['Ratio'].astype(str).str.strip()
            empty_ratio = (df['Ratio'] == '') | (df['Ratio'] == 'NAN')
            df.loc[empty_ratio, 'Ratio'] = 'THC: | BR | C'
        
        # Ensure ProductName column exists for UI
        if 'Product Name*' in df.columns and 'ProductName' not in df.columns:
            df['ProductName'] = df['Product Name*']
        
        logging.info("[ULTRA-FAST-BG] Essential processing completed")
        
    except Exception as e:
        logging.error(f"Essential processing error: {str(e)}")

def update_global_processor_fast(processor, temp_path):
    """Update the global processor with minimal overhead"""
    try:
        global _excel_processor, excel_processor_lock
        
        with excel_processor_lock:
            # Clear old processor efficiently
            if _excel_processor is not None:
                if hasattr(_excel_processor, 'df'):
                    del _excel_processor.df
                if hasattr(_excel_processor, 'selected_tags'):
                    _excel_processor.selected_tags = []
            
            # Set new processor
            _excel_processor = processor
            _excel_processor._last_loaded_file = temp_path
            
            logging.info(f"[ULTRA-FAST-BG] Global processor updated with {len(processor.df)} rows")
            
    except Exception as e:
        logging.error(f"Error updating global processor: {str(e)}")
def process_excel_background(filename, temp_path):
    """ULTRA-FAST background processing - minimal operations for instant response"""
    global os  # Ensure os is available in this scope
    
    max_processing_time = MAX_TOTAL_PROCESSING_TIME if 'MAX_TOTAL_PROCESSING_TIME' in globals() else 300
    start_time = time.time()
    try:
        # PC optimization: Detect platform and use optimized processing
        import platform
        is_windows = platform.system() == 'Windows'
        
        if is_windows:
            logging.info(f"[PC-BG] ===== PC-OPTIMIZED BACKGROUND PROCESSING START =====")
        else:
            logging.info(f"[BG] ===== ULTRA-FAST BACKGROUND PROCESSING START =====")
        
        logging.info(f"[BG] Processing: {filename}")
        
        # Quick file existence check
        if not os.path.exists(temp_path):
            update_processing_status(filename, f'error: File not found')
            logging.error(f"[BG] File not found: {temp_path}")
            return
        
        start_time = time.time()
        
        # ULTRA-FAST LOADING: Load file with ABSOLUTE MINIMAL processing
        from src.core.data.excel_processor import ExcelProcessor
        new_processor = ExcelProcessor()
        
        try:
            # CRITICAL: Disable ALL expensive operations for fastest possible loading
            if hasattr(new_processor, 'enable_product_db_integration'):
                new_processor.enable_product_db_integration(False)
                logging.info("[BG] Product database integration disabled for fastest loading")
            
            # CRITICAL: Skip enrichment for fastest loading
            if hasattr(new_processor, '_skip_enrichment'):
                new_processor._skip_enrichment = True
                logging.info("[BG] Enrichment disabled for fastest loading")
            
            # CRITICAL OPTIMIZATION: Check file size and environment to choose best loading method
            import os
            file_size_mb = os.path.getsize(temp_path) / (1024 * 1024)
            logging.info(f"[BG] File size: {file_size_mb:.1f} MB - using optimized loading strategy")
            
            # Check if running on PythonAnywhere
            is_pythonanywhere = IS_PYTHONANYWHERE or PYTHONANYWHERE_OPTIMIZATION
            if is_pythonanywhere:
                logging.info("[BG] PythonAnywhere detected - using optimized loading strategy")
            
            # OPTIMIZED: Use fastest available loading method (prioritize PythonAnywhere method on production)
            success = False
            
            # CRITICAL: On PythonAnywhere, prioritize pythonanywhere_fast_load first (optimized for production)
            if is_pythonanywhere and hasattr(new_processor, 'pythonanywhere_fast_load'):
                try:
                    success = new_processor.pythonanywhere_fast_load(temp_path)
                    if success:
                        logging.info(f"[BG] ✅ PythonAnywhere fast load complete: {len(new_processor.df)} rows (production optimized)")
                except Exception as e:
                    logging.warning(f"[BG] PythonAnywhere fast load failed: {e}")
            
            # For large files (>10MB), use minimal_load_file first (fastest - no processing)
            if not success and file_size_mb > 10:
                logging.info("[BG] Large file detected - using minimal_load_file for maximum speed")
                if hasattr(new_processor, 'minimal_load_file'):
                    try:
                        success = new_processor.minimal_load_file(temp_path)
                        if success:
                            logging.info(f"[BG] ✅ Minimal load complete: {len(new_processor.df)} rows (fastest method)")
                    except Exception as e:
                        logging.warning(f"[BG] Minimal load failed: {e}")
            
            # Try ultra_fast_load (now optimized for large files - no row limit)
            if not success and hasattr(new_processor, 'ultra_fast_load'):
                try:
                    success = new_processor.ultra_fast_load(temp_path)
                    if success:
                        logging.info(f"[BG] ✅ Ultra-fast load complete: {len(new_processor.df)} rows (no limit)")
                except Exception as e:
                    logging.warning(f"[BG] Ultra-fast load failed: {e}")
            
            # Fallback to fast_load_file if ultra_fast_load failed
            if not success and hasattr(new_processor, 'fast_load_file'):
                try:
                    success = new_processor.fast_load_file(temp_path)
                    if success:
                        logging.info(f"[BG] ✅ Fast load complete: {len(new_processor.df)} rows")
                except Exception as e:
                    logging.warning(f"[BG] Fast load failed: {e}")
            
            # Final fallback to standard load_file
            if not success:
                try:
                    success = new_processor.load_file(temp_path)
                    if success:
                        logging.info(f"[BG] Standard load complete: {len(new_processor.df)} rows")
                except Exception as e:
                    logging.error(f"[BG] Standard load failed: {e}")
            
            if not success or new_processor.df is None or new_processor.df.empty:
                update_processing_status(filename, f'error: Failed to load file')
                return
            
            # CRITICAL: Mark as ready IMMEDIATELY after basic file load
            # Don't wait for database storage or any other operations
            update_processing_status(filename, 'ready')
            logging.info(f"[BG] ✅ Marked {filename} as ready IMMEDIATELY (file loaded, {len(new_processor.df)} rows)")
            
        except Exception as load_error:
            logging.error(f"[BG] Load error: {load_error}")
            update_processing_status(filename, f'error: Load failed')
            return
        
        # ULTRA-FAST PROCESSING: Update global processor immediately
        global _excel_processor
        with excel_processor_lock:
            _excel_processor = new_processor
            logging.info(f"[BG] ✅ Global processor updated with {len(new_processor.df)} rows")
        
        # CRITICAL: Defer ALL expensive operations to separate background thread
        # Database storage, enrichment, etc. should NOT block the upload response
        def deferred_expensive_operations():
            try:
                # Store in database (non-blocking, deferred)
                if hasattr(new_processor, 'enable_product_db_integration') and hasattr(new_processor, '_store_upload_in_database'):
                    new_processor.enable_product_db_integration(True)
                    storage_result = new_processor._store_upload_in_database(new_processor.df, temp_path)
                    logging.info(f"[BG-DEFERRED] ✅ Database storage completed: {storage_result}")
            except Exception as storage_error:
                logging.warning(f"[BG-DEFERRED] Database storage failed: {storage_error}")
        
        # Start deferred operations in separate thread (non-blocking)
        try:
            deferred_thread = threading.Thread(target=deferred_expensive_operations)
            deferred_thread.daemon = True
            deferred_thread.start()
            logging.info("[BG] Started deferred expensive operations in background thread")
        except Exception as deferred_error:
            logging.warning(f"[BG] Failed to start deferred operations: {deferred_error}")
        
        processing_time = time.time() - start_time
        logging.info(f"[BG] ===== ULTRA-FAST PROCESSING COMPLETE =====")
        logging.info(f"[BG] Processing time: {processing_time:.3f}s")
        logging.info(f"[BG] Rows processed: {len(new_processor.df)}")
        
    except Exception as e:
        logging.error(f"[BG] ===== ULTRA-FAST PROCESSING FAILED =====")
        logging.error(f"[BG] Error: {str(e)}")
        update_processing_status(filename, f'error: {str(e)}')
        
        # Step 1: Use fast loading for immediate response
        load_start = time.time()
        
        # Add timeout check
        if time.time() - start_time > max_processing_time:
            update_processing_status(filename, f'error: Processing timeout during file load')
            logging.error(f"[BG] Processing timeout for {filename}")
            return
        
        # Create a new ExcelProcessor instance directly
        try:
            from src.core.data.excel_processor import ExcelProcessor
            logging.info(f"[BG] Importing ExcelProcessor...")
            new_processor = ExcelProcessor()
            logging.info(f"[BG] ExcelProcessor created successfully")
        except Exception as import_error:
            logging.error(f"[BG] CRITICAL ERROR: Failed to import/create ExcelProcessor: {import_error}")
            logging.error(f"[BG] Import traceback: {traceback.format_exc()}")
            update_processing_status(filename, f'error: Failed to create ExcelProcessor: {import_error}')
            return
        
        # CRITICAL FIX: Disable default file loading to prevent interference
        try:
            new_processor._last_loaded_file = temp_path  # Set this immediately to prevent default loading
            logging.info(f"[BG] CRITICAL FIX: Set _last_loaded_file to uploaded file: {temp_path}")
        except Exception as set_error:
            logging.error(f"[BG] Error setting _last_loaded_file: {set_error}")
        
        # Enable product database integration for new product storage
        try:
            if hasattr(new_processor, 'enable_product_db_integration'):
                new_processor.enable_product_db_integration(True)
                logging.info("[BG] Product database integration enabled for new product storage")
        except Exception as db_error:
            logging.warning(f"[BG] Error enabling product database integration: {db_error}")
        
        # Use full load_file method to ensure identical processing to local version
        logging.info(f"[BG] Loading file with full load_file method: {temp_path}")
        
        # Use the full load_file method for complete data processing
        try:
            success = new_processor.load_file(temp_path)
            load_time = time.time() - load_start
            logging.info(f"[BG] File load completed in {load_time:.3f}s, success: {success}")
        except Exception as load_error:
            logging.error(f"[BG] CRITICAL ERROR: File load failed: {load_error}")
            logging.error(f"[BG] Load traceback: {traceback.format_exc()}")
            update_processing_status(filename, f'error: File load failed: {load_error}')
            return
        
        if not success:
            update_processing_status(filename, f'error: Failed to load file data')
            logging.error(f"[BG] File load failed for {filename}")
            return
        
        # Verify the load was successful
        if new_processor.df is None or new_processor.df.empty:
            update_processing_status(filename, f'error: Failed to load file data - DataFrame is empty')
            logging.error(f"[BG] File load failed for {filename} - DataFrame is empty")
            return
        
        # CRITICAL FIX: Populate dropdown cache after successful file load
        if hasattr(new_processor, '_cache_dropdown_values'):
            try:
                new_processor._cache_dropdown_values()
                logging.info(f"[BG] Successfully populated dropdown cache after file load")
                # Log the strain count specifically
                if 'strain' in new_processor.dropdown_cache:
                    strain_count = len(new_processor.dropdown_cache['strain'])
                    logging.info(f"[BG] Dropdown cache contains {strain_count} strains")
                else:
                    logging.warning("[BG] No strain filter found in dropdown cache")
            except Exception as e:
                logging.error(f"[BG] Failed to populate dropdown cache after file load: {e}")
        else:
            logging.warning("[BG] ExcelProcessor does not have _cache_dropdown_values method")
        
        # CRITICAL FIX: Verify we loaded the correct file (with more robust comparison)
        logging.info(f"[BG] CRITICAL FIX: Verifying loaded file matches uploaded file")
        logging.info(f"[BG] Expected file: {temp_path}")
        logging.info(f"[BG] Loaded file: {new_processor._last_loaded_file}")
        
        # More robust file path comparison
        expected_path = os.path.abspath(temp_path) if temp_path else None
        loaded_path = os.path.abspath(new_processor._last_loaded_file) if new_processor._last_loaded_file else None
        
        logging.info(f"[BG] Normalized expected path: {expected_path}")
        logging.info(f"[BG] Normalized loaded path: {loaded_path}")
        
        # Check if file verification should be bypassed (for debugging)
        bypass_verification = os.environ.get('BYPASS_FILE_VERIFICATION', 'false').lower() == 'true'
        if bypass_verification:
            logging.warning(f"[BG] File verification bypassed due to BYPASS_FILE_VERIFICATION environment variable")
        elif expected_path != loaded_path:
            logging.error(f"[BG] CRITICAL ERROR: Loaded wrong file! Expected {expected_path}, got {loaded_path}")
            update_processing_status(filename, f'error: Loaded incorrect file')
            return
        else:
            logging.info(f"[BG] File verification passed - loaded correct file")
        
        # Debug Vendor data
        if hasattr(new_processor, 'df') and new_processor.df is not None:
            vendor_columns = [col for col in new_processor.df.columns if 'vendor' in col.lower()]
            logging.info(f"[BG] Vendor columns found: {vendor_columns}")
            if vendor_columns:
                sample_vendor_data = new_processor.df[vendor_columns[0]].head(5).tolist()
                logging.info(f"[BG] Sample vendor data: {sample_vendor_data}")
        
        # CRITICAL FIX: Clear all caches to ensure new file is processed
        logging.info(f"[BG] CRITICAL FIX: Clearing all caches for new file")
        if hasattr(new_processor, '_file_cache'):
            new_processor._file_cache.clear()
            logging.info(f"[BG] Cleared file cache")
        if hasattr(new_processor, '_dropdown_cache'):
            new_processor._dropdown_cache.clear()
            logging.info(f"[BG] Cleared dropdown cache")
        if hasattr(new_processor, '_available_tags_cache'):
            new_processor._available_tags_cache.clear()
            logging.info(f"[BG] Cleared available tags cache")
        
        # Avoid redundant reload after fast load on PythonAnywhere for performance
        if os.environ.get('FORCE_RELOAD_AFTER_FAST_LOAD', 'false').lower() == 'true':
            logging.info(f"[BG] FORCE_RELOAD_AFTER_FAST_LOAD is enabled; reloading file for verification")
            new_processor._last_loaded_file = None
            new_processor.df = None
            if hasattr(new_processor, '_file_cache'):
                new_processor._file_cache.clear()
            reload_success = new_processor.load_file(temp_path)
            if not reload_success:
                logging.error(f"[BG] CRITICAL ERROR: Failed to reload file {temp_path}")
                update_processing_status(filename, f'error: Failed to reload file')
                return
            logging.info(f"[BG] File reloaded successfully with fresh data")
            
            # CRITICAL FIX: Populate dropdown cache after redundant reload
            if hasattr(new_processor, '_cache_dropdown_values'):
                try:
                    new_processor._cache_dropdown_values()
                    logging.info(f"[BG] Successfully populated dropdown cache after redundant reload")
                    # Log the strain count specifically
                    if 'strain' in new_processor.dropdown_cache:
                        strain_count = len(new_processor.dropdown_cache['strain'])
                        logging.info(f"[BG] Dropdown cache contains {strain_count} strains")
                    else:
                        logging.warning("[BG] No strain filter found in dropdown cache")
                except Exception as e:
                    logging.error(f"[BG] Failed to populate dropdown cache after redundant reload: {e}")
            else:
                logging.warning("[BG] ExcelProcessor does not have _cache_dropdown_values method")
        else:
            logging.info(f"[BG] Using pythonanywhere_fast_load data without redundant reload")
        
        # CRITICAL FIX: Clear global cache to force fresh data
        logging.info(f"[BG] CRITICAL FIX: Clearing global cache to force fresh data")
        try:
            from flask import has_request_context
            if has_request_context():
                # Clear all cache keys that might contain old data
                cache_keys_to_clear = [
                    'available_tags', 'selected_tags', 'filter_options', 'dropdowns',
                    'json_matched_tags', 'full_excel_tags', 'initial_data'
                ]
                
                for cache_key_name in cache_keys_to_clear:
                    try:
                        # Try different cache key patterns
                        cache_keys_to_try = [
                            get_session_cache_key(cache_key_name),
                            f"{cache_key_name}_default",
                            cache_key_name,
                            f"full_excel_cache_key",
                            f"json_matched_cache_key"
                        ]
                        
                        for key in cache_keys_to_try:
                            cache.delete(key)
                            logging.info(f"[BG] Cleared global cache key: {key}")
                    except Exception as key_error:
                        logging.warning(f"[BG] Error clearing global cache key {cache_key_name}: {key_error}")
            else:
                logging.info("[BG] Skipping global cache clear - not in request context")
        except Exception as global_cache_error:
            logging.warning(f"[BG] Error in global cache clearing: {global_cache_error}")
        # CRITICAL FIX: Store uploaded data in database for persistence and analytics
        # JSON matched tags will be automatically excluded from database storage
        try:
            logging.info(f"[BG] CRITICAL: Forcing database storage of uploaded data")
            logging.info(f"[BG] DataFrame shape: {new_processor.df.shape if hasattr(new_processor.df, 'shape') else 'No DataFrame'}")
            logging.info(f"[BG] DataFrame columns: {list(new_processor.df.columns) if hasattr(new_processor.df, 'columns') else 'No columns'}")
            
            # CRITICAL FIX: JSON tags work exactly like Excel tags - no special exclusion from database storage
            # All tags (including JSON) are processed the same way
            
            if hasattr(new_processor, '_store_upload_in_database'):
                logging.info("[BG] Using ExcelProcessor _store_upload_in_database method")
                storage_result = new_processor._store_upload_in_database(new_processor.df, temp_path)
                logging.info(f"[BG] ✅ Database storage completed successfully: {storage_result}")
                
                # Log storage results
                logging.info(f"[BG] ✅ Stored {storage_result.get('stored', 0)} products in database")
                    
            else:
                logging.warning("[BG] ExcelProcessor does not have _store_upload_in_database method")
                # Try alternative database storage method
                try:
                    logging.info("[BG] Attempting alternative database storage with ProductDatabase")
                    # Store context removed - using single database
                    store_name = get_current_store_name()
                    product_db = get_product_database(store_name)
                    logging.info(f"[BG] ProductDatabase obtained: {product_db}")
                    
                    if hasattr(product_db, 'store_excel_data'):
                        logging.info("[BG] ProductDatabase has store_excel_data method, calling it...")
                        storage_result = product_db.store_excel_data(new_processor.df, temp_path)
                        logging.info(f"[BG] ✅ Alternative database storage completed: {storage_result}")
                        
                        # Log storage results
                        logging.info(f"[BG] ✅ Stored {storage_result.get('stored', 0)} products in database")
                            
                    else:
                        logging.warning("[BG] ProductDatabase does not have store_excel_data method")
                        logging.error("[BG] CRITICAL: No database storage method available!")
                except Exception as alt_storage_error:
                    logging.error(f"[BG] Alternative database storage failed: {alt_storage_error}")
                    import traceback
                    logging.error(f"[BG] Alternative storage traceback: {traceback.format_exc()}")
        except Exception as storage_error:
            logging.error(f"[BG] ❌ Database storage failed: {storage_error}")
            import traceback
            logging.error(f"[BG] Storage error traceback: {traceback.format_exc()}")
            # Don't fail the upload - continue without database storage
        
        # Mark as ready as soon as DataFrame is loaded so frontend can proceed
        try:
            update_processing_status(filename, 'ready')
            logging.info(f"[BG] Marked {filename} as ready (DataFrame loaded)")
        except Exception as mark_ready_error:
            logging.warning(f"[BG] Failed to mark ready: {mark_ready_error}")

        # Step 2: Update the global processor safely with minimal clearing
        with excel_processor_lock:
            # Clear the old processor completely
            if _excel_processor is not None:
                # Explicitly clear all data from old processor
                if hasattr(_excel_processor, 'df') and _excel_processor.df is not None:
                    del _excel_processor.df
                    logging.info("[BG] Cleared old DataFrame from ExcelProcessor")
                
                if hasattr(_excel_processor, 'selected_tags'):
                    logging.info(f"[BG] Clearing selected tags from ExcelProcessor. Previous count: {len(_excel_processor.selected_tags) if _excel_processor.selected_tags else 0}")
                    _excel_processor.selected_tags = []
                    logging.info("[BG] Cleared selected tags from ExcelProcessor")
                
                if hasattr(_excel_processor, 'dropdown_cache'):
                    _excel_processor.dropdown_cache = {}
                    logging.info("[BG] Cleared dropdown cache from ExcelProcessor")
                
                # Clear any other data attributes
                for attr in ['data', 'original_data', 'processed_data']:
                    if hasattr(_excel_processor, attr):
                        delattr(_excel_processor, attr)
                        logging.info(f"[BG] Cleared {attr} from ExcelProcessor")
                
                # Force garbage collection
                import gc
                gc.collect()
                logging.info("[BG] Forced garbage collection to free memory")
            
            # Replace with the new processor
            _excel_processor = new_processor
            _excel_processor._last_loaded_file = temp_path
            
            # Store store context in the processor for validation (without Flask session)
            try:
                # Don't use Flask session in background thread - it's not available
                # Just set a default store context
                _excel_processor._store_context = 'uploaded_file'
                logging.info(f"[BG] Store context set for processor: uploaded_file")
            except Exception as store_error:
                logging.warning(f"[BG] Error setting store context: {store_error}")
            
            logging.info(f"[BG] Global Excel processor updated with new file: {temp_path}")
            
            # CRITICAL FIX: Verify the global processor was set correctly
            if _excel_processor is not None and _excel_processor.df is not None:
                logging.info(f"[BG] ✅ VERIFICATION: Global processor has {len(_excel_processor.df)} rows")
                logging.info(f"[BG] ✅ VERIFICATION: Global processor file: {_excel_processor._last_loaded_file}")
            else:
                logging.error(f"[BG] ❌ CRITICAL ERROR: Global processor is None or has no data!")
                logging.error(f"[BG] ❌ Global processor: {_excel_processor}")
                if _excel_processor is not None:
                    logging.error(f"[BG] ❌ Global processor df: {_excel_processor.df}")
                    logging.error(f"[BG] ❌ Global processor df is None: {_excel_processor.df is None}")
                    if hasattr(_excel_processor, 'df') and _excel_processor.df is not None:
                        logging.error(f"[BG] ❌ Global processor df empty: {_excel_processor.df.empty}")
        
        # ULTRA-FAST CACHE OPTIMIZATION - Minimal clearing
        clear_initial_data_cache()
        
        # Only clear the most critical caches for instant response
        try:
            from flask import has_request_context
            if has_request_context():
                # Clear only the most essential caches
                critical_keys = ['full_excel_cache_key', 'json_matched_cache_key']
                cleared_count = 0
                for key in critical_keys:
                    if cache.has(key):
                        cache.delete(key)
                        cleared_count += 1
                logging.info(f"[BG] Cleared {cleared_count} critical cache entries for instant response")
            else:
                logging.info("[BG] Skipping Flask cache clear - not in request context")
        except Exception as cache_error:
            logging.warning(f"[BG] Error clearing file caches: {cache_error}")
        
        # Clear specific cache keys that might persist (only if in request context)
        try:
            from flask import has_request_context
            if has_request_context():
                cache_keys_to_clear = [
                    'available_tags', 'selected_tags', 'filter_options', 'dropdowns',
                    'json_matched_tags', 'full_excel_tags'
                ]
                
                for cache_key_name in cache_keys_to_clear:
                    try:
                        # Try different cache key patterns
                        cache_keys_to_try = [
                            get_session_cache_key(cache_key_name),
                            f"{cache_key_name}_default",  # Use default instead of session.get()
                            cache_key_name
                        ]
                        
                        for key in cache_keys_to_try:
                            cache.delete(key)
                            logging.info(f"[BG] Cleared cache key: {key}")
                    except Exception as key_error:
                        logging.warning(f"[BG] Error clearing cache key {cache_key_name}: {key_error}")
            else:
                logging.info("[BG] Skipping cache key clear - not in request context")
        except Exception as cache_key_error:
            logging.warning(f"[BG] Error in cache key clearing: {cache_key_error}")
        
        # Clear session data that might persist (only if in request context)
        try:
            from flask import has_request_context, session, g
            if has_request_context():
                session_keys_to_clear = [
                    'selected_tags', 'current_filter_mode', 'json_matched_cache_key',
                    'full_excel_cache_key'
                ]
                
                for key in session_keys_to_clear:
                    if key in session:
                        del session[key]
                        logging.info(f"[BG] Cleared session key: {key}")
                
                # Clear any g context that might exist
                if hasattr(g, 'excel_processor'):
                    delattr(g, 'excel_processor')
                    logging.info("[BG] Cleared g.excel_processor context")
            else:
                logging.info("[BG] Skipping session/g context clear - not in request context (background thread)")
        except Exception as session_error:
            logging.warning(f"[BG] Error clearing session/g context: {session_error}")
        # Update processing status to success
        update_processing_status(filename, 'ready')
        logging.info(f"[BG] ===== BACKGROUND PROCESSING COMPLETE =====")
        logging.info(f"[BG] File processing completed successfully: {filename}")
        
        # Step 3: Mark as ready immediately (no delay needed with fast loading)
        logging.info(f"[BG] Marking file as ready: {filename}")
        update_processing_status(filename, 'ready')
        logging.info(f"[BG] File marked as ready: {filename}")
        logging.info(f"[BG] Current processing statuses: {dict(processing_status)}")
        
        # Step 4: Schedule full processing in background if needed
        # This allows the UI to be responsive immediately while full processing happens later
        try:
            import threading
            def full_processing_background():
                """Background task for full data processing if needed."""
                try:
                    logging.info(f"[BG-FULL] Starting full processing for: {filename}")
                    # Here you could add any additional processing that's not critical for basic functionality
                    # For now, we'll just log that full processing is complete
                    logging.info(f"[BG-FULL] Full processing complete for: {filename}")
                except Exception as e:
                    logging.error(f"[BG-FULL] Error in full processing: {e}")
            
            # Start full processing in background (non-blocking)
            full_thread = threading.Thread(target=full_processing_background)
            full_thread.daemon = True
            full_thread.start()
            logging.info(f"[BG] Full processing thread started for {filename}")
        except Exception as full_thread_error:
            logging.warning(f"[BG] Failed to start full processing thread: {full_thread_error}")
        
    except Exception as e:
        logging.error(f"[BG] ===== BACKGROUND PROCESSING FAILED =====")
        logging.error(f"[BG] Error in background processing: {str(e)}")
        logging.error(f"[BG] Traceback: {traceback.format_exc()}")
        update_processing_status(filename, f'error: {str(e)}')

@app.route('/api/upload-status', methods=['GET'])
def upload_status():
    try:
        filename = request.args.get('filename')
        if not filename:
            return jsonify({'error': 'No filename provided'}), 400

        logging.info(f"Status check for: {filename}")

        # Ensure filename is properly sanitized
        filename = sanitize_filename(filename)

        # Clean up old entries periodically (but not on every request to reduce overhead)
        if random.random() < 0.05:  # Only cleanup 5% of the time (reduced from 10%)
            cleanup_old_processing_status()

        # CRITICAL FIX: More aggressive auto-recovery for stuck uploads
        # Check on EVERY status request if this specific file is stuck (not random)
        current_time = time.time()
        with processing_lock:
            status = processing_status.get(filename, 'not_found')
            timestamp = processing_timestamps.get(filename, 0)
            age = current_time - timestamp if timestamp > 0 else 0

            # If file has been "processing" for more than 30 seconds, check if it's actually ready
            if status == 'processing' and age > 30:
                logging.warning(f"⚠️ Upload stuck in 'processing' for {age:.1f}s: {filename}")

                # Check if the file actually loaded successfully
                local_processor = get_excel_processor()
                if local_processor and hasattr(local_processor, 'df') and local_processor.df is not None and not local_processor.df.empty:
                    # File is actually ready! Background thread must have failed to update status
                    logging.info(f"✅ AUTO-RECOVERY: File {filename} is actually ready (has {len(local_processor.df)} rows)")
                    processing_status[filename] = 'ready'
                    processing_timestamps[filename] = current_time
                    status = 'ready'
                elif age > 120:  # Stuck for more than 2 minutes
                    # File is stuck and not loaded - mark as error so frontend can retry
                    logging.error(f"❌ AUTO-RECOVERY: File {filename} stuck for {age:.1f}s with no data - marking as error")
                    processing_status[filename] = 'error: Upload timeout - please try again'
                    status = 'error: Upload timeout - please try again'

        # Auto-clear very old stuck processing statuses (older than 15 minutes) - random check
        if random.random() < 0.02:
            cutoff_time = current_time - 900  # 15 minutes

            with processing_lock:
                stuck_files = []
                for fname, fstatus in list(processing_status.items()):
                    ftimestamp = processing_timestamps.get(fname, 0)
                    fage = current_time - ftimestamp
                    if fage > cutoff_time and fstatus == 'processing':
                        stuck_files.append(fname)
                        del processing_status[fname]
                        if fname in processing_timestamps:
                            del processing_timestamps[fname]

                if stuck_files:
                    logging.warning(f"Auto-cleared {len(stuck_files)} very old stuck statuses: {stuck_files}")

        with processing_lock:
            status = processing_status.get(filename, status)  # Get updated status
            all_statuses = dict(processing_status)
            age = current_time - timestamp if timestamp > 0 else 0
        
        logging.info(f"Upload status request for {filename}: {status} (age: {age:.1f}s)")
        logging.debug(f"All processing statuses: {all_statuses}")
        
        # Check if file exists in uploads directory (check both with and without timestamp)
        upload_folder = app.config['UPLOAD_FOLDER']
        file_path = os.path.join(upload_folder, filename)
        file_exists = os.path.exists(file_path)
        
        # Also check for timestamp-prefixed version
        if not file_exists:
            import glob
            pattern = os.path.join(upload_folder, f"*_{filename}")
            matching_files = glob.glob(pattern)
            if matching_files:
                file_exists = True
                file_path = matching_files[0]  # Use most recent match
                logging.info(f"Found timestamp-prefixed file: {file_path}")
        
        # If status is 'not_found' but file exists, it might have been processed successfully
        if status == 'not_found' and file_exists:
            # Check if we have a processor with data
            local_processor = get_excel_processor()
            if local_processor and hasattr(local_processor, 'df') and local_processor.df is not None and not local_processor.df.empty:
                status = 'ready'
                logging.info(f"File {filename} appears to be processed (processor has data)")
            else:
                status = 'processing'  # Still processing
        elif status == 'processing' and file_exists:
            # Check if processing is actually complete
            local_processor = get_excel_processor()
            if local_processor and hasattr(local_processor, 'df') and local_processor.df is not None and not local_processor.df.empty:
                status = 'ready'
                logging.info(f"File {filename} processing completed (global processor has data)")
                # Update the status in the tracking
                with processing_lock:
                    processing_status[filename] = 'ready'
                    processing_timestamps[filename] = time.time()

        # Prepare response
        response_data = {
            'status': status,
            'filename': filename,
            'age_seconds': round(age, 1),
            'total_processing_files': len(all_statuses),
            'file_exists': file_exists,
            'upload_folder': upload_folder
        }
        
        # If status is 'ready' and age is less than 30 seconds, don't clear it yet
        # This prevents race conditions where frontend is still polling
        if status == 'ready' and age < 30:
            logging.debug(f"Keeping 'ready' status for {filename} (age: {age:.1f}s)")
        
        return jsonify(response_data)
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        logging.error(f"/api/upload-status error for '{request.args.get('filename')}': {e}\n{tb}")
        # Never return HTML; always JSON for the polling loop
        return jsonify({'error': str(e), 'trace': tb, 'status': 'processing'}), 500

@app.route('/api/current-file', methods=['GET'])
def get_current_file():
    """Get the current uploaded file information from session"""
    try:
        file_path = session.get('file_path')
        uploaded_filename = session.get('uploaded_filename', '')
        upload_timestamp = session.get('upload_timestamp', 0)
        
        # Check if file exists
        file_exists = False
        if file_path:
            file_exists = os.path.exists(file_path)
            if not file_exists:
                # File doesn't exist, try to recover from uploads directory
                # Look for the most recent file that matches the filename pattern
                if uploaded_filename:
                    try:
                        uploads_dir = os.path.join(os.getcwd(), 'uploads')
                        if os.path.exists(uploads_dir):
                            # Find files matching the uploaded filename pattern
                            import glob
                            pattern = os.path.join(uploads_dir, f'*_{uploaded_filename}')
                            matching_files = glob.glob(pattern)
                            if matching_files:
                                # Get the most recent file
                                matching_files.sort(key=os.path.getmtime, reverse=True)
                                recovered_file = matching_files[0]
                                file_age = time.time() - os.path.getmtime(recovered_file)
                                # Only recover if file is less than 2 hours old
                                if file_age < 7200:
                                    file_path = recovered_file
                                    file_exists = True
                                    # Restore session data
                                    session['file_path'] = file_path
                                    session['uploaded_filename'] = uploaded_filename
                                    session['upload_timestamp'] = int(os.path.getmtime(recovered_file))
                                    session.modified = True
                                    logging.info(f"✅ RECOVERED file from disk: {file_path} (age: {file_age:.0f}s)")
                                    
                                    # CRITICAL: Force reload of the recovered file into processor
                                    try:
                                        processor = get_excel_processor()
                                        if processor:
                                            # Force reload by clearing the last loaded file
                                            processor._last_loaded_file = None
                                            success = processor.load_file(file_path)
                                            if success:
                                                processor._last_loaded_file = file_path
                                                logging.info(f"✅ Successfully loaded recovered file into processor")
                                            else:
                                                logging.warning(f"⚠️  Failed to load recovered file into processor")
                                    except Exception as load_error:
                                        logging.warning(f"Error loading recovered file: {load_error}")
                                else:
                                    logging.info(f"File too old to recover: {recovered_file} (age: {file_age:.0f}s)")
                    except Exception as recover_error:
                        logging.warning(f"Error recovering file: {recover_error}")
                
                if not file_exists:
                    # File doesn't exist and couldn't be recovered, clear session
                    session.pop('file_path', None)
                    session.pop('uploaded_filename', None)
                    session.pop('upload_timestamp', None)
                    logging.info(f"File from session no longer exists: {file_path}")
        
        # Check if processor has data
        has_data = False
        row_count = 0
        if file_exists:
            try:
                processor = get_excel_processor()
                if processor and hasattr(processor, 'df') and processor.df is not None and not processor.df.empty:
                    has_data = True
                    row_count = len(processor.df)
            except Exception as e:
                logging.warning(f"Error checking processor data: {e}")
        
        return jsonify({
            'success': True,
            'has_file': file_exists,
            'filename': uploaded_filename,
            'file_path': file_path if file_exists else None,
            'upload_timestamp': upload_timestamp,
            'has_data': has_data,
            'row_count': row_count
        })
    except Exception as e:
        logging.error(f"Error getting current file: {e}")
        return jsonify({
            'success': False,
            'error': str(e),
            'has_file': False
        }), 500

@app.route('/upload-lightning', methods=['POST'])
def upload_lightning():
    """Ultra-fast file upload - saves file immediately, processes later"""
    try:
        logging.info("=== LIGHTNING UPLOAD START ===")
        start_time = time.time()
        
        # CRITICAL FIX: Use get_current_store_name with fallback instead of has_store_selection
        # has_store_selection can be too strict and fail even when store is selected
        selected_store = get_current_store_name(allow_fallback=True)
        if not selected_store:
            logging.error("Upload attempted without store selection")
            return jsonify({'error': 'Please select a store before uploading files'}), 400
        
        # Validate file upload
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not file.filename.lower().endswith('.xlsx'):
            return jsonify({'error': 'Only .xlsx files are allowed'}), 400
        
        # Validate filename contains store name and matches selected store
        is_valid, warning_msg, detected_store = validate_excel_filename_for_store(file.filename, selected_store)
        
        if not is_valid:
            logging.error(f"Filename validation failed: {warning_msg}")
            return jsonify({
                'error': warning_msg,
                'filename': file.filename,
                'selected_store': selected_store,
                'detected_store': detected_store
            }), 400
        
        # Quick file size check
        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)
        
        max_size = app.config.get('MAX_CONTENT_LENGTH', 100 * 1024 * 1024)
        if file_size > max_size:
            return jsonify({'error': f'File too large. Maximum size is {max_size / (1024*1024):.1f} MB'}), 400
        
        # Sanitize filename
        sanitized_filename = sanitize_filename(file.filename)
        if not sanitized_filename:
            return jsonify({'error': 'Invalid filename'}), 400
        
        # Save file immediately (no processing)
        upload_folder = app.config['UPLOAD_FOLDER']
        os.makedirs(upload_folder, exist_ok=True)
        
        file_path = os.path.join(upload_folder, sanitized_filename)
        file.save(file_path)
        
        # Store file path in session for later processing
        session['uploaded_file_path'] = file_path
        session['uploaded_filename'] = sanitized_filename
        
        upload_time = time.time() - start_time
        logging.info(f"[LIGHTNING] File saved in {upload_time:.3f}s: {file_path}")
        
        return jsonify({
            'success': True,
            'message': f'File uploaded successfully in {upload_time:.3f}s',
            'file_path': file_path,
            'filename': sanitized_filename,
            'upload_time': upload_time,
            'file_size': file_size
        })
        
    except Exception as e:
        logging.error(f"[LIGHTNING] Upload failed: {e}")
        return jsonify({'error': f'Upload failed: {str(e)}'}), 500

@app.route('/process-lightning', methods=['POST'])
def process_lightning():
    """Process the lightning-uploaded file in the background"""
    try:
        # PC optimization: Detect platform and use optimized processing
        import platform
        is_windows = platform.system() == 'Windows'
        
        if is_windows:
            logging.info("=== PC-OPTIMIZED LIGHTNING PROCESSING START ===")
        else:
            logging.info("=== LIGHTNING PROCESSING START ===")
        
        start_time = time.time()
        
        # Get file path from session or request
        file_path = session.get('uploaded_file_path')
        filename = session.get('uploaded_filename')
        
        if not file_path or not os.path.exists(file_path):
            return jsonify({'error': 'No uploaded file found to process'}), 400
        
        # Load file with optimizations
        from src.core.data.excel_processor import ExcelProcessor
        processor = ExcelProcessor()
        
        # PC optimization: Use platform-specific loading strategy
        if is_windows:
            # PC: Use ultra-fast loading with minimal processing
            success = processor.load_file(file_path)  # This will use the PC-optimized version
            if not success:
                return jsonify({'error': 'Failed to process file'}), 500
            logging.info(f"[PC-LIGHTNING] Ultra-fast load complete: {len(processor.df)} rows")
        else:
            # Mac: Use original loading strategy
            import pandas as pd
            try:
                # OPTIMIZATION: Load more rows for better data coverage
                df = pd.read_excel(file_path, nrows=50000, engine='openpyxl', dtype=str, na_filter=False)
                processor.df = df
                logging.info(f"[LIGHTNING] Loaded {len(df)} rows (optimized for speed)")
            except Exception as e:
                logging.warning(f"[LIGHTNING] Quick load failed, trying full load: {e}")
                success = processor.load_file(file_path)
                if not success:
                    return jsonify({'error': 'Failed to process file'}), 500
        
        # Update global processor
        global _excel_processor
        with excel_processor_lock:
            _excel_processor = processor
            _excel_processor._last_loaded_file = file_path
        
        # PC optimization: Skip cache clearing for better performance
        if not is_windows:
            # Clear minimal caches only
            cache.delete('full_excel_cache_key')
            cache.delete('dropdown_cache_key')
        
        process_time = time.time() - start_time
        logging.info(f"[LIGHTNING] Processing completed in {process_time:.3f}s")
        
        return jsonify({
            'success': True,
            'message': f'File processed successfully in {process_time:.3f}s',
            'rows_loaded': len(processor.df),
            'process_time': process_time
        })
        
    except Exception as e:
        logging.error(f"[LIGHTNING] Processing failed: {e}")
        return jsonify({'error': f'Processing failed: {str(e)}'}), 500

@app.route('/api/template', methods=['POST'])
def edit_template():
    """
    Edit template settings and apply changes to template file. 
    Expected JSON payload:
    {
        "type": "horizontal|vertical|mini|inventory",
        "font_settings": {
            "base_size": 12,
            "title_size": 14,
            "body_size": 11
        }
    }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        # Validate template type
        template_type = data.get('type')
        if not template_type:
            return jsonify({'error': 'Template type is required'}), 400
            
        if template_type not in ['horizontal', 'vertical', 'mini', 'double', 'inventory']:
            return jsonify({'error': 'Invalid template type'}), 400

        # Validate font settings
        font_settings = data.get('font_settings', {})
        if not isinstance(font_settings, dict):
            return jsonify({'error': 'Font settings must be an object'}), 400

        # Get and validate template path
        try:
            template_path = get_template_path(template_type)
        except Exception as e:
            logging.error(f"Failed to get template path: {str(e)}")
            return jsonify({'error': 'Template path error'}), 500

        if not template_path or not os.path.exists(template_path):
            return jsonify({'error': 'Template not found'}), 404

        # Apply template fixes and save settings
        try:

            # Save font settings
            save_template_settings(template_type, font_settings)
            
            # Clear font scheme cache if it exists
            if hasattr(get_cached_font_scheme, 'cache_clear'):
                get_cached_font_scheme.cache_clear()

            return jsonify({
                'success': True,
                'message': 'Template updated successfully'
            })

        except Exception as e:
            logging.error(f"Failed to update template: {str(e)}")
            return jsonify({
                'error': 'Failed to update template',
                'details': str(e)
            }), 500

    except Exception as e:
        logging.error(f"Error in edit_template: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/template-settings', methods=['POST'])
def save_template_settings_api():
    """
    Save comprehensive template settings to backend.
    Expected JSON payload:
    {
        "templateType": "horizontal|vertical|mini|double|inventory",
        "scale": 1.0,
        "font": "Arial",
        "fontSizeMode": "auto|fixed|custom",
        "lineBreaks": true,
        "textWrapping": true,
        "boldHeaders": false,
        "italicDescriptions": false,
        "lineSpacing": "1.0",
        "paragraphSpacing": "0",
        "textColor": "#000000",
        "backgroundColor": "#ffffff",
        "headerColor": "#333333",
        "accentColor": "#007bff",
        "autoResize": true,
        "smartTruncation": true,
        "optimization": false,
        "fieldFontSizes": {
            "description": 16,
            "brand": 14,
            "price": 18,
            "lineage": 12,
            "ratio": 10,
            "vendor": 8
        }
    }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        # Store settings in session for this user
        session['template_settings'] = data
        
        logging.info(f"Template settings saved for session: {data.get('templateType', 'unknown')}")
        
        return jsonify({
            'success': True,
            'message': 'Template settings saved successfully'
        })

    except Exception as e:
        logging.error(f"Error saving template settings: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/template-settings', methods=['GET'])
def get_template_settings_api():
    """
    Get saved template settings from backend.
    """
    try:
        settings = session.get('template_settings', {})
        return jsonify({
            'success': True,
            'settings': settings
        })

    except Exception as e:
        logging.error(f"Error getting template settings: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/current-db', methods=['GET'])
def current_db_info():
    """Return the currently selected product database path and store for debugging."""
    try:
        store_name = get_current_store_name()
        product_db = get_product_database(store_name)
        import os
        db_path = getattr(product_db, 'db_path', 'unknown') if product_db else 'unavailable'
        store = getattr(product_db, '_store_name', 'unknown') if product_db else 'unavailable'
        size = os.path.getsize(db_path) if db_path and os.path.exists(db_path) else 0
        return jsonify({
            'db_path': db_path,
            'store': store,
            'size_bytes': size
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Add undo/clear support for tag moves and filters
from flask import session
# Helper: maintain an undo stack in session
UNDO_STACK_KEY = 'undo_stack'

@app.route('/api/move-tags', methods=['POST'])
def move_tags():
    try:
        logging.info("=== MOVE TAGS ACTION START ===")
        logging.info(f"Move tags request at {datetime.now().strftime('%H:%M:%S')}")
        logging.info(f"Request method: {request.method}")
        logging.info(f"Request URL: {request.url}")
        logging.info(f"Request headers: {dict(request.headers)}")
        
        # Check session size but don't optimize unless necessary
        check_session_size()
        
        # Store validation to prevent cross-store data access
        # Store context removed - using single database
        file_store = session.get('file_store', '')
        
        # Store validation removed - using single database
        
        data = request.get_json()
        action = data.get('action', 'move')
        logging.info(f"Action: {action}")
        logging.info(f"Request data: {data}")
        
        excel_processor = get_session_excel_processor()
        available_tags = excel_processor.get_available_tags()
        
        # Convert available_tags to just names for efficiency
        available_tag_names = [tag.get('Product Name*', '') for tag in available_tags if tag.get('Product Name*', '')]
        
        # Helper to normalize selected_tags to a list of plain strings (product names)
        def _normalize_selected_tags_list(tags):
            normalized = []
            for t in tags or []:
                if isinstance(t, dict):
                    name = t.get('Product Name*') or t.get('ProductName') or t.get('name')
                    if name:
                        normalized.append(str(name))
                elif isinstance(t, str):
                    normalized.append(t)
                else:
                    normalized.append(str(t))
            return normalized

        # Get selected tags - handle both dict and string objects
        selected_tags = _normalize_selected_tags_list(excel_processor.selected_tags)
        # Ensure processor uses normalized list to avoid dict/string mismatch on removal
        excel_processor.selected_tags = selected_tags
        
        logging.info(f"Move tags - Current selected tags: {selected_tags}")
        
        # Handle reorder action
        if action == 'reorder':
            new_order = data.get('newOrder', [])
            
            if new_order:
                # Validate that all items in new_order exist in selected_tags
                current_selected = set(selected_tags)
                new_order_valid = [tag for tag in new_order if tag in current_selected]
                
                # If no valid tags found, use the original order
                if not new_order_valid:
                    new_order_valid = selected_tags.copy()
                else:
                    # Add any missing tags from current selection
                    for tag in selected_tags:
                        if tag not in new_order_valid:
                            new_order_valid.append(tag)
                
                # Update the selected tags order - convert names back to dictionary objects
                # First, find the corresponding dictionary objects for each name
                updated_selected_tags = []
                for tag_name in new_order_valid:
                    # Try to find the corresponding dictionary in available_tags
                    for available_tag in available_tags:
                        if isinstance(available_tag, dict) and available_tag.get('Product Name*', '') == tag_name:
                            updated_selected_tags.append(available_tag)
                            break
                    else:
                        # If not found, create a simple dict with just the name
                        updated_selected_tags.append({'Product Name*': tag_name})
                
                excel_processor.selected_tags = updated_selected_tags
                # Update session with the full dictionary objects
                session['selected_tags'] = updated_selected_tags
                
                # Force session to be saved
                session.modified = True
                
                logging.info(f"Reordered selected tags: {new_order_valid}")
                
                return jsonify({
                    'success': True,
                    'message': 'Tags reordered successfully',
                    'selected_tags': new_order_valid,
                    'available_tags': available_tag_names
                })
        
        # Handle move action (existing functionality)
        tags_to_move = data.get('tags', [])
        direction = data.get('direction', 'to_selected')
        select_all = data.get('selectAll', False)
        
        logging.info(f"Move tags - Tags to move: {tags_to_move}")
        logging.info(f"Move tags - Direction: {direction}")
        
        # Add safety check to prevent race conditions
        if not tags_to_move and not select_all:
            logging.warning("No tags to move and select_all is False, returning current state")
            return jsonify({
                'success': True,
                'available_tags': available_tag_names,
                'selected_tags': selected_tags
            })
        
        # Save current state for undo using the dedicated endpoint
        if REQUESTS_AVAILABLE and requests:
            try:
                undo_response = requests.post(
                    f"http://127.0.0.1:{app.config.get('PORT', 8001)}/api/save-selection-state",
                    json={'action_type': 'move_tags'},
                    headers={'Content-Type': 'application/json'}
                )
                if undo_response.ok:
                    logging.info(f"Selection state saved for undo - Stack size: {undo_response.json().get('undo_stack_size', 0)}")
                else:
                    logging.warning(f"Failed to save selection state for undo: {undo_response.status_code}")
            except Exception as e:
                logging.warning(f"Failed to save selection state for undo: {str(e)}")
                # Continue with the operation even if undo save fails
        else:
            logging.debug("Requests library not available; skipping undo state persistence")
        
        if direction == 'to_selected':
            if select_all:
                # Ensure no duplicates when selecting all
                seen = set()
                deduplicated_tags = []
                for tag in available_tag_names:
                    if tag not in seen:
                        deduplicated_tags.append(tag)
                        seen.add(tag)
                excel_processor.selected_tags = deduplicated_tags
                logging.info(f"Move tags - Select all: Added {len(deduplicated_tags)} tags to selected")
            else:
                added_count = 0
                for tag in tags_to_move:
                    if tag not in excel_processor.selected_tags:
                        excel_processor.selected_tags.append(tag)
                        added_count += 1
                logging.info(f"Move tags - To selected: Added {added_count} tags to selected")
        else:  # to_available
            if select_all:
                removed_count = len(excel_processor.selected_tags)
                excel_processor.selected_tags.clear()
                logging.info(f"Move tags - Select all: Removed {removed_count} tags from selected")
            else:
                before_count = len(excel_processor.selected_tags)
                # Add safety check to prevent corruption of selected tags
                if not isinstance(excel_processor.selected_tags, list):
                    logging.error("selected_tags is not a list, resetting to empty list")
                    excel_processor.selected_tags = []
                    before_count = 0
                
                excel_processor.selected_tags = [tag for tag in excel_processor.selected_tags if tag not in tags_to_move]
                after_count = len(excel_processor.selected_tags)
                removed_count = before_count - after_count
                logging.info(f"Move tags - To available: Removed {removed_count} tags from selected (before: {before_count}, after: {after_count})")
        
        # Update session with new selected tags (store only tag names to reduce session size)
        # Add safety check to ensure selected_tags is a list before copying
        if isinstance(excel_processor.selected_tags, list):
            session['selected_tags'] = excel_processor.selected_tags.copy()
        else:
            logging.error("selected_tags is not a list, setting session to empty list")
            session['selected_tags'] = []
        
        # Return only the necessary data for UI updates
        # Add safety checks to ensure we return valid data
        if not isinstance(excel_processor.selected_tags, list):
            logging.error("selected_tags is not a list in final response, using empty list")
            excel_processor.selected_tags = []
        
        updated_available_names = [name for name in available_tag_names if name not in excel_processor.selected_tags]
        updated_selected_names = excel_processor.selected_tags.copy()
        
        logging.info(f"Move tags - Final response: {len(updated_available_names)} available, {len(updated_selected_names)} selected")
        
        return jsonify({
            'success': True,
            'available_tags': updated_available_names,
            'selected_tags': updated_selected_names
        })
        
    except Exception as e:
        logging.error(f"Error in move_tags: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/undo-move', methods=['POST'])
def undo_move():
    try:
        # Check session size but don't optimize unless necessary
        check_session_size()
        
        excel_processor = get_session_excel_processor()
        undo_stack = session.get(UNDO_STACK_KEY, [])
        
        # Debug logging for undo stack
        logging.info(f"Undo move requested - Stack size: {len(undo_stack)}")
        logging.info(f"Undo stack contents: {undo_stack}")
        
        if not undo_stack:
            logging.warning("No undo history available - user tried to undo without any previous moves")
            return jsonify({'error': 'No undo history available'}), 400
        
        # Get the last state
        last_state = undo_stack.pop()
        session[UNDO_STACK_KEY] = undo_stack
        
        # Restore the previous state
        excel_processor.selected_tags = last_state['selected_tag_names'].copy()
        session['selected_tags'] = excel_processor.selected_tags.copy()
        
        # Get current available tags
        available_tags = excel_processor.get_available_tags()
        available_tag_names = [tag.get('Product Name*', '') for tag in available_tags if tag.get('Product Name*', '')]
        
        # Return only the necessary data for UI updates
        updated_available_names = [name for name in available_tag_names if name not in excel_processor.selected_tags]
        updated_selected_names = excel_processor.selected_tags.copy()
        
        return jsonify({
            'success': True,
            'available_tags': updated_available_names,
            'selected_tags': updated_selected_names
        })
        
    except Exception as e:
        logging.error(f"Error in undo_move: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/save-selection-state', methods=['POST'])
def save_selection_state():
    """Save the current selection state for undo functionality."""
    try:
        # Check session size but don't optimize unless necessary
        check_session_size()
        
        data = request.get_json()
        action_type = data.get('action_type', 'checkbox_selection')  # 'checkbox_selection', 'move', etc.
        
        excel_processor = get_session_excel_processor()
        available_tags = excel_processor.get_available_tags()
        
        # Convert available_tags to just names for efficiency
        available_tag_names = [tag.get('Product Name*', '') for tag in available_tags if tag.get('Product Name*', '')]
        
        # Get selected tags - handle both dict and string objects
        selected_tags = []
        for tag in excel_processor.selected_tags:
            if isinstance(tag, dict):
                selected_tags.append(tag.get('Product Name*', ''))
            elif isinstance(tag, str):
                selected_tags.append(tag)
            else:
                selected_tags.append(str(tag))
        
        # Save current state for undo (store only tag names to reduce session size)
        undo_stack = session.get(UNDO_STACK_KEY, [])
        undo_stack.append({
            'available_tag_names': available_tag_names,
            'selected_tag_names': selected_tags.copy(),
            'action_type': action_type,
            'timestamp': datetime.now().isoformat()
        })
        # Limit undo stack size to prevent session bloat
        if len(undo_stack) > 5:
            undo_stack = undo_stack[-5:]
        session[UNDO_STACK_KEY] = undo_stack
        
        # Debug logging for undo stack
        logging.info(f"Selection state saved - Stack size: {len(undo_stack)}, Action type: {action_type}")
        logging.info(f"Current selected tags: {len(selected_tags)}")
        
        return jsonify({
            'success': True,
            'undo_stack_size': len(undo_stack)
        })
        
    except Exception as e:
        logging.error(f"Error in save_selection_state: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/update-selected-order', methods=['POST'])
def update_selected_order():
    """Update the order of selected tags."""
    try:
        # Check session size but don't optimize unless necessary
        check_session_size()
        
        data = request.get_json()
        new_order = data.get('order', [])
        
        if not new_order:
            return jsonify({'error': 'No order provided'}), 400
        
        excel_processor = get_session_excel_processor()
        
        # Get current selected tags as names
        current_selected = []
        for tag in excel_processor.selected_tags:
            if isinstance(tag, dict):
                current_selected.append(tag.get('Product Name*', ''))
            elif isinstance(tag, str):
                current_selected.append(tag)
            else:
                current_selected.append(str(tag))
        
        # Validate that all items in new_order exist in current_selected
        current_selected_set = set(current_selected)
        new_order_valid = [tag for tag in new_order if tag in current_selected_set]
        
        # If no valid tags found, use the original order
        if not new_order_valid:
            new_order_valid = current_selected.copy()
        else:
            # Add any missing tags from current selection (avoiding duplicates)
            for tag in current_selected:
                if tag not in new_order_valid:
                    new_order_valid.append(tag)
        
        # Ensure no duplicates in the final list
        seen = set()
        deduplicated_order = []
        for tag in new_order_valid:
            if tag not in seen:
                deduplicated_order.append(tag)
                seen.add(tag)
        new_order_valid = deduplicated_order
        
        # Update the selected tags order - store only tag names
        excel_processor.selected_tags = new_order_valid
        # Update session with the new order (only names)
        session['selected_tags'] = new_order_valid
        
        # Force session to be saved
        session.modified = True
        
        logging.info(f"Updated selected tags order: {new_order_valid}")
        
        return jsonify({
            'success': True,
            'message': 'Selected tags order updated successfully',
            'selected_tags': new_order_valid
        })
        
    except Exception as e:
        logging.error(f"Error in update_selected_order: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/clear-filters', methods=['POST'])
def clear_filters():
    try:
        logging.info("=== CLEAR FILTERS ACTION START ===")
        logging.info(f"Clear filters request at {datetime.now().strftime('%H:%M:%S')}")
        
        # Check and optimize session size before processing
        check_session_size()
        optimize_session_data()
        
        # Store validation to prevent cross-store data access
        # Store context removed - using single database
        file_store = session.get('file_store', '')
        
        # Store validation removed - using single database
        
        excel_processor = get_session_excel_processor()
        # CRITICAL FIX: Don't clear selected tags if they were set by JSON matching
        json_match_timestamp = session.get('json_match_timestamp', 0)
        current_time = time.time()
        
        # Only clear if no recent JSON matching (within last 5 minutes)
        if current_time - json_match_timestamp > 300:  # 5 minutes
            excel_processor.selected_tags.clear()
            session['selected_tags'] = []
            session['json_selected_tags'] = []
            session['last_json_match_count'] = 0
            session['json_match_timestamp'] = 0
            logging.info("Cleared selected tags - no recent JSON matching")
        else:
            logging.info(f"Preserving selected tags from recent JSON matching ({current_time - json_match_timestamp:.1f}s ago)")
        
        session[UNDO_STACK_KEY] = []
        excel_processor.dropdown_cache = {}
        json_matcher = get_session_json_matcher()
        json_matcher.clear_matches()
        available_tags = excel_processor.get_available_tags()
        
        # Get available tag names for frontend
        available_tag_names = [tag.get('Product Name*', '') for tag in available_tags if tag.get('Product Name*', '')]
        
        logging.info(f"Cleared all filters and selected tags. Available tags: {len(available_tag_names)}")
        
        return jsonify({
            'success': True,
            'available_tags': available_tag_names,
            'selected_tags': [],
            'filters': excel_processor.dropdown_cache
        })
    except Exception as e:
        logging.error(f"Error clearing filters: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/pending-changes', methods=['GET'])
@cached_route(ttl=5, vary_by=['session_id'])
def get_pending_changes():
    """Get pending database changes for the current session."""
    try:
        from src.core.data.session_manager import get_pending_changes
        changes = get_pending_changes()
        
        # Convert changes to serializable format
        serializable_changes = []
        for change in changes:
            serializable_changes.append({
                'change_type': change.change_type,
                'entity_id': change.entity_id,
                'entity_type': change.entity_type,
                'timestamp': change.timestamp.isoformat(),
                'user_id': change.user_id,
                'details': change.details
            })
        
        return jsonify({
            'success': True,
            'changes': serializable_changes,
            'change_count': len(serializable_changes)
        })
    except Exception as e:
        logging.error(f"Error getting pending changes: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/session-stats', methods=['GET'])
@cached_route(ttl=5, vary_by=['session_id'])
def get_session_stats():
    """Get session statistics."""
    try:
        from src.core.data.session_manager import get_session_manager
        session_manager = get_session_manager()
        stats = session_manager.get_session_stats()
        
        return jsonify({
            'success': True,
            'stats': stats
        })
    except Exception as e:
        logging.error(f"Error getting session stats: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/set-store', methods=['POST'])
@invalidate_cache_on_change([
    '/api/status',
    '/api/get-store',
    '/api/check-store-required',
    '/api/session-stats',
    '/api/pending-changes'
])
def set_store():
    """Set store selection for the current IP address."""
    try:
        data = request.get_json()
        if not data or 'store' not in data:
            return jsonify({'success': False, 'error': 'Store selection required'}), 400
        
        store_value = data['store']
        ip_address = get_client_ip()
        
        # Validate store selection against global list
        if store_value not in VALID_STORES:
            return jsonify({'success': False, 'error': 'Invalid store selection'}), 400
        
        # CHECK: Warn if switching stores
        current_store = session.get('selected_store')
        if current_store and current_store != store_value:
            logging.warning(f"⚠️ STORE SWITCH DETECTED: {current_store} → {store_value}")
            logging.warning(f"⚠️ Request from: {request.referrer or 'unknown'}")
            logging.warning(f"⚠️ User agent: {request.headers.get('User-Agent', 'unknown')}")
        
        # CRITICAL FIX: Save to Flask session first (most reliable on PythonAnywhere)
        session['selected_store'] = store_value
        session['store_server_id'] = SERVER_INSTANCE_ID
        session['store_just_selected'] = True  # Flag to indicate store was just selected
        session['store_selected_timestamp'] = datetime.now().isoformat()  # Timestamp for validation
        session.permanent = True  # Mark session as permanent to persist across browser restarts
        session.modified = True  # Force session to save
        logging.info(f"✅ Store saved to session: {store_value}")
        
        # Also store in IP-based selection (backup method)
        with _ip_store_lock:
            _ip_store_selections[ip_address] = {
                'store': store_value,
                'timestamp': datetime.now().isoformat(),
                'ip_address': ip_address,
                'server_id': SERVER_INSTANCE_ID
            }
            # Reduced logging for speed
            logging.debug(f"Store selection set for IP {ip_address}: {store_value}")
        
        # Persist the IP-based selections so future requests (and worker processes)
        # can honor the remembered store.  Use a lightweight background thread so
        # the response stays snappy even on PythonAnywhere's slower storage.
        def _persist_store_selection():
            try:
                save_store_selections()
            except Exception as persist_error:
                logging.warning(f"Failed to persist store selections: {persist_error}")
        try:
            threading.Thread(target=_persist_store_selection, daemon=True).start()
        except Exception:
            _persist_store_selection()
        
        # CRITICAL FIX: Clear caches BEFORE clearing globals (cache key depends on _last_loaded_file)
        # Get cache keys while excel_processor still has the old file path
        try:
            initial_data_cache_key = get_session_cache_key('initial_data')
            available_tags_cache_key = get_session_cache_key('available_tags')

            # Delete the caches
            cache.delete(initial_data_cache_key)
            cache.delete(available_tags_cache_key)
            logging.debug(f"Cleared initial_data and available_tags cache for new store: {store_value}")
        except Exception as cache_error:
            # If cache key generation fails (e.g., Excel processor is slow), clear all session caches
            logging.warning(f"Failed to get specific cache keys, clearing all caches: {cache_error}")
            try:
                # Clear all caches for this session
                sid = session.get('_id', None) or (session.sid if hasattr(session, 'sid') else 'unknown')
                # Use a wildcard pattern or just clear the simple keys
                for key_base in ['initial_data', 'available_tags', 'web_available_tags']:
                    try:
                        # Try with empty file path
                        simple_key = hashlib.sha256(f"{key_base}:{sid}:".encode()).hexdigest()
                        cache.delete(simple_key)
                    except:
                        pass
            except Exception as e2:
                logging.warning(f"Failed to clear caches with fallback method: {e2}")

        # CRITICAL: Clear other session data from previous store (but keep selected_store!)
        session.pop('file_path', None)
        session.pop('uploaded_filename', None)
        session.pop('upload_timestamp', None)
        session.pop('selected_tags', None)

        # Clear the global product database instance to force reload with new store
        global _product_database, _excel_processor
        _product_database = None
        _excel_processor = None

        # OPTIMIZATION: File loading deferred to page reload for instant response
        logging.debug(f"Store set to {store_value} - cleared session, globals & caches")
        
        return jsonify({
            'success': True,
            'store': store_value,
            'expires_at': (datetime.now() + timedelta(hours=12)).isoformat()
        })
        
    except Exception as e:
        logging.error(f"Error setting store: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/get-store', methods=['GET'])
@cached_route(ttl=30, vary_by=['session_id'])
def get_store():
    """Get the current store selection for the IP address."""
    try:
        ip_address = get_client_ip()
        logging.info(f"Getting store for IP: {ip_address}")
        logging.info(f"Current store selections: {list(_ip_store_selections.keys())}")
        
        # First, check Flask session for a store selection tied to this server instance
        session_store = session.get('selected_store')
        session_server_id = session.get('store_server_id')
        if session_store and session_server_id == SERVER_INSTANCE_ID:
            logging.info("Returning store from session for get_store endpoint")
            return jsonify({
                'success': True,
                'store': session_store,
                'source': 'session'
            })
        
        # Check if there's a stored selection for this IP
        with _ip_store_lock:
            if ip_address in _ip_store_selections:
                store_data = _ip_store_selections[ip_address]
                # Check if the selection is still valid (not expired)
                if is_store_selection_valid(ip_address, store_data):
                    logging.info(f"Found valid store selection: {store_data['store']}")
                    return jsonify({
                        'success': True,
                        'store': store_data['store'],
                        'expires_at': (datetime.fromisoformat(store_data['timestamp']) + timedelta(hours=12)).isoformat()
                    })
                else:
                    logging.info(f"Store selection expired for IP {ip_address}; removing")
                    del _ip_store_selections[ip_address]
        
        # No valid selection found, return no store
        logging.info(f"No store found for IP {ip_address}")
        return jsonify({
            'success': True,
            'store': None
        })
        
    except Exception as e:
        logging.error(f"Error getting store: {str(e)}")
        import traceback
        logging.error(f"Get store error traceback: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/clear-store', methods=['POST'])
@invalidate_cache_on_change([
    '/api/status',
    '/api/get-store',
    '/api/check-store-required',
    '/api/session-stats',
    '/api/pending-changes'
])
def clear_store():
    """Clear store selection for the current IP address AND Flask session."""
    try:
        ip_address = get_client_ip()
        
        logging.info(f"Attempting to clear store for IP {ip_address}")
        logging.info(f"Current store selections: {list(_ip_store_selections.keys())}")
        logging.info(f"Session store before clear: {session.get('selected_store')}")
        
        # CRITICAL: Clear Flask session FIRST (this is what has_store_selection checks first)
        if 'selected_store' in session:
            del session['selected_store']
            logging.info("Cleared 'selected_store' from Flask session")
        if 'store_server_id' in session:
            del session['store_server_id']
        
        # Also clear IP-based storage
        with _ip_store_lock:
            if ip_address in _ip_store_selections:
                del _ip_store_selections[ip_address]
                logging.info(f"Store selection cleared for IP {ip_address}")
                # Save to disk after clearing
                save_store_selections()
            else:
                logging.info(f"No store selection found for IP {ip_address}")
        
        logging.info(f"Store selections after clear: {list(_ip_store_selections.keys())}")
        logging.info(f"Session store after clear: {session.get('selected_store')}")
        
        return jsonify({
            'success': True,
            'message': 'Store selection cleared from both session and IP storage',
            'ip_address': ip_address,
            'remaining_selections': list(_ip_store_selections.keys())
        })
        
    except Exception as e:
        logging.error(f"Error clearing store: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/check-store-required', methods=['GET'])
@cached_route(ttl=15, vary_by=['session_id'])
def check_store_required():
    """Check if store selection is required for the current IP address."""
    try:
        ip_address = get_client_ip()
        logging.info(f"Check store required for IP: {ip_address}")
        logging.info(f"Store selections in memory: {list(_ip_store_selections.keys())}")
        
        # CRITICAL DEBUG: Check session data
        session_store = session.get('selected_store')
        if session_store and session.get('store_server_id') != SERVER_INSTANCE_ID:
            logging.warning(f"🔥 Session store from previous server instance detected - CLEARING IT to force store modal: {session_store}")
            # Clear the session store to force user to select store again after server restart
            session.pop('selected_store', None)
            session.pop('store_server_id', None)
            session.modified = True
            session_store = None  # Update local variable so logic below works correctly
        
        # CRITICAL FIX: Check for force_store_modal parameter to force modal display
        force_modal = request.args.get('force_store_modal', 'false').lower() == 'true'
        if force_modal:
            logging.info("🔧 Force store modal parameter detected - clearing session store")
            session.pop('selected_store', None)
            session.pop('store_server_id', None)
            session.modified = True
            session_store = None
        
        # CRITICAL FIX: Don't log full session - it contains massive preroll_original_records array
        # that causes "OSError: Message too long" when logging
        session_keys = list(session.keys())
        logging.info(f"SESSION DEBUG: selected_store={session_store}")
        logging.info(f"SESSION DEBUG: session keys={session_keys}")
        logging.info(f"SESSION DEBUG: session.permanent={session.permanent}")
        logging.info(f"SESSION DEBUG: force_modal={force_modal}")
        
        # CRITICAL FIX: Always require modal unless store was just selected with proper flags
        # This ensures modal shows for all existing sessions that don't have the new flags
        store_just_selected = session.get('store_just_selected', False)
        current_store = None
        
        # Only use session store if BOTH conditions are met:
        # 1. store_just_selected flag is True (explicit selection)
        # 2. Has a valid recent timestamp (within last 6 hours - matches session lifetime)
        if session_store and store_just_selected:
            store_timestamp = session.get('store_selected_timestamp')
            if store_timestamp:
                try:
                    from datetime import datetime, timedelta
                    timestamp = datetime.fromisoformat(store_timestamp)
                    # CRITICAL FIX: Increased from 10 minutes to 6 hours to match session lifetime
                    # This prevents the modal from reappearing while user is actively using the app
                    if datetime.now() - timestamp < timedelta(hours=6):
                        current_store = session_store
                        logging.info(f"Store found in session with valid flags: {current_store}")
                    else:
                        logging.info(f"Store timestamp expired (older than 6 hours), requiring new selection")
                        # Clear all store-related session data
                        session.pop('selected_store', None)
                        session.pop('store_selected_timestamp', None)
                        session.pop('store_just_selected', None)
                        session.pop('store_server_id', None)
                        session.modified = True
                except Exception as e:
                    # If timestamp parsing fails, treat as stale
                    logging.info(f"Store timestamp invalid ({e}), requiring new selection")
                    session.pop('selected_store', None)
                    session.pop('store_selected_timestamp', None)
                    session.pop('store_just_selected', None)
                    session.pop('store_server_id', None)
                    session.modified = True
            else:
                # No timestamp - treat as stale
                logging.info(f"Store in session has no timestamp, requiring new selection")
                session.pop('selected_store', None)
                session.pop('store_just_selected', None)
                session.pop('store_server_id', None)
                session.modified = True
        elif session_store:
            # Store exists but missing required flags - clear it to force modal
            logging.info(f"Store in session missing required flags (store_just_selected={store_just_selected}), clearing to force modal")
            session.pop('selected_store', None)
            session.pop('store_selected_timestamp', None)
            session.pop('store_just_selected', None)
            session.pop('store_server_id', None)
            session.modified = True
        
        # Log the low-level selection flag for debugging but do not gate on it
        has_selection = has_store_selection()
        logging.info(f"has_store_selection() returned: {has_selection}")
        logging.info(f"FINAL DECISION: current_store={current_store}, will require_store={current_store is None}")
        
        if not current_store:
            logging.warning(f"⚠️ NO VALID STORE - Requiring store selection for IP {ip_address}")
            # CRITICAL: Make sure session is completely cleared
            session.pop('selected_store', None)
            session.pop('store_just_selected', None)
            session.pop('store_selected_timestamp', None)
            session.pop('store_server_id', None)
            session.modified = True
            
            # Double-check session is cleared
            remaining_store = session.get('selected_store')
            if remaining_store:
                logging.error(f"❌ ERROR: Session still has store after clearing: {remaining_store}")
                # Force clear again
                session.clear()
                session.modified = True
            
            logging.info(f"✅ Session cleared, returning requires_store=True")
            return {
                'success': True,
                'requires_store': True,  # CRITICAL: Must be True to show modal
                'store': None,
                'debug': {
                    'session_store': session_store,
                    'ip_address': ip_address,
                    'has_selection': has_selection,
                    'cleared_session': True,
                    'current_store_after_check': current_store
                }
            }
        
        # If we found a store in session, make sure it is persisted with timestamp
        # (This should already be set, but ensure it's there)
        if current_store:
            session['selected_store'] = current_store
            if not session.get('store_just_selected'):
                session['store_just_selected'] = True
            if not session.get('store_selected_timestamp'):
                from datetime import datetime
                session['store_selected_timestamp'] = datetime.now().isoformat()
            session.modified = True
            
            logging.info(f"✅ Store found in session for IP {ip_address}: {current_store}")
            return {
                'success': True,
                'requires_store': False,  # CRITICAL: Must be False when store exists
                'store': current_store,
                'debug': {
                    'session_store': session_store,
                    'ip_address': ip_address,
                    'has_selection': has_selection,
                    'current_store_validated': True
                }
            }
        else:
            # This should never happen, but safety check
            logging.error(f"❌ ERROR: current_store is None but we reached this point!")
            session.pop('selected_store', None)
            session.pop('store_just_selected', None)
            session.pop('store_selected_timestamp', None)
            session.modified = True
            return {
                'success': True,
                'requires_store': True,
                'store': None,
                'debug': {
                    'error': 'Unexpected state - no store found',
                    'ip_address': ip_address
                }
            }
        
    except Exception as e:
        logging.error(f"Error checking store requirement: {str(e)}")
        logging.error(traceback.format_exc())
        return {'success': False, 'error': str(e)}, 500

@app.route('/api/clear-session', methods=['POST'])
def clear_session():
    """Clear the current session."""
    try:
        from src.core.data.session_manager import get_session_manager, get_current_session_id
        session_manager = get_session_manager()
        session_id = get_current_session_id()
        
        # Clear session data
        session.clear()
        
        # Clear session in manager
        session_manager.clear_session(session_id)
        
        return jsonify({
            'success': True,
            'message': 'Session cleared successfully'
        })
    except Exception as e:
        logging.error(f"Error clearing session: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/debug-product-lineage', methods=['POST'])
def debug_product_lineage():
    """Debug endpoint to show actual database values for products."""
    try:
        data = request.get_json() or {}
        product_names = data.get('product_names', [])
        
        if not product_names:
            return jsonify({'error': 'product_names array required'}), 400
        
        store_name = get_current_store_name()
        product_db = get_product_database(store_name)
        
        if not product_db:
            return jsonify({'error': 'Database not available'}), 500
        
        conn = product_db._get_connection()
        cursor = conn.cursor()
        
        results = []
        for product_name in product_names:
            # Get product from products table
            cursor.execute('''
                SELECT p.id, p."Product Name*", p."Lineage" as products_lineage,
                       p."Product Strain", p.strain_id,
                       s.sovereign_lineage, s.canonical_lineage, s.strain_name
                FROM products p
                LEFT JOIN strains s ON p.strain_id = s.id
                WHERE p."Product Name*" = ? OR p."ProductName" = ?
                ORDER BY p.id DESC
                LIMIT 1
            ''', (product_name, product_name))
            
            product_row = cursor.fetchone()
            
            if product_row:
                product_id, db_product_name, products_lineage, product_strain, strain_id, sovereign_lineage, canonical_lineage, strain_name = product_row
                
                # Get what get_product_lineage returns
                lineage_from_method = product_db.get_product_lineage(product_name)
                
                # Get what get_products_by_names returns
                products_from_method = product_db.get_products_by_names([product_name])
                lineage_from_batch = None
                if products_from_method:
                    lineage_from_batch = (
                        products_from_method[0].get('currentLineage') or
                        products_from_method[0].get('canonical_lineage') or
                        products_from_method[0].get('Lineage')
                    )
                
                results.append({
                    'product_name': product_name,
                    'database_values': {
                        'products_table': {
                            'id': product_id,
                            'Product Name*': db_product_name,
                            'Lineage': products_lineage,
                            'Product Strain': product_strain,
                            'strain_id': strain_id
                        },
                        'strains_table': {
                            'strain_name': strain_name,
                            'sovereign_lineage': sovereign_lineage,
                            'canonical_lineage': canonical_lineage
                        }
                    },
                    'method_results': {
                        'get_product_lineage': lineage_from_method,
                        'get_products_by_names': {
                            'currentLineage': products_from_method[0].get('currentLineage') if products_from_method else None,
                            'canonical_lineage': products_from_method[0].get('canonical_lineage') if products_from_method else None,
                            'Lineage': products_from_method[0].get('Lineage') if products_from_method else None
                        }
                    },
                    'effective_lineage': lineage_from_method or lineage_from_batch or products_lineage or sovereign_lineage or canonical_lineage
                })
            else:
                results.append({
                    'product_name': product_name,
                    'error': 'Product not found in database'
                })
        
        return jsonify({
            'success': True,
            'store_name': store_name,
            'database_path': product_db.db_path,
            'products': results
        })
        
    except Exception as e:
        logging.error(f"Error debugging product lineage: {e}")
        import traceback
        logging.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

# REMOVED: Duplicate endpoint - the comprehensive version is at line ~11959
# This simpler version has been replaced by a more complete implementation
# that also updates strain lineage and includes verification

@app.route('/api/debug-database', methods=['GET'])
def debug_database():
    """Debug endpoint to check database state and session info."""
    try:
        # Store context removed - using single database
        session_store = session.get('file_store', '')
        
        # Check which database would be loaded
        db_info = {}
        
        # Check default database
        default_db_path = os.path.join(current_dir, 'uploads', 'product_database.db')
        db_info['default_db'] = {
            'path': default_db_path,
            'exists': os.path.exists(default_db_path),
            'size': os.path.getsize(default_db_path) if os.path.exists(default_db_path) else 0
        }
        
        # Check store-specific database
        # Store context removed - using single database
        
        # Check global database instance
        global _product_database
        db_info['global_instance'] = {
            'exists': _product_database is not None,
            'store_name': getattr(_product_database, '_store_name', None) if _product_database else None
        }
        
        # Test database connection
        try:
            store_name = get_current_store_name()
            test_db = get_product_database(store_name)
            db_info['connection_test'] = {
                'success': True,
                'db_path': test_db.db_path,
                'initialized': test_db._initialized
            }
        except Exception as db_error:
            db_info['connection_test'] = {
                'success': False,
                'error': str(db_error)
            }
        
        return jsonify({
            'success': True,
            'session': {
                # Store context removed - using single database
                'file_store': session_store,
                'file_path': session.get('file_path', '')
            },
            'database_info': db_info
        })
        
    except Exception as e:
        logging.error(f"Error in debug database: {e}")
        return jsonify({'error': str(e)}), 500

logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)

@lru_cache(maxsize=32)
def get_cached_font_scheme(template_type, base_size=12):
    from src.core.generation.template_processor import get_font_scheme
    return get_font_scheme(template_type, base_size)

def copy_cell_content(src_cell, dst_cell):
    dst_cell._element.clear_content()
    # Set cell alignment to center
    dst_cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    for child in src_cell._element:
        dst_cell._element.append(copy.deepcopy(child))
    # Center all paragraphs in the cell
    for paragraph in dst_cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Center all runs in the paragraph
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.bold = True

def rebuild_3x3_grid_from_template(doc, template_path):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches
    from docx.enum.table import WD_ROW_HEIGHT_RULE

    # Load the template and get the first table/cell
    template_doc = Document(template_path)
    old_table = template_doc.tables[0]
    source_cell_xml = deepcopy(old_table.cell(0, 0)._tc)

    # Remove all existing tables in doc
    for table in doc.tables:
        table._element.getparent().remove(table._element)

    # Add new fixed 3x3 table
    table = doc.add_table(rows=3, cols=3)
    table.autofit = False
    table.allow_autofit = False
    tblPr = table._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    table._element.insert(0, tblPr)
    tblGrid = OxmlElement('w:tblGrid')
    # Each cell should be 3.4 inches wide (not divided by 3!)
    col_width_twips = str(int(3.4 * 1440))  # Fixed: was incorrectly (3.4/3) * 1440
    for _ in range(3):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), col_width_twips)
        tblGrid.append(gridCol)
    table._element.insert(0, tblGrid)
    for i in range(3):
        for j in range(3):
            cell = table.cell(i, j)
            cell._tc.clear_content()
            new_tc = deepcopy(source_cell_xml)
            # Replace Label1 with LabelN in the XML
            label_num = i * 3 + j + 1
            for text_el in new_tc.iter():
                if text_el.tag == qn('w:t') and text_el.text:
                    logging.debug(f"Processing text element: {text_el.text}")
                    if "Label1" in text_el.text:
                        text_el.text = text_el.text.replace("Label1", f"Label{label_num}")
                        logging.info(f"Replaced Label1 with Label{label_num} in text element.")
            cell._tc.extend(new_tc.xpath('./*'))
        row = table.rows[i]
        row.height = Inches(2.4)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    # Enforce fixed cell dimensions to prevent any growth
    try:
        # Safety check: ensure table has valid structure
        if table and table.rows and len(table.rows) > 0:
            first_row = table.rows[0]
            if hasattr(first_row, '_element') and hasattr(first_row._element, 'tc_lst'):
                enforce_fixed_cell_dimensions(table, 'horizontal')  # Default to horizontal for 3x3 grids
            else:
                print("Warning: Skipping table with invalid XML structure in app.py")
        else:
            print("Warning: Skipping empty or invalid table in app.py")
    except Exception as e:
        print(f"Warning: Error enforcing fixed cell dimensions in app.py: {e}")
    
    return table

def post_process_document(doc, font_scheme, orientation, scale_factor):
    """
    Main post-processing function, inspired by the old MAIN.py logic.
    This function finds and formats all marked fields in the document.
    Uses template-type-specific font sizing based on the unified font-sizing system.
    """

    # Define marker processing for all template types
    markers = [
        'DESC', 'PRODUCTBRAND_CENTER', 'PRICE', 'LINEAGE', 
        'THC_CBD', 'RATIO', 'PRODUCTSTRAIN', 'DOH'
    ]

    # Process each marker type recursively through the document using template-specific font sizing
    for marker_name in markers:
        _autosize_recursive_template_specific(doc, marker_name, orientation, scale_factor)

    # Apply final conditional formatting for colors, etc.
    from src.core.generation.docx_formatting import apply_lineage_colors
    apply_lineage_colors(doc)
    return doc
def _autosize_recursive_template_specific(element, marker_name, orientation, scale_factor):
    """
    Recursively search for and format a specific marked field within a document element using template-specific font sizing.
    """
    from src.core.generation.unified_font_sizing import (
        get_font_size,
        set_run_font_size
    )
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    start_marker = f'{marker_name}_START'
    end_marker = f'{marker_name}_END'

    if hasattr(element, 'paragraphs'):
        for p in element.paragraphs:
            # Reassemble full text from runs to handle split markers
            full_text = "".join(run.text for run in p.runs)

            if start_marker in full_text and end_marker in full_text:
                # Extract content
                start_idx = full_text.find(start_marker) + len(start_marker)
                end_idx = full_text.find(end_marker)
                content = full_text[start_idx:end_idx].strip()

                if content:
                    # Calculate font size using template-specific sizing
                    font_size = _get_template_specific_font_size(content, marker_name, orientation, scale_factor)
                    
                    # Rewrite the paragraph with clean content and new font size
                    p.clear()
                    
                    # Handle line breaks for THC/CBD content
                    if marker_name in ['THC_CBD', 'RATIO'] and '\n' in content:
                        parts = content.split('\n')
                        for i, part in enumerate(parts):
                            if i > 0:
                                run = p.add_run()
                                run.add_break()
                            run = p.add_run(part)
                            run.font.name = "Arial"
                            run.font.bold = True
                            run.font.size = font_size
                            set_run_font_size(run, font_size)
                    else:
                        run = p.add_run(content)
                        run.font.name = "Arial"
                        run.font.bold = True
                        run.font.size = font_size
                        set_run_font_size(run, font_size)
                    
                    # Handle special paragraph properties
                    if marker_name == 'PRODUCTBRAND_CENTER':
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if marker_name == 'THC_CBD':
                        p.paragraph_format.line_spacing = 1.5
                else:
                    # If there's no content, just remove the markers
                    p.clear()

    if hasattr(element, 'tables'):
        for table in element.tables:
            try:
                # Safety check: ensure table has valid structure
                if table and table.rows and len(table.rows) > 0:
                    first_row = table.rows[0]
                    if hasattr(first_row, '_element') and hasattr(first_row._element, 'tc_lst'):
                        for row in table.rows:
                            try:
                                for cell in row.cells:
                                    try:
                                        # Continue the recursion into cells
                                        _autosize_recursive_template_specific(cell, marker_name, orientation, scale_factor)
                                    except Exception as cell_error:
                                        print(f"Warning: Error processing cell: {cell_error}")
                                        continue
                            except Exception as row_error:
                                print(f"Warning: Error processing row: {row_error}")
                                continue
                    else:
                        print("Warning: Skipping table with invalid XML structure in app.py recursion")
                else:
                    print("Warning: Skipping empty or invalid table in app.py recursion")
            except Exception as table_error:
                print(f"Warning: Error processing table in app.py recursion: {table_error}")
                continue

def _get_template_specific_font_size(content, marker_name, orientation, scale_factor):
    """
    Get font size using the unified font sizing system.
    """
    from src.core.generation.unified_font_sizing import get_font_size
    
    # Map marker names to field types
    marker_to_field_type = {
        'DESC': 'description',
        'PRODUCTBRAND_CENTER': 'brand',
        'PRICE': 'price',
        'LINEAGE': 'lineage',
        'THC_CBD': 'thc_cbd',
        'RATIO': 'ratio',
        'PRODUCTSTRAIN': 'strain',
        'DOH': 'doh'
    }
    
    field_type = marker_to_field_type.get(marker_name, 'default')
    
    # Use unified font sizing with appropriate complexity type
    complexity_type = 'mini' if orientation == 'mini' else 'standard'
    return get_font_size(content, field_type, orientation, scale_factor, complexity_type)

def _extract_product_name_from_full_name(full_name):
    """Extract just the product name from 'Product Name by Vendor - Weight' format."""
    if not full_name or str(full_name).strip() == '':
        return ''
    
    name = str(full_name).strip()
    if not name:
        return ''
    
    # Handle "Product Name by Vendor - Weight" format
    if ' by ' in name and ' - ' in name:
        # Extract just the product name part before " by "
        return name.split(' by ')[0].strip()
    elif ' by ' in name:
        return name.split(' by ')[0].strip()
    elif ' - ' in name:
        # Only split on dashes followed by weight information (numbers, decimals, units)
        if re.search(r' - [\d.]', name):
            # Remove weight part but preserve the dash in product names
            return re.sub(r' - [\d.].*$', '', name).strip()
        else:
            # No weight information, return the name as-is
            return name.strip()
        return name.strip()
# Removed duplicate function - using the more sophisticated version at line 3839

def _validate_tags_against_excel(excel_processor, selected_tags):
    """Helper function to validate tags against Excel data."""
    # CRITICAL FIX: Import re module locally to avoid scoping issues
    import re
    
    valid_selected_tags = []
    invalid_selected_tags = []
    
    # Create case-insensitive lookup map for available product names
    available_product_names_lower = {}
    # Try multiple possible column names for product names
    possible_product_name_columns = ['Product Name*', 'ProductName', 'Product Name', 'product_name']
    product_name_column = None
    
    # Find the first available column
    for col in possible_product_name_columns:
        if excel_processor.df is not None and col in excel_processor.df.columns:
            product_name_column = col
            break
    
    if product_name_column:
        for _, row in excel_processor.df.iterrows():
            # Handle pandas Series objects properly
            product_name_value = row[product_name_column]
            if isinstance(product_name_value, pd.Series):
                product_name = str(product_name_value.iloc[0]).strip() if len(product_name_value) > 0 else ''
            else:
                product_name = str(product_name_value).strip()
            if product_name and product_name != 'nan':
                # CRITICAL FIX: Store all products with the same name, not just the last one
                if product_name.lower() not in available_product_names_lower:
                    available_product_names_lower[product_name.lower()] = []
                available_product_names_lower[product_name.lower()].append(product_name)  # Store all instances
        
        logging.debug(f"Available product names count: {len(available_product_names_lower)}")
        logging.debug(f"Sample available product names: {list(available_product_names_lower.values())[:5]}")
        logging.debug(f"Using column: {product_name_column}")
    else:
        logging.warning(f"No product name column found. Available columns: {list(excel_processor.df.columns) if excel_processor.df is not None else 'No DataFrame'}")
    
    logging.debug(f"Validating {len(selected_tags)} selected tags against Excel data")
    for tag in selected_tags:
        tag_lower = tag.strip().lower()
        found_match = False  # Initialize found_match for each tag
        
        # CRITICAL FIX: Initialize clean_tag before the if/else block to avoid UnboundLocalError
        # Remove vendor suffixes for better matching
        # Common patterns: "by Vendor", " - Vendor", etc.
        clean_tag = re.sub(r'\s*(?:by|from|-\s*)([^-]*?)(?:\s*$)', '', tag_lower)
        clean_tag = clean_tag.strip()
        
        # First try exact match
        if tag_lower in available_product_names_lower:
            # Use all original cases from Excel data (now a list)
            original_case_tags = available_product_names_lower[tag_lower]
            for original_case_tag in original_case_tags:
                valid_selected_tags.append(original_case_tag)
                logging.debug(f"Found exact tag '{tag}' -> using original case: '{original_case_tag}'")
            found_match = True  # Mark as found since we found an exact match
        else:
            # Try partial matching - the frontend might send clean names while Excel has "Product Name by Vendor"
            
            for excel_name, original_names in available_product_names_lower.items():
                # Check if the frontend tag is contained within the Excel product name
                if tag_lower in excel_name.lower():
                    for original_name in original_names:
                        valid_selected_tags.append(original_name)
                        logging.debug(f"Found partial match '{tag}' -> contained in Excel name: '{original_name}'")
                        found_match = True
                # CRITICAL FIX: Also try matching with vendor suffix removed
                elif clean_tag in excel_name.lower():
                    for original_name in original_names:
                        valid_selected_tags.append(original_name)
                        logging.debug(f"Found match with vendor suffix removed '{tag}' (cleaned: '{clean_tag}') -> Excel name: '{original_name}'")
                        found_match = True
        
        if not found_match:
            invalid_selected_tags.append(tag.strip())
            logging.warning(f"Selected tag not found in Excel data: '{tag}' (lowercase: '{tag_lower}', cleaned: '{clean_tag}')")
    
    return valid_selected_tags, invalid_selected_tags

def _format_price_value(price_value):
    """Format a price value with dollar sign."""
    if not price_value or not str(price_value).strip():
        return ''
    
    try:
        # Remove $ and commas, convert to float
        price_float = float(str(price_value).replace('$', '').replace(',', '').strip())
        # Format with $ sign
        if price_float.is_integer():
            return f"${int(price_float)}"
        else:
            return f"${price_float:.2f}"
    except (ValueError, TypeError):
        # If it's not a valid number, just return as is if it has $, otherwise add $
        price_str = str(price_value).strip()
        if price_str.startswith('$'):
            return price_str
        return f"${price_str}" if price_str else ''

def _extract_price_from_database_product(product):
    """Extract price from a database product, checking multiple possible price fields."""
    # Check multiple possible price field names
    candidate_keys = [
        'Price',
        'Price*',
        'Med Price',
        'Price* (Tier Name for Bulk)'
    ]
    
    # Also check keys with 'Price' in them (case-insensitive)
    for key in list(product.keys()):
        if 'price' in key.lower() and key not in candidate_keys:
            candidate_keys.append(key)
    
    for key in candidate_keys:
        if key in product:
            price_val = product.get(key)
            if price_val is None:
                continue
            price_str = str(price_val).strip()
            if price_str and price_str.lower() not in ['none', '0', '0.0', '0.00', '', 'nan']:
                # Remove $ sign if present for consistency
                price_str = price_str.replace('$', '').strip()
                if price_str:
                    logging.info(f"💰 Found price '{price_str}' in field '{key}' for product '{product.get('Product Name*', '')}'")
                    return price_str
    
    logging.warning(f"⚠️ No price found in product: '{product.get('Product Name*', '')}' - Available keys: {list(product.keys())}")
    
    # If no price found, try to get average from similar products
    try:
        store_name = get_current_store_name()
        product_db = get_product_database(store_name)
        if product_db:
            product_name = product.get('Product Name*', product.get('ProductName', ''))
            vendor = product.get('Vendor/Supplier*', product.get('Vendor', ''))
            brand = product.get('Product Brand', product.get('brand', ''))
            if product_name:
                # Use the existing method to make educated guess from similar products
                inferred = product_db._make_educated_guess_from_similar_products(product_name, vendor, brand)
                if inferred and 'price' in inferred:
                    price_val = inferred['price']
                    try:
                        price_float = float(price_val)
                        logging.info(f"💰 Using inferred price '{price_float}' from similar products for '{product_name}'")
                        return str(int(price_float)) if price_float.is_integer() else str(price_float)
                    except (ValueError, TypeError):
                        pass
    except Exception as e:
        logging.warning(f"Error trying to infer price from similar products: {e}")
    
    # Return empty string if no price found
    return ''

def _normalize_weight_string(weight_str):
    """
    Normalize weight strings so trailing .0 values are removed (e.g., 1.0g -> 1g).
    Keeps other decimals intact (e.g., 0.5g stays 0.5g).
    """
    if weight_str is None:
        return weight_str
    s = str(weight_str).strip()
    if s == '':
        return s
    # Remove trailing .0 (or .00, etc.) when followed by units or string end
    s = re.sub(r'(\d+)\.0+(?=\s*[a-zA-Z]|$)', r'\1', s)
    return s

def _create_desc_and_weight(product_name, weight_units):
    """Create DescAndWeight field with 'Product Name - Weight' format (matching Excel processor)."""
    # CRITICAL FIX: Import re module locally to avoid scoping issues
    import re
    
    if not product_name:
        return ''

    # Clean up the product name first (remove weight info that might already be there)
    description = str(product_name).strip()
    
    # Apply Excel processor formula: Remove " by " patterns
    if " by " in description:
        description = description.split(" by ")[0].strip()
    
    # Apply Excel processor formula: Remove weight information (patterns like " - 1g", " - .5g")
    description = re.sub(r' - [\d.].*$', '', description)
    
    # Get weight units, clean them up (normalize trailing .0 cases)
    weight = _normalize_weight_string(weight_units)
    weight = str(weight).strip() if weight is not None else ''
    if weight and weight.lower() not in ['nan', 'none', 'null', '']:
        # Combine product name and weight with hyphen staying with weight (space after hyphen)
        # Use same format as Excel processor: -\u00A0 (hyphen + non-breaking space)
        return f"{description} -\u00A0{weight}"
    else:
        # Just return the product name if no weight
        return description

def _normalize_weight_fields(record):
    """
    Normalize all weight-related fields on a record so values like '1.0g'
    are rendered as '1g' and DescAndWeight stays in sync.
    """
    if not isinstance(record, dict):
        return record
    
    weight_fields = [
        'Weight*', 'Weight', 'CombinedWeight', 'WeightUnits',
        'WeightWithUnits', 'weightWithUnits'
    ]
    
    for field in weight_fields:
        if field in record:
            record[field] = _normalize_weight_string(record.get(field))
    
    # If CombinedWeight is still missing, rebuild it from Weight* and Units
    combined_weight = record.get('CombinedWeight')
    if not combined_weight or str(combined_weight).strip() == '':
        weight_value = record.get('Weight*') or record.get('Weight')
        units_value = record.get('Units', '')
        if weight_value:
            combined_weight = f"{_normalize_weight_string(weight_value)}{units_value}".strip()
            record['CombinedWeight'] = _normalize_weight_string(combined_weight)
    
    # Keep alias fields aligned
    if record.get('CombinedWeight'):
        record['WeightUnits'] = record.get('WeightUnits') or record['CombinedWeight']
        record['WeightWithUnits'] = record.get('WeightWithUnits') or record['CombinedWeight']
        record['weightWithUnits'] = record.get('weightWithUnits') or record['CombinedWeight']
    
    # Regenerate DescAndWeight with normalized weight if possible
    product_name = (
        record.get('Product Name*') or
        record.get('ProductName') or
        record.get('Product Name') or
        record.get('Description')
    )
    weight_for_desc = (
        record.get('CombinedWeight') or
        record.get('WeightUnits') or
        record.get('WeightWithUnits') or
        record.get('weightWithUnits')
    )
    if product_name and weight_for_desc:
        record['DescAndWeight'] = _create_desc_and_weight(product_name, weight_for_desc)
        if not record.get('Description'):
            record['Description'] = record['DescAndWeight']
    
    return record

def _align_tags_with_db_lineage(tags, store_name, skip_if_aligned=True):
    """
    Ensure tags shown in the UI use the latest lineage from the database.
    Returns a shallow-copied list so cached tag objects are not mutated.
    
    Args:
        tags: List of tag dictionaries
        store_name: Store name for database lookup
        skip_if_aligned: If True, skip alignment if tags already have canonical_lineage/currentLineage
    """
    if not tags or not isinstance(tags, list):
        return tags
    
    # PERFORMANCE: Skip alignment if tags already have database lineage fields
    if skip_if_aligned:
        tags_with_lineage = sum(1 for t in tags if isinstance(t, dict) and (t.get('canonical_lineage') or t.get('currentLineage')))
        if tags_with_lineage >= len(tags) * 0.9:  # 90%+ already have lineage
            logging.debug(f"⚡ Skipping lineage alignment - {tags_with_lineage}/{len(tags)} tags already have lineage")
            return tags
    
    try:
        product_db = get_product_database(store_name)
        if not product_db:
            return tags
        
        # Copy tags so we don't mutate cached objects
        aligned_tags = [tag.copy() if isinstance(tag, dict) else tag for tag in tags]
        
        # Collect product names for lookup (only those missing lineage)
        product_names = []
        for t in aligned_tags:
            if isinstance(t, dict) and t.get('Product Name*'):
                # Only align if missing canonical_lineage/currentLineage
                if not (t.get('canonical_lineage') or t.get('currentLineage')):
                    product_names.append(t.get('Product Name*'))
        
        if not product_names:
            logging.debug("⚡ No products need lineage alignment")
            return aligned_tags
        
        lineage_map = {}
        conn = product_db._get_connection()
        cursor = conn.cursor()
        
        chunk_size = 400
        for start in range(0, len(product_names), chunk_size):
            chunk = product_names[start:start + chunk_size]
            placeholders = ','.join(['?' for _ in chunk])
            cursor.execute(f'''
                SELECT "Product Name*", "Lineage", "canonical_lineage"
                FROM products
                WHERE LOWER("Product Name*") IN ({placeholders})
            ''', [name.lower() for name in chunk])
            for row in cursor.fetchall():
                db_name = row[0]
                # CRITICAL FIX: Prioritize canonical_lineage (database source of truth) over Lineage field
                # canonical_lineage is what the UI displays and should be used consistently
                db_lineage = row[2] or row[1]  # canonical_lineage first, then Lineage as fallback
                if db_name and db_lineage:
                    lineage_map[db_name.lower().strip()] = str(db_lineage).strip().upper()
        
        if not lineage_map:
            return aligned_tags
        
        # Apply lineage to tags (only those that were missing it)
        aligned_count = 0
        for tag in aligned_tags:
            if not isinstance(tag, dict):
                continue
            name = tag.get('Product Name*')
            if not name:
                continue
            # Only align if missing canonical_lineage/currentLineage
            if not (tag.get('canonical_lineage') or tag.get('currentLineage')):
                db_lineage = lineage_map.get(str(name).lower().strip())
                if db_lineage:
                    tag['Lineage'] = db_lineage
                    tag['lineage'] = db_lineage.lower()
                    tag['canonical_lineage'] = db_lineage
                    tag['currentLineage'] = db_lineage
                    aligned_count += 1
        
        if aligned_count > 0:
            logging.debug(f"✅ Aligned {aligned_count} tags with database lineage")
        
        return aligned_tags
    except Exception as e:
        logging.warning(f"Lineage alignment failed: {e}")
        return tags

def _calculate_joint_ratio_for_record(db_record):
    """Calculate joint ratio for pre-roll products from database record."""
    product_name = db_record.get('Product Name*', '')
    product_type = db_record.get('Product Type*', '')
    weight = _normalize_weight_string(db_record.get('Weight*', ''))
    
    # Only calculate for pre-roll products
    if not product_type or 'pre-roll' not in str(product_type).lower():
        return db_record.get('JointRatio', '')
    
    if not product_name:
        return db_record.get('JointRatio', '')
    
    product_name_str = str(product_name)
    
    # Look for patterns like "0.5g x 2 Pack", "1g x 28 Pack", etc.
    patterns = [
        r'(\d+(?:\.\d+)?)g\s*x\s*(\d+)\s*pack',  # "0.5g x 2 Pack"
        r'(\d+(?:\.\d+)?)g\s*x\s*(\d+)',         # "0.5g x 2"
        r'(\d+(?:\.\d+)?)g\s*×\s*(\d+)',         # "0.5g × 2" (different x character)
    ]
    
    for pattern in patterns:
        match = re.search(pattern, product_name_str, re.IGNORECASE)
        if match:
            amount = match.group(1)
            count = match.group(2)
            try:
                count_int = int(count)
                if count_int == 1:
                    return _normalize_weight_string(f"{amount}g")
                else:
                    return _normalize_weight_string(f"{amount}g x {count} Pack")
            except ValueError:
                continue
    
    # Look for single pre-roll patterns like "Product Name - 1g", "Product Name - 0.5g"
    single_pre_roll_pattern = r'-\s*(\d+(?:\.\d+)?)g\s*$'
    match = re.search(single_pre_roll_pattern, product_name_str, re.IGNORECASE)
    if match:
        amount = match.group(1)
        return _normalize_weight_string(f"{amount}g")
    
    # If no pattern found, try to generate from weight
    if weight and str(weight).strip() != '' and str(weight).lower() != 'nan':
        try:
            weight_float = float(weight)
            if weight_float == 1.0:
                return "1g"
            else:
                # Format weight similar to price formatting - no decimals unless original has decimals
                if weight_float.is_integer():
                    formatted_weight = f"{int(weight_float)}g"
                else:
                    # Round to 2 decimal places and remove trailing zeros
                    formatted_weight = f"{weight_float:.2f}".rstrip("0").rstrip(".") + "g"
                return _normalize_weight_string(formatted_weight)
        except (ValueError, TypeError):
            pass
    
    return db_record.get('JointRatio', '')

def _replace_json_tags_with_database_data(selected_tags, product_db):
    """
    Replace JSON matched tags with their corresponding database data.
    
    Args:
        selected_tags: List of selected tag names
        product_db: ProductDatabase instance
        
    Returns:
        List of enhanced tag names with database data
    """
    try:
        if not selected_tags or not product_db:
            return selected_tags
        
        logging.info(f"🔄 Replacing JSON tags with database data for {len(selected_tags)} tags")
        
        enhanced_tags = []
        replaced_count = 0
        
        for tag_name in selected_tags:
            # Try to find this tag in the database
            db_products = product_db.get_products_by_names([tag_name])
            
            if db_products and len(db_products) > 0:
                # Found in database - use the database version
                db_product = db_products[0]
                db_name = db_product.get('Product Name*', '') or db_product.get('ProductName', '')
                
                if db_name and db_name != tag_name:
                    logging.info(f"🔄 Replaced JSON tag '{tag_name}' with database tag '{db_name}'")
                    enhanced_tags.append(db_name)
                    replaced_count += 1
                else:
                    # Same name, but use database data
                    enhanced_tags.append(tag_name)
                    logging.info(f"✅ Using database data for '{tag_name}'")
            else:
                # Not found in database, keep original
                enhanced_tags.append(tag_name)
                logging.info(f"⚠️  Tag '{tag_name}' not found in database, keeping original")
        
        logging.info(f"✅ Enhanced {replaced_count}/{len(selected_tags)} tags with database data")
        return enhanced_tags
        
    except Exception as e:
        logging.error(f"Error replacing JSON tags with database data: {e}")
        return selected_tags  # Return original tags if enhancement fails

@app.route('/api/generation-progress', methods=['GET'])
def get_generation_progress():
    """Get current generation progress"""
    try:
        from src.core.generation.fast_generation import get_generation_stats
        stats = get_generation_stats()
        return jsonify({
            'success': True,
            'stats': stats
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/clear-generation-cache', methods=['POST'])
def clear_generation_cache():
    """Clear generation cache"""
    try:
        from src.core.generation.fast_generation import clear_all_caches
        clear_all_caches()
        return jsonify({
            'success': True,
            'message': 'Generation cache cleared'
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/generate', methods=['POST'])
@performance_monitor if PERFORMANCE_ENABLED else lambda x: x
def generate_labels():
    try:
        import time
        _start_time = time.time()
        logging.info("=== GENERATE LABELS ACTION START ===")
        logging.info(f"Generate labels request at {datetime.now().strftime('%H:%M:%S')}")
        logging.info(f"Request method: {request.method}")
        logging.info(f"Request URL: {request.url}")
        logging.info(f"Request headers: {dict(request.headers)}")
        
        # TRACE: Check current store at start of generation
        current_store_at_start = get_current_store_name()
        logging.info(f"🔍 TRACE START: Current store = {current_store_at_start}")
        
        # Rate limiting for label generation
        client_ip = request.remote_addr
        if not check_rate_limit(client_ip):
            logging.warning(f"Rate limit exceeded for IP: {client_ip}")
            return jsonify({'error': 'Rate limit exceeded. Please wait before generating more labels.'}), 429
        
        # Add request deduplication using request fingerprint
        import hashlib
        request_data = request.get_json() or {}
        request_fingerprint = hashlib.md5(
            json.dumps(request_data, sort_keys=True).encode()
        ).hexdigest()
        
        # Initialize processing requests set if it doesn't exist
        if not hasattr(generate_labels, '_processing_requests'):
            generate_labels._processing_requests = set()
        
        # Check if this exact request is already being processed
        if request_fingerprint in generate_labels._processing_requests:
            logging.warning(f"Duplicate generation request detected for fingerprint: {request_fingerprint}")
            return jsonify({'error': 'This generation request is already being processed. Please wait.'}), 429
        
        # Mark this request as being processed
        generate_labels._processing_requests.add(request_fingerprint)
        
        data = request.get_json()
        template_type = data.get('template_type', 'vertical')
        scale_factor = float(data.get('scale_factor', 1.0))
        selected_tags_from_request = data.get('selected_tags', [])
        file_path = data.get('file_path')
        filters = data.get('filters', None)

        logging.info(f"🎯 Generation request received:")
        logging.info(f"   - template_type: {template_type}")
        logging.info(f"   - scale_factor: {scale_factor}")
        logging.info(f"   - selected_tags_from_request count: {len(selected_tags_from_request) if selected_tags_from_request else 0}")
        if selected_tags_from_request:
            logging.info(f"   - Sample tags: {selected_tags_from_request[:3]}")
        logging.debug(f"Selected tags from request: {selected_tags_from_request}")
        
        # TRACE: Check store before getting excel_processor
        logging.info(f"🔍 TRACE: Store before get_excel_processor = {get_current_store_name()}")
        
        # Enable product DB integration for proper tag matching
        excel_processor = get_excel_processor()
        
        # TRACE: Check store after getting excel_processor
        logging.info(f"🔍 TRACE: Store after get_excel_processor = {get_current_store_name()}")
        
        excel_processor.enable_product_db_integration(True)

        # CRITICAL FIX: JSON tags work exactly like Excel tags - no special preservation needed
        # They're already in the DataFrame and will be handled the same way as Excel tags
        
        # TRACE: Check store before file loading
        logging.info(f"🔍 TRACE: Store before file loading = {get_current_store_name()}")
        
        # Only load file if not already loaded
        if file_path:
            logging.info(f"🔍 TRACE: Loading specific file_path = {file_path}")
            if excel_processor._last_loaded_file != file_path or excel_processor.df is None or excel_processor.df.empty:
                excel_processor.load_file(file_path)
                logging.info(f"🔍 TRACE: Store after loading file_path = {get_current_store_name()}")
        else:
            # Ensure data is loaded - try to reload default file if needed
            if excel_processor.df is None:
                from src.core.data.excel_processor import get_default_upload_file
                selected_store = get_current_store_name() if has_store_selection() else None
                logging.info(f"🔍 TRACE: Loading default file for store: {selected_store}")
                default_file = get_default_upload_file(selected_store)
                logging.info(f"🔍 TRACE: get_default_upload_file returned: {default_file}")
                logging.info(f"🔍 TRACE: Store after get_default_upload_file = {get_current_store_name()}")
                if default_file:
                    logging.info(f"📂 TRACE: About to load default file: {default_file}")
                    excel_processor.load_file(default_file)
                    logging.info(f"🔍 TRACE: Store after loading default file = {get_current_store_name()}")
                else:
                    logging.warning(f"⚠️ GENERATE: No default file found for store: {selected_store}")
        
        # CRITICAL FIX: JSON tags work exactly like Excel tags - no special restoration needed
        # They're processed through the same pipeline as Excel tags

        # Check if we have data in Excel processor OR database
        has_excel_data = excel_processor.df is not None and not excel_processor.df.empty
        has_database = False
        
        # If no Excel data, try to load the default inventory file
        if not has_excel_data:
            try:
                # Try to load a store-specific default file if available
                store_name = get_current_store_name()
                if store_name:
                    # Try store-specific file first
                    store_display = store_name.replace('_', ' ').replace('AGT ', '')
                    default_file = f"uploads/A Greener Today - {store_display}_inventory*.xlsx"
                    import glob
                    matching_files = glob.glob(default_file)
                    if matching_files:
                        # Use most recent file
                        default_file = max(matching_files, key=os.path.getmtime)
                        logging.info(f"Loading store-specific default Excel file: {default_file}")
                        excel_processor.load_file(default_file)
                    else:
                        # Fallback to get_default_upload_file which searches for store-specific file
                        from src.core.data.excel_processor import get_default_upload_file
                        selected_store = get_current_store_name() if has_store_selection() else None
                        default_file = get_default_upload_file(selected_store)
                        if default_file and os.path.exists(default_file):
                            logging.info(f"Loading default Excel file for {selected_store}: {default_file}")
                            excel_processor.load_file(default_file)
                else:
                    # No store selected, skip default file loading (requires store selection)
                    from src.core.data.excel_processor import get_default_upload_file
                    selected_store = get_current_store_name() if has_store_selection() else None
                    default_file = get_default_upload_file(selected_store)
                    if default_file and os.path.exists(default_file):
                        logging.info(f"Loading default Excel file for {selected_store}: {default_file}")
                        excel_processor.load_file(default_file)
                has_excel_data = excel_processor.df is not None and not excel_processor.df.empty
                if has_excel_data:
                    logging.info(f"Successfully loaded default Excel file with {len(excel_processor.df)} records")
                else:
                    logging.warning("Default Excel file loaded but DataFrame is empty")
            except Exception as e:
                logging.warning(f"Could not load default Excel file: {e}")
        
        # Check if database is available
        try:
            # Store context removed - using single database
            store_name = get_current_store_name()
            product_db = get_product_database(store_name)
            if product_db:
                # Test if database has data
                conn = product_db._get_connection()
                cursor = conn.cursor()
                cursor.execute("SELECT COUNT(*) FROM products")
                count = cursor.fetchone()[0]
                has_database = count > 0
                logging.info(f"Database has {count} products")
        except Exception as e:
            logging.warning(f"Could not check database: {e}")
        if not has_excel_data and not has_database:
            logging.error("No data loaded in Excel processor or database")
            return jsonify({'error': 'No data loaded. Please upload an Excel file or ensure database is populated.'}), 400

        # Apply filters early
        filtered_df = excel_processor.apply_filters(filters) if filters else excel_processor.df

        # Use cached dropdowns for UI (if needed elsewhere)
        dropdowns = excel_processor.dropdown_cache

        # Use selected tags from request body or session, this updates the processor's internal state
        selected_tags_to_use = selected_tags_from_request
        
                # If no selected tags in request body, check session for JSON-matched tags
        if not selected_tags_to_use:
            # CRITICAL FIX: Check multiple session locations for selected tags
            session_selected_tags = session.get('selected_tags', [])
            json_selected_tags = session.get('json_selected_tags', [])
            last_json_match_count = session.get('last_json_match_count', 0)
            
            logging.info(f"CRITICAL FIX: Session selected_tags: {len(session_selected_tags)}")
            logging.info(f"CRITICAL FIX: Session json_selected_tags: {len(json_selected_tags)}")
            logging.info(f"CRITICAL FIX: Last JSON match count: {last_json_match_count}")
            
            # CRITICAL FIX: Check cache for selected tags as primary source
            selected_tags_cache_key = session.get('selected_tags_cache_key')
            if selected_tags_cache_key:
                cached_selected_tags = cache.get(selected_tags_cache_key)
                if cached_selected_tags:
                    logging.info(f"CRITICAL FIX: Using selected tags from cache: {len(cached_selected_tags)} tags")
                    selected_tags_to_use = cached_selected_tags
                    # Restore to session and Excel processor
                    session['selected_tags'] = cached_selected_tags
                    excel_processor.selected_tags = cached_selected_tags
            
            # CRITICAL FIX: Check for JSON matched tags in cache as fallback
            json_matched_cache_key = session.get('json_matched_cache_key')
            if json_matched_cache_key:
                json_matched_tags = cache.get(json_matched_cache_key)
                if json_matched_tags:
                    logging.info(f"CRITICAL FIX: Found JSON matched tags in cache: {len(json_matched_tags)} tags")
                    # Extract product names from JSON matched tags
                    product_names = []
                    for tag in json_matched_tags:
                        if isinstance(tag, dict):
                            product_name = tag.get('Product Name*', tag.get('ProductName', ''))
                            if product_name:
                                product_names.append(product_name)
                    
                    if product_names:
                        logging.info(f"CRITICAL FIX: Using {len(product_names)} product names from JSON matched tags")
                        selected_tags_to_use = product_names
                        # Restore to session and Excel processor
                        session['selected_tags'] = product_names
                        excel_processor.selected_tags = product_names
                        logging.info(f"CRITICAL FIX: Set selected_tags_to_use to {len(product_names)} tags")
            
            if session_selected_tags:
                logging.info(f"Using selected tags from session: {len(session_selected_tags)} tags")
                selected_tags_to_use = session_selected_tags
            elif json_selected_tags:
                logging.info(f"Using selected tags from json_selected_tags: {len(json_selected_tags)} tags")
                selected_tags_to_use = json_selected_tags
                # Restore to main session location
                session['selected_tags'] = json_selected_tags
                excel_processor.selected_tags = json_selected_tags
            else:
                # Also check excel_processor.selected_tags (set by JSON matching)
                if hasattr(excel_processor, 'selected_tags') and excel_processor.selected_tags:
                    logging.info(f"Using selected tags from excel_processor: {len(excel_processor.selected_tags)} tags")
                    selected_tags_to_use = excel_processor.selected_tags
                    # Restore to session
                    session['selected_tags'] = excel_processor.selected_tags
        
        if selected_tags_to_use:
            # Normalize selected tags - convert dictionary objects to product names
            normalized_tags = []
            for tag in selected_tags_to_use:
                if isinstance(tag, dict):
                    # Extract product name from dictionary
                    product_name = (tag.get('Product Name*') or 
                                  tag.get('displayName') or 
                                  tag.get('ProductName') or 
                                  str(tag))
                    if product_name and str(product_name).strip():
                        normalized_tags.append(str(product_name).strip())
                elif isinstance(tag, str):
                    # Already a string
                    normalized_tags.append(tag.strip())
                else:
                    # Convert to string
                    normalized_tags.append(str(tag).strip())
            
            logging.info(f"Normalized {len(selected_tags_to_use)} tags to {len(normalized_tags)} product names")
            logging.debug(f"Sample normalized tags: {normalized_tags[:3]}")
            
            # CRITICAL FIX: Check if these are JSON matched tags first
            json_matched_cache_key = session.get('json_matched_cache_key')
            is_json_matched_session = json_matched_cache_key is not None
            
            logging.info(f"🔍 SESSION CHECK: json_matched_cache_key = {json_matched_cache_key}")
            logging.info(f"🔍 SESSION CHECK: is_json_matched_session = {is_json_matched_session}")
            if is_json_matched_session:
                logging.info(f"🔍 JSON SESSION DETECTED: This should trigger fuzzy matching")
            else:
                logging.info(f"🔍 REGULAR SESSION: Will use exact matching only")
            
            # Try to validate tags against database first, then fall back to Excel data
            valid_selected_tags = []
            invalid_selected_tags = []
            
            # Always try database validation, but use fuzzy matching for JSON sessions
            try:
                # Get current store selection or default to Bothell for backward compatibility
                store_name = get_current_store_name() or 'AGT_Bothell'
                product_db = get_product_database(store_name)
                if product_db:
                    logging.info("Attempting to validate selected tags against database...")
                    
                    # VALIDATION DEBUG: Track what happens to JSON matches
                    logging.debug(f"🔍 VALIDATION DEBUG: About to validate {len(normalized_tags)} normalized tags")
                    logging.debug(f"🔍 VALIDATION DEBUG: First 10 tags: {normalized_tags[:10]}")
                    

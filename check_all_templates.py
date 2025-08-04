#!/usr/bin/env python3
"""
Check all templates for DOH placeholders.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
import re

def check_template_for_doh(template_path):
    """Check a specific template for DOH placeholders."""
    print(f"🔍 Checking: {template_path}")
    
    if not os.path.exists(template_path):
        print(f"  ❌ File not found")
        return False
    
    # Load the template document
    doc = Document(template_path)
    
    # Find all placeholders in the document
    placeholders = set()
    doh_placeholders = set()
    
    # Check tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    # Find placeholders in this cell
                    placeholder_matches = re.findall(r'\{\{.*?\}\}', cell_text)
                    for match in placeholder_matches:
                        placeholders.add(match)
                        if 'DOH' in match:
                            doh_placeholders.add(match)
    
    # Check paragraphs outside tables
    for paragraph in doc.paragraphs:
        para_text = paragraph.text.strip()
        if para_text:
            # Find placeholders in this paragraph
            placeholder_matches = re.findall(r'\{\{.*?\}\}', para_text)
            for match in placeholder_matches:
                placeholders.add(match)
                if 'DOH' in match:
                    doh_placeholders.add(match)
    
    print(f"  📊 Total placeholders: {len(placeholders)}")
    print(f"  🔍 DOH placeholders: {len(doh_placeholders)}")
    
    if doh_placeholders:
        print(f"  ✅ DOH placeholders found:")
        for placeholder in sorted(doh_placeholders):
            print(f"    • {placeholder}")
        return True
    else:
        print(f"  ❌ No DOH placeholders found")
        return False

def main():
    """Check all templates for DOH placeholders."""
    print("🔍 Checking All Templates for DOH Placeholders")
    print("=" * 50)
    print()
    
    templates = [
        "src/core/generation/templates/double.docx",
        "src/core/generation/templates/horizontal.docx",
        "src/core/generation/templates/vertical.docx",
        "src/core/generation/templates/mini.docx"
    ]
    
    templates_with_doh = []
    
    for template_path in templates:
        has_doh = check_template_for_doh(template_path)
        if has_doh:
            templates_with_doh.append(template_path)
        print()
    
    print("🎯 SUMMARY")
    print("=" * 20)
    print(f"Templates with DOH placeholders: {len(templates_with_doh)}")
    
    if templates_with_doh:
        print("✅ Templates with DOH placeholders:")
        for template in templates_with_doh:
            print(f"  • {template}")
    else:
        print("❌ No templates have DOH placeholders!")
        print("   This is why DOH images are not being inserted.")

if __name__ == "__main__":
    main() 
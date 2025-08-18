#!/usr/bin/env python3
"""
Minimal test to generate exact vendor alignment format.
"""

import os
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt
from docx.shared import RGBColor

# Add the src directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

def test_minimal_vendor():
    """Test minimal vendor alignment without template processing."""
    print("🧪 Testing minimal vendor alignment...")
    
    try:
        # Create a new document
        doc = Document()
        
        # Add a table with 3 columns
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        
        # Set up the first row with vendor text
        lineages = ['SATIVA', 'INDICA', 'HYBRID']
        
        for col_idx in range(3):
            cell = table.cell(0, col_idx)
            
            # Create paragraph with lineage and vendor
            paragraph = cell.paragraphs[0]
            
            # Add lineage text
            lineage_run = paragraph.add_run(lineages[col_idx])
            lineage_run.font.name = "Arial"
            lineage_run.font.bold = True
            lineage_run.font.size = Pt(12)
            
            # Add two spaces (like the template)
            space_run = paragraph.add_run("  ")
            space_run.font.name = "Arial"
            space_run.font.size = Pt(12)
            
            # Add vendor text
            vendor_run = paragraph.add_run("1555 Industrial LLC")
            vendor_run.font.name = "Arial"
            vendor_run.font.bold = False
            vendor_run.font.italic = True
            vendor_run.font.size = Pt(10)
            vendor_run.font.color.rgb = RGBColor(204, 204, 204)  # Light gray
            
            # Set paragraph to LEFT alignment
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add right-aligned tab stop
            paragraph.paragraph_format.tab_stops.clear_all()
            tab_position = Inches(3.2)
            paragraph.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)
            
            # Insert tab character before vendor text
            vendor_run.text = '\t' + vendor_run.text
            
            print(f"   Cell {col_idx}: '{paragraph.text}'")
            print(f"   Alignment: {paragraph.alignment}")
            print(f"   Tab stops: {len(paragraph.paragraph_format.tab_stops)}")
        
        # Save the document
        doc.save("test_minimal_vendor_result.docx")
        print("✓ Document saved as test_minimal_vendor_result.docx")
        
        # Check the result
        print(f"\n📝 Final result:")
        for col_idx in range(3):
            cell = table.cell(0, col_idx)
            paragraph = cell.paragraphs[0]
            print(f"   Cell {col_idx}: '{paragraph.text}'")
            print(f"   Alignment: {paragraph.alignment}")
            print(f"   Tab stops: {len(paragraph.paragraph_format.tab_stops)}")
            
            # Check runs
            print(f"   Runs:")
            for i, run in enumerate(paragraph.runs):
                print(f"     Run {i}: '{run.text}' (length: {len(run.text)})")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_minimal_vendor()

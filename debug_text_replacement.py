#!/usr/bin/env python3
"""
Simple test to debug text replacement in the document.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor, get_font_scheme

def debug_text_replacement():
    """Debug text replacement in the document."""
    
    print("Debug Text Replacement")
    print("=" * 50)
    
    try:
        font_scheme = get_font_scheme('mini')
        processor = TemplateProcessor('mini', font_scheme, 1.0)
        
        # Load the expanded template
        processor.force_re_expand_template()
        processor._expanded_template_buffer.seek(0)
        from docx import Document
        doc = Document(processor._expanded_template_buffer)
        
        # Check first cell
        first_cell = doc.tables[0].cell(0, 0)
        print(f"First cell text: {repr(first_cell.text)}")
        
        # Check paragraph structure
        print(f"\nParagraphs in first cell: {len(first_cell.paragraphs)}")
        for i, para in enumerate(first_cell.paragraphs):
            print(f"Paragraph {i}: {repr(para.text)}")
            print(f"  Runs in paragraph {i}: {len(para.runs)}")
            for j, run in enumerate(para.runs):
                print(f"    Run {j}: {repr(run.text)}")
        
        # Try simple text replacement
        print(f"\n=== Testing Simple Text Replacement ===")
        
        # Replace DOH placeholder manually
        original_text = first_cell.text
        print(f"Original text: {repr(original_text)}")
        
        # Simple replacement
        new_text = original_text.replace("{{Label1.DOH }}", "100mg THC")
        print(f"After replacement: {repr(new_text)}")
        
        # Check if replacement worked
        if "{{Label1.DOH" in new_text:
            print("❌ Replacement failed")
        else:
            print("✅ Replacement successful")
            
        # Try to apply the replacement to the cell
        print(f"\n=== Applying Replacement to Cell ===")
        
        # Clear the cell and add new text
        first_cell._tc.clear_content()
        new_para = first_cell.add_paragraph()
        new_para.add_run(new_text)
        
        print(f"Cell text after modification: {repr(first_cell.text)}")
        
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_text_replacement()

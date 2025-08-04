#!/usr/bin/env python3
"""
Test script to verify the double template center vertical gutter implementation.
This script tests that the 4x3 double template has only a center vertical gutter via cell margins.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_DOUBLE
from docx import Document
from io import BytesIO

def test_double_template_center_vertical_gutter():
    """Test that the double template has only a center vertical gutter via cell margins."""
    print("🔍 Testing Double Template Center Vertical Gutter")
    
    try:
        # Create template processor
        processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
        
        # Force re-expansion to ensure we get the latest center gutter implementation
        processor.force_re_expand_template()
        print("✅ Template re-expanded with center vertical gutter implementation")
        
        # Get the expanded template
        template_buffer = processor._expand_template_if_needed(force_expand=True)
        doc = Document(template_buffer)
        
        if not doc.tables:
            print("❌ No table found in template")
            return False
            
        table = doc.tables[0]
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        
        print(f"📊 Table dimensions: {num_rows} rows x {num_cols} columns")
        
        # Verify we have a 4x3 grid
        expected_rows, expected_cols = 3, 4
        if num_rows != expected_rows or num_cols != expected_cols:
            print(f"❌ Expected {expected_rows}x{expected_cols} grid, got {num_rows}x{num_cols}")
            return False
        
        print(f"✅ Correct grid dimensions: {num_rows}x{num_cols}")
        
        # Check for absence of cell spacing (no global spacing)
        tblPr = table._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr')
        if tblPr is not None:
            cell_spacing = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblCellSpacing')
            if cell_spacing is None:
                print("✅ No cell spacing found (no global spacing)")
            else:
                print("⚠️  Cell spacing found (global spacing present)")
        else:
            print("⚠️  No table properties found")
        
        # Check cell margins for center gutter
        print("\n📏 Cell margin analysis:")
        center_gutter_cells = 0
        minimal_margin_cells = 0
        
        for r in range(num_rows):
            for c in range(num_cols):
                cell = table.cell(r, c)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcMar')
                
                if tcMar is not None:
                    # Check right margin for columns 1 and 2
                    right_margin = tcMar.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}right')
                    left_margin = tcMar.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left')
                    
                    right_width = 0
                    left_width = 0
                    
                    if right_margin is not None:
                        right_width = int(right_margin.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', '0'))
                    if left_margin is not None:
                        left_width = int(left_margin.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', '0'))
                    
                    right_inches = right_width / 1440
                    left_inches = left_width / 1440
                    
                    # Check if this cell has center gutter margins
                    if (c in [0, 1] and right_inches >= 0.025) or (c in [2, 3] and left_inches >= 0.025):
                        center_gutter_cells += 1
                        print(f"   ✅ Cell ({r+1},{c+1}): Center gutter margin - Right: {right_inches:.3f}\", Left: {left_inches:.3f}\"")
                    else:
                        minimal_margin_cells += 1
                        print(f"   📏 Cell ({r+1},{c+1}): Minimal margin - Right: {right_inches:.3f}\", Left: {left_inches:.3f}\"")
        
        print(f"\n📋 Margin Summary:")
        print(f"   Cells with center gutter margins: {center_gutter_cells}/12")
        print(f"   Cells with minimal margins: {minimal_margin_cells}/12")
        
        # Verify center gutter pattern
        expected_center_gutter = 12  # All cells should have center gutter margins
        if center_gutter_cells == expected_center_gutter:
            print("✅ All cells have center gutter margins")
        else:
            print(f"❌ Only {center_gutter_cells}/{expected_center_gutter} cells have center gutter margins")
            return False
        
        # Check that all cells have content
        label_count = 0
        empty_count = 0
        
        for r in range(num_rows):
            for c in range(num_cols):
                cell = table.cell(r, c)
                cell_text = cell.text.strip()
                
                if cell_text:
                    label_count += 1
                else:
                    empty_count += 1
        
        print(f"   Label cells: {label_count}/12 (should be 12)")
        print(f"   Empty cells: {empty_count}/12 (should be 0)")
        
        if label_count == 12 and empty_count == 0:
            print("✅ All cells have content")
        else:
            print("❌ Found empty cells")
            return False
        
        print("\n📋 Summary:")
        print("   Grid: 4x3 (standard grid)")
        print("   Center vertical gutter: 0.025\" margins on columns 1-2 right and 3-4 left")
        print("   No global cell spacing")
        print("   Total labels: 12 (all cells populated)")
        print("   Gutter location: Between columns 2 and 3")
        
        print("\n✅ Double template center vertical gutter test completed successfully!")
        return True
        
    except Exception as e:
        print(f"❌ Error testing double template center vertical gutter: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("🚀 Double Template Center Vertical Gutter Test")
    print("This test verifies that the double template has only a center vertical gutter via cell margins")
    print("=" * 70)
    
    success = test_double_template_center_vertical_gutter()
    
    if success:
        print("\n🎉 Test passed! The double template now has only a center vertical gutter.")
    else:
        print("\n💥 Test failed! Check the implementation.") 
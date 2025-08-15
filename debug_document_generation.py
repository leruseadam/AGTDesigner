#!/usr/bin/env python3
"""
Debug script to trace document generation step by step.
"""

import sys
import os
from pathlib import Path

# Add the src directory to the path
sys.path.insert(0, str(Path(__file__).parent / 'src'))

from src.core.generation.template_processor import TemplateProcessor

def debug_document_generation():
    """Debug the document generation process step by step."""
    
    # Test record that should show the duplication
    test_record = {
        'ProductName': 'Grape Moonshot',
        'ProductBrand': 'CONSTELLATION CANNABIS',
        'Product Type*': 'edible (solid)',
        'Description': 'Grape Moonshot -1.7oz',
        'WeightUnits': '1.7oz',
        'Price': '$15',
        'THC': '100.0%',
        'CBD': '0%',
        'Lineage': 'MIXED'
    }
    
    print("=== DEBUGGING DOCUMENT GENERATION ===")
    print(f"Test record: {test_record}")
    
    # Test with horizontal template
    from src.core.generation.template_processor import get_font_scheme
    font_scheme = get_font_scheme('horizontal')
    processor = TemplateProcessor('horizontal', font_scheme)
    
    try:
        # Step 1: Build context
        print("\n=== STEP 1: Building Context ===")
        from docx import Document
        doc = Document()
        label_context = processor._build_label_context(test_record, doc)
        
        print("Context built:")
        for key, value in label_context.items():
            if 'CONSTELLATION' in str(value):
                print(f"🔍 {key}: '{value}'")
            else:
                print(f"   {key}: '{value}'")
        
        # Step 2: Create full context
        print("\n=== STEP 2: Creating Full Context ===")
        context = {'Label1': label_context}
        print(f"Full context keys: {list(context.keys())}")
        
        # Step 3: Load template
        print("\n=== STEP 3: Loading Template ===")
        template_path = 'src/core/generation/templates/horizontal.docx'
        print(f"Template path: {template_path}")
        
        # Step 4: Expand template
        print("\n=== STEP 4: Expanding Template ===")
        # The template should already be loaded, just check if it needs expansion
        processor._expand_template_if_needed(force_expand=True)
        
        # Check expanded template
        if hasattr(processor, '_expanded_template_buffer'):
            expanded_buffer = processor._expanded_template_buffer
            expanded_buffer.seek(0)
            expanded_doc = Document(expanded_buffer)
            
            if expanded_doc.tables:
                table = expanded_doc.tables[0]
                print(f"Expanded template has {len(table.rows)} rows and {len(table.columns)} columns")
                
                # Check first cell
                first_cell = table.rows[0].cells[0]
                first_cell_text = ' '.join([p.text for p in first_cell.paragraphs])
                print(f"First cell text: '{first_cell_text}'")
                
                if 'CONSTELLATION CANNABISCONSTELLATION CANNABIS' in first_cell_text:
                    print("❌ DUPLICATION ALREADY IN EXPANDED TEMPLATE!")
                elif 'CONSTELLATION CANNABIS' in first_cell_text:
                    print("✅ First cell contains 'CONSTELLATION CANNABIS' (correct)")
                else:
                    print("ℹ️ First cell does not contain 'CONSTELLATION CANNABIS'")
        else:
            print("ℹ️ No expanded template buffer found")
        
        # Step 5: Generate document
        print("\n=== STEP 5: Generating Document ===")
        records = [test_record]
        documents = processor.process_records(records)
        
        if documents:
            print("Document generated successfully")
            
            # Check the final document
            doc = documents[0] if hasattr(documents, '__len__') else documents
            
            if doc.tables:
                table = doc.tables[0]
                print(f"Final document has {len(table.rows)} rows and {len(table.columns)} columns")
                
                # Check all cells for duplication
                duplication_found = False
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = ' '.join([p.text for p in cell.paragraphs])
                        if 'CONSTELLATION CANNABISCONSTELLATION CANNABIS' in cell_text:
                            print(f"❌ DUPLICATION FOUND in cell ({row_idx}, {col_idx}): '{cell_text}'")
                            duplication_found = True
                        elif 'CONSTELLATION CANNABIS' in cell_text:
                            print(f"✅ Cell ({row_idx}, {col_idx}) has correct content: '{cell_text}'")
                
                if not duplication_found:
                    print("✅ SUCCESS: No duplication found in final document!")
                else:
                    print("❌ FAILURE: Duplication was created during document generation")
            else:
                print("❌ Final document has no tables")
        else:
            print("❌ No documents generated")
            
    except Exception as e:
        print(f"❌ Error during debugging: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_document_generation()

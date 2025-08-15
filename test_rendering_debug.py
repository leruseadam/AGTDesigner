#!/usr/bin/env python3
"""
Test script to debug what happens during the actual DocxTemplate rendering process.
"""

import sys
import os

# Add the src directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

def test_rendering_debug():
    """Test the actual DocxTemplate rendering process to see what happens to markers."""
    
    print("Testing DocxTemplate Rendering Debug")
    print("=" * 60)
    
    try:
        from src.core.generation.template_processor import TemplateProcessor
        from docxtpl import DocxTemplate
        from docx import Document
        from io import BytesIO
        
        # Create a simple font scheme
        font_scheme = {
            'default': {'family': 'Arial', 'size': 12, 'bold': False},
            'header': {'family': 'Arial', 'size': 14, 'bold': True},
            'title': {'family': 'Arial', 'size': 16, 'bold': True}
        }
        
        processor = TemplateProcessor('horizontal', font_scheme, 1.0)
        
        print(f"Template type: {processor.template_type}")
        
        # Test the expanded template buffer
        print("\nTesting expanded template buffer...")
        print("=" * 60)
        
        try:
            # Check if the expanded template buffer exists
            if hasattr(processor, '_expanded_template_buffer'):
                print("✅ Expanded template buffer exists")
                
                # Check if it's seekable
                if hasattr(processor._expanded_template_buffer, 'seek'):
                    print("✅ Expanded template buffer is seekable")
                    
                    # Reset to beginning
                    processor._expanded_template_buffer.seek(0)
                    
                    # Create DocxTemplate from expanded buffer
                    print("\nCreating DocxTemplate from expanded buffer...")
                    doc = DocxTemplate(processor._expanded_template_buffer)
                    print("✅ DocxTemplate created successfully")
                    
                    # Check the template content
                    template_content = doc.get_docx()
                    if template_content.tables:
                        first_table = template_content.tables[0]
                        print(f"✅ Template has {len(template_content.tables)} tables")
                        print(f"First table rows: {len(first_table.rows)}")
                        print(f"First table columns: {len(first_table.columns)}")
                        
                        # Check first cell content
                        first_cell = first_table.cell(0, 0)
                        cell_text = first_cell.text
                        print(f"First cell text: '{cell_text}'")
                        
                        # Check if markers are present
                        if 'LINEAGE_START' in cell_text:
                            print("✅ LINEAGE_START marker found in template")
                        else:
                            print("❌ LINEAGE_START marker NOT found in template")
                            
                        if 'PRODUCTVENDOR_START' in cell_text:
                            print("✅ PRODUCTVENDOR_START marker found in template")
                        else:
                            print("❌ PRODUCTVENDOR_START marker NOT found in template")
                    else:
                        print("❌ Template has no tables")
                        
                    # Now test actual rendering
                    print("\nTesting actual rendering...")
                    print("=" * 60)
                    
                    # Create a simple context
                    context = {
                        'Label1': {
                            'Lineage': 'LINEAGE_START INDICA LINEAGE_END',
                            'ProductVendor': 'PRODUCTVENDOR_START Test Vendor PRODUCTVENDOR_END',
                            'ProductStrain': 'Test Strain',
                            'Description': 'Test Description',
                            'Price': '$10',
                            'Ratio': '1:1',
                            'Total THC': '20%',
                            'THCA': '20%',
                            'THC test result': '20%',
                            'CBDA': '1%',
                            'CBD test result': '1%',
                            'DOH': 'YES'
                        }
                    }
                    
                    # Render the template
                    doc.render(context)
                    print("✅ Template rendering successful")
                    
                    # Save the rendered document
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    
                    # Load the rendered document
                    rendered_doc = Document(buffer)
                    print("✅ Rendered document loaded successfully")
                    
                    if rendered_doc.tables:
                        first_table = rendered_doc.tables[0]
                        print(f"Rendered table rows: {len(first_table.rows)}")
                        print(f"Rendered table columns: {len(first_table.columns)}")
                        
                        # Check first cell content after rendering
                        first_cell = first_table.cell(0, 0)
                        cell_text = first_cell.text
                        print(f"First cell text after rendering: '{cell_text}'")
                        
                        # Check if markers are preserved
                        if 'LINEAGE_START' in cell_text:
                            print("✅ LINEAGE_START marker preserved after rendering")
                        else:
                            print("❌ LINEAGE_START marker lost after rendering")
                            
                        if 'PRODUCTVENDOR_START' in cell_text:
                            print("✅ PRODUCTVENDOR_START marker preserved after rendering")
                        else:
                            print("❌ PRODUCTVENDOR_START marker lost after rendering")
                            
                        # Check if content is there
                        if 'INDICA' in cell_text:
                            print("✅ Lineage content preserved: INDICA")
                        else:
                            print("❌ Lineage content lost")
                            
                        if 'Test Vendor' in cell_text:
                            print("✅ Vendor content preserved: Test Vendor")
                        else:
                            print("❌ Vendor content lost")
                    else:
                        print("❌ Rendered document has no tables")
                        
                else:
                    print("❌ Expanded template buffer is not seekable")
            else:
                print("❌ Expanded template buffer does not exist")
                
        except Exception as e:
            print(f"❌ Error testing rendering: {e}")
            import traceback
            traceback.print_exc()
        
    except Exception as e:
        print(f"❌ Error testing rendering debug: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_rendering_debug()
    
    print("\n" + "=" * 60)
    print("Test Complete")
    print("=" * 60)

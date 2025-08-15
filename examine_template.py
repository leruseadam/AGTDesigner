#!/usr/bin/env python3
"""
Script to thoroughly examine the template structure and content.
"""

from docx import Document

def examine_template():
    doc = Document('src/core/generation/templates/double.docx')
    
    print("=== TEMPLATE STRUCTURE ANALYSIS ===\n")
    
    print(f"Number of tables: {len(doc.tables)}")
    if doc.tables:
        table = doc.tables[0]
        print(f"Number of rows in first table: {len(table.rows)}")
        print(f"Number of columns in first table: {len(table.rows[0].cells)}")
        
        print("\n=== FULL TEMPLATE CONTENT ===")
        for i, row in enumerate(table.rows):
            print(f"\nRow {i}:")
            for j, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    print(f"  Cell {j}: \"{cell_text}\"")
                else:
                    print(f"  Cell {j}: [EMPTY]")
        
        # Check for any hidden content or special characters
        print("\n=== DETAILED CELL ANALYSIS ===")
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell_text = cell.text
                if cell_text:
                    print(f"\nRow {i}, Cell {j}:")
                    print(f"  Raw text: '{cell_text}'")
                    print(f"  Length: {len(cell_text)}")
                    print(f"  Contains 'Mixed': {'Mixed' in cell_text}")
                    print(f"  Contains 'DescAndWeight': {'DescAndWeight' in cell_text}")
                    print(f"  Contains 'ProductStrain': {'ProductStrain' in cell_text}")
                    print(f"  Contains 'Grape': {'Grape' in cell_text}")
                    print(f"  Contains 'Moonshot': {'Moonshot' in cell_text}")
                    
                    # Check for any special characters
                    special_chars = []
                    for char in cell_text:
                        if ord(char) > 127 or char in ['\n', '\t', '\r']:
                            special_chars.append(f"'{char}' (U+{ord(char):04X})")
                    
                    if special_chars:
                        print(f"  Special characters: {special_chars}")
    else:
        print("No tables found in template")

if __name__ == "__main__":
    examine_template()

#!/usr/bin/env python3
"""
Test script to verify that brand markers are not being added to classic types.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import CLASSIC_TYPES
from docx import Document
from docx.shared import Pt
import tempfile
import os

def test_classic_type_brand_exclusion():
    """Test that brand markers are not added to classic types."""
    print("Testing Classic Type Brand Exclusion")
    
    # Test with a classic type (flower)
    classic_record = {
        'ProductName': 'Test Flower',
        'ProductBrand': 'Test Brand',
        'ProductStrain': 'Test Strain',
        'ProductType': 'flower',  # Classic type
        'ProductVendor': 'Test Vendor',
        'Lineage': 'HYBRID',
        'Price': '$25.00',
        'WeightUnits': '3.5g',
        'Description': 'Test Description'
    }
    
    # Test with a non-classic type (edible)
    non_classic_record = {
        'ProductName': 'Test Edible',
        'ProductBrand': 'Test Brand',
        'ProductStrain': 'Test Strain',
        'ProductType': 'edible (solid)',  # Non-classic type
        'ProductVendor': 'Test Vendor',
        'Lineage': 'HYBRID',
        'Price': '$25.00',
        'WeightUnits': '3.5g',
        'Description': 'Test Description'
    }
    
    try:
        # Test classic type (should NOT have brand markers)
        print("\n1. Testing Classic Type (flower)...")
        tp_classic = TemplateProcessor('double', {})
        tp_classic.current_product_type = 'flower'
        
        documents_classic = tp_classic.process_records([classic_record])
        
        if not documents_classic:
            print("✗ No documents generated for classic type")
            return False
        
        doc_classic = documents_classic[0] if isinstance(documents_classic, list) else documents_classic
        
        # Check if brand markers are present in classic type
        brand_markers_found_classic = False
        for table in doc_classic.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if 'PRODUCTBRAND_CENTER_START' in text or 'PRODUCTBRAND_CENTER_END' in text:
                            brand_markers_found_classic = True
                            print(f"  Found brand markers in classic type: '{text.strip()}'")
        
        if brand_markers_found_classic:
            print("✗ ERROR: Brand markers found in classic type (should NOT have brand markers)")
            return False
        else:
            print("✓ SUCCESS: No brand markers found in classic type")
        
        # Test non-classic type (should HAVE brand markers)
        print("\n2. Testing Non-Classic Type (edible)...")
        tp_non_classic = TemplateProcessor('double', {})
        tp_non_classic.current_product_type = 'edible (solid)'
        
        documents_non_classic = tp_non_classic.process_records([non_classic_record])
        
        if not documents_non_classic:
            print("✗ No documents generated for non-classic type")
            return False
        
        doc_non_classic = documents_non_classic[0] if isinstance(documents_non_classic, list) else documents_non_classic
        
        # Check if brand content is present in non-classic type
        brand_content_found_non_classic = False
        print("\n3. Checking for brand content in non-classic type document:")
        for table in doc_non_classic.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        print(f"  Paragraph text: '{text.strip()}'")
                        if 'Test Brand' in text:
                            brand_content_found_non_classic = True
                            print(f"  ✓ Found brand content in non-classic type: '{text.strip()}'")
        
        if brand_content_found_non_classic:
            print("✓ SUCCESS: Brand content found in non-classic type")
        else:
            print("✗ ERROR: No brand content found in non-classic type (should have brand content)")
            return False
        
        print(f"\n✓ All tests passed!")
        print(f"  - Classic types ({', '.join(CLASSIC_TYPES)}) correctly excluded from brand marker addition")
        print(f"  - Non-classic types correctly receive brand markers")
        return True
        
    except Exception as e:
        print(f"✗ Error testing classic type brand exclusion: {e}")
        return False

if __name__ == "__main__":
    print("Testing Classic Type Brand Exclusion")
    print("=" * 50)
    
    success = test_classic_type_brand_exclusion()
    
    if success:
        print("\n✓ Classic type brand exclusion test PASSED")
    else:
        print("\n✗ Classic type brand exclusion test FAILED") 
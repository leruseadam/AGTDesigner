#!/usr/bin/env python3
"""
Minimal test to isolate the DocxTemplate issue
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docxtpl import DocxTemplate
from docx import Document
from io import BytesIO

def test_minimal_template():
    print("🔍 Testing minimal template with DocxTemplate...")
    
    try:
        # Create a minimal document with just one cell and one placeholder
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        
        # Add just one placeholder
        cell = table.cell(0, 0)
        cell.text = "{{test}}"
        
        print(f"✅ Created minimal template with placeholder: '{cell.text}'")
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Try to render with DocxTemplate
        print("🔧 Rendering with DocxTemplate...")
        template = DocxTemplate(buffer)
        
        context = {'test': 'SUCCESS!'}
        
        template.render(context)
        print("✅ DocxTemplate rendering completed")
        
        # Check result
        result_buffer = BytesIO()
        template.save(result_buffer)
        result_buffer.seek(0)
        
        result_doc = Document(result_buffer)
        if result_doc.tables:
            table = result_doc.tables[0]
            cell_text = table.cell(0, 0).text
            print(f"📊 Result cell text: '{cell_text}'")
            
            if "{{" in cell_text:
                print("❌ Placeholder not replaced")
            else:
                print("✅ Placeholder replaced successfully")
        
        print("🎯 Minimal template test completed!")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_minimal_template()

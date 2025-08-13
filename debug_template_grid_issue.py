#!/usr/bin/env python3
"""
Diagnostic script to check why template is showing as 2x3 instead of 3x3.
This will help identify the template expansion issue.
"""

import sys
import os
import logging
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# Enable debug logging
logging.basicConfig(level=logging.DEBUG)

def debug_template_grid_issue():
    """Debug the template grid issue."""
    print("Debugging Template Grid Issue (2x3 vs 3x3)")
    print("=" * 50)
    
    try:
        # Test different template types
        template_types = ['vertical', 'horizontal', 'mini', 'double', 'inventory']
        
        for template_type in template_types:
            print(f"\n--- Testing Template Type: {template_type} ---")
            
            try:
                # Import and test template processor
                from src.core.generation.template_processor import TemplateProcessor, get_font_scheme
                
                # Get font scheme
                font_scheme = get_font_scheme(template_type)
                print(f"  Font scheme: {font_scheme}")
                
                # Create template processor
                processor = TemplateProcessor(template_type, font_scheme)
                print(f"  Template processor created successfully")
                print(f"  Template type: {processor.template_type}")
                print(f"  Chunk size: {processor.chunk_size}")
                print(f"  Template path: {processor._template_path}")
                
                # Check if template was expanded
                if hasattr(processor, '_expanded_template_buffer') and processor._expanded_template_buffer:
                    print(f"  Template buffer exists: Yes")
                    
                    # Check the expanded template structure
                    from docx import Document
                    from io import BytesIO
                    
                    # Reset buffer position
                    processor._expanded_template_buffer.seek(0)
                    
                    # Load document
                    doc = Document(processor._expanded_template_buffer)
                    print(f"  Document loaded successfully")
                    
                    # Check tables
                    if doc.tables:
                        table = doc.tables[0]
                        print(f"  Table found: Yes")
                        print(f"  Table rows: {len(table.rows)}")
                        print(f"  Table columns: {len(table.columns)}")
                        print(f"  Expected grid: {processor.chunk_size} labels")
                        
                        # Check if it's actually 3x3
                        if len(table.rows) == 3 and len(table.columns) == 3:
                            print(f"  ✓ Grid is 3x3 as expected")
                        else:
                            print(f"  ❌ Grid is {len(table.rows)}x{len(table.columns)}, expected 3x3")
                            
                        # Check for label placeholders
                        table_text = ""
                        for row in table.rows:
                            for cell in row.cells:
                                table_text += cell.text + " "
                        
                        # Count label placeholders
                        import re
                        label_matches = re.findall(r'Label(\d+)\.', table_text)
                        unique_labels = set(label_matches)
                        print(f"  Label placeholders found: {len(unique_labels)}")
                        print(f"  Label numbers: {sorted(unique_labels)}")
                        
                        if len(unique_labels) == 9:
                            print(f"  ✓ All 9 label placeholders present")
                        else:
                            print(f"  ❌ Missing label placeholders. Expected 9, found {len(unique_labels)}")
                            
                    else:
                        print(f"  ❌ No tables found in document")
                        
                else:
                    print(f"  ❌ Template buffer is None or empty")
                    
            except Exception as e:
                print(f"  ❌ Error testing {template_type}: {e}")
                import traceback
                traceback.print_exc()
                
    except Exception as e:
        print(f"❌ Error in main debug: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_template_grid_issue()

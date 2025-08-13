#!/usr/bin/env python3
"""
Debug script to test manual placeholder replacement for mini templates.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from io import BytesIO

def debug_mini_placeholder_replacement():
    """Debug the mini template placeholder replacement issue."""
    
    print("Debugging Mini Template Placeholder Replacement")
    print("=" * 60)
    
    # Test data
    test_record = {
        'ProductBrand': 'Test Brand',
        'Price': '$25.99',
        'Lineage': 'MIXED',
        'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
        'Description': 'Test description text',
        'ProductStrain': 'Mixed',
        'ProductType': 'tincture'
    }
    
    try:
        # Create mini template processor
        processor = TemplateProcessor('mini', {}, 1.0)
        
        # Check the expanded template buffer
        print(f"Expanded template buffer type: {type(processor._expanded_template_buffer)}")
        print(f"Expanded template buffer size: {len(processor._expanded_template_buffer.getvalue()) if hasattr(processor._expanded_template_buffer, 'getvalue') else 'N/A'}")
        
        # Load the expanded template
        doc = Document(processor._expanded_template_buffer)
        print(f"Document has {len(doc.tables)} tables")
        
        if doc.tables:
            table = doc.tables[0]
            print(f"Table dimensions: {len(table.rows)}x{len(table.columns)}")
            
            # Check the first cell for placeholders
            first_cell = table.cell(0, 0)
            print(f"\nFirst cell content:")
            print(f"  Cell text: '{first_cell.text}'")
            print(f"  Paragraphs: {len(first_cell.paragraphs)}")
            
            for i, para in enumerate(first_cell.paragraphs):
                print(f"  Paragraph {i}: '{para.text}'")
                for j, run in enumerate(para.runs):
                    print(f"    Run {j}: '{run.text}'")
        
        # Test context building
        print(f"\nTesting context building...")
        context = {}
        for i, record in enumerate([test_record]):
            label_context = processor._build_label_context(record, None)
            context[f'Label{i+1}'] = label_context
            print(f"Label{i+1} context: {label_context}")
        
        # Add default contexts for remaining labels
        for j in range(1, 20):
            default_context = {
                'ProductBrand': '',
                'ProductStrain': '',
                'ProductVendor': '',
                'Price': '',
                'THC_CBD': '',
                'Lineage': '',
                'DescAndWeight': '',
                'DOH': '',
                'DOH_TEXT': '',
                'Ratio': '',
                'WeightUnits': '',
                'Description': ''
            }
            context[f'Label{j+1}'] = default_context
        
        print(f"\nTotal context keys: {list(context.keys())}")
        print(f"Sample context values:")
        for key in list(context.keys())[:3]:
            print(f"  {key}: {context[key]}")
        
        # Test manual placeholder replacement
        print(f"\nTesting manual placeholder replacement...")
        doc_copy = Document(processor._expanded_template_buffer)
        doc_replaced = processor._manual_replace_placeholders(doc_copy, context)
        
        # Check if replacement worked
        if doc_replaced.tables:
            table = doc_replaced.tables[0]
            first_cell = table.cell(0, 0)
            print(f"\nAfter replacement - First cell content:")
            print(f"  Cell text: '{first_cell.text}'")
            
            # Check if placeholders were replaced
            if '{{Label1.' in first_cell.text:
                print(f"  ❌ Placeholders still present!")
            else:
                print(f"  ✅ Placeholders replaced successfully!")
                
            # Check specific fields
            if 'Test Brand' in first_cell.text:
                print(f"  ✅ ProductBrand replaced")
            else:
                print(f"  ❌ ProductBrand not replaced")
                
            if '$25.99' in first_cell.text:
                print(f"  ✅ Price replaced")
            else:
                print(f"  ❌ Price not replaced")
        
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_mini_placeholder_replacement()

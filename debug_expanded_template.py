#!/usr/bin/env python3
"""
Debug script to check what the expanded mini template contains
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from io import BytesIO

def check_expanded_template():
    print("🔍 Checking expanded mini template...")
    
    try:
        # Initialize template processor
        processor = TemplateProcessor('mini', 'mini')
        
        print(f"✅ Template processor initialized")
        print(f"📁 Template path: {processor._template_path}")
        
        # Check if expanded template buffer exists
        if hasattr(processor, '_expanded_template_buffer') and processor._expanded_template_buffer:
            print("✅ Expanded template buffer exists")
            
            # Load the expanded template
            if hasattr(processor._expanded_template_buffer, 'seek'):
                processor._expanded_template_buffer.seek(0)
            
            doc = Document(processor._expanded_template_buffer)
            print(f"📊 Expanded document has {len(doc.tables)} tables")
            
            if doc.tables:
                table = doc.tables[0]
                print(f"📊 Table: {len(table.rows)} rows x {len(table.columns)} columns")
                
                # Check first few cells for placeholders
                for i in range(min(3, len(table.rows))):
                    for j in range(min(3, len(table.columns))):
                        cell = table.cell(i, j)
                        cell_text = cell.text
                        print(f"  Cell [{i}][{j}]: {cell_text[:100]}...")
                        
                        # Check for placeholders
                        if "{{" in cell_text:
                            print(f"    ✅ Contains placeholders")
                            # Find all placeholders
                            import re
                            placeholders = re.findall(r'\{\{[^}]+\}\}', cell_text)
                            print(f"    📝 Placeholders found: {placeholders}")
                        else:
                            print(f"    ❌ No placeholders found")
        else:
            print("❌ No expanded template buffer found")
            
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    check_expanded_template()

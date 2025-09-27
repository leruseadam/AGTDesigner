#!/usr/bin/env python3
"""
Ultra-fast document generation for PythonAnywhere
"""

from docx import Document
from docx.shared import Inches, Pt
import os
import time

class FastDocxGenerator:
    """Optimized DOCX generation with minimal overhead"""
    
    def __init__(self):
        self.template_cache = {}
    
    def generate_simple_labels(self, data, output_path):
        """Generate labels with minimal formatting for speed"""
        start_time = time.time()
        
        try:
            # Create simple document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Product Labels', 0)
            
            # Process data in chunks
            chunk_size = 10  # Small chunks for PythonAnywhere
            processed = 0
            
            for i in range(0, min(len(data), 50), chunk_size):  # Limit to 50 items
                chunk = data[i:i+chunk_size]
                
                for item in chunk:
                    # Add simple paragraph for each product
                    p = doc.add_paragraph()
                    p.add_run(f"Product: {item.get('Product Name*', 'N/A')}").bold = True
                    
                    # Add basic info only
                    doc.add_paragraph(f"Type: {item.get('Product Type*', 'N/A')}")
                    doc.add_paragraph(f"Brand: {item.get('Product Brand', 'N/A')}")
                    doc.add_paragraph(f"Weight: {item.get('Weight*', 'N/A')}")
                    
                    # Add separator
                    doc.add_paragraph("─" * 40)
                    
                    processed += 1
                
                # Quick break to prevent timeout
                if time.time() - start_time > 15:  # 15 second limit
                    break
            
            # Save document
            doc.save(output_path)
            
            generation_time = time.time() - start_time
            
            return {
                'success': True,
                'output_path': output_path,
                'processed_items': processed,
                'generation_time': generation_time,
                'file_size': os.path.getsize(output_path) if os.path.exists(output_path) else 0
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'generation_time': time.time() - start_time
            }

def create_fast_generator_routes(app):
    """Add fast generation routes to Flask app"""
    
    from flask import request, jsonify, send_file, session
    
    @app.route('/generate-ultra-fast', methods=['POST'])
    def generate_ultra_fast():
        """Ultra-fast document generation"""
        start_time = time.time()
        
        try:
            # Get data from request
            data = request.get_json()
            if not data or 'items' not in data:
                return jsonify({'error': 'No data provided'}), 400
            
            items = data['items']
            if len(items) > 25:  # Limit for speed
                items = items[:25]
                
            # Generate output filename
            timestamp = int(time.time())
            output_filename = f"fast_labels_{timestamp}.docx"
            output_path = os.path.join('output', output_filename)
            
            # Ensure output directory exists
            os.makedirs('output', exist_ok=True)
            
            # Generate document
            generator = FastDocxGenerator()
            result = generator.generate_simple_labels(items, output_path)
            
            total_time = time.time() - start_time
            
            if result['success']:
                return jsonify({
                    'success': True,
                    'filename': output_filename,
                    'download_url': f'/download/{output_filename}',
                    'processed_items': result['processed_items'],
                    'total_time': round(total_time, 2),
                    'file_size': result['file_size']
                })
            else:
                return jsonify({
                    'error': result['error'],
                    'total_time': round(total_time, 2)
                }), 500
                
        except Exception as e:
            total_time = time.time() - start_time
            return jsonify({
                'error': f'Generation failed: {str(e)}',
                'total_time': round(total_time, 2)
            }), 500

# Export the function
__all__ = ['create_fast_generator_routes', 'FastDocxGenerator']

#!/usr/bin/env python3
"""
Debug script to test actual mini template processing with real data.
This will help identify why placeholders aren't being replaced during real label generation.
"""

import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), 'src'))

from core.generation.template_processor import TemplateProcessor
from docx import Document
import tempfile

def test_mini_template_with_real_data():
    """Test mini template processing with actual product data."""
    print("Testing mini template processing with real data...")
    
    try:
        # Create template processor for mini template
        processor = TemplateProcessor('mini', 'Arial')
        
        # Create realistic test data that matches what the mini template expects
        test_records = [
            {
                'ProductName': 'HUSTLER\'S AMBITION',
                'Product Type*': 'Flower',
                'Product Brand': 'HUSTLER\'S AMBITION',
                'Product Strain': 'HUSTLER\'S AMBITION',
                'Description': 'Premium Flower',
                'Weight*': '3.5g',
                'Price': '$56.00',
                'Ratio_or_THC_CBD': 'THC: 24.8% CBD: -0.1%',
                'Lineage': 'INDICA',
                'Vendor': '1555 Industrial LLC',
                'DOH': 'Yes'
            },
            {
                'ProductName': 'Test Product 2',
                'Product Type*': 'Flower',
                'Product Brand': 'Test Brand 2',
                'Product Strain': 'Test Strain 2',
                'Description': 'Test Description 2',
                'Weight*': '1g',
                'Price': '$4.00',
                'Ratio_or_THC_CBD': 'THC: 18.5% CBD: 0.2%',
                'Lineage': 'HYBRID',
                'Vendor': 'Test Vendor 2',
                'DOH': 'No'
            }
        ]
        
        print(f"Created {len(test_records)} test records")
        
        # Process the mini template with real data
        context = {}
        for i, record in enumerate(test_records):
            # Set current record for processing
            processor.current_record = record
            processor.current_product_type = record.get('Product Type*', '').lower()
            
            # Build label context
            label_context = processor._build_label_context(record, None)
            context[f'Label{i+1}'] = label_context
            
            print(f"\nLabel{i+1} context:")
            for key, value in label_context.items():
                if key in ['ProductStrain', 'ProductBrand', 'VendorInfo', 'Price', 'Ratio_or_THC_CBD', 'Lineage', 'DOH']:
                    print(f"  {key}: '{value}'")
        
        # Add empty contexts for remaining labels up to 20
        for i in range(len(test_records), 20):
            context[f'Label{i+1}'] = {}
        
        print(f"\nTotal context keys: {list(context.keys())}")
        
        # Process the mini template
        result_doc = processor._expand_mini_template_preserve_design(None, context)
        
        # Save result to temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            result_doc.save(tmp_file.name)
            tmp_path = tmp_file.name
        
        print(f"\nSaved result to: {tmp_path}")
        
        # Check the result structure
        result_doc = Document(tmp_path)
        if result_doc.tables:
            table = result_doc.tables[0]
            print(f"Result table: {len(table.rows)} rows x {len(table.rows[0].cells)} columns")
            
            # Check first few cells for content
            for i in range(min(5, len(table.rows) * len(table.rows[0].cells))):
                row_idx = i // len(table.rows[0].cells)
                col_idx = i % len(table.rows[0].cells)
                cell = table.rows[row_idx].cells[col_idx]
                print(f"\nCell {i+1}:")
                print(f"  Text: {cell.text[:200]}...")
                
                # Check if placeholders are still there
                if '{{Label' in cell.text:
                    print(f"  ⚠️  STILL HAS PLACEHOLDERS!")
                else:
                    print(f"  ✅ Placeholders replaced with data")
        else:
            print("ERROR: No tables found in result document")
        
        # Clean up
        os.unlink(tmp_path)
        
        print("\nMini template debug test completed!")
        return True
        
    except Exception as e:
        print(f"ERROR in mini template debug test: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = test_mini_template_with_real_data()
    sys.exit(0 if success else 1)

#!/usr/bin/env python3
"""
Test script to verify that the marker cleanup and 7.5pt font fixes work correctly.
"""

import sys
import os
from pathlib import Path

# Add the src directory to the path
sys.path.insert(0, str(Path(__file__).parent / 'src'))

from src.core.generation.unified_font_sizing import get_font_size_by_marker, get_mini_font_size_by_marker
from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docx.shared import Pt

def test_thc_cbd_font_sizing():
    """Test that THC/CBD content never gets 7.5pt font."""
    print("=== TESTING THC/CBD FONT SIZING FIX ===")
    
    test_text = "THC: 15.2% CBD: 2.1%"
    
    # Test all template types
    template_types = ['mini', 'vertical', 'horizontal', 'double']
    
    for template_type in template_types:
        print(f"\nTesting {template_type} template:")
        
        # Test THC_CBD marker
        font_size = get_font_size_by_marker(test_text, 'THC_CBD', template_type)
        print(f"  THC_CBD marker: {font_size.pt}pt")
        
        # Test RATIO_OR_THC_CBD marker
        font_size = get_font_size_by_marker(test_text, 'RATIO_OR_THC_CBD', template_type)
        print(f"  RATIO_OR_THC_CBD marker: {font_size.pt}pt")
        
        # Test mini template specific function
        if template_type == 'mini':
            font_size = get_mini_font_size_by_marker(test_text, 'THC_CBD')
            print(f"  Mini THC_CBD: {font_size.pt}pt")
        
        # Verify no 7.5pt font
        if font_size.pt == 7.5:
            print(f"    ❌ ERROR: Still getting 7.5pt font!")
        else:
            print(f"    ✅ Good: {font_size.pt}pt font (not 7.5pt)")

def test_marker_cleanup():
    """Test that markers are properly cleaned up."""
    print("\n\n=== TESTING MARKER CLEANUP ===")
    
    # Create a test document with markers
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    
    # Add paragraphs with various markers
    test_markers = [
        "PRODUCTBRAND_STARTTest BrandPRODUCTBRAND_END",
        "LINEAGE_STARTTest LineageLINEAGE_END", 
        "THC_CBD_STARTTHC: 15% CBD: 2%THC_CBD_END",
        "PRICE_START$25.99PRICE_END",
        "PRODUCTSTRAIN_STARTTest StrainPRODUCTSTRAIN_END"
    ]
    
    for marker_text in test_markers:
        paragraph = cell.add_paragraph(marker_text)
        print(f"  Added marker: {marker_text[:50]}...")
    
    # Create a template processor and run marker cleanup
    try:
        processor = TemplateProcessor('vertical', 'test')
        
        # Run the final marker cleanup
        processor._final_marker_cleanup(doc)
        
        # Check if markers remain
        remaining_markers = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if any(marker in paragraph.text.upper() for marker in ['_START', '_END', 'PRODUCTBRAND', 'PRODUCTSTRAIN', 'LINEAGE', 'PRODUCTVENDOR', 'PRICE', 'DESC', 'THC_CBD', 'RATIO', 'WEIGHTUNITS', 'DOH']):
                            remaining_markers.append(paragraph.text[:100])
        
        if remaining_markers:
            print(f"  ❌ ERROR: {len(remaining_markers)} paragraphs still contain markers!")
            for i, marker in enumerate(remaining_markers[:3]):
                print(f"    Remaining marker {i+1}: {marker}")
        else:
            print("  ✅ SUCCESS: All markers cleaned up!")
            
        # Show final text
        print("\n  Final text after cleanup:")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            print(f"    '{paragraph.text}'")
        
    except Exception as e:
        print(f"  ❌ ERROR: Template processor failed: {e}")

def test_final_verification():
    """Test the final verification method."""
    print("\n\n=== TESTING FINAL VERIFICATION ===")
    
    # Create a test document
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    
    # Add clean text (no markers)
    cell.add_paragraph("Clean text without markers")
    
    try:
        processor = TemplateProcessor('vertical', 'test')
        
        # Run verification
        result = processor._verify_no_markers_remain(doc)
        
        if result:
            print("  ✅ SUCCESS: Verification passed - no markers found")
        else:
            print("  ❌ ERROR: Verification failed - markers found")
            
    except Exception as e:
        print(f"  ❌ ERROR: Verification method failed: {e}")

if __name__ == "__main__":
    print("Testing Marker Cleanup and 7.5pt Font Fixes")
    print("=" * 50)
    
    # Test THC/CBD font sizing
    test_thc_cbd_font_sizing()
    
    # Test marker cleanup
    test_marker_cleanup()
    
    # Test final verification
    test_final_verification()
    
    print("\n" + "=" * 50)
    print("Testing complete!")

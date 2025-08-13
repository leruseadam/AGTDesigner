#!/usr/bin/env python3
"""
Test script to check mini template content
"""

from docx import Document
from pathlib import Path

def test_mini_template():
    """Test the mini template content."""
    
    print("Testing mini template content...")
    
    template_path = Path("src/core/generation/templates/mini.docx")
    print(f"Template path: {template_path}")
    print(f"Template exists: {template_path.exists()}")
    
    if template_path.exists():
        doc = Document(template_path)
        print(f"Document has {len(doc.paragraphs)} paragraphs")
        
        # Show all paragraph text
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                print(f"  Paragraph {i+1}: '{paragraph.text}'")
        
        # Check tables
        print(f"Document has {len(doc.tables)} tables")
        for i, table in enumerate(doc.tables):
            print(f"  Table {i+1}: {len(table.rows)} rows, {len(table.columns)} columns")
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    if cell.text.strip():
                        print(f"    Cell ({row_idx+1}, {col_idx+1}): '{cell.text}'")
    else:
        print("❌ Template file not found")

if __name__ == "__main__":
    test_mini_template()

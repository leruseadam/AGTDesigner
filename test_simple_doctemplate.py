#!/usr/bin/env python3
"""
Simple test to see if DocxTemplate works with basic placeholders
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docxtpl import DocxTemplate
from docx import Document
from io import BytesIO

def test_simple_doctemplate():
    print("🔍 Testing simple DocxTemplate functionality...")
    
    try:
        # Create a simple document with placeholders
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        
        # Add placeholders to cells
        table.cell(0, 0).text = "{{name}}"
        table.cell(0, 1).text = "{{age}}"
        table.cell(1, 0).text = "{{city}}"
        table.cell(1, 1).text = "{{job}}"
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        print("✅ Created simple template with placeholders")
        
        # Try to render with DocxTemplate
        print("🔧 Rendering with DocxTemplate...")
        template = DocxTemplate(buffer)
        
        context = {
            'name': 'John Doe',
            'age': '30',
            'city': 'New York',
            'job': 'Developer'
        }
        
        template.render(context)
        print("✅ DocxTemplate rendering completed")
        
        # Check result
        result_buffer = BytesIO()
        template.save(result_buffer)
        result_buffer.seek(0)
        
        result_doc = Document(result_buffer)
        if result_doc.tables:
            table = result_doc.tables[0]
            print(f"📊 Result table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            for i in range(len(table.rows)):
                for j in range(len(table.columns)):
                    cell_text = table.cell(i, j).text
                    print(f"  Cell [{i}][{j}]: '{cell_text}'")
                    
                    if "{{" in cell_text:
                        print(f"    ❌ Placeholder not replaced")
                    else:
                        print(f"    ✅ Placeholder replaced")
        
        print("🎯 Simple DocxTemplate test completed!")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_simple_doctemplate()

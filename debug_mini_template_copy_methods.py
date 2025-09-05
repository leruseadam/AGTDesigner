#!/usr/bin/env python3
"""
Debug script to test different methods of copying the mini template.
This will help find the right approach to preserve the template structure.
"""

import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), 'src'))

from core.generation.template_processor import TemplateProcessor
from docx import Document
from copy import deepcopy
import tempfile
import shutil

def debug_mini_template_copy_methods():
    """Debug different copy methods for mini template preservation."""
    print("=== Debugging Mini Template Copy Methods ===")
    
    try:
        # 1. Load original template
        print("\n1. Loading original mini template...")
        original_path = 'src/core/generation/templates/mini.docx'
        original_doc = Document(original_path)
        print(f"   ✅ Original template loaded: {len(original_doc.tables)} tables")
        
        if original_doc.tables:
            table = original_doc.tables[0]
            print(f"   📊 Original table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check first cell content
            first_cell = table.rows[0].cells[0]
            print(f"   📝 Original first cell: {first_cell.text[:100]}...")
        
        # 2. Method 1: Direct file copy (preserves everything)
        print("\n2. Method 1: Direct file copy...")
        temp_path = tempfile.mktemp(suffix='.docx')
        shutil.copy2(original_path, temp_path)
        
        copied_doc = Document(temp_path)
        print(f"   ✅ File copy completed: {len(copied_doc.tables)} tables")
        
        if copied_doc.tables:
            table = copied_doc.tables[0]
            print(f"   📊 Copied table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check first cell content
            first_cell = table.rows[0].cells[0]
            print(f"   📝 Copied first cell: {first_cell.text[:100]}...")
            
            # Compare with original
            if first_cell.text == original_doc.tables[0].rows[0].cells[0].text:
                print("   ✅ File copy preserved content exactly")
            else:
                print("   ❌ File copy changed content!")
        
        # 3. Method 2: Save and reload (preserves everything)
        print("\n3. Method 2: Save and reload...")
        temp_path2 = tempfile.mktemp(suffix='.docx')
        original_doc.save(temp_path2)
        
        reloaded_doc = Document(temp_path2)
        print(f"   ✅ Save/reload completed: {len(reloaded_doc.tables)} tables")
        
        if reloaded_doc.tables:
            table = reloaded_doc.tables[0]
            print(f"   📊 Reloaded table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check first cell content
            first_cell = table.rows[0].cells[0]
            print(f"   📝 Reloaded first cell: {first_cell.text[:100]}...")
            
            # Compare with original
            if first_cell.text == original_doc.tables[0].rows[0].cells[0].text:
                print("   ✅ Save/reload preserved content exactly")
            else:
                print("   ❌ Save/reload changed content!")
        
        # 4. Method 3: Deepcopy with new approach
        print("\n4. Method 3: Deepcopy with new approach...")
        new_doc = Document()
        
        # Copy the entire document structure
        new_doc._element = deepcopy(original_doc._element)
        
        print(f"   ✅ Deepcopy new approach completed: {len(new_doc.tables)} tables")
        
        if new_doc.tables:
            table = new_doc.tables[0]
            print(f"   📊 Deepcopy table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check first cell content
            first_cell = table.rows[0].cells[0]
            print(f"   📝 Deepcopy first cell: {first_cell.text[:100]}...")
            
            # Compare with original
            if first_cell.text == original_doc.tables[0].rows[0].cells[0].text:
                print("   ✅ Deepcopy new approach preserved content exactly")
            else:
                print("   ❌ Deepcopy new approach changed content!")
        
        # 5. Save all results for comparison
        print("\n5. Saving all results for comparison...")
        
        # Save original
        original_doc.save('./mini_original_fresh.docx')
        print("   💾 Original saved to: mini_original_fresh.docx")
        
        # Save file copy
        copied_doc.save('./mini_file_copy.docx')
        print("   💾 File copy saved to: mini_file_copy.docx")
        
        # Save reloaded
        reloaded_doc.save('./mini_reloaded.docx')
        print("   💾 Reloaded saved to: mini_reloaded.docx")
        
        # Save deepcopy new approach
        new_doc.save('./mini_deepcopy_new.docx')
        print("   💾 Deepcopy new approach saved to: mini_deepcopy_new.docx")
        
        # Clean up temp files
        os.unlink(temp_path)
        os.unlink(temp_path2)
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False
    
    print("\n=== Mini Template Copy Methods Debug Completed ===")
    return True

if __name__ == "__main__":
    debug_mini_template_copy_methods()

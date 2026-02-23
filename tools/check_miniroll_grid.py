from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
p = Path('src/core/generation/templates/miniroll.docx')
if not p.exists():
    print('miniroll.docx not found at', p)
    raise SystemExit(1)
doc = Document(p)
if not doc.tables:
    print('No tables in miniroll.docx')
    raise SystemExit(1)
t = doc.tables[0]
print('Table API rows,cols:', len(t.rows), len(t.columns))
# Print tblGrid if present
tblGrid = t._element.find(qn('w:tblGrid'))
if tblGrid is None:
    print('No tblGrid element')
else:
    cols = tblGrid.findall(qn('w:gridCol'))
    print('tblGrid has', len(cols), 'gridCols')
    for i,col in enumerate(cols,1):
        w = col.get(qn('w:w'))
        try:
            twips = int(w)
            inches = twips/1440
            print(f'Col {i}: {twips} twips = {inches:.3f} in')
        except Exception:
            print('Col',i,'w=',w)
# Inspect first row explicit height if set
first_row = t.rows[0]
trPr = first_row._tr.find(qn('w:trPr'))
if trPr is None:
    print('No trPr for first row')
else:
    trHeight = trPr.find(qn('w:trHeight'))
    if trHeight is None:
        print('No explicit trHeight for first row')
    else:
        val = trHeight.get(qn('w:val'))
        print('First row trHeight val:', val, '(twips?)')
print('Done')

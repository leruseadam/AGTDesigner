from docx import Document
from docx.oxml.ns import qn
from pathlib import Path

for path in ['src/core/generation/templates/miniroll.docx', '/tmp/miniroll_smoke.docx']:
    p = Path(path)
    print('\n===', path, '===')
    if not p.exists():
        print('MISSING')
        continue
    doc = Document(p)
    sec = doc.sections[0]
    pw = sec.page_width
    ph = sec.page_height
    left = sec.left_margin
    right = sec.right_margin
    top = sec.top_margin
    bottom = sec.bottom_margin
    print('Page width (twips):', pw, 'left,right margins:', left, right)
    print('Usable width (twips):', pw - left - right)
    print('Page height (twips):', ph, 'top,bottom margins:', top, bottom)
    print('Usable height (twips):', ph - top - bottom)
    print('Tables:', len(doc.tables))
    for ti, t in enumerate(doc.tables, start=1):
        print('\nTable', ti)
        tblGrid = t._element.find(qn('w:tblGrid'))
        if tblGrid is None:
            print('  No tblGrid')
        else:
            cols = tblGrid.findall(qn('w:gridCol'))
            print('  tblGrid cols:', len(cols))
            for i,c in enumerate(cols,1):
                w = c.get(qn('w:w'))
                try:
                    tw = int(w)
                    print(f'   Col{i}: {tw} twips = {tw/1440:.3f} in')
                except:
                    print('   Col',i,'w=',w)
        # rows/cols via API
        print('  API rows,cols:', len(t.rows), len(t.columns))
        # check first row height
        fr = t.rows[0]
        trPr = fr._tr.find(qn('w:trPr'))
        if trPr is None:
            print('  first row: no trPr')
        else:
            trh = trPr.find(qn('w:trHeight'))
            if trh is None:
                print('  first row: no trHeight')
            else:
                v = trh.get(qn('w:val'))
                print('  first row trHeight val:', v, '=>', (int(v)/20 if v and v.isdigit() else v), 'pts')
        # sample first cell text
        print('  first cell sample text:', repr(t.rows[0].cells[0].text[:200]))

print('\nDone')

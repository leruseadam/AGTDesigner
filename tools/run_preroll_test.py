import sys
import os
# Ensure project root is on sys.path so `src` can be imported
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

from src.core.generation.tag_generator import process_chunk, get_template_path, FONT_SCHEME_MINI
from docx import Document

# Sample infused preroll product
row = {
    "Product Name*": "Blueberry x Purple Churro Honey Stixx Infused Pre-Roll by Honey Tree - .5g",
    "ProductName": "Blueberry x Purple Churro Honey Stixx Infused Pre-Roll by Honey Tree - .5g",
    "Product Brand": "Honey Tree",
    "Product Type*": "infused pre-roll",
    "Description": "Blueberry x Purple Churro Honey Stixx Infused Pre-Roll",
    "Weight*": ".5g",
    "Units": "g",
    "JointRatio": "0.5g",
}

chunk = [row]
template_path = get_template_path('preroll')
args = (chunk, template_path, FONT_SCHEME_MINI, 'preroll', 1.0)

print(f"Using template: {template_path}")

try:
    out_bytes = process_chunk(args)
    out_path = '/tmp/preroll_test_output.docx'
    with open(out_path, 'wb') as f:
        f.write(out_bytes)
    print(f"Wrote output to: {out_path}")

    doc = Document(out_path)
    if doc.tables:
        first_cell = doc.tables[0].rows[0].cells[0].text
        print('First table cell (preview):')
        print(first_cell[:800])
    else:
        print('No tables found in generated document')
except Exception as e:
    print('Error running process_chunk:', e)
    raise

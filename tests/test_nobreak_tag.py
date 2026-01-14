#!/usr/bin/env python3
"""
Test using Word's w:noBreak tag at run level to prevent breaking.
"""

import sys
sys.path.insert(0, '/Users/adamcordova/Desktop/labelMaker_ QR copy final')

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.core.generation.text_processing import make_nonbreaking_hyphens, make_nonbreaking_weight_units

# Process the text
test_text = "Blueberry - 7g"
processed = make_nonbreaking_hyphens(test_text)
processed = make_nonbreaking_weight_units(processed)

print(f"Testing different run-level approaches for: {repr(processed)}")

def create_test_with_nobreak_runs(filename, split_strategy):
    """Create a test file with different run splitting strategies."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]

    # Set cell width to 2 inches
    cell.width = Inches(2.0)
    tcPr = cell._element.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcPr.append(tcW)
    tcW.set(qn('w:w'), '2880')
    tcW.set(qn('w:type'), 'dxa')

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if split_strategy == "single_run":
        # All text in one run
        run = paragraph.add_run(processed)
        run.font.name = "Arial"
        run.font.bold = True
        run.font.size = Pt(10)

    elif split_strategy == "split_at_hyphen":
        # Split into two runs: "Blueberry " and "‑\xa07g"
        run1 = paragraph.add_run("Blueberry ")
        run1.font.name = "Arial"
        run1.font.bold = True
        run1.font.size = Pt(10)

        run2 = paragraph.add_run("‑\xa07g")
        run2.font.name = "Arial"
        run2.font.bold = True
        run2.font.size = Pt(10)

        # Add noBreak to second run
        rPr = run2._element.get_or_add_rPr()
        noBreak = OxmlElement('w:noBreak')
        rPr.append(noBreak)

    elif split_strategy == "nobreak_on_all":
        # Single run with noBreak
        run = paragraph.add_run(processed)
        run.font.name = "Arial"
        run.font.bold = True
        run.font.size = Pt(10)

        # Add noBreak
        rPr = run._element.get_or_add_rPr()
        noBreak = OxmlElement('w:noBreak')
        rPr.append(noBreak)

    doc.save(filename)
    print(f"✓ Created: {filename} ({split_strategy})")

# Create test files
create_test_with_nobreak_runs('/Users/adamcordova/Desktop/labelMaker_ QR copy final/test_6_single_run.docx', 'single_run')
create_test_with_nobreak_runs('/Users/adamcordova/Desktop/labelMaker_ QR copy final/test_7_split_at_hyphen_with_nobreak.docx', 'split_at_hyphen')
create_test_with_nobreak_runs('/Users/adamcordova/Desktop/labelMaker_ QR copy final/test_8_nobreak_entire_run.docx', 'nobreak_on_all')

print("\n" + "="*60)
print("Also test these files:")
print("  test_6_single_run.docx - All text in single run")
print("  test_7_split_at_hyphen_with_nobreak.docx - Weight unit in separate run with w:noBreak")
print("  test_8_nobreak_entire_run.docx - Entire text with w:noBreak")
print("="*60)

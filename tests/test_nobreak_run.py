#!/usr/bin/env python3
"""
Test using w:noProof to keep text together as a single unit.
"""

import sys
sys.path.insert(0, '/Users/adamcordova/Desktop/labelMaker_ QR copy final')

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_with_separate_runs(filename, description):
    """Split product name into two runs: 'Blueberry ' and '‑ 7g' with special formatting."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]

    # Set cell width to 2 inches (matching actual code)
    cell.width = Inches(2.0)
    tcPr = cell._element.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcPr.append(tcW)
    tcW.set(qn('w:w'), '2880')
    tcW.set(qn('w:type'), 'dxa')

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # First run: "Blueberry " (can wrap)
    run1 = paragraph.add_run("Blueberry ")
    run1.font.name = "Arial"
    run1.font.bold = True
    run1.font.size = Pt(10)

    # Second run: "‑ 7g" (MUST NOT WRAP)
    run2 = paragraph.add_run("‑ 7g")
    run2.font.name = "Arial"
    run2.font.bold = True
    run2.font.size = Pt(10)

    # Apply special formatting to second run based on test type
    rPr = run2._element.get_or_add_rPr()

    if "nobreak" in filename:
        # Add w:noBreak
        noBreak = OxmlElement('w:noBreak')
        rPr.append(noBreak)

    if "noproof" in filename:
        # Add w:noProof (tells Word to treat as single entity)
        noProof = OxmlElement('w:noProof')
        rPr.append(noProof)

    if "keepnext" in filename:
        # Add keepNext at paragraph level
        pPr = paragraph._element.get_or_add_pPr()
        keepNext = OxmlElement('w:keepNext')
        pPr.append(keepNext)

    if "keeplines" in filename:
        # Add keepLines at paragraph level
        pPr = paragraph._element.get_or_add_pPr()
        keepLines = OxmlElement('w:keepLines')
        pPr.append(keepLines)

    doc.save(filename)
    print(f"✓ Created: {filename}")
    print(f"  {description}")

# Create test files
print("="*80)
print("CREATING TEST FILES WITH DIFFERENT RUN-LEVEL APPROACHES")
print("="*80)
print()

create_test_with_separate_runs(
    '/Users/adamcordova/Desktop/labelMaker_ QR copy final/test_9_baseline_two_runs.docx',
    "Baseline: Two runs without special formatting"
)

create_test_with_separate_runs(
    '/Users/adamcordova/Desktop/labelMaker_ QR copy final/test_10_nobreak_on_weight.docx',
    "w:noBreak on '‑ 7g' run"
)

create_test_with_separate_runs(
    '/Users/adamcordova/Desktop/labelMaker_ QR copy final/test_11_noproof_on_weight.docx',
    "w:noProof on '‑ 7g' run"
)

create_test_with_separate_runs(
    '/Users/adamcordova/Desktop/labelMaker_ QR copy final/test_12_keeplines.docx',
    "w:keepLines at paragraph level"
)

create_test_with_separate_runs(
    '/Users/adamcordova/Desktop/labelMaker_ QR copy final/test_13_nobreak_noproof.docx',
    "Both w:noBreak and w:noProof on '‑ 7g' run"
)

print()
print("="*80)
print("Please test these files and tell me which one keeps '‑ 7g' together")
print("="*80)

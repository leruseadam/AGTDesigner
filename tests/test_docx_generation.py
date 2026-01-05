#!/usr/bin/env python3
"""
Test DOCX generation with "Blueberry - 7g" to debug weight unit splitting.
"""

import sys
sys.path.insert(0, '/Users/adamcordova/Desktop/labelMaker_ QR copy final')

from docx import Document
from docx.shared import Pt
from src.core.generation.text_processing import make_nonbreaking_hyphens, make_nonbreaking_weight_units
import zipfile
import os

# Test the text processing
test_text = "Blueberry - 7g"
print(f"Original text: {repr(test_text)}")

# Apply the processing functions in the same order as template_processor.py
processed = make_nonbreaking_hyphens(test_text)
print(f"After make_nonbreaking_hyphens: {repr(processed)}")

processed = make_nonbreaking_weight_units(processed)
print(f"After make_nonbreaking_weight_units: {repr(processed)}")

# Show the actual characters
print(f"\nCharacter breakdown:")
for i, char in enumerate(processed):
    print(f"  [{i}] {repr(char)} (U+{ord(char):04X})")

# Now create a simple DOCX with this text
doc = Document()
table = doc.add_table(rows=1, cols=1)
cell = table.rows[0].cells[0]

# Add the processed text to the cell
paragraph = cell.paragraphs[0]
run = paragraph.add_run(processed)
run.font.size = Pt(10)

# Save the document
output_file = '/Users/adamcordova/Desktop/labelMaker_ QR copy final/test_blueberry_7g.docx'
doc.save(output_file)
print(f"\n✓ Created DOCX: {output_file}")

# Now inspect the XML inside the DOCX
print("\n" + "="*60)
print("Inspecting DOCX XML content:")
print("="*60)

with zipfile.ZipFile(output_file, 'r') as zip_ref:
    # Read the main document XML
    with zip_ref.open('word/document.xml') as xml_file:
        xml_content = xml_file.read().decode('utf-8')

        # Find the text content in the XML
        import re
        text_matches = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)

        print(f"\nFound {len(text_matches)} text runs in XML:")
        for i, text in enumerate(text_matches):
            print(f"\n  Text run {i+1}: {repr(text)}")
            print(f"  Character breakdown:")
            for j, char in enumerate(text):
                print(f"    [{j}] {repr(char)} (U+{ord(char):04X})")

print("\n" + "="*60)
print("Open the file to visually inspect: test_blueberry_7g.docx")
print("="*60)

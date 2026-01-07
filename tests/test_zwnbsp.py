#!/usr/bin/env python3
"""
Test using Zero-Width No-Break Space (U+FEFF) - the BOM character.
This is the STRONGEST no-break character in Unicode.
"""

import sys
sys.path.insert(0, '/Users/adamcordova/Desktop/labelMaker_ QR copy final')

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

test_approaches = [
    ("test_14_zwnbsp.docx", "Blueberry ‑\ufeff7g", "Zero-Width No-Break Space (U+FEFF)"),
    ("test_15_narrow_nbsp.docx", "Blueberry ‑\u202f7g", "Narrow No-Break Space (U+202F)"),
    ("test_16_figure_space.docx", "Blueberry ‑\u2007g", "Figure Space (U+2007) - digit-width space"),
    ("test_17_plain_hyphen_7g.docx", "Blueberry - 7g", "Plain text (no special characters)"),
]

print("="*80)
print("TESTING ADDITIONAL UNICODE NO-BREAK APPROACHES")
print("="*80)
print()

for filename, text, description in test_approaches:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]

    # Set cell width to 2 inches
    cell.width = Inches(2.0)
    tcPr = cell._element.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcPr.append(tcW)
    tcW.set(qn('w:w'), '2880')
    tcW.set(qn('w:type'), 'dxa')

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add formatting like actual code
    pPr = paragraph._element.get_or_add_pPr()
    wordWrap = OxmlElement('w:wordWrap')
    pPr.append(wordWrap)
    wordWrap.set(qn('w:val'), '1')

    suppressAutoHyphens = OxmlElement('w:suppressAutoHyphens')
    pPr.append(suppressAutoHyphens)

    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.bold = True
    run.font.size = Pt(10)

    filepath = f'/Users/adamcordova/Desktop/labelMaker_ QR copy final/{filename}'
    doc.save(filepath)

    print(f"✓ {filename}")
    print(f"  {description}")
    print(f"  Text: {repr(text)}")
    for i, char in enumerate(text[text.index('‑') if '‑' in text else text.index('-'):]):
        if ord(char) > 127:
            print(f"    Char {i}: {repr(char)} (U+{ord(char):04X})")
    print()

print("="*80)
print("KEY:")
print("  U+FEFF = Zero-Width No-Break Space (BOM) - STRONGEST")
print("  U+202F = Narrow No-Break Space")
print("  U+2007 = Figure Space (same width as digits)")
print("="*80)

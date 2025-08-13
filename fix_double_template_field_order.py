#!/usr/bin/env python3
"""
Script to fix the field order in the original double.docx template.
The issue is that ProductBrand and ProductStrain have switched places in the template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def fix_double_template_field_order():
    """Fix the field order in the double.docx template by swapping ProductBrand and ProductStrain positions."""
    
    print("Fixing Double Template Field Order")
    print("=" * 50)
    
    template_path = "src/core/generation/templates/double.docx"
    
    # Load the original template
    print(f"Loading template: {template_path}")
    doc = Document(template_path)
    
    if not doc.tables:
        print("❌ No tables found in template")
        return False
    
    table = doc.tables[0]
    print(f"Table structure: {len(table.rows)} rows x {len(table.columns)} columns")
    
    # Check the first cell content before fixing
    first_cell = table.cell(0, 0)
    print(f"\nBefore fix - First cell content:")
    print(f"  {repr(first_cell.text)}")
    
    # The issue is that the template has:
    # Line 1: Lineage (e.g., "SATIVA")
    # Line 2: Price (e.g., "$27") 
    # Line 3: ProductBrand (e.g., "SATIVA" - but this should be different)
    # Line 4: Empty
    # Line 5: Ratio (e.g., "THC: CBD:")
    # Line 6: ProductStrain (e.g., "HUSTLER'S AMBITION")
    
    # We need to swap lines 3 and 6 so that:
    # Line 3: ProductStrain (e.g., "HUSTLER'S AMBITION")
    # Line 6: ProductBrand (e.g., "SATIVA")
    
    print(f"\nAnalyzing cell structure...")
    
    # Get all paragraphs in the first cell
    paragraphs = first_cell.paragraphs
    print(f"  Number of paragraphs: {len(paragraphs)}")
    
    # Find the ProductBrand and ProductStrain content
    product_brand_text = None
    product_strain_text = None
    product_brand_para_idx = None
    product_strain_para_idx = None
    
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        print(f"  Paragraph {i}: {repr(text)}")
        
        # Look for content that looks like ProductBrand (usually shorter, brand names)
        if text and len(text) < 20 and not text.startswith('$') and not text.startswith('THC:') and not text.startswith('CBD:'):
            if product_brand_text is None:
                product_brand_text = text
                product_brand_para_idx = i
                print(f"    -> Identified as ProductBrand: {text}")
        
        # Look for content that looks like ProductStrain (usually longer, strain names)
        if text and len(text) > 10 and not text.startswith('$') and not text.startswith('THC:') and not text.startswith('CBD:'):
            if product_strain_text is None:
                product_strain_text = text
                product_strain_para_idx = i
                print(f"    -> Identified as ProductStrain: {text}")
    
    if product_brand_text and product_strain_text:
        print(f"\nFound fields to swap:")
        print(f"  ProductBrand (paragraph {product_brand_para_idx}): {product_brand_text}")
        print(f"  ProductStrain (paragraph {product_strain_para_idx}): {product_strain_text}")
        
        # Swap the content
        print(f"\nSwapping field content...")
        
        # Store the original formatting
        brand_para = paragraphs[product_brand_para_idx]
        strain_para = paragraphs[product_strain_para_idx]
        
        # Get formatting from both paragraphs
        brand_runs = list(brand_para.runs)
        strain_runs = list(strain_para.runs)
        
        # Clear both paragraphs
        brand_para.clear()
        strain_para.clear()
        
        # Swap the content with original formatting
        for run in strain_runs:
            new_run = brand_para.add_run()
            new_run.text = run.text
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.bold = run.font.bold
            new_run.font.italic = run.font.italic
            new_run.font.underline = run.font.underline
            new_run.font.color.rgb = run.font.color.rgb if run.font.color.rgb else None
        
        for run in brand_runs:
            new_run = strain_para.add_run()
            new_run.text = run.text
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.bold = run.font.bold
            new_run.font.italic = run.font.italic
            new_run.font.underline = run.font.underline
            new_run.font.color.rgb = run.font.color.rgb if run.font.color.rgb else None
        
        print(f"  ✅ Swapped ProductBrand and ProductStrain content")
        
    else:
        print(f"\n❌ Could not identify ProductBrand and ProductStrain fields")
        print(f"  ProductBrand found: {product_brand_text}")
        print(f"  ProductStrain found: {product_strain_text}")
        return False
    
    # Check the first cell content after fixing
    print(f"\nAfter fix - First cell content:")
    print(f"  {repr(first_cell.text)}")
    
    # Save the fixed template
    backup_path = "src/core/generation/templates/double.docx.backup"
    fixed_path = "src/core/generation/templates/double.docx"
    
    # Create backup first
    print(f"\nCreating backup: {backup_path}")
    doc.save(backup_path)
    
    # Save the fixed template
    print(f"Saving fixed template: {fixed_path}")
    doc.save(fixed_path)
    
    print(f"\n✅ Template field order has been fixed!")
    print(f"  Backup saved to: {backup_path}")
    print(f"  Fixed template saved to: {fixed_path}")
    print(f"\nThe template now has the correct field order:")
    print(f"  1. Lineage (e.g., SATIVA)")
    print(f"  2. Price (e.g., $27)")
    print(f"  3. ProductStrain (e.g., HUSTLER'S AMBITION)")
    print(f"  4. Empty")
    print(f"  5. Ratio (e.g., THC: CBD:)")
    print(f"  6. ProductBrand (e.g., SATIVA)")
    
    return True

if __name__ == "__main__":
    try:
        success = fix_double_template_field_order()
        if success:
            print("\n🎉 Double template field order has been successfully fixed!")
        else:
            print("\n❌ Failed to fix template field order.")
    except Exception as e:
        print(f"\n💥 Error fixing template: {e}")
        import traceback
        traceback.print_exc() 
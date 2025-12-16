#!/usr/bin/env python3
"""
Debug script to examine the vertical template structure
"""

from docx import Document
import os

def examine_template():
    """Examine the vertical template to understand its structure."""
    template_path = "/Users/adamcordova/Desktop/labelMaker_ QR copy final/src/core/generation/templates/vertical.docx"
    
    if not os.path.exists(template_path):
        print(f"❌ Template not found: {template_path}")
        return
    
    print(f"🔍 Examining vertical template: {template_path}")
    
    try:
        doc = Document(template_path)
        
        print(f"\n📊 Document Structure:")
        print(f"  Paragraphs: {len(doc.paragraphs)}")
        print(f"  Tables: {len(doc.tables)}")
        
        # Check paragraphs
        print(f"\n📝 Paragraphs:")
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                print(f"  Paragraph {i}: '{text}'")
        
        # Check tables
        print(f"\n📋 Tables:")
        for table_idx, table in enumerate(doc.tables):
            print(f"  Table {table_idx}: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check each cell
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"    Cell [{row_idx}][{col_idx}]: '{cell_text}'")
                    
                    # Check for placeholders
                    if '{{' in cell_text and '}}' in cell_text:
                        print(f"    🎯 PLACEHOLDER FOUND: '{cell_text}'")
                    
                    # Check runs for split placeholders
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if '{{' in run.text or '}}' in run.text:
                                print(f"    🔍 POTENTIAL SPLIT PLACEHOLDER: '{run.text}'")
        
        print(f"\n✅ Template examination complete")
        
    except Exception as e:
        print(f"❌ Error examining template: {e}")

if __name__ == "__main__":
    examine_template()
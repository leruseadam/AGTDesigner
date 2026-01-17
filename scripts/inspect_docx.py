from docx import Document
from zipfile import ZipFile
import re

p = 'test_preroll_output.docx'
print('Opening', p)

try:
    doc = Document(p)
except Exception as e:
    print('Failed to open with python-docx:', e)
    raise

print('paragraphs:', len(doc.paragraphs))
for i, p_par in enumerate(doc.paragraphs[:40]):
    print('P[{}] text={} runs={}'.format(i, repr(p_par.text), len(p_par.runs)))
    for j, r in enumerate(p_par.runs):
        print('  run[{}]: {}'.format(j, repr(r.text)))

print('\nTables:', len(doc.tables))
for ti, table in enumerate(doc.tables):
    cols = len(table.rows[0].cells) if table.rows else 0
    print('Table', ti, 'rows', len(table.rows), 'cols', cols)
    for ri, row in enumerate(table.rows[:5]):
        for ci, cell in enumerate(row.cells[:10]):
            texts = [pp.text for pp in cell.paragraphs]
            print(' T{} R{}C{}: {}'.format(ti, ri, ci, texts))

print('\nReading raw word/document.xml ...')
with ZipFile(p) as z:
    names = z.namelist()
    print('archive entries:', len(names))
    if 'word/document.xml' in names:
        s = z.read('word/document.xml').decode('utf-8')
        wt = re.findall(r'<w:t[^>]*>(.*?)</w:t>', s, flags=re.DOTALL)
        print('w:t count:', len(wt))
        nonempt = [t for t in wt if t.strip()]
        print('non-empty w:t count:', len(nonempt))
        print('first 40 non-empty w:t samples:')
        for t in nonempt[:40]:
            print('-', repr(t))
    if 'word/settings.xml' in names:
        s2 = z.read('word/settings.xml').decode('utf-8')
        print('\nword/settings.xml contains documentProtection?:', 'documentProtection' in s2)

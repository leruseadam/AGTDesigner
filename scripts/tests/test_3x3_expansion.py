#!/usr/bin/env python3
"""
Test the 3x3 expansion for vertical template to debug placeholder issues
"""

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
import tempfile
import os

def test_3x3_expansion():
    """Test the 3x3 expansion and examine the resulting placeholders."""
    print("🧪 Testing 3x3 expansion for vertical template...")
    
    try:
        # Create a TemplateProcessor for vertical template
        processor = TemplateProcessor('vertical', 'default')
        
        # Call the 3x3 expansion method
        print("🔧 Calling _expand_template_to_3x3_fixed()...")
        expanded_buffer = processor._expand_template_to_3x3_fixed()
        
        # Save to temporary file to examine
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(expanded_buffer.getvalue())
            temp_path = temp_file.name
        
        print(f"💾 Expanded template saved to: {temp_path}")
        
        # Load and examine the expanded template
        doc = Document(temp_path)
        
        print(f"\n📊 Expanded Document Structure:")
        print(f"  Paragraphs: {len(doc.paragraphs)}")
        print(f"  Tables: {len(doc.tables)}")
        
        if doc.tables:
            table = doc.tables[0]
            print(f"  Table dimensions: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check each cell for placeholders
            placeholder_count = 0
            found_placeholders = set()
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    print(f"\n📋 Cell [{row_idx}][{col_idx}]: '{cell_text}'")
                    
                    # Look for placeholders
                    if '{{' in cell_text and '}}' in cell_text:
                        print(f"  🎯 PLACEHOLDERS FOUND!")
                        import re
                        placeholders = re.findall(r'{{[^}]+}}', cell_text)
                        for placeholder in placeholders:
                            found_placeholders.add(placeholder)
                            placeholder_count += 1
                            print(f"    • {placeholder}")
                    
                    # Check runs for split placeholders
                    has_split_placeholders = False
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if '{{' in run.text or '}}' in run.text:
                                if not has_split_placeholders:
                                    print(f"  🔍 Split placeholder components:")
                                    has_split_placeholders = True
                                print(f"    • Run: '{run.text}'")
            
            print(f"\n📊 Summary:")
            print(f"  Total placeholders found: {placeholder_count}")
            print(f"  Unique placeholders: {sorted(found_placeholders)}")
        
        # Clean up
        os.unlink(temp_path)
        print(f"✅ Test complete!")
        
    except Exception as e:
        print(f"❌ Error during 3x3 expansion test: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_3x3_expansion()
#!/usr/bin/env python3
"""
Generate a professional DOCX user manual from the markdown content.
Creates a formatted Word document with cover page, TOC, and proper styling.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()

def set_cell_borders(cell, top=True, bottom=True, left=True, right=True):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    borders = OxmlElement('w:tcBorders')
    
    for border_name, border_val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if border_val:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            borders.append(border)
    
    tcPr.append(borders)

def create_manual():
    """Create the professional DOCX manual."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ===== COVER PAGE =====
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('AGT Designer')
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()  # Spacing
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run('User Manual')
    subtitle_run.font.size = Pt(28)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Add vertical space
    for _ in range(8):
        doc.add_paragraph()
    
    # Document info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run('Professional Cannabis Label Generation')
    info_run.font.size = Pt(14)
    info_run.italic = True
    
    doc.add_paragraph()
    
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version_para.add_run(f'Version 1.0\n{datetime.now().strftime("%B %Y")}')
    version_run.font.size = Pt(12)
    
    # Page break
    add_page_break(doc)
    
    # ===== TABLE OF CONTENTS =====
    toc_title = doc.add_heading('Table of Contents', 1)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc_items = [
        'Overview',
        'Installation',
        'Starting the Application',
        'First-Time Setup: Store Selection',
        'Main Workflow: Creating Labels',
        'Uploading an Excel File',
        'Matching Products from a JSON URL',
        'Filtering and Selecting Tags',
        'Generating and Downloading Labels',
        'Database and Data Tools',
        'Troubleshooting',
        'Quick Reference'
    ]
    
    for i, item in enumerate(toc_items, 1):
        para = doc.add_paragraph()
        para.add_run(f'{i}. ').bold = True
        para.add_run(item)
        para.paragraph_format.left_indent = Inches(0.5)
    
    add_page_break(doc)
    
    # ===== SECTION 1: OVERVIEW =====
    doc.add_heading('1. Overview', 1)
    
    doc.add_paragraph(
        'AGT Designer is a web application for generating professional cannabis product labels. '
        'It uses store-specific Excel files and/or a product database, and can match products from '
        'external JSON URLs (e.g., Cultivera inventory transfers).'
    )
    
    doc.add_heading('What the application does', 2)
    
    # Table for functions
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'Function'
    headers[1].text = 'Description'
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_borders(cell)
    
    data = [
        ('Load product data', 'From store-specific Excel files or the product database.'),
        ('Match products', 'From external JSON URLs (e.g., Cultivera) to your Excel/database products.'),
        ('Select products', 'Use filters and checkboxes to choose which products (tags) to include.'),
        ('Generate labels', 'Produce DOCX label documents in multiple templates.')
    ]
    
    for i, (func, desc) in enumerate(data, 1):
        row = table.rows[i].cells
        row[0].text = func
        row[1].text = desc
        for cell in row:
            set_cell_borders(cell)
    
    doc.add_heading('Available label templates', 2)
    para = doc.add_paragraph()
    para.add_run('Horizontal • Vertical • Mini • Double • Preroll • Inventory')
    para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('Typical workflow', 2)
    workflow_steps = [
        'Select a store.',
        'Load data (upload Excel, use database only, or match a JSON URL).',
        'Filter and select the tags you want.',
        'Choose a template and click Generate Tags.',
        'Download the generated DOCX file.'
    ]
    for step in workflow_steps:
        para = doc.add_paragraph(step, style='List Number')
        para.paragraph_format.left_indent = Inches(0.5)
    
    add_page_break(doc)
    
    # ===== SECTION 2: INSTALLATION =====
    doc.add_heading('2. Installation', 1)
    
    doc.add_heading('2.1 Requirements', 2)
    req_para = doc.add_paragraph()
    req_para.add_run('Python: ').bold = True
    req_para.add_run('3.x')
    req_para = doc.add_paragraph()
    req_para.add_run('Dependencies: ').bold = True
    req_para.add_run('As listed in requirements.txt')
    
    doc.add_heading('2.2 Install dependencies', 2)
    doc.add_paragraph('Choose one method.')
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'Method'
    headers[1].text = 'Command'
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_borders(cell)
    
    methods = [
        ('Automated (recommended)', './install_requirements.sh'),
        ('Cross-platform', 'python3 install_requirements.py'),
        ('Manual', 'pip3 install --user -r requirements.txt\nthen python3 patch_docxcompose.py if needed')
    ]
    
    for i, (method, cmd) in enumerate(methods, 1):
        row = table.rows[i].cells
        row[0].text = method
        row[1].text = cmd
        row[1].paragraphs[0].runs[0].font.name = 'Courier New'
        for cell in row:
            set_cell_borders(cell)
    
    doc.add_heading('2.3 Verify installation', 2)
    doc.add_paragraph('Run:')
    code_para = doc.add_paragraph('python3 -c \'import app; print("Ready to run")\'')
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_para.runs[0].font.name = 'Courier New'
    doc.add_paragraph('If you see Ready to run, the installation is successful.')
    
    add_page_break(doc)
    
    # ===== SECTION 3: STARTING THE APPLICATION =====
    doc.add_heading('3. Starting the Application', 1)
    
    doc.add_heading('3.1 Run the application', 2)
    doc.add_paragraph('From the project root directory:')
    code_para = doc.add_paragraph('python3 app.py')
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_para.runs[0].font.name = 'Courier New'
    doc.add_paragraph('or')
    code_para = doc.add_paragraph('python app.py')
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_para.runs[0].font.name = 'Courier New'
    
    doc.add_heading('3.2 URLs and access', 2)
    doc.add_paragraph('After startup, the console will display:')
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'Access type'
    headers[1].text = 'URL example'
    headers[2].text = 'Notes'
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_borders(cell)
    
    access_data = [
        ('Local (this computer)', 'http://127.0.0.1:8001', 'Port may be 8001–8010 if 8001 is in use.'),
        ('Network (other devices)', 'http://<your-IP>:8001', 'Shown in console; allow the port in your firewall if needed.')
    ]
    
    for i, (access, url, notes) in enumerate(access_data, 1):
        row = table.rows[i].cells
        row[0].text = access
        row[1].text = url
        row[1].paragraphs[0].runs[0].font.name = 'Courier New'
        row[2].text = notes
        for cell in row:
            set_cell_borders(cell)
    
    doc.add_heading('3.3 Optional settings', 2)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'Goal'
    headers[1].text = 'Action'
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_borders(cell)
    
    settings = [
        ('Restrict to this computer only', 'Before starting: export HOST=127.0.0.1 then run the app.'),
        ('Use a different port', 'Before starting: export FLASK_PORT=5001 (or another free port) then run the app.')
    ]
    
    for i, (goal, action) in enumerate(settings, 1):
        row = table.rows[i].cells
        row[0].text = goal
        row[1].text = action
        row[1].paragraphs[0].runs[0].font.name = 'Courier New'
        for cell in row:
            set_cell_borders(cell)
    
    doc.add_heading('3.4 Open the application', 2)
    doc.add_paragraph(
        'In a web browser, go to the URL shown in the console (e.g. http://127.0.0.1:8001). '
        'If other computers cannot connect, allow the application\'s port in your firewall.'
    )
    
    add_page_break(doc)
    
    # ===== SECTION 4: STORE SELECTION =====
    doc.add_heading('4. First-Time Setup: Store Selection', 1)
    
    doc.add_paragraph(
        'A store must be selected before uploading files or generating labels. '
        'Each store has its own Excel data and product database.'
    )
    
    doc.add_heading('Procedure: Select a store', 2)
    steps = [
        'Open the application.',
        'When the store selection modal appears, choose your store from the list.',
        'Click the store name. The page will reload with that store\'s context.',
        'The choice is remembered for your session and can persist across reloads.'
    ]
    for step in steps:
        para = doc.add_paragraph(step, style='List Number')
        para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('Available stores', 2)
    stores_para = doc.add_paragraph()
    stores_para.add_run('AGT Bothell • AGT Burien • AGT Goldbar • AGT Lynnwood • AGT Seattle • AGT Shoreline • AGT Walla Walla • Test')
    stores_para.paragraph_format.left_indent = Inches(0.5)
    
    # Important note
    note_para = doc.add_paragraph()
    note_para.add_run('Important: ').bold = True
    note_para.add_run(
        'You cannot upload an Excel file until a store is selected. The file name should match '
        'the selected store (e.g., a file for Bothell should indicate "Bothell" in the name).'
    )
    note_para.paragraph_format.left_indent = Inches(0.5)
    note_para.paragraph_format.right_indent = Inches(0.5)
    note_para.runs[0].font.color.rgb = RGBColor(204, 0, 0)
    
    add_page_break(doc)
    
    # ===== SECTION 5: MAIN WORKFLOW =====
    doc.add_heading('5. Main Workflow: Creating Labels', 1)
    
    doc.add_heading('Step-by-step workflow', 2)
    
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'Step'
    headers[1].text = 'Action'
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_borders(cell)
    
    workflow_data = [
        ('1', 'Select store (if not already selected).'),
        ('2', 'Load product data: upload Excel, use database only, or use Match JSON with a URL.'),
        ('3', 'Refine the list using the filter dropdowns (Vendor, Brand, Product Type, Lineage, etc.).'),
        ('4', 'Select tags in Current Inventory (left). Selected items appear in Selected (right).'),
        ('5', 'Choose a template from the TEMPLATE dropdown.'),
        ('6', 'Click Generate Tags.'),
        ('7', 'Download the generated DOCX when the process completes.')
    ]
    
    for i, (step, action) in enumerate(workflow_data, 1):
        row = table.rows[i].cells
        row[0].text = step
        row[1].text = action
        for cell in row:
            set_cell_borders(cell)
    
    add_page_break(doc)
    
    # ===== SECTION 6: UPLOADING EXCEL =====
    doc.add_heading('6. Uploading an Excel File', 1)
    
    doc.add_heading('6.1 When to upload', 2)
    doc.add_paragraph('• You have a store-specific Excel inventory file (e.g., from your POS or export).')
    doc.add_paragraph('• You want labels to reflect that file\'s products and data.')
    
    doc.add_heading('6.2 Procedure: Upload an Excel file', 2)
    upload_steps = [
        'Ensure the correct store is selected.',
        'In the filter bar at the top, click Upload Excel.',
        'In the file picker, select an .xlsx or .xls file whose name matches the selected store.',
        'Wait for processing. The file name will appear near the upload button.',
        'When loading finishes, Current Inventory (left column) will list product tags. Use search and filters as needed.'
    ]
    for step in upload_steps:
        para = doc.add_paragraph(step, style='List Number')
        para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('6.3 Rules', 2)
    rules_para = doc.add_paragraph()
    rules_para.add_run('Accepted formats: ').bold = True
    rules_para.add_run('Excel only (.xlsx, .xls).')
    rules_para = doc.add_paragraph()
    rules_para.add_run('File name: ').bold = True
    rules_para.add_run('Must be valid for the selected store (validation runs on upload).')
    rules_para = doc.add_paragraph()
    rules_para.add_run('Session: ').bold = True
    rules_para.add_run('One file per store per session; a new upload replaces the previous one.')
    
    doc.add_heading('6.4 Without an Excel file', 2)
    doc.add_paragraph('• You can still use Match JSON; matching will use the product database when no Excel file is present.')
    doc.add_paragraph('• You can also use database-only mode: no file upload, with the product list coming from the store\'s product database.')
    
    add_page_break(doc)
    
    # ===== SECTION 7: JSON MATCHING =====
    doc.add_heading('7. Matching Products from a JSON URL', 1)
    
    doc.add_paragraph(
        'This feature matches products from an external JSON URL (e.g., Cultivera inventory transfer) '
        'to your Excel/database products and can add them to Selected in one step.'
    )
    
    doc.add_heading('7.1 When to use it', 2)
    doc.add_paragraph('• You have a JSON URL that lists products (e.g., a Cultivera inventory transfer URL).')
    doc.add_paragraph('• You want to match those products to your Excel/database and add them to Selected.')
    
    doc.add_heading('7.2 Procedure: Match products from a JSON URL', 2)
    json_steps = [
        '(Recommended) Select a store and, if you use Excel, upload an Excel file first.',
        'In the filter bar, click Match JSON.',
        'In the JSON Product Matching modal:',
        '  • Paste the full JSON URL into the input (e.g. https://...).',
        '  • Click Match JSON.',
        'Wait for matching to complete (up to about two minutes for large data). Progress may appear in the modal or browser console.',
        'When finished:',
        '  • The modal shows how many products were matched and selected.',
        '  • Matched products are added to Selected (right column).',
        '  • Current Inventory may update with new tags from the match.',
        'Close the modal and continue with filters, template selection, and Generate Tags as needed.'
    ]
    for step in json_steps:
        if step.startswith('  •'):
            para = doc.add_paragraph(step, style='List Bullet')
            para.paragraph_format.left_indent = Inches(1.0)
        else:
            para = doc.add_paragraph(step, style='List Number')
            para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('7.3 URL requirements', 2)
    req_para = doc.add_paragraph()
    req_para.add_run('Protocol: ').bold = True
    req_para.add_run('HTTP or HTTPS (or data:). The app may add https:// if omitted.')
    req_para = doc.add_paragraph()
    req_para.add_run('Content: ').bold = True
    req_para.add_run('The URL must return inventory transfer / product list JSON that the application can parse.')
    
    doc.add_heading('7.4 Detailed match view', 2)
    doc.add_paragraph(
        'Some flows offer a detailed match or "Before & After" view. You can review matches, '
        'use Accept All Matches or accept per item, then Save to apply the selection.'
    )
    
    doc.add_heading('7.5 JSON Inventory Slips (separate feature)', 2)
    doc.add_paragraph(
        'JSON Inventory Slips (e.g., under Data Tools) is separate from Match JSON. '
        'There you paste a JSON URL to generate inventory slip documents, not to match and select tags for label generation.'
    )
    
    add_page_break(doc)
    
    # ===== SECTION 8: FILTERING AND SELECTING =====
    doc.add_heading('8. Filtering and Selecting Tags', 1)
    
    doc.add_heading('8.1 Screen layout', 2)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'Area'
    headers[1].text = 'Content'
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_borders(cell)
    
    layout_data = [
        ('Left', 'Current Inventory — product tags from Excel and/or database (and from JSON match).'),
        ('Center', 'Template selector, Generate Tags button, and controls (Undo, Redo, Clear, Export, Data & Analytics, Reset Cache, Lineage Editor).'),
        ('Right', 'Selected — tags chosen for label generation. Items can be reordered by dragging.')
    ]
    
    for i, (area, content) in enumerate(layout_data, 1):
        row = table.rows[i].cells
        row[0].text = area
        row[1].text = content
        for cell in row:
            set_cell_borders(cell)
    
    doc.add_heading('8.2 Filter dropdowns (top bar)', 2)
    doc.add_paragraph('Use these to narrow Current Inventory:')
    filters = ['Vendor', 'Brand', 'Product Type', 'Lineage (e.g., Sativa, Indica, Hybrid)', 'Weight', 'Price', 'DOH Compliance', 'High CBD']
    for filter_name in filters:
        para = doc.add_paragraph(filter_name, style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph('You can combine multiple filters.')
    
    doc.add_heading('8.3 Selecting tags', 2)
    selecting_items = [
        'Manual: In Current Inventory, check the products you want; they appear in Selected.',
        'Select All: Use when available for a vendor/category to select a whole group.',
        'Search: Use the search box above Current Inventory to find products by name or other text.',
        'Order: Drag items in the Selected column to change order; order can affect layout in the generated document.'
    ]
    for item in selecting_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('8.4 Clearing and resetting', 2)
    doc.add_paragraph('• Clear & Reset (or "Clear Filters"): Clears selected tags and resets filters.')
    doc.add_paragraph('• Undo / Redo: Reverses or re-applies the last selection change.')
    
    add_page_break(doc)
    
    # ===== SECTION 9: GENERATING LABELS =====
    doc.add_heading('9. Generating and Downloading Labels', 1)
    
    doc.add_heading('9.1 Before you generate', 2)
    doc.add_paragraph('• At least one tag must be in Selected (right column).')
    doc.add_paragraph('• Choose the template you want from the TEMPLATE dropdown.')
    
    doc.add_heading('9.2 Template options', 2)
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'Template'
    headers[1].text = 'Use'
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_borders(cell)
    
    templates = [
        ('Horizontal', 'Horizontal label layout.'),
        ('Vertical', 'Vertical label layout.'),
        ('Mini', 'Smaller labels.'),
        ('Double', 'Two-up or double layout.'),
        ('Preroll', 'For pre-roll products.'),
        ('Inventory', 'Inventory-style layout.')
    ]
    
    for i, (template, use) in enumerate(templates, 1):
        row = table.rows[i].cells
        row[0].text = template
        row[1].text = use
        for cell in row:
            set_cell_borders(cell)
    
    doc.add_heading('9.3 Procedure: Generate labels', 2)
    generate_steps = [
        'Set TEMPLATE to the desired layout.',
        'Click Generate Tags.',
        'Wait for generation (seconds to minutes depending on set size). A progress or status message may appear.',
        'When complete, the application will prompt or auto-download a DOCX file (e.g., "Labels.docx"). Save it to your computer.'
    ]
    for step in generate_steps:
        para = doc.add_paragraph(step, style='List Number')
        para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('9.4 After generation', 2)
    doc.add_paragraph('• Use Export Data to download the selected tags as Excel (separate from the label DOCX).')
    doc.add_paragraph('• To change layout, select a different TEMPLATE and click Generate Tags again.')
    
    add_page_break(doc)
    
    # ===== SECTION 10: DATABASE TOOLS =====
    doc.add_heading('10. Database and Data Tools', 1)
    
    doc.add_paragraph('Tools are available from the center column or via Data & Analytics.')
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'Tool'
    headers[1].text = 'Purpose'
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_borders(cell)
    
    tools = [
        ('Export Data', 'Download selected tags as an Excel file.'),
        ('Data & Analytics', 'Open the product database manager: browse products, run analytics, Edit DB (add/edit/delete products).'),
        ('Reset Cache', 'Clear cached data so the next load uses fresh data from Excel/database. Use when the list seems stale.'),
        ('Lineage Editor', 'Manage strain names and lineage (e.g., Sativa/Indica/Hybrid) and related display.')
    ]
    
    for i, (tool, purpose) in enumerate(tools, 1):
        row = table.rows[i].cells
        row[0].text = tool
        row[1].text = purpose
        for cell in row:
            set_cell_borders(cell)
    
    doc.add_heading('10.1 Database manager (Edit DB)', 2)
    db_items = [
        'View the product table, search, and run analytics.',
        'Edit existing products (name, vendor, type, lineage, weight, price, DOH, etc.).',
        'Add new products and Delete products you no longer need.',
        'Changes affect Current Inventory and matching (including JSON match).'
    ]
    for item in db_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('10.2 Backups and health', 2)
    doc.add_paragraph(
        'From the database/analytics UI you may backup the database, restore from backup, '
        'and check database health. Use these for safety and troubleshooting.'
    )
    
    add_page_break(doc)
    
    # ===== SECTION 11: TROUBLESHOOTING =====
    doc.add_heading('11. Troubleshooting', 1)
    
    doc.add_heading('11.1 Store and upload', 2)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'Issue'
    headers[1].text = 'Solution'
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_borders(cell)
    
    store_issues = [
        ('"Please select a store before uploading"', 'Choose a store from the modal first, then upload.'),
        ('Upload rejected (filename/store)', 'Ensure the Excel file name matches the selected store and the file is .xlsx or .xls.')
    ]
    
    for i, (issue, solution) in enumerate(store_issues, 1):
        row = table.rows[i].cells
        row[0].text = issue
        row[1].text = solution
        for cell in row:
            set_cell_borders(cell)
    
    doc.add_heading('11.2 JSON match', 2)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'Issue'
    headers[1].text = 'Solution'
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_borders(cell)
    
    json_issues = [
        ('"Please enter a JSON URL first"', 'Paste a full URL in the JSON Match modal and click Match JSON.'),
        ('Match fails or times out', 'Confirm the URL is reachable in a browser; allow 1–2 minutes for large payloads; ensure store is set and, if using Excel, that a file is uploaded.')
    ]
    
    for i, (issue, solution) in enumerate(json_issues, 1):
        row = table.rows[i].cells
        row[0].text = issue
        row[1].text = solution
        for cell in row:
            set_cell_borders(cell)
    
    doc.add_heading('11.3 Tags and generation', 2)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'Issue'
    headers[1].text = 'Solution'
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_borders(cell)
    
    tag_issues = [
        ('No tags in Current Inventory', 'Confirm store is selected; upload Excel or ensure the product database has products for that store; try Reset Cache and reload.'),
        ('Generate does nothing or no download', 'Ensure at least one tag is in Selected; check browser download settings and pop-up blocker; check app logs or browser console (F12).')
    ]
    
    for i, (issue, solution) in enumerate(tag_issues, 1):
        row = table.rows[i].cells
        row[0].text = issue
        row[1].text = solution
        for cell in row:
            set_cell_borders(cell)
    
    doc.add_heading('11.4 Performance and cache', 2)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'Issue'
    headers[1].text = 'Solution'
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_borders(cell)
    
    perf_issues = [
        ('Slow or stale data', 'Use Reset Cache and reload; for heavy use, restart the application.'),
        ('Port already in use', 'Use the URL shown in the console (app may use 8001–8010). Or set FLASK_PORT to a free port and restart.')
    ]
    
    for i, (issue, solution) in enumerate(perf_issues, 1):
        row = table.rows[i].cells
        row[0].text = issue
        row[1].text = solution
        for cell in row:
            set_cell_borders(cell)
    
    doc.add_heading('11.5 Network (other computers cannot connect)', 2)
    network_items = [
        'Use the network URL printed at startup (e.g. http://<IP>:8001).',
        'Allow the application\'s port in your firewall (Windows/macOS/Linux).',
        'For local-only use, set HOST=127.0.0.1 before starting; use port forwarding or a reverse proxy for remote access.'
    ]
    for item in network_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('11.6 Logs', 2)
    logs_para = doc.add_paragraph()
    logs_para.add_run('Server logs: ').bold = True
    logs_para.add_run('Log viewer (e.g. under /logs or linked from the UI).')
    logs_para = doc.add_paragraph()
    logs_para.add_run('Browser: ').bold = True
    logs_para.add_run('Developer Tools (F12) → Console and Network for front-end and API errors.')
    
    add_page_break(doc)
    
    # ===== SECTION 12: QUICK REFERENCE =====
    doc.add_heading('12. Quick Reference', 1)
    
    table = doc.add_table(rows=11, cols=2)
    table.style = 'Light Grid Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'Task'
    headers[1].text = 'Action'
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_borders(cell)
    
    quick_ref = [
        ('Start application', 'python3 app.py'),
        ('Open application', 'Browser → http://127.0.0.1:8001 (or URL shown in console)'),
        ('Select store', 'Use modal on first load; required before upload'),
        ('Load products', 'Upload Excel or Match JSON with a URL'),
        ('Filter list', 'Use Vendor, Brand, Product Type, Lineage, etc.'),
        ('Select for labels', 'Check items in Current Inventory; review Selected on the right'),
        ('Generate labels', 'Choose TEMPLATE → Generate Tags → download DOCX'),
        ('Export selection', 'Export Data (Excel of selected tags)'),
        ('Manage products', 'Data & Analytics → Edit DB, backups, analytics'),
        ('Fix stale data', 'Reset Cache and/or reload page')
    ]
    
    for i, (task, action) in enumerate(quick_ref, 1):
        row = table.rows[i].cells
        row[0].text = task
        row[1].text = action
        if action.startswith('python') or action.startswith('Browser') or 'http' in action:
            row[1].paragraphs[0].runs[0].font.name = 'Courier New'
        for cell in row:
            set_cell_borders(cell)
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer_para = doc.add_paragraph('Related documentation:')
    footer_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para = doc.add_paragraph('Installation details: INSTALLATION.md')
    footer_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para = doc.add_paragraph('Performance and API: QUICK_START_GUIDE.md, api_endpoints_summary.md')
    footer_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    final_para = doc.add_paragraph('AGT Designer — User Manual v1.0')
    final_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_para.runs[0].italic = True
    
    return doc

if __name__ == '__main__':
    print('Generating professional DOCX user manual...')
    doc = create_manual()
    
    output_path = os.path.join(os.path.dirname(__file__), '..', 'docs', 'AGT_Designer_User_Manual.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    doc.save(output_path)
    print(f'✅ Manual created successfully: {output_path}')
    print(f'   File size: {os.path.getsize(output_path) / 1024:.1f} KB')

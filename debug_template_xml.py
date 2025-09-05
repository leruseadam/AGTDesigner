#!/usr/bin/env python3
"""
Debug script to examine the XML structure of the expanded template
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docx.oxml.ns import qn

def debug_template_xml():
    print("🔍 Examining expanded template XML structure...")
    
    try:
        # Initialize template processor
        processor = TemplateProcessor('mini', 'mini')
        print("✅ Template processor initialized")
        
        # Get expanded template
        if hasattr(processor, '_expanded_template_buffer') and processor._expanded_template_buffer:
            if hasattr(processor._expanded_template_buffer, 'seek'):
                processor._expanded_template_buffer.seek(0)
            
            doc = Document(processor._expanded_template_buffer)
            print(f"📊 Document has {len(doc.tables)} tables")
            
            if doc.tables:
                table = doc.tables[0]
                print(f"📊 Table: {len(table.rows)} rows x {len(table.columns)} columns")
                
                # Examine first cell in detail
                print("\n🔍 Examining first cell XML structure:")
                first_cell = table.cell(0, 0)
                print(f"Cell text: '{first_cell.text}'")
                
                # Get the XML structure
                cell_xml = first_cell._tc.xml
                print(f"\nCell XML length: {len(cell_xml)} characters")
                
                # Look for placeholders in XML
                import re
                placeholders = re.findall(r'\{\{[^}]+\}\}', cell_xml)
                print(f"\nPlaceholders found in XML: {len(placeholders)}")
                for i, ph in enumerate(placeholders[:5]):
                    print(f"  {i+1}. {ph}")
                
                # Check if placeholders are in text elements
                print(f"\n🔍 Checking text elements:")
                text_elements = first_cell._tc.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                print(f"Found {len(text_elements)} text elements")
                
                for i, text_el in enumerate(text_elements[:5]):
                    text_content = text_el.text if text_el.text else ""
                    print(f"  Text element {i+1}: '{text_content}'")
                    if "{{" in text_content:
                        print(f"    ✅ Contains placeholder")
                    else:
                        print(f"    ❌ No placeholder")
                
                # Check for content controls
                print(f"\n🔍 Checking for content controls:")
                content_controls = first_cell._tc.findall('.//w:sdt', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                print(f"Found {len(content_controls)} content controls")
                
                if content_controls:
                    for i, cc in enumerate(content_controls[:3]):
                        print(f"  Content control {i+1}: {cc.tag}")
                        # Check if content control has text
                        cc_text = cc.find('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                        if cc_text is not None and cc_text.text:
                            print(f"    Text: '{cc_text.text}'")
                        else:
                            print(f"    No text content")
                
        else:
            print("❌ No expanded template buffer")
            
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_template_xml()

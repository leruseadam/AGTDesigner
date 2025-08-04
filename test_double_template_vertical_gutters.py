#!/usr/bin/env python3
"""
Test script to verify the double template has both horizontal and vertical gutters.
This script tests that the 6x5 double template now has vertical gutters after every second column.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from io import BytesIO

def test_double_template_vertical_gutters():
    """Test that the double template has both horizontal and vertical gutters."""
    print("🔍 Testing Double Template Vertical and Horizontal Gutters")
    print("=" * 60)
    
    try:
        # Create template processor for double template
        processor = TemplateProcessor('double', {})
        print("✅ Template processor created successfully")
        
        # Force re-expand the template
        processor.force_re_expand_template()
        print("✅ Template re-expanded with vertical and horizontal gutter implementation")
        
        # Get the expanded template
        template_buffer = processor._expand_template_to_4x3_fixed_double()
        template_buffer.seek(0)
        
        # Load the document
        doc = Document(template_buffer)
        
        # Check table structure
        if not doc.tables:
            print("❌ No tables found in document")
            return False
            
        table = doc.tables[0]
        rows = len(table.rows)
        cols = len(table.columns)
        
        print(f"✅ Found table with {rows} rows and {cols} columns")
        
        # Verify dimensions
        expected_rows = 5  # 3 label rows + 2 gutter rows
        expected_cols = 6  # 4 label columns + 2 gutter columns
        
        if rows != expected_rows:
            print(f"❌ Expected {expected_rows} rows, got {rows}")
            return False
        if cols != expected_cols:
            print(f"❌ Expected {expected_cols} columns, got {cols}")
            return False
            
        print(f"✅ Table dimensions are correct ({rows} rows x {cols} columns)")
        
        # Check row heights
        print("\n📏 Row Height Analysis:")
        label_row_height_expected = 2.5 * 72  # 2.5 inches in points
        gutter_row_height_expected = 0.10 * 72  # 0.10 inches in points
        
        for i, row in enumerate(table.rows):
            height_pts = row.height.pt if row.height else 0
            if i in [1, 3]:  # Gutter rows
                expected = gutter_row_height_expected
                row_type = "Gutter"
                tolerance = 1.0  # Allow 1 point tolerance
            else:  # Label rows
                expected = label_row_height_expected
                row_type = "Label"
                tolerance = 1.0  # Allow 1 point tolerance
                
            print(f"   Row {i+1}: {height_pts:.1f}pt ({row_type}) - Expected: {expected:.1f}pt")
            
            if abs(height_pts - expected) > tolerance:
                print(f"   ❌ Height incorrect for row {i+1}")
                return False
            else:
                print(f"   ✅ Height correct for row {i+1}")
        
        # Check column widths
        print("\n📏 Column Width Analysis:")
        label_col_width_expected = 1.125 * 1440  # 1.125 inches in twips
        gutter_col_width_expected = 0.05 * 1440  # 0.05 inches in twips
        
        for i, col in enumerate(table.columns):
            # Get column width from the grid
            grid_cols = table._element.findall('.//w:gridCol', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if i < len(grid_cols):
                width_twips = int(grid_cols[i].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', 0))
                width_inches = width_twips / 1440
                
                if i in [2, 5]:  # Gutter columns
                    expected = gutter_col_width_expected
                    col_type = "Gutter"
                    tolerance = 10  # Allow 10 twips tolerance
                else:  # Label columns
                    expected = label_col_width_expected
                    col_type = "Label"
                    tolerance = 10  # Allow 10 twips tolerance
                
                print(f"   Column {i+1}: {width_inches:.3f}\" ({col_type}) - Expected: {expected/1440:.3f}\"")
                
                if abs(width_twips - expected) > tolerance:
                    print(f"   ❌ Width incorrect for column {i+1}")
                    return False
                else:
                    print(f"   ✅ Width correct for column {i+1}")
        
        # Check cell content
        print("\n📋 Cell Content Analysis:")
        label_cells = 0
        gutter_cells = 0
        
        for r in range(rows):
            for c in range(cols):
                cell = table.cell(r, c)
                cell_text = cell.text.strip()
                
                # Check if this is a gutter cell (gutter row or gutter column)
                is_gutter_row = r in [1, 3]
                is_gutter_col = c in [2, 5]
                
                if is_gutter_row or is_gutter_col:  # Gutter cells
                    if cell_text == "":
                        gutter_cells += 1
                        print(f"   ✅ Gutter cell ({r+1},{c+1}): Empty (correct)")
                    else:
                        print(f"   ❌ Gutter cell ({r+1},{c+1}): Contains text '{cell_text}' (should be empty)")
                        return False
                else:  # Label cells
                    if "Label" in cell_text:
                        label_cells += 1
                        print(f"   ✅ Label cell ({r+1},{c+1}): Contains label placeholder")
                    else:
                        print(f"   ❌ Label cell ({r+1},{c+1}): Missing label placeholder")
                        return False
        
        print(f"\n📊 Summary:")
        print(f"   Label cells: {label_cells}/12 (should be 12)")
        print(f"   Gutter cells: {gutter_cells}/18 (should be 18)")
        
        if label_cells == 12 and gutter_cells == 18:
            print("   ✅ All cells correctly configured")
        else:
            print("   ❌ Cell count mismatch")
            return False
        
        # Visual layout representation
        print("\n🎨 Layout Structure:")
        print("   ┌─────────┬─────────┬───┬─────────┬─────────┬───┐")
        print("   │ Label1  │ Label2  │   │ Label3  │ Label4  │   │")
        print("   │         │         │   │         │         │   │")
        print("   ├─────────┼─────────┼───┼─────────┼─────────┼───┤")
        print("   │         │         │   │         │         │   │  ← 0.10\" horizontal gutter")
        print("   ├─────────┼─────────┼───┼─────────┼─────────┼───┤")
        print("   │ Label5  │ Label6  │   │ Label7  │ Label8  │   │")
        print("   │         │         │   │         │         │   │")
        print("   ├─────────┼─────────┼───┼─────────┼─────────┼───┤")
        print("   │         │         │   │         │         │   │  ← 0.10\" horizontal gutter")
        print("   ├─────────┼─────────┼───┼─────────┼─────────┼───┤")
        print("   │ Label9  │ Label10 │   │ Label11 │ Label12 │   │")
        print("   │         │         │   │         │         │   │")
        print("   └─────────┴─────────┴───┴─────────┴─────────┴───┘")
        print("   │ 1.125\" │ 1.125\" │0.05\"│ 1.125\" │ 1.125\" │0.05\"│")
        print("   ↑        ↑        ↑   ↑        ↑        ↑")
        print("   Label   Label   Vert Label   Label   Vert")
        print("   Col1    Col2    Gut  Col3    Col4    Gut")
        
        print("\n✅ Double template vertical and horizontal gutter test completed successfully!")
        return True
        
    except Exception as e:
        print(f"❌ Error testing double template gutters: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("🚀 Double Template Vertical and Horizontal Gutter Test")
    print("This test verifies that the double template has both horizontal and vertical gutters")
    print("=" * 70)
    
    success = test_double_template_vertical_gutters()
    
    if success:
        print("\n🎉 Test passed! The double template now has both horizontal and vertical gutters.")
    else:
        print("\n💥 Test failed! Check the implementation.") 
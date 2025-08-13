#!/usr/bin/env python3
"""
Test script to check template content directly
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from pathlib import Path

def test_template_content():
    """Test the actual content of template files."""
    
    print("Testing template content directly...")
    
    # Check vertical template
    print("\n1. Checking vertical template...")
    try:
        template_path = Path(__file__).resolve().parent / "src" / "core" / "generation" / "templates" / "vertical.docx"
        print(f"Template path: {template_path}")
        print(f"Template exists: {template_path.exists()}")
        
        if template_path.exists():
            doc = Document(template_path)
            print(f"Document has {len(doc.paragraphs)} paragraphs")
            
            # Show all paragraph text
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    print(f"  Paragraph {i+1}: '{paragraph.text}'")
                    
                    # Show runs
                    for j, run in enumerate(paragraph.runs):
                        if run.text.strip():
                            print(f"    Run {j+1}: '{run.text}' - Font: {run.font.name}, Bold: {run.font.bold}")
            
            # Check tables
            print(f"Document has {len(doc.tables)} tables")
            for i, table in enumerate(doc.tables):
                print(f"  Table {i+1}: {len(table.rows)} rows, {len(table.columns)} columns")
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            print(f"    Cell ({row_idx+1}, {col_idx+1}): '{cell.text}'")
        else:
            print("❌ Template file not found")
            
    except Exception as e:
        print(f"❌ Error reading vertical template: {e}")
        import traceback
        traceback.print_exc()
    
    # Check horizontal template
    print("\n2. Checking horizontal template...")
    try:
        template_path = Path(__file__).resolve().parent / "src" / "core" / "generation" / "templates" / "horizontal.docx"
        print(f"Template path: {template_path}")
        print(f"Template exists: {template_path.exists()}")
        
        if template_path.exists():
            doc = Document(template_path)
            print(f"Document has {len(doc.paragraphs)} paragraphs")
            
            # Show all paragraph text
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    print(f"  Paragraph {i+1}: '{paragraph.text}'")
                    
                    # Show runs
                    for j, run in enumerate(paragraph.runs):
                        if run.text.strip():
                            print(f"    Run {j+1}: '{run.text}' - Font: {run.font.name}, Bold: {run.font.bold}")
            
            # Check tables
            print(f"Document has {len(doc.tables)} tables")
            for i, table in enumerate(doc.tables):
                print(f"  Table {i+1}: {len(table.rows)} rows, {len(table.columns)} columns")
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            print(f"    Cell ({row_idx+1}, {col_idx+1}): '{cell.text}'")
        else:
            print("❌ Template file not found")
            
    except Exception as e:
        print(f"❌ Error reading horizontal template: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_template_content()

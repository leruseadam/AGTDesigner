#!/usr/bin/env python3
"""
Debug script to test placeholder replacement step by step.
This will help identify where the template corruption occurs.
"""

import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), 'src'))

from core.generation.template_processor import TemplateProcessor
from docx import Document
import tempfile

def debug_placeholder_replacement():
    """Debug the placeholder replacement process step by step."""
    print("=== Debugging Placeholder Replacement ===")
    
    try:
        # 1. Load original template
        print("\n1. Loading original mini template...")
        original_path = 'src/core/generation/templates/mini.docx'
        original_doc = Document(original_path)
        print(f"   ✅ Original template loaded: {len(original_doc.tables)} tables")
        
        if original_doc.tables:
            table = original_doc.tables[0]
            print(f"   📊 Original table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check first cell content
            first_cell = table.rows[0].cells[0]
            print(f"   📝 Original first cell: {first_cell.text[:100]}...")
        
        # 2. Create a copy using the working method
        print("\n2. Creating copy using working method...")
        new_doc = Document()
        new_doc._element = original_doc._element
        
        print(f"   ✅ Copy created: {len(new_doc.tables)} tables")
        
        if new_doc.tables:
            table = new_doc.tables[0]
            print(f"   📊 Copied table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check first cell content
            first_cell = table.rows[0].cells[0]
            print(f"   📝 Copied first cell: {first_cell.text[:100]}...")
        
        # 3. Test placeholder replacement manually
        print("\n3. Testing placeholder replacement manually...")
        
        # Create test context
        context = {
            'Label1': {
                'ProductStrain': 'TEST STRAIN',
                'ProductBrand': 'TEST BRAND', 
                'VendorInfo': 'TEST VENDOR',
                'Price': '$10.00',
                'Ratio_or_THC_CBD': 'THC: 20%',
                'Lineage': 'TEST'
            }
        }
        
        print(f"   📋 Created context: {context}")
        
        # 4. Replace placeholders manually step by step
        print("\n4. Replacing placeholders step by step...")
        
        # Process each table in the document
        for table in new_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # Process runs for placeholder replacement
                        for run in paragraph.runs:
                            if run.text:
                                # Handle both double braces {{Label1.Field}} and triple braces {{{Label1.Field}}}
                                text = run.text
                                original_text = text
                                
                                # Find all placeholder patterns
                                import re
                                placeholder_patterns = [
                                    r'\{\{(\w+)\.(\w+)\}\}',  # {{Label1.Field}}
                                    r'\{\{\{(\w+)\.(\w+)\}\}\}'  # {{{Label1.Field}}}
                                ]
                                
                                for pattern in placeholder_patterns:
                                    matches = re.findall(pattern, text)
                                    for match in matches:
                                        label_key, field_name = match
                                        placeholder = f"{{{{{label_key}.{field_name}}}}}"
                                        triple_placeholder = f"{{{{{{{label_key}.{field_name}}}}}}}"
                                        
                                        # Get the value from context
                                        if label_key in context and field_name in context[label_key]:
                                            value = context[label_key][field_name]
                                            if value is None:
                                                value = ""
                                            else:
                                                value = str(value)
                                            
                                            # Replace both placeholder formats
                                            text = text.replace(placeholder, value)
                                            text = text.replace(triple_placeholder, value)
                                            
                                            print(f"      🔄 Replaced {placeholder} with '{value}'")
                                        else:
                                            print(f"      ⚠️  Placeholder not found: {label_key}.{field_name}")
                                
                                # Update the run text if it changed
                                if text != original_text:
                                    run.text = text
                                    print(f"      ✅ Updated run text from '{original_text[:30]}...' to '{text[:30]}...'")
        
        print("   ✅ Placeholder replacement completed")
        
        # 5. Check the result
        print("\n5. Checking result after placeholder replacement...")
        
        if new_doc.tables:
            table = new_doc.tables[0]
            print(f"   📊 Result table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check first cell content
            first_cell = table.rows[0].cells[0]
            print(f"   📝 Result first cell: {first_cell.text[:100]}...")
            
            # Check if placeholders were replaced
            if 'TEST STRAIN' in first_cell.text and 'TEST BRAND' in first_cell.text:
                print("   ✅ Placeholders replaced successfully")
            else:
                print("   ❌ Placeholders not replaced correctly!")
        
        # 6. Save results for comparison
        print("\n6. Saving results for comparison...")
        
        # Save original
        original_doc.save('./mini_original_step_by_step.docx')
        print("   💾 Original saved to: mini_original_step_by_step.docx")
        
        # Save result
        new_doc.save('./mini_result_step_by_step.docx')
        print("   💾 Result saved to: mini_result_step_by_step.docx")
        
        # 7. Compare file sizes
        print("\n7. Comparing file sizes...")
        original_size = os.path.getsize('./mini_original_step_by_step.docx')
        result_size = os.path.getsize('./mini_result_step_by_step.docx')
        
        print(f"   📏 Original size: {original_size} bytes")
        print(f"   📏 Result size: {result_size} bytes")
        print(f"   📊 Size difference: {result_size - original_size} bytes")
        
        if abs(result_size - original_size) < 100:
            print("   ✅ File sizes are similar (good)")
        else:
            print("   ⚠️  File sizes differ significantly (potential issue)")
            
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False
    
    print("\n=== Placeholder Replacement Debug Completed ===")
    return True

if __name__ == "__main__":
    debug_placeholder_replacement()

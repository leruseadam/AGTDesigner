#!/usr/bin/env python3
"""
Simple test to verify vendor justification is working.
"""

import os
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add the src directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from src.core.generation.template_processor import TemplateProcessor

def test_simple_vendor():
    """Test simple vendor justification."""
    print("🧪 Testing simple vendor justification...")
    
    # Simple test data
    test_records = [
        {
            'ProductName': 'Test Product',
            'ProductType': 'Flower',
            'Lineage': 'SATIVA',
            'ProductVendor': '1555 Industrial LLC',
            'Price': '$25.00'
        }
    ]
    
    try:
        # Process
        processor = TemplateProcessor('horizontal', {}, 1.0)
        result_doc = processor.process_records(test_records)
        
        # Save
        result_doc.save("test_simple_vendor_result.docx")
        print("✓ Document saved")
        
        # Check the first cell
        if result_doc.tables:
            table = result_doc.tables[0]
            cell = table.cell(0, 0)
            paragraph = cell.paragraphs[0]
            
            print(f"📝 First cell text: '{paragraph.text}'")
            print(f"🔧 Paragraph alignment: {paragraph.alignment}")
            print(f"📏 Left indent: {paragraph.paragraph_format.left_indent}")
            print(f"📌 Tab stops: {len(paragraph.paragraph_format.tab_stops)}")
            
            # Check runs
            print(f"🏃 Runs: {len(paragraph.runs)}")
            for i, run in enumerate(paragraph.runs):
                print(f"   Run {i}: '{run.text}' (font: {run.font.name}, size: {run.font.size}, bold: {run.font.bold})")
            
            # Check if vendor is properly positioned
            if '1555 Industrial LLC' in paragraph.text:
                print("✅ Vendor text found")
                if paragraph.alignment == 3:  # WD_ALIGN_PARAGRAPH.JUSTIFY
                    print("✅ Correct alignment (JUSTIFY)")
                else:
                    print("❌ Wrong alignment")
                
                if paragraph.paragraph_format.left_indent is None or paragraph.paragraph_format.left_indent == 0:
                    print("✅ No left indent (good)")
                else:
                    print("❌ Left indent present (bad)")
                
                if len(paragraph.paragraph_format.tab_stops) > 0:
                    print("✅ Tab stops present")
                else:
                    print("❌ No tab stops")
                
                # Try to manually apply right alignment to vendor text
                print("\n🔧 Attempting manual vendor alignment...")
                try:
                    # Find the vendor run and make it right-aligned
                    for run in paragraph.runs:
                        if '1555 Industrial LLC' in run.text:
                            print(f"   Found vendor run: '{run.text}'")
                            # Set paragraph to justified alignment
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            print(f"   Set paragraph alignment to JUSTIFY")
                            break
                    
                    # Save the modified document
                    result_doc.save("test_simple_vendor_aligned.docx")
                    print("   ✓ Modified document saved as test_simple_vendor_aligned.docx")
                    
                except Exception as e:
                    print(f"   ❌ Error applying alignment: {e}")
                
                # Try a different approach: LEFT alignment with tab stops
                print("\n🔧 Trying LEFT alignment with tab stops...")
                try:
                    # Set to LEFT alignment
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    print(f"   Set paragraph alignment to LEFT")
                    
                    # Add tab stops to push vendor text right
                    from docx.shared import Inches
                    from docx.enum.text import WD_TAB_ALIGNMENT
                    
                    # Clear existing tab stops
                    paragraph.paragraph_format.tab_stops.clear_all()
                    
                    # Add right-aligned tab stop
                    tab_position = Inches(3.2)
                    paragraph.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)
                    print(f"   Added tab stop at {tab_position}")
                    
                    # Insert tab character before vendor text
                    vendor_run = None
                    for run in paragraph.runs:
                        if '1555 Industrial LLC' in run.text:
                            vendor_run = run
                            break
                    
                    if vendor_run:
                        # Insert tab character before vendor text
                        vendor_run.text = '\t' + vendor_run.text
                        print(f"   Added tab character before vendor text")
                    
                    # Save the modified document
                    result_doc.save("test_simple_vendor_tabbed.docx")
                    print("   ✓ Tabbed document saved as test_simple_vendor_tabbed.docx")
                    
                except Exception as e:
                    print(f"   ❌ Error applying tab stops: {e}")
            else:
                print("❌ Vendor text not found")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_simple_vendor()

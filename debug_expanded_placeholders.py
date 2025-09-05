#!/usr/bin/env python3
"""
Debug script to check what placeholders are in the expanded mini template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
import re

def debug_expanded_placeholders():
    """Debug the placeholders in the expanded mini template."""
    print("🔍 Debugging expanded mini template placeholders...")
    
    try:
        # Create a mini template processor
        processor = TemplateProcessor('mini', {}, 1.0)
        
        print(f"✅ Template processor created")
        
        # Check the expanded template buffer
        if hasattr(processor._expanded_template_buffer, 'seek'):
            processor._expanded_template_buffer.seek(0)
        
        # Load the expanded template
        expanded_doc = Document(processor._expanded_template_buffer)
        print(f"✅ Expanded template loaded")
        print(f"📊 Expanded template has {len(expanded_doc.tables)} tables")
        
        if expanded_doc.tables:
            table = expanded_doc.tables[0]
            print(f"📊 Main table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Get all text from the table
            all_text = table._element.xml
            print(f"📄 Table XML length: {len(all_text)} characters")
            
            # Check for placeholders
            placeholder_patterns = [
                r'\{\{([^}]+)\}\}',  # Standard placeholders
                r'Label(\d+)\.',     # Label patterns
                r'Label(\d+)',       # Just Label numbers
            ]
            
            for i, pattern in enumerate(placeholder_patterns):
                matches = re.findall(pattern, all_text)
                print(f"🔍 Pattern {i+1} '{pattern}': {len(matches)} matches")
                if matches:
                    print(f"  Matches: {matches[:10]}")  # Show first 10
            
            # Check specific cells for placeholders
            print(f"\n🔍 Checking specific cells for placeholders:")
            for i in range(min(3, len(table.rows))):
                for j in range(min(3, len(table.columns))):
                    cell = table.cell(i, j)
                    cell_text = cell.text.strip()
                    print(f"  Cell [{i}][{j}]: '{cell_text}'")
                    
                    # Check for placeholders in this cell
                    placeholders = re.findall(r'\{\{([^}]+)\}\}', cell_text)
                    if placeholders:
                        print(f"    → Found placeholders: {placeholders}")
                    
                    # Check inner tables
                    if cell.tables:
                        print(f"    📋 Has {len(cell.tables)} inner tables")
                        for inner_idx, inner_table in enumerate(cell.tables):
                            for inner_row_idx, inner_row in enumerate(inner_table.rows):
                                for inner_cell_idx, inner_cell in enumerate(inner_row.cells):
                                    inner_text = inner_cell.text.strip()
                                    if inner_text:
                                        print(f"      Inner cell [{inner_row_idx}][{inner_cell_idx}]: '{inner_text}'")
                                        inner_placeholders = re.findall(r'\{\{([^}]+)\}\}', inner_text)
                                        if inner_placeholders:
                                            print(f"        → Found placeholders: {inner_placeholders}")
            
            # Show a sample of the XML content
            print(f"\n📄 Sample XML content (first 1000 chars):")
            print(all_text[:1000])
            
        else:
            print("❌ No tables found in expanded template!")
        
        print("✅ Expanded template placeholder debug completed!")
        
    except Exception as e:
        print(f"❌ An error occurred during debugging: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_expanded_placeholders()

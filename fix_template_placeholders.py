#!/usr/bin/env python3
"""
Script to fix the double template placeholders so they include markers that DocxTemplate will preserve.
This will ensure that when DocxTemplate replaces placeholders, the markers are preserved for font sizing.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def fix_template_placeholders():
    """Fix the double template placeholders to include markers."""
    
    print("Fixing Double Template Placeholders with Markers")
    print("=" * 60)
    
    template_path = "src/core/generation/templates/double.docx"
    
    # Load the current template
    doc = Document(template_path)
    print(f"Loaded template: {template_path}")
    
    if not doc.tables:
        print("ERROR: No tables found in template")
        return
    
    table = doc.tables[0]
    print(f"Table structure: {len(table.rows)} rows x {len(table.columns)} columns")
    
    # Process the first cell (main content cell)
    cell = table.cell(0, 0)
    print("Processing first cell...")
    
    # Clear the cell and rebuild with proper placeholders that include markers
    cell._tc.clear_content()
    
    # Add paragraphs with placeholders that include markers
    # The key is to put the markers INSIDE the placeholder so DocxTemplate preserves them
    # Use ProductBrand_Center instead of ProductBrand for proper centering
    paragraphs = [
        ("{{Label1.Lineage_START}}{{Label1.Lineage}}{{Label1.Lineage_END}}", "Lineage"),
        ("{{Label1.ProductStrain_START}}{{Label1.ProductStrain}}{{Label1.ProductStrain_END}}", "ProductStrain"),
        ("", "Empty"),
        ("{{Label1.Price_START}}{{Label1.Price}}{{Label1.Price_END}}", "Price"),
        ("{{Label1.Description_START}}{{Label1.Description}}{{Label1.Description_END}}", "Description"),
        ("{{Label1.THC_CBD_START}}{{Label1.Ratio_or_THC_CBD}}{{Label1.THC_CBD_END}}", "THC_CBD"),
        ("{{Label1.ProductBrand_Center_START}}{{Label1.ProductBrand_Center}}{{Label1.ProductBrand_Center_END}}", "ProductBrand_Center"),
        ("", "Empty")
    ]
    
    for i, (content, description) in enumerate(paragraphs):
        if content.strip():  # Only add paragraphs with content
            para = cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add the content with proper formatting
            run = para.add_run(content)
            run.font.name = "Arial"
            run.font.size = Pt(12)  # Default size, will be overridden by font sizing system
            
            print(f"  Added paragraph {i}: {description}")
            print(f"    Content: {repr(content)}")
    
    # Save the fixed template
    doc.save(template_path)
    print(f"\n✅ Template placeholders have been fixed!")
    print(f"  Template saved to: {template_path}")
    print("\nThe template now has placeholders that include markers:")
    print("  1. {{Label1.Lineage_START}}{{Label1.Lineage}}{{Label1.Lineage_END}}")
    print("  2. {{Label1.ProductStrain_START}}{{Label1.ProductStrain}}{{Label1.ProductStrain_END}}")
    print("  3. {{Label1.Price_START}}{{Label1.Price}}{{Label1.Price_END}}")
    print("  4. {{Label1.Description_START}}{{Label1.Description}}{{Label1.Description_END}}")
    print("  5. {{Label1.THC_CBD_START}}{{Label1.Ratio_or_THC_CBD}}{{Label1.THC_CBD_END}}")
    print("  6. {{Label1.ProductBrand_Center_START}}{{Label1.ProductBrand_Center}}{{Label1.ProductBrand_Center_END}}")
    print("\nWhen DocxTemplate processes these placeholders, it will:")
    print("  1. Replace {{Label1.Lineage_START}} with 'LINEAGE_START'")
    print("  2. Replace {{Label1.Lineage}} with the actual lineage value")
    print("  3. Replace {{Label1.Lineage_END}} with 'LINEAGE_END'")
    print("  4. Preserve the markers for the font sizing system")
    print("  5. Use ProductBrand_Center for proper brand centering and font sizing")
    print("\n🎉 Double template placeholders have been successfully fixed!")

if __name__ == "__main__":
    fix_template_placeholders() 
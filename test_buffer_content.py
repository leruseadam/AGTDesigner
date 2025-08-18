#!/usr/bin/env python3
"""
Test script to check the expanded template buffer content.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor

def test_buffer_content():
    """Test the expanded template buffer content."""
    
    try:
        print("Testing expanded template buffer content...")
        processor = TemplateProcessor('horizontal', {}, 1.0)
        
        print(f"Template type: {processor.template_type}")
        print(f"Chunk size: {processor.chunk_size}")
        
        # Check if we have an expanded template buffer
        if hasattr(processor, '_expanded_template_buffer'):
            print(f"\nExpanded template buffer exists")
            print(f"Buffer type: {type(processor._expanded_template_buffer)}")
            
            # Check buffer position
            if hasattr(processor._expanded_template_buffer, 'tell'):
                current_pos = processor._expanded_template_buffer.tell()
                print(f"Current buffer position: {current_pos}")
                
                # Reset buffer position
                processor._expanded_template_buffer.seek(0)
                print(f"Buffer position after seek(0): {processor._expanded_template_buffer.tell()}")
            
            # Try to load the document from the buffer
            try:
                from docx import Document
                doc = Document(processor._expanded_template_buffer)
                
                print(f"\nDocument loaded from buffer:")
                print(f"  Tables: {len(doc.tables)}")
                print(f"  Paragraphs: {len(doc.paragraphs)}")
                
                if doc.tables:
                    for i, table in enumerate(doc.tables):
                        print(f"  Table {i}: {len(table.rows)} rows x {len(table.columns)} columns")
                        
                        # Check cell content
                        for row_idx, row in enumerate(table.rows):
                            for col_idx, cell in enumerate(row.cells):
                                cell_text = cell.text.strip()
                                if cell_text:
                                    print(f"    Cell ({row_idx},{col_idx}): '{cell_text[:100]}...'")
                                else:
                                    print(f"    Cell ({row_idx},{col_idx}): Empty")
                
                # Save the document for inspection
                doc.save("test_buffer_content_result.docx")
                print("\n✓ Document from buffer saved as: test_buffer_content_result.docx")
                
            except Exception as e:
                print(f"❌ Error loading document from buffer: {e}")
                import traceback
                traceback.print_exc()
                
        else:
            print("❌ No expanded template buffer found")
            return False
            
        return True
        
    except Exception as e:
        print(f"❌ Error in buffer content test: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    test_buffer_content()

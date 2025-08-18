#!/usr/bin/env python3
"""
Simple test to verify the tab character fix.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

def test_tab_fix():
    """Test the tab character fix."""
    
    print("Testing Tab Character Fix")
    print("=" * 40)
    
    # Create a test document
    doc = Document()
    
    # Add a paragraph
    paragraph = doc.add_paragraph()
    
    # Set paragraph to justified alignment
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Add lineage
    lineage_run = paragraph.add_run("SATIVA")
    lineage_run.font.name = "Arial"
    lineage_run.font.bold = True
    lineage_run.font.size = Pt(12)
    
    # Add tab character using chr(9)
    tab_run = paragraph.add_run(chr(9))
    tab_run.font.name = "Arial"
    tab_run.font.bold = True
    tab_run.font.size = Pt(12)
    
    # Add vendor
    vendor_run = paragraph.add_run("155 Industrial LLC")
    vendor_run.font.name = "Arial"
    vendor_run.font.bold = False
    vendor_run.font.italic = True
    
    # Set tab stops
    paragraph.paragraph_format.tab_stops.clear_all()
    tab_position = Inches(2.8)  # Adjusted position
    paragraph.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)
    
    # Add backup tab stop
    backup_tab_position = Inches(3.0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(backup_tab_position, WD_TAB_ALIGNMENT.RIGHT)
    
    print(f"Paragraph text: '{paragraph.text}'")
    print(f"Paragraph alignment: {paragraph.alignment}")
    print(f"Number of tab stops: {len(paragraph.paragraph_format.tab_stops)}")
    
    # Check tab stop properties
    for i, tab_stop in enumerate(paragraph.paragraph_format.tab_stops):
        print(f"Tab stop {i+1}: position={tab_stop.position}, alignment={tab_stop.alignment}")
    
    # Check runs
    print(f"Number of runs: {len(paragraph.runs)}")
    for i, run in enumerate(paragraph.runs):
        print(f"Run {i+1}: text='{run.text}', bold={run.font.bold}, italic={run.font.italic}")
        if run.text == chr(9):
            print(f"  Run {i+1} contains tab character (ASCII 9)")
    
    # Save the test document
    output_path = "test_tab_fix_output.docx"
    doc.save(output_path)
    print(f"\nTest document saved to: {output_path}")
    
    print("\nTest completed!")
    print("Open the output document to see if the tab stops are working correctly.")

if __name__ == "__main__":
    test_tab_fix()

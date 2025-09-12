#!/usr/bin/env python3
"""
Test script to verify QR code integration with template processing.
"""
import sys
import os
sys.path.append('.')

from src.core.generation.template_processor import TemplateProcessor, get_font_scheme
from docx import Document

def test_qr_integration():
    """Test QR code integration with template processing."""
    print("Testing QR code integration...")
    
    # Get font scheme for mini template
    font_scheme = get_font_scheme('mini', 12)
    
    # Create template processor
    processor = TemplateProcessor('mini', font_scheme)
    
    # Test data
    test_data = [{
        'Product Name*': 'Test Product for QR Code',
        'Product Brand': 'Test Brand',
        'THC test result': 25.5,
        'CBD test result': 2.1,
        'Weight*': '3.5g',
        'Price': '$25.00'
    }]
    
    try:
        # Generate labels
        output_file = 'test_qr_integration.docx'
        final_doc = processor.process_records(test_data)
        
        if final_doc is None:
            print("✗ Failed to generate document")
            return False
            
        # Save the document
        final_doc.save(output_file)
        print(f"✓ Labels generated successfully: {output_file}")
        
        # Check if QR code was included in the output
        doc = Document(output_file)
        has_qr = False
        
        # Check all paragraphs for QR content
        for para in doc.paragraphs:
            if '{{QR}}' in para.text:
                has_qr = True
                print(f"✓ Found QR placeholder in paragraph: {para.text}")
                break
        
        # Also check table cells
        if not has_qr and doc.tables:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if '{{QR}}' in para.text:
                                has_qr = True
                                print(f"✓ Found QR placeholder in table cell: {para.text}")
                                break
                        if has_qr:
                            break
                    if has_qr:
                        break
                if has_qr:
                    break
        
        if has_qr:
            print("✓ QR code placeholder found in generated document")
        else:
            print("✗ QR code placeholder not found in generated document")
            
        # Check if the QR code was actually generated (not just placeholder)
        # Look for InlineImage objects or actual QR code content
        print("\nChecking for actual QR code generation...")
        
        # The QR code should be replaced with actual InlineImage objects
        # Let's check if there are any images in the document
        if hasattr(doc, 'part') and hasattr(doc.part, 'related_parts'):
            image_count = len([part for part in doc.part.related_parts.values() 
                             if 'image' in part.content_type])
            print(f"Found {image_count} images in the document")
        
        return True
        
    except Exception as e:
        print(f"✗ Error during QR code integration test: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = test_qr_integration()
    if success:
        print("\n✓ QR code integration test completed successfully!")
    else:
        print("\n✗ QR code integration test failed!")
#!/usr/bin/env python3
from flask import Flask
from src.core.generation.preroll_tag_generator import generate_preroll_tags
from docx import Document
import os

app = Flask(__name__)
app.secret_key = 'test-secret'

with app.test_request_context('/'):
    records = [
        {
            'Product Name*': 'Khalifa Kush - 7g',
            'Description': 'Khalifa Kush - 7g',
            'Lineage': 'Indica\u00A0\u00AD   Kush   Family',
            'Price': '$14',
            'Vendor': 'TestVendor'
        }
    ]
    cache = type('C', (), {'_d':{}, 'set': lambda self,k,v,timeout=None: self._d.__setitem__(k,v), 'get': lambda self,k,default=None: self._d.get(k, default)})()
    grouped = generate_preroll_tags(records, cache)
    doc = Document()
    doc.add_heading('Preroll Label Test', level=1)
    for rep in grouped:
        doc.add_paragraph(f"ProductName: {rep.get('ProductName')}")
        doc.add_paragraph(f"Lineage: {rep.get('Lineage')}")
        doc.add_paragraph(f"Price: {rep.get('Price')}")
    out = 'test_preroll_simple.docx'
    doc.save(out)
    print('Wrote', out)

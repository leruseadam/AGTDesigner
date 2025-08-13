#!/usr/bin/env python3
"""
Script to manually fix the double template with all required placeholders and markers.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def fix_template_manually():
    """Manually fix the double template with proper placeholders."""
    
    print("Manually Fixing Double Template")
    print("=" * 40)
    
    template_path = "src/core/generation/templates/double.docx"
    
    # Load the template
    doc = Document(template_path)
    print(f"Loaded template: {template_path}")
    
    if not doc.tables:
        print("ERROR: No tables found in template")
        return
    
    table = doc.tables[0]
    print(f"Table structure: {len(table.rows)} rows x {len(table.columns)} columns")
    
    # Process the first cell (main content cell)
    cell = table.cell(0, 0)
    print("Processing first cell...")
    
    # Clear the cell completely
    cell._tc.clear_content()
    
    # Add all required paragraphs with proper placeholders
    paragraphs = [
        ("{{Label1.Lineage_START}}{{Label1.Lineage}}{{Label1.Lineage_END}}", "Lineage"),
        ("{{Label1.ProductStrain_START}}{{Label1.ProductStrain}}{{Label1.ProductStrain_END}}", "ProductStrain"),
        ("", "Empty"),
        ("{{Label1.Price_START}}{{Label1.Price}}{{Label1.Price_END}}", "Price"),
        ("{{Label1.Description_START}}{{Label1.Description}}{{Label1.Description_END}}", "Description"),
        ("{{Label1.THC_CBD_START}}{{Label1.Ratio_or_THC_CBD}}{{Label1.THC_CBD_END}}", "THC_CBD"),
        ("{{Label1.ProductBrand_Center_START}}{{Label1.ProductBrand_Center}}{{Label1.ProductBrand_Center_END}}", "ProductBrand_Center"),
        ("", "Empty")
    ]
    
    for i, (content, description) in enumerate(paragraphs):
        if content.strip():  # Only add paragraphs with content
            para = cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add the content with proper formatting
            run = para.add_run(content)
            run.font.name = "Arial"
            run.font.size = Pt(12)  # Default size, will be overridden by font sizing system
            
            print(f"  Added paragraph {i}: {description}")
            print(f"    Content: {repr(content)}")
        else:
            # Add empty paragraph for spacing
            para = cell.add_paragraph()
            print(f"  Added empty paragraph {i} for spacing")
    
    # Save the fixed template
    doc.save(template_path)
    print(f"\n✅ Template has been manually fixed!")
    print(f"  Template saved to: {template_path}")
    
    # Verify the template structure
    print("\nVerifying template structure...")
    doc2 = Document(template_path)
    table2 = doc2.tables[0]
    cell2 = table2.cell(0, 0)
    
    print(f"Final template has {len(cell2.paragraphs)} paragraphs:")
    for i, para in enumerate(cell2.paragraphs):
        if para.text.strip():
            print(f"  Paragraph {i}: {repr(para.text)}")
        else:
            print(f"  Paragraph {i}: [EMPTY]")
    
    print("\n🎉 Double template has been successfully fixed with all required placeholders!")

if __name__ == "__main__":
    fix_template_manually() 
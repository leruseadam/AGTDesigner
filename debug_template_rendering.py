#!/usr/bin/env python3
"""
Debug script to see exactly what's happening during template rendering.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
import tempfile
import os

def debug_template_rendering():
    """Debug template rendering to see where content is getting mixed up."""
    print("Debugging Template Rendering")
    print("=" * 50)
    
    # Create test record
    test_record = {
        'ProductName': 'Test Product',
        'ProductBrand': 'Alpha Crux, LLC',
        'ProductStrain': 'HYBRID',
        'ProductType': 'flower',
        'ProductVendor': 'Test Vendor',
        'Lineage': 'HYBRID',
        'Price': '$25.00',
        'WeightUnits': '3.5g',
        'Description': 'Test Description'
    }
    
    # Test horizontal template
    print("\n🔍 Testing HORIZONTAL template rendering:")
    print("-" * 40)
    
    try:
        processor = TemplateProcessor(template_type='horizontal', font_scheme='Arial')
        
        # Check what context is being built
        print("📋 Context being built:")
        context = processor._build_label_context(test_record, None)
        print(f"  Lineage: {repr(context.get('Lineage', ''))}")
        print(f"  ProductStrain: {repr(context.get('ProductStrain', ''))}")
        print(f"  ProductVendor: {repr(context.get('ProductVendor', ''))}")
        print(f"  ProductBrand: {repr(context.get('ProductBrand', ''))}")
        
        # Check template structure
        print("\n📄 Template structure:")
        template_path = processor._get_template_path()
        print(f"  Template path: {template_path}")
        
        from docx import Document
        doc = Document(template_path)
        table = doc.tables[0]
        cell = table.cell(0, 0)
        print(f"  Cell text: {repr(cell.text)}")
        print(f"  Paragraphs: {len(cell.paragraphs)}")
        for i, p in enumerate(cell.paragraphs):
            print(f"    Para {i}: {repr(p.text)}")
        
        # Now try to render with the context
        print("\n🎯 Attempting template rendering:")
        try:
            result = processor.process_records([test_record])
            if result:
                print("  ✅ Template rendered successfully")
                
                # Check the output
                if hasattr(result, 'tables') and result.tables:
                    table = result.tables[0]
                    cell = table.cell(0, 0)
                    print(f"  📤 Output cell text: {repr(cell.text)}")
                    
                    # Check each paragraph
                    for i, p in enumerate(cell.paragraphs):
                        print(f"    Para {i}: {repr(p.text)}")
                        for j, run in enumerate(p.runs):
                            print(f"      Run {j}: {repr(run.text)} (font: {run.font.name}, size: {run.font.size})")
                else:
                    print("  ❌ No tables in output")
            else:
                print("  ❌ Template rendering failed")
                
        except Exception as e:
            print(f"  ❌ Template rendering error: {e}")
            import traceback
            traceback.print_exc()
            
    except Exception as e:
        print(f"❌ Error setting up processor: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_template_rendering() 
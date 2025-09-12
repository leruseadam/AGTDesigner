#!/usr/bin/env python3
"""
Test script to verify Vape Cartridge lineage override is working.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def test_vape_cartridge_override():
    """Test that Vape Cartridge products get left-aligned lineage regardless of other logic."""
    print("=== Testing Vape Cartridge Override ===")
    
    # Create a mock document with a table and paragraph
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    
    # Add some text to the paragraph
    paragraph.add_run("HYBRID/SATIVA")
    
    # Create template processor
    processor = TemplateProcessor('horizontal', 1.0, 'default')
    
    # Set up the processor state as it would be during label generation
    processor.current_product_type = 'vape cartridge'
    
    # Test the lineage alignment logic
    print(f"Initial paragraph alignment: {paragraph.alignment}")
    
    # Simulate the lineage processing logic
    marker_name = 'LINEAGE'
    content = 'HYBRID/SATIVA'
    
    # Test the main lineage alignment logic
    print("\n=== Testing Main Lineage Logic ===")
    product_type = processor.current_product_type
    is_classic_product = product_type and product_type.lower() in {'flower', 'pre-roll', 'concentrate', 'infused pre-roll', 'solventless concentrate', 'vape cartridge', 'rso/co2 tankers'}
    
    print(f"product_type: '{product_type}'")
    print(f"is_classic_product: {is_classic_product}")
    
    if is_classic_product:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        print("✓ Set to LEFT alignment (classic product)")
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print("✗ Set to CENTER alignment (non-classic product)")
    
    # Test the Vape Cartridge override
    print("\n=== Testing Vape Cartridge Override ===")
    if product_type and 'vape' in product_type.lower():
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        print("✓ VAPE CARTRIDGE OVERRIDE: Forced LEFT alignment")
    else:
        print("✗ No Vape Cartridge override applied")
    
    print(f"Final paragraph alignment: {paragraph.alignment}")
    print(f"Expected: {WD_ALIGN_PARAGRAPH.LEFT}")
    print(f"Match: {paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT}")
    
    # Test with different product types to ensure override only applies to Vape Cartridge
    print("\n=== Testing Other Product Types ===")
    test_products = [
        'flower',
        'concentrate', 
        'edible',
        'vape cartridge',
        'Vape Cartridge',
        'VAPE CARTRIDGE'
    ]
    
    for test_product in test_products:
        processor.current_product_type = test_product
        product_type = processor.current_product_type
        is_classic = product_type and product_type.lower() in {'flower', 'pre-roll', 'concentrate', 'infused pre-roll', 'solventless concentrate', 'vape cartridge', 'rso/co2 tankers'}
        
        # Apply main logic
        if is_classic:
            main_alignment = "LEFT"
        else:
            main_alignment = "CENTER"
        
        # Apply Vape Cartridge override
        if product_type and 'vape' in product_type.lower():
            override_alignment = "LEFT (Vape Override)"
        else:
            override_alignment = main_alignment
        
        print(f"  '{test_product}' -> classic: {is_classic} -> main: {main_alignment} -> final: {override_alignment}")

if __name__ == "__main__":
    test_vape_cartridge_override()

#!/usr/bin/env python3
"""
Debug script to see what the original template cell contains
"""

import os
import sys
from pathlib import Path

# Add the src directory to the path
sys.path.insert(0, str(Path(__file__).parent / "src"))

from docx import Document

def debug_original_template():
    """Debug what the original template cell contains."""
    
    print("Debugging original template cell content...")
    
    try:
        # Load the original vertical template
        template_path = "src/core/generation/templates/vertical.docx"
        print(f"Template path: {template_path}")
        
        if not os.path.exists(template_path):
            print(f"❌ Template file not found: {template_path}")
            return
        
        # Load the template document
        doc = Document(template_path)
        
        if not doc.tables:
            print("❌ No tables found in template")
            return
        
        # Check the first table and first cell
        table = doc.tables[0]
        print(f"Template has {len(table.rows)} rows and {len(table.columns)} columns")
        
        if table.rows and table.columns:
            cell = table.cell(0, 0)
            print(f"\nFirst cell content:")
            print(f"  Raw text: '{cell.text}'")
            
            # Check individual paragraphs and runs
            for para_idx, paragraph in enumerate(cell.paragraphs):
                print(f"  Paragraph {para_idx + 1}: '{paragraph.text}'")
                for run_idx, run in enumerate(paragraph.runs):
                    print(f"    Run {run_idx + 1}: '{run.text}'")
            
            # Check for specific placeholders
            cell_text = cell.text
            import re
            placeholders = re.findall(r'\{\{.*?\}\}', cell_text)
            
            print(f"\nPlaceholders found in first cell:")
            for placeholder in placeholders:
                print(f"  - {placeholder}")
            
            # Check if DescAndWeight is missing
            if not any('DescAndWeight' in p for p in placeholders):
                print("\n❌ DescAndWeight placeholder is missing from original template")
                print("This is why preroll descriptions aren't working!")
            else:
                print("\n✅ DescAndWeight placeholder found in original template")
                
        else:
            print("❌ Template table is empty")
            
    except Exception as e:
        print(f"❌ Error during debugging: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_original_template()

#!/usr/bin/env python3
"""
Debug script to test mini template expansion step by step.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from io import BytesIO

def debug_mini_expansion():
    """Debug the mini template expansion process."""
    print("🔍 Debugging mini template expansion...")
    
    try:
        # Create a mini template processor
        processor = TemplateProcessor('mini', {}, 1.0)
        
        print(f"✅ Template processor created")
        print(f"📁 Template path: {processor._template_path}")
        
        # Check if template exists
        if not os.path.exists(processor._template_path):
            print(f"❌ Template file not found: {processor._template_path}")
            return
        
        print(f"✅ Template file exists")
        
        # Check the expanded template buffer
        print(f"🔧 Checking expanded template buffer...")
        if hasattr(processor._expanded_template_buffer, 'seek'):
            processor._expanded_template_buffer.seek(0)
        
        # Load the expanded template
        expanded_doc = Document(processor._expanded_template_buffer)
        print(f"✅ Expanded template loaded")
        print(f"📊 Expanded template has {len(expanded_doc.tables)} tables")
        
        if expanded_doc.tables:
            table = expanded_doc.tables[0]
            print(f"📊 Main table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check first few cells
            for i in range(min(3, len(table.rows))):
                for j in range(min(3, len(table.columns))):
                    cell = table.cell(i, j)
                    cell_text = cell.text.strip()
                    print(f"🔍 Cell [{i}][{j}]: '{cell_text[:50]}{'...' if len(cell_text) > 50 else ''}'")
        else:
            print("❌ No tables found in expanded template!")
            
            # Check what's in the document
            print(f"📄 Document has {len(expanded_doc.paragraphs)} paragraphs")
            for i, para in enumerate(expanded_doc.paragraphs[:5]):
                print(f"  Paragraph {i}: '{para.text[:100]}{'...' if len(para.text) > 100 else ''}'")
        
        print("✅ Mini template expansion debug completed!")
        
    except Exception as e:
        print(f"❌ An error occurred during debugging: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_mini_expansion()

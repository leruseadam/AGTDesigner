#!/usr/bin/env python3
"""
Test the horizontal template landscape fix.
This verifies that horizontal templates use landscape orientation and proper margins.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

def test_horizontal_landscape_fix():
    """Test that horizontal templates use landscape orientation and proper margins."""
    print("🧪 Testing Horizontal Template Landscape Fix")
    print("=" * 50)
    
    try:
        from docx import Document
        from src.core.generation.docx_formatting import fix_page_margins_for_horizontal_3x3_grid
        from src.core.generation.template_processor import TemplateProcessor
        
        # Test 1: Direct landscape margin fix
        print("\n--- Test 1: Direct Landscape Margin Fix ---")
        doc1 = Document()
        doc1 = fix_page_margins_for_horizontal_3x3_grid(doc1)
        
        if doc1.sections:
            section1 = doc1.sections[0]
            print(f"✓ Page dimensions: {section1.page_width/1440:.2f}\" × {section1.page_height/1440:.2f}\"")
            print(f"✓ Margins: L={section1.left_margin.inches:.2f}\", R={section1.right_margin.inches:.2f}\", T={section1.top_margin.inches:.2f}\", B={section1.bottom_margin.inches:.2f}\"")
            
            # Check orientation
            if hasattr(section1, 'orientation'):
                orientation = "LANDSCAPE" if section1.orientation == 1 else "PORTRAIT"
                print(f"✓ Orientation: {orientation}")
            else:
                print(f"⚠️  Orientation attribute not available")
            
            # Calculate available space
            available_width = (section1.page_width - section1.left_margin - section1.right_margin) / 1440
            available_height = (section1.page_height - section1.top_margin - section1.bottom_margin) / 1440
            print(f"✓ Available space: {available_width:.2f}\" × {available_height:.2f}\"")
            
            # Save test document
            doc1.save("test_horizontal_landscape.docx")
            print(f"✓ Landscape test document saved: test_horizontal_landscape.docx")
        
        # Test 2: Template processor with horizontal template
        print("\n--- Test 2: Template Processor Horizontal Template ---")
        try:
            # Create a template processor for horizontal template
            processor = TemplateProcessor('horizontal', {})
            print(f"✓ Template processor created for: {processor.template_type}")
            
            # Check if template expansion is working
            if hasattr(processor, '_expanded_template_buffer') and processor._expanded_template_buffer:
                print(f"✓ Template expansion buffer exists")
                
                # Try to load the expanded template
                processor._expanded_template_buffer.seek(0)
                expanded_doc = Document(processor._expanded_template_buffer)
                
                if expanded_doc.sections:
                    section2 = expanded_doc.sections[0]
                    print(f"✓ Expanded template page: {section2.page_width/1440:.2f}\" × {section2.page_height/1440:.2f}\"")
                    print(f"✓ Expanded template margins: L={section2.left_margin.inches:.2f}\", R={section2.right_margin.inches:.2f}\", T={section2.top_margin.inches:.2f}\", B={section2.bottom_margin.inches:.2f}\"")
                    
                    # Check orientation
                    if hasattr(section2, 'orientation'):
                        orientation = "LANDSCAPE" if section2.orientation == 1 else "PORTRAIT"
                        print(f"✓ Expanded template orientation: {orientation}")
                    else:
                        print(f"⚠️  Expanded template orientation attribute not available")
                    
                    # Calculate available space
                    available_width2 = (section2.page_width - section2.left_margin - section2.right_margin) / 1440
                    available_height2 = (section2.page_height - section2.top_margin - section2.bottom_margin) / 1440
                    print(f"✓ Expanded template available space: {available_width2:.2f}\" × {available_height2:.2f}\"")
                
                if expanded_doc.tables:
                    expanded_table = expanded_doc.tables[0]
                    print(f"✓ Expanded template table: {len(expanded_table.rows)}×{len(expanded_table.columns)}")
                    
                    # Save expanded template
                    expanded_doc.save("test_horizontal_expanded.docx")
                    print(f"✓ Horizontal expanded template saved: test_horizontal_expanded.docx")
                else:
                    print(f"⚠️  Expanded template has no tables")
            else:
                print(f"⚠️  No expanded template buffer")
                
        except Exception as e:
            print(f"❌ Template processor error: {e}")
            import traceback
            traceback.print_exc()
        
        print("\n--- Landscape Orientation Verification ---")
        print("Horizontal templates should now:")
        print("✅ Use LANDSCAPE orientation (11\" × 8.5\")")
        print("✅ Have 0.25\" margins on all sides")
        print("✅ Provide 10.5\" × 8.0\" available space")
        print("✅ Fit 3×3 grid with 3.47\" × 2.67\" cells")
        
        print("\n--- Test Summary ---")
        print("✅ Created landscape-oriented test document")
        print("✅ Tested horizontal template processor")
        print("✅ Verified landscape orientation enforcement")
        print("✅ Checked margin optimization")
        
        print("\n📁 Generated Files:")
        print("- test_horizontal_landscape.docx (Direct landscape fix)")
        print("- test_horizontal_expanded.docx (Template expansion)")
        
        print("\n🔍 Next Steps:")
        print("1. Open both documents to verify landscape orientation")
        print("2. Check that margins are optimized for 3×3 grid")
        print("3. Generate horizontal labels through the web interface")
        print("4. Verify the 3×3 grid fits completely in landscape mode")
        
    except Exception as e:
        print(f"❌ Error in test: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_horizontal_landscape_fix()

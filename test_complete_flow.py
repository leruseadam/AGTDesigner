#!/usr/bin/env python3
"""
Test script to test the complete mini template flow
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from io import BytesIO

def test_complete_flow():
    print("🔍 Testing complete mini template flow...")
    
    # Test data that matches what the template expects
    test_data = [
        {
            'ProductName': 'Test Product 1',
            'ProductBrand': 'Test Brand 1',
            'Price': '$25.00',
            'Description': 'Test Description 1',
            'WeightUnits': '1g',
            'Ratio': 'THC: 20%',
            'ProductType': 'Flower',
            'DOH': 'YES'
        },
        {
            'ProductName': 'Test Product 2',
            'ProductBrand': 'Test Brand 2', 
            'Price': '$30.00',
            'Description': 'Test Description 2',
            'WeightUnits': '3.5g',
            'Ratio': 'THC: 25%',
            'ProductType': 'Flower',
            'DOH': 'NO'
        }
    ]
    
    try:
        # Initialize template processor
        print("1. Initializing template processor...")
        processor = TemplateProcessor('mini', 'mini')
        print("✅ Template processor initialized")
        
        # Check expanded template
        print("\n2. Checking expanded template...")
        if hasattr(processor, '_expanded_template_buffer') and processor._expanded_template_buffer:
            print("✅ Expanded template buffer exists")
            
            # Load expanded template
            if hasattr(processor._expanded_template_buffer, 'seek'):
                processor._expanded_template_buffer.seek(0)
            
            doc = Document(processor._expanded_template_buffer)
            print(f"📊 Expanded template: {len(doc.tables)} tables")
            
            if doc.tables:
                table = doc.tables[0]
                print(f"📊 Table: {len(table.rows)} rows x {len(table.columns)} columns")
                
                # Check first few cells for placeholders
                for i in range(min(2, len(table.rows))):
                    for j in range(min(2, len(table.columns))):
                        cell = table.cell(i, j)
                        cell_text = cell.text
                        print(f"  Cell [{i}][{j}]: {cell_text[:50]}...")
        else:
            print("❌ No expanded template buffer")
            return
        
        # Process the data
        print("\n3. Processing data...")
        print(f"📊 Input data: {len(test_data)} records")
        for i, record in enumerate(test_data):
            print(f"  Record {i+1}: {record['ProductBrand']} - {record['Price']}")
        
        # Process chunk
        print("\n4. Processing chunk...")
        result_doc = processor._process_chunk(test_data)
        print("✅ Chunk processing completed")
        
        # Check result
        print("\n5. Checking result...")
        if hasattr(result_doc, 'tables'):
            print(f"📊 Result has {len(result_doc.tables)} tables")
            
            if result_doc.tables:
                table = result_doc.tables[0]
                print(f"📊 Result table: {len(table.rows)} rows x {len(table.columns)} columns")
                
                # Check first few cells for content
                for i in range(min(2, len(table.rows))):
                    for j in range(min(2, len(table.columns))):
                        cell = table.cell(i, j)
                        cell_text = cell.text
                        print(f"  Cell [{i}][{j}]: {cell_text[:100]}...")
                        
                        if cell_text.strip():
                            print(f"    ✅ Has content: {len(cell_text)} characters")
                        else:
                            print(f"    ❌ Empty cell")
        else:
            print("❌ Result has no tables")
        
        # Save result for inspection
        print("\n6. Saving result...")
        output_path = "test_output.docx"
        result_doc.save(output_path)
        print(f"✅ Saved result to {output_path}")
        
        print("\n🎯 Test completed!")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_complete_flow()

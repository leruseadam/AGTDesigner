#!/usr/bin/env python3
"""
Script to restore the double.docx template with proper markers and correct field order.
This will recreate the template with PRODUCTSTRAIN_START/END markers so font sizing works.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def restore_double_template_markers():
    """Restore the double.docx template with proper markers and correct field order."""
    
    print("Restoring Double Template with Proper Markers")
    print("=" * 50)
    
    template_path = "src/core/generation/templates/double.docx"
    
    # Create a new document
    doc = Document()
    
    # Create a table with 3 rows x 4 columns
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    
    # Set up the first cell with proper markers and field order
    first_cell = table.cell(0, 0)
    
    # Clear any existing content
    first_cell._tc.clear_content()
    
    # Add content with proper markers in correct order:
    # 1. Lineage
    # 2. Price  
    # 3. ProductStrain
    # 4. Empty
    # 5. Ratio
    # 6. ProductBrand
    
    # Lineage (first paragraph)
    lineage_para = first_cell.add_paragraph()
    lineage_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lineage_run = lineage_para.add_run("LINEAGE_START SATIVA LINEAGE_END")
    lineage_run.font.name = "Arial"
    lineage_run.font.size = Pt(16)
    lineage_run.font.bold = True
    
    # Price (second paragraph)
    price_para = first_cell.add_paragraph()
    price_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    price_run = price_para.add_run("PRICE_START $27 PRICE_END")
    price_run.font.name = "Arial"
    price_run.font.size = Pt(18)
    price_run.font.bold = True
    
    # ProductStrain (third paragraph) - this should get 1pt font
    strain_para = first_cell.add_paragraph()
    strain_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    strain_run = strain_para.add_run("PRODUCTSTRAIN_START HUSTLER'S AMBITION PRODUCTSTRAIN_END")
    strain_run.font.name = "Arial"
    strain_run.font.size = Pt(16)  # This will be overridden by the font sizing system
    strain_run.font.bold = True
    
    # Empty paragraph (fourth paragraph)
    empty_para = first_cell.add_paragraph()
    empty_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Ratio (fifth paragraph)
    ratio_para = first_cell.add_paragraph()
    ratio_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ratio_run = ratio_para.add_run("THC_CBD_START THC: CBD: THC_CBD_END")
    ratio_run.font.name = "Arial"
    ratio_run.font.size = Pt(6.5)
    ratio_run.font.bold = True
    
    # ProductBrand (sixth paragraph)
    brand_para = first_cell.add_paragraph()
    brand_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_run = brand_para.add_run("PRODUCTBRAND_START SATIVA PRODUCTBRAND_END")
    brand_run.font.name = "Arial"
    brand_run.font.size = Pt(16)
    brand_run.font.bold = True
    
    # Empty paragraph (seventh paragraph)
    final_empty_para = first_cell.add_paragraph()
    final_empty_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set table properties
    table.allow_autofit = False
    table.autofit = False
    
    # Set cell dimensions
    for row in table.rows:
        row.height = Pt(120)
        for cell in row.cells:
            cell.width = Pt(120)
    
    # Save the restored template
    backup_path = "src/core/generation/templates/double.docx.backup2"
    restored_path = "src/core/generation/templates/double.docx"
    
    # Create backup of current template
    print(f"Creating backup of current template: {backup_path}")
    if os.path.exists(restored_path):
        import shutil
        shutil.copy2(restored_path, backup_path)
    
    # Save the restored template
    print(f"Saving restored template: {restored_path}")
    doc.save(restored_path)
    
    print(f"\n✅ Template has been restored with proper markers!")
    print(f"  Backup saved to: {backup_path}")
    print(f"  Restored template saved to: {restored_path}")
    print(f"\nThe template now has:")
    print(f"  1. LINEAGE_START/END markers")
    print(f"  2. PRICE_START/END markers")
    print(f"  3. PRODUCTSTRAIN_START/END markers (will get 1pt font)")
    print(f"  4. Empty space")
    print(f"  5. THC_CBD_START/END markers")
    print(f"  6. PRODUCTBRAND_START/END markers")
    print(f"\nNow the font sizing system can properly apply 1pt font to ProductStrain!")
    
    return True

if __name__ == "__main__":
    try:
        success = restore_double_template_markers()
        if success:
            print("\n🎉 Double template markers have been successfully restored!")
        else:
            print("\n❌ Failed to restore template markers.")
    except Exception as e:
        print(f"\n💥 Error restoring template: {e}")
        import traceback
        traceback.print_exc() 
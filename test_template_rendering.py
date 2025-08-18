#!/usr/bin/env python3
"""
Test template rendering without custom processing.
"""

import os
import sys
from docx import Document

# Add the src directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from src.core.generation.template_processor import TemplateProcessor

def test_template_rendering():
    """Test how the template renders naturally."""
    print("🧪 Testing template rendering without custom processing...")
    
    # Simple test data
    test_records = [
        {
            'ProductName': 'Test Product',
            'ProductType': 'Flower',
            'Lineage': 'SATIVA',
            'ProductVendor': '1555 Industrial LLC',
            'Price': '$25.00'
        }
    ]
    
    try:
        # Process with minimal interference
        processor = TemplateProcessor('horizontal', {}, 1.0)
        
        # Get the expanded template
        processor.force_re_expand_template()
        
        # Create a context with all 9 labels (template expects them)
        context = {}
        for i in range(1, 10):
            if i == 1:
                context[f'Label{i}'] = {
                    'Lineage': 'SATIVA',
                    'ProductVendor': '1555 Industrial LLC',
                    'ProductStrain': 'Test Strain'
                }
            else:
                context[f'Label{i}'] = {}
        
        # Render with DocxTemplate
        from docxtpl import DocxTemplate
        doc = DocxTemplate(processor._expanded_template_buffer)
        doc.render(context)
        
        # Save result
        doc.save("test_template_rendering_result.docx")
        print("✓ Document saved")
        
        # Check the result
        result_doc = Document("test_template_rendering_result.docx")
        if result_doc.tables:
            table = result_doc.tables[0]
            cell = table.cell(0, 0)
            paragraph = cell.paragraphs[0]
            
            print(f"📝 First cell text: '{paragraph.text}'")
            print(f"🔧 Paragraph alignment: {paragraph.alignment}")
            print(f"📏 Left indent: {paragraph.paragraph_format.left_indent}")
            print(f"📌 Tab stops: {len(paragraph.paragraph_format.tab_stops)}")
            
            # Check runs
            print(f"🏃 Runs: {len(paragraph.runs)}")
            for i, run in enumerate(paragraph.runs):
                print(f"   Run {i}: '{run.text}' (font: {run.font.name}, size: {run.font.size}, bold: {run.font.bold})")
            
            # Check if spaces are preserved
            if '  ' in paragraph.text:
                print("✅ Double spaces preserved")
            else:
                print("❌ Double spaces not preserved")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_template_rendering()

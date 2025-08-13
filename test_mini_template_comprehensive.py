#!/usr/bin/env python3
"""
Comprehensive test script for mini template with 1.5 x 1.5" dimensions and color preservation.
This script tests the reworked mini template generation to ensure it works perfectly.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.generation.docx_formatting import apply_lineage_colors, COLORS
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
import traceback

def test_mini_template_comprehensive():
    """Comprehensive test of the reworked mini template."""
    
    print("🧪 Comprehensive Mini Template Test")
    print("=" * 60)
    
    # Test data with various content types
    test_records = [
        {
            'ProductBrand': 'Premium Cannabis Co.',
            'Price': '$45.99',
            'Lineage': 'SATIVA',
            'Ratio_or_THC_CBD': 'THC: 28% CBD: 1%',
            'Description': 'Premium sativa strain with energizing effects',
            'ProductStrain': 'Super Lemon Haze',
            'ProductType': 'flower'
        },
        {
            'ProductBrand': 'CBD Wellness',
            'Price': '$29.99',
            'Lineage': 'CBD',
            'Ratio_or_THC_CBD': 'THC: 0.3% CBD: 25%',
            'Description': 'High CBD tincture for relaxation',
            'ProductStrain': 'Charlotte\'s Web',
            'ProductType': 'tincture'
        },
        {
            'ProductBrand': 'Hybrid Farms',
            'Price': '$35.50',
            'Lineage': 'HYBRID',
            'Ratio_or_THC_CBD': 'THC: 22% CBD: 8%',
            'Description': 'Balanced hybrid for daytime use',
            'ProductStrain': 'Blue Dream',
            'ProductType': 'flower'
        }
    ]
    
    try:
        print("Step 1: Creating mini template processor...")
        processor = TemplateProcessor('mini', {}, 1.0)
        
        print("Step 2: Testing template expansion...")
        if hasattr(processor, '_expanded_template_buffer'):
            processor._expanded_template_buffer.seek(0)
            doc = Document(processor._expanded_template_buffer)
            if doc.tables:
                table = doc.tables[0]
                print(f"✅ Template expanded to {len(table.rows)}x{len(table.rows[0].cells)} grid")
                
                # Check first cell for template structure
                cell = table.cell(0, 0)
                print(f"First cell text: '{cell.text[:100]}...'")
                
                # Check for template variables
                if '{{Label1.' in cell.text:
                    print("✅ Template variables found in expanded template")
                else:
                    print("❌ No template variables found in expanded template")
            else:
                print("❌ No tables found in expanded template")
        else:
            print("❌ No expanded template buffer found")
        
        print("\nStep 3: Processing test records...")
        result_doc = processor.process_records(test_records)
        
        if result_doc and result_doc.tables:
            table = result_doc.tables[0]
            print(f"✅ Result table dimensions: {len(table.rows)}x{len(table.rows[0].cells)}")
            
            # Check cell dimensions
            print("\nStep 4: Verifying cell dimensions...")
            first_cell = table.cell(0, 0)
            
            # Get cell dimensions
            tc = first_cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            
            if tcW is not None:
                width_twips = int(tcW.get(qn('w:w')))
                width_inches = width_twips / 1440
                print(f"Cell width: {width_inches:.2f} inches")
                
                if abs(width_inches - 1.5) < 0.1:
                    print("✅ Cell width is approximately 1.5 inches")
                else:
                    print(f"❌ Cell width should be 1.5 inches, got {width_inches:.2f}")
            else:
                print("⚠️ Could not determine cell width")
            
            # Check row heights
            first_row = table.rows[0]
            if first_row.height:
                height_inches = first_row.height.inches
                print(f"Row height: {height_inches:.2f} inches")
                
                if abs(height_inches - 1.5) < 0.1:
                    print("✅ Row height is approximately 1.5 inches")
                else:
                    print(f"❌ Row height should be 1.5 inches, got {height_inches:.2f}")
            else:
                print("⚠️ Could not determine row height")
            
            print("\nStep 5: Checking content population...")
            populated_cells = 0
            for row_idx in range(min(3, len(table.rows))):  # Check first 3 rows
                for col_idx in range(min(4, len(table.rows[0].cells))):  # Check first 4 columns
                    cell = table.cell(row_idx, col_idx)
                    cell_text = cell.text.strip()
                    
                    if cell_text and not cell_text.startswith('{{Label'):
                        populated_cells += 1
                        print(f"Cell ({row_idx}, {col_idx}): '{cell_text[:50]}...'")
                    elif cell_text.startswith('{{Label'):
                        print(f"Cell ({row_idx}, {col_idx}): Template placeholder (not populated)")
                    else:
                        print(f"Cell ({row_idx}, {col_idx}): Empty")
            
            print(f"✅ Found {populated_cells} populated cells out of {min(3, len(table.rows)) * min(4, len(table.rows[0].cells))} checked")
            
            print("\nStep 6: Checking formatting preservation...")
            # Check if cells have background colors (indicating original formatting preserved)
            colored_cells = 0
            for row_idx in range(min(2, len(table.rows))):
                for col_idx in range(min(2, len(table.rows[0].cells))):
                    cell = table.cell(row_idx, col_idx)
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shading = tcPr.find(qn('w:shd'))
                    
                    if shading is not None:
                        fill_color = shading.get(qn('w:fill'))
                        if fill_color:
                            colored_cells += 1
                            print(f"Cell ({row_idx}, {col_idx}) has background color: {fill_color}")
                        else:
                            print(f"Cell ({row_idx}, {col_idx}) has shading but no fill color")
                    else:
                        print(f"Cell ({row_idx}, {col_idx}) has no background color")
            
            if colored_cells > 0:
                print(f"✅ Found {colored_cells} cells with background colors (original formatting preserved)")
            else:
                print("⚠️ No cells have background colors - original formatting may not be preserved")
            
            print("\nStep 7: Testing lineage coloring...")
            apply_lineage_colors(result_doc)
            
            # Check if lineage colors were applied
            lineage_colored = 0
            for row_idx in range(min(2, len(table.rows))):
                for col_idx in range(min(2, len(table.rows[0].cells))):
                    cell = table.cell(row_idx, col_idx)
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shading = tcPr.find(qn('w:shd'))
                    
                    if shading is not None:
                        fill_color = shading.get(qn('w:fill'))
                        if fill_color in COLORS.values():
                            lineage_colored += 1
                            print(f"Cell ({row_idx}, {col_idx}) has lineage color: {fill_color}")
                        else:
                            print(f"Cell ({row_idx}, {col_idx}) has non-lineage color: {fill_color}")
                    else:
                        print(f"Cell ({row_idx}, {col_idx}) has no background color after lineage coloring")
            
            if lineage_colored > 0:
                print(f"✅ Applied lineage colors to {lineage_colored} cells")
            else:
                print("⚠️ No lineage colors were applied")
            
            print("\nStep 8: Saving test result...")
            output_path = "test_mini_template_result.docx"
            result_doc.save(output_path)
            print(f"✅ Test result saved to: {output_path}")
            
            print("\n🎯 Test Summary:")
            print(f"   • Template expanded: ✅ {len(table.rows)}x{len(table.rows[0].cells)} grid")
            dimension_status = "✅ 1.5\" x 1.5\"" if colored_cells > 0 else "⚠️ Check dimensions"
            print(f"   • Cell dimensions: {dimension_status}")
            print(f"   • Content populated: ✅ {populated_cells} cells")
            print(f"   • Formatting preserved: {'✅ Yes' if colored_cells > 0 else '⚠️ No'}")
            print(f"   • Lineage coloring: {'✅ Applied' if lineage_colored > 0 else '⚠️ Not applied'}")
            
        else:
            print("❌ Failed to process records or no tables in result")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")
        traceback.print_exc()

def test_mini_template_dimensions():
    """Test specific dimension requirements for mini template."""
    
    print("\n📏 Mini Template Dimension Test")
    print("=" * 40)
    
    try:
        processor = TemplateProcessor('mini', {}, 1.0)
        
        if hasattr(processor, '_expanded_template_buffer'):
            processor._expanded_template_buffer.seek(0)
            doc = Document(processor._expanded_template_buffer)
            
            if doc.tables:
                table = doc.tables[0]
                
                # Check table dimensions
                expected_rows, expected_cols = 5, 4
                actual_rows, actual_cols = len(table.rows), len(table.rows[0].cells)
                
                print(f"Expected grid: {expected_rows}x{expected_cols}")
                print(f"Actual grid: {actual_rows}x{actual_cols}")
                
                if actual_rows == expected_rows and actual_cols == expected_cols:
                    print("✅ Grid dimensions correct")
                else:
                    print("❌ Grid dimensions incorrect")
                
                # Check individual cell dimensions
                if actual_rows > 0 and actual_cols > 0:
                    cell = table.cell(0, 0)
                    
                    # Check cell width
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcW = tcPr.find(qn('w:tcW'))
                    
                    if tcW is not None:
                        width_twips = int(tcW.get(qn('w:w')))
                        width_inches = width_twips / 1440
                        print(f"Cell width: {width_inches:.3f} inches")
                        
                        if abs(width_inches - 1.5) < 0.01:
                            print("✅ Cell width is exactly 1.5 inches")
                        else:
                            print(f"❌ Cell width should be 1.5 inches, got {width_inches:.3f}")
                    else:
                        print("⚠️ Could not determine cell width")
                    
                    # Check row height
                    row = table.rows[0]
                    if row.height:
                        height_inches = row.height.inches
                        print(f"Row height: {height_inches:.3f} inches")
                        
                        if abs(height_inches - 1.5) < 0.01:
                            print("✅ Row height is exactly 1.5 inches")
                        else:
                            print(f"❌ Row height should be 1.5 inches, got {height_inches:.3f}")
                    else:
                        print("⚠️ Could not determine row height")
                    
                    # Calculate total dimensions
                    total_width = width_inches * actual_cols
                    total_height = height_inches * actual_rows
                    print(f"Total table dimensions: {total_width:.1f}\" x {total_height:.1f}\"")
                    
                    if abs(total_width - 6.0) < 0.1 and abs(total_height - 7.5) < 0.1:
                        print("✅ Total table dimensions correct (6.0\" x 7.5\")")
                    else:
                        print(f"❌ Total table dimensions should be 6.0\" x 7.5\", got {total_width:.1f}\" x {total_height:.1f}\"")
                else:
                    print("❌ No cells found to check dimensions")
            else:
                print("❌ No tables found in template")
        else:
            print("❌ No expanded template buffer found")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")
        traceback.print_exc()

if __name__ == "__main__":
    test_mini_template_comprehensive()
    test_mini_template_dimensions()

#!/usr/bin/env python3
"""
Improved DOH Image Centering Fix
This script provides a better solution for centering DOH placeholder InlineImages.
"""

import re

def print_fix_plan():
    """Print the improved fix plan."""
    
    print("🔧 Improved DOH Image Centering Fix")
    print("=" * 45)
    print()
    
    print("📋 ISSUE IDENTIFIED:")
    print("-" * 25)
    print("• DOH placeholder InlineImages are not properly centered")
    print("• Current centering method clears cell content")
    print("• InlineImage objects need special handling")
    print("• Need to preserve existing content while centering")
    print()
    
    print("🛠️ IMPROVED FIX PLAN:")
    print("-" * 25)
    print("1. Improve DOH image detection without clearing cells")
    print("2. Apply centering at paragraph level only")
    print("3. Preserve InlineImage objects intact")
    print("4. Add proper spacing and alignment")
    print("5. Handle both DOH.png and HighCBD.png images")
    print()

def apply_improved_doh_centering():
    """Apply the improved DOH centering fix."""
    
    print("🔧 Applying Improved DOH Centering Fix...")
    
    # Read the current template processor
    with open('src/core/generation/template_processor.py', 'r') as f:
        content = f.read()
    
    # Define the improved DOH centering function
    improved_function = '''    def _ensure_doh_image_centering(self, doc):
        """
        Ensure DOH images are properly centered in all cells.
        This method provides improved centering for InlineImage objects without clearing content.
        """
        try:
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Check if this cell contains a DOH image
                        has_doh_image = False
                        image_paragraph = None
                        
                        # Improved image detection without clearing content
                        for paragraph in cell.paragraphs:
                            paragraph_text = paragraph.text.strip()
                            
                            # Check for DOH placeholder text
                            if 'DOH' in paragraph_text or '{{' in paragraph_text:
                                # Look for drawing elements (InlineImage) in this paragraph
                                for run in paragraph.runs:
                                    if hasattr(run, '_element'):
                                        # Check for drawing elements (InlineImage)
                                        if run._element.find(qn('w:drawing')) is not None:
                                            has_doh_image = True
                                            image_paragraph = paragraph
                                            break
                                        # Check for picture elements
                                        elif run._element.find(qn('w:pict')) is not None:
                                            has_doh_image = True
                                            image_paragraph = paragraph
                                            break
                                
                                if has_doh_image:
                                    break
                        
                        if has_doh_image and image_paragraph:
                            self.logger.debug("Found DOH image, applying improved centering")
                            
                            # Apply centering at paragraph level without clearing content
                            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Set minimal spacing
                            image_paragraph.paragraph_format.space_before = Pt(0)
                            image_paragraph.paragraph_format.space_after = Pt(0)
                            image_paragraph.paragraph_format.line_spacing = 1.0
                            
                            # Set cell vertical alignment to center
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            
                            # Ensure proper XML-level centering
                            pPr = image_paragraph._element.get_or_add_pPr()
                            
                            # Set paragraph justification to center
                            jc = pPr.find(qn('w:jc'))
                            if jc is None:
                                jc = OxmlElement('w:jc')
                                pPr.append(jc)
                            jc.set(qn('w:val'), 'center')
                            
                            # Remove any existing spacing that might interfere
                            existing_spacing = pPr.find(qn('w:spacing'))
                            if existing_spacing is not None:
                                pPr.remove(existing_spacing)
                            
                            # Add minimal spacing
                            spacing = OxmlElement('w:spacing')
                            spacing.set(qn('w:before'), '0')
                            spacing.set(qn('w:after'), '0')
                            spacing.set(qn('w:line'), '240')  # 1.0 line spacing
                            spacing.set(qn('w:lineRule'), 'auto')
                            pPr.append(spacing)
                            
                            # Ensure proper indentation
                            ind = pPr.find(qn('w:ind'))
                            if ind is None:
                                ind = OxmlElement('w:ind')
                                pPr.append(ind)
                            ind.set(qn('w:left'), '0')
                            ind.set(qn('w:right'), '0')
                            ind.set(qn('w:firstLine'), '0')
                            ind.set(qn('w:hanging'), '0')
                            
                            # Ensure the image run itself is properly formatted
                            for run in image_paragraph.runs:
                                if hasattr(run, '_element') and (run._element.find(qn('w:drawing')) is not None or run._element.find(qn('w:pict')) is not None):
                                    # Ensure the run has proper text content
                                    if not run.text or run.text.strip() == '':
                                        run.text = '\\u00A0'  # Non-breaking space
                                    break
                            
                            self.logger.debug("Applied improved DOH image centering")
                                
        except Exception as e:
            self.logger.warning(f"Error in improved DOH image centering: {e}")'''
    
    # Replace the existing function with the improved version
    pattern = r'def _ensure_doh_image_centering\(self, doc\):.*?(?=\n    def|\n\n    def|\Z)'
    
    if re.search(pattern, content, re.DOTALL):
        # Replace the existing function
        content = re.sub(pattern, improved_function, content, flags=re.DOTALL)
        print("✅ Replaced existing DOH centering function with improved version")
    else:
        # Add the function if it doesn't exist
        # Find a good place to insert it (after the _post_process_and_replace_content method)
        insert_point = content.find('def _post_process_and_replace_content(self, doc):')
        if insert_point != -1:
            # Find the end of this method
            method_end = content.find('\n    def ', insert_point)
            if method_end == -1:
                method_end = len(content)
            
            # Insert the new function before the next method
            content = content[:method_end] + '\n\n' + improved_function + '\n' + content[method_end:]
            print("✅ Added improved DOH centering function")
        else:
            print("❌ Could not find insertion point for DOH centering function")
            return False
    
    # Write the updated content back
    with open('src/core/generation/template_processor.py', 'w') as f:
        f.write(content)
    
    print("✅ Improved DOH centering fix applied successfully")
    return True

def create_test_script():
    """Create a test script to verify the improved DOH centering."""
    
    test_script = '''#!/usr/bin/env python3
"""
Test script for improved DOH image centering.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_DOUBLE
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def test_improved_doh_centering():
    """Test the improved DOH image centering."""
    print("🧪 Testing Improved DOH Image Centering")
    print("=" * 45)
    
    # Create a test record with DOH image
    test_record = {
        'Description': 'Test Product with DOH',
        'DOH': 'YES',
        'Product Type*': 'Flower',
        'Product Name*': 'Test Product',
        'Brand': 'Test Brand',
        'Price': '$10.00',
        'THC': '15.5%',
        'CBD': '0.5%',
        'Lineage': 'SATIVA'
    }
    
    # Create template processor
    processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
    
    # Process the test record
    try:
        result = processor.process_records([test_record])
        print("✅ Document generated successfully")
        
        # Check if DOH images are properly centered
        doh_images_found = 0
        centered_images = 0
        
        for table in result.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # Check for DOH images
                        for run in paragraph.runs:
                            if hasattr(run, '_element'):
                                if run._element.find(qn('w:drawing')) is not None or run._element.find(qn('w:pict')) is not None:
                                    doh_images_found += 1
                                    
                                    # Check if paragraph is centered
                                    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                                        centered_images += 1
                                        print(f"✅ Found centered DOH image in cell")
                                    else:
                                        print(f"❌ Found DOH image but not centered")
        
        print(f"\\n📊 Results:")
        print(f"• DOH images found: {doh_images_found}")
        print(f"• Centered images: {centered_images}")
        
        if doh_images_found > 0 and centered_images == doh_images_found:
            print("✅ SUCCESS: All DOH images are properly centered")
            return True
        else:
            print("❌ FAILURE: Some DOH images are not properly centered")
            return False
            
    except Exception as e:
        print(f"❌ Error testing DOH centering: {e}")
        return False

if __name__ == "__main__":
    print("🎯 Improved DOH Image Centering Test")
    print("=" * 40)
    
    # Test the improved DOH centering
    test_passed = test_improved_doh_centering()
    
    if test_passed:
        print("\\n✅ TEST PASSED: Improved DOH image centering is working correctly")
    else:
        print("\\n❌ TEST FAILED: Improved DOH image centering needs more work")
'''
    
    with open('test_improved_doh_centering.py', 'w') as f:
        f.write(test_script)
    
    print("✅ Created test script: test_improved_doh_centering.py")

def main():
    """Main function to apply the improved DOH centering fix."""
    
    print_fix_plan()
    
    # Apply the improved fix
    if apply_improved_doh_centering():
        # Create test script
        create_test_script()
        
        print("\\n✅ IMPROVED DOH CENTERING FIX COMPLETED!")
        print("=" * 45)
        print()
        print("📋 WHAT WAS IMPROVED:")
        print("• Better DOH image detection without clearing cells")
        print("• Preserved InlineImage objects intact")
        print("• Applied centering at paragraph level only")
        print("• Added proper spacing and alignment")
        print("• Improved error handling")
        print()
        print("🚀 NEXT STEPS:")
        print("1. Test the fix: python test_improved_doh_centering.py")
        print("2. Generate a document with DOH images")
        print("3. Check that images are properly centered")
        print("4. Verify no content is lost during centering")
        print()
        print("📁 FILES CREATED:")
        print("• test_improved_doh_centering.py")
        print()
        print("🔧 The improved DOH centering should now work better with InlineImage objects!")

if __name__ == "__main__":
    main() 
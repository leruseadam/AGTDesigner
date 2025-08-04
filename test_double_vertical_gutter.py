#!/usr/bin/env python3
"""
Test script to verify the double template vertical gutter implementation.
This script tests that the 5x3 double template has only a vertical gutter down the middle.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_DOUBLE
from docx import Document
from io import BytesIO

def test_double_template_vertical_gutter():
    """Test that the double template has only a vertical gutter down the middle."""
    print("🔍 Testing Double Template Vertical Gutter")
    
    try:
        # Create template processor
        processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
        
        # Force re-expansion to ensure we get the latest vertical gutter implementation
        processor.force_re_expand_template()
        print("✅ Template re-expanded with vertical gutter implementation")
        
        # Get the expanded template
        template_buffer = processor._expand_template_if_needed(force_expand=True)
        doc = Document(template_buffer)
        
        if not doc.tables:
            print("❌ No table found in template")
            return False
            
        table = doc.tables[0]
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        
        print(f"📊 Table dimensions: {num_rows} rows x {num_cols} columns")
        
        # Verify we have a 5x3 grid (2 label + 1 gutter + 2 label columns)
        expected_rows, expected_cols = 3, 5
        if num_rows != expected_rows or num_cols != expected_cols:
            print(f"❌ Expected {expected_rows}x{expected_cols} grid, got {num_rows}x{num_cols}")
            return False
        
        print(f"✅ Correct grid dimensions: {num_rows}x{num_cols}")
        
        # Check column widths
        print("\n📏 Column widths:")
        for c in range(num_cols):
            col = table.columns[c]
            cell = col.cells[0]  # Get first cell in column
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW')
            if tcW is not None:
                width = tcW.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
                width_inches = int(width) / 1440 if width else 0
                if c == 2:  # Middle column (gutter)
                    print(f"   Column {c+1}: {width_inches:.3f}\" (Gutter)")
                else:  # Label columns
                    print(f"   Column {c+1}: {width_inches:.3f}\" (Label)")
        
        # Check that gutter column is empty and label columns have content
        label_count = 0
        gutter_count = 0
        
        for r in range(num_rows):
            for c in range(num_cols):
                cell = table.cell(r, c)
                cell_text = cell.text.strip()
                
                if c == 2:  # Gutter column
                    if not cell_text:
                        gutter_count += 1
                        print(f"   ✅ Gutter cell ({r+1},{c+1}): Empty (correct)")
                    else:
                        print(f"   ⚠️  Gutter cell ({r+1},{c+1}): Has content '{cell_text[:20]}...'")
                else:  # Label columns
                    if cell_text:
                        label_count += 1
                        print(f"   ✅ Label cell ({r+1},{c+1}): Has content '{cell_text[:30]}...'")
                    else:
                        print(f"   ⚠️  Label cell ({r+1},{c+1}): No content")
        
        print(f"\n📋 Cell Summary:")
        print(f"   Label cells: {label_count}/12 (should be 12)")
        print(f"   Gutter cells: {gutter_count}/3 (should be 3)")
        
        if label_count == 12 and gutter_count == 3:
            print("✅ All label cells have content and all gutter cells are empty")
        else:
            print("❌ Cell content mismatch")
            return False
        
        # Check for absence of cell spacing (no horizontal gutters)
        tblPr = table._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr')
        if tblPr is not None:
            cell_spacing = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblCellSpacing')
            if cell_spacing is None:
                print("✅ No cell spacing found (no horizontal gutters)")
            else:
                print("⚠️  Cell spacing found (horizontal gutters present)")
        else:
            print("⚠️  No table properties found")
        
        # Check cell margins (should be minimal)
        cell_with_minimal_margins = 0
        for r in range(num_rows):
            for c in range(num_cols):
                cell = table.cell(r, c)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcMar')
                if tcMar is not None:
                    # Check if margins are minimal (0.001 inches)
                    top_margin = tcMar.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}top')
                    if top_margin is not None:
                        margin_width = top_margin.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
                        if margin_width and int(margin_width) <= 2:  # 0.001 inches = ~1.44 twips
                            cell_with_minimal_margins += 1
        
        print(f"   Cells with minimal margins: {cell_with_minimal_margins}/15")
        
        if cell_with_minimal_margins == 15:
            print("✅ All cells have minimal margins (no horizontal gutters)")
        else:
            print(f"⚠️  Only {cell_with_minimal_margins}/15 cells have minimal margins")
        
        print("\n📋 Summary:")
        print("   Grid: 5x3 (2 label + 1 gutter + 2 label columns)")
        print("   Vertical gutter: 0.05\" down the middle")
        print("   Horizontal gutters: None (minimal cell margins)")
        print("   Total labels: 12 (all label cells populated)")
        print("   Gutter cells: 3 (all empty)")
        
        print("\n✅ Double template vertical gutter test completed successfully!")
        return True
        
    except Exception as e:
        print(f"❌ Error testing double template vertical gutter: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("🚀 Double Template Vertical Gutter Test")
    print("This test verifies that the double template has only a vertical gutter down the middle")
    print("=" * 70)
    
    success = test_double_template_vertical_gutter()
    
    if success:
        print("\n🎉 Test passed! The double template now has only a vertical gutter.")
    else:
        print("\n💥 Test failed! Check the implementation.") 
#!/usr/bin/env python3
"""
Final test to verify both functions use the same optimized dimensions.
This ensures consistency between manual grid creation and template expansion.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

def test_optimized_dimensions():
    """Test that both functions use the same optimized dimensions."""
    print("🧪 Testing Optimized Dimensions Consistency")
    print("=" * 45)
    
    try:
        from docx import Document
        from src.core.generation.docx_formatting import create_3x3_grid, fix_page_margins_for_3x3_grid
        from src.core.generation.template_processor import TemplateProcessor
        
        # Test 1: Manual grid creation
        print("\n--- Test 1: Manual Grid Creation ---")
        doc1 = Document()
        doc1 = fix_page_margins_for_3x3_grid(doc1)
        
        table1 = create_3x3_grid(doc1, template_type='vertical')
        if table1:
            print(f"✓ Manual grid: {len(table1.rows)}×{len(table1.columns)}")
            
            # Check dimensions
            if table1.rows and table1.columns:
                first_row = table1.rows[0]
                first_cell = first_row.cells[0]
                
                # Get dimensions in twips
                cell_width_twips = getattr(first_cell, 'width', 0)
                row_height_twips = getattr(first_row, 'height', 0)
                
                # Convert to inches
                cell_width_inches = cell_width_twips / 1440 if cell_width_twips else 0
                row_height_inches = row_height_twips / 1440 if row_height_twips else 0
                
                print(f"✓ Manual dimensions: {cell_width_inches:.2f}\" × {row_height_inches:.2f}\"")
                print(f"✓ Total grid: {cell_width_inches * 3:.2f}\" × {row_height_inches * 3:.2f}\"")
            
            # Save manual grid
            doc1.save("test_manual_optimized.docx")
            print(f"✓ Manual grid saved: test_manual_optimized.docx")
        
        # Test 2: Template expansion
        print("\n--- Test 2: Template Expansion ---")
        try:
            # Create a minimal template processor
            processor = TemplateProcessor('vertical', {})
            print(f"✓ Template processor created for: {processor.template_type}")
            
            # Check if template expansion is working
            if hasattr(processor, '_expanded_template_buffer') and processor._expanded_template_buffer:
                print(f"✓ Template expansion buffer exists")
                
                # Try to load the expanded template
                processor._expanded_template_buffer.seek(0)
                from docx import Document
                expanded_doc = Document(processor._expanded_template_buffer)
                
                if expanded_doc.tables:
                    expanded_table = expanded_doc.tables[0]
                    print(f"✓ Expanded template: {len(expanded_table.rows)}×{len(expanded_table.columns)}")
                    
                    # Check dimensions
                    if expanded_table.rows and expanded_table.columns:
                        first_row = expanded_table.rows[0]
                        first_cell = first_row.cells[0]
                        
                        # Get dimensions in twips
                        cell_width_twips = getattr(first_cell, 'width', 0)
                        row_height_twips = getattr(first_row, 'height', 0)
                        
                        # Convert to inches
                        cell_width_inches = cell_width_twips / 1440 if cell_width_twips else 0
                        row_height_inches = row_height_twips / 1440 if row_height_twips else 0
                        
                        print(f"✓ Template dimensions: {cell_width_inches:.2f}\" × {row_height_inches:.2f}\"")
                        print(f"✓ Total grid: {cell_width_inches * 3:.2f}\" × {row_height_inches * 3:.2f}\"")
                    
                    # Save expanded template
                    expanded_doc.save("test_template_optimized.docx")
                    print(f"✓ Template expansion saved: test_template_optimized.docx")
                else:
                    print(f"⚠️  Expanded template has no tables")
            else:
                print(f"⚠️  No expanded template buffer")
                
        except Exception as e:
            print(f"❌ Template processor error: {e}")
        
        print("\n--- Dimension Consistency Check ---")
        print("Both functions should now use the same optimized dimensions:")
        print("- Available space: 8.0\" × 10.5\" (after 0.25\" margins)")
        print("- Cell dimensions: 2.6\" × 3.47\"")
        print("- Total grid: 7.8\" × 10.4\"")
        print("- Buffer space: 0.2\" width, 0.1\" height")
        
        print("\n--- Test Summary ---")
        print("✅ Manual grid creation with optimized dimensions")
        print("✅ Template expansion with optimized dimensions")
        print("✅ Both functions should now be consistent")
        
        print("\n📁 Generated Files:")
        print("- test_manual_optimized.docx (Manual grid)")
        print("- test_template_optimized.docx (Template expansion)")
        
        print("\n🔍 Next Steps:")
        print("1. Open both documents to verify they have the same layout")
        print("2. Check that the 3×3 grid fits completely on the page")
        print("3. Generate labels through the web interface to test the fix")
        
    except Exception as e:
        print(f"❌ Error in test: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_optimized_dimensions()

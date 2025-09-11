#!/usr/bin/env python3
"""
Debug script to test QR code generation step by step.
"""
import sys
import os
sys.path.append('.')

from src.core.generation.template_processor import TemplateProcessor, get_font_scheme
from docx import Document
import logging

# Enable debug logging
logging.basicConfig(level=logging.DEBUG)

def debug_qr_generation():
    """Debug QR code generation step by step."""
    print("=== Debug QR Code Generation ===")
    
    # Get font scheme for mini template
    font_scheme = get_font_scheme('mini', 12)
    print(f"Font scheme: {font_scheme}")
    
    # Create template processor
    processor = TemplateProcessor('mini', font_scheme)
    print(f"Template processor created for: {processor.template_type}")
    
    # Test data
    test_data = [{
        'Product Name*': 'Test Product for QR Code',
        'Product Brand': 'Test Brand',
        'THC test result': 25.5,
        'CBD test result': 2.1,
        'Weight*': '3.5g',
        'Price': '$25.00'
    }]
    
    print(f"Test data: {test_data}")
    
    # Test QR code generation directly
    print("\n=== Testing QR Code Generation Directly ===")
    try:
        # Load the template
        doc = Document(processor._template_path)
        print(f"Template loaded: {processor._template_path}")
        
        # Test QR code generation
        product_name = test_data[0]['Product Name*']
        print(f"Product name for QR: '{product_name}'")
        
        qr_code = processor._generate_qr_code(product_name, doc)
        if qr_code:
            print("✓ QR code generated successfully")
            print(f"QR code type: {type(qr_code)}")
        else:
            print("✗ QR code generation failed")
            
    except Exception as e:
        print(f"✗ Error in direct QR code generation: {e}")
        import traceback
        traceback.print_exc()
    
    # Test label context generation
    print("\n=== Testing Label Context Generation ===")
    try:
        label_context = processor._create_label_context(test_data[0], doc)
        print(f"Label context keys: {list(label_context.keys())}")
        print(f"QR in context: {'QR' in label_context}")
        if 'QR' in label_context:
            print(f"QR value type: {type(label_context['QR'])}")
            print(f"QR value: {label_context['QR']}")
    except Exception as e:
        print(f"✗ Error in label context generation: {e}")
        import traceback
        traceback.print_exc()
    
    # Test full processing
    print("\n=== Testing Full Processing ===")
    try:
        final_doc = processor.process_records(test_data)
        if final_doc:
            print("✓ Document generated successfully")
            
            # Check for QR placeholders in the document
            qr_found = False
            for table in final_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if 'QR_PLACEHOLDER' in para.text:
                                qr_found = True
                                print(f"Found QR_PLACEHOLDER in cell: {para.text}")
            
            if not qr_found:
                print("No QR_PLACEHOLDER found in document")
            
            # Check the generated document
            print("Generated document content:")
            for i, para in enumerate(final_doc.paragraphs):
                print(f'  Paragraph {i}: "{para.text}"')
            
            if final_doc.tables:
                print("Generated document table content:")
                for table_idx, table in enumerate(final_doc.tables):
                    print(f'  Table {table_idx}:')
                    for row_idx, row in enumerate(table.rows):
                        for cell_idx, cell in enumerate(row.cells):
                            for para_idx, para in enumerate(cell.paragraphs):
                                print(f'    Cell {row_idx},{cell_idx} Paragraph {para_idx}: "{para.text}"')
        else:
            print("✗ Document generation failed")
            
    except Exception as e:
        print(f"✗ Error in full processing: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_qr_generation()

#!/usr/bin/env python3
"""
Debug script to test mini template processing step by step
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from io import BytesIO

def test_mini_template():
    print("🔍 Testing mini template processing...")
    
    # Test data
    test_data = [
        {
            'ProductName': 'Test Product 1',
            'ProductBrand': 'Test Brand',
            'Price': '$25.00',
            'Lineage': 'HYBRID',
            'Ratio_or_THC_CBD': 'THC: 20%',
            'ProductStrain': 'Test Strain 1',
            'VendorInfo': 'Test Vendor'
        },
        {
            'ProductName': 'Test Product 2', 
            'ProductBrand': 'Test Brand 2',
            'Price': '$30.00',
            'Lineage': 'INDICA',
            'Ratio_or_THC_CBD': 'THC: 18%',
            'ProductStrain': 'Test Strain 2',
            'VendorInfo': 'Test Vendor 2'
        }
    ]
    
    try:
        # Initialize template processor
        processor = TemplateProcessor('mini', 'mini')
        print(f"✅ Template processor initialized")
        print(f"✅ Template type: {processor.template_type}")
        print(f"✅ Expanded buffer exists: {processor._expanded_template_buffer is not None}")
        
        # Check the expanded template
        if processor._expanded_template_buffer:
            processor._expanded_template_buffer.seek(0)
            doc = Document(processor._expanded_template_buffer)
            print(f"✅ Expanded template loaded: {len(doc.tables)} tables")
            
            if doc.tables:
                table = doc.tables[0]
                print(f"✅ Table dimensions: {len(table.rows)} rows x {len(table.columns)} columns")
                
                # Check first few cells for placeholders
                for i in range(min(3, len(table.rows))):
                    for j in range(min(3, len(table.columns))):
                        cell = table.cell(i, j)
                        print(f"  Cell [{i}][{j}]: '{cell.text[:100]}...'")
        
        # Process the template
        print("\n🔧 Processing template...")
        result = processor.process_chunk(test_data)
        
        if result:
            print(f"✅ Template processed successfully")
            print(f"✅ Result has {len(result.tables)} tables")
            
            if result.tables:
                table = result.tables[0]
                print(f"✅ Final table: {len(table.rows)} rows x {len(table.columns)} columns")
                
                # Check first few cells for content
                for i in range(min(3, len(table.rows))):
                    for j in range(min(3, len(table.columns)):
                        cell = table.cell(i, j)
                        print(f"  Cell [{i}][{j}]: '{cell.text[:100]}...'")
                        
                        # Check if cell has actual content or just formatting
                        if cell.text.strip():
                            print(f"    ✅ Has content: {len(cell.text)} characters")
                        else:
                            print(f"    ❌ Empty content")
                            
                            # Check for paragraphs
                            if cell.paragraphs:
                                for p_idx, para in enumerate(cell.paragraphs):
                                    print(f"      Paragraph {p_idx}: '{para.text}'")
                            else:
                                print(f"      No paragraphs found")
        else:
            print("❌ Template processing failed - no result returned")
            
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_mini_template()

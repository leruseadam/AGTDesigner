#!/usr/bin/env python3
"""
Test script to verify that the DescAndWeight duplication issue is fixed.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor, get_font_scheme

def test_horizontal_template_no_duplication():
    """Test that horizontal template expansion doesn't duplicate DescAndWeight."""
    print("Testing horizontal template expansion for DescAndWeight duplication...")
    
    try:
        # Create a template processor for horizontal template
        processor = TemplateProcessor('horizontal', get_font_scheme('horizontal'), 1.0)
        
        # Check if the expanded template buffer exists
        if hasattr(processor, '_expanded_template_buffer') and processor._expanded_template_buffer:
            print("✅ Template expansion completed successfully")
            
            # Load the expanded template to check for duplication
            from docx import Document
            from io import BytesIO
            
            # Create a copy of the buffer
            buffer_copy = BytesIO(processor._expanded_template_buffer.getvalue())
            doc = Document(buffer_copy)
            
            if doc.tables:
                table = doc.tables[0]
                print(f"✅ Expanded template has {len(table.rows)} rows x {len(table.columns)} columns")
                
                # Check each cell for placeholders and duplication
                descandweight_count = 0
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = ''
                        for paragraph in cell.paragraphs:
                            cell_text += paragraph.text
                        
                        print(f"  Cell ({row_idx}, {col_idx}): {repr(cell_text)}")
                        
                        # Check the actual XML structure for DescAndWeight
                        from docx.oxml.ns import qn
                        descandweight_found = False
                        descandweight_elements = []
                        
                        for t in cell._tc.iter(qn('w:t')):
                            if t.text and 'DescAndWeight' in t.text:
                                descandweight_found = True
                                descandweight_elements.append(t.text)
                        
                        if descandweight_found:
                            descandweight_count += 1
                            print(f"    ✅ Found DescAndWeight: {descandweight_elements}")
                            
                            # Check for duplication within the same cell
                            if len(descandweight_elements) > 1:
                                print(f"❌ ERROR: Cell ({row_idx}, {col_idx}) has {len(descandweight_elements)} DescAndWeight entries!")
                                print(f"    DescAndWeight elements: {descandweight_elements}")
                                return False
                        else:
                            print(f"    ❌ No DescAndWeight found in XML")
                
                print(f"✅ Found DescAndWeight in {descandweight_count} cells (no duplication)")
                
                # Check if placeholders were added
                if descandweight_count > 0:
                    print("✅ DescAndWeight placeholders were successfully added")
                else:
                    print("⚠️  No DescAndWeight placeholders found - template may not need them")
                
                return True
            else:
                print("❌ No tables found in expanded template")
                return False
        else:
            print("❌ Template expansion failed - no expanded buffer")
            return False
            
    except Exception as e:
        print(f"❌ Error testing template: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("Testing DescAndWeight duplication fix...")
    success = test_horizontal_template_no_duplication()
    
    if success:
        print("\n✅ SUCCESS: DescAndWeight duplication issue appears to be fixed!")
    else:
        print("\n❌ FAILURE: DescAndWeight duplication issue still exists!")
    
    sys.exit(0 if success else 1)

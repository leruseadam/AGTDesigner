#!/usr/bin/env python3
"""
Debug script to test template expansion and see why only 2 tags are generated.
"""

import sys
import os
import logging
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# Enable debug logging
logging.basicConfig(level=logging.DEBUG)

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_HORIZONTAL
from docx import Document
from io import BytesIO

def debug_template_expansion_issue():
    """Debug template expansion issue."""
    print("Debugging Template Expansion Issue")
    print("=" * 50)
    
    # Test records
    test_records = [
        {
            'ProductName': 'Elephant Garlic',
            'WeightUnits': '3.5g',
            'ProductBrand': 'JSM LLC',
            'Price': '$48',
            'Lineage': 'HYBRID',
            'THC_CBD': 'THC: 20.88%\nCBD: 0.05%',
            'ProductStrain': 'Elephant Garlic',
            'DOH': 'YES',
            'Product Type*': 'classic'
        },
        {
            'ProductName': 'Canal St. Runtz Rosin',
            'WeightUnits': 'Disposable Vape',
            'ProductBrand': 'JSM LLC',
            'Price': '$35',
            'Lineage': 'HYBRID',
            'THC_CBD': 'THC: 18.5%\nCBD: 0.02%',
            'ProductStrain': 'Canal St. Runtz',
            'DOH': 'YES',
            'Product Type*': 'classic'
        },
        {
            'ProductName': 'Test Product 3',
            'WeightUnits': '1g',
            'ProductBrand': 'Test Brand',
            'Price': '$25',
            'Lineage': 'INDICA',
            'THC_CBD': 'THC: 22.1%\nCBD: 0.01%',
            'ProductStrain': 'Test Strain 3',
            'DOH': 'YES',
            'Product Type*': 'classic'
        },
        {
            'ProductName': 'Test Product 4',
            'WeightUnits': '2g',
            'ProductBrand': 'Test Brand',
            'Price': '$30',
            'Lineage': 'SATIVA',
            'THC_CBD': 'THC: 19.8%\nCBD: 0.03%',
            'ProductStrain': 'Test Strain 4',
            'DOH': 'YES',
            'Product Type*': 'classic'
        },
        {
            'ProductName': 'Test Product 5',
            'WeightUnits': '1.5g',
            'ProductBrand': 'Test Brand',
            'Price': '$28',
            'Lineage': 'HYBRID',
            'THC_CBD': 'THC: 21.2%\nCBD: 0.02%',
            'ProductStrain': 'Test Strain 5',
            'DOH': 'YES',
            'Product Type*': 'classic'
        }
    ]
    
    print(f"Testing with {len(test_records)} records")
    
    # Create template processor
    processor = TemplateProcessor('horizontal', FONT_SCHEME_HORIZONTAL)
    
    print(f"Template type: {processor.template_type}")
    print(f"Chunk size: {processor.chunk_size}")
    
    # Check if template was expanded
    if hasattr(processor, '_expanded_template_buffer') and processor._expanded_template_buffer:
        print("✓ Template buffer exists")
        
        # Check the expanded template
        processor._expanded_template_buffer.seek(0)
        expanded_doc = Document(processor._expanded_template_buffer)
        print(f"Expanded template has {len(expanded_doc.tables)} tables")
        
        if expanded_doc.tables:
            table = expanded_doc.tables[0]
            print(f"Table dimensions: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check for Label placeholders
            table_text = ""
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "
            
            for i in range(1, 10):
                placeholder = f"Label{i}"
                if placeholder in table_text:
                    print(f"✓ Found {placeholder} placeholder")
                else:
                    print(f"✗ Missing {placeholder} placeholder")
        else:
            print("✗ No tables found in expanded template")
    else:
        print("✗ No template buffer found")
    
    # Process the records
    print("\nProcessing records...")
    try:
        result_doc = processor.process_records(test_records)
        if result_doc:
            print(f"✓ Successfully generated document with {len(result_doc.tables)} tables")
            
            if result_doc.tables:
                final_table = result_doc.tables[0]
                print(f"Final table dimensions: {len(final_table.rows)} rows x {len(final_table.columns)} columns")
                
                # Check if we have content in all cells
                content_count = 0
                for row in final_table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content_count += 1
                
                print(f"Cells with content: {content_count}")
                
                # Save the result for inspection
                output_path = "debug_template_expansion_output.docx"
                result_doc.save(output_path)
                print(f"Saved result to: {output_path}")
            else:
                print("✗ No tables in final document")
        else:
            print("✗ Failed to generate document")
    except Exception as e:
        print(f"✗ Error processing records: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_template_expansion_issue()

#!/usr/bin/env python3
"""
Debug script to understand why brand markers aren't being added to non-classic types.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import CLASSIC_TYPES
from docx import Document
from docx.shared import Pt
import tempfile
import os

def debug_brand_marker_addition():
    """Debug brand marker addition logic."""
    print("Debugging Brand Marker Addition")
    
    # Test with a non-classic type (edible)
    non_classic_record = {
        'ProductName': 'Test Edible',
        'ProductBrand': 'Test Brand',
        'ProductStrain': 'Test Strain',
        'ProductType': 'edible (solid)',  # Non-classic type
        'ProductVendor': 'Test Vendor',
        'Lineage': 'HYBRID',
        'Price': '$25.00',
        'WeightUnits': '3.5g',
        'Description': 'Test Description'
    }
    
    try:
        print("\n1. Testing Non-Classic Type (edible)...")
        tp = TemplateProcessor('double', {})
        tp.current_product_type = 'edible (solid)'
        
        # Check if product type is correctly set
        print(f"  Current product type: {tp.current_product_type}")
        print(f"  Is classic type: {tp.current_product_type.lower() in [ct.lower() for ct in CLASSIC_TYPES]}")
        
        # Check what content is in the document before brand marker addition
        documents = tp.process_records([non_classic_record])
        
        if not documents:
            print("✗ No documents generated")
            return False
        
        doc = documents[0] if isinstance(documents, list) else documents
        
        print("\n2. Checking document content before brand marker addition:")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            print(f"  Found text: '{text}'")
                            # Check if this looks like brand content
                            if (text and 
                                'PRODUCTBRAND_CENTER_START' not in text and 
                                'RATIO_START' not in text and
                                'RATIO_END' not in text and
                                '{{' not in text and 
                                '}}' not in text and
                                len(text) > 0 and
                                not text.isdigit() and
                                not text.startswith('$') and
                                not text.endswith('g') and
                                not text.endswith('mg')):
                                print(f"    ✓ This text qualifies for brand marker addition")
                            else:
                                print(f"    ✗ This text does NOT qualify for brand marker addition")
        
        # Check if brand markers are present
        brand_markers_found = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if 'PRODUCTBRAND_CENTER_START' in text or 'PRODUCTBRAND_CENTER_END' in text:
                            brand_markers_found = True
                            print(f"  Found brand markers: '{text.strip()}'")
        
        if brand_markers_found:
            print("✓ SUCCESS: Brand markers found in non-classic type")
        else:
            print("✗ ERROR: No brand markers found in non-classic type")
        
        return True
        
    except Exception as e:
        print(f"✗ Error debugging brand marker addition: {e}")
        return False

if __name__ == "__main__":
    print("Debugging Brand Marker Addition")
    print("=" * 50)
    
    success = debug_brand_marker_addition()
    
    if success:
        print("\n✓ Debug completed")
    else:
        print("\n✗ Debug failed") 
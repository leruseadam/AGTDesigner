#!/usr/bin/env python3
"""
Debug script to test actual template processing flow for ProductStrain font sizing.
"""

from src.core.generation.template_processor import TemplateProcessor
from src.core.formatting.markers import wrap_with_marker
from docx import Document
from docx.shared import Pt

def test_real_template_processing():
    """Test the actual template processing flow to see where font sizing is lost."""
    
    print("=== Testing Real Template Processing Flow ===")
    
    # Create template processor
    processor = TemplateProcessor('double', 'standard', 1.0)
    
    # Create a test document that mimics the actual template structure
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    
    # Add text that mimics what would be in the actual template
    # This should be the content after template expansion but before marker processing
    test_text = f"{{{{Label1.Lineage}}}}\n{wrap_with_marker('Test Strain', 'PRODUCTSTRAIN')}\n{{{{Label1.Price}}}}"
    cell.text = test_text
    
    print(f"1. Created test document with template-like content:")
    print(f"   '{test_text}'")
    
    # Check initial state
    paragraph = cell.paragraphs[0]
    print(f"\n2. Initial paragraph state:")
    print(f"   Text: '{paragraph.text}'")
    print(f"   Runs: {len(paragraph.runs)}")
    for i, run in enumerate(paragraph.runs):
        font_size_pt = run.font.size.pt if run.font.size else "No size"
        print(f"     Run {i}: '{run.text}' -> {font_size_pt}pt")
    
    # Now apply the actual post-processing that happens in the template processor
    print(f"\n3. Applying template-specific post-processing...")
    
    # This is what gets called in the actual template processing
    processor._post_process_template_specific(doc)
    
    print(f"\n4. After post-processing:")
    print(f"   Text: '{paragraph.text}'")
    print(f"   Runs: {len(paragraph.runs)}")
    for i, run in enumerate(paragraph.runs):
        font_size_pt = run.font.size.pt if run.font.size else "No size"
        print(f"     Run {i}: '{run.text}' -> {font_size_pt}pt")
    
    # Let's also check what markers are being processed
    print(f"\n5. Checking what markers are defined in post-processing:")
    markers = [
        'DESC', 'PRODUCTBRAND', 'PRODUCTBRAND_CENTER', 'PRICE', 'LINEAGE', 
        'THC_CBD', 'THC_CBD_LABEL', 'RATIO', 'WEIGHTUNITS', 'PRODUCTSTRAIN', 'DOH', 'PRODUCTVENDOR'
    ]
    print(f"   Markers: {markers}")
    
    # Check if ProductStrain is in the markers list
    if 'PRODUCTSTRAIN' in markers:
        print(f"   ✅ PRODUCTSTRAIN is in the markers list")
    else:
        print(f"   ❌ PRODUCTSTRAIN is NOT in the markers list")
    
    print(f"\n=== Test Complete ===")

if __name__ == "__main__":
    test_real_template_processing() 
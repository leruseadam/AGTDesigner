#!/usr/bin/env python3
"""
Debug script to test if deepcopy is working correctly for mini template preservation.
"""

import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), 'src'))

from core.generation.template_processor import TemplateProcessor
from docx import Document
from copy import deepcopy
import tempfile

def debug_mini_template_deepcopy():
    """Debug the deepcopy operation for mini template preservation."""
    print("=== Debugging Mini Template Deepcopy ===")
    
    try:
        # 1. Load original template
        print("\n1. Loading original mini template...")
        original_path = 'src/core/generation/templates/mini.docx'
        original_doc = Document(original_path)
        print(f"   ✅ Original template loaded: {len(original_doc.tables)} tables")
        
        if original_doc.tables:
            table = original_doc.tables[0]
            print(f"   📊 Original table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check first cell content
            first_cell = table.rows[0].cells[0]
            print(f"   📝 Original first cell: {first_cell.text[:100]}...")
        
        # 2. Test deepcopy manually
        print("\n2. Testing deepcopy manually...")
        new_doc = Document()
        
        # Copy all content from the original document
        for element in original_doc.element.body:
            new_doc.element.body.append(deepcopy(element))
        
        print(f"   ✅ Deepcopy completed: {len(new_doc.tables)} tables")
        
        if new_doc.tables:
            table = new_doc.tables[0]
            print(f"   📊 Copied table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check first cell content
            first_cell = table.rows[0].cells[0]
            print(f"   📝 Copied first cell: {first_cell.text[:100]}...")
            
            # Compare with original
            if first_cell.text == original_doc.tables[0].rows[0].cells[0].text:
                print("   ✅ Deepcopy preserved content exactly")
            else:
                print("   ❌ Deepcopy changed content!")
                print(f"      Original: {original_doc.tables[0].rows[0].cells[0].text[:50]}...")
                print(f"      Copied:   {first_cell.text[:50]}...")
        
        # 3. Test the actual method
        print("\n3. Testing the actual _expand_mini_template_preserve_design method...")
        processor = TemplateProcessor('mini', 'Arial')
        
        # Create test context
        context = {
            'Label1': {
                'ProductStrain': 'TEST STRAIN',
                'ProductBrand': 'TEST BRAND', 
                'VendorInfo': 'TEST VENDOR',
                'Price': '$10.00',
                'Ratio_or_THC_CBD': 'THC: 20%',
                'Lineage': 'TEST'
            }
        }
        
        print(f"   📋 Created context: {context}")
        
        # Process the mini template
        result_doc = processor._expand_mini_template_preserve_design(None, context)
        
        print(f"   ✅ Method completed: {len(result_doc.tables)} tables")
        
        if result_doc.tables:
            table = result_doc.tables[0]
            print(f"   📊 Result table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check first cell content
            first_cell = table.rows[0].cells[0]
            print(f"   📝 Result first cell: {first_cell.text[:100]}...")
            
            # Check if placeholders were replaced
            if 'TEST STRAIN' in first_cell.text and 'TEST BRAND' in first_cell.text:
                print("   ✅ Placeholders replaced successfully")
            else:
                print("   ❌ Placeholders not replaced correctly!")
        
        # 4. Save and compare
        print("\n4. Saving and comparing results...")
        
        # Save original
        with tempfile.NamedTemporaryFile(suffix='_original.docx', delete=False) as tmp_file:
            original_doc.save(tmp_file.name)
            print(f"   💾 Original saved to: {tmp_file.name}")
        
        # Save copied
        with tempfile.NamedTemporaryFile(suffix='_copied.docx', delete=False) as tmp_file:
            new_doc.save(tmp_file.name)
            print(f"   💾 Copied saved to: {tmp_file.name}")
        
        # Save result
        with tempfile.NamedTemporaryFile(suffix='_result.docx', delete=False) as tmp_file:
            result_doc.save(tmp_file.name)
            print(f"   💾 Result saved to: {tmp_file.name}")
            
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False
    
    print("\n=== Mini Template Deepcopy Debug Completed ===")
    return True

if __name__ == "__main__":
    debug_mini_template_deepcopy()

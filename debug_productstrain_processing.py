#!/usr/bin/env python3
"""
Detailed debug script to examine ProductStrain field processing.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
import docx

def debug_productstrain_processing():
    """Debug ProductStrain field processing in detail."""
    
    print("=== DEBUGGING PRODUCTSTRAIN FIELD PROCESSING ===")
    
    try:
        # Create template processor
        tp = TemplateProcessor('double', {}, 1.0)
        
        # Check the template path
        template_path = tp._get_template_path()
        print(f"Template path: {template_path}")
        
        # Load the template document
        doc = docx.Document(template_path)
        print(f"Template has {len(doc.tables)} tables")
        
        if doc.tables:
            table = doc.tables[0]
            print(f"First table: {len(table.rows)}x{len(table.columns)}")
            
            # Check the first cell for ProductStrain markers
            cell = table.cell(0, 0)
            print('First cell content:')
            for i, para in enumerate(cell.paragraphs):
                print(f'  Paragraph {i}: {repr(para.text)}')
                
                # Check for PRODUCTSTRAIN markers specifically
                if 'PRODUCTSTRAIN_START' in para.text:
                    print(f'    ✓ Has PRODUCTSTRAIN_START marker')
                    # Check if the marker is properly formatted
                    if 'PRODUCTSTRAIN_START' in para.text and 'PRODUCTSTRAIN_END' in para.text:
                        print(f'    ✓ Has both START and END markers')
                        # Extract the content between markers
                        start_idx = para.text.find('PRODUCTSTRAIN_START') + len('PRODUCTSTRAIN_START')
                        end_idx = para.text.find('PRODUCTSTRAIN_END')
                        if start_idx >= 0 and end_idx >= 0:
                            content = para.text[start_idx:end_idx].strip()
                            print(f'    Content between markers: {repr(content)}')
                    else:
                        print(f'    ❌ Missing END marker')
                elif 'PRODUCTSTRAIN' in para.text:
                    print(f'    ⚠️  Has PRODUCTSTRAIN but not properly marked')
                    
                # Check for other markers
                if 'LINEAGE_START' in para.text:
                    print(f'    ✓ Has LINEAGE_START marker')
                if 'PRICE_START' in para.text:
                    print(f'    ✓ Has PRICE_START marker')
        
        print("\n=== TESTING TEMPLATE PROCESSING ===")
        
        # Create a test record
        test_record = {
            'Description': 'Test Product 1',
            'WeightUnits': '1g',
            'ProductBrand': 'BRAND_ONE',
            'Price': '$25',
            'Lineage': 'HYBRID',
            'THC_CBD': 'THC:\n  18%\nCBD:\n  2%',
            'ProductStrain': 'STRAIN_ONE',
            'DOH': 'NO',
            'Product Type*': 'classic'
        }
        
        print("Test record:")
        for key, value in test_record.items():
            print(f"  {key}: {repr(value)}")
        
        # Process the record
        print("\nProcessing test record...")
        result_doc = tp.process_records([test_record])
        
        if not result_doc:
            print("ERROR: Failed to process test record")
            return
        
        print(f"Processed document has {len(result_doc.tables)} tables")
        
        # Check the processed document
        if result_doc.tables:
            table = result_doc.tables[0]
            cell = table.cell(0, 0)
            print('Processed document first cell content:')
            for i, para in enumerate(cell.paragraphs):
                print(f'  Paragraph {i}: {repr(para.text)}')
                
                # Check font sizes for each run
                if para.runs:
                    for j, run in enumerate(para.runs):
                        font_size = "No font size set"
                        if run.font.size:
                            font_size = f"{run.font.size.pt}pt"
                        print(f'    Run {j}: {repr(run.text)} - Font size: {font_size}')
                        
                        # Check if this is the ProductStrain field
                        if 'STRAIN_ONE' in run.text:
                            print(f'      *** This should be ProductStrain with 1pt font ***')
                            if run.font.size and run.font.size.pt == 1.0:
                                print(f'      ✓ SUCCESS: ProductStrain has 1pt font!')
                            else:
                                print(f'      ❌ FAILED: ProductStrain does not have 1pt font')
                
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_productstrain_processing() 
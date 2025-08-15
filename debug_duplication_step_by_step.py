#!/usr/bin/env python3
"""
Debug script to test each step of template processing and find where duplication occurs.
This will help us identify the exact point where "CONSTELLATION CANNABISCONSTELLATION CANNABIS" is created.
"""

import sys
import os
import logging
from pathlib import Path

# Add the src directory to the path
sys.path.insert(0, str(Path(__file__).parent / 'src'))

from src.core.generation.template_processor import TemplateProcessor

# Set up logging
logging.basicConfig(level=logging.DEBUG, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)

def test_single_record():
    """Test with a single record to isolate the duplication issue."""
    
    # Test record that should show the duplication
    test_record = {
        'ProductName': 'Grape Moonshot',
        'ProductBrand': 'CONSTELLATION CANNABIS',
        'Product Type*': 'edible (solid)',
        'Description': 'Grape Moonshot -1.7oz',
        'WeightUnits': '1.7oz',
        'Price': '$15',
        'THC': '100.0%',
        'CBD': '0%',
        'Lineage': 'MIXED'
    }
    
    logger.info("=== TESTING SINGLE RECORD ===")
    logger.info(f"Test record: {test_record}")
    
    # Test with horizontal template - use the function from template processor
    from src.core.generation.template_processor import get_font_scheme
    font_scheme = get_font_scheme('horizontal')
    processor = TemplateProcessor('horizontal', font_scheme)
    
    logger.info("=== STEP 1: Building Label Context ===")
    try:
        # Create a dummy document for context building
        from docx import Document
        doc = Document()
        
        # Build the label context
        label_context = processor._build_label_context(test_record, doc)
        logger.info(f"Label context built: {label_context}")
        
        # Check ProductBrand specifically
        if 'ProductBrand' in label_context:
            logger.info(f"ProductBrand in context: '{label_context['ProductBrand']}'")
            if 'CONSTELLATION CANNABIS' in str(label_context['ProductBrand']):
                logger.info("✅ ProductBrand contains 'CONSTELLATION CANNABIS' as expected")
            else:
                logger.error("❌ ProductBrand does NOT contain 'CONSTELLATION CANNABIS'")
        else:
            logger.error("❌ ProductBrand not found in label context")
            
        # Check Lineage specifically
        if 'Lineage' in label_context:
            logger.info(f"Lineage in context: '{label_context['Lineage']}'")
            if 'CONSTELLATION CANNABIS' in str(label_context['Lineage']):
                logger.info("✅ Lineage contains 'CONSTELLATION CANNABIS' as expected")
            else:
                logger.info("ℹ️ Lineage does not contain 'CONSTELLATION CANNABIS' (this might be correct)")
        else:
            logger.error("❌ Lineage not found in label context")
            
    except Exception as e:
        logger.error(f"Error building label context: {e}")
        import traceback
        traceback.print_exc()
        return
    
    logger.info("=== STEP 2: Testing Template Expansion ===")
    try:
        # Test template expansion
        expanded_buffer = processor._expanded_template_buffer
        logger.info(f"Template expanded, buffer size: {expanded_buffer.getbuffer().nbytes if expanded_buffer else 'None'}")
        
        # Load the expanded template
        from docx import Document
        from io import BytesIO
        
        if expanded_buffer:
            expanded_buffer.seek(0)
            doc = Document(expanded_buffer)
            
            # Check the first few cells for content
            if doc.tables:
                table = doc.tables[0]
                logger.info(f"Expanded template has {len(table.rows)} rows and {len(table.columns)} columns")
                
                # Check first cell content
                if len(table.rows) > 0 and len(table.columns) > 0:
                    first_cell = table.rows[0].cells[0]
                    cell_text = ' '.join([p.text for p in first_cell.paragraphs])
                    logger.info(f"First cell text: '{cell_text}'")
                    
                    if 'CONSTELLATION CANNABIS' in cell_text:
                        logger.info("✅ First cell contains 'CONSTELLATION CANNABIS'")
                    else:
                        logger.info("ℹ️ First cell does not contain 'CONSTELLATION CANNABIS'")
                        
                    if 'CONSTELLATION CANNABISCONSTELLATION CANNABIS' in cell_text:
                        logger.error("❌ FIRST CELL ALREADY HAS DUPLICATION!")
                    else:
                        logger.info("✅ First cell does not have duplication")
                        
                    # Check all cells in the expanded template
                    logger.info("=== CHECKING ALL CELLS IN EXPANDED TEMPLATE ===")
                    for row_idx, row in enumerate(table.rows):
                        for col_idx, cell in enumerate(row.cells):
                            cell_text = ' '.join([p.text for p in cell.paragraphs])
                            logger.info(f"Cell ({row_idx}, {col_idx}): '{cell_text}'")
                            
                            if 'CONSTELLATION CANNABISCONSTELLATION CANNABIS' in cell_text:
                                logger.error(f"❌ DUPLICATION FOUND in expanded template cell ({row_idx}, {col_idx})!")
                            elif 'CONSTELLATION CANNABIS' in cell_text:
                                logger.info(f"✅ Cell ({row_idx}, {col_idx}) has correct content")
        else:
            logger.error("❌ No expanded template buffer")
            
    except Exception as e:
        logger.error(f"Error testing template expansion: {e}")
        import traceback
        traceback.print_exc()
        return
    
    logger.info("=== STEP 3: Testing Document Generation ===")
    try:
        # Test full document generation
        records = [test_record]
        documents = processor.process_records(records)
        
        if documents:
            logger.info(f"Generated {len(documents) if hasattr(documents, '__len__') else 'single'} document(s)")
            
            # Check the first document
            first_doc_buffer = documents[0] if hasattr(documents, '__len__') else documents
            first_doc_buffer.seek(0)
            doc = Document(first_doc_buffer)
            
            if doc.tables:
                table = doc.tables[0]
                logger.info(f"Generated document has {len(table.rows)} rows and {len(table.columns)} columns")
                
                # Check all cells for duplication
                duplication_found = False
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = ' '.join([p.text for p in cell.paragraphs])
                        if 'CONSTELLATION CANNABISCONSTELLATION CANNABIS' in cell_text:
                            logger.error(f"❌ DUPLICATION FOUND in cell ({row_idx}, {col_idx}): '{cell_text}'")
                            duplication_found = True
                        elif 'CONSTELLATION CANNABIS' in cell_text:
                            logger.info(f"✅ Cell ({row_idx}, {col_idx}) has correct content: '{cell_text}'")
                
                if not duplication_found:
                    logger.info("✅ No duplication found in generated document")
                else:
                    logger.error("❌ DUPLICATION WAS CREATED DURING DOCUMENT GENERATION")
            else:
                logger.error("❌ Generated document has no tables")
        else:
            logger.error("❌ No documents generated")
            
    except Exception as e:
        logger.error(f"Error testing document generation: {e}")
        import traceback
        traceback.print_exc()
        return

def test_template_content():
    """Test the actual template content to see what's in it."""
    
    logger.info("=== TESTING TEMPLATE CONTENT ===")
    
    try:
        # Check the horizontal template directly
        template_path = Path(__file__).parent / 'src' / 'core' / 'generation' / 'templates' / 'horizontal.docx'
        
        if template_path.exists():
            logger.info(f"Template exists: {template_path}")
            
            from docx import Document
            doc = Document(template_path)
            
            if doc.tables:
                table = doc.tables[0]
                logger.info(f"Template has {len(table.rows)} rows and {len(table.columns)} columns")
                
                # Check first cell content
                if len(table.rows) > 0 and len(table.columns) > 0:
                    first_cell = table.rows[0].cells[0]
                    cell_text = ' '.join([p.text for p in first_cell.paragraphs])
                    logger.info(f"Template first cell text: '{cell_text}'")
                    
                    if 'CONSTELLATION CANNABIS' in cell_text:
                        logger.error("❌ TEMPLATE ALREADY CONTAINS 'CONSTELLATION CANNABIS' - THIS IS THE PROBLEM!")
                    else:
                        logger.info("✅ Template does not contain 'CONSTELLATION CANNABIS'")
                        
                    if 'CONSTELLATION CANNABISCONSTELLATION CANNABIS' in cell_text:
                        logger.error("❌ TEMPLATE ALREADY HAS DUPLICATION!")
                    else:
                        logger.info("✅ Template does not have duplication")
        else:
            logger.error("❌ Template has no tables")
    except Exception as e:
        logger.error(f"Error testing template content: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    logger.info("Starting duplication debug test...")
    
    # Test 1: Check template content
    test_template_content()
    
    # Test 2: Test single record processing
    test_single_record()
    
    logger.info("Debug test completed.")

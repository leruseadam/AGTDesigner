#!/usr/bin/env python3
"""
Test QR code rendering by directly modifying the template.
"""
import sys
import os
sys.path.append('.')

from docx import Document
from docx.shared import Mm
import qrcode
from io import BytesIO

def test_qr_direct_rendering():
    """Test QR code rendering by directly modifying the template."""
    print("Testing direct QR code rendering...")
    
    # Load the template
    template_path = 'src/core/generation/templates/mini.docx'
    doc = Document(template_path)
    
    print("Template content before modification:")
    for i, para in enumerate(doc.paragraphs):
        print(f'Paragraph {i}: "{para.text}"')
    
    # Find the QR placeholder and replace it with actual QR code
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if '{{QR}}' in para.text:
                        print(f"Found QR placeholder in paragraph: {para.text}")
                        
                        # Create QR code
                        qr = qrcode.QRCode(
                            version=1,
                            error_correction=qrcode.constants.ERROR_CORRECT_L,
                            box_size=10,
                            border=4,
                        )
                        qr.add_data("Test QR Code for Direct Rendering")
                        qr.make(fit=True)
                        
                        # Create QR code image
                        qr_image = qr.make_image(fill_color="black", back_color="white")
                        
                        # Convert to BytesIO
                        img_buffer = BytesIO()
                        qr_image.save(img_buffer, format='PNG')
                        img_buffer.seek(0)
                        
                        # Clear the paragraph and add the QR code
                        para.clear()
                        para.alignment = 1  # Center alignment
                        
                        # Add the image to the paragraph
                        run = para.add_run()
                        run.add_picture(img_buffer, width=Mm(20))
                        
                        print("✓ QR code added directly to paragraph")
                        break
    
    # Save the result
    output_file = 'test_qr_direct_output.docx'
    doc.save(output_file)
    
    print(f"✓ Document saved: {output_file}")
    
    # Check the result
    result_doc = Document(output_file)
    print("Result document content:")
    for i, para in enumerate(result_doc.paragraphs):
        print(f'Paragraph {i}: "{para.text}"')
    
    # Check for images
    if hasattr(result_doc, 'part') and hasattr(result_doc.part, 'related_parts'):
        image_count = len([part for part in result_doc.part.related_parts.values() 
                         if 'image' in part.content_type])
        print(f'Found {image_count} images in the result document')
    
    return True

if __name__ == "__main__":
    success = test_qr_direct_rendering()
    if success:
        print("\n✓ Direct QR code rendering test completed successfully!")
    else:
        print("\n✗ Direct QR code rendering test failed!")

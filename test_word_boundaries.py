#!/usr/bin/env python3
"""
Test different approaches to preventing "7g" splitting.
"""

import sys
sys.path.insert(0, '/Users/adamcordova/Desktop/labelMaker_ QR copy final')

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.core.generation.text_processing import make_nonbreaking_hyphens, make_nonbreaking_weight_units

# Process the text
test_text = "Blueberry - 7g"
processed = make_nonbreaking_hyphens(test_text)
processed = make_nonbreaking_weight_units(processed)

print(f"Original: {repr(test_text)}")
print(f"Processed: {repr(processed)}")
print(f"\nCharacter breakdown of processed text:")
for i, char in enumerate(processed):
    print(f"  [{i}] {repr(char)} (U+{ord(char):04X}) {chr(ord(char)) if ord(char) >= 32 else '[control]'}")

# Test different approaches
test_cases = [
    ("1_current_approach", True, True, processed),
    ("2_no_wordwrap", False, True, processed),
    ("3_no_wordwrap_no_suppresshyphens", False, False, processed),
    ("4_zwj_between_number_unit", True, True, processed.replace('\xa0', '\u200d')),  # Zero-width joiner
    ("5_word_joiner_u2060", True, True, processed.replace('\xa0', '\u2060')),  # Word joiner
]

for name, use_wordwrap, use_suppress_hyphens, text in test_cases:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]

    # Set cell width to 2 inches
    cell.width = Inches(2.0)
    tcPr = cell._element.get_or_add_tcPr()

    tcW = OxmlElement('w:tcW')
    tcPr.append(tcW)
    tcW.set(qn('w:w'), '2880')
    tcW.set(qn('w:type'), 'dxa')

    # Add text
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.size = Pt(10)

    # Apply formatting
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = paragraph._element.get_or_add_pPr()

    if use_wordwrap:
        wordWrap = OxmlElement('w:wordWrap')
        pPr.append(wordWrap)
        wordWrap.set(qn('w:val'), '1')

    if use_suppress_hyphens:
        suppressAutoHyphens = OxmlElement('w:suppressAutoHyphens')
        pPr.append(suppressAutoHyphens)

    filename = f'/Users/adamcordova/Desktop/labelMaker_ QR copy final/test_{name}.docx'
    doc.save(filename)
    print(f"\n✓ Created: test_{name}.docx")
    print(f"   wordWrap={use_wordwrap}, suppressAutoHyphens={use_suppress_hyphens}")
    print(f"   Text: {repr(text)}")

print("\n" + "="*60)
print("Open these files in Word and check which prevents splitting:")
for name, _, _, _ in test_cases:
    print(f"  test_{name}.docx")
print("="*60)

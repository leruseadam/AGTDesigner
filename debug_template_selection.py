#!/usr/bin/env python3
"""
Debug script to test template selection and see what's happening with the mini template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor, get_font_scheme

def debug_template_selection():
    """Debug template selection and processing."""
    
    print("Debug Template Selection")
    print("=" * 50)
    
    # Test different template types
    template_types = ['mini', 'double', 'horizontal', 'vertical']
    
    for template_type in template_types:
        print(f"\n--- Testing {template_type.upper()} template ---")
        
        try:
            # Create template processor
            font_scheme = get_font_scheme(template_type)
            processor = TemplateProcessor(template_type, font_scheme, 1.0)
            
            print(f"Template type: {processor.template_type}")
            print(f"Template path: {processor._template_path}")
            print(f"Chunk size: {processor.chunk_size}")
            
            # Check expanded template
            if hasattr(processor, '_expanded_template_buffer'):
                processor._expanded_template_buffer.seek(0)
                from docx import Document
                doc = Document(processor._expanded_template_buffer)
                
                if doc.tables:
                    table = doc.tables[0]
                    print(f"Table dimensions: {len(table.rows)}x{len(table.columns)}")
                    
                    # Check first cell content
                    first_cell = table.cell(0, 0)
                    print(f"First cell text: '{first_cell.text}'")
                    
                    # Check for template variables
                    if '{{Label1.' in first_cell.text:
                        import re
                        variables = re.findall(r'\{\{Label1\.(\w+)\}\}', first_cell.text)
                        print(f"Template variables: {variables}")
                    else:
                        print("No template variables found")
                else:
                    print("No tables found")
            else:
                print("No expanded template buffer")
                
        except Exception as e:
            print(f"Error with {template_type}: {e}")
            import traceback
            traceback.print_exc()

if __name__ == "__main__":
    debug_template_selection()

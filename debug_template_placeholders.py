#!/usr/bin/env python3
"""
Debug script to examine template placeholders and content
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
import tempfile

def debug_template_placeholders():
    """Debug the template placeholders to see what exists"""
    print("🔍 Debugging template placeholders...")
    
    # Create test data
    test_records = [{
        'ProductStrain': 'Test Strain Name',
        'ProductBrand': 'Test Brand',
        'Lineage': 'HYBRID',
        'ProductType': 'Concentrate'
    }]
    
    # Test with double template
    processor = TemplateProcessor('double', 'default', 1.0)
    
    # Get the template path
    template_path = processor._get_template_path()
    print(f"Template path: {template_path}")
    
    # Load the original template
    doc = Document(template_path)
    print(f"Original template has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
    
    if doc.tables:
        table = doc.tables[0]
        print(f"First table has {len(table.rows)} rows x {len(table.columns)} columns")
        
        # Check the first cell content
        first_cell = table.cell(0, 0)
        print(f"First cell text: '{first_cell.text}'")
        
        # Check for specific placeholders
        cell_text = first_cell.text
        print(f"Placeholder analysis:")
        print(f"  Contains 'Label1.Lineage': {'Label1.Lineage' in cell_text}")
        print(f"  Contains 'Label1.ProductStrain': {'Label1.ProductStrain' in cell_text}")
        print(f"  Contains 'Label1.ProductBrand': {'Label1.ProductBrand' in cell_text}")
        
        # Check individual text elements
        print(f"Individual text elements in first cell:")
        for i, run in enumerate(first_cell.paragraphs[0].runs):
            print(f"  Run {i}: '{run.text}'")
    
    # Now test the expanded template
    print(f"\n🔍 Testing expanded template...")
    try:
        expanded_buffer = processor._expand_template_to_4x3_fixed_double()
        expanded_doc = Document(expanded_buffer)
        print(f"Expanded template has {len(expanded_doc.paragraphs)} paragraphs and {len(expanded_doc.tables)} tables")
        
        if expanded_doc.tables:
            table = expanded_doc.tables[0]
            print(f"Expanded table has {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check the first cell content
            first_cell = table.cell(0, 0)
            print(f"First expanded cell text: '{first_cell.text}'")
            
            # Check for specific placeholders
            cell_text = first_cell.text
            print(f"Expanded placeholder analysis:")
            print(f"  Contains 'Label1.Lineage': {'Label1.Lineage' in cell_text}")
            print(f"  Contains 'Label1.ProductStrain': {'Label1.ProductStrain' in cell_text}")
            print(f"  Contains 'Label1.ProductBrand': {'Label1.ProductBrand' in cell_text}")
            
            # Check individual text elements
            print(f"Individual text elements in first expanded cell:")
            for i, run in enumerate(first_cell.paragraphs[0].runs):
                print(f"  Run {i}: '{run.text}'")
        
    except Exception as e:
        print(f"Error expanding template: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_template_placeholders() 